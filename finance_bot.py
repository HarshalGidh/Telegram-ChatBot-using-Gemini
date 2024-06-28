#best code so far correctly takes the assessment and gives financial advice along 
#with returns with proper disclaimer and in brief

import os
import filetype
import docx
import PyPDF2
import re
from aiogram import Bot, Dispatcher, executor, types
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.memory import ConversationSummaryMemory
import asyncio
from aiogram.utils.exceptions import NetworkError, RetryAfter, TelegramAPIError
import google.generativeai as genai

# Import things that are needed generically
from langchain.pydantic_v1 import BaseModel, Field
from langchain.tools import BaseTool, StructuredTool, tool

load_dotenv()

TOKEN = os.getenv("TOKEN")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

os.environ["LANGCHAIN_TRACING_V2"]="true"
os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# Configure generativeai with your API key
genai.configure(api_key=GOOGLE_API_KEY)

# Initialize bot
bot = Bot(token=TOKEN)
dispatcher = Dispatcher(bot)

rag_on = False
retriever = None  # Store retriever globally
summary = ""
investment_personality = ""
chat_history = ""

class Reference:
    def __init__(self):
        self.response = ""


reference = Reference()


def clear_past():
    reference.response = ""


@dispatcher.message_handler(commands=['clear'])
async def clear(message: types.Message):
    """
    A handler to clear the previous conversation and context.
    """
    clear_past()
    await message.reply("I've cleared the past conversation and context.")


# Store user states
states = {}
# Dictionary to hold question-answer pairs
user_responses = {}

# Define Questions for assessment
questions = [
    """
        1. Singapore plans to build a new observation tower called 'The Rook'.
        How many steps do you think it will take to go to the top floor?

        a) Less than 500 
        b) More than 500

    """,
    "2. Now Guess the number of steps" ,
    """
    3. How confident are you that the real number is in the range you have selected? 
    Answer within a range of 100.  
    """,
    """ 
    4. You and your friend are betting on a series of coin tosses.

    He always bets ₹2,000 on Heads

    You always bet ₹2,000 on Tails

    Winner of last 8 turns

    You lost ₹8,000 in the last 4 turns!

    If you were to bet one last time, what would you bet on heads or tails ?
    """ ,
    """
    5. How confident are you that your bet will win this time?
    Answer how confident you are. 
    (Example: Not confident at all, somewhat confident, confident, or Very confident)
    """,
    """
    6. What do you think are the chances that you will achieve and maintain your financial goals within the next 
    10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
    Answer how likely you are to achieve your goal.
    (Example: Less likely than others, likely than others, or More likely than others)
    """,
    """
    7. Imagine you are a contestant in a game show, and you are presented the following choices.

    What would you prefer?
    a) 50 percent chance of winning 15 gold coins 
    b) 100 percent chance of winning 8 gold coins
    """,
    """
    8. Ok, one last choice...

    What would you prefer?
    a) 50 percent chance of winning 15 gold coins 
    b) 100 percent chance of winning 2 gold coins
    """,
    """
    9. In general, how would your best friend describe your risk-taking tendencies?
    a) A real gambler
    b) Willing to take risks after completing adequate research
    c) Cautious
    d) Avoids risk as much as possible
    """,
    """
    10. Suppose you could replace your current investment portfolio with this new one:
    50 percent chance of Gaining 35 percent or 50 percent chance of Loss
    In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
    Answer between the range of -5 to -35.
    """,
    """
    11. Suppose that in the next 7 years,

    YOUR INCOME

    grows 8% each year

    VS
    INFLATION

    grows 10% a year

    At the end of 7 years, how much will you be able to buy with your income?
    Options:
    a) More than today
    b) Exactly the same
    c) Less than today
    d) Cannot say
    """,
    """
    12. If somebody buys a bond of Company B, which of the following statements seems correct:
    a) She owns part of Company B
    b) She has lent money to Company B
    c) She is liable for Company B's debt
    d) Cannot say
    """,
    """
    13. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
    a) More than ₹2 lakhs
    b) Less than ₹2 lakhs
    c) Exactly ₹2 lakhs
    d) Cannot say
    """,
    """
    14. When an investor spreads money across different asset classes, what happens to the risk of losing money:
    a) Increases
    b) Decreases
    c) Stays the same
    d) Cannot say
    """,
    """
    15. When a country's central bank reduces interest rates, it makes:

    a) Borrowing more attractive and saving less attractive
    b) Borrowing less attractive and saving more attractive
    c) Both borrowing and saving less attractive
    d) Cannot say
    """,
    """
    16. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
    a) Option A : Min Value is 9.2L and Max Value is 11.8L
    b) Option B : Min Value is 8.8L and Max Value is 12.3L
    c) Option C : Min Value is 8.5L and Max Value is 12.8L
    d) Option D : Min Value is 8.1L and Max Value is 13.3L
    e) Option E : Min Value is 7.8L and Max Value is 13.8L
    """,
    """
    17. From Sept 2008 to Nov 2008, Stock market went down by 31%.

    If you owned a stock investment that lost about 31 percent in 3 months, you would:
    a) Sell all of the remaining investment
    b) Sell a portion of the remaining investment
    c) Hold on to the investment and sell nothing
    d) Buy little
    e) Buy more of the investment
    """,
    """
    18. Over any 1-year period, what would be the maximum drop in the value of your investment 
    portfolio that you would be comfortable with?
    a) <5%
    b) 5 - 10%
    c) 10 - 15%
    d) 15 - 20%
    e) >20%
    """,
    """
    19. When investing, what do you consider the most?

    a) Risk 
    b) Return
    """,
    """
    20. What best describes your attitude?

    a) Prefer reasonable returns, can take reasonable risk
    b) Like higher returns, can take slightly higher risk
    c) Want to maximize returns, can take significant high risk
    """,
    """
    21. How much monthly investment you want to do?
    """,
    """
    22. What is the time horizon for your investment?
    You can answer in any range, example 1-5 years."""  
]


# Handler for /start command
@dispatcher.message_handler(commands=['start'])
async def handle_start(message: types.Message):
    """
    This handler receives messages with /start command
    """
    chat_id = message.chat.id
    # Start asking questions
    await start_assessment(chat_id)


# Function to start the assessment
async def start_assessment(chat_id):
    await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
    await ask_next_question(chat_id, 0)

# Function to ask the next question
async def ask_next_question(chat_id, question_index):
    if question_index < len(questions):
        # Ask the next question
        await bot.send_message(chat_id, questions[question_index])
        # Update state to indicate the next expected answer
        states[chat_id] = question_index
    else:
        # No more questions, finish assessment
        await finish_assessment(chat_id)

# Handler for receiving assessment answers
assessment_in_progress = True


async def finish_assessment(chat_id):
    if chat_id in states and states[chat_id] == len(questions):
        # All questions have been answered, now process the assessment
        await bot.send_message(chat_id, "Assessment completed. Thank you!")

        # Determine investment personality based on collected responses
        global investment_personality
        investment_personality = await determine_investment_personality(user_responses)

        # Inform the user about their investment personality
        await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

        # Summarize collected information
        global summary
        summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
        summary = summary + "\n" + "Your investment personality:" + investment_personality
        # Ensure to await the determination of investment personality
        await send_summary_chunks(chat_id, summary)
        global assessment_in_progress 
        assessment_in_progress = False
        # Prompt the user to begin financial advice process
        await bot.send_message(chat_id, "To receive financial advice, please upload a document with your financial details using /begin.")

async def send_summary_chunks(chat_id, summary):
    # Split the summary into chunks that fit within Telegram's message limits
    chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

    # Send each chunk as a separate message
    for chunk in chunks:
        await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


async def determine_investment_personality(assessment_data):
    try:
        # Prepare input text for the chatbot based on assessment data
        input_text = "User Profile:\n"
        for question, answer in assessment_data.items():
            input_text += f"{question}: {answer}\n"

        # Introduce the chatbot's task and prompt for classification
        input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
                      "- Conservative Investor\n" \
                      "- Moderate Investor\n" \
                      "- Aggressive Investor"

        # Use your generative AI model to generate a response
        # print(input_text)
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(input_text)

        # Determine the investment personality from the chatbot's response
        response_text = response.text.lower()
        if "conservative" in response_text:
            personality = "Conservative Investor"
        elif "moderate" in response_text:
            personality = "Moderate Investor"
        elif "aggressive" in response_text:
            personality = "Aggressive Investor"
        else:
            personality = "Unknown"

        return personality
        # Send the determined investment personality back to the user
        #await bot.send_message(chat_id, f"Investment Personality: {personality}")

    except Exception as e:
        print(f"Error generating response: {e}")
        #await bot.send_message(chat_id, "Error processing investment personality classification.")




@dispatcher.message_handler(commands=['help'])
async def helper(message: types.Message):
    """
    A handler to display the help menu.
    """
    help_command = """
    Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
    /start - to start the investment personality assessment.
    /clear - to clear the past conversation and context.
    /help - to get this help menu.
    /begin - to start the Financial Suggestion Conversation with the ChatBot.
    I hope this helps. :)
    """
    await message.reply(help_command)

# Handler for /begin command to initiate financial advice
@dispatcher.message_handler(commands=['begin'])
async def handle_begin(message: types.Message):
    chat_id = message.chat.id
    file_instructions ="""Hello there!My name is Finbot and I am a Wealth Management Advisor Chatbot.I need more details related to your Financial Profile so that I can give you 
    personalised Financial Advice.Please follow this drive link : https://docs.google.com/document/d/1JaqWUXcq3MNrPTCvEy_2RngbGIY0rybX/edit?usp=sharing&ouid=101252223236130106095&rtpof=true&sd=true
    ,there you will find a Word Document.Please fill in your details with correct Information and then upload the Document in the chat. """

    # """
    # Hi there! I'm now in Financial Advisor mode. Please upload a document with your financial details.
    # """
    await message.reply(file_instructions)
    #await bot.send_message(chat_id, "Please upload a document with your financial details.")

# Handler for document upload
@dispatcher.message_handler(content_types=['document'])
async def handle_document(message: types.Message):
    global summary,investment_personality  
    
    # Obtain file information
    file_id = message.document.file_id
    file = await bot.get_file(file_id)
    file_path = file.file_path
    
    # Download the file
    await bot.download_file(file_path, "data/uploaded_file")
    
    # Process the uploaded document
    extracted_text = await process_document("data/uploaded_file")
    
    if extracted_text:
        # Load vector database (assuming this is part of setting up the retriever)
        retriever = await load_vector_db("data/uploaded_file")
        file_path = 'data/uploaded_file'
        client_name, validation_errors = await process_document(file_path)

        # Print results
        print(f"Client Name: {client_name}")
        if validation_errors:
            print("**Validation Errors:**")
            for error in validation_errors:
                print(error)
        else:
            print("All fields are filled correctly.")

        await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")
 
        if retriever is None:
            await message.reply("The retrieval chain is not set up. Please upload a document first.")
            return

        # Check if a valid chain can be created
        chain = await make_retrieval_chain(retriever)
        if chain is None:
            await message.reply("Failed to create the retrieval chain.")
            return
        
        try:     
            query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
            response = chain.invoke({"input": query})
            print(response['answer'])
            global chat_history
            chat_history = response['answer'] 
            print(f"\n Chat History : {chat_history}")
            await message.reply(response['answer'])

        except Exception as e:
            print(f"Error invoking retrieval chain on attempt : {e}")
   
    else:
        await message.reply("Failed to process the uploaded file.")



async def process_document(file_path):
    try:
        file_type = filetype.guess(file_path)
        if file_type is not None:
            if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # return extract_text_from_word(file_path)
                return extract_text_and_tables_from_word(file_path)
            elif file_type.mime == "application/pdf":
                return extract_text_from_pdf(file_path)
        return None
    except Exception as e:
        print(f"Error processing document: {e}")
        return None

def extract_text_from_pdf(pdf_file_path):
    try:
        with open(pdf_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            text_content = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text_content.append(page.extract_text())
            return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None


import re
import docx

def extract_text_and_tables_from_word(docx_file_path):
    """
    Extracts text and tables from a Word document (.docx).

    Args:
        docx_file_path (str): Path to the Word document file.

    Returns:
        tuple: Extracted text content and tables from the document.
    """
    try:
        doc = docx.Document(docx_file_path)
        text_content = []
        tables_content = []

        for para in doc.paragraphs:
            text_content.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)

        return "\n".join(text_content), tables_content
    except Exception as e:
        print(f"Error extracting text and tables from Word document: {e}")
        return None, None

def validate_document_content(text, tables):
    """
    Validates the content of the document.

    Args:
        text (str): Extracted text content from the document.
        tables (list): Extracted tables content from the document.

    Returns:
        tuple: Client name and validation errors.
    """
    errors = []
    
    # Extract client name
    client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
    client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

    # Define required sections
    required_sections = [
        "YOUR RETIREMENT GOAL",
        "YOUR OTHER MAJOR GOALS",
        "YOUR ASSETS AND LIABILITIES",
        "MY LIABILITIES",
        "YOUR CURRENT ANNUAL INCOME"
    ]

    # Check for the presence of required sections
    for section in required_sections:
        if section not in text:
            errors.append(f"* {section} section missing.")
    
    # Define table field checks
    table_checks = {
        "YOUR RETIREMENT GOAL": [
            r"When do you plan to retire\? \(age or date\)",
            r"Social Security Benefit \(include expected start date\)",
            r"Pension Benefit \(include expected start date\)",
            r"Other Expected Income \(rental, part-time work, etc.\)",
            r"Estimated Annual Retirement Expense"
        ],
        "YOUR OTHER MAJOR GOALS": [
            r"GOAL", r"COST", r"WHEN"
        ],
        "YOUR ASSETS AND LIABILITIES": [
            r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
            r"Current Value", r"Annual Contributions"
        ],
        "MY LIABILITIES": [
            r"Balance", r"Interest Rate", r"Monthly Payment"
        ]
    }

    # Validate table content
    for section, checks in table_checks.items():
        section_found = False
        for table in tables:
            table_text = "\n".join(["\t".join(row) for row in table])
            if section in table_text:
                section_found = True
                for check in checks:
                    if not re.search(check, table_text, re.IGNORECASE):
                        errors.append(f"* Missing or empty field in {section} section: {check}")
                break
        if not section_found:
            errors.append(f"* {section} section missing.")

    return client_name, errors

async def process_document(file_path):
    try:
        text, tables = extract_text_and_tables_from_word(file_path)
        if text is not None and tables is not None:
            client_name, errors = validate_document_content(text, tables)
            return client_name, errors
        return None, ["Error processing document."]
    except Exception as e:
        print(f"Error processing document: {e}")
        return None, [f"Error processing document: {e}"]



async def load_vector_db(file_path):
    try:
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        return vector_store.as_retriever(search_kwargs={"k": 1})
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None


async def make_retrieval_chain(retriever):
    """
    Create a retrieval chain using the provided retriever.

    Args:
        retriever (RetrievalQA): A retriever object.

    Returns:
        RetrievalQA: A retrieval chain object.
    """
    try:
        global investment_personality,summary
        llm = ChatGoogleGenerativeAI(
            model="gemini-pro",
            temperature=0.7,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )

        prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
                Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
                Also give the user detailed information about the investment how to invest,where to invest and how much they
                should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
                investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
                investment.Also explain the user why you are giving them that particular
                investment suggestion.Answer in 3-4 lines.\n
                <context>
                {context}
                </context>
                Question: {input}"""

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        # Update combine_docs_chain with your actual document combining logic
        combine_docs_chain = None  # Replace this with your combine_docs_chain

        if retriever is not None :  #and combine_docs_chain is not None:
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in creating chain: {e}")
        return None


@dispatcher.message_handler()
async def main_bot(message: types.Message):
    global retriever, extracted_text,investment_personality,summary,chat_history

    # Handle the first tasks assessments answers from the user
    chat_id = message.chat.id

    if chat_id in states and states[chat_id] < len(questions):
        # Retrieve the index of the current question
        question_index = states[chat_id]

        # Save the user's response to the current question
        answer = message.text
        user_responses[questions[question_index]] = answer
        states[chat_id] += 1  # Move to the next question

        # Ask the next question
        await ask_next_question(chat_id, question_index + 1)
    else:
        # Handle q&a chat messages using your Gemini model (llm)
        try:

            task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
            Also give the user detailed information about the investment how to invest,where to invest and how much they
            should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
            investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
            investment.Also explain the user why you are giving them that particular
            investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
            User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
            So dont mention user to consult to a financial expert."""
        
            model = genai.GenerativeModel('gemini-pro')
            print(investment_personality)
            # query = task + "\n" + investment_personality + "\n" + summary + "\n" +  extracted_text + "\n"  +   message.text

            # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +  extracted_text + "\n"  +   message.text
            query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
            print(f"\nQuery : {query}")
            response = model.generate_content(query)
            await message.reply(response.text) #(response['answer']) 
           
        except (NetworkError, RetryAfter) as e:
            print(f"Network error: {e}. Retrying...")
            await asyncio.sleep(5)  # Wait before retrying
            await main_bot(message)  # Retry the message handling
        except TelegramAPIError as e:
            print(f"Telegram API error: {e}")
            await message.reply("An error occurred while communicating with Telegram. Please try again later.")
        except Exception as e:
            print(f"Error processing general chat message: {e}")
            await message.reply("Failed to process your request.")
        


if __name__ == "__main__":
    executor.start_polling(dispatcher, skip_updates=True)




# import re

# def validate_document_content(file_path):
#     text = extract_text_from_word(file_path)

#     if not text:
#         return "Unknown", ["Failed to extract text from the document."]

#     # Extract client name
#     # client_name_match = re.search(r"Client Name:\s*(.*)", text, re.IGNORECASE)
#     client_name_match = re.search(r"Client Name:\s*([^\n,]*)", text, re.IGNORECASE)
#     client_name = client_name_match.group(1).strip() if client_name_match else "Unknown"

#     errors = []

#     # YOUR RETIREMENT GOAL section
#     retirement_goal_section = re.search(r"YOUR\s+RETIREMENT\s+GOAL\s*([\s\S]*?)\s*(?:YOUR\s+OTHER\s+MAJOR\s+GOALS|YOUR\s+ASSETS\s+AND\s+LIABILITIES|MY\s+LIABILITIES)", text, re.IGNORECASE)
#     if retirement_goal_section:
#         retirement_goal_text = retirement_goal_section.group(1).strip()
#         required_fields_retirement = [
#             "Client",
#             "Co-Client",
#             "When do you plan to retire?",
#             "Social Security Benefit",
#             "Pension Benefit",
#             "Other Expected Income",
#             "Estimated Annual Retirement Expense"
#         ]
#         for field in required_fields_retirement:
#             if not re.search(rf"{field}.*?\S+", retirement_goal_text, re.IGNORECASE):
#                 errors.append(f"Missing or empty field in YOUR RETIREMENT GOAL section: {field}")
#     else:
#         errors.append("YOUR RETIREMENT GOAL section missing.")

#     # YOUR OTHER MAJOR GOALS section
#     major_goals_section = re.search(r"YOUR\s+OTHER\s+MAJOR\s+GOALS\s*([\s\S]*?)\s*(?:YOUR\s+ASSETS\s+AND\s+LIABILITIES|MY\s+LIABILITIES)", text, re.IGNORECASE)
#     if major_goals_section:
#         major_goals_text = major_goals_section.group(1).strip()
#         if not re.search(r"Goal.*?\S+", major_goals_text, re.IGNORECASE):
#             errors.append("At least one goal must be filled in YOUR OTHER MAJOR GOALS section.")
#         if not re.search(r"Cost.*?\S+", major_goals_text, re.IGNORECASE):
#             errors.append("Cost must be filled for at least one goal in YOUR OTHER MAJOR GOALS section.")
#         if not re.search(r"When.*?\S+", major_goals_text, re.IGNORECASE):
#             errors.append("When must be filled for at least one goal in YOUR OTHER MAJOR GOALS section.")
#     else:
#         errors.append("YOUR OTHER MAJOR GOALS section missing.")

#     # YOUR ASSETS AND LIABILITIES section
#     assets_liabilities_section = re.search(r"YOUR\s+ASSETS\s+AND\s+LIABILITIES\s*([\s\S]*?)\s*(?:MY\s+LIABILITIES|YOUR\s+NET\s+WORTH)", text, re.IGNORECASE)
#     if assets_liabilities_section:
#         assets_liabilities_text = assets_liabilities_section.group(1).strip()
#         required_fields_assets = [
#             "Cash/bank accounts",
#             "Home",
#             "Other Real Estate",
#             "Business",
#             "Current Value",
#             "Annual Contributions"
#         ]
#         for field in required_fields_assets:
#             if not re.search(rf"{field}.*?\S+", assets_liabilities_text, re.IGNORECASE):
#                 errors.append(f"Missing or empty field in YOUR ASSETS AND LIABILITIES section: {field}")
#     else:
#         errors.append("YOUR ASSETS AND LIABILITIES section missing.")

#     # MY LIABILITIES section
#     liabilities_section = re.search(r"MY\s+LIABILITIES\s*([\s\S]*?)\s*(?:YOUR\s+NET\s+WORTH|$)", text, re.IGNORECASE)
#     if liabilities_section:
#         liabilities_text = liabilities_section.group(1).strip()
#         required_fields_liabilities = [
#             "Mortgage",
#             "Credit Card",
#             "Student Loans",
#             "Other Loans"
#         ]
#         for field in required_fields_liabilities:
#             if not re.search(rf"{field}.*?\S+", liabilities_text, re.IGNORECASE):
#                 errors.append(f"Missing or empty field in MY LIABILITIES section: {field}")
#     else:
#         errors.append("MY LIABILITIES section missing.")

#     return client_name, errors



# validate with text :

# def validate_extracted_text(text):
#     errors = []
    
#     # Extract information from text
#     lines = text.split('\n')
    
#     # Validate client and co-client information
#     client_info = [line for line in lines if "Client Name" in line]
#     if client_info:
#         client_info = client_info[0]
#         if "Client Name" not in client_info or "Age" not in client_info:
#             errors.append("Client Name or Age missing.")
#         if "Co-Client Name" not in client_info or "Age" not in client_info:
#             errors.append("Co-Client Name or Age missing.")
#     else:
#         errors.append("Client and Co-Client information missing.")
    
#     # Validate Retirement Goal
#     retirement_goal_section = [line for line in lines if "YOUR RETIREMENT GOAL" in line]
#     if retirement_goal_section:
#         start_index = lines.index(retirement_goal_section[0])
#         end_index = start_index + 6
#         retirement_goal_lines = lines[start_index:end_index]

#         required_fields = [
#             "When do you plan to retire? (age or date)",
#             "Social Security Benefit (include expected start date)",
#             "Pension Benefit (include expected start date)",
#             "Other Expected Income (rental, part-time work, etc.)",
#             "Estimated Annual Retirement Expense ($ or % of current salary)"
#         ]
        
#         for field in required_fields:
#             if not any(field in line for line in retirement_goal_lines):
#                 errors.append(f"Missing or incomplete data for: {field}")
#     else:
#         errors.append("YOUR RETIREMENT GOAL section missing.")
    
#     # Validate Other Major Goals
#     other_goals_section = [line for line in lines if "YOUR OTHER MAJOR GOALS" in line]
#     if other_goals_section:
#         goal_lines = lines[lines.index(other_goals_section[0])+1:lines.index(other_goals_section[0])+6]
#         goals_filled = any("GOAL" in goal for goal in goal_lines)
#         costs_filled = any("COST" in cost for cost in goal_lines)
#         when_filled = any("WHEN" in when for when in goal_lines)
#         if not (goals_filled and costs_filled and when_filled):
#             errors.append("At least one goal, along with its cost and timeframe, must be filled in YOUR OTHER MAJOR GOALS.")
#     else:
#         errors.append("YOUR OTHER MAJOR GOALS section missing.")
    
#     # Validate Assets and Liabilities
#     assets_section = [line for line in lines if "YOUR ASSETS AND LIABILITIES" in line]
#     if assets_section:
#         asset_lines = lines[lines.index(assets_section[0])+1:lines.index(assets_section[0])+11]
#         required_assets = ["Cash/bank accounts", "Home", "Other Real Estate", "Business"]
#         for asset in required_assets:
#             if not any(asset in line for line in asset_lines):
#                 errors.append(f"{asset} details missing in YOUR ASSETS AND LIABILITIES.")
#     else:
#         errors.append("YOUR ASSETS AND LIABILITIES section missing.")
    
#     # Validate Liabilities
#     liabilities_section = [line for line in lines if "MY LIABILITIES" in line]
#     if liabilities_section:
#         liability_lines = lines[lines.index(liabilities_section[0])+1:lines.index(liabilities_section[0])+4]
#         required_liabilities = ["Mortgage(s)", "Credit Card(s)", "Other loans"]
#         for liability in required_liabilities:
#             if not any(liability in line for line in liability_lines):
#                 errors.append(f"{liability} details missing in MY LIABILITIES.")
#     else:
#         errors.append("MY LIABILITIES section missing.")
    
#     return "\n".join(errors) if errors else None













# # Able to identify the investment personality of the user 

# import os
# import filetype
# import docx
# import PyPDF2
# from aiogram import Bot, Dispatcher, executor, types
# from dotenv import load_dotenv
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
# from langchain_community.document_loaders import Docx2txtLoader
# from langchain_core.prompts import ChatPromptTemplate
# from langchain.chains import create_retrieval_chain
# from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# from langchain.chains.combine_documents import create_stuff_documents_chain
# from langchain.memory import ConversationSummaryMemory

# import google.generativeai as genai

# load_dotenv()

# TOKEN = os.getenv("TOKEN")
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # Configure generativeai with your API key
# genai.configure(api_key=GOOGLE_API_KEY)

# # Initialize bot
# bot = Bot(token=TOKEN)
# dispatcher = Dispatcher(bot)

# rag_on = False
# retriever = None  # Store retriever globally


# class Reference:
#     def __init__(self):
#         self.response = ""


# reference = Reference()


# def clear_past():
#     reference.response = ""


# @dispatcher.message_handler(commands=['clear'])
# async def clear(message: types.Message):
#     """
#     A handler to clear the previous conversation and context.
#     """
#     clear_past()
#     await message.reply("I've cleared the past conversation and context.")


# # Store user states
# states = {}
# # Dictionary to hold question-answer pairs
# user_responses = {}

# # Define Questions for assessment
# questions = [
#     "1. What is your name?",
#     "2. What is your age?",
#     "3. What is your gender?",
#     "4. What is your marital status?",
#     """
#         5. Singapore plans to build a new observation tower called 'The Rook'.
#         How many steps do you think it will take to go to the top floor?

#         a) Less than 500 
#         b) More than 500

#     """,
#     "6. Now Guess the number of steps" ,
#     """
#     7. How confident are you that the real number is in the range you have selected? 
#     Answer within a range of 100.  
#     """,
#     """ 
#     8. You and your friend are betting on a series of coin tosses.

#     He always bets ₹2,000 on Heads

#     You always bet ₹2,000 on Tails

#     Winner of last 8 turns

#     You lost ₹8,000 in the last 4 turns!

#     If you were to bet one last time, what would you bet on heads or tails ?
#     """ ,
#     """
#     9. How confident are you that your bet will win this time?
#     Answer how confident you are. 
#     (Example: Not confident at all, somewhat confident, confident, or Very confident)
#     """,
#     """
#     10. What do you think are the chances that you will achieve and maintain your financial goals within the next 
#     10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
#     Answer how likely you are to achieve your goal.
#     (Example: Less likely than others, likely than others, or More likely than others)
#     """,
#     """
#     11. Imagine you are a contestant in a game show, and you are presented the following choices.

#     What would you prefer?
#     a) 50 percent chance of winning 15 gold coins 
#     b) 100 percent chance of winning 8 gold coins
#     """,
#     """
#     12. Ok, one last choice...

#     What would you prefer?
#     a) 50 percent chance of winning 15 gold coins 
#     b) 100 percent chance of winning 2 gold coins
#     """,
#     """
#     13. In general, how would your best friend describe your risk-taking tendencies?
#     a) A real gambler
#     b) Willing to take risks after completing adequate research
#     c) Cautious
#     d) Avoids risk as much as possible
#     """,
#     """
#     14. Suppose you could replace your current investment portfolio with this new one:
#     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
#     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
#     Answer between the range of -5 to -35.
#     """,
#     """
#     15. Suppose that in the next 7 years,

#     YOUR INCOME

#     grows 8% each year

#     VS
#     INFLATION

#     grows 10% a year

#     At the end of 7 years, how much will you be able to buy with your income?
#     Options:
#     a) More than today
#     b) Exactly the same
#     c) Less than today
#     d) Cannot say
#     """,
#     """
#     16. If somebody buys a bond of Company B, which of the following statements seems correct:
#     a) She owns part of Company B
#     b) She has lent money to Company B
#     c) She is liable for Company B's debt
#     d) Cannot say
#     """,
#     """
#     17. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
#     a) More than ₹2 lakhs
#     b) Less than ₹2 lakhs
#     c) Exactly ₹2 lakhs
#     d) Cannot say
#     """,
#     """
#     18. When an investor spreads money across different asset classes, what happens to the risk of losing money:
#     a) Increases
#     b) Decreases
#     c) Stays the same
#     d) Cannot say
#     """,
#     """
#     19. When a country's central bank reduces interest rates, it makes:

#     a) Borrowing more attractive and saving less attractive
#     b) Borrowing less attractive and saving more attractive
#     c) Both borrowing and saving less attractive
#     d) Cannot say
#     """,
#     """
#     20. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
#     a) Option A : Min Value is 9.2L and Max Value is 11.8L
#     b) Option B : Min Value is 8.8L and Max Value is 12.3L
#     c) Option C : Min Value is 8.5L and Max Value is 12.8L
#     d) Option D : Min Value is 8.1L and Max Value is 13.3L
#     e) Option E : Min Value is 7.8L and Max Value is 13.8L
#     """,
#     """
#     21. From Sept 2008 to Nov 2008, Stock market went down by 31%.

#     If you owned a stock investment that lost about 31 percent in 3 months, you would:
#     a) Sell all of the remaining investment
#     b) Sell a portion of the remaining investment
#     c) Hold on to the investment and sell nothing
#     d) Buy little
#     e) Buy more of the investment
#     """,
#     """
#     22. Over any 1-year period, what would be the maximum drop in the value of your investment 
#     portfolio that you would be comfortable with?
#     a) <5%
#     b) 5 - 10%
#     c) 10 - 15%
#     d) 15 - 20%
#     e) >20%
#     """,
#     """
#     23. When investing, what do you consider the most?

#     a) Risk 
#     b) Return
#     """,
#     """
#     24. What best describes your attitude?

#     a) Prefer reasonable returns, can take reasonable risk
#     b) Like higher returns, can take slightly higher risk
#     c) Want to maximize returns, can take significant high risk
#     """,
#     """
#     25. What is your approximate monthly take-home salary/income?
#     You may enter the salary that gets credited to your bank account or a rough estimate of your monthly income.
#     """,
#     """
#     26. What is your monthly expenditure?

#     How much do you end up roughly spending every month? Do not include EMIs.
#     """,
#     """
#     27. How much are your assets worth?

#     Some common assets include investments, deposits, real estate, etc.
#     NOTE: For this assessment, please do not include the house you live in and the gold you use under "assets".
#     """,
#     """
#     28. What is your total liability?

#     A liability is something that you owe. Some common liabilities are home loans, car loans, debt, etc.
#     """,
#     """
#     29. What is your total EMI amount?

#     You usually pay EMIs for the loans you take for things such as education, home, car, etc.
#     """
# ]




# @dispatcher.message_handler(commands=['start'])
# async def welcome(message: types.Message):
#     """
#     This handler receives messages with /start command
#     """
#     await message.reply("Hi! I am a Chat Bot. Let's start a quick personality assessment.")

#     # Start asking questions
#     await ask_next_question(message.chat.id, 0)

# async def ask_next_question(chat_id, question_index):
#     if question_index < len(questions):
#         # Ask the next question
#         await bot.send_message(chat_id, questions[question_index])
#         # Update state to indicate the next expected answer
#         states[chat_id] = question_index
#     else:
#         # No more questions, finish assessment
#         await process_assessment(chat_id)

# @dispatcher.message_handler()
# async def handle_message(message: types.Message):
#     chat_id = message.chat.id

#     if chat_id in states:
#         # Retrieve the index of the current question
#         question_index = states[chat_id]

#         # Save the user's response to the current question
#         answer = message.text
#         user_responses[questions[question_index]] = answer  # Store question-answer pair
#         states[chat_id] += 1  # Move to the next question

#         # Ask the next question
#         await ask_next_question(chat_id, question_index + 1)
#     else:
#         # If the state is not found, it means the user needs to start the assessment
#         await welcome(message)

# async def process_assessment(chat_id):
#     # Process the collected responses
#     if chat_id in states and states[chat_id] == len(questions):
#         # All questions have been answered, now process the assessment
#         await bot.send_message(chat_id, "Assessment completed. Thank you!")

#         # Prepare data to pass to the chatbot for processing
#         assessment_data = user_responses  # Dictionary containing question-answer pairs

#         # Call the chatbot for further processing (example)
#         await process_chatbot_request(chat_id, assessment_data)

# async def process_chatbot_request(chat_id, assessment_data):
#     try:
#         # Prepare input text for the chatbot based on assessment data
#         input_text = "User Profile:\n"
#         for question, answer in assessment_data.items():
#             input_text += f"{question}: {answer}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
#                       "- Conservative Investor\n" \
#                       "- Moderate Investor\n" \
#                       "- Aggressive Investor"

#         # Use your generative AI model to generate a response
#         print(input_text)
#         model = genai.GenerativeModel('gemini-pro')
#         response = model.generate_content(input_text)

#         # Determine the investment personality from the chatbot's response
#         response_text = response.text.lower()
#         if "conservative" in response_text:
#             personality = "Conservative Investor"
#         elif "moderate" in response_text:
#             personality = "Moderate Investor"
#         elif "aggressive" in response_text:
#             personality = "Aggressive Investor"
#         else:
#             personality = "Unknown"

#         # Send the determined investment personality back to the user
#         await bot.send_message(chat_id, f"Investment Personality: {personality}")

#     except Exception as e:
#         print(f"Error generating response: {e}")
#         await bot.send_message(chat_id, "Error processing investment personality classification.")


# # async def process_chatbot_request(chat_id, assessment_data):
# #     try:
# #         # Prepare input for the chatbot based on assessment data
# #         # You can format the input according to your chatbot's requirements
# #         input_text = f"User Profile:\n"
# #         for question, answer in assessment_data.items():
# #             input_text += f"{question}: {answer}\n"

# #         input_text = input_text + "\n" + "You are an investment Personality identifier.You must classify the users investment personality" 
# #         input_text = input_text + "\n" + """You should classify the personality as "Conservative Investor"
# #                                     or "Moderate Investor" or "Aggressive Investor" """
        
# #         # Initialize your llm 
# #         # llm = ChatGoogleGenerativeAI(
# #         #     model="gemini-pro",
# #         #     google_api_key=GOOGLE_API_KEY,
# #         #     temperature=0.7,
# #         #     top_p=0.85
# #         # )
# #         # Generate response from the chatbot based on the input text
# #         # response = await llm.generate_text(input_text)

# #         model = genai.GenerativeModel('gemini-pro')
# #         response = model.generate_content(question)#response = await model.generate_content(question)
# #         print(response.text)
# #         await bot.send_message(chat_id, response.text)
        

# #         # Determine investment personality based on the chatbot's response
        
# #         # if "conservative" in response.lower():
# #         #     personality = "Conservative Investor"
# #         # elif "moderate" in response.lower():
# #         #     personality = "Moderate Investor"
# #         # elif "aggressive" in response.lower():
# #         #     personality = "Aggressive Investor"
# #         # else:
# #         #     personality = "Unknown"

# #         # Send the determined investment personality back to the user
# #         # await bot.send_message(chat_id, f"Investment Personality: {personality}")

# #     except Exception as e:
# #         # Handle any errors that may occur during chatbot processing
# #         print(f"Error generating response: {e}")
# #         await bot.send_message(chat_id, f"Error generating response: {e}")




# @dispatcher.message_handler(commands=['help'])
# async def helper(message: types.Message):
#     """
#     A handler to display the help menu.
#     """
#     help_command = """
#     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
#     /start - to start the conversation
#     /clear - to clear the past conversation and context.
#     /help - to get this help menu.
#     /begin - to start the Financial Suggestion Conversation with the ChatBot.
#     I hope this helps. :)
#     """
#     await message.reply(help_command)



# async def process_document(file_path):
#     try:
#         file_type = filetype.guess(file_path)
#         if file_type is not None:
#             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 return extract_text_from_word(file_path)
#             elif file_type.mime == "application/pdf":
#                 return extract_text_from_pdf(file_path)
#         return None
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None

# def extract_text_from_word(docx_file_path):
#     """
#     Extracts text content from a Word document (.docx).

#     Args:
#         docx_file_path (str): Path to the Word document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         doc = docx.Document(docx_file_path)
#         text_content = []
#         for para in doc.paragraphs:
#             text_content.append(para.text)
#         return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from Word document: {e}")
#         return None

# def extract_text_from_pdf(pdf_file_path):
#     """
#     Extracts text content from a PDF file.

#     Args:
#         pdf_file_path (str): Path to the PDF file.

#     Returns:
#         str: Extracted text content from the PDF.
#     """
#     try:
#         with open(pdf_file_path, "rb") as pdf_file:
#             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#             text_content = []
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text_content.append(page.extract_text())
#             return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None

# async def load_vector_db(file_path):
#     try:
#         loader = Docx2txtLoader(file_path)
#         documents = loader.load()
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         text_chunks = text_splitter.split_documents(documents)
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
#         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
#         return vector_store.as_retriever(search_kwargs={"k": 1})
#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None



# async def make_retrieval_chain(retriever):
#     """
#     Create a retrieval chain using the provided retriever.

#     Args:
#         retriever (RetrievalQA): A retriever object.

#     Returns:
#         RetrievalQA: A retrieval chain object.
#     """
#     try:
#         llm = ChatGoogleGenerativeAI(
#             model="gemini-pro",
#             temperature=0.7,
#             top_p=0.85,
#             google_api_key=GOOGLE_API_KEY
#         )

#         prompt_template = """You are a Financial assistant for question-answering tasks.
#                 You are supposed to give financial advice to the user related to the details shared to you in the file
#                 that can help them gain good knowledge about investing.
#                 You are supposed to tell about the risk taking personality with respect to investment of the user. 
#                 If you don't know the answer, just say that you don't know.
#                 Use five sentences maximum and keep the answer concise.\n
#                 <context>
#                 {context}
#                 </context>
#                 Question: {input}"""
        
#                 #Question: {question}\nContext: {context}\nAnswer:"""

#         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

#         document_chain = create_stuff_documents_chain(llm, llm_prompt)
#         # Update combine_docs_chain with your actual document combining logic
#         combine_docs_chain = None  # Replace this with your combine_docs_chain

#         if retriever is not None :  #and combine_docs_chain is not None:
#             retriever_chain = create_retrieval_chain(retriever,document_chain) 
#             #combine_docs_chain)
#             # response = retriever_chain.invoke({"input":"Give me client detail "})
#             # print(response['answer'])
#             print(retriever_chain)
#             return retriever_chain
#         else:
#             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
#             return None

#     except Exception as e:
#         print(f"Error in creating chain: {e}")
#         return None



# @dispatcher.message_handler(commands=['begin'])
# async def finance_bot_activate(message: types.Message):
#     global rag_on, retriever
#     rag_on = True
#     file_instructions = """
#     Hi there! I'm now in Financial Advisor mode. Please upload a document with your financial details.
#     """
#     await message.reply(file_instructions)

# @dispatcher.message_handler(content_types=['document'])
# async def get_file(message: types.Message):
#     global retriever
#     file_id = message.document.file_id
#     file = await bot.get_file(file_id)
#     file_path = file.file_path
#     await bot.download_file(file_path, "data/uploaded_file")
#     extracted_text = await process_document("data/uploaded_file")
#     if extracted_text:
#         retriever = await load_vector_db("data/uploaded_file")
#         if retriever:
#             await message.reply("Processed file, please ask your question.")
#         else:
#             await message.reply("Failed to load vector database.")
#     else:
#         await message.reply("Failed to process the uploaded file.")

# @dispatcher.message_handler()
# async def main_bot(message: types.Message):
#     global rag_on, retriever,extracted_text

#     if rag_on:
#         if retriever is None:
#             await message.reply("The retrieval chain is not set up. Please upload a document first.")
#             return

#         # Check if a valid chain can be created
#         chain = await make_retrieval_chain(retriever)
#         if chain is None:
#             await message.reply("Failed to create the retrieval chain.")
#             return
#         # Prepare the input data for the retrieval chain
#         # input_data = {
#         #     "context": extracted_text,  # Use extracted text as the context
#         #     "question": message.text
#         # }
#         # Invoke the retrieval chain with the user's message
#         try:
#             response =  chain.invoke({"input":message.text}) 

#             # response = await chain.invoke({"input":message.text}) 
#             #(input_data) 
#             #(({"context":chain,"question":message.text}))
#             #({"context":retriever,"question":message.text}) 
#             #({"input":message.text}) #(message.text)
#             print(response['answer'])
#             await message.reply(response['answer']) #(response.text)
#         except Exception as e:
#             print(f"Error invoking retrieval chain: {e}")
#             await message.reply("Failed to process your request.")
#     else:
#         # Handle general chat messages using your Gemini model (llm)
#         print(f">>> USER: \n\t{message.text}")
#         # Use your Gemini model here for general chat responses
#         try:
#             # Example: response = await get_gemini_response(message.text)
#             response = "Hello! How can I assist you today?"
#             print(f">>> gemini: \n\t{response}")
#             await bot.send_message(chat_id=message.chat.id, text=response)
#         except Exception as e:
#             print(f"Error processing general chat message: {e}")
#             await message.reply("Failed to process your request.")

# # async def get_gemini_response(question):
#     # model = genai.GenerativeModel('gemini-pro')
#     # response = model.generate_content(question)
#     # return response.text

# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)







# Best Code it is performing well for the second part needs more information and context information

# import os
# import filetype
# import docx
# import PyPDF2
# from aiogram import Bot, Dispatcher, executor, types
# from dotenv import load_dotenv
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
# from langchain_community.document_loaders import Docx2txtLoader
# from langchain_core.prompts import ChatPromptTemplate
# from langchain.chains import create_retrieval_chain
# from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# from langchain.chains.combine_documents import create_stuff_documents_chain
# from langchain.memory import ConversationSummaryMemory

# load_dotenv()

# TOKEN = os.getenv("TOKEN")
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# bot = Bot(token=TOKEN)
# dispatcher = Dispatcher(bot)

# rag_on = False
# retriever = None  # Store retriever globally

# async def process_document(file_path):
#     try:
#         file_type = filetype.guess(file_path)
#         if file_type is not None:
#             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 return extract_text_from_word(file_path)
#             elif file_type.mime == "application/pdf":
#                 return extract_text_from_pdf(file_path)
#         return None
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None

# def extract_text_from_word(docx_file_path):
#     """
#     Extracts text content from a Word document (.docx).

#     Args:
#         docx_file_path (str): Path to the Word document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         doc = docx.Document(docx_file_path)
#         text_content = []
#         for para in doc.paragraphs:
#             text_content.append(para.text)
#         return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from Word document: {e}")
#         return None

# def extract_text_from_pdf(pdf_file_path):
#     """
#     Extracts text content from a PDF file.

#     Args:
#         pdf_file_path (str): Path to the PDF file.

#     Returns:
#         str: Extracted text content from the PDF.
#     """
#     try:
#         with open(pdf_file_path, "rb") as pdf_file:
#             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#             text_content = []
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text_content.append(page.extract_text())
#             return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None

# async def load_vector_db(file_path):
#     try:
#         loader = Docx2txtLoader(file_path)
#         documents = loader.load()
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         text_chunks = text_splitter.split_documents(documents)
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
#         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
#         return vector_store.as_retriever(search_kwargs={"k": 1})
#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None



# async def make_retrieval_chain(retriever):
#     """
#     Create a retrieval chain using the provided retriever.

#     Args:
#         retriever (RetrievalQA): A retriever object.

#     Returns:
#         RetrievalQA: A retrieval chain object.
#     """
#     try:
#         llm = ChatGoogleGenerativeAI(
#             model="gemini-pro",
#             temperature=0.7,
#             top_p=0.85,
#             google_api_key=GOOGLE_API_KEY
#         )

#         prompt_template = """You are a Financial assistant for question-answering tasks.
#                 You are supposed to give financial advice to the user related to the details shared to you in the file
#                 that can help them gain good knowledge about investing.
#                 You are supposed to tell about the risk taking personality with respect to investment of the user. 
#                 If you don't know the answer, just say that you don't know.
#                 Use five sentences maximum and keep the answer concise.\n
#                 <context>
#                 {context}
#                 </context>
#                 Question: {input}"""
        
#                 #Question: {question}\nContext: {context}\nAnswer:"""

#         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

#         document_chain = create_stuff_documents_chain(llm, llm_prompt)
#         # Update combine_docs_chain with your actual document combining logic
#         combine_docs_chain = None  # Replace this with your combine_docs_chain

#         if retriever is not None :  #and combine_docs_chain is not None:
#             retriever_chain = create_retrieval_chain(retriever,document_chain) #combine_docs_chain)
#             # response = retriever_chain.invoke({"input":"Give me client detail "})
#             # print(response['answer'])
#             print(retriever_chain)
#             return retriever_chain
#         else:
#             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
#             return None

#     except Exception as e:
#         print(f"Error in creating chain: {e}")
#         return None



# @dispatcher.message_handler(commands=['begin'])
# async def finance_bot_activate(message: types.Message):
#     global rag_on, retriever
#     rag_on = True
#     file_instructions = """
#     Hi there! I'm now in Financial Advisor mode. Please upload a document with your financial details.
#     """
#     await message.reply(file_instructions)

# @dispatcher.message_handler(content_types=['document'])
# async def get_file(message: types.Message):
#     global retriever
#     file_id = message.document.file_id
#     file = await bot.get_file(file_id)
#     file_path = file.file_path
#     await bot.download_file(file_path, "data/uploaded_file")
#     extracted_text = await process_document("data/uploaded_file")
#     if extracted_text:
#         retriever = await load_vector_db("data/uploaded_file")
#         if retriever:
#             await message.reply("Processed file, please ask your question.")
#         else:
#             await message.reply("Failed to load vector database.")
#     else:
#         await message.reply("Failed to process the uploaded file.")

# @dispatcher.message_handler()
# async def main_bot(message: types.Message):
#     global rag_on, retriever,extracted_text

#     if rag_on:
#         if retriever is None:
#             await message.reply("The retrieval chain is not set up. Please upload a document first.")
#             return

#         # Check if a valid chain can be created
#         chain = await make_retrieval_chain(retriever)
#         if chain is None:
#             await message.reply("Failed to create the retrieval chain.")
#             return
#         # Prepare the input data for the retrieval chain
#         # input_data = {
#         #     "context": extracted_text,  # Use extracted text as the context
#         #     "question": message.text
#         # }
#         # Invoke the retrieval chain with the user's message
#         try:
#             response =  chain.invoke({"input":message.text}) 

#             # response = await chain.invoke({"input":message.text}) 
#             #(input_data) 
#             #(({"context":chain,"question":message.text}))
#             #({"context":retriever,"question":message.text}) 
#             #({"input":message.text}) #(message.text)
#             print(response['answer'])
#             await message.reply(response['answer']) #(response.text)
#         except Exception as e:
#             print(f"Error invoking retrieval chain: {e}")
#             await message.reply("Failed to process your request.")
#     else:
#         # Handle general chat messages using your Gemini model (llm)
#         print(f">>> USER: \n\t{message.text}")
#         # Use your Gemini model here for general chat responses
#         try:
#             # Example: response = await get_gemini_response(message.text)
#             response = "Hello! How can I assist you today?"
#             print(f">>> gemini: \n\t{response}")
#             await bot.send_message(chat_id=message.chat.id, text=response)
#         except Exception as e:
#             print(f"Error processing general chat message: {e}")
#             await message.reply("Failed to process your request.")


# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)




# Message error and chain errors 

# import os
# import filetype
# import docx
# import PyPDF2
# from aiogram import Bot, Dispatcher, executor, types
# from dotenv import load_dotenv
# import google.generativeai as genai
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
# from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# from langchain.memory import ConversationSummaryMemory
# from langchain.chains import ConversationalRetrievalChain
# from langchain_core.prompts import PromptTemplate
# from langchain.schema import StrOutputParser
# from langchain_community.document_loaders import Docx2txtLoader

# from langchain.schema.runnable import RunnablePassthrough
# from langchain.chains import RetrievalQA
# from langchain_community.document_loaders import DirectoryLoader
# from langchain_community.document_loaders import TextLoader
# from langchain_core.prompts import ChatPromptTemplate
# from langchain.chains.combine_documents import create_stuff_documents_chain
# from langchain.chains import create_retrieval_chain

# load_dotenv()

# TOKEN = os.getenv("TOKEN")
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # Configure generativeai with your API key
# genai.configure(api_key=GOOGLE_API_KEY)

# # Initialize bot
# bot = Bot(token=TOKEN)
# dispatcher = Dispatcher(bot)

# rag_on = False

# class Reference:
#     def __init__(self):
#         self.response = ""


# reference = Reference()


# def clear_past():
#     reference.response = ""


# @dispatcher.message_handler(commands=['clear'])
# async def clear(message: types.Message):
#     """
#     A handler to clear the previous conversation and context.
#     """
#     clear_past()
#     await message.reply("I've cleared the past conversation and context.")


# @dispatcher.message_handler(commands=['start'])
# async def welcome(message: types.Message):
#     """
#     This handler receives messages with /start or /help command
#     """
#     await message.reply("Hi! I am a Chat Bot created by Harshal Gidh. How can I assist you?")


# @dispatcher.message_handler(commands=['help'])
# async def helper(message: types.Message):
#     """
#     A handler to display the help menu.
#     """
#     help_command = """
#     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
#     /start - to start the conversation
#     /clear - to clear the past conversation and context.
#     /help - to get this help menu.
#     /begin - to start the Financial Suggestion Conversation with the ChatBot.
#     I hope this helps. :)
#     """
#     await message.reply(help_command)


# @dispatcher.message_handler(commands=['begin'])
# async def finance_bot_activate(message: types.Message):
#     """
#     A handler to display the help menu and activate financial mode.
#     """
    
    
#     help_command = """
#     Hi there! I'm now in Financial Advisor mode. Please follow these instructions:
#     1. Follow this Link (link to upload document) and download the word file.
#     2. Fill in your financial details in the downloaded file.
#     3. Submit the document in the chat.
#     4. Ask your financial questions related to the data you submitted.
#     I will try my best to answer your questions based on the information you provided.
#     I hope this helps. :)
#     """
#     await message.reply(help_command)

    

#     def process_document(file_path):
#         """
#         Processes a document file (either Word document or PDF).

#         Args:
#             file_path (str): Path to the document file.

#         Returns:
#             str: Extracted text content from the document.
#         """
#         try:
#             file_type = filetype.guess(file_path)
#             if file_type is not None:
#                 if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                     print("Processing Word document...")
#                     return extract_text_from_word(file_path)
#                 elif file_type.mime == "application/pdf":
#                     print("Processing PDF document...")
#                     return extract_text_from_pdf(file_path)
#                 else:
#                     print(f"Unsupported file type: {file_type.mime}")
#                     return None
#             else:
#                 print("Could not determine file type.")
#                 return None
#         except Exception as e:
#             print(f"Error processing document: {e}")
#             return None


#     def extract_text_from_word(docx_file_path):
#         """
#         Extracts text content from a Word document (.docx).

#         Args:
#             docx_file_path (str): Path to the Word document file.

#         Returns:
#             str: Extracted text content from the document.
#         """
#         try:
#             doc = docx.Document(docx_file_path)
#             text_content = []
#             for para in doc.paragraphs:
#                 text_content.append(para.text)
#             return "\n".join(text_content)
#         except Exception as e:
#             print(f"Error extracting text from Word document: {e}")
#             return None


#     def extract_text_from_pdf(pdf_file_path):
#         """
#         Extracts text content from a PDF file.

#         Args:
#             pdf_file_path (str): Path to the PDF file.

#         Returns:
#             str: Extracted text content from the PDF.
#         """
#         try:
#             with open(pdf_file_path, "rb") as pdf_file:
#                 pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#                 text_content = []
#                 for page_num in range(pdf_reader.numPages):
#                     page = pdf_reader.getPage(page_num)
#                     text_content.append(page.extract_text())
#                 return "\n".join(text_content)
#         except Exception as e:
#             print(f"Error extracting text from PDF: {e}")
#             return None


#     def load_vector_db(file_path):  #(extracted_text):
#         try:
#                 # Load text from the document
#             loader = Docx2txtLoader(file_path) 
#             # loader = DirectoryLoader(file_path, glob="./*.txt", loader_cls=TextLoader)
#             documents = loader.load()
#             # Split text into chunks
#             text_splitter = RecursiveCharacterTextSplitter(
#                 chunk_size=500,
#                 chunk_overlap=50
#             )
#             text_chunks = text_splitter.split_documents(documents)     #(extracted_text)

#             # Load the embedding model
#             embeddings = GoogleGenerativeAIEmbeddings(
#                 model="models/embedding-001",
#                 google_api_key=GOOGLE_API_KEY
#             )
#             persist_directory = "./chroma_db"
#             # Create a Chroma vector store from text chunks
#             vector_store = Chroma.from_documents(
#                 documents=text_chunks,
#                 embedding=embeddings
#                 #persist_directory= persist_directory
#             )
#             # vectordb = vector_store.persist()
#             # # vectordb = None
#             # vectordb = Chroma(persist_directory=persist_directory,
#             #           embedding_function=embeddings)
#             retriever = vector_store.as_retriever(search_kwargs={"k": 1})
#             print(retriever)
#             return retriever

#         except Exception as e:
#             print(f"Error loading vector database: {e}")
#             return None

#     # Combine data from documents to readable string format.
#     def format_docs(docs):
#         return "\n\n".join(doc.page_content for doc in docs)


#     async def gemini_rag_chain(retriever):
#         llm = ChatGoogleGenerativeAI(
#             model="gemini-pro",
#             temperature=0.7,
#             top_p=0.85,
#             google_api_key=GOOGLE_API_KEY
#         )
#         memory = ConversationSummaryMemory(llm=llm, memory_key="chat_history", return_messages=True)

#         # create the chain to answer questions
#         # qa_chain = RetrievalQA.from_chain_type( llm=llm,
#         #                             chain_type="stuff",
#         #                             retriever=retriever,
#         #                             return_source_documents=True)
        
#         prompt_template = """You are a Financial assistant for question-answering tasks.
#                 You are supposed to give Financial advice to the user that can help the user to gain good knowledge about investing.
#                 If you don't know the answer, just say that you don't know.
#                 Use five sentences maximum and keep the answer concise.\n
#                 Question: {question}\nContext: {context}\nAnswer:"""
        
#         # qa_chain = RetrievalQA.from_chain_type(
#         #                 PromptTemplate(prompt_template),  
#         #                 llm=llm,
#         #                 chain_type="stuff",
#         #                 retriever=retriever,
#         #                 return_source_documents=True
#         #             )


#         # print(qa_chain)

#         # return qa_chain
        
        
#         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

#         document_chain=create_stuff_documents_chain(llm,llm_prompt)

        
#         retrieval_chain=create_retrieval_chain(retriever,document_chain)
#         print(retrieval_chain)
#         return retrieval_chain

#         rag_chain = (
#         {"context": retriever | format_docs, "question": RunnablePassthrough()}
#         | llm_prompt
#         | llm
#         | StrOutputParser()
#         )

#         input_data = {
#             "input_variables": {
#                 "context": await retriever.load() if retriever else None,
#                 "question": question
#             }
#         }

         
#         return rag_chain 
        

#     async def get_gemini_rag_response(chain,question):
#         try:
#             response = chain.invoke(question) #rag_chain.invoke(input_data)
#             print(response.text)
#             return response
#         except Exception as e:
#             print(f"Error in get_gemini_rag_response: {e}")
#             return None
        
#     @dispatcher.message_handler()
#     async def get_query(message: types.Message):
#         return  message.text  # Return the text directly

#     @dispatcher.message_handler(content_types=['document'])
#     async def get_file(message: types.Message):
#         """
#         A handler to process the user's input and generate a response using the Gemini model.
#         """
#         file_id = message.document.file_id
#         file = await bot.get_file(file_id)
#         file_path = file.file_path
#         await bot.download_file(file_path, "data/uploaded_file")

#         print(f"Received file: {file_path}")

#         # Process the uploaded document to extract text content
#         extracted_text = process_document("data/uploaded_file")

#         if extracted_text:
#             # Load the vector database
#             retriever = load_vector_db("data/uploaded_file")

#             if retriever:
#                 # Generate response using RAG model
#                 await message.reply("Processed file, Please ask your question")
#                 question = await get_query(message)  # Await the result of get_query
#                 chain = await gemini_rag_chain(retriever)
#                 print(f">>> Client: \n\t{question}")
#                 response = await get_gemini_rag_response(chain, question)

#             else:
#                 response = "Failed to load vector database."

#             # Send the response to the user
#             print(f">>> rag_gemini: \n\t{response}")
#             await bot.send_message(chat_id=message.chat.id, text=response)
#         else:
#             await bot.send_message(chat_id=message.chat.id, text="Failed to process the uploaded file.")

#     # @dispatcher.message_handler()
#     # async def get_query(message: types.Message):
#     #     query = message.text 
#     #     return await query 

#     # @dispatcher.message_handler(content_types=['document'])
#     # async def get_file(message: types.Message):
#     #     """
#     #     A handler to process the user's input and generate a response using the Gemini model.
#     #     """
#     #     file_id = message.document.file_id
#     #     file = await bot.get_file(file_id)
#     #     file_path = file.file_path
#     #     await bot.download_file(file_path, "data/uploaded_file")

#     #     print(f"Received file: {file_path}")

#     #     # Process the uploaded document to extract text content
#     #     extracted_text = process_document("data/uploaded_file")

#     #     if extracted_text:
#     #         # Load the vector database
#     #         retriever = load_vector_db("data/uploaded_file")  #(extracted_text)

#     #         if retriever:
#     #             # Generate response using RAG model
#     #             await message.reply("Processed file, Please ask your question")
#     #             question =  get_query(message)
#     #             chain = await gemini_rag_chain(retriever)
#     #             response = await get_gemini_rag_response(chain, question)
                
#     #         else:
#     #             response = "Failed to load vector database."

#     #         # Send the response to the user
#     #         await bot.send_message(chat_id=message.chat.id, text=response)
#     #     else:
#     #         await bot.send_message(chat_id=message.chat.id, text="Failed to process the uploaded file.")



# async def get_gemini_response(question):
#     model = genai.GenerativeModel('gemini-pro')
#     response = model.generate_content(question)
#     return response.text


# if not rag_on :
#     @dispatcher.message_handler()
#     async def main_bot(message: types.Message):
#         """
#         A handler to process the user's input and generate a response using the Gemini model.
#         """
 #       # print(f">>> USER: \n\t{message.text}")
 #       # response = await get_gemini_response(message.text)
  #      # print(f">>> gemini: \n\t{response}")
   #     # await bot.send_message(chat_id=message.chat.id, text=response)


# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)







# Nice imporved code with more chat functionality, fails in rag response generation

# import os
# import filetype
# import docx
# import PyPDF2
# from aiogram import Bot, Dispatcher, executor, types
# from dotenv import load_dotenv
# import google.generativeai as genai
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
# from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# from langchain.memory import ConversationSummaryMemory
# from langchain.chains import ConversationalRetrievalChain
# from langchain_core.prompts import PromptTemplate
# from langchain.schema import StrOutputParser
# from langchain_community.document_loaders import Docx2txtLoader

# # from langchain.document_loaders import DirectoryLoader
# # from langchain.document_loaders import TextLoader
# from langchain.chains import RetrievalQA
# from langchain_community.document_loaders import DirectoryLoader
# from langchain_community.document_loaders import TextLoader
# load_dotenv()

# TOKEN = os.getenv("TOKEN")
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # Configure generativeai with your API key
# genai.configure(api_key=GOOGLE_API_KEY)

# # Initialize bot
# bot = Bot(token=TOKEN)
# dispatcher = Dispatcher(bot)

# rag_on = False

# class Reference:
#     def __init__(self):
#         self.response = ""


# reference = Reference()


# def clear_past():
#     reference.response = ""


# @dispatcher.message_handler(commands=['clear'])
# async def clear(message: types.Message):
#     """
#     A handler to clear the previous conversation and context.
#     """
#     clear_past()
#     await message.reply("I've cleared the past conversation and context.")


# @dispatcher.message_handler(commands=['start'])
# async def welcome(message: types.Message):
#     """
#     This handler receives messages with /start or /help command
#     """
#     await message.reply("Hi! I am a Chat Bot created by Harshal Gidh. How can I assist you?")


# @dispatcher.message_handler(commands=['help'])
# async def helper(message: types.Message):
#     """
#     A handler to display the help menu.
#     """
#     help_command = """
#     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
#     /start - to start the conversation
#     /clear - to clear the past conversation and context.
#     /help - to get this help menu.
#     /begin - to start the Financial Suggestion Conversation with the ChatBot.
#     I hope this helps. :)
#     """
#     await message.reply(help_command)


# @dispatcher.message_handler(commands=['begin'])
# async def finance_bot_activate(message: types.Message):
#     """
#     A handler to display the help menu.
#     """
#     help_command = """
#     Hi there! I'm a bot created by Harshal Gidh. Please follow these instructions:
#     Follow this Link : link 
#     and Download the word file ,then fill the details and then submit the document in the chat.
#     /begin - to start the Financial Suggestion Conversation with the ChatBot.
#     I hope this helps. :)
#     """
#     rag_on = True
#     await message.reply(help_command)


# def process_document(file_path):
#     """
#     Processes a document file (either Word document or PDF).

#     Args:
#         file_path (str): Path to the document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         file_type = filetype.guess(file_path)
#         if file_type is not None:
#             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 print("Processing Word document...")
#                 return extract_text_from_word(file_path)
#             elif file_type.mime == "application/pdf":
#                 print("Processing PDF document...")
#                 return extract_text_from_pdf(file_path)
#             else:
#                 print(f"Unsupported file type: {file_type.mime}")
#                 return None
#         else:
#             print("Could not determine file type.")
#             return None
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None


# def extract_text_from_word(docx_file_path):
#     """
#     Extracts text content from a Word document (.docx).

#     Args:
#         docx_file_path (str): Path to the Word document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         doc = docx.Document(docx_file_path)
#         text_content = []
#         for para in doc.paragraphs:
#             text_content.append(para.text)
#         return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from Word document: {e}")
#         return None


# def extract_text_from_pdf(pdf_file_path):
#     """
#     Extracts text content from a PDF file.

#     Args:
#         pdf_file_path (str): Path to the PDF file.

#     Returns:
#         str: Extracted text content from the PDF.
#     """
#     try:
#         with open(pdf_file_path, "rb") as pdf_file:
#             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#             text_content = []
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text_content.append(page.extract_text())
#             return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None


# def load_vector_db(file_path):  #(extracted_text):
#     try:
#             # Load text from the document
#         loader = Docx2txtLoader(file_path) 
#         # loader = DirectoryLoader(file_path, glob="./*.txt", loader_cls=TextLoader)
#         documents = loader.load()
#         # Split text into chunks
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=500,
#             chunk_overlap=50
#         )
#         text_chunks = text_splitter.split_documents(documents)     #(extracted_text)

#         # Load the embedding model
#         embeddings = GoogleGenerativeAIEmbeddings(
#             model="models/embedding-001",
#             google_api_key=GOOGLE_API_KEY
#         )
#         persist_directory = "./chroma_db"
#         # Create a Chroma vector store from text chunks
#         vector_store = Chroma.from_documents(
#             documents=text_chunks,
#             embedding=embeddings,
#             persist_directory= persist_directory
#         )
#         # vectordb = vector_store.persist()
#         # # vectordb = None
#         # vectordb = Chroma(persist_directory=persist_directory,
#         #           embedding_function=embeddings)
#         retriever = vector_store.as_retriever(search_kwargs={"k": 1})
#         return retriever

#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None


# async def gemini_rag_chain(retriever):
#     llm = ChatGoogleGenerativeAI(
#         model="gemini-pro",
#         temperature=0.7,
#         top_p=0.85,
#         google_api_key=GOOGLE_API_KEY
#     )
#     memory = ConversationSummaryMemory(llm=llm, memory_key="chat_history", return_messages=True)

#     # create the chain to answer questions
#     qa_chain = RetrievalQA.from_chain_type(
#         # PromptTemplate(
#         #     prompt="""You are a Financial assistant for question-answering tasks.
#         #     You are supposed to give Financial advice to the user that can help the user to gain good knowledge about investing.
#         #     If you don't know the answer, just say that you don't know.
#         #     Use five sentences maximum and keep the answer concise.\n
#         #     Question: {question}\nContext: {context}\nAnswer:"""
#         # ) ,
#                             llm=llm,
#                                   chain_type="stuff",
#                                   retriever=retriever,
#                                   return_source_documents=True)
#     # rag_chain = (
#         # PromptTemplate(
#         #     prompt="""You are a Financial assistant for question-answering tasks.
#         #     You are supposed to give Financial advice to the user that can help the user to gain good knowledge about investing.
#         #     If you don't know the answer, just say that you don't know.
#         #     Use five sentences maximum and keep the answer concise.\n
#         #     Question: {question}\nContext: {context}\nAnswer:"""
#         # )
#     #     | llm
#     #     | StrOutputParser()
#     # )

#     # input_data = {
#     #     "input_variables": {
#     #         "context": await retriever.load() if retriever else None,
#     #         "question": question
#     #     }
#     # }
    

# async def get_gemini_rag_response(chain,question):
#     try:
#         response = chain.invoke(question) #rag_chain.invoke(input_data)
#         return response
#     except Exception as e:
#         print(f"Error in get_gemini_rag_response: {e}")
#         return None


# async def get_gemini_response(question):
#     model = genai.GenerativeModel('gemini-pro')
#     response = model.generate_content(question)
#     return response.text


# @dispatcher.message_handler(content_types=['document'])
# async def get_file(message: types.Message):
#     """
#     A handler to process the user's input and generate a response using the Gemini model.
#     """
#     file_id = message.document.file_id
#     file = await bot.get_file(file_id)
#     file_path = file.file_path
#     await bot.download_file(file_path, "data/uploaded_file")

#     print(f"Received file: {file_path}")

#     # Process the uploaded document to extract text content
#     extracted_text = process_document("data/uploaded_file")

#     if extracted_text:
#         # Load the vector database
#         retriever = load_vector_db("data/uploaded_file")  #(extracted_text)

#         if retriever:
#             # Generate response using RAG model
#             await message.reply("Processed file, Please ask your question")
#             chain = await gemini_rag_chain(retriever)
#             response = await get_gemini_rag_response(chain, message.text)
            
#         else:
#             response = "Failed to load vector database."

#         # Send the response to the user
#         await bot.send_message(chat_id=message.chat.id, text=response)
#     else:
#         await bot.send_message(chat_id=message.chat.id, text="Failed to process the uploaded file.")

# if not rag_on :
#     @dispatcher.message_handler()
#     async def main_bot(message: types.Message):
#         """
#         A handler to process the user's input and generate a response using the Gemini model.
#         """
#         print(f">>> USER: \n\t{message.text}")
#         response = await get_gemini_response(message.text)
#         print(f">>> gemini: \n\t{response}")
#         await bot.send_message(chat_id=message.chat.id, text=response)


# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)








# Just Problem with the chain rest code is working fine and is able to handle the file document

# import os
# import filetype
# import docx
# import PyPDF2
# from aiogram import Bot, Dispatcher, executor, types
# from dotenv import load_dotenv
# import google.generativeai as genai
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
# from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# from langchain.memory import ConversationSummaryMemory
# from langchain.chains import ConversationalRetrievalChain
# from langchain_core.prompts import PromptTemplate
# from langchain.schema import StrOutputParser
# from langchain_community.document_loaders import Docx2txtLoader

# # from langchain.document_loaders import DirectoryLoader
# # from langchain.document_loaders import TextLoader
# from langchain.chains import RetrievalQA
# from langchain_community.document_loaders import DirectoryLoader
# from langchain_community.document_loaders import TextLoader
# load_dotenv()

# TOKEN = os.getenv("TOKEN")
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # Configure generativeai with your API key
# genai.configure(api_key=GOOGLE_API_KEY)

# # Initialize bot
# bot = Bot(token=TOKEN)
# dispatcher = Dispatcher(bot)


# class Reference:
#     def __init__(self):
#         self.response = ""


# reference = Reference()


# def clear_past():
#     reference.response = ""


# @dispatcher.message_handler(commands=['clear'])
# async def clear(message: types.Message):
#     """
#     A handler to clear the previous conversation and context.
#     """
#     clear_past()
#     await message.reply("I've cleared the past conversation and context.")


# @dispatcher.message_handler(commands=['start'])
# async def welcome(message: types.Message):
#     """
#     This handler receives messages with /start or /help command
#     """
#     await message.reply("Hi! I am a Chat Bot created by Harshal Gidh. How can I assist you?")


# @dispatcher.message_handler(commands=['help'])
# async def helper(message: types.Message):
#     """
#     A handler to display the help menu.
#     """
#     help_command = """
#     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
#     /start - to start the conversation
#     /clear - to clear the past conversation and context.
#     /help - to get this help menu.
#     I hope this helps. :)
#     """
#     await message.reply(help_command)


# def process_document(file_path):
#     """
#     Processes a document file (either Word document or PDF).

#     Args:
#         file_path (str): Path to the document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         file_type = filetype.guess(file_path)
#         if file_type is not None:
#             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 print("Processing Word document...")
#                 return extract_text_from_word(file_path)
#             elif file_type.mime == "application/pdf":
#                 print("Processing PDF document...")
#                 return extract_text_from_pdf(file_path)
#             else:
#                 print(f"Unsupported file type: {file_type.mime}")
#                 return None
#         else:
#             print("Could not determine file type.")
#             return None
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None


# def extract_text_from_word(docx_file_path):
#     """
#     Extracts text content from a Word document (.docx).

#     Args:
#         docx_file_path (str): Path to the Word document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         doc = docx.Document(docx_file_path)
#         text_content = []
#         for para in doc.paragraphs:
#             text_content.append(para.text)
#         return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from Word document: {e}")
#         return None


# def extract_text_from_pdf(pdf_file_path):
#     """
#     Extracts text content from a PDF file.

#     Args:
#         pdf_file_path (str): Path to the PDF file.

#     Returns:
#         str: Extracted text content from the PDF.
#     """
#     try:
#         with open(pdf_file_path, "rb") as pdf_file:
#             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#             text_content = []
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text_content.append(page.extract_text())
#             return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None


# def load_vector_db(file_path):  #(extracted_text):
#     try:
#             # Load text from the document
#         loader = Docx2txtLoader(file_path) 
#         # loader = DirectoryLoader(file_path, glob="./*.txt", loader_cls=TextLoader)
#         documents = loader.load()
#         # Split text into chunks
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=500,
#             chunk_overlap=50
#         )
#         text_chunks = text_splitter.split_documents(documents)     #(extracted_text)

#         # Load the embedding model
#         embeddings = GoogleGenerativeAIEmbeddings(
#             model="models/embedding-001",
#             google_api_key=GOOGLE_API_KEY
#         )
#         persist_directory = "./chroma_db"
#         # Create a Chroma vector store from text chunks
#         vector_store = Chroma.from_documents(
#             documents=text_chunks,
#             embedding=embeddings,
#             persist_directory= persist_directory
#         )
#         # vectordb = vector_store.persist()
#         # # vectordb = None
#         # vectordb = Chroma(persist_directory=persist_directory,
#         #           embedding_function=embeddings)
#         retriever = vector_store.as_retriever(search_kwargs={"k": 1})
#         return retriever

#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None


# async def gemini_rag_response(retriever, question):
#     llm = ChatGoogleGenerativeAI(
#         model="gemini-pro",
#         temperature=0.7,
#         top_p=0.85,
#         google_api_key=GOOGLE_API_KEY
#     )
#     memory = ConversationSummaryMemory(llm=llm, memory_key="chat_history", return_messages=True)

#     # create the chain to answer questions
#     qa_chain = RetrievalQA.from_chain_type(
#         # PromptTemplate(
#         #     prompt="""You are a Financial assistant for question-answering tasks.
#         #     You are supposed to give Financial advice to the user that can help the user to gain good knowledge about investing.
#         #     If you don't know the answer, just say that you don't know.
#         #     Use five sentences maximum and keep the answer concise.\n
#         #     Question: {question}\nContext: {context}\nAnswer:"""
#         # ) ,
#                             llm=llm,
#                                   chain_type="stuff",
#                                   retriever=retriever,
#                                   return_source_documents=True)
#     # rag_chain = (
#         # PromptTemplate(
#         #     prompt="""You are a Financial assistant for question-answering tasks.
#         #     You are supposed to give Financial advice to the user that can help the user to gain good knowledge about investing.
#         #     If you don't know the answer, just say that you don't know.
#         #     Use five sentences maximum and keep the answer concise.\n
#         #     Question: {question}\nContext: {context}\nAnswer:"""
#         # )
#     #     | llm
#     #     | StrOutputParser()
#     # )

#     # input_data = {
#     #     "input_variables": {
#     #         "context": await retriever.load() if retriever else None,
#     #         "question": question
#     #     }
#     # }
    
#     try:
#         response = await qa_chain.invoke(question) #rag_chain.invoke(input_data)
#         return response
#     except Exception as e:
#         print(f"Error in gemini_rag_response: {e}")
#         return None


# async def get_gemini_response(question):
#     model = genai.GenerativeModel('gemini-pro')
#     response = model.generate_content(question)
#     return response.text


# @dispatcher.message_handler(content_types=['document'])
# async def get_file(message: types.Message):
#     """
#     A handler to process the user's input and generate a response using the Gemini model.
#     """
#     file_id = message.document.file_id
#     file = await bot.get_file(file_id)
#     file_path = file.file_path
#     await bot.download_file(file_path, "data/uploaded_file")

#     print(f"Received file: {file_path}")

#     # Process the uploaded document to extract text content
#     extracted_text = process_document("data/uploaded_file")

#     if extracted_text:
#         # Load the vector database
#         retriever = load_vector_db("data/uploaded_file")  #(extracted_text)

#         if retriever:
#             # Generate response using RAG model
#             response = await gemini_rag_response(retriever, extracted_text)
#         else:
#             response = "Failed to load vector database."

#         # Send the response to the user
#         await bot.send_message(chat_id=message.chat.id, text=response)
#     else:
#         await bot.send_message(chat_id=message.chat.id, text="Failed to process the uploaded file.")


# @dispatcher.message_handler()
# async def main_bot(message: types.Message):
#     """
#     A handler to process the user's input and generate a response using the Gemini model.
#     """
#     print(f">>> USER: \n\t{message.text}")
#     response = await get_gemini_response(message.text)
#     print(f">>> gemini: \n\t{response}")
#     await bot.send_message(chat_id=message.chat.id, text=response)


# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)







# Best Code but not working with rag system :

# import os
# import filetype
# import docx
# import PyPDF2
# from aiogram import Bot, Dispatcher, executor, types
# from dotenv import load_dotenv
# import google.generativeai as genai

# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.memory import ConversationSummaryMemory
# from langchain.chains import ConversationalRetrievalChain
# from langchain import PromptTemplate
# from langchain.schema import StrOutputParser
# from langchain.schema.runnable import RunnablePassthrough
# from langchain_google_genai import GoogleGenerativeAIEmbeddings

# load_dotenv()

# TOKEN = os.getenv("TOKEN")
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # Configure generativeai with your API key
# genai.configure(api_key=GOOGLE_API_KEY)

# # Initialize bot
# bot = Bot(token=TOKEN)
# dispatcher = Dispatcher(bot)


# class Reference:
#     def _init_(self):
#         self.response = ""


# reference = Reference()


# def clear_past():
#     reference.response = ""


# @dispatcher.message_handler(commands=['clear'])
# async def clear(message: types.Message):
#     """
#     A handler to clear the previous conversation and context.
#     """
#     clear_past()
#     await message.reply("I've cleared the past conversation and context.")


# @dispatcher.message_handler(commands=['start'])
# async def welcome(message: types.Message):
#     """This handler receives messages with /start or  `/help `command

#     Args:
#         message (types.Message): description
#     """
#     await message.reply("Hi\nI am a Chat Bot! Created by Harshal Gidh!. How can i assist you?")


# @dispatcher.message_handler(commands=['help'])
# async def helper(message: types.Message):
#     """
#     A handler to display the help menu.
#     """
#     help_command = """
#     Hi There, I'm a bot created by Harshal Gidh! Please follow these commands - 
#     /start - to start the conversation
#     /clear - to clear the past conversation and context.
#     /help - to get this help menu.
#     I hope this helps. :)
#     """
#     await message.reply(help_command)


# def process_document(file_path):
#     """
#     Processes a document file (either Word document or PDF).

#     Args:
#         file_path (str): Path to the document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         file_type = filetype.guess(file_path)
#         if file_type is not None:
#             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 print("Processing Word document...")
#                 return extract_text_from_word(file_path)
#             elif file_type.mime == "application/pdf":
#                 print("Processing PDF document...")
#                 return extract_text_from_pdf(file_path)
#             else:
#                 print(f"Unsupported file type: {file_type.mime}")
#                 return None
#         else:
#             print("Could not determine file type.")
#             return None
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None


# def extract_text_from_word(docx_file_path):
#     """
#     Extracts text content from a Word document (.docx).

#     Args:
#         docx_file_path (str): Path to the Word document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         doc = docx.Document(docx_file_path)
#         text_content = []
#         for para in doc.paragraphs:
#             text_content.append(para.text)
#         return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from Word document: {e}")
#         return None


# def extract_text_from_pdf(pdf_file_path):
#     """
#     Extracts text content from a PDF file.

#     Args:
#         pdf_file_path (str): Path to the PDF file.

#     Returns:
#         str: Extracted text content from the PDF.
#     """
#     try:
#         with open(pdf_file_path, "rb") as pdf_file:
#             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#             text_content = []
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text_content.append(page.extract_text())
#             return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None


# def load_vectordb(extracted_text):
#     try:
#         # Load text from the document
#         # loader = Docx2txtLoader(file_path)
#         # documents = loader.load()

#         # Split text into chunks
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=500,
#             chunk_overlap=50
#         )
#         text_chunks = text_splitter.split_documents(extracted_text)

#         # Load the embedding model
#         embeddings = GoogleGenerativeAIEmbeddings(
#             model="models/embedding-001",
#             google_api_key=GOOGLE_API_KEY
#         )

#         # Create a Chroma vector store from text chunks
#         vector_store = Chroma.from_documents(
#             documents=text_chunks,
#             embedding=embeddings,
#             persist_directory="./chroma_db"
#         )
#         # Load from disk
#         vectorstore_disk = Chroma(
#                         persist_directory="./chroma_db",       # Directory of db
#                         embedding_function=embeddings   # Embedding model
#                    )
#         retriever = vectorstore_disk.as_retriever(search_kwargs={"k": 1})
#         return retriever

#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None


# # def format_docs(docs):
# #     return "\n\n".join(doc.page_content for doc in docs)

# async def gemini_rag_response(retriever):
#     llm = ChatGoogleGenerativeAI(
#         model="gemini-pro",
#         temperature=0.7,
#         top_p=0.85,
#         google_api_key=GOOGLE_API_KEY
#     )
#     memory = ConversationSummaryMemory(llm=llm, memory_key="chat_history", return_messages=True)

#     rag_chain = (
#         PromptTemplate(
#             prompt="""You are a Financial assistant for question-answering tasks.
#             You are supposed to give Financial advice to the user that can help the user to gain good knowledge about investing.
#             If you don't know the answer, just say that you don't know.
#             Use five sentences maximum and keep the answer concise.\n
#             Question: {question}\nContext: {context}\nAnswer:"""
#         )
#         | llm
#         | StrOutputParser()
#     )

#     context = retriever #format_docs(await retriever.load()) if retriever else None
#     input_data = {
#         "input_variables": {
#             "context": context,
#             "question": RunnablePassthrough() #None  # Placeholder for the question
#         }
#     }
    
#     try:
#         response = await rag_chain.invoke(input_data)
#         return response
#     except Exception as e:
#         print(f"Error in gemini_rag_response: {e}")
#         return None


# async def get_gemini_response(question):
#     model = genai.GenerativeModel('gemini-pro')
#     response = model.generate_content(question)
#     return response.text


# @dispatcher.message_handler(content_types=['document'])
# async def get_file(message: types.Message):
#     """
#     A handler to process the user's input and generate a response using the Gemini model.
#     """
#     file_id = message.document.file_id
#     file = await bot.get_file(file_id)
#     file_path = file.file_path
#     await bot.download_file(file_path, "data/uploaded_file")

#     print(f"Received file: {file_path}")

#     # Process the uploaded document to extract text content
#     extracted_text = process_document("data/uploaded_file")

#     if extracted_text:
#         # Use the extracted text as input to the Gemini model
#         retriver = load_vectordb(extracted_text)
#         response = await get_gemini_response(extracted_text)

#         # Send the response to the user
#         await bot.send_message(chat_id=message.chat.id, text=response)
#     else:
#         await bot.send_message(chat_id=message.chat.id, text="Failed to process the uploaded file.")


# @dispatcher.message_handler()
# async def main_bot(message: types.Message):
#     """
#     A handler to process the user's input and generate a response using the Gemini model.
#     """
#     print(f">>> USER: \n\t{message.text}")
#     response = await get_gemini_response(message.text)
#     print(f">>> gemini: \n\t{response}")
#     await bot.send_message(chat_id=message.chat.id, text=response)


# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)



# working code :

# import os
# import filetype
# import docx
# import PyPDF2
# from aiogram import Bot, Dispatcher, executor, types
# from dotenv import load_dotenv
# import google.generativeai as genai

# load_dotenv()

# TOKEN = os.getenv("TOKEN")
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # Configure generativeai with your API key
# genai.configure(api_key=GOOGLE_API_KEY)

# # Initialize bot
# bot = Bot(token=TOKEN)
# dispatcher = Dispatcher(bot)


# class Reference:
#     def _init_(self):
#         self.response = ""


# reference = Reference()


# def clear_past():
#     reference.response = ""


# @dispatcher.message_handler(commands=['clear'])
# async def clear(message: types.Message):
#     """
#     A handler to clear the previous conversation and context.
#     """
#     clear_past()
#     await message.reply("I've cleared the past conversation and context.")


# @dispatcher.message_handler(commands=['start'])
# async def welcome(message: types.Message):
#     """This handler receives messages with /start or  `/help `command

#     Args:
#         message (types.Message): description
#     """
#     await message.reply("Hi\nI am a Chat Bot! Created by Harshal Gidh!. How can i assist you?")


# @dispatcher.message_handler(commands=['help'])
# async def helper(message: types.Message):
#     """
#     A handler to display the help menu.
#     """
#     help_command = """
#     Hi There, I'm a bot created by Harshal Gidh! Please follow these commands - 
#     /start - to start the conversation
#     /clear - to clear the past conversation and context.
#     /help - to get this help menu.
#     I hope this helps. :)
#     """
#     await message.reply(help_command)


# def process_document(file_path):
#     """
#     Processes a document file (either Word document or PDF).

#     Args:
#         file_path (str): Path to the document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         file_type = filetype.guess(file_path)
#         if file_type is not None:
#             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 print("Processing Word document...")
#                 return extract_text_from_word(file_path)
#             elif file_type.mime == "application/pdf":
#                 print("Processing PDF document...")
#                 return extract_text_from_pdf(file_path)
#             else:
#                 print(f"Unsupported file type: {file_type.mime}")
#                 return None
#         else:
#             print("Could not determine file type.")
#             return None
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None


# def extract_text_from_word(docx_file_path):
#     """
#     Extracts text content from a Word document (.docx).

#     Args:
#         docx_file_path (str): Path to the Word document file.

#     Returns:
#         str: Extracted text content from the document.
#     """
#     try:
#         doc = docx.Document(docx_file_path)
#         text_content = []
#         for para in doc.paragraphs:
#             text_content.append(para.text)
#         return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from Word document: {e}")
#         return None


# def extract_text_from_pdf(pdf_file_path):
#     """
#     Extracts text content from a PDF file.

#     Args:
#         pdf_file_path (str): Path to the PDF file.

#     Returns:
#         str: Extracted text content from the PDF.
#     """
#     try:
#         with open(pdf_file_path, "rb") as pdf_file:
#             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#             text_content = []
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text_content.append(page.extract_text())
#             return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None


# async def get_gemini_response(question):
#     model = genai.GenerativeModel('gemini-pro')
#     response = model.generate_content(question)
#     return response.text


# @dispatcher.message_handler(content_types=['document'])
# async def get_file(message: types.Message):
#     """
#     A handler to process the user's input and generate a response using the Gemini model.
#     """
#     file_id = message.document.file_id
#     file = await bot.get_file(file_id)
#     file_path = file.file_path
#     await bot.download_file(file_path, "data/uploaded_file")

#     print(f"Received file: {file_path}")

#     # Process the uploaded document to extract text content
#     extracted_text = process_document("data/uploaded_file")

#     if extracted_text:
#         # Use the extracted text as input to the Gemini model
#         response = await get_gemini_response(extracted_text)

#         # Send the response to the user
#         await bot.send_message(chat_id=message.chat.id, text=response)
#     else:
#         await bot.send_message(chat_id=message.chat.id, text="Failed to process the uploaded file.")


# @dispatcher.message_handler()
# async def main_bot(message: types.Message):
#     """
#     A handler to process the user's input and generate a response using the Gemini model.
#     """
#     print(f">>> USER: \n\t{message.text}")
#     response = await get_gemini_response(message.text)
#     print(f">>> gemini: \n\t{response}")
#     await bot.send_message(chat_id=message.chat.id, text=response)


# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)