#best code so far now it can upload files and receive various forms of messages as well and provide us graphs that we want 
# and also reply to images, give stock informations ,give stock analysis information and a prediction of the price

import os
import filetype
import docx
import PyPDF2
import re
from aiogram import Bot, Dispatcher, types
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.memory import ConversationSummaryMemory
import asyncio
import numpy as np
import json
import re
import google.generativeai as genai
import pathlib
# Import things that are needed generically
from langchain.pydantic_v1 import BaseModel, Field
from langchain.tools import BaseTool, StructuredTool, tool

from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import CommandStart
from aiogram.types import Message
from aiogram import F
from aiogram import Router
import logging
import sys
from aiogram.filters import Command
from aiogram.types import FSInputFile
# from aiogram.utils import executor
import io
import matplotlib.pyplot as plt
import seaborn as sns
import aiohttp
from aiogram.types import InputFile , BufferedInputFile
import PIL.Image

router = Router(name=__name__)

load_dotenv()

TOKEN = os.getenv("TOKEN")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

os.environ["LANGCHAIN_TRACING_V2"]="true"
os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# Configure generativeai with your API key
genai.configure(api_key=GOOGLE_API_KEY)

# Initialize bot
bot = Bot(token=TOKEN)
dp = Dispatcher()

# Glbal variables
rag_on = False
retriever = None  # Store retriever globally
summary = ""
investment_personality = ""
# history = []
previous_suggestions = ""

CHAT_HISTORY_FILE = 'chat_history.json'

def read_chat_history(chat_id):
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, 'r') as file:
            chat_history = json.load(file)
            return chat_history.get(str(chat_id), [])
    return []

def write_chat_history(chat_id, message):
    chat_history = {}
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, 'r') as file:
            chat_history = json.load(file)
    if str(chat_id) not in chat_history:
        chat_history[str(chat_id)] = []
    chat_history[str(chat_id)].append(message)
    with open(CHAT_HISTORY_FILE, 'w') as file:
        json.dump(chat_history, file)

class Reference:
    def __init__(self):
        self.response = ""


reference = Reference()


def clear_past():
    reference.response = ""


@router.message(F.text == "clear")
async def clear(message: types.Message):
    """
    A handler to clear the previous conversation and context.
    """
    clear_past()
    await message.reply("I've cleared the past conversation and context.")

#Global Variables :

# Store user states
states = {}

# Dictionary to hold question-answer pairs
user_responses = {}
#
user_images = {}
# Define Questions for assessment
questions = [
    """ 
    1. You and your friend are betting on a series of coin tosses.

    He always bets ₹2,000 on Heads

    You always bet ₹2,000 on Tails

    Winner of last 8 turns

    You lost ₹8,000 in the last 4 turns!

    If you were to bet one last time, what would you bet on:
    a) heads or b) tails ?
    """ ,
    """
    2. Imagine you are a contestant in a game show, and you are presented the following choices.

    What would you prefer?
    a) 50 percent chance of winning 15 gold coins 
    b) 100 percent chance of winning 8 gold coins
    """,
    """
    3. Ok, one more choice...

    What would you prefer?
    a) 50 percent chance of winning 15 gold coins 
    b) 100 percent chance of winning 2 gold coins
    """,
    """
    4. In general, how would your best friend describe your risk-taking tendencies?
    a) A real gambler
    b) Willing to take risks after completing adequate research
    c) Cautious
    d) Avoids risk as much as possible
    """,
    """
    5. Suppose you could replace your current investment portfolio with this new one:
    50 percent chance of Gaining 35 percent or 50 percent chance of Loss
    In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
    a)-5 to -10
    b)-10 to -15
    c)-15 to -20
    d)-20 to -25
    e)-25 to -30
    f)-30 to -35
    """,
    """
    6. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
    a) Option A : Min Value is 9.2L and Max Value is 11.8L
    b) Option B : Min Value is 8.8L and Max Value is 12.3L
    c) Option C : Min Value is 8.5L and Max Value is 12.8L
    d) Option D : Min Value is 8.1L and Max Value is 13.3L
    e) Option E : Min Value is 7.8L and Max Value is 13.8L
    """,
    """
    7. From Sept 2008 to Nov 2008, Stock market went down by 31%.

    If you owned a stock investment that lost about 31 percent in 3 months, you would:
    a) Sell all of the remaining investment
    b) Sell a portion of the remaining investment
    c) Hold on to the investment and sell nothing
    d) Buy little
    e) Buy more of the investment
    """,
    """
    8. Over any 1-year period, what would be the maximum drop in the value of your investment 
    portfolio that you would be comfortable with?
    a) <5%
    b) 5 - 10%
    c) 10 - 15%
    d) 15 - 20%
    e) >20%
    """,
    """
    9. When investing, what do you consider the most?

    a) Risk 
    b) Return
    """,
    """
    10. What best describes your attitude?

    a) Prefer reasonable returns, can take reasonable risk
    b) Like higher returns, can take slightly higher risk
    c) Want to maximize returns, can take significant high risk
    """,
    """
    11. How much monthly investment you want to do?
    """,
    """
    12. What is the time horizon for your investment?
    You can answer in any range, example 1-5 years."""  
]


# Handler for /start command
@dp.message(CommandStart())
async def handle_start(message: types.Message):
    """
    This handler receives messages with /start command
    """
    chat_id = message.chat.id
    # Start asking questions
    await start_assessment(chat_id)


# Function to start the assessment
async def start_assessment(chat_id):
    await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
    await ask_next_question(chat_id, 0)

# Function to ask the next question
async def ask_next_question(chat_id, question_index):
    if question_index < len(questions):
        # Ask the next question
        await bot.send_message(chat_id, questions[question_index])
        # Update state to indicate the next expected answer
        states[chat_id] = question_index
    else:
        # No more questions, finish assessment
        await finish_assessment(chat_id)

# Handler for receiving assessment answers
assessment_in_progress = True

from aiogram.types import FSInputFile
async def finish_assessment(chat_id):
    if chat_id in states and states[chat_id] == len(questions):
        # All questions have been answered, now process the assessment
        await bot.send_message(chat_id, "Assessment completed. Thank you!")

        # Determine investment personality based on collected responses
        global investment_personality
        investment_personality = await determine_investment_personality(user_responses)

        # Inform the user about their investment personality
        await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

        # Store the response in chat history
        write_chat_history(chat_id, {'role': 'bot', 'message': investment_personality})

        # Summarize collected information
        global summary
        summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
        summary = summary + "\n" + "Your investment personality:" + investment_personality
        # Ensure to await the determination of investment personality
        await send_summary_chunks(chat_id, summary)
        global assessment_in_progress 
        assessment_in_progress = False
       
        await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
        file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

        await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
        # await bot.send_message(chat_id,file)

async def send_summary_chunks(chat_id, summary):
    # Split the summary into chunks that fit within Telegram's message limits
    chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

    # Send each chunk as a separate message
    for chunk in chunks:
        await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


async def determine_investment_personality(assessment_data):
    try:
        # Prepare input text for the chatbot based on assessment data
        input_text = "User Profile:\n"
        for question, answer in assessment_data.items():
            input_text += f"{question}: {answer}\n"

        # Introduce the chatbot's task and prompt for classification
        input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
                      "- Conservative Investor\n" \
                      "- Moderate Investor\n" \
                      "- Aggressive Investor"

        # Use your generative AI model to generate a response
        # print(input_text)
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(input_text)

        # Determine the investment personality from the chatbot's response
        response_text = response.text.lower()
        if "conservative" in response_text:
            personality = "Conservative Investor"
        elif "moderate" in response_text:
            personality = "Moderate Investor"
        elif "aggressive" in response_text:
            personality = "Aggressive Investor"
        else:
            personality = "Unknown"

        return personality
        # Send the determined investment personality back to the user
        #await bot.send_message(chat_id, f"Investment Personality: {personality}")

    except Exception as e:
        print(f"Error generating response: {e}")
        #await bot.send_message(chat_id, "Error processing investment personality classification.")


# Handler for document upload
async def load_vector_db(file_path):
    try:
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        return vector_store.as_retriever(search_kwargs={"k": 1})
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None



async def make_retrieval_chain(retriever):
    """
    Create a retrieval chain using the provided retriever.

    Args:
        retriever (RetrievalQA): A retriever object.

    Returns:
        RetrievalQA: A retrieval chain object.
    """
    try:
        global investment_personality,summary
        llm = ChatGoogleGenerativeAI(
            #model="gemini-pro",
            model = "gemini-1.5-flash",
            temperature=0.7,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )

        prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
                Respond to the client by the client name.
                Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
                Also give the user detailed information about the investment how to invest,where to invest and how much they
                should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
                Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
                investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
                Also give the user minimum and maximum expected growth in dollars for the time horizon .
                Also explain the user why you are giving them that particular investment suggestion.
                Give the client suggestions of Investment based on their investment personality that can also help them manage their mortages and other liablilities if they have any,ignore if they dont have any liabilities.
                Answer in 3-4 lines.\n
                <context>
                {context}
                </context>
                Question: {input}"""

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in creating chain: {e}")
        return None

from aiogram.filters import Filter

# @router.message(F.document)
@dp.message(F.document)
async def handle_document(message: types.Message):
    global summary,investment_personality  

    chat_id = message.chat.id

    # Obtain file information
    file_id = message.document.file_id
    file = await bot.get_file(file_id)
    file_path = file.file_path
    
    # Download the file
    await bot.download_file(file_path, "data/uploaded_file")
    
    # Process the uploaded document
    extracted_text = await process_document("data/uploaded_file")
    
    if extracted_text:
        # Load vector database (assuming this is part of setting up the retriever)
        retriever = await load_vector_db("data/uploaded_file")
        file_path = 'data/uploaded_file'
        client_name, validation_errors = await process_document(file_path)

        # Print results
        print(f"Client Name: {client_name}")
        if validation_errors:
            print("**Validation Errors:**")
            for error in validation_errors:
                print(error)
        else:
            print("All fields are filled correctly.")
        if client_name == None:
            try:
                await message.reply("Processing the uploaded image")
                await handle_image(message) 
                return 
            except Exception as e:
                await message.reply("error processing uploaded image")
                print(e)
        await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")

        if retriever is None:
            await message.reply("The retrieval chain is not set up. Please upload a document first.")
            return

        # Check if a valid chain can be created
        chain = await make_retrieval_chain(retriever)
        if chain is None:
            await message.reply("Failed to create the retrieval chain.")
            return
        
        try:     
            query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
        
            response = chain.invoke({"input": query})
            print(response['answer'])
            global chat_history
            chat_history = response['answer'] 
            print(f"\n Chat History : {chat_history}")
            format_response = markdown_to_text(response['answer'])

            # Store the extracted_text in chat history
            write_chat_history(chat_id, {'role': 'bot', 'message': extracted_text})
        
            # Store the response in chat history
            write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

            await message.reply(format_response)
            # await message.reply(response['answer'])

            try:
                graph_query = (
                    # extracted_text + "\n" + 
                    summary + "\n" + "investment_personality" + "\n" + response['answer'] + "\n" +
                    "Please provide the following information in JSON format:\n" +
                    "{\n"
                    "  'growth_investment_min': <minimum percentage of growth-oriented investments>,\n"
                    "  'growth_investment_max': <maximum percentage of growth-oriented investments>,\n"
                    "  'fixed_income_min': <minimum percentage of fixed-income investments>,\n"
                    "  'fixed_income_max': <maximum percentage of fixed-income investments>,\n"
                    "  'cash_min': <minimum percentage of cash investments>,\n"
                    "  'cash_max': <maximum percentage of cash investments>,\n"
                    "  'return_growth_investment': <Maximum percentage returns of growth-oriented investments>,\n"
                    "  'return_fixed_income': <Maximum percentage of returns of fixed-income investments>,\n"
                    "  'return_cash_max': <maximum percentage of retruns of cash investments>,\n"
                    "  'return_min': <minimum expected annual return percentage>,\n"
                    "  'return_max': <maximum expected annual return percentage>,\n"
                    "  'growth_min': <minimum expected growth in dollars>,\n"
                    "  'growth_max': <maximum expected growth in dollars>,\n"
                    "  'initial_investment': <initial monthly investment>,\n"
                    "  'time_horizon': <time horizon in years>\n"
                    "}"
                )
                graph_response = chain.invoke({"input": graph_query})
                print(graph_response['answer'])
                await handle_graph(graph_response['answer'],chat_id)
                try:
                    await refine_investment_suggestions(chat_id, response['answer'])  # Generate refined suggestions
                except Exception as e:
                    print(f"Error refining investment suggestions: {e}")
                    await bot.send_message(chat_id, "Error refining investment suggestions. Please try again later.")

            except Exception as e:
                print(f"Error plotting graph : {e}")
        except Exception as e:
            print(f"Error invoking retrieval chain on attempt : {e}")
            await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

    else:
        await message.reply("Failed to process the uploaded file.")
    

# Function to extract data from LLM response
def extract_data_from_response(response):
    try:
        # Locate the JSON-like data in the response
        json_start = response.find("{")
        json_end = response.rfind("}") + 1
        
        if json_start == -1 or json_end == -1:
            raise ValueError("No JSON data found in the response.")
        
        json_data = response[json_start:json_end]
        
        # Parse the JSON data
        data = json.loads(json_data.replace("'", "\""))
        print(data)
        return data
    except Exception as e:
        logging.error(f"Error extracting data: {e}")
        return None


async def handle_graph(response, chat_id):
    try:
        data = extract_data_from_response(response)
        if not data:
            await bot.send_message(chat_id, "Failed to extract data from the response.")
            return
        
        # Log extracted data for debugging
        logging.info(f"Extracted data: {data}")
        
        # Create a pie chart for investment allocation
        labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
        sizes = [
            (data['growth_investment_min'] + data['growth_investment_max']) / 2,
            (data['fixed_income_min'] + data['fixed_income_max']) / 2,
            (data['cash_min'] + data['cash_max']) / 2
        ]
        fig1, ax1 = plt.subplots()
        ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
        pie_chart_buffer = io.BytesIO()
        plt.savefig(pie_chart_buffer, format='png')
        pie_chart_buffer.seek(0)
        pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

        # Create a bar graph for potential returns
        allocations = {
            'Growth-Oriented Investments': (data['growth_investment_min'] + data['growth_investment_max']) / 2,
            'Fixed-Income Investments': (data['fixed_income_min'] + data['fixed_income_max']) / 2,
            'Cash and Cash Equivalents': (data['cash_min'] + data['cash_max']) / 2
        }

        returns = {
            'Growth-Oriented Investments': (data['return_growth_investment'] ) / 100,  # Annual return
            'Fixed-Income Investments': (data['return_fixed_income'] ) / 100 ,
            'Cash and Cash Equivalents': (data['return_cash_max'] ) / 100
        }
        
        time_horizon_years = data['time_horizon']
        initial_monthly_investment = data['initial_investment']
        
        # Calculate total returns for each category
        total_investment = initial_monthly_investment * 12 * time_horizon_years
        total_returns = {category: total_investment * (1 + returns[category]) ** time_horizon_years for category in allocations.keys()}
        
        # Create the bar chart
        fig2, ax2 = plt.subplots(figsize=(10, 6))
        bars = ax2.bar(allocations.keys(), total_returns.values(), color=['#FF9999', '#66B2FF', '#99FF99'])

        # Add labels to the top of each bar
        for bar in bars:
            yval = bar.get_height()
            ax2.text(bar.get_x() + bar.get_width()/2, yval + 1000, f'₹{round(yval, 2)}', ha='center', va='bottom')

        # Add titles and labels
        plt.title('Investment Performance Over Time')
        plt.xlabel('Investment Categories')
        plt.ylabel('Total Returns (₹)')
        plt.grid(True, axis='y', linestyle='--', linewidth=0.7)
        
        # Save and display the plot
        bar_chart_buffer = io.BytesIO()
        plt.savefig(bar_chart_buffer, format='png')
        bar_chart_buffer.seek(0)
        bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='investment_performance.png')

        await bot.send_document(chat_id, pie_chart_file, caption="Investment Allocation Chart")
        await bot.send_document(chat_id, bar_chart_file, caption="Investment Performance Over Time")

    except Exception as e:
        logging.error(f"Error plotting graph: {e}")
        await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")


 # Detailed Stock Suggestions :        
async def refine_investment_suggestions(chat_id, suggestions):
    try:
        task = """You are a Stock Market Expert. You know everything about stock market trends and patterns.
                  You are an expert Financial Advisor and wealth manager.
                  Based on the following investment suggestions, provide more detailed investment advice specifically 
                  on Investment Allocation of stocks:
                  How to invest in these stocks, what percentage of the investment should go into each stock,
                  where to invest in terms of sector and how much,always give some examples of these  
                  suggested sectors stocks/etfs/mutual funds that the user can invest or look at ,
                  Give a detailed allocation of each stock/etfs/mutual funds investment the user should do per month based on their 
                  monthly investment amount,give the users expected returns on their investments from each allocations ,
                  then explain the reasoning behind each suggestion to the user. 
                  Also, include any potential risks and recommendations for monitoring the investments.
                  Give the response in brief in around 15-20 lines."""

        query = task + "\n" + suggestions
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(query)
        await bot.send_message(chat_id, "Here are more detailed stock investment suggestions:")
        format_response = markdown_to_text(response.text) 

        # Store the response in chat history
        write_chat_history(chat_id, {'role': 'bot', 'message': format_response})
        await bot.send_message(chat_id, format_response) #response.text

        try:
            # Read chat history again to include the latest response
            chat_history = read_chat_history(chat_id)
            # Extract allocations and create pie chart
            allocations = extract_allocations_from_json(chat_history,chat_id) #extract_allocations(format_response)
            if allocations:
                chart_path = create_pie_chart(allocations, chat_id)
                if chart_path:
                    await bot.send_photo(chat_id, photo=FSInputFile(chart_path))
                else:
                    await bot.send_message(chat_id,"Could not generate a pie chart. Please check the investment allocations provided.")
            else:
                await bot.send_message(chat_id,"No valid allocations found to generate a pie chart.")

            chart_path = create_pie_chart(allocations, chat_id)
            
            # await bot.send_photo(chat_id, photo=FSInputFile(chart_path))
            # if chart_path:
            #     await bot.send_photo(chat_id, photo=FSInputFile(chart_path))
            # else:
            #     await bot.send_message(chat_id,"Could not generate a pie chart. Please check the investment allocations provided.")

        except Exception as e:
            logging.error(f"Error generating pi chart : {e}")
            await bot.send_message(chat_id,"An error occurred while generating pi chart")

        task2 = """You are a Stock Market Expert. You know everything about stock market trends and patterns.
                You are an expert Financial Advisor and wealth manager.
                Based on the response geenerated,write ticker names of the stocks/etfs/mutual funds/anything that was 
                suggested in the response.The ticker names must be accurate,if there arent proper stocks/etfs/mfs suggestions 
                then you provide suggestions with ticker names based on the suggestion for the user.
                Just give the user names by saying here are some ticker names that are recommended for you."""
        
        query2 = task2 + "\n" + format_response
        response2 = model.generate_content(query2)
        # await bot.send_message(chat_id, "Here are some stock names that are recommended for you:")
        format_response2 = markdown_to_text(response2.text)
        await bot.send_message(chat_id, format_response2) #response2.text
        await bot.send_message(chat_id,"Get more detailed information about these stocks by typing the ticker names in the chat.")
        

    except Exception as e:
        logging.error(f"Error generating refined investment suggestions: {e}")
        await bot.send_message(chat_id, "There was an error generating more detailed investment suggestions. Please try again later.")


# List of ticker symbols mentioned in the suggestion
suggested_tickers = ['SPY', 'VOO', 'IJH', 'VO', 'IJR', 'VB']

def extract_allocations_from_json(json_data,chat_id):
    allocations = {}
    for entry in json_data.get(str(chat_id), []):
        if entry['role'] == 'bot':
            message = entry['message']
            lines = message.split('\n')
            current_category = None

            for line in lines:
                match = re.match(r'^(.*?):\s*(\d+)%$', line)
                if match:
                    category, percent = match.groups()
                    allocations[category] = []
                    current_category = category
                elif current_category and re.match(r'.*\d+%', line):
                    subcategory_match = re.match(r'^(.*?)(\d+)%$', line)
                    if subcategory_match:
                        subcategory, percent = subcategory_match.groups()
                        allocations[current_category].append((subcategory.strip(), float(percent)))

    return allocations


def create_pie_chart(allocations, chat_id):
    labels = []
    sizes = []
    for category, subcategories in allocations.items():
        for subcategory, percent in subcategories:
            labels.append(f"{category} - {subcategory}")
            sizes.append(percent)
    
    if sizes:
        fig, ax = plt.subplots()
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
        ax.axis('equal')
        
        plt.title("Investment Allocation")
        chart_path = f"data/investment_allocation_{chat_id}.png"
        plt.savefig(chart_path)
        plt.close()
        
        return chart_path
    else:
        return None
    
# def extract_allocations(llm_response):
#     allocations = {}
#     current_category = None
    
#     for line in llm_response.split('\n'):
#         # Match categories and percentages
#         match = re.match(r'^(.*?):\s*(\d+%)$', line)
#         if match:
#             category, percent = match.groups()
#             allocations[category] = []
#             current_category = category
#         # Match subcategories and percentages
#         elif current_category and re.match(r'.*\d+%', line):
#             subcategory_match = re.match(r'^(.*?)(\d+)%$', line)
#             if subcategory_match:
#                 subcategory, percent = subcategory_match.groups()
#                 allocations[current_category].append((subcategory.strip(), float(percent)))
    
#     return allocations

# def create_pie_chart(allocations, chat_id):
#     labels = []
#     sizes = []
#     for category, subcategories in allocations.items():
#         for subcategory, percent in subcategories:
#             labels.append(f"{category} - {subcategory}")
#             sizes.append(percent)
    
#     fig, ax = plt.subplots()
#     ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
#     ax.axis('equal')
    
#     plt.title("Investment Allocation")
#     chart_path = f"data/investment_allocation_{chat_id}.png"
#     plt.savefig(chart_path)
#     plt.close()
    
#     return chart_path



async def process_document(file_path):
    try:
        file_type = filetype.guess(file_path)
        if file_type is not None:
            if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # return extract_text_from_word(file_path)
                return extract_text_and_tables_from_word(file_path)
            elif file_type.mime == "application/pdf":
                return extract_text_from_pdf(file_path)
        return None
    except Exception as e:
        print(f"Error processing document: {e}")
        return None

def extract_text_from_pdf(pdf_file_path):
    try:
        with open(pdf_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            text_content = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text_content.append(page.extract_text())
            return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None


import re
import docx

def extract_text_and_tables_from_word(docx_file_path):
    """
    Extracts text and tables from a Word document (.docx).

    Args:
        docx_file_path (str): Path to the Word document file.

    Returns:
        tuple: Extracted text content and tables from the document.
    """
    try:
        doc = docx.Document(docx_file_path)
        text_content = []
        tables_content = []

        for para in doc.paragraphs:
            text_content.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)

        return "\n".join(text_content), tables_content
    except Exception as e:
        print(f"Error extracting text and tables from Word document: {e}")
        return None, None

def validate_document_content(text, tables):
    """
    Validates the content of the document.

    Args:
        text (str): Extracted text content from the document.
        tables (list): Extracted tables content from the document.

    Returns:
        tuple: Client name and validation errors.
    """
    errors = []
    
    # Extract client name
    client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
    client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

    # Define required sections
    required_sections = [
        "YOUR RETIREMENT GOAL",
        "YOUR OTHER MAJOR GOALS",
        "YOUR ASSETS AND LIABILITIES",
        "MY LIABILITIES",
        "YOUR CURRENT ANNUAL INCOME"
    ]

    # Check for the presence of required sections
    for section in required_sections:
        if section not in text:
            errors.append(f"* {section} section missing.")
    
    # Define table field checks
    table_checks = {
        "YOUR RETIREMENT GOAL": [
            r"When do you plan to retire\? \(age or date\)",
            r"Social Security Benefit \(include expected start date\)",
            r"Pension Benefit \(include expected start date\)",
            r"Other Expected Income \(rental, part-time work, etc.\)",
            r"Estimated Annual Retirement Expense"
        ],
        "YOUR OTHER MAJOR GOALS": [
            r"GOAL", r"COST", r"WHEN"
        ],
        "YOUR ASSETS AND LIABILITIES": [
            r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
            r"Current Value", r"Annual Contributions"
        ],
        "MY LIABILITIES": [
            r"Balance", r"Interest Rate", r"Monthly Payment"
        ]
    }

    # Validate table content
    for section, checks in table_checks.items():
        section_found = False
        for table in tables:
            table_text = "\n".join(["\t".join(row) for row in table])
            if section in table_text:
                section_found = True
                for check in checks:
                    if not re.search(check, table_text, re.IGNORECASE):
                        errors.append(f"* Missing or empty field in {section} section: {check}")
                break
        if not section_found:
            errors.append(f"* {section} section missing.")

    return client_name, errors

async def process_document(file_path):
    try:
        text, tables = extract_text_and_tables_from_word(file_path)
        if text is not None and tables is not None:
            client_name, errors = validate_document_content(text, tables)
            return client_name, errors
        return None, ["Error processing document."]
    except Exception as e:
        print(f"Error processing document: {e}")
        return None, [f"Error processing document: {e}"]

@dp.message()
async def main_bot(message: types.Message):
    global retriever, extracted_text, investment_personality, summary, chat_history

    chat_id = message.chat.id

    if chat_id in states and states[chat_id] < len(questions):
        question_index = states[chat_id]
        answer = message.text
        user_responses[questions[question_index]] = answer
        states[chat_id] += 1
        await ask_next_question(chat_id, question_index + 1)
    elif message.text:
        if message.text.upper().isalpha() and len(message.text) <= 5:
            ticker = message.text.upper()
            try:
                hist, data, excel_file = await fetch_stock_data(message, ticker)
                mean_price, analysis = await analyze_stock_data(hist, data, excel_file,chat_id)
                await message.reply(f"{mean_price}")
                await message.reply(f"Stock analysis for {ticker}:\n{analysis}")
            except Exception as e:
                logging.error(f"Error fetching stock data for {ticker}: {e}")
                await message.reply("Failed to fetch stock data. Please try again.")
        else:
            try:
                task = """You are a Financial Expert and Wealth Advisor.
                    You also a Stock Market Expert. You know everything about stock market trends and patterns.
                    Provide financial advice or Stock Related advice and suggestions based on the user's query.
                    Consider user's investment personality and Financial Details if provided.
                    Address the user by their name(client_name: Emily in our case but if any other name is give refer to that) if provided.
                    Include detailed information about the investment, where to invest, how much to invest, 
                    expected returns, and why you are giving this advice.
                    As you are a Wealth Advisor if user asks queries related to saving taxes or calculating taxes refer to the 
                    US Tax Laws given by the IRS and based on that information calculate the taxes for the user 
                    consider the information shared by the user such as their annual income and their monthly investment if provided,
                    also give advice to the user on how they can save their taxes.
                    Include a disclaimer to monitor and rebalance investments regularly based on risk tolerance."""
                
                # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + message.text
                # query = chat_history_text + "\n" + query

                # Include chat history
                chat_history = read_chat_history(chat_id)
                chat_history_text = '\n'.join([f"{entry['role']}: {entry['message']}" for entry in chat_history])
                # history.append(chat_history_text)
                query = task + "\n" + investment_personality + "\n" + chat_history_text + "\n" + message.text

                model = genai.GenerativeModel('gemini-1.5-flash')
                chat = model.start_chat(history=[])
                response = chat.send_message(query)

                # Enhanced logging for debugging
                logging.info(f"Model response: {response}")
                format_response = markdown_to_text(response.text) #(response_text) #response.result

                # Store the response in chat history
                write_chat_history(chat_id, {'role': 'bot', 'message': format_response})
                await message.reply(format_response)

            except Exception as e:
                logging.error(f"Error processing general chat message: {e}")
                await message.reply("Failed to process your request.")

import yfinance as yf
import pandas as pd
import requests
from aiogram import types
from aiogram.types import InputFile
import os
import logging

NEWS_API_KEY = os.getenv('NEWS_API_KEY')

async def fetch_stock_data(message, ticker):
    stock = yf.Ticker(ticker)
    chat_id = message.chat.id
    data = {}

    company_details = stock.info.get('longBusinessSummary', 'No details available')
    data['Company Details'] = company_details
    sector = stock.info.get('sector', 'No sector information available')
    data['Sector'] = sector
    prev_close = stock.info.get('previousClose', 'No previous close price available')
    data['Previous Closing Price'] = prev_close
    open_price = stock.info.get('open', 'No opening price available')
    data['Today Opening Price'] = open_price

    hist = stock.history(period="5d")
    if not hist.empty and 'Close' in hist.columns:
        if hist.index[-1].date() == yf.download(ticker, period="1d").index[-1].date():
            close_price = hist['Close'].iloc[-1]
            data['Todays Closing Price'] = close_price
        else:
            data['Todays Closing Price'] = "Market is open, there is no closing price available yet."
    else:
        data['Todays Closing Price'] = "No historical data available for closing price."

    day_high = stock.info.get('dayHigh', 'No high price available')
    data['Today High Price'] = day_high
    day_low = stock.info.get('dayLow', 'No low price available')
    data['Today Low Price'] = day_low
    volume = stock.info.get('volume', 'No volume information available')
    data['Today Volume'] = volume
    dividends = stock.info.get('dividendRate', 'No dividend information available')
    data['Today Dividends'] = dividends
    splits = stock.info.get('lastSplitFactor', 'No stock split information available')
    data['Today Stock Splits'] = splits
    pe_ratio = stock.info.get('trailingPE', 'No P/E ratio available')
    data['P/E Ratio'] = pe_ratio
    market_cap = stock.info.get('marketCap', 'No market cap available')
    data['Market Cap'] = market_cap

    income_statement = stock.financials
    balance_sheet = stock.balance_sheet
    cashflow = stock.cashflow

    news_url = f'https://newsapi.org/v2/everything?q={ticker}&apiKey={NEWS_API_KEY}&pageSize=3'
    news_response = requests.get(news_url)
    if news_response.status_code == 200:
        news_data = news_response.json()
        articles = news_data.get('articles', [])
        if articles:
            top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
            data['Top News'] = top_news
        else:
            data['Top News'] = "No news articles found."
    else:
        data['Top News'] = "Failed to fetch news articles."

    for key, value in data.items():
        await message.reply(f"{key}: {value}")

    graph_url = f"https://finance.yahoo.com/chart/{ticker}"
    data['graph_url'] = graph_url
    await message.reply(f"Stock Chart: \n{graph_url}")

    file_path = os.path.join('data', f'{ticker}_financial_data.xlsx')
    with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
        income_statement.to_excel(writer, sheet_name='Income Statement')
        balance_sheet.to_excel(writer, sheet_name='Balance Sheet')
        cashflow.to_excel(writer, sheet_name='Cashflow')

    excel_file = FSInputFile(file_path, filename=f'{ticker}_financial_data.xlsx')
    await bot.send_document(chat_id, excel_file, caption=f'{ticker} Financial Data')

    data_list = list(data.items())
    data_str = str(data_list)

    return hist, data_str, file_path


async def analyze_stock_data(hist, data, excel_file,chat_id):
    avg_close = hist['Close'].mean()
    formatted_data = extract_excel_data(excel_file)

    task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.
                Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
                Predict the stock price range for the next week and provide reasons for your prediction.
                Advise whether to buy this stock now or not, with reasons for your advice."""

    query = task + "\nStock Data: " + data + "\nFinancial Data: " + formatted_data
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(query)
    
    # Log the response object to understand its structure
    logging.info(f"Model response: {response}")
    
     # Extract the text content from the response
    try:
        response_text = response.text #result.candidates[0].content.parts[0].text  #response.candidates[0]['content']['parts'][0]['text']
        format_response = markdown_to_text(response_text)
    except Exception as e:
        logging.error(f"Error extracting text from response: {e}")
        return f"Average closing price for the last 3 months: ${avg_close:.2f}", response
    

    # Store the response in chat history
    write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

    return f"Average closing price for the last 3 months: ${avg_close:.2f}", format_response


def extract_excel_data(file_path):
    financial_data = ""
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        financial_data += f"\n\nSheet: {sheet_name}\n"
        financial_data += df.to_string()
    return financial_data

def markdown_to_text(md):
    # Simple conversion for markdown to plain text
    md = md.replace('**', '')
    md = md.replace('*', '')
    md = md.replace('_', '')
    md = md.replace('#', '')
    md = md.replace('`', '')
    return md.strip()


from aiogram.types.input_file import BufferedInputFile
from aiogram import BaseMiddleware
# from aiogram.dispatcher.router import Router
from PIL import Image

# Function to handle image messages
# @dp.message(F.photo)
# @router.message(F.photo)
import PIL.Image

async def handle_image(message: types.Message):
    global investment_personality, chat_history

    chat_id = message.chat.id
    # Handle image inputs
    try:
        # Obtain file information
        try:
            photo_id = message.document.file_id
            photo = await bot.get_file(photo_id)
            photo_path = photo.file_path
            # Download the file
            photo_file = await bot.download_file(photo_path, "data/uploaded_image.png")

        except Exception as e:
            print(f"Error downloading image: {e}")
            await bot.send_message(chat_id, "Error processing image. Please try again.")
            return
        
        # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
        #             Also give the user detailed information about the investment how to invest, where to invest and how much they
        #             should invest in terms of percentage of their investment amount. Give the user detailed information about the returns on their 
        #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compounded returns on their 
        #             investment. Also explain the user why you are giving them that particular
        #             investment suggestion. Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
        #             User should also invest as per their risk tolerance level. Since you are the financial advisor don't ask user to consult anyone else.
        #             So don't mention user to consult to a financial expert."""

        task = """You are a Financial Expert.You will be provided with a Financial Form from Boston Harbor.
                If you recieve any other image tell the user to Upload the Images of the form or upload the word document of the form.
                You are supposed to Respond to the user's Image query and If they ask for any information provide them the information in Detail.
                Be helpful and informative.Give proper information of any Financial terms the user may ask you.Address the user by their Client Name if provided.
                Also provide the user helpful links so that they can refer to the link for more information.
                If the image provided is not related to Finance then just answer about the image and any caption if provided.
                """

        prompt = message.caption if message.caption else ""  # Use the photo caption if available
        # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + prompt
        query = task + prompt 

        image =  PIL.Image.open('data/uploaded_image.png') #(photo_file) 
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(image)
        await bot.send_message(chat_id,"I will describe the image that was uploaded")
        format_response = markdown_to_text(response.text)
        await message.reply(format_response)
        # await message.reply(response.text)

        # chat = model.start_chat(history=[])
        # response = chat.send_message(query)
        # format_response = markdown_to_text(response.result)
        # await message.reply(format_response)

        response = model.generate_content([query, image])
        format_response = markdown_to_text(response.text)

        # Store the response in chat history
        write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

        await message.reply(format_response)
        # await message.reply(response.text) 
    except Exception as e:
        logging.error(f"Error generating response for the image: {e}")
        await message.reply("There was an error generating response for the image. Please try again later.")
    # await message.reply("Cant process the image")
    # return



from aiogram.filters import command
from aiogram.types import bot_command
import markdown
from bs4 import BeautifulSoup

def markdown_to_text(markdown_text):
    # Convert markdown to HTML
    html = markdown.markdown(markdown_text)
    # Parse the HTML
    soup = BeautifulSoup(html, 'html.parser')
    # Extract plain text
    text = soup.get_text()
    return text


# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)

async def main() -> None:
    # Initialize Bot instance with default bot properties which will be passed to all API calls
    bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

    # And the run events dispatching
    await dp.start_polling(bot)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, stream=sys.stdout)
    asyncio.run(main())


