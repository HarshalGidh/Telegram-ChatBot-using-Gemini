import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt

import os
import filetype
import docx
import PyPDF2
import re
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.memory import ConversationSummaryMemory
import asyncio
import numpy as np
import json

import google.generativeai as genai
import pathlib
import logging
import sys
import io
import matplotlib.pyplot as plt
import seaborn as sns
# Import things that are needed generically
from langchain.pydantic_v1 import BaseModel, Field
from langchain.tools import BaseTool, StructuredTool, tool
# Define functions to generate investment suggestions :

load_dotenv()
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

from flask import Flask, request, jsonify

app = Flask(__name__)

# Configure generativeai with your API key
genai.configure(api_key=GOOGLE_API_KEY)

import markdown
# def convert_to_markdown(raw_text):
#     # Replace specific text patterns with markdown syntax
#     formatted_text = raw_text.replace('\n', '\n\n')  # Ensure newlines create paragraphs
    
#     # Convert text into markdown format
#     html = markdown.markdown(formatted_text)

#     return html

def markdown_table_to_html(md_table):
    # Split the markdown table by lines
    lines = md_table.strip().split("\n")
    
    # Extract headers and rows
    headers = lines[0].strip('|').split('|')
    rows = [line.strip('|').split('|') for line in lines[2:]]  # Skip the separator line

    # Start creating the HTML table
    html_table = "<table>\n"
    
    # Add headers
    html_table += "  <thead>\n    <tr>\n"
    for header in headers:
        html_table += f"      <th>{header.strip()}</th>\n"
    html_table += "    </tr>\n  </thead>\n"
    
    # Add rows
    html_table += "  <tbody>\n"
    for row in rows:
        html_table += "    <tr>\n"
        for col in row:
            html_table += f"      <td>{col.strip()}</td>\n"
        html_table += "    </tr>\n"
    html_table += "  </tbody>\n</table>"

    return html_table



import markdown2
from bs4 import BeautifulSoup

def markdown_to_readable_text(md_text):
    # Convert markdown to HTML
    html = markdown2.markdown(md_text)

    # Parse the HTML
    soup = BeautifulSoup(html, "html.parser")

    # Function to format plain text from tags
    def format_text_from_html(soup):
        formatted_text = ''
        for element in soup:
            if element.name == "h1":
                formatted_text += f"\n\n# {element.text.upper()} #\n\n"
            elif element.name == "h2":
                formatted_text += f"\n\n## {element.text} ##\n\n"
            elif element.name == "h3":
                formatted_text += f"\n\n### {element.text} ###\n\n"
            elif element.name == "strong":
                formatted_text += f"**{element.text}**"
            elif element.name == "em":
                formatted_text += f"_{element.text}_"
            elif element.name == "ul":
                for li in element.find_all("li"):
                    formatted_text += f"\n - {li.text}"
            elif element.name == "ol":
                for idx, li in enumerate(element.find_all("li"), 1):
                    formatted_text += f"\n {idx}. {li.text}"
            elif element.name == "table":
                # Convert markdown table to HTML table
                formatted_text += "<table>\n"
                rows = element.find_all("tr")
                for row in rows:
                    formatted_text += "<tr>\n"
                    cols = row.find_all(["th", "td"])
                    for col in cols:
                        tag = 'th' if col.name == "th" else 'td'
                        formatted_text += f"<{tag}>{col.text.strip()}</{tag}>\n"
                    formatted_text += "</tr>\n"
                formatted_text += "</table>\n"
            else:
                formatted_text += element.text

        return formatted_text.strip()

    return format_text_from_html(soup)


# def markdown_to_readable_text(md_text):
#     # Convert markdown to HTML
#     html = markdown2.markdown(md_text)

#     # Parse the HTML
#     soup = BeautifulSoup(html, "html.parser")

#     # Function to format plain text from tags
#     def format_text_from_html(soup):
#         formatted_text = ''
#         for element in soup:
#             if element.name == "h1":
#                 formatted_text += f"\n\n# {element.text.upper()} #\n\n"
#             elif element.name == "h2":
#                 formatted_text += f"\n\n## {element.text} ##\n\n"
#             elif element.name == "h3":
#                 formatted_text += f"\n\n### {element.text} ###\n\n"
#             elif element.name == "strong":
#                 formatted_text += f"**{element.text}**"
#             elif element.name == "em":
#                 formatted_text += f"_{element.text}_"
#             elif element.name == "ul":
#                 for li in element.find_all("li"):
#                     formatted_text += f"\n - {li.text}"
#             elif element.name == "ol":
#                 for idx, li in enumerate(element.find_all("li"), 1):
#                     formatted_text += f"\n {idx}. {li.text}"
#             elif element.name == "table":
#                 rows = element.find_all("tr")
#                 for row in rows:
#                     cols = row.find_all(["th", "td"])
#                     row_text = ' | '.join(col.text.strip() for col in cols)
#                     formatted_text += f"{row_text}\n"
#                 formatted_text += "\n"
#             else:
#                 formatted_text += element.text

#         return formatted_text.strip()

#     return format_text_from_html(soup)

def markdown_to_text(md): # og solution code 
    # Simple conversion for markdown to plain text
    md = md.replace('**', '')
    md = md.replace('*', '')
    md = md.replace('_', '')
    md = md.replace('#', '')
    md = md.replace('`', '')
    return md.strip()


# import docx

# def extract_responses_from_docx(personality_file):
#     """
#     Extracts responses from a Word document (.docx) where answers are typed in.

#     Args:
#         personality_file (UploadedFile): The file object uploaded via Streamlit.

#     Returns:
#         dict: A dictionary containing the questions and the typed answers.
#     """
#     try:
#         doc = docx.Document(personality_file)
#         responses = {}
#         current_question = None

#         # Check paragraphs
#         for para in doc.paragraphs:
#             text = para.text.strip()
#             if text:
#                 # Check if the paragraph contains a question
#                 if "?" in text or text.endswith(":"):
#                     current_question = text
#                     st.write(f"Identified question: {current_question}")  # Debugging log
#                 else:
#                     # This is a typed answer
#                     typed_answer = text.strip()
#                     st.write(f"Identified typed answer: {typed_answer}")  # Debugging log
#                     if current_question:
#                         # If the question already has an answer, append to it (handles multiple responses)
#                         if current_question in responses:
#                             responses[current_question] += "; " + typed_answer
#                         else:
#                             responses[current_question] = typed_answer

#             # Debugging log to understand document structure
#             st.write(f"Processing paragraph: {text}")  # Console log for local testing

#         # Check tables for additional responses
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     text = cell.text.strip()
#                     if text:
#                         if "?" in text or text.endswith(":"):
#                             current_question = text
#                             st.write(f"Identified question in table: {current_question}")  # Debugging log
#                         else:
#                             typed_answer = text.strip()
#                             st.write(f"Identified typed answer in table: {typed_answer}")  # Debugging log
#                             if current_question:
#                                 if current_question in responses:
#                                     responses[current_question] += "; " + typed_answer
#                                 else:
#                                     responses[current_question] = typed_answer

#         if responses:
#             st.write("Extracted Responses:")
#             for question, answer in responses.items():
#                 st.write(f"**{question}**: {answer}")
#         else:
#             st.write("No responses captured. Please check the document formatting or symbols used.")

#         return responses

#     except Exception as e:
#         st.write(f"Error extracting responses: {e}")  # Console log for local testing
#         return None

# def determine_investment_personality(responses):
#     """
#     Determines the investment personality based on extracted responses.

#     Args:
#         responses (dict): A dictionary containing the questions and the selected answers.

#     Returns:
#         str: The determined investment personality.
#     """
#     try:
#         # Prepare input text for the chatbot based on extracted responses
#         input_text = "User Profile:\n"
#         for question, response in responses.items():
#             input_text += f"{question}: {response}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYour task is to determine the investment personality based on the above profile."

#         # Here you would send the input_text to your chatbot or classification model
#         # For demonstration, we'll just return the input_text
#         return input_text

#     except Exception as e:
#         st.write(f"Error determining investment personality: {e}")  # Console log for local testing
#         return None

# def extract_responses_from_docx(personality_file):
#     try:
#         doc = docx.Document(personality_file)
#         responses = {}
#         current_question = None

#         # Check paragraphs
#         for para in doc.paragraphs:
#             text = para.text.strip()
#             if text:
#                 # Check if the paragraph contains a question
#                 if "?" in text or text.endswith(":"):
#                     current_question = text
#                 else:
#                     # This is a typed answer
#                     typed_answer = text.strip()
#                     if current_question:
#                         # If the question already has an answer, append to it (handles multiple responses)
#                         if current_question in responses:
#                             responses[current_question] += "; " + typed_answer
#                         else:
#                             responses[current_question] = typed_answer

#         # Check tables for additional responses
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     text = cell.text.strip()
#                     if text:
#                         if "?" in text or text.endswith(":"):
#                             current_question = text
#                         else:
#                             typed_answer = text.strip()
#                             if current_question:
#                                 if current_question in responses:
#                                     responses[current_question] += "; " + typed_answer
#                                 else:
#                                     responses[current_question] = typed_answer

#         return responses

#     except Exception as e:
#         print(f"Error extracting responses: {e}")
#         return None

import docx

# # GET Method for me POST method for Frontend
def extract_responses_from_docx(personality_file): # Using text responses parsing
    """
    Extracts responses from a Word document (.docx) where the selected answers are listed as text after the options.

    Args:
        personality_file (str): Path to the Word document file.

    Returns:
        dict: A dictionary containing the questions and the selected answers.
    """
    try:
        doc = docx.Document(personality_file)
        responses = {}
        current_question = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect the beginning of a question
                if "?" in text:
                    current_question = text
                # Detect a chosen response (assuming it follows the question and options)
                elif current_question and not text.startswith(("a.", "b.", "c.", "d.")):
                    selected_answer = text
                    responses[current_question] = selected_answer
                    current_question = None  # Reset for the next question

        if responses:
            print(responses)
            # st.write(responses)
        else:
            print("\nNo responses captured")
            st.write("No responses captured")
        return responses
    except Exception as e:
        print(f"Error extracting responses: {e}")
        return None

# def extract_responses_from_assessment(personality_file): # using boxes
#     # Load the document
#     # doc = Document(docx_filename)
#     doc = docx.Document(personality_file)
    
#     # Initialize a list to store responses
#     responses = []
    
#     # Iterate through each paragraph in the document
#     for para in doc.paragraphs:
#         text = para.text.strip()
#         # Check if the paragraph contains a checkbox
#         if '☒' in text or '☐' in text:
#             # Extract the response marked with ☒
#             if '☒' in text:
#                 response = text.split('☒')[1].strip()
#                 responses.append(response)
    
#     return responses

# import asyncio
# # from some_generative_ai_library import GenerativeModel  # Replace with actual import

# async def determine_investment_personality(assessment_data):
#     try:
#         # Prepare input text for the chatbot based on assessment data
#         input_text = "User Profile:\n"
#         for question, answer in assessment_data.items():
#             input_text += f"{question}: {answer}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
#                       "- Conservative Investor\n" \
#                       "- Moderate Investor\n" \
#                       "- Aggressive Investor\n\n" \
#                       "Please provide the classification below:\n"

#         # Use your generative AI model to generate a response
#         model = GenerativeModel('gemini-1.5-flash')
#         response = await model.generate_content(input_text)

#         # Determine the investment personality from the chatbot's response
#         response_text = response.text.lower()

#         if "conservative investor" in response_text:
#             personality = "Conservative Investor"
#         elif "moderate investor" in response_text:
#             personality = "Moderate Investor"
#         elif "aggressive investor" in response_text:
#             personality = "Aggressive Investor"
#         else:
#             personality = "Unknown"

#         return personality
#     except Exception as e:
#         print(f"Error generating response: {e}")
#         return "Unknown"


# GET Method
async def determine_investment_personality(assessment_data): # proper code 
    try:
        # Prepare input text for the chatbot based on assessment data
        input_text = "User Profile:\n"
        for question, answer in assessment_data.items():
            input_text += f"{question}: {answer}\n"

        # Introduce the chatbot's task and prompt for classification
        input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
                      "- Conservative Investor\n" \
                      "- Moderate Investor\n" \
                      "- Aggressive Investor\n\n" \
                      "Please provide the classification below:\n"

        # Use your generative AI model to generate a response
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(input_text)

        # Determine the investment personality from the chatbot's response
        response_text = response.text.lower()

        if "conservative investor" in response_text:
            personality = "Conservative Investor"
        elif "moderate investor" in response_text:
            personality = "Moderate Investor"
        elif "aggressive investor" in response_text:
            personality = "Aggressive Investor"
        else:
            personality = "Unknown"

        return personality
    except Exception as e:
        print(f"Error generating response: {e}")
        return "Unknown"




#Load the Vector DataBase : # current version :
async def load_vector_db(file_path): # # GET Method 
    try:
        print("Loading vector database...")
        # file_path = os.path.basename(file_path)
        
        # Verify the file path
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        print(f"File path: {file_path}")
        
        # Check file permissions
        if not os.access(file_path, os.R_OK):
            raise PermissionError(f"File is not readable: {file_path}")
        
        # print(file_path)
        
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
        vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

        # vector_store = FAISS(
        #     embedding_function=embeddings,
        #     index=index,
        #     docstore=InMemoryDocstore(),
        #     index_to_docstore_id={},
        # )
        
        print("Vector database loaded successfully.") 
        return vector_store.as_retriever(search_kwargs={"k": 1})
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None

# import os

# async def load_vector_db(file_storage): 
#     try:
#         # Define the destination folder and ensure it exists
#         destination_folder = 'path/to/your/destination/folder'
#         if not os.path.exists(destination_folder):
#             os.makedirs(destination_folder)
        
#         # Construct the destination file path
#         file_path = os.path.join(destination_folder, file_storage.filename)
        
#         # Save the file to the destination folder
#         file_storage.save(file_path)
        
#         print("Loading vector database...")
#         print(f"File path: {file_path}")
        
#         loader = Docx2txtLoader(file_path)
#         documents = loader.load()
        
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         text_chunks = text_splitter.split_documents(documents)
        
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
#         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        
#         print("Vector database loaded successfully.") 
#         return vector_store.as_retriever(search_kwargs={"k": 1})
#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None



# investment_personality = "Moderate Investor"
async def make_retrieval_chain(retriever,investmentPersonality,clientName,monthly_investment=10000,investment_period=3): # GET Method
    """
    Create a retrieval chain using the provided retriever.

    Args:
        retriever (RetrievalQA): A retriever object.

    Returns:
        RetrievalQA: A retrieval chain object.
    """
    try:
        # global investment_personality #,summary
        
        print(f"{retriever}\n {investmentPersonality}\n {clientName}\n {monthly_investment}")
        # try:
        #     print(type(investmentPersonality))
        # except Exception as e:
        #     print(f"Error in personality: {e}")
        #     return None
        
        # print(clientName)
        
        llm = ChatGoogleGenerativeAI(
            #model="gemini-pro",
            model = "gemini-1.5-flash",
            temperature = 0.45,
            # temperature=0.7,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )
        # New Template 
        investment_period = str(investment_period)
        print(investmentPersonality)
        monthly_investment = str(monthly_investment)
        print(monthly_investment)
        print(investment_period)
        
        # New Prompt Template :
        
        prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
                Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
                Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
                Also give the user detailed information about the investment how to invest,where to invest and how much they
                should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
                Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
                investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
                Also give the user minimum and maximum expected growth in dollars for the time horizon .
                Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
                
                You are a Financial Advisor for question-answering tasks related to the document. Based on the client's investment personality and financial details provided, generate responsible investment suggestions to achieve their financial goals while managing debts.

                Step-by-Step Guidance:
                1. Assets: Calculate total assets by analyzing the provided financial document in the My Assets section. Ensure you include cash, real estate, retirement accounts, brokerage accounts, and any other relevant asset types from the document.
                2. Liabilities: Calculate total liabilities by analyzing the provided financial document in the My Liabilities section. Consider mortgages, credit card debts, student loans, car loans, and other liabilities. 
                3. Monthly Investment Feasibility: Use the client's assets and liabilities to assess whether their planned monthly investment is feasible. If not feasible, suggest a more realistic monthly investment amount.
                4. Analyze Liabilities: Determine if the client's monthly investment plan is feasible after covering liabilities and expected expenses and also considering some amount for savings. If the client's monthly investment plan is not feasible after covering expenses and savings, generate investment suggestions on a smaller monthly investment plan amount if it can help the client else mention amount is too small for the client's requirementys to be made.
                5. Investment Strategy: Suggest a strategy where monthly investments can both generate returns and pay off debts effectively and helps client to achieve their financial goals.
                6. Allocation: Provide detailed allocations between growth-oriented investments and conservative investments, ensuring the client can meet their monthly debt obligations and save for their future financial goals.
                7. Returns: Include minimum and maximum compounded returns over 5-10 years, along with inflation-adjusted returns for clarity.
                8. Suggestions: Offer advice on how to use remaining funds to build wealth after clearing liabilities and achive their financial goal.
                
                
                Here's an example for the required Output Format(if there are comments indicated by # in the example output format then thats a side note for your reference dont write it in the response that will be generated ) :
                
                Client's Financial Information :(# This is a header line have it in bold) 
                
                
                Client Name: """ + clientName + """(# have the client name in underline)

                Financial Overview: (#the data presented is just an example for your reference do not consider it as factual refere to the document provided to you and generate data based on the provided data and only when nothing is provided assume some data for analysis, This is a header line have it in bold. The data below it should be displayed in a table format so make sure of that data.There must be 2 columns 1 for Category and second for Value.List down all the assets and liabilities along with its values and then Total of assets,liabilities,etc.)
                
                - Total Assets: (# Sum of all client assets and Annual Income . Mention all assets and their respected values.if non consider the example assets)
                
                - Total Liabilities: (# Sum of all liabilities. Mention all liabilities and their respected values if non consider the example liabilities)
                
                
                - Monthly Liabilities: (# Monthly payments derived from liabilities)
                
                - Total Annual Income : (# Sum of all client's anual income)
                
                - Monthly Investment Amount : """ + monthly_investment + """ (# if no specific amount is specified to you then only assume  10,000 else consider the amount mention to you and just display the amount)
                
                - Investment Period : """ + investment_period + """  (# if no specific period is specified to you then only assume 3 years else consider the period mention to you and just display the period)


                Financial Analysis :(#Analyse the assets and liabilities and based on that give a suggestion for analysis generate suggestions for one of the following conditions:)
                (#1st condition : Everything is Positive)Based on the given Financial Conditions the client is having a good and stable income with great assets and manageable debt and liabilities.
                Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the clients monthly income.
                (# if this condition is true then ignore the other conditions and start with the Investment Suggestions)
                
                (#2nd condition : Everything is temporarily Negative) Based on the given Financial Conditions the client is facing a low income for now but have great assets and manageable debt and liabilities.
                Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the client's monthly income but the client might not be able to sustain the monthly investment amount that they are planning.)
                Instead I would like to recommend this amount to the client for their monthly investment : (#Mention a feasible amount to the client for monthly investment and start suggesting investments based on this amount and not the previous amount being taken into consideration)
                
                (#3rd condition : Everything is Negative) Based on the given Financial Conditions the client is facing a low income and doesnt have good assets to manage the debts and liabilities of the client and in such a condition this monthly investment amount is not feasible.
                Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is not manageable for the client's monthly income and so the client might not be able to sustain the monthly investment amount that they are planning to do.)
                I would like to recommend this amount to the client for monthly investment : (# Mention a minimum amount to the client for monthly investment if possible else just say the client should first prioritize on savings and generating more income to manage their debts and liabilities first and so dont give any investment suggestions to the client.)
                
                (#If the financial is 1 or 2 only then give investment suggestions to the client)
                
                
                Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

                Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

                Investment Allocation: (#remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

                Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
                How to Invest: Diversify across various asset classes like:  (#Give allocations % as well)
                
                Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                
                ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                
                Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                
                Consider investing in blue-chip companies or growth sectors like technology. 
                
                
                Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


                Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
                Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

                Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
                How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
                Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
                Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
                Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
                Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
                Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
                Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.


                Time Horizon and Expected Returns:

                Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
                Minimum Expected Annual Return: 4% - 6% 
                
                
                Maximum Expected Annual Return: 8% - 10% 
                
                
                Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, (# consider the monthly investment amount and give returns based on that only) $10,000 could grow to approximately 17,908 in 10 years.
                Minimum Expected Growth in Dollars: 
                
                4,000−6,000 (over 10 years) 
                
                
                Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
                Inflation Adjusted Returns:(#do not write this part inside the bracket just give answer,assume US inflation rate assume 3% if you dont know, and give the investment returns value that was suggested by you for the considered monthly investment amount after 3,5,10years of growth mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
                Rationale for Investment Suggestions:

                This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
                Important Considerations:

                Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

                Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

                Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
                <context>
                {context}
                </context>
                Question: {input}"""
    
        print("Retriever Created ")
        print(f"Investment Personality :{investmentPersonality}")
        
        # Good but inaccurate in terms of assets and liability data calculation
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
        #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
        #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
        #         Also give the user detailed information about the investment how to invest,where to invest and how much they
        #         should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
        #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
        #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
        #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
        #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
                
        #         You are a Financial Advisor for question-answering tasks related to the document. Based on the client's investment personality and financial details provided, generate responsible investment suggestions to achieve their financial goals while managing debts.

        #         Step-by-Step Guidance:
        #         1. Analyze Liabilities: Determine if the client's monthly investment plan is feasible after covering liabilities and expected expenses and also considering some amount for savings. If the client's monthly investment plan is not feasible after covering expenses and savings, generate investment suggestions on a smaller monthly investment plan amount if it can help the client else mention amount is too small for the client's requirementys to be made.
        #         2. Investment Strategy: Suggest a strategy where monthly investments can both generate returns and pay off debts effectively and helps client to achieve their financial goals.
        #         3. Allocation: Provide detailed allocations between growth-oriented investments and conservative investments, ensuring the client can meet their monthly debt obligations and save for their future financial goals.
        #         4. Returns: Include minimum and maximum compounded returns over 5-10 years, along with inflation-adjusted returns for clarity.
        #         5. Suggestions: Offer advice on how to use remaining funds to build wealth after clearing liabilities.
                
                
        #         Here's an example for the required Output Format(if there are comments indicated by # in the example output format then thats a side note for your reference dont write it in the response that will be generated ) :
                
        #         Investment Suggestions : 
                
                
        #         Client Name: """ + clientName + """

        #         Financial Overview: (#the data presented is just an example for your reference do not consider it as factual refere to the document provided to you and generate data based on the provided data only)
                
        #         - Total Assets: (# Consider the data available to you and use it  for display. For ex : $100,000 (cash), $150,000 (home), $12,000 (other assets) )
                
        #         - Liabilities: (# Consider the data available to you and use it  for display. For ex : $200,000 mortgage at 12% interest, $400 credit card debt at 3.5%, $15,000 other loans at 10%.)
                
        #         - Total Monthly Liabilities: (# Consider the data available to you and use it  for display. For ex : $1,650 (Mortgage, Credit Card, Other Loans).)
                
        #         - Monthly Investment Amount : """ + monthly_investment + """ (# if no specific amount is specified to you then only assume  10,000 else consider the amount mention to you and just display the amount)
                
        #         - Investment Period : """ + investment_period + """  (# if no specific period is specified to you then only assume 3 years else consider the period mention to you and just display the period)

        #         Financial Analysis :(#Analyse the assets and liabilities and based on that give a suggestion for analysis generate suggestions for one of the following conditions:)
        #         (#1st condition : Everything is Positive)Based on the given Financial Conditions the client is having a good and stable income with great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the clients monthly income.
        #         (# if this condition is true then ignore the other conditions and start with the Investment Suggestions)
                
        #         (#2nd condition : Everything is temporarily Negative) Based on the given Financial Conditions the client is facing a bad income for now but have great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the client's monthly income but the client might not be able to sustain the monthly investment amount that they are planning.)
        #         Instead I would like to recommend this amount to the client for their monthly investment : (#Mention a feasible amount to the client for monthly investment and start suggesting investments based on this amount and not the previous amount being taken into consideration)
                
        #         (#3rd condition : Everything is Negative) Based on the given Financial Conditions the client is facing a bad income and doesnt have good assets to manage the debts and liabilities of the client and in such a condition this monthly investment amount is not feasible.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is not manageable for the client's monthly income and so the client might not be able to sustain the monthly investment amount that they are planning to do.)
        #         I would like to recommend this amount to the client for monthly investment : (# Mention a minimum amount to the client for monthly investment if possible else just say the client should first prioritize on savings and generating more income to manage their debts and liabilities first and so dont give any investment suggestions to the client.)
                
        #         (#If the financial is 1 or 2 only then give investment suggestions to the client)
                
                
        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

        #         Investment Allocation: (#remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (#Give allocations % as well)
                
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                
        #         Consider investing in blue-chip companies or growth sectors like technology. 
                
                
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.


        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, (# consider the monthly investment amount and give returns based on that only) $10,000 could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(#do not write this part inside the bracket just give answer,assume US inflation rate assume 3% if you dont know, and give the investment returns value that was suggested by you for the considered monthly investment amount after 3,5,10years of growth mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""
        
        
                
        # # latest version gives some info about the client conditions but not in detail 
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
        #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
        #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
        #         Also give the user detailed information about the investment how to invest,where to invest and how much they
        #         should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
        #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
        #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
        #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
        #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
        #         Here's an example for the required Output Format :

        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

        #         Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
        #         Consider investing in blue-chip companies or growth sectors like technology. 
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, a 10,000 investment could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(do not write this part inside the bracket just give answer,assume US inflation rate, and give the investment returns value that was suggested by you  for $10k investment after 3,5,10years of growth  mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""
        
        
        
        
        # # Working code but gives clientname in brackets 
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
        #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality provided to you.
        #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
        #         Also give the user detailed information about the investment how to invest,where to invest and how much they
        #         should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
        #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
        #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
        #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
        #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
        #         Here's an example for the required Output Format :

        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

        #         Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
        #         Consider investing in blue-chip companies or growth sectors like technology. 
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, a 10,000 investment could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(do not write this part inside the bracket just give answer,assume US inflation rate, and give the investment returns value that was suggested by you  for $10k investment after 3,5,10years of growth  mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""
                

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            # print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in creating chain: {e}")
        return None


import json
import io


async def process_document(file_path): # GET Method
    try:
        print("Processing the document")
        file_type = filetype.guess(file_path)
        if file_type is not None:
            if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Await the coroutine to extract text and tables
                return await extract_text_and_tables_from_word(file_path)
            elif file_type.mime == "application/pdf":
                return await extract_text_from_pdf(file_path)
        return None
    except Exception as e:
        print(f"Error processing document: {e}")
        return None

# Async function to extract text from a PDF file
async def extract_text_from_pdf(pdf_file_path): # GET Method
    try:
        print("Processing pdf file")
        with open(pdf_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            text_content = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text_content.append(page.extract_text())
            return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

# Async function to extract text and tables from a Word document
async def extract_text_and_tables_from_word(docx_file_path): # GET Method
    try:
        print("Extracting text and tables from word file")
        doc = docx.Document(docx_file_path)
        text_content = []
        tables_content = []

        for para in doc.paragraphs:
            text_content.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)
        print("Extracted text from word file")
        return "\n".join(text_content), tables_content
    except Exception as e:
        print(f"Error extracting text and tables from Word document: {e}")
        return None, None



async def validate_document_content(text, tables):
    """
    Validates the content of the document.

    Args:
        text (str): Extracted text content from the document.
        tables (list): Extracted tables content from the document.

    Returns:
        tuple: Client name and validation errors.
    """
    errors = []
    
    # Extract client name
    client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
    client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

    # Define required sections
    required_sections = [
        "YOUR RETIREMENT GOAL",
        "YOUR OTHER MAJOR GOALS",
        "YOUR ASSETS AND LIABILITIES",
        "MY LIABILITIES",
        "YOUR CURRENT ANNUAL INCOME"
    ]

    # Check for the presence of required sections
    for section in required_sections:
        if section not in text:
            errors.append(f"* {section} section missing.")
    
    # Define table field checks
    table_checks = {
        "YOUR RETIREMENT GOAL": [
            r"When do you plan to retire\? \(age or date\)",
            r"Social Security Benefit \(include expected start date\)",
            r"Pension Benefit \(include expected start date\)",
            r"Other Expected Income \(rental, part-time work, etc.\)",
            r"Estimated Annual Retirement Expense"
        ],
        "YOUR OTHER MAJOR GOALS": [
            r"GOAL", r"COST", r"WHEN"
        ],
        "YOUR ASSETS AND LIABILITIES": [
            r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
            r"Current Value", r"Annual Contributions"
        ],
        "MY LIABILITIES": [
            r"Balance", r"Interest Rate", r"Monthly Payment"
        ]
    }

    # Validate table content
    for section, checks in table_checks.items():
        section_found = False
        for table in tables:
            table_text = "\n".join(["\t".join(row) for row in table])
            if section in table_text:
                section_found = True
                for check in checks:
                    if not re.search(check, table_text, re.IGNORECASE):
                        errors.append(f"* Missing or empty field in {section} section: {check}")
                break
        if not section_found:
            errors.append(f"* {section} section missing.")

    return client_name, errors


# RUN Button :
# async def generate_investment_suggestions(investment_personality, context): # # GET Method for py , for front end its Post API
    
#     # retriever = asyncio.run(load_vector_db("uploaded_file"))

#     retriever = await load_vector_db("uploaded_file")
#     # retriever = await load_vector_db("data\Financial_Investment_1.docx") 

#     chain = await make_retrieval_chain(retriever)

#     # chain = asyncio.run(make_retrieval_chain(retriever))
    
#     if chain is not None:
#         # summary = context
#         # query = summary + "\n" + investment_personality
#         query = str(investment_personality)
#         response = chain.invoke({"input": query})
#         format_response = markdown_to_text(response['answer'])
#         return format_response
#         # st.write(format_response)

#         # handle_graph(response['answer'])

#     else:
#         st.error("Failed to create the retrieval chain. Please upload a valid document.")



# Generate Infographics : Best Code so far:

import re
from collections import defaultdict
import numpy as np



def extract_numerical_data(response): # curr version but cant capture annual return 
    # Define patterns to match different sections and their respective allocations
    patterns = {
        'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
        'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
        'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
    }

    data = defaultdict(dict)

    for section, pattern in patterns.items():
        match = pattern.search(response)
        if match:
            investments_text = match.group(1)
            # Extract individual investment types and their allocations
            investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
            for investment_match in investment_pattern.findall(investments_text):
                investment_type, min_allocation, max_allocation = investment_match
                data[section][investment_type.strip()] = {
                    'min': min_allocation,
                    'max': max_allocation
                }

    # Extract time horizon and expected returns
    time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
    min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
    max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
    min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
    max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

    time_horizon_match = time_horizon_pattern.search(response)
    min_return_match = min_return_pattern.search(response)
    max_return_match = max_return_pattern.search(response)
    min_growth_match = min_growth_pattern.search(response)
    max_growth_match = max_growth_pattern.search(response)

    if time_horizon_match:
        data['Time Horizon'] = {
            'min_years': time_horizon_match.group(1),
            'max_years': time_horizon_match.group(2)
        }

    if min_return_match:
        data['Expected Annual Return'] = {
            'min': min_return_match.group(1),
            'max': min_return_match.group(2)
        }

    if max_return_match:
        data['Expected Annual Return'] = {
            'min': max_return_match.group(1),
            'max': max_return_match.group(2)
        }

    if min_growth_match:
        data['Expected Growth in Dollars'] = {
            'min': min_growth_match.group(1),
            'max': min_growth_match.group(2)
        }

    if max_growth_match:
        data['Expected Growth in Dollars'] = {
            'min': max_growth_match.group(1),
            'max': max_growth_match.group(2)
        }

    return data

def normalize_allocations(allocations):
    total = sum(allocations)
    if total == 100:
        return allocations
    return [round((allocation / total) * 100, 2) for allocation in allocations]


import datetime  # Import the datetime module to get the current year
# uodated to have current year
def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
    try:
        # Get the current year
        curr_year = datetime.datetime.now().year

        # Print data_extracted to debug the structure
        print("Data extracted:", data_extracted)

        # Check if 'Expected Annual Return' and 'Time Horizon' exist and have the expected keys
        if 'Expected Annual Return' not in data_extracted:
            print("'Expected Annual Return' missing in data_extracted")
            data_extracted['Expected Annual Return']['min'] = 6
            data_extracted['Expected Annual Return']['max'] = 8
            min_return = 6
            max_return = 8
        else:
            min_return = float(data_extracted['Expected Annual Return'].get('min', '0').strip('%'))
            max_return = float(data_extracted['Expected Annual Return'].get('max', '0').strip('%'))

        min_years = int(data_extracted['Time Horizon'].get('min_years', 1))  # Default to 1 year if missing
        max_years = int(data_extracted['Time Horizon'].get('max_years', 10))  # Default to 10 years if missing

        def calculate_compounded_return(principal, rate, years):
            return principal * (1 + rate / 100) ** years

        def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
            return nominal_return / (1 + inflation_rate / 100) ** years

        # Create labels for the next 10 years starting from the current year
        labels = list(range(curr_year, curr_year + max_years))

        min_compounded = []
        max_compounded = []
        min_inflation_adjusted = []
        max_inflation_adjusted = []

        for year in range(1, max_years + 1):
            # Calculate nominal compounded returns
            min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
            max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

            # Calculate inflation-adjusted compounded returns
            min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
            max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

            # Append results
            min_compounded.append(min_compounded_value)
            max_compounded.append(max_compounded_value)
            min_inflation_adjusted.append(min_inflation_value)
            max_inflation_adjusted.append(max_inflation_value)

        # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
        combined_chart_data = {
            'labels': labels,  # Current year and the next 10 years
            'datasets': [
                {
                    'label': 'Minimum Compounded Return',
                    'data': min_compounded,
                    'borderColor': 'rgb(255, 99, 132)',  # Red color
                    'fill': False
                },
                {
                    'label': 'Maximum Compounded Return',
                    'data': max_compounded,
                    'borderColor': 'rgb(54, 162, 235)',  # Blue color
                    'fill': False
                },
                {
                    'label': 'Min Inflation Adjusted Return',
                    'data': min_inflation_adjusted,
                    'borderColor': 'rgb(75, 192, 192)',  # Light blue
                    'borderDash': [5, 5],  # Dashed line for distinction
                    'fill': False
                },
                {
                    'label': 'Max Inflation Adjusted Return',
                    'data': max_inflation_adjusted,
                    'borderColor': 'rgb(153, 102, 255)',  # Light purple
                    'borderDash': [5, 5],  # Dashed line for distinction
                    'fill': False
                }
            ]
        }
    except KeyError as e:
        print(f"KeyError occurred: {e}")
        return jsonify({'message': f'Key Error: {e}'}), 400
    except Exception as e:
        print(f"Error occurred while preparing data for combined line chart: {e}")
        return jsonify({'message': 'Internal Server Error in creating line chart'}), 500

    return combined_chart_data



# def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
#     try:
#         min_return = float(data_extracted['Expected Annual Return']['min'].strip('%'))
#         max_return = float(data_extracted['Expected Annual Return']['max'].strip('%'))
#         min_years = int(data_extracted['Time Horizon']['min_years'])
#         max_years = int(data_extracted['Time Horizon']['max_years'])

#         def calculate_compounded_return(principal, rate, years):
#             return principal * (1 + rate / 100) ** years

#         def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
#             return nominal_return / (1 + inflation_rate / 100) ** years

#         labels = list(range(1, max_years + 1))  # Years for the x-axis
#         min_compounded = []
#         max_compounded = []
#         min_inflation_adjusted = []
#         max_inflation_adjusted = []

#         for year in labels:
#             # Calculate nominal compounded returns
#             min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
#             max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

#             # Calculate inflation-adjusted compounded returns
#             min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
#             max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

#             # Append results
#             min_compounded.append(min_compounded_value)
#             max_compounded.append(max_compounded_value)
#             min_inflation_adjusted.append(min_inflation_value)
#             max_inflation_adjusted.append(max_inflation_value)

#         # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
#         combined_chart_data = {
#             'labels': labels,
#             'datasets': [
#                 {
#                     'label': 'Minimum Compounded Return',
#                     'data': min_compounded,
#                     'borderColor': 'rgb(255, 99, 132)',  # Red color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Maximum Compounded Return',
#                     'data': max_compounded,
#                     'borderColor': 'rgb(54, 162, 235)',  # Blue color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Min Inflation Adjusted Return',
#                     'data': min_inflation_adjusted,
#                     'borderColor': 'rgb(75, 192, 192)',  # Light blue
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 },
#                 {
#                     'label': 'Max Inflation Adjusted Return',
#                     'data': max_inflation_adjusted,
#                     'borderColor': 'rgb(153, 102, 255)',  # Light purple
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 }
#             ]
#         }
#     except Exception as e:
#         print(f"Error occurred while preparing data for combined line chart: {e}")
#         return jsonify({'message': 'Internal Server Error in creating line chart'}), 500
#     return combined_chart_data



def plot_investment_allocations(data):
    # Create subplots with a large figure size
    fig, axes = plt.subplots(2, 1, figsize= (16,10)) #(28, 15))  # Adjust size as needed

    # Plot Growth-Oriented Investments
    growth_data = data['Growth-Oriented Investments']
    growth_labels = list(growth_data.keys())
    growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
    growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

    axes[0].bar(growth_labels, growth_min, color='skyblue', label='Min Allocation')
    axes[0].bar(growth_labels, growth_max, bottom=growth_min, color='lightgreen', label='Max Allocation')
    axes[0].set_title('Growth-Oriented Investments', fontsize=16)
    axes[0].set_ylabel('Percentage Allocation', fontsize=14)
    axes[0].set_xlabel('Investment Types', fontsize=14)
    axes[0].tick_params(axis='x', rotation=45, labelsize=12)
    axes[0].tick_params(axis='y', labelsize=12)
    axes[0].legend()

    # Plot Conservative Investments
    conservative_data = data['Conservative Investments']
    conservative_labels = list(conservative_data.keys())
    conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
    conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

    axes[1].bar(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
    axes[1].bar(conservative_labels, conservative_max, bottom=conservative_min, color='lightgreen', label='Max Allocation')
    axes[1].set_title('Conservative Investments', fontsize=16)
    axes[1].set_ylabel('Percentage Allocation', fontsize=14)
    axes[1].set_xlabel('Investment Types', fontsize=14)
    axes[1].tick_params(axis='x', rotation=45, labelsize=12)
    axes[1].tick_params(axis='y', labelsize=12)
    axes[1].legend()

    # Tight layout for better spacing
    plt.tight_layout()
    plt.show()
    return fig


def plot_pie_chart(data):
    fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

    # Combine all investment data for pie chart
    all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
    labels = list(all_data.keys())
    sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
    colors = plt.cm.Paired(range(len(labels)))

    wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    ax.set_title('Investment Allocation')

    # Add legend
    ax.legend(wedges, labels, title="Investment Types", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

    return fig



def bar_chart(data):
    fig, ax = plt.subplots(figsize=(12, 8))  # Increased size

    # Data for plotting
    categories = list(data.keys())
    values_min = [int(data[cat]['min'].strip('%')) for cat in categories]
    values_max = [int(data[cat]['max'].strip('%')) for cat in categories]

    x = range(len(categories))

    ax.bar(x, values_min, width=0.4, label='Min Allocation', color='skyblue', align='center')
    ax.bar(x, values_max, width=0.4, label='Max Allocation', color='lightgreen', align='edge')

    ax.set_xticks(x)
    ax.set_xticklabels(categories, rotation=45, ha='right')
    ax.set_xlabel('Investment Categories')
    ax.set_ylabel('Percentage Allocation')
    ax.set_title('Investment Allocation')
    ax.legend()

    plt.tight_layout()
    return fig


import random

def generate_colors(n):
    """
    Generate 'n' random RGB colors.

    Args:
        n (int): Number of colors to generate.
    
    Returns:
        list: A list of RGB colors in 'rgb(r, g, b)' format.
    """
    colors = []
    for _ in range(n):
        r = random.randint(0, 255)
        g = random.randint(0, 255)
        b = random.randint(0, 255)
        colors.append(f'rgb({r}, {g}, {b})')
    
    return colors


import plotly.graph_objects as go
import numpy as np

 
# def client_form():
#     st.title("Client Details Form")

#     with st.form("client_form"):
#         st.header("Personal Information")
#         client_name = st.text_input("Client Name")
#         co_client_name = st.text_input("Co-Client Name")
#         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
#         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
#         today_date = st.date_input("Today's Date")
        
#         st.header("Financial Information")
#         current_assets = st.text_area("Current Assets (e.g., type and value)")
#         liabilities = st.text_area("Liabilities (e.g., type and amount)")
#         annual_income = st.text_area("Current Annual Income (source and amount)")
#         annual_contributions = st.text_area("Annual Contributions (e.g., retirement savings)")

#         st.header("Insurance Information")
#         life_insurance = st.text_input("Life Insurance (e.g., coverage amount)")
#         disability_insurance = st.text_input("Disability Insurance (e.g., coverage amount)")
#         long_term_care = st.text_input("Long-Term Care Insurance (e.g., coverage amount)")

#         st.header("Estate Planning")
#         will_status = st.radio("Do you have a will?", ["Yes", "No"])
#         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
#         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
#         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

#         # Submit button
#         submitted = st.form_submit_button("Submit")

#         if submitted:
#             # Save form data
#             form_data = {
#                 "Client Name": client_name,
#                 "Co-Client Name": co_client_name,
#                 "Client Age": client_age,
#                 "Co-Client Age": co_client_age,
#                 "Today's Date": str(today_date),
#                 "Current Assets": current_assets,
#                 "Liabilities": liabilities,
#                 "Annual Income": annual_income,
#                 "Annual Contributions": annual_contributions,
#                 "Life Insurance": life_insurance,
#                 "Disability Insurance": disability_insurance,
#                 "Long-Term Care Insurance": long_term_care,
#                 "Will Status": will_status,
#                 "Trust Status": trust_status,
#                 "Power of Attorney": power_of_attorney,
#                 "Healthcare Proxy": healthcare_proxy,
#             }
            
#             # Save to a file or database
#             with open("client_data.txt", "a") as f:
#                 f.write(str(form_data) + "\n")
            
#             st.success("Form submitted successfully!")
#             st.session_state.page = "main"  # Redirect back to main page after form submission


from datetime import date  # Make sure to import the date class


# Function to parse financial data from the text
import re

def parse_financial_data(text_content):
    assets = []
    liabilities = []

    # Define regex patterns to capture text following headings
    asset_pattern = re.compile(r"MY ASSETS:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)
    liability_pattern = re.compile(r"LIABILITIES:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)

    # Extract assets
    asset_matches = asset_pattern.findall(text_content)
    if asset_matches:
        asset_text = asset_matches[0]
        # Further processing to extract individual asset values if they are detailed
        asset_lines = asset_text.split('\n')
        for line in asset_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                asset_value = float(match.group().replace(",", ""))
                assets.append(asset_value)

    # Extract liabilities
    liability_matches = liability_pattern.findall(text_content)
    if liability_matches:
        liability_text = liability_matches[0]
        # Further processing to extract individual liability values if they are detailed
        liability_lines = liability_text.split('\n')
        for line in liability_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                liability_value = float(match.group().replace(",", ""))
                liabilities.append(liability_value)

    print("Assets Found:", assets)
    print("Liabilities Found:", liabilities)

    return assets, liabilities



# Function to extract numerical values from a text input
def extract_numeric(value):
    try:
        return float(re.sub(r'[^\d.]', '', value))  # Remove non-numeric characters and convert to float
    except ValueError:
        return 0


# plots graph from the details of the form :


def is_float(value):
    try:
        float(value)
        return True
    except ValueError:
        return False



def plot_assets_liabilities_pie_chart(assets, liabilities, threshold=50): # best plot 
    """
    Plots separate pie charts for assets and liabilities. If there are any categories
    below a specified threshold, they are plotted in an additional small pie chart.
    
    Parameters:
    - assets: dict, keys are asset names, values are their amounts.
    - liabilities: dict, keys are liability names, values are their amounts.
    - threshold: int, percentage threshold below which segments are considered small.
    """
    # Update matplotlib settings to increase the font size globally
    # plt.rcParams.update({'font.size': 32})

    plt.rcParams.update({'font.size': 16})

    def plot_pie(data, title):
        # Filter out zero values and create a summary for small segments
        total = sum(data.values())
        filtered_data = {k: v for k, v in data.items() if (v / total) >= threshold / 100}
        small_segments = {k: v for k, v in data.items() if (v / total) < threshold / 100}
        small_total = sum(small_segments.values())

        # Plotting logic
        if small_segments:
            fig, (ax_main, ax_small) = plt.subplots(1, 2, figsize=(30, 15))  # Side-by-side layout
        else:
            fig, ax_main = plt.subplots(figsize=(30, 20))  # Only main chart with larger size

            # fig, ax_main = plt.subplots(figsize=(10, 10))  # Only main chart with larger size

        # Plot main pie chart
        labels_main = list(filtered_data.keys()) + ([f"Other small {title}"] if small_segments else [])
        values_main = list(filtered_data.values()) + ([small_total] if small_segments else [])
        wedges_main, texts_main, autotexts_main = ax_main.pie(
            values_main, labels=labels_main, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
            startangle=140, textprops={'fontsize': 28} #18}  # Larger font size for labels
        )

        ax_main.set_title(title, fontsize=20)
        # Position legend to the right of the plot to avoid overlapping
        ax_main.legend(wedges_main, labels_main, title="Categories", loc="upper right", bbox_to_anchor=(0.001, 0.9), fontsize= 28)#14)

        if small_segments:
            # Plot additional small pie chart for small segments
            labels_small = list(small_segments.keys())
            values_small = list(small_segments.values())
            wedges_small, texts_small, autotexts_small = ax_small.pie(
                values_small, labels=labels_small, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
                startangle=140, textprops={'fontsize': 24} #14}  # Consistent label size for small chart
            )
            ax_small.set_title(f"Small Segments of {title}", fontsize=20)
            # Position legend to the right of the small pie chart but slightly lower to avoid overlap with the main chart's legend
            ax_small.legend(wedges_small, labels_small, title="Small Categories", loc="center left", bbox_to_anchor=(1.2, 0.3), fontsize= 22)#12)

        st.pyplot(fig)

    # Convert valid entries to float, ensuring only numeric values are considered
    assets = {k: float(v) for k, v in assets.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}
    liabilities = {k: float(v) for k, v in liabilities.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}

    # Plot pie charts
    plot_pie(assets, 'Distribution of Assets')
    plot_pie(liabilities, 'Distribution of Liabilities')


# def plot_assets_liabilities_pie_chart(assets, liabilities):# properly plots a big and 1 small pie chart for both assets and liability
#     # Filter and convert values to float, handle non-numeric or empty inputs
#     filtered_assets = {k: float(v) for k, v in assets.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}
#     filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}

#     # Combine assets and liabilities for total calculation
#     all_values = {**filtered_assets, **filtered_liabilities}
#     total_value = sum(all_values.values())

#     # Separate main and small segments
#     main_segments = {k: v for k, v in all_values.items() if (v / total_value) >= 0.05}
#     small_segments = {k: v for k, v in all_values.items() if (v / total_value) < 0.05}
#     small_total = sum(small_segments.values())

#     # Prepare data for main pie chart
#     main_labels = list(main_segments.keys()) + (["Others"] if small_segments else [])
#     main_values = list(main_segments.values()) + ([small_total] if small_segments else [])

#     # Prepare data for small pie chart (only if there are small segments)
#     small_labels = list(small_segments.keys())
#     small_values = list(small_segments.values())

#     fig, ax = plt.subplots(figsize=(8, 6))

#     # Plot main pie chart
#     wedges, texts, autotexts = ax.pie(
#         main_values,
#         labels=main_labels,
#         autopct='%1.1f%%',
#         startangle=140,
#         colors=plt.cm.Paired.colors,
#     )

#     # Explode the "Others" slice
#     if small_segments:
#         others_index = main_labels.index("Others")
#         wedges[others_index].set_edgecolor('white')
#         # wedges[others_index].set_linestyle('--')
#         wedges[others_index].set_linewidth(2)
#         wedges[others_index].set_hatch('/')

#     ax.set_title('Assets and Liabilities Distribution')

#     # Draw a second pie chart for "Others"
#     if small_segments:
#         fig2, ax2 = plt.subplots(figsize=(8, 6))
#         wedges_small, texts_small, autotexts_small = ax2.pie(
#             small_values,
#             labels=small_labels,
#             autopct='%1.1f%%',
#             startangle=140,
#             colors=plt.cm.Pastel1.colors
#         )

#         ax2.set_title('Detailed View of "Others" Categories')

#     plt.tight_layout()
#     st.pyplot(fig)
#     if small_segments:
#         st.pyplot(fig2)



def save_data_to_file(form_data):
    file_path = 'client_data.txt'
    with open(file_path, 'a') as file:
        file.write(str(form_data) + "\n")
    # st.success(f"Form data saved to {file_path}")
    print(f"Form data saved to {file_path}")
    

def client_form():
    st.title("Client Details Form")

    with st.form("client_form"):
        st.header("Personal Information")
        client_name = st.text_input("Client Name")
        co_client_name = st.text_input("Co-Client Name")
        client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
        co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
        today_date = st.date_input("Today's Date")

        st.header("Your Assets (in $)")

        assets = {
            # 'Annual Income': st.text_input("Annual Income (e.g. , Your Annual Salary Income or other source of income) "),
            'Cash/Bank Account': st.text_input("Cash/Bank Account"),
            '401(k), 403(b), 457 Plans': st.text_input("Your 401(k), 403(b), 457 Plans "),
            'Traditional, SEP and SIMPLE IRAs': st.text_input("Traditional, SEP and SIMPLE IRAs "),
            'Roth IRA,Roth 401(k)': st.text_input("Roth IRA, Roth 401(k)"),
            'Brokerage/non-qualified accounts': st.text_input("Brokerage/non-qualified accounts"),
            'Annuities': st.text_input("Annuities"),
            '529 Plans': st.text_input("529 Plans"),
            'Home': st.text_input("Home"),
            'Other Real Estate': st.text_input("Other Real Estate"),
            'Business': st.text_input("Business"),
            'Other': st.text_input("Other")
        }
        st.header("Your Liabilities (in $)")

        liabilities = {
            'Mortgage': st.text_input("Mortgage"),
            # 'Annual Mortgage Interest Rate': st.number_input("Annual Mortgage Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=12.0, step=0.5),
            # 'Mortagage Time Period': st.number_input("Mortagage Time Period (Mention the time period of the Mortgage in years)", min_value=0, max_value=100,value=10,step=1),

            'Home Loans': st.text_input("Home Loans"),
            # 'Home Loans Interest Rate': st.number_input("Home Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=10.0, step=0.5),
            # 'Home Loans Time Period': st.number_input("Home Loans Time Period (Mention the time period of the Home Loan in years)", min_value=0, max_value=100,value=15,step=1),

            'Vehicle Loans': st.text_input("Vehicle Loans"),
            # 'Vehicle Loans Interest Rate': st.number_input("Vehicle Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
            # 'Vehicle Loans Time Period': st.number_input("Vehicle Loans Time Period (Mention the time period of the Car/Vehicle Loan in years)", min_value=0, max_value=100,value=15,step=1),

            'Education Loans': st.text_input("Education Loans"),
            # 'Education Loans Interest Rate' : st.number_input("Education Loans Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
            # 'Education Loans Time Period': st.number_input("Education Loans Time Period (Mention the time period of the Education Loan in years)", min_value=0, max_value=100,value=15,step=1),

            # 'Credit Card': st.text_input("Monthly Credit Card Debt (Mention Amount)"),
            # 'Credit Card Debt Interest Rate': st.number_input("Credit Card Debt Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),

            'Miscellaneous': st.text_input("Miscellaneous"),
        }

        st.header("Your Retirement Goal")
        retirement_age = st.number_input("At what age do you plan to retire?", min_value=0, max_value=120, value=65, step=1)
        retirement_income = st.text_input("Desired annual retirement income")

        st.header("Your Other Goals")
        goal_name = st.text_input("Name of the Goal (e.g . , Dream House, Travel, Educational, etc.)")
        goal_amount = st.text_input("Amount needed for the goal (in $)")
        goal_timeframe = st.number_input("Timeframe to achieve the goal (in years)", min_value=0, max_value=100, value=5, step=1)

        st.header("Insurance Information")
        life_insurance_Benefit = st.text_input("Life Insurance-Benefit")
        life_insurance_Premium = st.text_input("Life Insurance-Premium")
        disability_insurance_Benefit = st.text_input("Disability Insurance-Benefit")
        disability_insurance_Premium = st.text_input("Disability Insurance-Premium")
        long_term_care_benefit = st.text_input("Long-Term Care Insurance-Benefit")
        long_term_care_premium = st.text_input("Long-Term Care Insurance-Premium")


        st.header("Estate Planning")
        will_status = st.radio("Do you have a will?", ["Yes", "No"])
        trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
        power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
        healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

        submitted = st.form_submit_button("Submit")

        if submitted:
            form_data = {
                "Client Name": client_name,
                "Co-Client Name": co_client_name,
                "Client Age": client_age,
                "Co-Client Age": co_client_age,
                "Today's Date": str(today_date),
                "Assets": assets,
                "Liabilities": liabilities,
                "Retirement Age": retirement_age,
                "Desired Retirement Income": retirement_income,
                "Goal Name": goal_name,
                "Goal Amount": goal_amount,
                "Goal Timeframe": goal_timeframe,
                "Life Insurance Benefit": life_insurance_Benefit,
                "Life Insurance Premium": life_insurance_Premium,
                "Disability Insurance Benefit": disability_insurance_Benefit,
                "Disability Insurance Premium": disability_insurance_Premium,
                "Long-Term Care Insurance Benefit": long_term_care_benefit,
                "Long-Term Care Insurance Premium": long_term_care_premium,
                "Will Status": will_status,
                "Trust Status": trust_status,
                "Power of Attorney": power_of_attorney,
                "Healthcare Proxy": healthcare_proxy,
            }

            save_data_to_file(form_data)
            
            # # Plot the pie chart
            # st.subheader("Assets and Liabilities Breakdown")
            # plot_assets_liabilities_pie_chart(assets, liabilities)

            # Store data in session state and redirect to main
            st.session_state.assets = assets
            st.session_state.liabilities = liabilities
            st.session_state.total_assets, st.session_state.total_liabilities = calculate_totals(assets, liabilities)
            st.session_state.page = "main"
            st.success("Data submitted!\nThank You for filling the form !\nReturning to main portal...")

import math
def calculate_compounded_amount(principal, rate, time):
    """
    Calculates the compounded amount using the formula:
    A = P * (1 + r/n)^(nt)
    Assuming n (compounding frequency) is 1 for simplicity (annually).
    """
    if principal == 0 or rate == 0 or time == 0:
        return principal
    else:
        # Using annual compounding
        return principal * (1 + rate / 100) ** time
    
def calculate_totals(assets, liabilities):
    total_assets = sum(extract_numeric(v) for v in assets.values())
    print(f"Total Assets : {total_assets}")
    total_liabilities = 0
    total_liabilities = sum(extract_numeric(v) for v in liabilities.values() )

    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Mortgage']),
    #     liabilities['Annual Mortgage Interest Rate'],
    #     liabilities['Mortagage Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Home Loans']),
    #     liabilities['Home Loans Interest Rate'],
    #     liabilities['Home Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Vehicle Loans']),
    #     liabilities['Vehicle Loans Interest Rate'],
    #     liabilities['Vehicle Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Education Loans']),
    #     liabilities['Education Loans Interest Rate'],
    #     liabilities['Education Loans Time Period']
    # )
    
    # For credit card debt, only calculate compounded amount if interest rate > 0

    # credit_card_balance = extract_numeric(liabilities['Credit Card'])
    # credit_card_interest = liabilities['Credit Card Debt Interest Rate']
    # if credit_card_interest > 0:
    #     # Assuming the time period for credit card debt is 1 year for compounding
    #     total_liabilities += calculate_compounded_amount(credit_card_balance, credit_card_interest, 1)
    # else:
    #     total_liabilities += credit_card_balance
    
    # Miscellaneous debts are taken directly as is
    total_liabilities += extract_numeric(liabilities['Miscellaneous'])
    rounded_liabilities = round(total_liabilities,2)

    print(f"Total liabilities :{total_liabilities}")
    print(f"Rounded of Total liabilities :{rounded_liabilities}")

    return total_assets, rounded_liabilities #total_liabilities

def create_financial_summary_table(assets, liabilities):
    # Filter out items with zero value
    filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
    filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

    # Create DataFrames for assets and liabilities with indices starting from 1
    assets_df = pd.DataFrame(
        list(filtered_assets.items()), 
        columns=['Assets', 'Amount ($)'], 
        index=range(1, len(filtered_assets) + 1)
    )
    liabilities_df = pd.DataFrame(
        list(filtered_liabilities.items()), 
        columns=['Liabilities', 'Amount ($)'], 
        index=range(1, len(filtered_liabilities) + 1)
    )

    # Calculate total
    total_assets, total_liabilities = calculate_totals(assets, liabilities)

    # Add total row with index incremented by 1
    total_assets_row = pd.DataFrame(
        [['TOTAL', total_assets]], 
        columns=['Assets', 'Amount ($)'], 
        index=[len(assets_df) + 1]
    )
    total_liabilities_row = pd.DataFrame(
        [['TOTAL', total_liabilities]], 
        columns=['Liabilities', 'Amount ($)'], 
        index=[len(liabilities_df) + 1]
    )

    # Append total rows to DataFrames
    assets_df = pd.concat([assets_df, total_assets_row])
    liabilities_df = pd.concat([liabilities_df, total_liabilities_row])

    # Display tables with formatted values
    st.subheader("Assets")
    st.table(assets_df.style.format({'Amount ($)': '{:,.2f}'}))

    st.subheader("Liabilities")
    st.table(liabilities_df.style.format({'Amount ($)': '{:,.2f}'}))


def plot_bar_graphs(assets, liabilities):
    # Filter out items with zero values
    filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
    filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

    # Calculate compounded liabilities
    # compounded_liabilities = {} 

    # for k, v in filtered_liabilities.items():
        # if 'Interest Rate' in k or 'Time Period' in k:
        #     continue  # Skip non-monetary entries

        # if k == 'Credit Card Payment' and liabilities['Credit Card Debt Interest Rate'] == 0.0:
        #     continue  # Skip if credit card interest rate is zero

        # if k == 'Mortgage':
        #     interest_rate = liabilities['Annual Mortgage Interest Rate']
        #     time_period = liabilities['Mortagage Time Period']

        # elif k == 'Home Loans':
        #     interest_rate = liabilities['Home Loans Interest Rate']
        #     time_period = liabilities['Home Loans Time Period']

        # elif k == 'Car/Vehicle Loans':
        #     interest_rate = liabilities['Car/Vehicle Loans Interest Rate']
        #     time_period = liabilities['Car/Vehicle Loans Time Period']

        # elif k == 'Education Loans':
        #     interest_rate = liabilities['Education Loans Interest Rate']
        #     time_period = liabilities['Education Loans Time Period']

        # elif k == 'Credit Card Payment':
        #     interest_rate = liabilities['Credit Card Debt Interest Rate']
        #     time_period = 1  # Assuming interest is calculated yearly

        # if interest_rate > 0:
        #     compounded_amount = float(v) * (1 + float(interest_rate) / 100) ** float(time_period)
        #     compounded_liabilities[k] = compounded_amount
        # else:
        #     compounded_liabilities[k] = float(v)

    # Plot bar graph for assets
    st.write("### All Assets ")
    fig1, ax1 = plt.subplots()
    ax1.bar(filtered_assets.keys(), filtered_assets.values(), color='green')
    ax1.set_ylabel('Amount ($)')
    ax1.set_xlabel('Asset Type')
    ax1.set_title(' All Assets ')
    plt.xticks(rotation=45)
    st.pyplot(fig1)

    # Plot bar graph for liabilities
    st.write("### All Liabilities ")
    # st.write("### All Liabilities with Compounded Interest")
    fig2, ax2 = plt.subplots()
    # ax2.bar(compounded_liabilities.keys(), compounded_liabilities.values(), color='red')
    ax2.bar(filtered_liabilities.keys(), filtered_liabilities.values(), color='red')    
    ax2.set_ylabel('Amount ($)')
    ax2.set_xlabel('Liability Type')
    ax2.set_title(' All Liabilities ')

    # ax2.set_title(' All Liabilities with Compounded Interest')
    plt.xticks(rotation=45)
    st.pyplot(fig2)


from docx import Document
# Define a helper function to read and extract text from a DOCX file
def read_docx(file_path):
    document = Document(file_path)
    extracted_text = "\n".join([para.text for para in document.paragraphs])
    return extracted_text



class TrieNode:
    def __init__(self):
        self.children = {}
        self.client_ids = []
        self.end_of_name = False  # Marks the end of a client's name

class Trie:
    def __init__(self):
        self.root = TrieNode()

    def insert(self, name, client_id):
        node = self.root
        for char in name:
            if char not in node.children:
                node.children[char] = TrieNode()
            node = node.children[char]
        node.client_ids.append(client_id)
        node.end_of_name = True

    def search(self, prefix):
        node = self.root
        for char in prefix:
            if char in node.children:
                node = node.children[char]
            else:
                return []  # Prefix not found
        return self._get_all_names_from_node(prefix, node)

    def _get_all_names_from_node(self, prefix, node):
        suggestions = []
        if node.end_of_name:
            suggestions.append((prefix, node.client_ids))
        for char, child_node in node.children.items():
            suggestions.extend(self._get_all_names_from_node(prefix + char, child_node))
        return suggestions



def preload_trie():
    trie = Trie()
    clients = {
        "John Doe": "C001",
        "Jane Smith": "C002",
        "James Brown": "C003",
        "Jill Johnson": "C004",
        "Jake White": "C005"
    }
    for name, client_id in clients.items():
        trie.insert(name.lower(), client_id)  # Insert in lowercase for case-insensitive search
    return trie

async def generate_investment_suggestions_for_investor(investment_personality,clientName,financial_data,financial_file,monthly_investment=10000,investment_period=3): # # GET Method for py , for front end its Post API
    
    # retriever = asyncio.run(load_vector_db("uploaded_file"))

    # retriever =  await load_vector_db("uploaded_file")
    try:
        retriever =  await load_vector_db(financial_file)
    except Exception as e :
        print(f"Error : {e}")
        return jsonify("Error : Failed to load vector database and to generate suggestions : {e}"),400
    
    if not retriever:
        # await load_vector_db("data\Financial_Investment_new.docx")
        await load_vector_db("data\EW2400.docx")
        # await load_vector_db("data\Financial_Investment_1_new.docx") # doesnt works
        # await load_vector_db("data\Financial_Investment_1.docx")
        if not retriever:
            raise Exception("Failed to load vector database.")
    
    print("VectorDB is created successfully")
    # retriever = await load_vector_db("data\Financial_Investment_1.docx") 
    
    try:
        chain = await make_retrieval_chain(retriever,investment_personality,clientName,monthly_investment,investment_period)
    except Exception as e :
        print(f"Error : {e}")
        return jsonify("Error : Failed to create retrieval chain and generate suggestions : {e}"),400
    
    if not chain:
        raise Exception("Failed to create retrieval chain.")
    print("Chain is created to generate suggestions ")
    
    # chain = asyncio.run(make_retrieval_chain(retriever))
    
    print(f"Financial Data : {financial_data}")
    try :
        print(type(financial_data))
        query = f"The Investment Personality of {clientName} is : {investment_personality}" + f"Consider the Monthly Investment as {monthly_investment} and Investment period as {investment_period}" + f"Financial Data of client is : {financial_data[0]}"
        print(query)
    except Exception as e :
        print(f"Error : {e}")
        return "Error : Failed to load financial data"
    
    if chain is not None:
        # summary = context
        # query = summary + "\n" + investment_personality
        
        # query = str(investment_personality)
        response = chain.invoke({"input": query})
        
        # format_response = markdown_to_text(response['answer'])
        # return format_response
        
        # html_output = markdown.markdown(response['answer'])
        # return html_output
        
        # readable_text = markdown_to_readable_text(response['answer'])
        # print(readable_text)
        # return readable_text

        # format_text = convert_to_markdown(response['answer'])
        # return format_text
        
        return response['answer']
    
        

        # handle_graph(response['answer'])

    else:
        logging.INFO("response is not generated by llm model")
        return jsonify("response is not generated by llm model"),500
        # st.error("Failed to create the retrieval chain. Please upload a valid document.")

from flask import Flask, request, jsonify, send_file
import asyncio
from flask_cors import CORS
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})
# CORS(app,resources={r"/api/*":{"origins":"*"}})
# CORS(app)

# Initialize the Trie with preloaded clients
trie = preload_trie()

@app.route('/')
def home():
    return "Wealth Advisor Chatbot API"

@app.route('/investment-suggestions', methods=['POST'])
def investment_suggestions():
    # Get the input data (new or existing client)
    data = request.get_json()

    # Determine if it's a new client or existing client
    client_type = data.get("client_type")

    if client_type == "New Client":
        # Get form details and perform investment suggestions

        # Check if assets and liabilities are provided
        assets = data.get('assets', None)
        liabilities = data.get('liabilities', None)

        if assets and liabilities:
            financial_summary = create_financial_summary_table(assets, liabilities)
            bar_graphs = plot_bar_graphs(assets, liabilities)
            pie_chart = plot_assets_liabilities_pie_chart(assets, liabilities)

            return jsonify({
                "financial_summary": financial_summary,
                "bar_graphs": "Bar graphs generated.",
                "pie_chart": "Pie chart generated."
            })

        return jsonify({"message": "Please fill in the client details to view the assets and liabilities breakdown."})

    elif client_type == "Existing Client":
        # Search for an existing client in the Trie
        search_query = data.get("search_query", "").lower()
        matching_names = trie.search(search_query)

        if matching_names:
            suggestions = [{"name": name, "client_ids": client_ids} for name, client_ids in matching_names]
            return jsonify({"suggestions": suggestions})
        else:
            return jsonify({"message": "No matching clients found."})
    
    return jsonify({"message": "Invalid client type."})



import random

# Generate unique client ID

# def generate_unique_id(name):
#     name_parts = name.split(" ")
#     first_initial = name_parts[0][0] if len(name_parts) > 0 else ""
#     last_initial = name_parts[1][0] if len(name_parts) > 1 else ""
#     random_number = random.randint(1000, 9999)
#     unique_id = f"{first_initial}{last_initial}{random_number}"
#     return unique_id

# # Save details in a Word file
import docx
import os
# #Curr version :
def save_to_word_file(data, file_name):
    doc = docx.Document()
    doc.add_heading('Client Details', 0)

    # Adding client details
    client_details = data.get('clientDetail', {})
    doc.add_paragraph(f"Client Name: {client_details.get('clientName', '')}")
    doc.add_paragraph(f"Client Mobile: {client_details.get('clientMoNo', '')}")
    doc.add_paragraph(f"Client Age: {client_details.get('clientAge', '')}")
    doc.add_paragraph(f"Co-Client Name: {client_details.get('coClientName', '')}")
    doc.add_paragraph(f"Co-Client Mobile: {client_details.get('coMobileNo', '')}")
    doc.add_paragraph(f"Co-Client Age: {client_details.get('coClientAge', '')}")

    # Retirement Plan
    retirement_goal = data.get('retirementGoal', {})
    retirement_plan = retirement_goal.get('retirementPlan', {})
    doc.add_paragraph(f"Retirement Plan Client Age: {retirement_plan.get('retirementAgeClient', '')}")
    doc.add_paragraph(f"Retirement Plan Co-Client Age: {retirement_plan.get('retirementAgeCoClient', '')}")
    
    social_benefit = retirement_goal.get('socialBenefit', {})
    doc.add_paragraph(f"Social Benefit Client: {social_benefit.get('socialBenefitClient', '')}")
    doc.add_paragraph(f"Social Benefit Co-Client: {social_benefit.get('socialBenefitCoClient', '')}")
    
    pension_benefit = retirement_goal.get('pensionBenefit', {})
    doc.add_paragraph(f"Pension Benefit Client: {pension_benefit.get('pensionBenefitClient', '')}")
    doc.add_paragraph(f"Pension Benefit Co-Client: {pension_benefit.get('pensionBenefitCoClient', '')}")
    
    otherIncome = retirement_goal.get('otherIncome', {})
    doc.add_paragraph(f"Other IncomeClient Client: {otherIncome.get('otherIncomeClient', '')}")
    doc.add_paragraph(f"Other IncomeClient Co-Client: {otherIncome.get('otherIncomeCoClient', '')}")
   
    # Estimated Annual Retirement Expense ($ or % of current salary)
    annualRetirement = retirement_goal.get('annualRetirement', {})
    doc.add_paragraph(f"Estimated Annual Retirement Expense ($ or % of current salary) Client: {annualRetirement.get('annualRetireClient', '')}")
    doc.add_paragraph(f"Estimated Annual Retirement Expense ($ or % of current salary) Co-Client: {annualRetirement.get('annualRetireCoClient', '')}")
    

    # Assets and Liabilities
    assets_liabilities = data.get('assetsLiabilities', {})
    
    # Assets
    
    for asset_key, asset_info in assets_liabilities.items():
        current_value_key = [key for key in asset_info.keys() if key.startswith("current")][0]
        annual_value_key = [key for key in asset_info.keys() if key.startswith("annual")][0]
        assets_name_key = "assetsName"
        doc.add_paragraph(f"Assets - {asset_info[assets_name_key]} : Current Value - {asset_info[current_value_key]} , Annual Contributions - {asset_info[annual_value_key]}")
        
    # Liabilities
    myLiabilities = data.get('myLiabilities', {})
    for liability_key, liability_info in myLiabilities.items():
        balance_key = [key for key in liability_info.keys() if key.endswith("Balance")][0]
        interest_key = [key for key in liability_info.keys() if key.endswith("Interest")][0]
        monthly_key = [key for key in liability_info.keys() if key.endswith("Monthly")][0]
        liability_name_key = "liabilityName"
        doc.add_paragraph(f"Liabilities - {liability_info[liability_name_key]} : Balance - {liability_info[balance_key]} , Interest - {liability_info[interest_key]} , Monthly - {liability_info[monthly_key]}")
        
    # my_liabilities = data.get('myLiabilities', {})
    # for liability_type, liability_info in my_liabilities.items():
    #     doc.add_paragraph(f"Liabilities - {liability_info.get('liabilityName', '')}: Balance - {liability_info.get('mortgageBalance', '')} Interest - {liability_info.get('mortgageInterest', '')} Monthly - {liability_info.get('mortgageMonthly', '')}")

    # Protection Plan
    protection_plan = data.get('protectionPlan', {})
    doc.add_paragraph(f"Check Will: {protection_plan.get('checkWill', False)}")
    doc.add_paragraph(f"Check Healthcare: {protection_plan.get('checkHealthCare', False)}")
    doc.add_paragraph(f"Check Attorney: {protection_plan.get('checkAttorney', False)}")
    doc.add_paragraph(f"Check Trust: {protection_plan.get('checkTrust', False)}")

    # Insurance Coverage
    insurance_coverage = data.get('insuranceCoverage', {})
    life_insurance_client = insurance_coverage.get('lifeInsuranceClient', {})
    doc.add_paragraph(f"Life Insurance Client: Benefit - {life_insurance_client.get('benefitLIClient', '')} Monthly Pay - {life_insurance_client.get('monthlyPayLIClient', '')}")
    
    life_insurance_co_client = insurance_coverage.get('lifeInsuranceCoClient', {})
    doc.add_paragraph(f"Life Insurance Co-Client: Benefit - {life_insurance_co_client.get('benefitLICoClient', '')} Monthly Pay - {life_insurance_co_client.get('monthlyPayLICoClient', '')}")
 
    disableIncome = insurance_coverage.get('disableIncomeClient', {})
    disableIncomeClient = insurance_coverage.get('disableIncomeClient',{})
    doc.add_paragraph(f"Disable Income Client - {disableIncomeClient.get('benefitDisableClient', '')}")
    
    disableIncomeCoClient = insurance_coverage.get('disableIncomeCoClient', {})
    doc.add_paragraph(f"Disable Income Co-Client - {disableIncomeCoClient.get('benefitDisableCoClient', '')}")
    
    longTermCoClient = insurance_coverage.get('longTermCoClient')
    doc.add_paragraph(f"Long Term Client: Benefit - {longTermCoClient.get('benefitLongTermClient', '')} Monthly Pay - {longTermCoClient.get('monthlyPayLongTermClient', '')}")
    
    # Goal Fields
    goal_fields = data.get('goalFields', [])
    for goal in goal_fields:
        doc.add_paragraph(f"Goal: {goal.get('goal', '')} Cost: {goal.get('cost', '')} When: {goal.get('when', '')}")

    # Income Fields
    income_fields = data.get('incomeFields', [])
    for income in income_fields:
        doc.add_paragraph(f"Income Source: {income.get('sourceIncome', '')} Amount: {income.get('amountIncome', '')}")

    # Save file
    file_name = os.path.join("data", file_name)
    doc.save(f"{file_name}.docx")


@app.route('/submit-client-data', methods=['POST'])
def submit_client_data():
    try:
        data = request.get_json()
        print(data)
        # Generate the unique ID
        client_name = data['clientDetail']['clientName']
        
        # unique_id = generate_unique_id(client_name)
        # unique_id = data['clientDetail']['uniqueId']
        
        unique_id = data['uniqueId']
        print(unique_id)
        
        data['uniqueId'] = unique_id
        
        # Save the data to a Word file
        
        file_name = unique_id
        # save_to_word_file(data, file_name)
        save_to_word_file(data, file_name)
        
        return jsonify({
            'message': 'Client data received and saved successfully.'
        }), 200
    except Exception as e:
        return jsonify({'message': f"An error occurred: {e}"}), 500


# Determine Investment personality through the investor assesmnet tab : 
@app.route('/investor-personality-assessment', methods=['POST'])
def investor_personality_assessment():
    try:
        # Collecting client name and assessment data
        data = request.json  # Expecting JSON input
        # client_name = data.get('client_name')
        client_id = data.get('client_id')
        assessment_data = data.get('assessment_data')  
        
        # if not client_id or not assessment_data:
        #     return jsonify({'message': 'Client name and assessment data are required.'}), 400
        
        logging.info(f"Received assessment data for client with client id : {client_id}")

        # Pass the assessment data to determine the investment personality
        personality = asyncio.run(determine_investment_personality(assessment_data))
        logging.info(f"Determined personality for {client_id}: {personality}")

        # Return the personality and client id in response
        return jsonify({
            'client_id': client_id,
            'investment_personality': personality
        }), 200
    
    except Exception as e:
        logging.error(f"Error processing investor assessment: {e}")
        return jsonify({'message': 'Internal Server Error'}), 500
    

import logging
# global investmentPersonality  # Global Variable
# investmentPersonality = ""

async def make_suggestions(investmentPersonality,clientName,financial_file="data\Financial_Investment_1.docx",monthly_investment=10000,investment_period=3):
    try:
        try:
            # financial_file = financial_file
            print(f"Processing the File for the client: {clientName} and the file : {financial_file}")
            # financial_data = asyncio.run(process_document(financial_file))
            
            financial_data = await process_document(f"data\{financial_file}")
            
            print(f"Data passed : {financial_data}")
            
            financial_file = f"data\{financial_file}"
            
            print(f"Finished processing the File for the client : {financial_file}")
            
            # suggestions = await generate_investment_suggestions_for_investor(investmentPersonality,clientName, financial_file,monthly_investment,investment_period)
            
            suggestions = await generate_investment_suggestions_for_investor(investmentPersonality,clientName,financial_data,financial_file,monthly_investment,investment_period)
            
            
            # print(f"Finished processing the suggestions : {suggestions}")
           
            htmlSuggestions = markdown.markdown(suggestions)
            # ans = markdown_to_readable_text(htmlSuggestions)
            # formatSuggestions = ans
            # formatSuggestions = markdown_to_text(htmlSuggestions)
            # print(f"HTML Suggestions: {htmlSuggestions}")
            formatSuggestions = markdown.markdown(htmlSuggestions)
            
            # ---------------------------------=----------
            # htmlSuggestions = markdown.markdown(suggestions)
            # htmlSuggestions = markdown2.markdown(suggestions)
            
            # print(f"HTML Suggestions: {htmlSuggestions}")
            
            # formatSuggestions = markdown_to_text(suggestions)
            
            # print(f"The suggestions generated for the client are :\n {formatSuggestions}")
            
            # need to change the data extraction process : 
            data_extracted = extract_numerical_data(suggestions)
            
            min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                            [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
            max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                            [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

            # Normalize allocations
            min_allocations = normalize_allocations(min_allocations)
            max_allocations = normalize_allocations(max_allocations)

            # Update Bar Chart Data
            bar_chart_data = {
                'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
                'datasets': [{
                    'label': 'Min Allocation',
                    'data': min_allocations,
                    'backgroundColor': 'skyblue'
                },
                {
                    'label': 'Max Allocation',
                    'data': max_allocations,
                    'backgroundColor': 'lightgreen'
                }]
            }

            # Similar changes can be made for the Pie Chart Data:
            all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
            num_labels = len(all_labels)
            max_allocations_for_pie = normalize_allocations(
                [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
                [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
            )
            
            # Generate colors based on the number of labels
            dynamic_colors = generate_colors(num_labels)

            # Update Pie Chart Data
            pie_chart_data = {
                'labels': all_labels,
                'datasets': [{
                    'label': 'Investment Allocation',
                    'data': max_allocations_for_pie,
                    'backgroundColor': dynamic_colors,
                    'hoverOffset': 4
                }]
            }
            
            
            # Prepare the data for the line chart with inflation adjustment
            initial_investment = 10000
            # compounded_chart_data, inflation_adjusted_chart_data = prepare_line_chart_data_with_inflation(data_extracted, initial_investment)
            combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
            print(f"\nThe combined chart data is : {combined_chart_data}")
            
            # return htmlSuggestions, pie_chart_data, bar_chart_data, combined_chart_data
            
            print(f"Format suggestions : {formatSuggestions}")
            
            return formatSuggestions, pie_chart_data, bar_chart_data, combined_chart_data
            
            
        except Exception as e:
            logging.info(f"Error occurred while generating investment suggestions: {e}")
            return jsonify({'message': f'Error occurred while considering preuploaded file : {e}'}), 500
    

        
        # return jsonify({
        #     "status": 200,
        #     "message": "Success",
        #     "investmentSuggestions": htmlSuggestions,
        #     "pieChartData": pie_chart_data,
        #     "barChartData": bar_chart_data,
        #     "compoundedChartData":combined_chart_data
        # }), 200
        
    except Exception as e:
        logging.error(f"Error processing personality assessment: {e}")
        return jsonify({'message': 'Error in generating suggestions with personality'}), 500
        

@app.route('/personality-assessment', methods=['POST'])
def personality_selected():
    try:
        data = request.json
        clientName = data.get('clientName')
        try :
            # clientId = data.get('client_id')
            clientId = data.get('clientId')
            investmentPersonality = data.get('investmentPersonality') # investment_personality
            financial_file = f"{clientId}.docx"
            # financial_file = f"data\{clientId}.docx" # data\EW2400.docx
            print(f"The clients ClientName is : {clientName} and their ClientId is : {clientId}")
            print(f"InvestmentPersonality received is : {investmentPersonality}")
            logging.info('Recieved Values')
            
        except Exception as e:
            logging.info(f"Error occurred while retrieving client id: {e}")
            return jsonify({'message': f'Error occurred while retrieving client id: {e}'}), 400

        try:
            # monthly_investment= data.get('monthly_investment') #10000
            # investment_period= data.get('investment_period')  #3
            monthly_investment= 10000
            investment_period= 3
            formatSuggestions,pie_chart_data,bar_chart_data,combined_chart_data = asyncio.run(make_suggestions(investmentPersonality,clientName,financial_file,monthly_investment,investment_period))
            answer = markdown_table_to_html(formatSuggestions)
            print(answer)
            # htmlSuggestions,pie_chart_data,bar_chart_data,combined_chart_data = asyncio.run(make_suggestions(investmentPersonality,clientName,financial_file,monthly_investment,investment_period))
            
        except Exception as e:
            logging.info(f"Error occurred while processing investment data: {e}")
            return jsonify({'message': f'Error occurred while processing investment data: {e}'}), 400
        
        # htmlSuggestions,pie_chart_data,bar_chart_data,combined_chart_data = asyncio.run(make_suggestions(investmentPersonality,clientName))
        
        # return jsonify({
        #     "status": 200,
        #     "message": "Success",
        #     "investmentSuggestions": htmlSuggestions,
        #     "pieChartData": pie_chart_data,
        #     "barChartData": bar_chart_data,
        #     "compoundedChartData":combined_chart_data
        # }), 200
        return jsonify({
            "status": 200,
            "message": "Success",
            "investmentSuggestions": answer, #formatSuggestions,
            "pieChartData": pie_chart_data,
            "barChartData": bar_chart_data,
            "compoundedChartData":combined_chart_data
        }), 200
                    
                    
        # return jsonify({'message':'Sab thik'}),200
        
        # if investmentPersonality == 'aggressiveInvestor':
        #     pass
    
    except Exception as e:
        logging.info(f"Error in personality assessment: {e}")
        print(f"Error occured in Investor Personality while collecting data :\n{e}")
        return jsonify({'message': 'Internal Server Error in Investor Personality'}), 500



# Route to handle generating investment suggestions
import shutil
import os

def save_file_to_folder(file_storage, destination_folder):
    try:
        # Ensure the destination folder exists
        if not os.path.exists(destination_folder):
            os.makedirs(destination_folder)
        
        # Construct the destination file path
        destination_file_path = os.path.join(destination_folder, file_storage.filename)
        
        # Check if the file already exists
        if not os.path.exists(destination_file_path):
            # Save the file
            file_storage.save(destination_file_path)
            print(f"File saved to {destination_file_path}")
            return destination_file_path
        else:
            print(f"File already exists at {destination_file_path}")
            return destination_file_path
        
    except Exception as e:
        print(f"Error saving file: {e}")



# #Stock analysis code :

from flask import Flask, request, jsonify
import yfinance as yf
import pandas as pd
import requests
import os
import logging


NEWS_API_KEY = os.getenv('NEWS_API_KEY')


# Included Additional KPI's except for ROCE and Order Booking
# @app.route('/analyze_stock', methods=['POST'])
# def analyze_stock():
#     try:
#         ticker = request.json.get('ticker')
        
#         if not ticker:
#             return jsonify({"error": "Ticker is required"}), 400

#         query = request.json.get('query')
#         if not query:
#             query = "Generate Stock Analysis and give some predictions on its price "
        
#         # Step 1: Fetch Stock Data
#         stock = yf.Ticker(ticker)
#         data = {}

#         company_details = stock.info.get('longBusinessSummary', 'No details available')
#         data['Company Details'] = company_details
#         sector = stock.info.get('sector', 'No sector information available')
#         data['Sector'] = sector
#         prev_close = stock.info.get('previousClose', 'No previous close price available')
#         data['Previous Closing Price'] = prev_close
#         open_price = stock.info.get('open', 'No opening price available')
#         data['Today Opening Price'] = open_price

#         # Additional KPIs
#         data['EPS'] = stock.info.get('trailingEps', 'No EPS information available')
#         data['Book_Value'] = stock.info.get('bookValue', 'No book value available')
#         data['ROE'] = stock.info.get('returnOnEquity', 'No ROE information available')
#         data['ROCE'] = stock.info.get('returnOnAssets', 'No ROCE information available')  # ROCE is not available directly

        # # Revenue Growth (CAGR) and Earnings Growth would need to be calculated based on historical data
        # earnings_growth = stock.info.get('earningsGrowth', 'No earnings growth available')
        # revenue_growth = stock.info.get('revenueGrowth', 'No revenue growth available')

        # data['Earnings Growth'] = earnings_growth
        # data['Revenue Growth'] = revenue_growth

#         # Fetch historical financials to calculate CAGR or additional metrics if required
#         financials = stock.financials

#         # Process news related to stock
#         news_url = f'https://newsapi.org/v2/everything?q={ticker}&apiKey={NEWS_API_KEY}&pageSize=3'
#         news_response = requests.get(news_url)
#         if news_response.status_code == 200:
#             news_data = news_response.json()
#             articles = news_data.get('articles', [])
#             if articles:
#                 top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
#                 data['Top News'] = top_news
#             else:
#                 data['Top News'] = "No news articles found."
#         else:
#             data['Top News'] = "Failed to fetch news articles."
    
#     except Exception as e:
#         logging.error(f"Error occurred while collecting stock data: {e}")
#         return jsonify({'message': 'Internal Server Error in Stock Data Collection'}), 500

#     # Save Financial Data to Excel
#     try:
#         file_path = os.path.join('data', f'{ticker}_financial_data.xlsx')
#         with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
#             stock.financials.to_excel(writer, sheet_name='Income Statement')
#             stock.balance_sheet.to_excel(writer, sheet_name='Balance Sheet')
#             stock.cashflow.to_excel(writer, sheet_name='Cashflow')
    
#     except Exception as e:
#         logging.error(f"Error occurred while saving financial data: {e}")
#         return jsonify({'message': 'Internal Server Error in saving data to Excel'}), 500

#     # Perform stock analysis using AI model
#     try:
#         task = f"""You are a Stock Market Expert. Analyze the stock's performance and predict its price for the next week. 
#                    Consider the following stock data: {data}"""
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content(task)
#         html_suggestions = markdown.markdown(response.text)
        
#     except Exception as e:
#         logging.error(f"Error performing analysis with AI model: {e}")
#         return jsonify({"error": "Failed to analyze stock data"}), 500

#     # Return response
#     return jsonify({
#         "data": data,
#         "analysis": markdown_to_text(html_suggestions),
#         "news": data['Top News'],
#         "graph_url": f"https://finance.yahoo.com/chart/{ticker}"
#     })



# # Best version 
@app.route('/analyze_stock', methods=['POST'])
def analyze_stock():
    try:
        ticker = request.json.get('ticker')
        
        if not ticker:
            print("Ticker is required")
            # return jsonify({"error": "Ticker is required"}), 400

        query = request.json.get('query')
        if not query:
            query = "Generate Stock Analysis and give some predictions on its price "
        # Step 1: Fetch Stock Data
        stock = yf.Ticker(ticker)
        data = {}

        company_details = stock.info.get('longBusinessSummary', 'No details available')
        data['Company_Details'] = company_details
        sector = stock.info.get('sector', 'No sector information available')
        data['Sector'] = sector
        prev_close = stock.info.get('previousClose', 'No previous close price available')
        data['Previous_Closing_Price'] = prev_close
        open_price = stock.info.get('open', 'No opening price available')
        data['Today_Opening_Price'] = open_price
         
        hist = stock.history(period="5d")
        if not hist.empty and 'Close' in hist.columns:
            if hist.index[-1].date() == yf.download(ticker, period="1d").index[-1].date():
                close_price = hist['Close'].iloc[-1]
                data['Todays_Closing_Price'] = close_price
            else:
                data['Todays_Closing_Price'] = "Market is open, no closing price available yet."
        else:
            data['Todays_Closing_Price'] = "No historical data available for closing price."

        day_high = stock.info.get('dayHigh', 'No high price available')
        data['Today_High_Price'] = day_high
        day_low = stock.info.get('dayLow', 'No low price available')
        data['Today_Low_Price'] = day_low
        volume = stock.info.get('volume', 'No volume information available')
        data['Today_Volume'] = volume
        dividends = stock.info.get('dividendRate', 'No dividend information available')
        data['Today_Dividends'] = dividends
        splits = stock.info.get('lastSplitFactor', 'No stock split information available')
        data['Today_Stock_Splits'] = splits
        pe_ratio = stock.info.get('trailingPE', 'No P/E ratio available')
        data['PE_Ratio'] = pe_ratio
        market_cap = stock.info.get('marketCap', 'No market cap available')
        data['Market_Cap'] = market_cap

        # Additional KPIs
        data['EPS'] = stock.info.get('trailingEps', 'No EPS information available')
        data['Book_Value'] = stock.info.get('bookValue', 'No book value available')
        data['ROE'] = stock.info.get('returnOnEquity', 'No ROE information available')
        data['ROCE'] = stock.info.get('returnOnAssets', 'No ROCE information available')  # ROCE is not available directly
        
        # Revenue Growth (CAGR) and Earnings Growth would need to be calculated based on historical data
        earnings_growth = stock.info.get('earningsGrowth', 'No earnings growth available')
        revenue_growth = stock.info.get('revenueGrowth', 'No revenue growth available')

        data['Earnings_Growth'] = earnings_growth
        data['Revenue_Growth'] = revenue_growth
        
        
        income_statement = stock.financials
        balance_sheet = stock.balance_sheet
        cashflow = stock.cashflow

        # Step 2: Get News Related to Stock
        news_url = f'https://newsapi.org/v2/everything?q={ticker}&apiKey={NEWS_API_KEY}&pageSize=3'
        news_response = requests.get(news_url)
        if news_response.status_code == 200:
            news_data = news_response.json()
            articles = news_data.get('articles', [])
            if articles:
                top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
                data['Top_News'] = top_news
            else:
                data['Top_News'] = "No news articles found."
        else:
            data['Top_News'] = "Failed to fetch news articles."
    except Exception as e:
        logging.info(f"Error occurred while collecting stock data: {e}")
        print(f"Error occurred while collecting stock data: :\n{e}")
        return jsonify({'message': 'Internal Server Error in Stock Data Collection'}), 500
    
    print(data['Top_News'])
    
    try:
            
        # Step 3: Save Financial Data to Excel
        file_path = os.path.join('data', f'{ticker}_financial_data.xlsx')
        with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
            income_statement.to_excel(writer, sheet_name='Income Statement')
            balance_sheet.to_excel(writer, sheet_name='Balance Sheet')
            cashflow.to_excel(writer, sheet_name='Cashflow')

        # Step 4: Perform Analysis
        avg_close = hist['Close'].mean()
        formatted_data = extract_excel_data(file_path)
    
    except Exception as e:
        logging.info(f"Error occurred while performing analysis: {e}")
        print(f"Error occurred while performing analysis :\n{e}")
        return jsonify({'message': 'Internal Server Error in Stock Analysis'}), 500

    try:
        
        task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.
                    Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
                    Predict the stock price range for the next week and provide reasons for your prediction.
                    Advise whether to buy this stock now or not, with reasons for your advice."""

        query = task + "\nStock Data: " + str(data) + "\nFinancial Data: " + formatted_data

        # Use your generative AI model for analysis (example with 'gemini-1.5-flash')
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(query)
        print(response.text)
        print(data)
    
    except Exception as e:
        logging.error(f"Error performing analysis with generative AI: {e}")
        return jsonify({f"error": "Failed to give analysis of stock data : {e}"}), 500
    
    # Extract response from the model
    try:
        html_suggestions = markdown.markdown(response.text)
        
        print(f"Html Suggestions : {html_suggestions}")
        
        logging.info(f"Suggestions for stock: \n{response.text}")
        
        # format_suggestions = markdown_to_text(response)
        print(f"Html Suggestions : {html_suggestions}")
        format_suggestions = markdown_to_text(html_suggestions)
        
    except Exception as e:
        logging.error(f"Error extracting text from response: {e}")
        print(f"Error extracting text from response : {e}")
        return jsonify({"error": "Failed to analyze stock data"}), 500

    # Return all collected and analyzed data
    return jsonify({
        "data": data,
        "average_closing_price": f"${avg_close:.2f}",
        "analysis": format_suggestions,
        "news": data['Top_News'],
        "graph_url": f"https://finance.yahoo.com/chart/{ticker}"
    })

def extract_excel_data(file_path):
    financial_data = ""
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        financial_data += f"\n\nSheet: {sheet_name}\n"
        financial_data += df.to_string()
    return financial_data


# @app.route('/fetch_stock_data', methods=['GET'])
# @app.route('/fetch_stock_data', methods=['POST'])
# def fetch_stock_data():
#     ticker = request.args.get('ticker')
#     if not ticker:
#         return jsonify({"error": "Ticker is required"}), 400

#     stock = yf.Ticker(ticker)
#     data = {}

#     company_details = stock.info.get('longBusinessSummary', 'No details available')
#     data['Company Details'] = company_details
#     sector = stock.info.get('sector', 'No sector information available')
#     data['Sector'] = sector
#     prev_close = stock.info.get('previousClose', 'No previous close price available')
#     data['Previous Closing Price'] = prev_close
#     open_price = stock.info.get('open', 'No opening price available')
#     data['Today Opening Price'] = open_price

#     hist = stock.history(period="5d")
#     if not hist.empty and 'Close' in hist.columns:
#         if hist.index[-1].date() == yf.download(ticker, period="1d").index[-1].date():
#             close_price = hist['Close'].iloc[-1]
#             data['Todays Closing Price'] = close_price
#         else:
#             data['Todays Closing Price'] = "Market is open, there is no closing price available yet."
#     else:
#         data['Todays Closing Price'] = "No historical data available for closing price."

#     day_high = stock.info.get('dayHigh', 'No high price available')
#     data['Today High Price'] = day_high
#     day_low = stock.info.get('dayLow', 'No low price available')
#     data['Today Low Price'] = day_low
#     volume = stock.info.get('volume', 'No volume information available')
#     data['Today Volume'] = volume
#     dividends = stock.info.get('dividendRate', 'No dividend information available')
#     data['Today Dividends'] = dividends
#     splits = stock.info.get('lastSplitFactor', 'No stock split information available')
#     data['Today Stock Splits'] = splits
#     pe_ratio = stock.info.get('trailingPE', 'No P/E ratio available')
#     data['P/E Ratio'] = pe_ratio
#     market_cap = stock.info.get('marketCap', 'No market cap available')
#     data['Market Cap'] = market_cap

#     income_statement = stock.financials
#     balance_sheet = stock.balance_sheet
#     cashflow = stock.cashflow

#     news_url = f'https://newsapi.org/v2/everything?q={ticker}&apiKey={NEWS_API_KEY}&pageSize=3'
#     news_response = requests.get(news_url)
#     if news_response.status_code == 200:
#         news_data = news_response.json()
#         articles = news_data.get('articles', [])
#         if articles:
#             top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
#             data['Top News'] = top_news
#         else:
#             data['Top News'] = "No news articles found."
#     else:
#         data['Top News'] = "Failed to fetch news articles."

#     graph_url = f"https://finance.yahoo.com/chart/{ticker}"
#     data['graph_url'] = graph_url

#     file_path = os.path.join('data', f'{ticker}_financial_data.xlsx')
#     with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
#         income_statement.to_excel(writer, sheet_name='Income Statement')
#         balance_sheet.to_excel(writer, sheet_name='Balance Sheet')
#         cashflow.to_excel(writer, sheet_name='Cashflow')

#     data_list = list(data.items())
#     data_str = str(data_list)

#     return jsonify({
#         "data": data,
#         "file_path": file_path,
#         "data_str": data_str
#     })

# # @app.route('/analyze_stock_data', methods=['GET'])
# @app.route('/analyze_stock_data', methods=['POST'])
# def analyze_stock_data():
#     ticker = request.args.get('ticker')
#     if not ticker:
#         return jsonify({"error": "Ticker is required"}), 400

#     query = "Give Stock Analysis :"
#     query1 = request.args.get('query')
#     print(query1)
#     print(type(query1))
#     if query1:
#         query = query + query1
#     hist, data_str, file_path = fetch_stock_data(ticker)
#     avg_close = hist['Close'].mean()
#     formatted_data = extract_excel_data(file_path)

#     task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.
#                 Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
#                 Predict the stock price range for the next week and provide reasons for your prediction.
#                 Advise whether to buy this stock now or not, with reasons for your advice."""

#     query = task + "\nStock Data: " + data_str + "\nFinancial Data: " + formatted_data
#     model = genai.GenerativeModel('gemini-1.5-flash')
#     response = model.generate_content(query)
    
#     # Log the response object to understand its structure
#     logging.info(f"Model response: {response['answer']}")
    
#     # Extract the text content from the response
#     try:
#         # response_text = response.text
#         # format_response = markdown_to_text(response_text)
#         htmlSuggestions = markdown.markdown(response['answer'])
#         logging.info(f"Suggestions for stock: \n{response['answer']}")
        
#         formatSuggestions = markdown_to_text(response)
        
#     except Exception as e:
#         logging.error(f"Error extracting text from response: {e}")
#         return jsonify({"error": "Failed to analyze stock data"}), 500

#     return jsonify({
#         "average_closing_price": f"${avg_close:.2f}",
#         "analysis": formatSuggestions
#     })

# def extract_excel_data(file_path):
#     financial_data = ""
#     xls = pd.ExcelFile(file_path)
#     for sheet_name in xls.sheet_names:
#         df = pd.read_excel(xls, sheet_name=sheet_name)
#         financial_data += f"\n\nSheet: {sheet_name}\n"
#         financial_data += df.to_string()
#     return financial_data




# #Working for both the methods :
@app.route('/generate-investment-suggestions', methods=['POST'])
def generate_investment_suggestions():
    try:
        assessment_file = request.files['assessmentFile']
        financial_file = request.files['financialFile']
        logging.info("Requested files")
        
        responses = extract_responses_from_docx(assessment_file)
        if not responses:
            raise Exception("Failed to extract responses from assessment file.")
        
        destination_folder = 'data'
        file_path = save_file_to_folder(financial_file, destination_folder)
        if not file_path:
            raise Exception("Failed to save financial file.")
        
        financial_data = asyncio.run(process_document(file_path))
        if not financial_data:
            raise Exception("Failed to process financial file.")
        
        logging.info(f"Received Responses from the file {responses}")
        
        personality = asyncio.run(determine_investment_personality(responses))
        if not personality:
            raise Exception("Failed to determine personality.")
        
        logging.info(f"Personality of the user is: {personality}")
        
        clientName = "Emilly Watts"
        suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality, clientName, financial_data, file_path))
        if "Error" in suggestions:
            raise Exception(suggestions)
        
        htmlSuggestions = markdown.markdown(suggestions)
        logging.info(f"Suggestions for investor: \n{suggestions}")
        
        formatSuggestions = markdown_to_text(htmlSuggestions)
        answer = markdown_table_to_html(formatSuggestions)
        print(answer)
        data_extracted = extract_numerical_data(suggestions)
        
        min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                          [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
        max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                          [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

        # Normalize allocations
        min_allocations = normalize_allocations(min_allocations)
        max_allocations = normalize_allocations(max_allocations)

        # Update Bar Chart Data
        bar_chart_data = {
            'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
            'datasets': [{
                'label': 'Min Allocation',
                'data': min_allocations,
                'backgroundColor': 'skyblue'
            },
            {
                'label': 'Max Allocation',
                'data': max_allocations,
                'backgroundColor': 'lightgreen'
            }]
        }

        # Similar changes can be made for the Pie Chart Data:
        all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
        num_labels = len(all_labels)
        max_allocations_for_pie = normalize_allocations(
            [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
            [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
        )
        
        # Generate colors based on the number of labels
        dynamic_colors = generate_colors(num_labels)

        # Update Pie Chart Data
        pie_chart_data = {
            'labels': all_labels,
            'datasets': [{
                'label': 'Investment Allocation',
                'data': max_allocations_for_pie,
                'backgroundColor': dynamic_colors,
                'hoverOffset': 4
            }]
        }
        
        # Prepare the data for the line chart with inflation adjustment
        initial_investment = 10000
        combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
        print(f"\nThe combined chart data is: {combined_chart_data}")
        
        return jsonify({
            "status": 200,
            "message": "Success",
            "investmentSuggestions":  answer, #formatSuggestions,
            "pieChartData": pie_chart_data,
            "barChartData": bar_chart_data,
            "compoundedChartData": combined_chart_data
        }), 200
        
        # return jsonify({
        #     "status": 200,
        #     "message": "Success",
        #     "investmentSuggestions": htmlSuggestions,
        #     "pieChartData": pie_chart_data,
        #     "barChartData": bar_chart_data,
        #     "compoundedChartData": combined_chart_data
        # }), 200

    except Exception as e:
        logging.info(f"Error in generating investment suggestions: {e}")
        return jsonify({'message': f'Internal Server Error in Generating responses : {e}'}), 500


# #Current error code
# @app.route('/generate-investment-suggestions', methods=['POST'])
# def generate_investment_suggestions():
#     try:
#         try :
#             assessment_file = request.files['assessmentFile']
#             financial_file = request.files['financialFile']
#             logging.info(" Requested files")
#         except Exception as e:
#             logging.info(" Requested files not passed")
#             return jsonify({'message': f'Error occurred while retrieving files12: {e}'}), 400
                   
#         try :
#             try:
#                 responses = extract_responses_from_docx(assessment_file)
#             except Exception as e:
#                 logging.info(f"Failed to extract responses from assessment file: {e}")
#                 return jsonify({'message': 'Failed to extract responses from assessment file.'}), 400
            
#             try:
#                 financial_data = asyncio.run(process_document(financial_file))
#                 print(financial_data)
#                 print(f"The Financial Data is of type : {type(financial_data)}")
#                 print(financial_file)
#                 destination_folder = 'data'
#                 file_path = save_file_to_folder(financial_file, destination_folder)
#                 print(file_path)
#             except Exception as e:
#                 logging.info(f"Failed to process financial file: {e}")
#                 return jsonify({'message': 'Failed to process financial file.'}), 400

#             logging.info(f"Received Responses from the file {responses}")
#         except Exception as e:
#             logging.info("Failed to process files")
#             return jsonify({'message': f'Error occurred while processing files: {e}'}), 400

#         try:
#             # Determine investment personality
#             # monthly_investment= 10000
#             # investment_period= 3
#             # personality,monthly_investment,investment_period = asyncio.run(determine_investment_personality(responses))
            
#             personality= asyncio.run(determine_investment_personality(responses))
            
#             logging.info(f"\nPersonality of the user is : {personality}")
#             print(f"\nPersonality of the user is : {personality}")
#             print(f"Type of Investment Personality is : {type(personality)}")
            
#         except Exception as e:
#             logging.info("Failed to determine personality")
#             return jsonify({'message': f'Error occurred while determining personality: {e}'}), 400
        
#         try:
#             # Generate investment suggestions based on personality and financial data
#             clientName = "Emilly Watts"
#             print(f"{type(personality)} \n {type(financial_data)} \n {type(financial_file)} \n {type(file_path)}")
#             try:
#                 suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality,clientName, financial_data,file_path)) #monthly_investment,investment_period))
#             except Exception as e:
#                 logging.info(f"Failed to generate investment suggestions for investor: {e}")
#                 print(f"Failed to generate investment suggestions for investor: {e}")
#                 return jsonify({'message': f'Failed to generate investment suggestions for investor. : {e}'}), 400
            
#             # suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality,clientName, financial_data,monthly_investment,investment_period))
#             htmlSuggestions = markdown.markdown(suggestions)
#             logging.info(f"\Suggestions for investor : \n{suggestions}")
            
#         except Exception as e:
#             logging.info("Failed to generate suggestions")
#             return jsonify({'message': f'Error occurred while generating suggestions: {e}'}), 400


#         logging.info("Successfully generated")
#         formatSuggestions = markdown_to_text(suggestions)
#         data_extracted = extract_numerical_data(suggestions)
        
#         min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                         [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
#         max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                         [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

#         # Normalize allocations
#         min_allocations = normalize_allocations(min_allocations)
#         max_allocations = normalize_allocations(max_allocations)

#         # Update Bar Chart Data
#         bar_chart_data = {
#             'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
#             'datasets': [{
#                 'label': 'Min Allocation',
#                 'data': min_allocations,
#                 'backgroundColor': 'skyblue'
#             },
#             {
#                 'label': 'Max Allocation',
#                 'data': max_allocations,
#                 'backgroundColor': 'lightgreen'
#             }]
#         }

#         # Similar changes can be made for the Pie Chart Data:
#         all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
#         num_labels = len(all_labels)
#         max_allocations_for_pie = normalize_allocations(
#             [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
#             [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
#         )
        
#         # Generate colors based on the number of labels
#         dynamic_colors = generate_colors(num_labels)

#         # Update Pie Chart Data
#         pie_chart_data = {
#             'labels': all_labels,
#             'datasets': [{
#                 'label': 'Investment Allocation',
#                 'data': max_allocations_for_pie,
#                 'backgroundColor': dynamic_colors,
#                 'hoverOffset': 4
#             }]
#         }
        
        
#         # Prepare the data for the line chart with inflation adjustment
#         initial_investment = 10000
#         # compounded_chart_data, inflation_adjusted_chart_data = prepare_line_chart_data_with_inflation(data_extracted, initial_investment)
#         combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
#         print(f"\nThe combined chart data is : {combined_chart_data}")
        
#         return jsonify({
#             "status": 200,
#             "message": "Success",
#             "investmentSuggestions": htmlSuggestions,
#             "pieChartData": pie_chart_data,
#             "barChartData": bar_chart_data,
#             "compoundedChartData":combined_chart_data
#         }), 200

#     except Exception as e:
#         logging.info(f"Error in generating investment suggestions: {e}")
#         return jsonify({'message': 'Internal Server Error in Generating responses'}), 500




# Run the Flask application
if __name__ == '__main__':
    app.run(host='0.0.0.0',debug=True)
