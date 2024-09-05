import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt

import os
import filetype
import docx
import PyPDF2
import re
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.memory import ConversationSummaryMemory
import asyncio
import numpy as np
import json

import google.generativeai as genai
import pathlib
import logging
import sys
import io
import matplotlib.pyplot as plt
import seaborn as sns
# Import things that are needed generically
from langchain.pydantic_v1 import BaseModel, Field
from langchain.tools import BaseTool, StructuredTool, tool
# Define functions to generate investment suggestions :

load_dotenv()
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

# Configure generativeai with your API key
genai.configure(api_key=GOOGLE_API_KEY)

def markdown_to_text(md):
    # Simple conversion for markdown to plain text
    md = md.replace('**', '')
    md = md.replace('*', '')
    md = md.replace('_', '')
    md = md.replace('#', '')
    md = md.replace('`', '')
    return md.strip()


# import docx

# def extract_responses_from_docx(personality_file):
#     """
#     Extracts responses from a Word document (.docx) where answers are typed in.

#     Args:
#         personality_file (UploadedFile): The file object uploaded via Streamlit.

#     Returns:
#         dict: A dictionary containing the questions and the typed answers.
#     """
#     try:
#         doc = docx.Document(personality_file)
#         responses = {}
#         current_question = None

#         # Check paragraphs
#         for para in doc.paragraphs:
#             text = para.text.strip()
#             if text:
#                 # Check if the paragraph contains a question
#                 if "?" in text or text.endswith(":"):
#                     current_question = text
#                     st.write(f"Identified question: {current_question}")  # Debugging log
#                 else:
#                     # This is a typed answer
#                     typed_answer = text.strip()
#                     st.write(f"Identified typed answer: {typed_answer}")  # Debugging log
#                     if current_question:
#                         # If the question already has an answer, append to it (handles multiple responses)
#                         if current_question in responses:
#                             responses[current_question] += "; " + typed_answer
#                         else:
#                             responses[current_question] = typed_answer

#             # Debugging log to understand document structure
#             st.write(f"Processing paragraph: {text}")  # Console log for local testing

#         # Check tables for additional responses
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     text = cell.text.strip()
#                     if text:
#                         if "?" in text or text.endswith(":"):
#                             current_question = text
#                             st.write(f"Identified question in table: {current_question}")  # Debugging log
#                         else:
#                             typed_answer = text.strip()
#                             st.write(f"Identified typed answer in table: {typed_answer}")  # Debugging log
#                             if current_question:
#                                 if current_question in responses:
#                                     responses[current_question] += "; " + typed_answer
#                                 else:
#                                     responses[current_question] = typed_answer

#         if responses:
#             st.write("Extracted Responses:")
#             for question, answer in responses.items():
#                 st.write(f"**{question}**: {answer}")
#         else:
#             st.write("No responses captured. Please check the document formatting or symbols used.")

#         return responses

#     except Exception as e:
#         st.write(f"Error extracting responses: {e}")  # Console log for local testing
#         return None

# def determine_investment_personality(responses):
#     """
#     Determines the investment personality based on extracted responses.

#     Args:
#         responses (dict): A dictionary containing the questions and the selected answers.

#     Returns:
#         str: The determined investment personality.
#     """
#     try:
#         # Prepare input text for the chatbot based on extracted responses
#         input_text = "User Profile:\n"
#         for question, response in responses.items():
#             input_text += f"{question}: {response}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYour task is to determine the investment personality based on the above profile."

#         # Here you would send the input_text to your chatbot or classification model
#         # For demonstration, we'll just return the input_text
#         return input_text

#     except Exception as e:
#         st.write(f"Error determining investment personality: {e}")  # Console log for local testing
#         return None

# def extract_responses_from_docx(personality_file):
#     try:
#         doc = docx.Document(personality_file)
#         responses = {}
#         current_question = None

#         # Check paragraphs
#         for para in doc.paragraphs:
#             text = para.text.strip()
#             if text:
#                 # Check if the paragraph contains a question
#                 if "?" in text or text.endswith(":"):
#                     current_question = text
#                 else:
#                     # This is a typed answer
#                     typed_answer = text.strip()
#                     if current_question:
#                         # If the question already has an answer, append to it (handles multiple responses)
#                         if current_question in responses:
#                             responses[current_question] += "; " + typed_answer
#                         else:
#                             responses[current_question] = typed_answer

#         # Check tables for additional responses
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     text = cell.text.strip()
#                     if text:
#                         if "?" in text or text.endswith(":"):
#                             current_question = text
#                         else:
#                             typed_answer = text.strip()
#                             if current_question:
#                                 if current_question in responses:
#                                     responses[current_question] += "; " + typed_answer
#                                 else:
#                                     responses[current_question] = typed_answer

#         return responses

#     except Exception as e:
#         print(f"Error extracting responses: {e}")
#         return None

import docx

# # GET Method for me POST method for Frontend
def extract_responses_from_docx(personality_file): # Using text responses parsing
    """
    Extracts responses from a Word document (.docx) where the selected answers are listed as text after the options.

    Args:
        personality_file (str): Path to the Word document file.

    Returns:
        dict: A dictionary containing the questions and the selected answers.
    """
    try:
        doc = docx.Document(personality_file)
        responses = {}
        current_question = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect the beginning of a question
                if "?" in text:
                    current_question = text
                # Detect a chosen response (assuming it follows the question and options)
                elif current_question and not text.startswith(("a.", "b.", "c.", "d.")):
                    selected_answer = text
                    responses[current_question] = selected_answer
                    current_question = None  # Reset for the next question

        if responses:
            print(responses)
            # st.write(responses)
        else:
            print("\nNo responses captured")
            st.write("No responses captured")
        return responses
    except Exception as e:
        print(f"Error extracting responses: {e}")
        return None

# def extract_responses_from_assessment(personality_file): # using boxes
#     # Load the document
#     # doc = Document(docx_filename)
#     doc = docx.Document(personality_file)
    
#     # Initialize a list to store responses
#     responses = []
    
#     # Iterate through each paragraph in the document
#     for para in doc.paragraphs:
#         text = para.text.strip()
#         # Check if the paragraph contains a checkbox
#         if '☒' in text or '☐' in text:
#             # Extract the response marked with ☒
#             if '☒' in text:
#                 response = text.split('☒')[1].strip()
#                 responses.append(response)
    
#     return responses

# import asyncio
# # from some_generative_ai_library import GenerativeModel  # Replace with actual import

# async def determine_investment_personality(assessment_data):
#     try:
#         # Prepare input text for the chatbot based on assessment data
#         input_text = "User Profile:\n"
#         for question, answer in assessment_data.items():
#             input_text += f"{question}: {answer}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
#                       "- Conservative Investor\n" \
#                       "- Moderate Investor\n" \
#                       "- Aggressive Investor\n\n" \
#                       "Please provide the classification below:\n"

#         # Use your generative AI model to generate a response
#         model = GenerativeModel('gemini-1.5-flash')
#         response = await model.generate_content(input_text)

#         # Determine the investment personality from the chatbot's response
#         response_text = response.text.lower()

#         if "conservative investor" in response_text:
#             personality = "Conservative Investor"
#         elif "moderate investor" in response_text:
#             personality = "Moderate Investor"
#         elif "aggressive investor" in response_text:
#             personality = "Aggressive Investor"
#         else:
#             personality = "Unknown"

#         return personality
#     except Exception as e:
#         print(f"Error generating response: {e}")
#         return "Unknown"


# GET Method
async def determine_investment_personality(assessment_data): # proper code 
    try:
        # Prepare input text for the chatbot based on assessment data
        input_text = "User Profile:\n"
        for question, answer in assessment_data.items():
            input_text += f"{question}: {answer}\n"

        # Introduce the chatbot's task and prompt for classification
        input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
                      "- Conservative Investor\n" \
                      "- Moderate Investor\n" \
                      "- Aggressive Investor\n\n" \
                      "Please provide the classification below:\n"

        # Use your generative AI model to generate a response
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(input_text)

        # Determine the investment personality from the chatbot's response
        response_text = response.text.lower()

        if "conservative investor" in response_text:
            personality = "Conservative Investor"
        elif "moderate investor" in response_text:
            personality = "Moderate Investor"
        elif "aggressive investor" in response_text:
            personality = "Aggressive Investor"
        else:
            personality = "Unknown"

        return personality
    except Exception as e:
        print(f"Error generating response: {e}")
        return "Unknown"



# #Load the Vector DataBase :
async def load_vector_db(file_path): # # GET Method 
    try:
        print("Loading vector database...")
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
        vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

        # vector_store = FAISS(
        #     embedding_function=embeddings,
        #     index=index,
        #     docstore=InMemoryDocstore(),
        #     index_to_docstore_id={},
        # )
        
        print("Vector database loaded successfully.") 
        return vector_store.as_retriever(search_kwargs={"k": 1})
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None


# async def load_vector_db(file_path="client_data.txt"):
#     try:
#         print("Loading vector database...")
#         with open(file_path, "r") as file:
#             text = file.read()
        
#         loader = Docx2txtLoader(file_path)
#         documents = loader.load()
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         text_chunks = text_splitter.split_documents(documents) #([Document(text=text)])
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        
#         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        
#         print("Vector database loaded successfully.") 
#         return vector_store.as_retriever(search_kwargs={"k": 1})
#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None


investment_personality = "Moderate Investor"
async def make_retrieval_chain(retriever): # GET Method
    """
    Create a retrieval chain using the provided retriever.

    Args:
        retriever (RetrievalQA): A retriever object.

    Returns:
        RetrievalQA: A retrieval chain object.
    """
    try:
        global investment_personality #,summary
        llm = ChatGoogleGenerativeAI(
            #model="gemini-pro",
            model = "gemini-1.5-flash",
            temperature=0.7,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )
                                                    # \n + summary 
        prompt_template = investment_personality +  "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
                Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality.
                Also give the user detailed information about the investment how to invest,where to invest and how much they
                should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
                Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
                investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
                Also give the user minimum and maximum expected growth in dollars for the time horizon .
                Also explain the user why you are giving them that particular investment suggestion.
                Here's an example for the required Output Format :

                Investment Suggestions for a Moderate Investor(This is for a Moderate Investor but you need to generate for any investor)

                Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

                Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

                Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
                How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
                Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                Consider investing in blue-chip companies or growth sectors like technology. 
                Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


                Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
                Percentage Allocation: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

                Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
                How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
                Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
                Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
                Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
                Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
                Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
                Percentage Allocation: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

                Time Horizon and Expected Returns:

                Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                Minimum Expected Annual Return: 4% - 6% 
                Maximum Expected Annual Return: 8% - 10% 
                Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, a 10,000 investment could grow to approximately 17,908 in 10 years.
                Minimum Expected Growth in Dollars: 
                
                4,000−6,000 (over 10 years) Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
                Rationale for Investment Suggestions:

                This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
                Important Considerations:

                Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

                Inflation Adjusted Returns:(do not write this part inside the bracket just give answer,assume US inflation rate, and give the investment returns value that was suggested by you  for $10k investment after 3,5,10years of growth  mention the values before adjusting and after adjusting with inflation )

                Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

                <context>
                {context}
                </context>
                Question: {input}"""

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in creating chain: {e}")
        return None

import streamlit as st
import json
import matplotlib.pyplot as plt
import io



# async def process_document(file_path):
#     try:
#         print("Processing the document")
#         file_type = filetype.guess(file_path)
#         if file_type is not None:
#             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 # return extract_text_from_word(file_path)
#                 return extract_text_and_tables_from_word(file_path)
#             elif file_type.mime == "application/pdf":
#                 return extract_text_from_pdf(file_path)
#         return None
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None


# async def extract_text_from_pdf(pdf_file_path):
#     try:
#         print("Processing pdf file")
#         with open(pdf_file_path, "rb") as pdf_file:
#             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#             text_content = []
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text_content.append(page.extract_text())
#             return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None



# async def extract_text_and_tables_from_word(docx_file_path):
#     """
#     Extracts text and tables from a Word document (.docx).

#     Args:
#         docx_file_path (str): Path to the Word document file.

#     Returns:
#         tuple: Extracted text content and tables from the document.
#     """
#     try:
#         print("Extracting text and tables from word file")
#         doc = docx.Document(docx_file_path)
#         text_content = []
#         tables_content = []

#         for para in doc.paragraphs:
#             text_content.append(para.text)

#         for table in doc.tables:
#             table_data = []
#             for row in table.rows:
#                 row_data = []
#                 for cell in row.cells:
#                     row_data.append(cell.text.strip())
#                 table_data.append(row_data)
#             tables_content.append(table_data)
#         print("Extracted text from word file")
#         return "\n".join(text_content), tables_content
#     except Exception as e:
#         print(f"Error extracting text and tables from Word document: {e}")
#         return None, None


# def parse_financial_data(text_content, tables_content):
#     assets = {}
#     liabilities = {}

#     # Extract assets and liabilities from tables
#     if tables_content:
#         for table in tables_content:
#             for row in table:
#                 # Basic keyword matching (e.g., 'Asset' or 'Liability' detection)
#                 if any("asset" in cell.lower() for cell in row):
#                     key = row[0]
#                     value = float(row[1].replace('$', '').replace(',', ''))
#                     assets[key] = value
#                 elif any("liability" in cell.lower() for cell in row):
#                     key = row[0]
#                     value = float(row[1].replace('$', '').replace(',', ''))
#                     liabilities[key] = value

#     # Additional parsing logic can be added here based on the document structure

#     return assets, liabilities

# # Step 3: Plot pie chart
# def plot_pie_chart(assets, liabilities):
#     total_assets = sum(assets.values())
#     total_liabilities = sum(liabilities.values())

#     labels = ['Assets', 'Liabilities']
#     sizes = [total_assets, total_liabilities]

#     fig, ax = plt.subplots(figsize=(8, 8))
#     ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=['green', 'red'])
#     ax.set_title('Assets vs Liabilities Distribution')

#     st.pyplot(fig)

async def process_document(file_path): # GET Method
    try:
        print("Processing the document")
        file_type = filetype.guess(file_path)
        if file_type is not None:
            if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Await the coroutine to extract text and tables
                return await extract_text_and_tables_from_word(file_path)
            elif file_type.mime == "application/pdf":
                return await extract_text_from_pdf(file_path)
        return None
    except Exception as e:
        print(f"Error processing document: {e}")
        return None

# Async function to extract text from a PDF file
async def extract_text_from_pdf(pdf_file_path): # GET Method
    try:
        print("Processing pdf file")
        with open(pdf_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            text_content = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text_content.append(page.extract_text())
            return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

# Async function to extract text and tables from a Word document
async def extract_text_and_tables_from_word(docx_file_path): # GET Method
    try:
        print("Extracting text and tables from word file")
        doc = docx.Document(docx_file_path)
        text_content = []
        tables_content = []

        for para in doc.paragraphs:
            text_content.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)
        print("Extracted text from word file")
        return "\n".join(text_content), tables_content
    except Exception as e:
        print(f"Error extracting text and tables from Word document: {e}")
        return None, None



async def validate_document_content(text, tables):
    """
    Validates the content of the document.

    Args:
        text (str): Extracted text content from the document.
        tables (list): Extracted tables content from the document.

    Returns:
        tuple: Client name and validation errors.
    """
    errors = []
    
    # Extract client name
    client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
    client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

    # Define required sections
    required_sections = [
        "YOUR RETIREMENT GOAL",
        "YOUR OTHER MAJOR GOALS",
        "YOUR ASSETS AND LIABILITIES",
        "MY LIABILITIES",
        "YOUR CURRENT ANNUAL INCOME"
    ]

    # Check for the presence of required sections
    for section in required_sections:
        if section not in text:
            errors.append(f"* {section} section missing.")
    
    # Define table field checks
    table_checks = {
        "YOUR RETIREMENT GOAL": [
            r"When do you plan to retire\? \(age or date\)",
            r"Social Security Benefit \(include expected start date\)",
            r"Pension Benefit \(include expected start date\)",
            r"Other Expected Income \(rental, part-time work, etc.\)",
            r"Estimated Annual Retirement Expense"
        ],
        "YOUR OTHER MAJOR GOALS": [
            r"GOAL", r"COST", r"WHEN"
        ],
        "YOUR ASSETS AND LIABILITIES": [
            r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
            r"Current Value", r"Annual Contributions"
        ],
        "MY LIABILITIES": [
            r"Balance", r"Interest Rate", r"Monthly Payment"
        ]
    }

    # Validate table content
    for section, checks in table_checks.items():
        section_found = False
        for table in tables:
            table_text = "\n".join(["\t".join(row) for row in table])
            if section in table_text:
                section_found = True
                for check in checks:
                    if not re.search(check, table_text, re.IGNORECASE):
                        errors.append(f"* Missing or empty field in {section} section: {check}")
                break
        if not section_found:
            errors.append(f"* {section} section missing.")

    return client_name, errors


# RUN Button :
async def generate_investment_suggestions(investment_personality, context): # # GET Method for py , for front end its Post API
    
    # retriever = asyncio.run(load_vector_db("uploaded_file"))

    retriever = await load_vector_db("uploaded_file")
    # retriever = await load_vector_db("data\Financial_Investment_1.docx") 

    chain = await make_retrieval_chain(retriever)

    # chain = asyncio.run(make_retrieval_chain(retriever))
    
    if chain is not None:
        # summary = context
        # query = summary + "\n" + investment_personality
        query = str(investment_personality)
        response = chain.invoke({"input": query})
        format_response = markdown_to_text(response['answer'])
        return format_response
        # st.write(format_response)

        # handle_graph(response['answer'])

    else:
        st.error("Failed to create the retrieval chain. Please upload a valid document.")



# Generate Infographics : Best Code so far:


import seaborn as sns
import re
from collections import defaultdict
import matplotlib.pyplot as plt
import streamlit as st
import numpy as np

def extract_numerical_data(response):
    # Define patterns to match different sections and their respective allocations
    patterns = {
        'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
        'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
        'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
    }

    data = defaultdict(dict)

    for section, pattern in patterns.items():
        match = pattern.search(response)
        if match:
            investments_text = match.group(1)
            # Extract individual investment types and their allocations
            investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
            for investment_match in investment_pattern.findall(investments_text):
                investment_type, min_allocation, max_allocation = investment_match
                data[section][investment_type.strip()] = {
                    'min': min_allocation,
                    'max': max_allocation
                }

    # Extract time horizon and expected returns
    time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
    min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
    max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
    min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
    max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

    time_horizon_match = time_horizon_pattern.search(response)
    min_return_match = min_return_pattern.search(response)
    max_return_match = max_return_pattern.search(response)
    min_growth_match = min_growth_pattern.search(response)
    max_growth_match = max_growth_pattern.search(response)

    if time_horizon_match:
        data['Time Horizon'] = {
            'min_years': time_horizon_match.group(1),
            'max_years': time_horizon_match.group(2)
        }

    if min_return_match:
        data['Expected Annual Return'] = {
            'min': min_return_match.group(1),
            'max': min_return_match.group(2)
        }

    if max_return_match:
        data['Expected Annual Return'] = {
            'min': max_return_match.group(1),
            'max': max_return_match.group(2)
        }

    if min_growth_match:
        data['Expected Growth in Dollars'] = {
            'min': min_growth_match.group(1),
            'max': min_growth_match.group(2)
        }

    if max_growth_match:
        data['Expected Growth in Dollars'] = {
            'min': max_growth_match.group(1),
            'max': max_growth_match.group(2)
        }

    return data


def plot_investment_allocations(data):
    # Create subplots with a large figure size
    fig, axes = plt.subplots(2, 1, figsize= (16,10)) #(28, 15))  # Adjust size as needed

    # Plot Growth-Oriented Investments
    growth_data = data['Growth-Oriented Investments']
    growth_labels = list(growth_data.keys())
    growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
    growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

    axes[0].bar(growth_labels, growth_min, color='skyblue', label='Min Allocation')
    axes[0].bar(growth_labels, growth_max, bottom=growth_min, color='lightgreen', label='Max Allocation')
    axes[0].set_title('Growth-Oriented Investments', fontsize=16)
    axes[0].set_ylabel('Percentage Allocation', fontsize=14)
    axes[0].set_xlabel('Investment Types', fontsize=14)
    axes[0].tick_params(axis='x', rotation=45, labelsize=12)
    axes[0].tick_params(axis='y', labelsize=12)
    axes[0].legend()

    # Plot Conservative Investments
    conservative_data = data['Conservative Investments']
    conservative_labels = list(conservative_data.keys())
    conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
    conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

    axes[1].bar(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
    axes[1].bar(conservative_labels, conservative_max, bottom=conservative_min, color='lightgreen', label='Max Allocation')
    axes[1].set_title('Conservative Investments', fontsize=16)
    axes[1].set_ylabel('Percentage Allocation', fontsize=14)
    axes[1].set_xlabel('Investment Types', fontsize=14)
    axes[1].tick_params(axis='x', rotation=45, labelsize=12)
    axes[1].tick_params(axis='y', labelsize=12)
    axes[1].legend()

    # Tight layout for better spacing
    plt.tight_layout()
    plt.show()
    return fig


def plot_pie_chart(data):
    fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

    # Combine all investment data for pie chart
    all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
    labels = list(all_data.keys())
    sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
    colors = plt.cm.Paired(range(len(labels)))

    wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    ax.set_title('Investment Allocation')

    # Add legend
    ax.legend(wedges, labels, title="Investment Types", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

    return fig



def bar_chart(data):
    fig, ax = plt.subplots(figsize=(12, 8))  # Increased size

    # Data for plotting
    categories = list(data.keys())
    values_min = [int(data[cat]['min'].strip('%')) for cat in categories]
    values_max = [int(data[cat]['max'].strip('%')) for cat in categories]

    x = range(len(categories))

    ax.bar(x, values_min, width=0.4, label='Min Allocation', color='skyblue', align='center')
    ax.bar(x, values_max, width=0.4, label='Max Allocation', color='lightgreen', align='edge')

    ax.set_xticks(x)
    ax.set_xticklabels(categories, rotation=45, ha='right')
    ax.set_xlabel('Investment Categories')
    ax.set_ylabel('Percentage Allocation')
    ax.set_title('Investment Allocation')
    ax.legend()

    plt.tight_layout()
    return fig



import plotly.graph_objects as go
import streamlit as st
import numpy as np

def plot_3d_bar_graph(data):
    # Initialize a Plotly figure
    fig = go.Figure()

    if not data:
        st.write("No data available to plot.")
        return

    # Extracting data for plotting
    x = []
    y1 = []  # Min values for plotting
    y2 = []  # Max values for plotting
    categories = []

    for i, (key, value) in enumerate(data.items()):
        categories.append(key)  # Categories
        x.append(i)  # X-axis value
        min_val = int(value.get('min', '0').replace('%', ''))  # Default to 0 if 'min' is missing
        max_val = int(value.get('max', '0').replace('%', ''))  # Default to 0 if 'max' is missing
        y1.append(min_val)
        y2.append(max_val)

    # Adding bars for Min Compounded Returns
    for i in range(len(x)):
        fig.add_trace(go.Scatter3d(
            x=[x[i], x[i]],
            y=[0, 0.5],
            z=[0, y1[i]],
            mode='lines',
            line=dict(color='red', width=10),
            name='Min Compounded Returns'
        ))

    # Adding bars for Max Compounded Returns
    for i in range(len(x)):
        fig.add_trace(go.Scatter3d(
            x=[x[i] + 0.5, x[i] + 0.5],
            y=[0, 0.5],
            z=[0, y2[i]],
            mode='lines',
            line=dict(color='blue', width=10),
            name='Max Compounded Returns'
        ))

    # Update layout to match the desired appearance
    fig.update_layout(
        scene=dict(
            xaxis=dict(
                tickvals=np.arange(len(categories)) + 0.25,
                ticktext=categories,
                title='Investment Types'
            ),
            yaxis=dict(title=''),
            zaxis=dict(title='Allocation (%)')
        ),
        title='Investment Allocation 3D Bar Graph',
        legend=dict(x=0.1, y=0.9)
    )

    st.plotly_chart(fig)




# def plot_3d_bar_graph(data):
#     # Initialize a Plotly figure
#     fig = go.Figure()

#     if not data:
#         print("No data available to plot.")
#         return

#     # Extracting data for plotting
#     x = []
#     y1 = []  # Min values for plotting
#     y2 = []  # Max values for plotting
#     categories = []

#     for i, (key, value) in enumerate(data.items()):
#         categories.append(key)  # Categories
#         x.append(i)  # X-axis value
#         min_val = int(value.get('min', '0').replace('%', ''))  # Default to 0 if 'min' is missing
#         max_val = int(value.get('max', '0').replace('%', ''))  # Default to 0 if 'max' is missing
#         y1.append(min_val)
#         y2.append(max_val)

#     # Adding bars for Min Compounded Returns
#     for i in range(len(x)):
#         fig.add_trace(go.Scatter3d(
#             x=[x[i], x[i]],
#             y=[0, 0.5],
#             z=[0, y1[i]],
#             mode='lines',
#             line=dict(color='red', width=10),
#             name='Min Compounded Returns'
#         ))

#     # Adding bars for Max Compounded Returns
#     for i in range(len(x)):
#         fig.add_trace(go.Scatter3d(
#             x=[x[i] + 0.5, x[i] + 0.5],
#             y=[0, 0.5],
#             z=[0, y2[i]],
#             mode='lines',
#             line=dict(color='blue', width=10),
#             name='Max Compounded Returns'
#         ))

#     # Update layout to match the desired appearance
#     fig.update_layout(
#         scene=dict(
#             xaxis=dict(
#                 tickvals=np.arange(len(categories)) + 0.25,
#                 ticktext=categories,
#                 title='Investment Types'
#             ),
#             yaxis=dict(title=''),
#             zaxis=dict(title='Allocation (%)')
#         ),
#         title='Investment Allocation 3D Bar Graph',
#         legend=dict(x=0.1, y=0.9)
#     )

#     fig.show()
#     return fig


 
# def client_form():
#     st.title("Client Details Form")

#     with st.form("client_form"):
#         st.header("Personal Information")
#         client_name = st.text_input("Client Name")
#         co_client_name = st.text_input("Co-Client Name")
#         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
#         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
#         today_date = st.date_input("Today's Date")
        
#         st.header("Financial Information")
#         current_assets = st.text_area("Current Assets (e.g., type and value)")
#         liabilities = st.text_area("Liabilities (e.g., type and amount)")
#         annual_income = st.text_area("Current Annual Income (source and amount)")
#         annual_contributions = st.text_area("Annual Contributions (e.g., retirement savings)")

#         st.header("Insurance Information")
#         life_insurance = st.text_input("Life Insurance (e.g., coverage amount)")
#         disability_insurance = st.text_input("Disability Insurance (e.g., coverage amount)")
#         long_term_care = st.text_input("Long-Term Care Insurance (e.g., coverage amount)")

#         st.header("Estate Planning")
#         will_status = st.radio("Do you have a will?", ["Yes", "No"])
#         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
#         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
#         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

#         # Submit button
#         submitted = st.form_submit_button("Submit")

#         if submitted:
#             # Save form data
#             form_data = {
#                 "Client Name": client_name,
#                 "Co-Client Name": co_client_name,
#                 "Client Age": client_age,
#                 "Co-Client Age": co_client_age,
#                 "Today's Date": str(today_date),
#                 "Current Assets": current_assets,
#                 "Liabilities": liabilities,
#                 "Annual Income": annual_income,
#                 "Annual Contributions": annual_contributions,
#                 "Life Insurance": life_insurance,
#                 "Disability Insurance": disability_insurance,
#                 "Long-Term Care Insurance": long_term_care,
#                 "Will Status": will_status,
#                 "Trust Status": trust_status,
#                 "Power of Attorney": power_of_attorney,
#                 "Healthcare Proxy": healthcare_proxy,
#             }
            
#             # Save to a file or database
#             with open("client_data.txt", "a") as f:
#                 f.write(str(form_data) + "\n")
            
#             st.success("Form submitted successfully!")
#             st.session_state.page = "main"  # Redirect back to main page after form submission


from datetime import date  # Make sure to import the date class


# Function to parse financial data from the text
import re

def parse_financial_data(text_content):
    assets = []
    liabilities = []

    # Define regex patterns to capture text following headings
    asset_pattern = re.compile(r"MY ASSETS:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)
    liability_pattern = re.compile(r"LIABILITIES:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)

    # Extract assets
    asset_matches = asset_pattern.findall(text_content)
    if asset_matches:
        asset_text = asset_matches[0]
        # Further processing to extract individual asset values if they are detailed
        asset_lines = asset_text.split('\n')
        for line in asset_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                asset_value = float(match.group().replace(",", ""))
                assets.append(asset_value)

    # Extract liabilities
    liability_matches = liability_pattern.findall(text_content)
    if liability_matches:
        liability_text = liability_matches[0]
        # Further processing to extract individual liability values if they are detailed
        liability_lines = liability_text.split('\n')
        for line in liability_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                liability_value = float(match.group().replace(",", ""))
                liabilities.append(liability_value)

    print("Assets Found:", assets)
    print("Liabilities Found:", liabilities)

    return assets, liabilities



# Function to extract numerical values from a text input
def extract_numeric(value):
    try:
        return float(re.sub(r'[^\d.]', '', value))  # Remove non-numeric characters and convert to float
    except ValueError:
        return 0


# plots graph from the details of the form :


def is_float(value):
    try:
        float(value)
        return True
    except ValueError:
        return False


# def plot_assets_liabilities_pie_chart(assets, liabilities, threshold=5): two pie charts but legends overlap
#     """
#     Plots separate pie charts for assets and liabilities. If there are any categories
#     below a specified threshold, they are plotted in an additional small pie chart.
    
#     Parameters:
#     - assets: dict, keys are asset names, values are their amounts.
#     - liabilities: dict, keys are liability names, values are their amounts.
#     - threshold: int, percentage threshold below which segments are considered small.
#     """
#     # Update matplotlib settings to increase the font size globally
#     plt.rcParams.update({'font.size': 16})

#     def plot_pie(data, title):
#         # Filter out zero values and create a summary for small segments
#         total = sum(data.values())
#         filtered_data = {k: v for k, v in data.items() if (v / total) >= threshold / 100}
#         small_segments = {k: v for k, v in data.items() if (v / total) < threshold / 100}
#         small_total = sum(small_segments.values())

#         # Plotting logic
#         if small_segments:
#             fig, (ax_main, ax_small) = plt.subplots(1, 2, figsize=(15, 8))  # Side-by-side layout with larger size
#         else:
#             fig, ax_main = plt.subplots(figsize=(15, 15))  # Only main chart with larger size

#         # Plot main pie chart
#         labels_main = list(filtered_data.keys()) + ([f"Other small {title}"] if small_segments else [])
#         values_main = list(filtered_data.values()) + ([small_total] if small_segments else [])
#         wedges_main, texts_main, autotexts_main = ax_main.pie(
#             values_main, labels=labels_main, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
#             startangle=140, textprops={'fontsize': 18}  # Larger font size for labels
#         )

#         ax_main.set_title(title, fontsize=20)
#         # Position legend to the right of the plot to avoid overlapping
#         ax_main.legend(wedges_main, labels_main, title="Categories", loc="upper left", bbox_to_anchor=(1, 1), fontsize=14)

#         if small_segments:
#             # Plot additional small pie chart for small segments
#             labels_small = list(small_segments.keys())
#             values_small = list(small_segments.values())
#             wedges_small, texts_small, autotexts_small = ax_small.pie(
#                 values_small, labels=labels_small, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
#                 startangle=140, textprops={'fontsize': 14}  # Consistent label size for small chart
#             )
#             ax_small.set_title(f"Small Segments of {title}", fontsize=20)
#             # Position legend below the small pie chart to avoid overlapping
#             ax_small.legend(wedges_small, labels_small, title="Small Categories", loc="upper right", bbox_to_anchor=(1, 0.5), fontsize=12)

#         st.pyplot(fig)

#     # Convert valid entries to float, ensuring only numeric values are considered
#     assets = {k: float(v) for k, v in assets.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}
#     liabilities = {k: float(v) for k, v in liabilities.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}

#     # Plot pie charts
#     plot_pie(assets, 'Distribution of Assets')
#     plot_pie(liabilities, 'Distribution of Liabilities')


def plot_assets_liabilities_pie_chart(assets, liabilities, threshold=50): # best plot 
    """
    Plots separate pie charts for assets and liabilities. If there are any categories
    below a specified threshold, they are plotted in an additional small pie chart.
    
    Parameters:
    - assets: dict, keys are asset names, values are their amounts.
    - liabilities: dict, keys are liability names, values are their amounts.
    - threshold: int, percentage threshold below which segments are considered small.
    """
    # Update matplotlib settings to increase the font size globally
    # plt.rcParams.update({'font.size': 32})

    plt.rcParams.update({'font.size': 16})

    def plot_pie(data, title):
        # Filter out zero values and create a summary for small segments
        total = sum(data.values())
        filtered_data = {k: v for k, v in data.items() if (v / total) >= threshold / 100}
        small_segments = {k: v for k, v in data.items() if (v / total) < threshold / 100}
        small_total = sum(small_segments.values())

        # Plotting logic
        if small_segments:
            fig, (ax_main, ax_small) = plt.subplots(1, 2, figsize=(30, 15))  # Side-by-side layout
        else:
            fig, ax_main = plt.subplots(figsize=(30, 20))  # Only main chart with larger size

            # fig, ax_main = plt.subplots(figsize=(10, 10))  # Only main chart with larger size

        # Plot main pie chart
        labels_main = list(filtered_data.keys()) + ([f"Other small {title}"] if small_segments else [])
        values_main = list(filtered_data.values()) + ([small_total] if small_segments else [])
        wedges_main, texts_main, autotexts_main = ax_main.pie(
            values_main, labels=labels_main, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
            startangle=140, textprops={'fontsize': 28} #18}  # Larger font size for labels
        )

        ax_main.set_title(title, fontsize=20)
        # Position legend to the right of the plot to avoid overlapping
        ax_main.legend(wedges_main, labels_main, title="Categories", loc="upper right", bbox_to_anchor=(0.001, 0.9), fontsize= 28)#14)

        if small_segments:
            # Plot additional small pie chart for small segments
            labels_small = list(small_segments.keys())
            values_small = list(small_segments.values())
            wedges_small, texts_small, autotexts_small = ax_small.pie(
                values_small, labels=labels_small, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
                startangle=140, textprops={'fontsize': 24} #14}  # Consistent label size for small chart
            )
            ax_small.set_title(f"Small Segments of {title}", fontsize=20)
            # Position legend to the right of the small pie chart but slightly lower to avoid overlap with the main chart's legend
            ax_small.legend(wedges_small, labels_small, title="Small Categories", loc="center left", bbox_to_anchor=(1.2, 0.3), fontsize= 22)#12)

        st.pyplot(fig)

    # Convert valid entries to float, ensuring only numeric values are considered
    assets = {k: float(v) for k, v in assets.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}
    liabilities = {k: float(v) for k, v in liabilities.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}

    # Plot pie charts
    plot_pie(assets, 'Distribution of Assets')
    plot_pie(liabilities, 'Distribution of Liabilities')

# def plot_assets_liabilities_pie_chart(assets, liabilities):# properly plots a big and 1 small pie chart for both assets and liability
#     # Filter and convert values to float, handle non-numeric or empty inputs
#     filtered_assets = {k: float(v) for k, v in assets.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}
#     filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}

#     # Combine assets and liabilities for total calculation
#     all_values = {**filtered_assets, **filtered_liabilities}
#     total_value = sum(all_values.values())

#     # Separate main and small segments
#     main_segments = {k: v for k, v in all_values.items() if (v / total_value) >= 0.05}
#     small_segments = {k: v for k, v in all_values.items() if (v / total_value) < 0.05}
#     small_total = sum(small_segments.values())

#     # Prepare data for main pie chart
#     main_labels = list(main_segments.keys()) + (["Others"] if small_segments else [])
#     main_values = list(main_segments.values()) + ([small_total] if small_segments else [])

#     # Prepare data for small pie chart (only if there are small segments)
#     small_labels = list(small_segments.keys())
#     small_values = list(small_segments.values())

#     fig, ax = plt.subplots(figsize=(8, 6))

#     # Plot main pie chart
#     wedges, texts, autotexts = ax.pie(
#         main_values,
#         labels=main_labels,
#         autopct='%1.1f%%',
#         startangle=140,
#         colors=plt.cm.Paired.colors,
#     )

#     # Explode the "Others" slice
#     if small_segments:
#         others_index = main_labels.index("Others")
#         wedges[others_index].set_edgecolor('white')
#         # wedges[others_index].set_linestyle('--')
#         wedges[others_index].set_linewidth(2)
#         wedges[others_index].set_hatch('/')

#     ax.set_title('Assets and Liabilities Distribution')

#     # Draw a second pie chart for "Others"
#     if small_segments:
#         fig2, ax2 = plt.subplots(figsize=(8, 6))
#         wedges_small, texts_small, autotexts_small = ax2.pie(
#             small_values,
#             labels=small_labels,
#             autopct='%1.1f%%',
#             startangle=140,
#             colors=plt.cm.Pastel1.colors
#         )

#         ax2.set_title('Detailed View of "Others" Categories')

#     plt.tight_layout()
#     st.pyplot(fig)
#     if small_segments:
#         st.pyplot(fig2)



def save_data_to_file(form_data):
    file_path = 'client_data.txt'
    with open(file_path, 'a') as file:
        file.write(str(form_data) + "\n")
    # st.success(f"Form data saved to {file_path}")
    print(f"Form data saved to {file_path}")
    

def client_form():
    st.title("Client Details Form")

    with st.form("client_form"):
        st.header("Personal Information")
        client_name = st.text_input("Client Name")
        co_client_name = st.text_input("Co-Client Name")
        client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
        co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
        today_date = st.date_input("Today's Date")

        st.header("Your Assets (in $)")

        assets = {
            # 'Annual Income': st.text_input("Annual Income (e.g. , Your Annual Salary Income or other source of income) "),
            'Cash/Bank Account': st.text_input("Cash/Bank Account"),
            '401(k), 403(b), 457 Plans': st.text_input("Your 401(k), 403(b), 457 Plans "),
            'Traditional, SEP and SIMPLE IRAs': st.text_input("Traditional, SEP and SIMPLE IRAs "),
            'Roth IRA,Roth 401(k)': st.text_input("Roth IRA, Roth 401(k)"),
            'Brokerage/non-qualified accounts': st.text_input("Brokerage/non-qualified accounts"),
            'Annuities': st.text_input("Annuities"),
            '529 Plans': st.text_input("529 Plans"),
            'Home': st.text_input("Home"),
            'Other Real Estate': st.text_input("Other Real Estate"),
            'Business': st.text_input("Business"),
            'Other': st.text_input("Other")
        }
        st.header("Your Liabilities (in $)")

        liabilities = {
            'Mortgage': st.text_input("Mortgage"),
            # 'Annual Mortgage Interest Rate': st.number_input("Annual Mortgage Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=12.0, step=0.5),
            # 'Mortagage Time Period': st.number_input("Mortagage Time Period (Mention the time period of the Mortgage in years)", min_value=0, max_value=100,value=10,step=1),

            'Home Loans': st.text_input("Home Loans"),
            # 'Home Loans Interest Rate': st.number_input("Home Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=10.0, step=0.5),
            # 'Home Loans Time Period': st.number_input("Home Loans Time Period (Mention the time period of the Home Loan in years)", min_value=0, max_value=100,value=15,step=1),

            'Vehicle Loans': st.text_input("Vehicle Loans"),
            # 'Vehicle Loans Interest Rate': st.number_input("Vehicle Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
            # 'Vehicle Loans Time Period': st.number_input("Vehicle Loans Time Period (Mention the time period of the Car/Vehicle Loan in years)", min_value=0, max_value=100,value=15,step=1),

            'Education Loans': st.text_input("Education Loans"),
            # 'Education Loans Interest Rate' : st.number_input("Education Loans Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
            # 'Education Loans Time Period': st.number_input("Education Loans Time Period (Mention the time period of the Education Loan in years)", min_value=0, max_value=100,value=15,step=1),

            # 'Credit Card': st.text_input("Monthly Credit Card Debt (Mention Amount)"),
            # 'Credit Card Debt Interest Rate': st.number_input("Credit Card Debt Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),

            'Miscellaneous': st.text_input("Miscellaneous"),
        }

        st.header("Your Retirement Goal")
        retirement_age = st.number_input("At what age do you plan to retire?", min_value=0, max_value=120, value=65, step=1)
        retirement_income = st.text_input("Desired annual retirement income")

        st.header("Your Other Goals")
        goal_name = st.text_input("Name of the Goal (e.g . , Dream House, Travel, Educational, etc.)")
        goal_amount = st.text_input("Amount needed for the goal (in $)")
        goal_timeframe = st.number_input("Timeframe to achieve the goal (in years)", min_value=0, max_value=100, value=5, step=1)

        st.header("Insurance Information")
        life_insurance_Benefit = st.text_input("Life Insurance-Benefit")
        life_insurance_Premium = st.text_input("Life Insurance-Premium")
        disability_insurance_Benefit = st.text_input("Disability Insurance-Benefit")
        disability_insurance_Premium = st.text_input("Disability Insurance-Premium")
        long_term_care_benefit = st.text_input("Long-Term Care Insurance-Benefit")
        long_term_care_premium = st.text_input("Long-Term Care Insurance-Premium")


        st.header("Estate Planning")
        will_status = st.radio("Do you have a will?", ["Yes", "No"])
        trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
        power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
        healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

        submitted = st.form_submit_button("Submit")

        if submitted:
            form_data = {
                "Client Name": client_name,
                "Co-Client Name": co_client_name,
                "Client Age": client_age,
                "Co-Client Age": co_client_age,
                "Today's Date": str(today_date),
                "Assets": assets,
                "Liabilities": liabilities,
                "Retirement Age": retirement_age,
                "Desired Retirement Income": retirement_income,
                "Goal Name": goal_name,
                "Goal Amount": goal_amount,
                "Goal Timeframe": goal_timeframe,
                "Life Insurance Benefit": life_insurance_Benefit,
                "Life Insurance Premium": life_insurance_Premium,
                "Disability Insurance Benefit": disability_insurance_Benefit,
                "Disability Insurance Premium": disability_insurance_Premium,
                "Long-Term Care Insurance Benefit": long_term_care_benefit,
                "Long-Term Care Insurance Premium": long_term_care_premium,
                "Will Status": will_status,
                "Trust Status": trust_status,
                "Power of Attorney": power_of_attorney,
                "Healthcare Proxy": healthcare_proxy,
            }

            save_data_to_file(form_data)
            
            # # Plot the pie chart
            # st.subheader("Assets and Liabilities Breakdown")
            # plot_assets_liabilities_pie_chart(assets, liabilities)

            # Store data in session state and redirect to main
            st.session_state.assets = assets
            st.session_state.liabilities = liabilities
            st.session_state.total_assets, st.session_state.total_liabilities = calculate_totals(assets, liabilities)
            st.session_state.page = "main"
            st.success("Data submitted!\nThank You for filling the form !\nReturning to main portal...")

import math
def calculate_compounded_amount(principal, rate, time):
    """
    Calculates the compounded amount using the formula:
    A = P * (1 + r/n)^(nt)
    Assuming n (compounding frequency) is 1 for simplicity (annually).
    """
    if principal == 0 or rate == 0 or time == 0:
        return principal
    else:
        # Using annual compounding
        return principal * (1 + rate / 100) ** time
    
def calculate_totals(assets, liabilities):
    total_assets = sum(extract_numeric(v) for v in assets.values())
    print(f"Total Assets : {total_assets}")
    total_liabilities = 0
    total_liabilities = sum(extract_numeric(v) for v in liabilities.values() )

    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Mortgage']),
    #     liabilities['Annual Mortgage Interest Rate'],
    #     liabilities['Mortagage Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Home Loans']),
    #     liabilities['Home Loans Interest Rate'],
    #     liabilities['Home Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Vehicle Loans']),
    #     liabilities['Vehicle Loans Interest Rate'],
    #     liabilities['Vehicle Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Education Loans']),
    #     liabilities['Education Loans Interest Rate'],
    #     liabilities['Education Loans Time Period']
    # )
    
    # For credit card debt, only calculate compounded amount if interest rate > 0

    # credit_card_balance = extract_numeric(liabilities['Credit Card'])
    # credit_card_interest = liabilities['Credit Card Debt Interest Rate']
    # if credit_card_interest > 0:
    #     # Assuming the time period for credit card debt is 1 year for compounding
    #     total_liabilities += calculate_compounded_amount(credit_card_balance, credit_card_interest, 1)
    # else:
    #     total_liabilities += credit_card_balance
    
    # Miscellaneous debts are taken directly as is
    total_liabilities += extract_numeric(liabilities['Miscellaneous'])
    rounded_liabilities = round(total_liabilities,2)

    print(f"Total liabilities :{total_liabilities}")
    print(f"Rounded of Total liabilities :{rounded_liabilities}")

    return total_assets, rounded_liabilities #total_liabilities

def create_financial_summary_table(assets, liabilities):
    # Filter out items with zero value
    filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
    filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

    # Create DataFrames for assets and liabilities with indices starting from 1
    assets_df = pd.DataFrame(
        list(filtered_assets.items()), 
        columns=['Assets', 'Amount ($)'], 
        index=range(1, len(filtered_assets) + 1)
    )
    liabilities_df = pd.DataFrame(
        list(filtered_liabilities.items()), 
        columns=['Liabilities', 'Amount ($)'], 
        index=range(1, len(filtered_liabilities) + 1)
    )

    # Calculate total
    total_assets, total_liabilities = calculate_totals(assets, liabilities)

    # Add total row with index incremented by 1
    total_assets_row = pd.DataFrame(
        [['TOTAL', total_assets]], 
        columns=['Assets', 'Amount ($)'], 
        index=[len(assets_df) + 1]
    )
    total_liabilities_row = pd.DataFrame(
        [['TOTAL', total_liabilities]], 
        columns=['Liabilities', 'Amount ($)'], 
        index=[len(liabilities_df) + 1]
    )

    # Append total rows to DataFrames
    assets_df = pd.concat([assets_df, total_assets_row])
    liabilities_df = pd.concat([liabilities_df, total_liabilities_row])

    # Display tables with formatted values
    st.subheader("Assets")
    st.table(assets_df.style.format({'Amount ($)': '{:,.2f}'}))

    st.subheader("Liabilities")
    st.table(liabilities_df.style.format({'Amount ($)': '{:,.2f}'}))


def plot_bar_graphs(assets, liabilities):
    # Filter out items with zero values
    filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
    filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

    # Calculate compounded liabilities
    # compounded_liabilities = {} 

    # for k, v in filtered_liabilities.items():
        # if 'Interest Rate' in k or 'Time Period' in k:
        #     continue  # Skip non-monetary entries

        # if k == 'Credit Card Payment' and liabilities['Credit Card Debt Interest Rate'] == 0.0:
        #     continue  # Skip if credit card interest rate is zero

        # if k == 'Mortgage':
        #     interest_rate = liabilities['Annual Mortgage Interest Rate']
        #     time_period = liabilities['Mortagage Time Period']

        # elif k == 'Home Loans':
        #     interest_rate = liabilities['Home Loans Interest Rate']
        #     time_period = liabilities['Home Loans Time Period']

        # elif k == 'Car/Vehicle Loans':
        #     interest_rate = liabilities['Car/Vehicle Loans Interest Rate']
        #     time_period = liabilities['Car/Vehicle Loans Time Period']

        # elif k == 'Education Loans':
        #     interest_rate = liabilities['Education Loans Interest Rate']
        #     time_period = liabilities['Education Loans Time Period']

        # elif k == 'Credit Card Payment':
        #     interest_rate = liabilities['Credit Card Debt Interest Rate']
        #     time_period = 1  # Assuming interest is calculated yearly

        # if interest_rate > 0:
        #     compounded_amount = float(v) * (1 + float(interest_rate) / 100) ** float(time_period)
        #     compounded_liabilities[k] = compounded_amount
        # else:
        #     compounded_liabilities[k] = float(v)

    # Plot bar graph for assets
    st.write("### All Assets ")
    fig1, ax1 = plt.subplots()
    ax1.bar(filtered_assets.keys(), filtered_assets.values(), color='green')
    ax1.set_ylabel('Amount ($)')
    ax1.set_xlabel('Asset Type')
    ax1.set_title(' All Assets ')
    plt.xticks(rotation=45)
    st.pyplot(fig1)

    # Plot bar graph for liabilities
    st.write("### All Liabilities ")
    # st.write("### All Liabilities with Compounded Interest")
    fig2, ax2 = plt.subplots()
    # ax2.bar(compounded_liabilities.keys(), compounded_liabilities.values(), color='red')
    ax2.bar(filtered_liabilities.keys(), filtered_liabilities.values(), color='red')    
    ax2.set_ylabel('Amount ($)')
    ax2.set_xlabel('Liability Type')
    ax2.set_title(' All Liabilities ')

    # ax2.set_title(' All Liabilities with Compounded Interest')
    plt.xticks(rotation=45)
    st.pyplot(fig2)


from docx import Document
# Define a helper function to read and extract text from a DOCX file
def read_docx(file_path):
    document = Document(file_path)
    extracted_text = "\n".join([para.text for para in document.paragraphs])
    return extracted_text



# # Ask for investment personality
# investment_personality = st.selectbox(
#     "Select the investment personality of the client:",
#     ("Conservative Investor", "Moderate Investor", "Aggressive Investor")
# )

# # Step 4: Generate investment suggestions
# if st.button("Generate Investment Suggestions"):
#     st.write("Generating investment suggestions...")

#     # Replace this with the path to your DOCX file
#     file_path = "data/Financial_Investment_1.docx"

#     # New logic to read the file directly
#     try:
#         st.write("Loading client document from predefined path...")
#         extracted_text = read_docx(file_path)
#         st.write("Client document loaded successfully.")
#     except FileNotFoundError:
#         st.write("Client document file not found.")
#         st.stop()  # Stop the Streamlit app if the file is not found


#     # Assuming generate_investment_suggestions is an async function
    
#     suggestions = asyncio.run(generate_investment_suggestions(investment_personality, extracted_text))
#     st.write(suggestions)

#     data_extracted = extract_numerical_data(suggestions)

#     # Streamlit app for visualizing investment allocation
#     st.title('Investment Allocation Infographics')

#     st.write('## Investment Allocation Charts')
#     fig = plot_investment_allocations(data_extracted)
#     st.pyplot(fig)

#     st.write('## Pie Chart of Investment Allocation')
#     fig = plot_pie_chart(data_extracted)
#     st.pyplot(fig)

#     st.write('## Bar Chart of Compounded Returns')
#     fig = bar_chart(data_extracted['Growth-Oriented Investments'])
#     st.pyplot(fig)

#     plot_3d_bar_graph(data_extracted)


# class TrieNode:
#     def __init__(self):
#         self.children = {}
#         self.client_ids = []  # List to store client IDs for this prefix



class TrieNode:
    def __init__(self):
        self.children = {}
        self.client_ids = []
        self.end_of_name = False  # Marks the end of a client's name

class Trie:
    def __init__(self):
        self.root = TrieNode()

    def insert(self, name, client_id):
        node = self.root
        for char in name:
            if char not in node.children:
                node.children[char] = TrieNode()
            node = node.children[char]
        node.client_ids.append(client_id)
        node.end_of_name = True

    def search(self, prefix):
        node = self.root
        for char in prefix:
            if char in node.children:
                node = node.children[char]
            else:
                return []  # Prefix not found
        return self._get_all_names_from_node(prefix, node)

    def _get_all_names_from_node(self, prefix, node):
        suggestions = []
        if node.end_of_name:
            suggestions.append((prefix, node.client_ids))
        for char, child_node in node.children.items():
            suggestions.extend(self._get_all_names_from_node(prefix + char, child_node))
        return suggestions

# pip install streamlit-autocomplete
# from streamlit_autocomplete import st_autocomplete

def preload_trie():
    trie = Trie()
    clients = {
        "John Doe": "C001",
        "Jane Smith": "C002",
        "James Brown": "C003",
        "Jill Johnson": "C004",
        "Jake White": "C005"
    }
    for name, client_id in clients.items():
        trie.insert(name.lower(), client_id)  # Insert in lowercase for case-insensitive search
    return trie



# class TrieNode:
#     def __init__(self):
#         self.children = {}
#         self.client_ids = []
#         self.end_of_name = False  # Marks the end of a client's name


# class Trie:
#     def __init__(self):
#         self.root = TrieNode()

#     def insert(self, name, client_id):
#         node = self.root
#         for char in name:
#             if char not in node.children:
#                 node.children[char] = TrieNode()
#             node = node.children[char]
#         node.client_ids.append(client_id)
#         node.end_of_name = True

#     def search(self, prefix):
#         node = self.root
#         for char in prefix:
#             if char in node.children:
#                 node = node.children[char]
#             else:
#                 return []  # Prefix not found
#         return self._get_all_names_from_node(prefix, node)

#     def _get_all_names_from_node(self, prefix, node):
#         suggestions = []

#         if node.end_of_name:
#             suggestions.append((prefix, node.client_ids))

#         for char, child_node in node.children.items():
#             suggestions.extend(self._get_all_names_from_node(prefix + char, child_node))

#         return suggestions


# def preload_trie():
#     trie = Trie()
#     clients = {
#         "John Doe": "C001",
#         "Jane Smith": "C002",
#         "James Brown": "C003",
#         "Jill Johnson": "C004",
#         "Jake White": "C005"
#     }
#     for name, client_id in clients.items():
#         trie.insert(name.lower(), client_id)  # Insert in lowercase for case-insensitive search
#     return trie





# def main(): # main file that uses Client.txt directly 
#     if st.session_state.page == "main":
#         st.title("Wealth Advisor Chatbot")

#         # Step 1: Choose between Investment Suggestions or Stock Analysis
#         task_choice = st.radio("Select the task you want to perform:", ["Investment Suggestions", "Stock Analysis"])

#         # Step 2: Choose between New Client or Existing Client (for Investment Suggestions)
#         if task_choice == "Investment Suggestions":
#             client_type = st.radio("Is this for a new client or an existing client?", ["New Client", "Existing Client"])

#             if client_type == "New Client":
#                 # Button to redirect to the form page
#                 if st.button("Fill in the Client Details"):
#                     st.session_state.page = "form"
                
#                 # Display the pie chart if assets and liabilities data is available
#                 if 'assets' in st.session_state and 'liabilities' in st.session_state:
#                     st.subheader("Assets and Liabilities Breakdown")

#                     # Create Summary table for all Assets and Liabilities:
#                     create_financial_summary_table(st.session_state.assets, st.session_state.liabilities)

#                     # Plot the bar graphs for assets and liabilities
#                     plot_bar_graphs(st.session_state.assets, st.session_state.liabilities)

#                     # Plot the pie chart
#                     plot_assets_liabilities_pie_chart(st.session_state.assets, st.session_state.liabilities)
                    
#                 else:
#                     st.info("Please fill in the client details to view the assets and liabilities breakdown.")

#                 # Load data from client_data.txt instead of file upload
#                 # uploaded_file = 
#                 # try:
#                 #     with open("client_data.txt", "r") as file:
#                 #         extracted_text = file.read()
#                 #     st.write("Client data loaded successfully.")
#                 # except FileNotFoundError:
#                 #     # st.write("Client data file not found.")
#                 #     print("Client data file not found.")

#                 #     return

#                 # Ask for investment personality
#                 investment_personality = st.selectbox(
#                     "Select the investment personality of the client:",
#                     ("Conservative Investor", "Moderate Investor", "Aggressive Investor")
#                 )

#                 # Step 4: Generate investment suggestions
#                 if st.button("Generate Investment Suggestions"):
#                     st.write("Generating investment suggestions...")

#                     # Replace this with the path to your DOCX file
#                     # file_path =  "C:\Users\Harshal\OneDrive\Desktop\Wealth_Management_ChatBot\Telegram-ChatBot-using-Gemini\data\Financial_Investment_1.docx"    #"data/Financial_Investment_1.docx"

#                     # # New logic to read the file directly
#                     # try:
#                     #     st.write("Loading client document from predefined path...")
#                     #     extracted_text = read_docx(file_path)
#                     #     st.write("Client document loaded successfully.")
#                     # except FileNotFoundError:
#                     #     st.write("Client document file not found.")
#                     #     st.stop()  # Stop the Streamlit app if the file is not found



#                     # Assuming generate_investment_suggestions is an async function
                    
#                     suggestions = asyncio.run(generate_investment_suggestions(investment_personality, extracted_text))
#                     st.write(suggestions)

#                     data_extracted = extract_numerical_data(suggestions)

#                     # Streamlit app for visualizing investment allocation
#                     st.title('Investment Allocation Infographics')

#                     st.write('## Investment Allocation Charts')
#                     fig = plot_investment_allocations(data_extracted)
#                     st.pyplot(fig)

#                     st.write('## Pie Chart of Investment Allocation')
#                     fig = plot_pie_chart(data_extracted)
#                     st.pyplot(fig)

#                     st.write('## Bar Chart of Compounded Returns')
#                     fig = bar_chart(data_extracted['Growth-Oriented Investments'])
#                     st.pyplot(fig)

#                     plot_3d_bar_graph(data_extracted)

#             elif client_type == "Existing Client":
#                 st.write("Fetching existing client data...")

#         elif task_choice == "Stock Analysis":
#             st.write("Performing stock analysis...")







def main():

    # Initialize the Trie with preloaded clients
    trie = preload_trie()

    if st.session_state.page == "main":
        st.title("Wealth Advisor Chatbot")

        # Step 1: Choose between Investment Suggestions or Stock Analysis
        task_choice = st.radio("Select the task you want to perform:", ["Investment Suggestions", "Stock Analysis"])

        # Step 2: Choose between New Client or Existing Client (for Investment Suggestions)
        if task_choice == "Investment Suggestions":
            client_type = st.radio("Is this for a new client or an existing client?", ["New Client", "Existing Client"])
            

            if client_type == "New Client":
                # Button to redirect to the form page
                if st.button("Fill in the Client Details"):
                    st.session_state.page = "form"
                
                # Display the pie chart if assets and liabilities data is available
                if 'assets' in st.session_state and 'liabilities' in st.session_state:
                    st.subheader("Assets and Liabilities Breakdown")

                    # Create Summary table for all Assets and Liabilities:
                    create_financial_summary_table(st.session_state.assets, st.session_state.liabilities)

                    # Plot the bar graphs for assets and liabilities
                    plot_bar_graphs(st.session_state.assets, st.session_state.liabilities)

                    # Plot the pie chart
                    plot_assets_liabilities_pie_chart(st.session_state.assets, st.session_state.liabilities)
                    
                else:
                    st.info("Please fill in the client details to view the assets and liabilities breakdown.")


                uploaded_file = st.file_uploader("Upload the client's document", type=["docx", "pdf"])

                if uploaded_file is not None:
                    # st.write("Extracting text from the document...")
                    st.write("Received Personal Details File ")
                    document_data = asyncio.run(process_document(uploaded_file))

                    if document_data:
                        if isinstance(document_data, tuple) and len(document_data) == 2:
                            extracted_text, tables_content = document_data
                            # st.write("Text extracted from the document")

                            # Print extracted text for debugging
                            print("Extracted Text Content:", extracted_text)

                        
                        else:
                            st.write("Unexpected data format returned from document processing.")

                    # Ask for investment personality
                    st.title("Investment Personality Assessment")
    
                    # Use file_uploader to allow users to upload the document
                    personality_file = st.file_uploader("Upload the client's Investment Assessment File", type=["docx", "pdf"])
                    
                    if personality_file is not None:
                        st.write("Extracting responses from the assessment document...")
                        responses = extract_responses_from_docx(personality_file)
                        # responses = extract_responses_from_assessment(personality_file)
                        if responses:
                            st.write("Responses extracted successfully.")
                            
                            # Determine investment personality based on extracted responses
                            # personality = determine_investment_personality(responses)
                            personality = asyncio.run(determine_investment_personality(responses)
)
                            st.write(f"The Client's Investment Personality is: {personality}")

                    # if personality_file is not None:
                    #     st.write("Extracting responses from the assessment document...")
                    #     responses = extract_responses_from_assessment(personality_file)
                    #     if responses:
                    #         st.write("Responses extracted successfully.")
                            
                    #         # Convert responses to a dictionary format expected by the AI model
                    #         assessment_data = {f"Question {i+1}": response for i, response in enumerate(responses)}
                            
                    #         # Determine investment personality based on extracted responses
                    #         personality = asyncio.run(determine_investment_personality(assessment_data))
                    #         st.write(f"Determined Investment Personality: {personality}")


                            if personality == "Conservative Investor":
                                investment_personality = st.selectbox(
                                    "To toggle with Personality, Select the investment personality of the client:",
                                    ("Conservative Investor", "Moderate Investor", "Aggressive Investor")
                                )
                            elif personality == "Moderate Investor":
                                investment_personality = st.selectbox(
                                    "To toggle with Personality, Select the investment personality of the client:",
                                    ( "Moderate Investor","Conservative Investor", "Aggressive Investor")
                                )
                            else:
                                investment_personality = st.selectbox(
                                    "To toggle with Personality, Select the investment personality of the client:",
                                    ( "Moderate Investor","Conservative Investor", "Aggressive Investor")
                                )


                    #     else:
                    #         st.write("No responses found in the document or an error occurred.")

                    # if personality == "Conservative Investor":
                    #     investment_personality = st.selectbox(
                    #         "To toggle with Personality, Select the investment personality of the client:",
                    #         ("Conservative Investor", "Moderate Investor", "Aggressive Investor")
                    #     )
                    # elif personality == "Moderate Investor":
                    #     investment_personality = st.selectbox(
                    #         "To toggle with Personality, Select the investment personality of the client:",
                    #         ( "Moderate Investor","Conservative Investor", "Aggressive Investor")
                    #     )
                    # else:
                    #     investment_personality = st.selectbox(
                    #         "To toggle with Personality, Select the investment personality of the client:",
                    #         ( "Moderate Investor","Conservative Investor", "Aggressive Investor")
                    #     )

                    # Step 4: Generate investment suggestions
                    if st.button("Generate Investment Suggestions"):
                        st.write("Generating investment suggestions...")
                        # Assuming generate_investment_suggestions is an async function
                        
                        suggestions = asyncio.run(generate_investment_suggestions(investment_personality, extracted_text))
                        st.write(suggestions)

                        # # Generate infographics
                        # generate_infographics(suggestions, investment_personality)

                        data_extracted = extract_numerical_data(suggestions)

                        # Streamlit app
                        st.title('Investment Allocation Infographics')

                        st.write('## Investment Allocation Charts')
                        fig = plot_investment_allocations(data_extracted)
                        st.pyplot(fig)

                        st.write('## Pie Chart of Investment Allocation')

                        # plot_pie_chart(data_extracted)
                        fig = plot_pie_chart(data_extracted)
                        st.pyplot(fig)

                        

            elif client_type == "Existing Client":
                st.subheader("Search for Existing Client")

                # def get_suggestions(prefix):
                #     return [name for name, _ in trie.search(prefix.lower())]

                # search_term = st_autocomplete("Search for a client:", get_suggestions)

                # if search_term:
                #     suggestions = trie.search(search_term.lower())
                #     if suggestions:
                #         st.write("Suggestions:")
                #         for suggestion, client_ids in suggestions:
                #             st.write(f"{suggestion} (Client IDs: {', '.join(client_ids)})")
                #     else:
                #         st.write("No suggestions found.")


                search_query = st.text_input("Search by Client Name")
                if search_query:
                    search_query = search_query.lower()
                    matching_names = trie.search(search_query)
                    if matching_names:
                        st.write("Suggestions:")
                        for name, client_ids in matching_names:
                            for client_id in client_ids:
                                st.write(f"{name.title()} - Client ID: {client_id}")
                    else:
                        st.write("No matching clients found.")

                

        elif task_choice == "Stock Analysis":
            st.write("Performing stock analysis...")



# def main():
    
#     if st.session_state.page == "main":
#         st.title("Wealth Advisor Chatbot")

#         # Step 1: Choose between Investment Suggestions or Stock Analysis
#         task_choice = st.radio("Select the task you want to perform:", ["Investment Suggestions", "Stock Analysis"])

#         # Step 2: Choose between New Client or Existing Client (for Investment Suggestions)
#         if task_choice == "Investment Suggestions":
#             client_type = st.radio("Is this for a new client or an existing client?", ["New Client", "Existing Client"])
            

#             if client_type == "New Client":
#                 # Button to redirect to the form page
#                 if st.button("Fill in the Client Details"):
#                     st.session_state.page = "form"
                
#                 # Display the pie chart if assets and liabilities data is available
#                 if 'assets' in st.session_state and 'liabilities' in st.session_state:
#                     st.subheader("Assets and Liabilities Breakdown")

#                     # Create Summary table for all Assets and Liabilities:
#                     create_financial_summary_table(st.session_state.assets, st.session_state.liabilities)

#                     # Plot the bar graphs for assets and liabilities
#                     plot_bar_graphs(st.session_state.assets, st.session_state.liabilities)

#                     # Plot the pie chart
#                     plot_assets_liabilities_pie_chart(st.session_state.assets, st.session_state.liabilities)
                    
#                 else:
#                     st.info("Please fill in the client details to view the assets and liabilities breakdown.")

#             # if client_type == "New Client":
#             #     # Button to redirect to the form page
#             #     if st.button("Fill in the Client Details"):
#             #         st.session_state.page = "form"
                
#             #         # Display the pie chart if assets and liabilities data is available

#             #     if 'assets' in st.session_state and 'liabilities' in st.session_state:
#             #         st.subheader("Assets and Liabilities Breakdown")

#             #         # Create Summary table for all Assets and Liabilities:
#             #         create_financial_summary_table(st.session_state.assets, st.session_state.liabilities)

#             #         # Plot the bar graphs for assets and liabilities
#             #         plot_bar_graphs(st.session_state.assets, st.session_state.liabilities)

#             #         # Plot the pie chart
#             #         # st.subheader("Assets and Liabilities Breakdown")
#             #         # plot_assets_liabilities_pie_chart(assets, liabilities)
#             #         plot_assets_liabilities_pie_chart(st.session_state.assets, st.session_state.liabilities)
                    
#             #         # Display total Calculated assets and liabilities : No need to display them 
#             #         # st.write(f"Total Assets: {st.session_state.total_assets}")
#             #         # st.write(f"Total Liabilities: {st.session_state.total_liabilities}")

#             #     else:
#             #         st.info("Please fill in the client details to view the assets and liabilities breakdown.")

#                 uploaded_file = st.file_uploader("Upload the client's document", type=["docx", "pdf"])

#                 if uploaded_file is not None:
#                     # st.write("Extracting text from the document...")
#                     st.write("Received Personal Details File ")
#                     document_data = asyncio.run(process_document(uploaded_file))

#                     if document_data:
#                         if isinstance(document_data, tuple) and len(document_data) == 2:
#                             extracted_text, tables_content = document_data
#                             # st.write("Text extracted from the document")

#                             # Print extracted text for debugging
#                             print("Extracted Text Content:", extracted_text)

                        
#                         else:
#                             st.write("Unexpected data format returned from document processing.")

#                     # Ask for investment personality
#                     st.title("Investment Personality Assessment")
    
#                     # Use file_uploader to allow users to upload the document
#                     personality_file = st.file_uploader("Upload the client's Investment Assessment File", type=["docx", "pdf"])
                    
#                     if personality_file is not None:
#                         st.write("Extracting responses from the assessment document...")
#                         responses = extract_responses_from_docx(personality_file)
                        
#                         if responses:
#                             st.write("Responses extracted successfully.")
                            
#                             # Determine investment personality based on extracted responses
#                             personality = determine_investment_personality(responses)
#                             st.write(f"The Client's Investment Personality is: {personality}")
#                         else:
#                             st.write("No responses found in the document or an error occurred.")


#                     investment_personality = st.selectbox(
#                         "To toggle with Personality, Select the investment personality of the client:",
#                         ("Conservative Investor", "Moderate Investor", "Aggressive Investor")
#                     )

#                     # Step 4: Generate investment suggestions
#                     if st.button("Generate Investment Suggestions"):
#                         st.write("Generating investment suggestions...")
#                         # Assuming generate_investment_suggestions is an async function
                        
#                         suggestions = asyncio.run(generate_investment_suggestions(investment_personality, extracted_text))
#                         st.write(suggestions)

#                         # # Generate infographics
#                         # generate_infographics(suggestions, investment_personality)

#                         data_extracted = extract_numerical_data(suggestions)

#                         # Streamlit app
#                         st.title('Investment Allocation Infographics')

#                         st.write('## Investment Allocation Charts')
#                         fig = plot_investment_allocations(data_extracted)
#                         st.pyplot(fig)

#                         st.write('## Pie Chart of Investment Allocation')

#                         # plot_pie_chart(data_extracted)
#                         fig = plot_pie_chart(data_extracted)
#                         st.pyplot(fig)

#                         # st.write('## Bar Chart of Compounded Returns')
#                         # fig = bar_chart(data_extracted['Growth-Oriented Investments'])
#                         # st.pyplot(fig)

#                         # plot_3d_bar_graph(data_extracted)
#                         # fig = plot_3d_bar_graph(data_extracted)
#                         # st.pyplot(fig)

#             elif client_type == "Existing Client":
#                 st.write("Fetching existing client data...")

#         elif task_choice == "Stock Analysis":
#             st.write("Performing stock analysis...")

    

if __name__ == "__main__":
    if 'page' not in st.session_state:
        st.session_state.page = "main"
    
    if st.session_state.page == "form":
        client_form()
    else:
        main()

# if __name__ == "__main__":
#     main()

