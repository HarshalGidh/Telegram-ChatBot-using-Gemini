#best code so far now it can upload files and receive various forms of messages as well and provide us graphs that we want 
# and also reply to images, give stock informations ,give stock analysis information and a prediction of the price

import os
import filetype
import docx
import PyPDF2
import re
from aiogram import Bot, Dispatcher, types
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.memory import ConversationSummaryMemory
import asyncio
import numpy as np
import json
import re
import google.generativeai as genai
import pathlib
# Import things that are needed generically
from langchain.pydantic_v1 import BaseModel, Field
from langchain.tools import BaseTool, StructuredTool, tool

from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import CommandStart
from aiogram.types import Message
from aiogram import F
from aiogram import Router
import logging
import sys
from aiogram.filters import Command
from aiogram.types import FSInputFile
# from aiogram.utils import executor
import io
import matplotlib.pyplot as plt
import seaborn as sns
import aiohttp
from aiogram.types import InputFile , BufferedInputFile
import PIL.Image

router = Router(name=__name__)

load_dotenv()

TOKEN = os.getenv("TOKEN")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

os.environ["LANGCHAIN_TRACING_V2"]="true"
os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# Configure generativeai with your API key
genai.configure(api_key=GOOGLE_API_KEY)

# Initialize bot
bot = Bot(token=TOKEN)
dp = Dispatcher()

# Glbal variables
rag_on = False
retriever = None  # Store retriever globally
summary = ""
investment_personality = ""
chat_history = ""
previous_suggestions = ""

class Reference:
    def __init__(self):
        self.response = ""


reference = Reference()


def clear_past():
    reference.response = ""


@router.message(F.text == "clear")
async def clear(message: types.Message):
    """
    A handler to clear the previous conversation and context.
    """
    clear_past()
    await message.reply("I've cleared the past conversation and context.")

#Global Variables :

# Store user states
states = {}

# Dictionary to hold question-answer pairs
user_responses = {}
#
user_images = {}
# Define Questions for assessment
questions = [
    """ 
    1. You and your friend are betting on a series of coin tosses.

    He always bets ₹2,000 on Heads

    You always bet ₹2,000 on Tails

    Winner of last 8 turns

    You lost ₹8,000 in the last 4 turns!

    If you were to bet one last time, what would you bet on:
    a) heads or b) tails ?
    """ ,
    """
    2. Imagine you are a contestant in a game show, and you are presented the following choices.

    What would you prefer?
    a) 50 percent chance of winning 15 gold coins 
    b) 100 percent chance of winning 8 gold coins
    """,
    """
    3. Ok, one more choice...

    What would you prefer?
    a) 50 percent chance of winning 15 gold coins 
    b) 100 percent chance of winning 2 gold coins
    """,
    """
    4. In general, how would your best friend describe your risk-taking tendencies?
    a) A real gambler
    b) Willing to take risks after completing adequate research
    c) Cautious
    d) Avoids risk as much as possible
    """,
    """
    5. Suppose you could replace your current investment portfolio with this new one:
    50 percent chance of Gaining 35 percent or 50 percent chance of Loss
    In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
    a)-5 to -10
    b)-10 to -15
    c)-15 to -20
    d)-20 to -25
    e)-25 to -30
    f)-30 to -35
    """,
    """
    6. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
    a) Option A : Min Value is 9.2L and Max Value is 11.8L
    b) Option B : Min Value is 8.8L and Max Value is 12.3L
    c) Option C : Min Value is 8.5L and Max Value is 12.8L
    d) Option D : Min Value is 8.1L and Max Value is 13.3L
    e) Option E : Min Value is 7.8L and Max Value is 13.8L
    """,
    """
    7. From Sept 2008 to Nov 2008, Stock market went down by 31%.

    If you owned a stock investment that lost about 31 percent in 3 months, you would:
    a) Sell all of the remaining investment
    b) Sell a portion of the remaining investment
    c) Hold on to the investment and sell nothing
    d) Buy little
    e) Buy more of the investment
    """,
    """
    8. Over any 1-year period, what would be the maximum drop in the value of your investment 
    portfolio that you would be comfortable with?
    a) <5%
    b) 5 - 10%
    c) 10 - 15%
    d) 15 - 20%
    e) >20%
    """,
    """
    9. When investing, what do you consider the most?

    a) Risk 
    b) Return
    """,
    """
    10. What best describes your attitude?

    a) Prefer reasonable returns, can take reasonable risk
    b) Like higher returns, can take slightly higher risk
    c) Want to maximize returns, can take significant high risk
    """,
    """
    11. How much monthly investment you want to do?
    """,
    """
    12. What is the time horizon for your investment?
    You can answer in any range, example 1-5 years."""  
]


# Handler for /start command
@dp.message(CommandStart())
async def handle_start(message: types.Message):
    """
    This handler receives messages with /start command
    """
    chat_id = message.chat.id
    # Start asking questions
    await start_assessment(chat_id)


# Function to start the assessment
async def start_assessment(chat_id):
    await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
    await ask_next_question(chat_id, 0)

# Function to ask the next question
async def ask_next_question(chat_id, question_index):
    if question_index < len(questions):
        # Ask the next question
        await bot.send_message(chat_id, questions[question_index])
        # Update state to indicate the next expected answer
        states[chat_id] = question_index
    else:
        # No more questions, finish assessment
        await finish_assessment(chat_id)

# Handler for receiving assessment answers
assessment_in_progress = True

from aiogram.types import FSInputFile
async def finish_assessment(chat_id):
    if chat_id in states and states[chat_id] == len(questions):
        # All questions have been answered, now process the assessment
        await bot.send_message(chat_id, "Assessment completed. Thank you!")

        # Determine investment personality based on collected responses
        global investment_personality
        investment_personality = await determine_investment_personality(user_responses)

        # Inform the user about their investment personality
        await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

        # Summarize collected information
        global summary
        summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
        summary = summary + "\n" + "Your investment personality:" + investment_personality
        # Ensure to await the determination of investment personality
        await send_summary_chunks(chat_id, summary)
        global assessment_in_progress 
        assessment_in_progress = False
       
        await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
        file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

        await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
        # await bot.send_message(chat_id,file)

async def send_summary_chunks(chat_id, summary):
    # Split the summary into chunks that fit within Telegram's message limits
    chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

    # Send each chunk as a separate message
    for chunk in chunks:
        await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


async def determine_investment_personality(assessment_data):
    try:
        # Prepare input text for the chatbot based on assessment data
        input_text = "User Profile:\n"
        for question, answer in assessment_data.items():
            input_text += f"{question}: {answer}\n"

        # Introduce the chatbot's task and prompt for classification
        input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
                      "- Conservative Investor\n" \
                      "- Moderate Investor\n" \
                      "- Aggressive Investor"

        # Use your generative AI model to generate a response
        # print(input_text)
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(input_text)

        # Determine the investment personality from the chatbot's response
        response_text = response.text.lower()
        if "conservative" in response_text:
            personality = "Conservative Investor"
        elif "moderate" in response_text:
            personality = "Moderate Investor"
        elif "aggressive" in response_text:
            personality = "Aggressive Investor"
        else:
            personality = "Unknown"

        return personality
        # Send the determined investment personality back to the user
        #await bot.send_message(chat_id, f"Investment Personality: {personality}")

    except Exception as e:
        print(f"Error generating response: {e}")
        #await bot.send_message(chat_id, "Error processing investment personality classification.")


# Handler for document upload
async def load_vector_db(file_path):
    try:
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        return vector_store.as_retriever(search_kwargs={"k": 1})
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None



async def make_retrieval_chain(retriever):
    """
    Create a retrieval chain using the provided retriever.

    Args:
        retriever (RetrievalQA): A retriever object.

    Returns:
        RetrievalQA: A retrieval chain object.
    """
    try:
        global investment_personality,summary
        llm = ChatGoogleGenerativeAI(
            #model="gemini-pro",
            model = "gemini-1.5-flash",
            temperature=0.7,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )

        prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
                Respond to the client by the client name.
                Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
                Also give the user detailed information about the investment how to invest,where to invest and how much they
                should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
                Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
                investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
                Also give the user minimum and maximum expected growth in dollars for the time horizon .
                Also explain the user why you are giving them that particular investment suggestion.
                Give the client suggestions of Investment based on their investment personality that can also help them manage their mortages and other liablilities if they have any,ignore if they dont have any liabilities.
                Answer in 3-4 lines.\n
                <context>
                {context}
                </context>
                Question: {input}"""

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in creating chain: {e}")
        return None

from aiogram.filters import Filter

# @router.message(F.document)
@dp.message(F.document)
async def handle_document(message: types.Message):
    global summary,investment_personality  

    chat_id = message.chat.id

    # Obtain file information
    file_id = message.document.file_id
    file = await bot.get_file(file_id)
    file_path = file.file_path
    
    # Download the file
    await bot.download_file(file_path, "data/uploaded_file")
    
    # Process the uploaded document
    extracted_text = await process_document("data/uploaded_file")
    
    if extracted_text:
        # Load vector database (assuming this is part of setting up the retriever)
        retriever = await load_vector_db("data/uploaded_file")
        file_path = 'data/uploaded_file'
        client_name, validation_errors = await process_document(file_path)

        # Print results
        print(f"Client Name: {client_name}")
        if validation_errors:
            print("**Validation Errors:**")
            for error in validation_errors:
                print(error)
        else:
            print("All fields are filled correctly.")
        if client_name == None:
            try:
                await message.reply("Processing the uploaded image")
                await handle_image(message) 
                return 
            except Exception as e:
                await message.reply("error processing uploaded image")
                print(e)
        await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")

        if retriever is None:
            await message.reply("The retrieval chain is not set up. Please upload a document first.")
            return

        # Check if a valid chain can be created
        chain = await make_retrieval_chain(retriever)
        if chain is None:
            await message.reply("Failed to create the retrieval chain.")
            return
        
        try:     
            query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
        
            response = chain.invoke({"input": query})
            print(response['answer'])
            global chat_history
            chat_history = response['answer'] 
            print(f"\n Chat History : {chat_history}")
            format_response = markdown_to_text(response['answer'])
            await message.reply(format_response)
            # await message.reply(response['answer'])

            try:
                graph_query = (
                    # extracted_text + "\n" + 
                    summary + "\n" + "investment_personality" + "\n" + response['answer'] + "\n" +
                    "Please provide the following information in JSON format:\n" +
                    "{\n"
                    "  'growth_investment_min': <minimum percentage of growth-oriented investments>,\n"
                    "  'growth_investment_max': <maximum percentage of growth-oriented investments>,\n"
                    "  'fixed_income_min': <minimum percentage of fixed-income investments>,\n"
                    "  'fixed_income_max': <maximum percentage of fixed-income investments>,\n"
                    "  'cash_min': <minimum percentage of cash investments>,\n"
                    "  'cash_max': <maximum percentage of cash investments>,\n"
                    "  'return_growth_investment': <Maximum percentage returns of growth-oriented investments>,\n"
                    "  'return_fixed_income': <Maximum percentage of returns of fixed-income investments>,\n"
                    "  'return_cash_max': <maximum percentage of retruns of cash investments>,\n"
                    "  'return_min': <minimum expected annual return percentage>,\n"
                    "  'return_max': <maximum expected annual return percentage>,\n"
                    "  'growth_min': <minimum expected growth in dollars>,\n"
                    "  'growth_max': <maximum expected growth in dollars>,\n"
                    "  'initial_investment': <initial monthly investment>,\n"
                    "  'time_horizon': <time horizon in years>\n"
                    "}"
                )
                graph_response = chain.invoke({"input": graph_query})
                print(graph_response['answer'])
                await handle_graph(graph_response['answer'],chat_id)
                try:
                    await refine_investment_suggestions(chat_id, response['answer'])  # Generate refined suggestions
                except Exception as e:
                    print(f"Error refining investment suggestions: {e}")
                    await bot.send_message(chat_id, "Error refining investment suggestions. Please try again later.")

            except Exception as e:
                print(f"Error plotting graph : {e}")
        except Exception as e:
            print(f"Error invoking retrieval chain on attempt : {e}")
            await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

    else:
        await message.reply("Failed to process the uploaded file.")
    

async def refine_investment_suggestions(chat_id, suggestions):
    try:
        task = """Based on the following investment suggestions, provide more detailed investment advice specifically 
                  on Investment Allocation of stocks:
                  How to invest in these stocks, what percentage of the investment should go into each stock, 
                  the expected returns, and the reasoning behind each suggestion. 
                  Also, include any potential risks and recommendations for monitoring the investments.
                  Give the response in brief in around 15-20 lines."""

        query = task + "\n" + suggestions
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(query)
        await bot.send_message(chat_id, "Here are more detailed stock investment suggestions:")
        format_response = markdown_to_text(response.text)
        await bot.send_message(chat_id, format_response) #response.text
    
    except Exception as e:
        logging.error(f"Error generating refined investment suggestions: {e}")
        await bot.send_message(chat_id, "There was an error generating more detailed investment suggestions. Please try again later.")

# Function to extract data from LLM response
def extract_data_from_response(response):
    try:
        # Locate the JSON-like data in the response
        json_start = response.find("{")
        json_end = response.rfind("}") + 1
        
        if json_start == -1 or json_end == -1:
            raise ValueError("No JSON data found in the response.")
        
        json_data = response[json_start:json_end]
        
        # Parse the JSON data
        data = json.loads(json_data.replace("'", "\""))
        print(data)
        return data
    except Exception as e:
        logging.error(f"Error extracting data: {e}")
        return None


async def handle_graph(response, chat_id):
    try:
        data = extract_data_from_response(response)
        if not data:
            await bot.send_message(chat_id, "Failed to extract data from the response.")
            return
        
        # Log extracted data for debugging
        logging.info(f"Extracted data: {data}")
        
        # Create a pie chart for investment allocation
        labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
        sizes = [
            (data['growth_investment_min'] + data['growth_investment_max']) / 2,
            (data['fixed_income_min'] + data['fixed_income_max']) / 2,
            (data['cash_min'] + data['cash_max']) / 2
        ]
        fig1, ax1 = plt.subplots()
        ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
        pie_chart_buffer = io.BytesIO()
        plt.savefig(pie_chart_buffer, format='png')
        pie_chart_buffer.seek(0)
        pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

        # Create a bar graph for potential returns
        allocations = {
            'Growth-Oriented Investments': (data['growth_investment_min'] + data['growth_investment_max']) / 2,
            'Fixed-Income Investments': (data['fixed_income_min'] + data['fixed_income_max']) / 2,
            'Cash and Cash Equivalents': (data['cash_min'] + data['cash_max']) / 2
        }

        returns = {
            'Growth-Oriented Investments': (data['return_growth_investment'] ) / 100,  # Annual return
            'Fixed-Income Investments': (data['return_fixed_income'] ) / 100 ,
            'Cash and Cash Equivalents': (data['return_cash_max'] ) / 100
        }
        
        time_horizon_years = data['time_horizon']
        initial_monthly_investment = data['initial_investment']
        
        # Calculate total returns for each category
        total_investment = initial_monthly_investment * 12 * time_horizon_years
        total_returns = {category: total_investment * (1 + returns[category]) ** time_horizon_years for category in allocations.keys()}
        
        # Create the bar chart
        fig2, ax2 = plt.subplots(figsize=(10, 6))
        bars = ax2.bar(allocations.keys(), total_returns.values(), color=['#FF9999', '#66B2FF', '#99FF99'])

        # Add labels to the top of each bar
        for bar in bars:
            yval = bar.get_height()
            ax2.text(bar.get_x() + bar.get_width()/2, yval + 1000, f'₹{round(yval, 2)}', ha='center', va='bottom')

        # Add titles and labels
        plt.title('Investment Performance Over Time')
        plt.xlabel('Investment Categories')
        plt.ylabel('Total Returns (₹)')
        plt.grid(True, axis='y', linestyle='--', linewidth=0.7)
        
        # Save and display the plot
        bar_chart_buffer = io.BytesIO()
        plt.savefig(bar_chart_buffer, format='png')
        bar_chart_buffer.seek(0)
        bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='investment_performance.png')

        await bot.send_document(chat_id, pie_chart_file, caption="Investment Allocation Chart")
        await bot.send_document(chat_id, bar_chart_file, caption="Investment Performance Over Time")

    except Exception as e:
        logging.error(f"Error plotting graph: {e}")
        await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")


async def process_document(file_path):
    try:
        file_type = filetype.guess(file_path)
        if file_type is not None:
            if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # return extract_text_from_word(file_path)
                return extract_text_and_tables_from_word(file_path)
            elif file_type.mime == "application/pdf":
                return extract_text_from_pdf(file_path)
        return None
    except Exception as e:
        print(f"Error processing document: {e}")
        return None

def extract_text_from_pdf(pdf_file_path):
    try:
        with open(pdf_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            text_content = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text_content.append(page.extract_text())
            return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None


import re
import docx

def extract_text_and_tables_from_word(docx_file_path):
    """
    Extracts text and tables from a Word document (.docx).

    Args:
        docx_file_path (str): Path to the Word document file.

    Returns:
        tuple: Extracted text content and tables from the document.
    """
    try:
        doc = docx.Document(docx_file_path)
        text_content = []
        tables_content = []

        for para in doc.paragraphs:
            text_content.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)

        return "\n".join(text_content), tables_content
    except Exception as e:
        print(f"Error extracting text and tables from Word document: {e}")
        return None, None

def validate_document_content(text, tables):
    """
    Validates the content of the document.

    Args:
        text (str): Extracted text content from the document.
        tables (list): Extracted tables content from the document.

    Returns:
        tuple: Client name and validation errors.
    """
    errors = []
    
    # Extract client name
    client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
    client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

    # Define required sections
    required_sections = [
        "YOUR RETIREMENT GOAL",
        "YOUR OTHER MAJOR GOALS",
        "YOUR ASSETS AND LIABILITIES",
        "MY LIABILITIES",
        "YOUR CURRENT ANNUAL INCOME"
    ]

    # Check for the presence of required sections
    for section in required_sections:
        if section not in text:
            errors.append(f"* {section} section missing.")
    
    # Define table field checks
    table_checks = {
        "YOUR RETIREMENT GOAL": [
            r"When do you plan to retire\? \(age or date\)",
            r"Social Security Benefit \(include expected start date\)",
            r"Pension Benefit \(include expected start date\)",
            r"Other Expected Income \(rental, part-time work, etc.\)",
            r"Estimated Annual Retirement Expense"
        ],
        "YOUR OTHER MAJOR GOALS": [
            r"GOAL", r"COST", r"WHEN"
        ],
        "YOUR ASSETS AND LIABILITIES": [
            r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
            r"Current Value", r"Annual Contributions"
        ],
        "MY LIABILITIES": [
            r"Balance", r"Interest Rate", r"Monthly Payment"
        ]
    }

    # Validate table content
    for section, checks in table_checks.items():
        section_found = False
        for table in tables:
            table_text = "\n".join(["\t".join(row) for row in table])
            if section in table_text:
                section_found = True
                for check in checks:
                    if not re.search(check, table_text, re.IGNORECASE):
                        errors.append(f"* Missing or empty field in {section} section: {check}")
                break
        if not section_found:
            errors.append(f"* {section} section missing.")

    return client_name, errors

async def process_document(file_path):
    try:
        text, tables = extract_text_and_tables_from_word(file_path)
        if text is not None and tables is not None:
            client_name, errors = validate_document_content(text, tables)
            return client_name, errors
        return None, ["Error processing document."]
    except Exception as e:
        print(f"Error processing document: {e}")
        return None, [f"Error processing document: {e}"]

@dp.message()
async def main_bot(message: types.Message):
    global retriever, extracted_text, investment_personality, summary, chat_history

    chat_id = message.chat.id

    if chat_id in states and states[chat_id] < len(questions):
        question_index = states[chat_id]
        answer = message.text
        user_responses[questions[question_index]] = answer
        states[chat_id] += 1
        await ask_next_question(chat_id, question_index + 1)
    elif message.text:
        if message.text.upper().isalpha() and len(message.text) <= 5:
            ticker = message.text.upper()
            try:
                hist, data, excel_file = await fetch_stock_data(message, ticker)
                mean_price, analysis = await analyze_stock_data(hist, data, excel_file)
                await message.reply(f"{mean_price}")
                await message.reply(f"Stock analysis for {ticker}:\n{analysis}")
            except Exception as e:
                logging.error(f"Error fetching stock data for {ticker}: {e}")
                await message.reply("Failed to fetch stock data. Please try again.")
        else:
            try:
                task = """Provide financial advice based on the user's investment personality.
                Include detailed information about the investment, where to invest, how much to invest, 
                expected returns, and why you are giving this advice. 
                Include a disclaimer to monitor and rebalance investments regularly based on risk tolerance."""
                
                query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + message.text
                model = genai.GenerativeModel('gemini-1.5-flash')
                chat = model.start_chat(history=[])
                response = chat.send_message(query)
                format_response = markdown_to_text(response.result)
                await message.reply(format_response)
            
            except Exception as e:
                logging.error(f"Error processing general chat message: {e}")
                await message.reply("Failed to process your request.")

import yfinance as yf
import pandas as pd
import requests
from aiogram import types
from aiogram.types import InputFile
import os
import logging

NEWS_API_KEY = os.getenv('NEWS_API_KEY')

async def fetch_stock_data(message, ticker):
    stock = yf.Ticker(ticker)
    chat_id = message.chat.id
    data = {}

    company_details = stock.info.get('longBusinessSummary', 'No details available')
    data['Company Details'] = company_details
    sector = stock.info.get('sector', 'No sector information available')
    data['Sector'] = sector
    prev_close = stock.info.get('previousClose', 'No previous close price available')
    data['Previous Closing Price'] = prev_close
    open_price = stock.info.get('open', 'No opening price available')
    data['Today Opening Price'] = open_price

    hist = stock.history(period="5d")
    if not hist.empty and 'Close' in hist.columns:
        if hist.index[-1].date() == yf.download(ticker, period="1d").index[-1].date():
            close_price = hist['Close'].iloc[-1]
            data['Todays Closing Price'] = close_price
        else:
            data['Todays Closing Price'] = "Market is open, there is no closing price available yet."
    else:
        data['Todays Closing Price'] = "No historical data available for closing price."

    day_high = stock.info.get('dayHigh', 'No high price available')
    data['Today High Price'] = day_high
    day_low = stock.info.get('dayLow', 'No low price available')
    data['Today Low Price'] = day_low
    volume = stock.info.get('volume', 'No volume information available')
    data['Today Volume'] = volume
    dividends = stock.info.get('dividendRate', 'No dividend information available')
    data['Today Dividends'] = dividends
    splits = stock.info.get('lastSplitFactor', 'No stock split information available')
    data['Today Stock Splits'] = splits
    pe_ratio = stock.info.get('trailingPE', 'No P/E ratio available')
    data['P/E Ratio'] = pe_ratio
    market_cap = stock.info.get('marketCap', 'No market cap available')
    data['Market Cap'] = market_cap

    income_statement = stock.financials
    balance_sheet = stock.balance_sheet
    cashflow = stock.cashflow

    news_url = f'https://newsapi.org/v2/everything?q={ticker}&apiKey={NEWS_API_KEY}&pageSize=3'
    news_response = requests.get(news_url)
    if news_response.status_code == 200:
        news_data = news_response.json()
        articles = news_data.get('articles', [])
        if articles:
            top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
            data['Top News'] = top_news
        else:
            data['Top News'] = "No news articles found."
    else:
        data['Top News'] = "Failed to fetch news articles."

    for key, value in data.items():
        await message.reply(f"{key}: {value}")

    graph_url = f"https://finance.yahoo.com/chart/{ticker}"
    await message.reply(f"Stock Chart: \n{graph_url}")

    file_path = os.path.join('data', f'{ticker}_financial_data.xlsx')
    with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
        income_statement.to_excel(writer, sheet_name='Income Statement')
        balance_sheet.to_excel(writer, sheet_name='Balance Sheet')
        cashflow.to_excel(writer, sheet_name='Cashflow')

    excel_file = FSInputFile(file_path, filename=f'{ticker}_financial_data.xlsx')
    await bot.send_document(chat_id, excel_file, caption=f'{ticker} Financial Data')

    data_list = list(data.items())
    data_str = str(data_list)

    return hist, data_str, file_path


async def analyze_stock_data(hist, data, excel_file):
    avg_close = hist['Close'].mean()
    formatted_data = extract_excel_data(excel_file)

    task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.
                Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
                Predict the stock price range for the next month and provide reasons for your prediction.
                Advise whether to buy this stock now or not, with reasons for your advice."""

    query = task + "\nStock Data: " + data + "\nFinancial Data: " + formatted_data
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(query)
    
    # Log the response object to understand its structure
    logging.info(f"Model response: {response}")
    
     # Extract the text content from the response
    try:
        response_text = response.text #result.candidates[0].content.parts[0].text  #response.candidates[0]['content']['parts'][0]['text']
        format_response = markdown_to_text(response_text)
    except Exception as e:
        logging.error(f"Error extracting text from response: {e}")
        return f"Average closing price for the last 3 months: ${avg_close:.2f}", response
    
    return f"Average closing price for the last 3 months: ${avg_close:.2f}", format_response


def extract_excel_data(file_path):
    financial_data = ""
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        financial_data += f"\n\nSheet: {sheet_name}\n"
        financial_data += df.to_string()
    return financial_data

def markdown_to_text(md):
    # Simple conversion for markdown to plain text
    md = md.replace('**', '')
    md = md.replace('*', '')
    md = md.replace('_', '')
    md = md.replace('#', '')
    md = md.replace('`', '')
    return md.strip()


# @dispatcher.message_handler()
# @dp.message()
# async def main_bot(message: types.Message):
#     global retriever, extracted_text,investment_personality,summary,chat_history

#     # Handle the first tasks assessments answers from the user
#     chat_id = message.chat.id

#     if chat_id in states and states[chat_id] < len(questions):
#         # Retrieve the index of the current question
#         question_index = states[chat_id]

#         # Save the user's response to the current question
#         answer = message.text
#         user_responses[questions[question_index]] = answer
#         states[chat_id] += 1  # Move to the next question

#         # Ask the next question
#         await ask_next_question(chat_id, question_index + 1)
#     # Handle normal text messages
#     elif message.text:
#         # Check if the message is a stock ticker
#         if message.text.upper().isalpha() and len(message.text) <= 5:
#             ticker = message.text.upper()
#             try:
#                 hist,data,excel_file = await fetch_stock_data(message,ticker) #hist,data 
#                 mean_price,analysis = await analyze_stock_data(hist,data,excel_file)
#                 await message.reply(f"{mean_price}")
#                 await message.reply(f"Stock analysis for {ticker}:\n{analysis}")
#             except Exception as e:
#                 logging.error(f"Error fetching stock data for {ticker}: {e}")
#                 await message.reply("Failed to fetch stock data. Please try again.")

#         # Handle q&a chat messages using your Gemini model (llm)
#         else:
#             try:

#                 task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
#                 Also give the user detailed information about the investment how to invest,where to invest and how much they
#                 should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
#                 investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
#                 investment.Also explain the user why you are giving them that particular
#                 investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
#                 User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
#                 So dont mention user to consult to a financial expert."""
            
#                 # model = genai.GenerativeModel('gemini-1.5-flash') #('gemini-pro')
#                 print(investment_personality)
#                 query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
#                 print(f"\nQuery : {query}")
#                 # response = model.generate_content(query)
#                 model = genai.GenerativeModel('gemini-1.5-flash')
#                 chat = model.start_chat(history=[])
#                 print(f"Chat Conversation History : \n {chat.history}")
#                 response = chat.send_message(query)
#                 format_response = markdown_to_text(response.text)
#                 await message.reply(format_response) #(response.text) #(response['answer']) 
            
#             except Exception as e:
#                 print(f"Error processing general chat message: {e}")
#                 await message.reply("Failed to process your request.")


# # Stock data functions
# import yfinance as yf
# import requests
# import pandas as pd
# import io
# import matplotlib.pyplot as plt
# import requests
# from aiogram import types
# from aiogram.types import InputFile

# # Define the API key for the news API (replace with your own API key)
# NEWS_API_KEY = os.getenv('NEWS_API_KEY')
# import logging

# # Stock data functions
# async def fetch_stock_data(message, ticker):
#     stock = yf.Ticker(ticker)
#     chat_id = message.chat.id

#     # Dictionary to store stock info for messages
#     data = {}

#     # Stock Summary
#     company_details = stock.info.get('longBusinessSummary', 'No details available')
#     data['Company Details'] = company_details

#     # Stock Sector
#     sector = stock.info.get('sector', 'No sector information available')
#     data['Sector'] = sector

#     # Previous Closing Price
#     prev_close = stock.info.get('previousClose', 'No previous close price available')
#     data['Previous Closing Price'] = prev_close

#     # Today's Opening Price
#     open_price = stock.info.get('open', 'No opening price available')
#     data['Today Opening Price'] = open_price

#     # Fetch historical data for the last 5 days
#     hist = stock.history(period="5d")
#     if not hist.empty and 'Close' in hist.columns:
#         if hist.index[-1].date() == yf.download(ticker, period="1d").index[-1].date():
#             close_price = hist['Close'].iloc[-1]
#             data['Todays Closing Price'] = close_price
#         else:
#             data['Todays Closing Price'] = "Market is open, there is no closing price available yet."
#     else:
#         data['Todays Closing Price'] = "No historical data available for closing price."

#     # Today's High Price
#     day_high = stock.info.get('dayHigh', 'No high price available')
#     data['Today High Price'] = day_high

#     # Today's Low Price
#     day_low = stock.info.get('dayLow', 'No low price available')
#     data['Today Low Price'] = day_low

#     # Today's Volume
#     volume = stock.info.get('volume', 'No volume information available')
#     data['Today Volume'] = volume

#     # Today's Dividends
#     dividends = stock.info.get('dividendRate', 'No dividend information available')
#     data['Today Dividends'] = dividends

#     # Today's Stock Splits
#     splits = stock.info.get('lastSplitFactor', 'No stock split information available')
#     data['Today Stock Splits'] = splits

#     # P/E Ratio
#     pe_ratio = stock.info.get('trailingPE', 'No P/E ratio available')
#     data['P/E Ratio'] = pe_ratio

#     # Market Cap
#     market_cap = stock.info.get('marketCap', 'No market cap available')
#     data['Market Cap'] = market_cap

#     # Fetch Income Statement, Balance Sheet, and Cashflow data
#     income_statement = stock.financials
#     balance_sheet = stock.balance_sheet
#     cashflow = stock.cashflow

#     # Fetch top 3 news articles
#     news_url = f'https://newsapi.org/v2/everything?q={ticker}&apiKey={NEWS_API_KEY}&pageSize=3'
#     news_response = requests.get(news_url)
#     if news_response.status_code == 200:
#         news_data = news_response.json()
#         articles = news_data.get('articles', [])
#         if articles:
#             top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
#             data['Top News'] = top_news
#         else:
#             data['Top News'] = "No news articles found."
#     else:
#         data['Top News'] = "Failed to fetch news articles."

#     # Send collected data as messages
#     for key, value in data.items():
#         await message.reply(f"{key}: {value}")

#     # Stock Graph for the last 5 days using Yahoo Finance URL
#     graph_url = f"https://finance.yahoo.com/chart/{ticker}"
#     await message.reply(f"Stock Chart: \n{graph_url}")

#     # Create a DataFrame for the financial data and save to an Excel file
#     file_path = os.path.join('data', f'{ticker}_financial_data.xlsx')
#     with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
#         income_statement.to_excel(writer, sheet_name='Income Statement')
#         balance_sheet.to_excel(writer, sheet_name='Balance Sheet')
#         cashflow.to_excel(writer, sheet_name='Cashflow')
    
#     # Send the Excel file to the user
#     excel_file = FSInputFile(file_path, filename=f'{ticker}_financial_data.xlsx')
#     await bot.send_document(chat_id, excel_file, caption=f'{ticker} Financial Data')

#     # Convert Dictionary to List using items()
#     data_list = list(data.items())
#     #Converting List to string 
#     data_str = str(data_list)

#     return hist,data_str,file_path #,excel_file

# async def analyze_stock_data(hist, data, excel_file):
#     avg_close = hist['Close'].mean()
#     formatted_data = extract_excel_data(excel_file)

#     task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.
#                 Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
#                 Predict the stock price range for the next month and provide reasons for your prediction.
#                 Advise whether to buy this stock now or not, with reasons for your advice."""

#     query = task + "\nStock Data: " + data + "\nFinancial Data: " + formatted_data
#     model = genai.GenerativeModel('gemini-1.5-flash')
#     response = model.generate_content(query)
#     if response is None or not response.candidates:
#         return f"Average closing price for the last 3 months : ${avg_close:.2f}", "No valid response received from the model."

#     format_response = markdown_to_text(response.candidates[0].text)
#     return f"Average closing price for the last 3 months : ${avg_close:.2f}", format_response


# def extract_excel_data(file_path):
#     financial_data = ""
#     xls = pd.ExcelFile(file_path)
#     for sheet_name in xls.sheet_names:
#         df = pd.read_excel(xls, sheet_name=sheet_name)
#         financial_data += f"\n\nSheet: {sheet_name}\n"
#         financial_data += df.to_string()
#     return financial_data






# def analyze_stock_data(hist,data,excel_file):
#     # Simple analysis: Calculate average closing price
#     avg_close = hist['Close'].mean()

#     task = f"""You are a Stock Market Expert.You know everything about Stock Market its's trends and Patterns.
#                 You will be provided with data of a stock.
#                 Based on the data provided to you about a stock in a excel file : {excel_file}, provide more detailed analysis specifically 
#                 on the performance of the stock and whether the stock is performing well or not,whether it is overvalued or
#                 undervalued. Based on this analysis predict the stock price for the next month.
#                 Give a range of what you think the Stock Price will be based on the data provided to you.
#                 Give Reasons why you have predicted this range of the stock price.
#                 Also tell whether you will advice someone to buy this stock now or not.Give reasons why you gave this advice.
#                 """

#     query = task + data 
#     model = genai.GenerativeModel('gemini-1.5-flash')
#     response = model.generate_content(query)
#     print(response.text)
#     format_response = markdown_to_text(response.text)
#     # await message.reply(format_response)
#     return f"Average closing price for the last 3 months : ${avg_close:.2f}",format_response


from aiogram.types.input_file import BufferedInputFile
from aiogram import BaseMiddleware
# from aiogram.dispatcher.router import Router
from PIL import Image

# Function to handle image messages
# @dp.message(F.photo)
# @router.message(F.photo)
import PIL.Image

async def handle_image(message: types.Message):
    global investment_personality, chat_history

    chat_id = message.chat.id
    # Handle image inputs
    try:
        # Obtain file information
        try:
            photo_id = message.document.file_id
            photo = await bot.get_file(photo_id)
            photo_path = photo.file_path
            # Download the file
            photo_file = await bot.download_file(photo_path, "data/uploaded_image.png")

        except Exception as e:
            print(f"Error downloading image: {e}")
            await bot.send_message(chat_id, "Error processing image. Please try again.")
            return
        
        # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
        #             Also give the user detailed information about the investment how to invest, where to invest and how much they
        #             should invest in terms of percentage of their investment amount. Give the user detailed information about the returns on their 
        #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compounded returns on their 
        #             investment. Also explain the user why you are giving them that particular
        #             investment suggestion. Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
        #             User should also invest as per their risk tolerance level. Since you are the financial advisor don't ask user to consult anyone else.
        #             So don't mention user to consult to a financial expert."""

        task = """You are a Financial Expert.You will be provided with a Financial Form from Boston Harbor.
                If you recieve any other image tell the user to Upload the Images of the form or upload the word document of the form.
                You are supposed to Respond to the user's Image query and If they ask for any information provide them the information in Detail.
                Be helpful and informative.Give proper information of any Financial terms the user may ask you.Address the user by their Client Name if provided.
                Also provide the user helpful links so that they can refer to the link for more information.
                If the image provided is not related to Finance then just answer about the image and any caption if provided.
                """

        prompt = message.caption if message.caption else ""  # Use the photo caption if available
        # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + prompt
        query = task + prompt 

        image =  PIL.Image.open('data/uploaded_image.png') #(photo_file) 
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(image)
        await bot.send_message(chat_id,"I will describe the image that was uploaded")
        format_response = markdown_to_text(response.text)
        await message.reply(format_response)
        # await message.reply(response.text)

        # chat = model.start_chat(history=[])
        # response = chat.send_message(query)
        # format_response = markdown_to_text(response.result)
        # await message.reply(format_response)

        response = model.generate_content([query, image])
        format_response = markdown_to_text(response.text)
        await message.reply(format_response)
        # await message.reply(response.text) 
    except Exception as e:
        logging.error(f"Error generating response for the image: {e}")
        await message.reply("There was an error generating response for the image. Please try again later.")
    # await message.reply("Cant process the image")
    # return



from aiogram.filters import command
from aiogram.types import bot_command


# Function to process the analysis result
def process_analysis_result(analysis_result):
    return {
        'initial_investment': analysis_result['initial_investment'],
        'monthly_investment': analysis_result['monthly_investment'],
        'return_min': analysis_result['return_min'],
        'return_max': analysis_result['return_max'],
        'growth_investment_min': analysis_result['growth_investment_min'],
        'growth_investment_max': analysis_result['growth_investment_max'],
        'fixed_income_min': analysis_result['fixed_income_min'],
        'fixed_income_max': analysis_result['fixed_income_max'],
        'cash_min': analysis_result['cash_min'],
        'cash_max': analysis_result['cash_max'],
        'time_horizon': analysis_result['time_horizon'],
    }

# Function to generate performance graphs
def generate_performance_graph(data, years):
    labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
    sizes = [
        (data['growth_investment_min'] + data['growth_investment_max']) / 2,
        (data['fixed_income_min'] + data['fixed_income_max']) / 2,
        (data['cash_min'] + data['cash_max']) / 2
    ]
    
    fig1, ax1 = plt.subplots()
    ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
    ax1.axis('equal')
    pie_chart_buffer = io.BytesIO()
    plt.savefig(pie_chart_buffer, format='png')
    pie_chart_buffer.seek(0)

    initial_investment = data['initial_investment']
    monthly_investment = data['monthly_investment']
    annual_return_rate = (data['return_min'] + data['return_max']) / 2 / 100

    investment_values = [initial_investment]
    for month in range(1, years * 12 + 1):
        investment_values.append(investment_values[-1] * (1 + annual_return_rate / 12) + monthly_investment)
    
    fig2, ax2 = plt.subplots()
    ax2.plot(range(len(investment_values)), investment_values, marker='o')
    ax2.set(title='Investment Growth Over Time', xlabel='Month', ylabel='Investment Value ($)')
    ax2.grid(True)
    investment_growth_buffer = io.BytesIO()
    plt.savefig(investment_growth_buffer, format='png')
    investment_growth_buffer.seek(0)

    return pie_chart_buffer, investment_growth_buffer
  

import markdown
from bs4 import BeautifulSoup

def markdown_to_text(markdown_text):
    # Convert markdown to HTML
    html = markdown.markdown(markdown_text)
    # Parse the HTML
    soup = BeautifulSoup(html, 'html.parser')
    # Extract plain text
    text = soup.get_text()
    return text


# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)

async def main() -> None:
    # Initialize Bot instance with default bot properties which will be passed to all API calls
    bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

    # And the run events dispatching
    await dp.start_polling(bot)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, stream=sys.stdout)
    asyncio.run(main())



# #best code so far now it can upload files and receive various forms of messages as well and provide us graphs that we want 
# # and also reply to images, give stock informations 

# import os
# import filetype
# import docx
# import PyPDF2
# import re
# from aiogram import Bot, Dispatcher, types
# from dotenv import load_dotenv
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
# from langchain_community.document_loaders import Docx2txtLoader
# from langchain_core.prompts import ChatPromptTemplate
# from langchain.chains import create_retrieval_chain
# from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# from langchain.chains.combine_documents import create_stuff_documents_chain
# from langchain.memory import ConversationSummaryMemory
# import asyncio
# import numpy as np
# import json
# import re
# import google.generativeai as genai
# import pathlib
# # Import things that are needed generically
# from langchain.pydantic_v1 import BaseModel, Field
# from langchain.tools import BaseTool, StructuredTool, tool

# from aiogram.client.default import DefaultBotProperties
# from aiogram.enums import ParseMode
# from aiogram.filters import CommandStart
# from aiogram.types import Message
# from aiogram import F
# from aiogram import Router
# import logging
# import sys
# from aiogram.filters import Command
# from aiogram.types import FSInputFile
# # from aiogram.utils import executor
# import io
# import matplotlib.pyplot as plt
# import seaborn as sns
# import aiohttp
# from aiogram.types import InputFile , BufferedInputFile
# import PIL.Image

# router = Router(name=__name__)

# load_dotenv()

# TOKEN = os.getenv("TOKEN")
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# os.environ["LANGCHAIN_TRACING_V2"]="true"
# os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # Configure generativeai with your API key
# genai.configure(api_key=GOOGLE_API_KEY)

# # Initialize bot
# bot = Bot(token=TOKEN)
# dp = Dispatcher()

# # Glbal variables
# rag_on = False
# retriever = None  # Store retriever globally
# summary = ""
# investment_personality = ""
# chat_history = ""
# previous_suggestions = ""

# class Reference:
#     def __init__(self):
#         self.response = ""


# reference = Reference()


# def clear_past():
#     reference.response = ""


# @router.message(F.text == "clear")
# async def clear(message: types.Message):
#     """
#     A handler to clear the previous conversation and context.
#     """
#     clear_past()
#     await message.reply("I've cleared the past conversation and context.")

# #Global Variables :

# # Store user states
# states = {}

# # Dictionary to hold question-answer pairs
# user_responses = {}
# #
# user_images = {}
# # Define Questions for assessment
# questions = [
#     """ 
#     1. You and your friend are betting on a series of coin tosses.

#     He always bets ₹2,000 on Heads

#     You always bet ₹2,000 on Tails

#     Winner of last 8 turns

#     You lost ₹8,000 in the last 4 turns!

#     If you were to bet one last time, what would you bet on:
#     a) heads or b) tails ?
#     """ ,
#     """
#     2. Imagine you are a contestant in a game show, and you are presented the following choices.

#     What would you prefer?
#     a) 50 percent chance of winning 15 gold coins 
#     b) 100 percent chance of winning 8 gold coins
#     """,
#     """
#     3. Ok, one more choice...

#     What would you prefer?
#     a) 50 percent chance of winning 15 gold coins 
#     b) 100 percent chance of winning 2 gold coins
#     """,
#     """
#     4. In general, how would your best friend describe your risk-taking tendencies?
#     a) A real gambler
#     b) Willing to take risks after completing adequate research
#     c) Cautious
#     d) Avoids risk as much as possible
#     """,
#     """
#     5. Suppose you could replace your current investment portfolio with this new one:
#     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
#     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
#     a)-5 to -10
#     b)-10 to -15
#     c)-15 to -20
#     d)-20 to -25
#     e)-25 to -30
#     f)-30 to -35
#     """,
#     """
#     6. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
#     a) Option A : Min Value is 9.2L and Max Value is 11.8L
#     b) Option B : Min Value is 8.8L and Max Value is 12.3L
#     c) Option C : Min Value is 8.5L and Max Value is 12.8L
#     d) Option D : Min Value is 8.1L and Max Value is 13.3L
#     e) Option E : Min Value is 7.8L and Max Value is 13.8L
#     """,
#     """
#     7. From Sept 2008 to Nov 2008, Stock market went down by 31%.

#     If you owned a stock investment that lost about 31 percent in 3 months, you would:
#     a) Sell all of the remaining investment
#     b) Sell a portion of the remaining investment
#     c) Hold on to the investment and sell nothing
#     d) Buy little
#     e) Buy more of the investment
#     """,
#     """
#     8. Over any 1-year period, what would be the maximum drop in the value of your investment 
#     portfolio that you would be comfortable with?
#     a) <5%
#     b) 5 - 10%
#     c) 10 - 15%
#     d) 15 - 20%
#     e) >20%
#     """,
#     """
#     9. When investing, what do you consider the most?

#     a) Risk 
#     b) Return
#     """,
#     """
#     10. What best describes your attitude?

#     a) Prefer reasonable returns, can take reasonable risk
#     b) Like higher returns, can take slightly higher risk
#     c) Want to maximize returns, can take significant high risk
#     """,
#     """
#     11. How much monthly investment you want to do?
#     """,
#     """
#     12. What is the time horizon for your investment?
#     You can answer in any range, example 1-5 years."""  
# ]


# # Handler for /start command
# @dp.message(CommandStart())
# async def handle_start(message: types.Message):
#     """
#     This handler receives messages with /start command
#     """
#     chat_id = message.chat.id
#     # Start asking questions
#     await start_assessment(chat_id)


# # Function to start the assessment
# async def start_assessment(chat_id):
#     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
#     await ask_next_question(chat_id, 0)

# # Function to ask the next question
# async def ask_next_question(chat_id, question_index):
#     if question_index < len(questions):
#         # Ask the next question
#         await bot.send_message(chat_id, questions[question_index])
#         # Update state to indicate the next expected answer
#         states[chat_id] = question_index
#     else:
#         # No more questions, finish assessment
#         await finish_assessment(chat_id)

# # Handler for receiving assessment answers
# assessment_in_progress = True

# from aiogram.types import FSInputFile
# async def finish_assessment(chat_id):
#     if chat_id in states and states[chat_id] == len(questions):
#         # All questions have been answered, now process the assessment
#         await bot.send_message(chat_id, "Assessment completed. Thank you!")

#         # Determine investment personality based on collected responses
#         global investment_personality
#         investment_personality = await determine_investment_personality(user_responses)

#         # Inform the user about their investment personality
#         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

#         # Summarize collected information
#         global summary
#         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
#         summary = summary + "\n" + "Your investment personality:" + investment_personality
#         # Ensure to await the determination of investment personality
#         await send_summary_chunks(chat_id, summary)
#         global assessment_in_progress 
#         assessment_in_progress = False
       
#         await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
#         file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

#         await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
#         # await bot.send_message(chat_id,file)

# async def send_summary_chunks(chat_id, summary):
#     # Split the summary into chunks that fit within Telegram's message limits
#     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

#     # Send each chunk as a separate message
#     for chunk in chunks:
#         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# async def determine_investment_personality(assessment_data):
#     try:
#         # Prepare input text for the chatbot based on assessment data
#         input_text = "User Profile:\n"
#         for question, answer in assessment_data.items():
#             input_text += f"{question}: {answer}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
#                       "- Conservative Investor\n" \
#                       "- Moderate Investor\n" \
#                       "- Aggressive Investor"

#         # Use your generative AI model to generate a response
#         # print(input_text)
#         model = genai.GenerativeModel('gemini-pro')
#         response = model.generate_content(input_text)

#         # Determine the investment personality from the chatbot's response
#         response_text = response.text.lower()
#         if "conservative" in response_text:
#             personality = "Conservative Investor"
#         elif "moderate" in response_text:
#             personality = "Moderate Investor"
#         elif "aggressive" in response_text:
#             personality = "Aggressive Investor"
#         else:
#             personality = "Unknown"

#         return personality
#         # Send the determined investment personality back to the user
#         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

#     except Exception as e:
#         print(f"Error generating response: {e}")
#         #await bot.send_message(chat_id, "Error processing investment personality classification.")


# # Handler for document upload
# async def load_vector_db(file_path):
#     try:
#         loader = Docx2txtLoader(file_path)
#         documents = loader.load()
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         text_chunks = text_splitter.split_documents(documents)
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
#         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
#         return vector_store.as_retriever(search_kwargs={"k": 1})
#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None



# async def make_retrieval_chain(retriever):
#     """
#     Create a retrieval chain using the provided retriever.

#     Args:
#         retriever (RetrievalQA): A retriever object.

#     Returns:
#         RetrievalQA: A retrieval chain object.
#     """
#     try:
#         global investment_personality,summary
#         llm = ChatGoogleGenerativeAI(
#             #model="gemini-pro",
#             model = "gemini-1.5-flash",
#             temperature=0.7,
#             top_p=0.85,
#             google_api_key=GOOGLE_API_KEY
#         )

#         prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
#                 Respond to the client by the client name.
#                 Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
#                 Also give the user detailed information about the investment how to invest,where to invest and how much they
#                 should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
#                 Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
#                 investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
#                 Also give the user minimum and maximum expected growth in dollars for the time horizon .
#                 Also explain the user why you are giving them that particular investment suggestion.
#                 Give the client suggestions of Investment based on their investment personality that can also help them manage their mortages and other liablilities if they have any,ignore if they dont have any liabilities.
#                 Answer in 3-4 lines.\n
#                 <context>
#                 {context}
#                 </context>
#                 Question: {input}"""

#         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

#         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
#         combine_docs_chain = None  

#         if retriever is not None :  
#             retriever_chain = create_retrieval_chain(retriever,document_chain) 
#             print(retriever_chain)
#             return retriever_chain
#         else:
#             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
#             return None

#     except Exception as e:
#         print(f"Error in creating chain: {e}")
#         return None

# from aiogram.filters import Filter

# # @router.message(F.document)
# @dp.message(F.document)
# async def handle_document(message: types.Message):
#     global summary,investment_personality  

#     chat_id = message.chat.id

#     # Obtain file information
#     file_id = message.document.file_id
#     file = await bot.get_file(file_id)
#     file_path = file.file_path
    
#     # Download the file
#     await bot.download_file(file_path, "data/uploaded_file")
    
#     # Process the uploaded document
#     extracted_text = await process_document("data/uploaded_file")
    
#     if extracted_text:
#         # Load vector database (assuming this is part of setting up the retriever)
#         retriever = await load_vector_db("data/uploaded_file")
#         file_path = 'data/uploaded_file'
#         client_name, validation_errors = await process_document(file_path)

#         # Print results
#         print(f"Client Name: {client_name}")
#         if validation_errors:
#             print("**Validation Errors:**")
#             for error in validation_errors:
#                 print(error)
#         else:
#             print("All fields are filled correctly.")
#         if client_name == None:
#             try:
#                 await message.reply("Processing the uploaded image")
#                 await handle_image(message) 
#                 return 
#             except Exception as e:
#                 await message.reply("error processing uploaded image")
#                 print(e)
#         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")

#         if retriever is None:
#             await message.reply("The retrieval chain is not set up. Please upload a document first.")
#             return

#         # Check if a valid chain can be created
#         chain = await make_retrieval_chain(retriever)
#         if chain is None:
#             await message.reply("Failed to create the retrieval chain.")
#             return
        
#         try:     
#             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
        
#             response = chain.invoke({"input": query})
#             print(response['answer'])
#             global chat_history
#             chat_history = response['answer'] 
#             print(f"\n Chat History : {chat_history}")
#             await message.reply(response['answer'])

#             try:
#                 graph_query = (
#                     # extracted_text + "\n" + 
#                     summary + "\n" + "investment_personality" + "\n" + response['answer'] + "\n" +
#                     "Please provide the following information in JSON format:\n" +
#                     "{\n"
#                     "  'growth_investment_min': <minimum percentage of growth-oriented investments>,\n"
#                     "  'growth_investment_max': <maximum percentage of growth-oriented investments>,\n"
#                     "  'fixed_income_min': <minimum percentage of fixed-income investments>,\n"
#                     "  'fixed_income_max': <maximum percentage of fixed-income investments>,\n"
#                     "  'cash_min': <minimum percentage of cash investments>,\n"
#                     "  'cash_max': <maximum percentage of cash investments>,\n"
#                     "  'return_growth_investment': <Maximum percentage returns of growth-oriented investments>,\n"
#                     "  'return_fixed_income': <Maximum percentage of returns of fixed-income investments>,\n"
#                     "  'return_cash_max': <maximum percentage of retruns of cash investments>,\n"
#                     "  'return_min': <minimum expected annual return percentage>,\n"
#                     "  'return_max': <maximum expected annual return percentage>,\n"
#                     "  'growth_min': <minimum expected growth in dollars>,\n"
#                     "  'growth_max': <maximum expected growth in dollars>,\n"
#                     "  'initial_investment': <initial monthly investment>,\n"
#                     "  'time_horizon': <time horizon in years>\n"
#                     "}"
#                 )
#                 graph_response = chain.invoke({"input": graph_query})
#                 print(graph_response['answer'])
#                 await handle_graph(graph_response['answer'],chat_id)
#                 try:
#                     await refine_investment_suggestions(chat_id, response['answer'])  # Generate refined suggestions
#                 except Exception as e:
#                     print(f"Error refining investment suggestions: {e}")
#                     await bot.send_message(chat_id, "Error refining investment suggestions. Please try again later.")

#             except Exception as e:
#                 print(f"Error plotting graph : {e}")
#         except Exception as e:
#             print(f"Error invoking retrieval chain on attempt : {e}")
#             await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

#     else:
#         await message.reply("Failed to process the uploaded file.")
    

# async def refine_investment_suggestions(chat_id, suggestions):
#     try:
#         task = """Based on the following investment suggestions, provide more detailed investment advice specifically 
#                   on Investment Allocation of stocks:
#                   How to invest in these stocks, what percentage of the investment should go into each stock, 
#                   the expected returns, and the reasoning behind each suggestion. 
#                   Also, include any potential risks and recommendations for monitoring the investments.
#                   Give the response in brief in around 15-20 lines."""

#         query = task + "\n" + suggestions
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content(query)
#         await bot.send_message(chat_id, "Here are more detailed stock investment suggestions:")
#         await bot.send_message(chat_id, response.text)
    
#     except Exception as e:
#         logging.error(f"Error generating refined investment suggestions: {e}")
#         await bot.send_message(chat_id, "There was an error generating more detailed investment suggestions. Please try again later.")

# # Function to extract data from LLM response
# def extract_data_from_response(response):
#     try:
#         # Locate the JSON-like data in the response
#         json_start = response.find("{")
#         json_end = response.rfind("}") + 1
        
#         if json_start == -1 or json_end == -1:
#             raise ValueError("No JSON data found in the response.")
        
#         json_data = response[json_start:json_end]
        
#         # Parse the JSON data
#         data = json.loads(json_data.replace("'", "\""))
#         print(data)
#         return data
#     except Exception as e:
#         logging.error(f"Error extracting data: {e}")
#         return None


# async def handle_graph(response, chat_id):
#     try:
#         data = extract_data_from_response(response)
#         if not data:
#             await bot.send_message(chat_id, "Failed to extract data from the response.")
#             return
        
#         # Log extracted data for debugging
#         logging.info(f"Extracted data: {data}")
        
#         # Create a pie chart for investment allocation
#         labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
#         sizes = [
#             (data['growth_investment_min'] + data['growth_investment_max']) / 2,
#             (data['fixed_income_min'] + data['fixed_income_max']) / 2,
#             (data['cash_min'] + data['cash_max']) / 2
#         ]
#         fig1, ax1 = plt.subplots()
#         ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
#         ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
#         pie_chart_buffer = io.BytesIO()
#         plt.savefig(pie_chart_buffer, format='png')
#         pie_chart_buffer.seek(0)
#         pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

#         # Create a bar graph for potential returns
#         allocations = {
#             'Growth-Oriented Investments': (data['growth_investment_min'] + data['growth_investment_max']) / 2,
#             'Fixed-Income Investments': (data['fixed_income_min'] + data['fixed_income_max']) / 2,
#             'Cash and Cash Equivalents': (data['cash_min'] + data['cash_max']) / 2
#         }

#         returns = {
#             'Growth-Oriented Investments': (data['return_growth_investment'] ) / 100,  # Annual return
#             'Fixed-Income Investments': (data['return_fixed_income'] ) / 100 ,
#             'Cash and Cash Equivalents': (data['return_cash_max'] ) / 100
#         }
        
#         time_horizon_years = data['time_horizon']
#         initial_monthly_investment = data['initial_investment']
        
#         # Calculate total returns for each category
#         total_investment = initial_monthly_investment * 12 * time_horizon_years
#         total_returns = {category: total_investment * (1 + returns[category]) ** time_horizon_years for category in allocations.keys()}
        
#         # Create the bar chart
#         fig2, ax2 = plt.subplots(figsize=(10, 6))
#         bars = ax2.bar(allocations.keys(), total_returns.values(), color=['#FF9999', '#66B2FF', '#99FF99'])

#         # Add labels to the top of each bar
#         for bar in bars:
#             yval = bar.get_height()
#             ax2.text(bar.get_x() + bar.get_width()/2, yval + 1000, f'₹{round(yval, 2)}', ha='center', va='bottom')

#         # Add titles and labels
#         plt.title('Investment Performance Over Time')
#         plt.xlabel('Investment Categories')
#         plt.ylabel('Total Returns (₹)')
#         plt.grid(True, axis='y', linestyle='--', linewidth=0.7)
        
#         # Save and display the plot
#         bar_chart_buffer = io.BytesIO()
#         plt.savefig(bar_chart_buffer, format='png')
#         bar_chart_buffer.seek(0)
#         bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='investment_performance.png')

#         await bot.send_document(chat_id, pie_chart_file, caption="Investment Allocation Chart")
#         await bot.send_document(chat_id, bar_chart_file, caption="Investment Performance Over Time")

#     except Exception as e:
#         logging.error(f"Error plotting graph: {e}")
#         await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")


# async def process_document(file_path):
#     try:
#         file_type = filetype.guess(file_path)
#         if file_type is not None:
#             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 # return extract_text_from_word(file_path)
#                 return extract_text_and_tables_from_word(file_path)
#             elif file_type.mime == "application/pdf":
#                 return extract_text_from_pdf(file_path)
#         return None
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None

# def extract_text_from_pdf(pdf_file_path):
#     try:
#         with open(pdf_file_path, "rb") as pdf_file:
#             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#             text_content = []
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text_content.append(page.extract_text())
#             return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None


# import re
# import docx

# def extract_text_and_tables_from_word(docx_file_path):
#     """
#     Extracts text and tables from a Word document (.docx).

#     Args:
#         docx_file_path (str): Path to the Word document file.

#     Returns:
#         tuple: Extracted text content and tables from the document.
#     """
#     try:
#         doc = docx.Document(docx_file_path)
#         text_content = []
#         tables_content = []

#         for para in doc.paragraphs:
#             text_content.append(para.text)

#         for table in doc.tables:
#             table_data = []
#             for row in table.rows:
#                 row_data = []
#                 for cell in row.cells:
#                     row_data.append(cell.text.strip())
#                 table_data.append(row_data)
#             tables_content.append(table_data)

#         return "\n".join(text_content), tables_content
#     except Exception as e:
#         print(f"Error extracting text and tables from Word document: {e}")
#         return None, None

# def validate_document_content(text, tables):
#     """
#     Validates the content of the document.

#     Args:
#         text (str): Extracted text content from the document.
#         tables (list): Extracted tables content from the document.

#     Returns:
#         tuple: Client name and validation errors.
#     """
#     errors = []
    
#     # Extract client name
#     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
#     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

#     # Define required sections
#     required_sections = [
#         "YOUR RETIREMENT GOAL",
#         "YOUR OTHER MAJOR GOALS",
#         "YOUR ASSETS AND LIABILITIES",
#         "MY LIABILITIES",
#         "YOUR CURRENT ANNUAL INCOME"
#     ]

#     # Check for the presence of required sections
#     for section in required_sections:
#         if section not in text:
#             errors.append(f"* {section} section missing.")
    
#     # Define table field checks
#     table_checks = {
#         "YOUR RETIREMENT GOAL": [
#             r"When do you plan to retire\? \(age or date\)",
#             r"Social Security Benefit \(include expected start date\)",
#             r"Pension Benefit \(include expected start date\)",
#             r"Other Expected Income \(rental, part-time work, etc.\)",
#             r"Estimated Annual Retirement Expense"
#         ],
#         "YOUR OTHER MAJOR GOALS": [
#             r"GOAL", r"COST", r"WHEN"
#         ],
#         "YOUR ASSETS AND LIABILITIES": [
#             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
#             r"Current Value", r"Annual Contributions"
#         ],
#         "MY LIABILITIES": [
#             r"Balance", r"Interest Rate", r"Monthly Payment"
#         ]
#     }

#     # Validate table content
#     for section, checks in table_checks.items():
#         section_found = False
#         for table in tables:
#             table_text = "\n".join(["\t".join(row) for row in table])
#             if section in table_text:
#                 section_found = True
#                 for check in checks:
#                     if not re.search(check, table_text, re.IGNORECASE):
#                         errors.append(f"* Missing or empty field in {section} section: {check}")
#                 break
#         if not section_found:
#             errors.append(f"* {section} section missing.")

#     return client_name, errors

# async def process_document(file_path):
#     try:
#         text, tables = extract_text_and_tables_from_word(file_path)
#         if text is not None and tables is not None:
#             client_name, errors = validate_document_content(text, tables)
#             return client_name, errors
#         return None, ["Error processing document."]
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None, [f"Error processing document: {e}"]



# # @dispatcher.message_handler()
# @dp.message()
# async def main_bot(message: types.Message):
#     global retriever, extracted_text,investment_personality,summary,chat_history

#     # Handle the first tasks assessments answers from the user
#     chat_id = message.chat.id

#     if chat_id in states and states[chat_id] < len(questions):
#         # Retrieve the index of the current question
#         question_index = states[chat_id]

#         # Save the user's response to the current question
#         answer = message.text
#         user_responses[questions[question_index]] = answer
#         states[chat_id] += 1  # Move to the next question

#         # Ask the next question
#         await ask_next_question(chat_id, question_index + 1)
#     # Handle normal text messages
#     elif message.text:

#         # Check if message contains a stock ticker
#         # text = message.text.strip()
#         # if text.isupper() and len(text) <= 5:  # Assuming ticker symbols are in uppercase and up to 5 characters
#         #     hist = fetch_stock_data(text)
#         #     analysis = analyze_stock_data(hist)
#         #     await message.reply(f"Stock analysis for {text}:\n{analysis}")
#         #     return

#         # Check if the message is a stock ticker
#         if message.text.upper().isalpha() and len(message.text) <= 5:
#             ticker = message.text.upper()
#             try:
#                 hist = await fetch_stock_data(message,chat_id,ticker)
#                 analysis = analyze_stock_data(hist)
#                 await message.reply(f"Stock analysis for {ticker}:\n{analysis}")
#             except Exception as e:
#                 logging.error(f"Error fetching stock data for {ticker}: {e}")
#                 await message.reply("Failed to fetch stock data. Please try again.")

#         # Handle q&a chat messages using your Gemini model (llm)
#         else:
#             try:

#                 task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
#                 Also give the user detailed information about the investment how to invest,where to invest and how much they
#                 should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
#                 investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
#                 investment.Also explain the user why you are giving them that particular
#                 investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
#                 User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
#                 So dont mention user to consult to a financial expert."""
            
#                 # model = genai.GenerativeModel('gemini-1.5-flash') #('gemini-pro')
#                 print(investment_personality)
#                 query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
#                 print(f"\nQuery : {query}")
#                 # response = model.generate_content(query)
#                 model = genai.GenerativeModel('gemini-1.5-flash')
#                 chat = model.start_chat(history=[])
#                 print(f"Chat Conversation History : \n {chat.history}")
#                 response = chat.send_message(query)
#                 await message.reply(response.text) #(response['answer']) 
            
#             except Exception as e:
#                 print(f"Error processing general chat message: {e}")
#                 await message.reply("Failed to process your request.")


# # Stock data functions
# import yfinance as yf
# import requests
# import io
# async def fetch_stock_data(message,chat_id,ticker):
#     stock = yf.Ticker(ticker)
#     # Stock Summary :
#     company_details = stock.info['longBusinessSummary']
#     await message.reply(f"Comapany Details : \n{company_details}")

#     # Stock Sector :
#     sector = stock.info['sector'] 
#     await message.reply(f"Sector : \n{sector}")

#     #Previous Closing :
#     prev_close = stock.info.get('previousClose', 'No previous close price available')
#     await message.reply(f"Previous Closing Price: \n${prev_close}")

#     # Today's Opening Price
#     open_price = stock.info.get('open', 'No opening price available')
#     await message.reply(f"Today's Opening Price: \n${open_price}")
    
#     # Today's Closing Price

#     # Today's Closing Price
#     hist = stock.history(period="5d")
#     if not hist.empty and 'Close' in hist.columns:
#         if hist.index[-1].date() == yf.download(ticker, period="1d").index[-1].date():
#             close_price = hist['Close'].iloc[-1]
#             await message.reply(f"Today's Closing Price: \n${close_price}")
#         else:
#             await message.reply("Market is open, there is no closing price available yet.")
#     else:
#         await message.reply("No historical data available for closing price.")

#     # Today's High Price
#     day_high = stock.info.get('dayHigh', 'No high price available')
#     await message.reply(f"Today's High Price: \n${day_high}")
    
#     # Today's Low Price
#     day_low = stock.info.get('dayLow', 'No low price available')
#     await message.reply(f"Today's Low Price: \n${day_low}")
    
#     # Today's Volume
#     volume = stock.info.get('volume', 'No volume information available')
#     await message.reply(f"Today's Volume: \n{volume}")
    
#     # Today's Dividends
#     dividends = stock.info.get('dividendRate', 'No dividend information available')
#     await message.reply(f"Today's Dividends: \n${dividends}")
    
#     # Today's Stock Splits
#     splits = stock.info.get('lastSplitFactor', 'No stock split information available')
#     await message.reply(f"Today's Stock Splits: \n{splits}")
    
#      # Stock Graph for the last 5 days using Yahoo Finance URL
#     graph_url = f"https://finance.yahoo.com/chart/{ticker}"
#     await message.reply(f"Stock Chart: \n{graph_url}")
    
#     # # Generate Stock Graph for the last 5 days :
#     # try:
#     #     # Generate the graph
#     #     hist['Close'].plot(title=f'{ticker} Closing Prices - Last 5 Days')
#     #     plt.xlabel("Date")
#     #     plt.ylabel("Closing Price")
        
#     #     # Save the plot to a BytesIO object
#     #     buf = io.BytesIO()
#     #     plt.savefig(buf, format='png')
#     #     buf.seek(0)
#     #     plt.close()

#     #     # Send the image as a document
#     #     await bot.send_document(chat_id, ('stock_graph.png', buf), caption=f'{ticker} Closing Prices - Last 5 Days')
#     # except Exception as e:
#     #     await message.reply(f"Failed to generate stock graph: {e}")

#     hist = stock.history(period="3mo")  # Fetch data for the past 3 months
#     return hist
    

# def analyze_stock_data(hist):
#     # Simple analysis: Calculate average closing price
#     avg_close = hist['Close'].mean()
#     return f"Average closing price for the last 3 months : ${avg_close:.2f}"

# from aiogram.types.input_file import BufferedInputFile
# from aiogram import BaseMiddleware
# # from aiogram.dispatcher.router import Router
# from PIL import Image

# # Function to handle image messages
# # @dp.message(F.photo)
# # @router.message(F.photo)
# import PIL.Image

# async def handle_image(message: types.Message):
#     global investment_personality, chat_history

#     chat_id = message.chat.id
#     # Handle image inputs
#     try:
#         # Obtain file information
#         try:
#             photo_id = message.document.file_id
#             photo = await bot.get_file(photo_id)
#             photo_path = photo.file_path
#             # Download the file
#             photo_file = await bot.download_file(photo_path, "data/uploaded_image.png")

#         except Exception as e:
#             print(f"Error downloading image: {e}")
#             await bot.send_message(chat_id, "Error processing image. Please try again.")
#             return
        
#         # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
#         #             Also give the user detailed information about the investment how to invest, where to invest and how much they
#         #             should invest in terms of percentage of their investment amount. Give the user detailed information about the returns on their 
#         #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compounded returns on their 
#         #             investment. Also explain the user why you are giving them that particular
#         #             investment suggestion. Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
#         #             User should also invest as per their risk tolerance level. Since you are the financial advisor don't ask user to consult anyone else.
#         #             So don't mention user to consult to a financial expert."""

#         task = """You are a Financial Expert.You will be provided with a Financial Form from Boston Harbor.
#                 If you recieve any other image tell the user to Upload the Images of the form or upload the word document of the form.
#                 You are supposed to Respond to the user's Image query and If they ask for any information provide them the information in Detail.
#                 Be helpful and informative.Give proper information of any Financial terms the user may ask you.Address the user by their Client Name if provided.
#                 Also provide the user helpful links so that they can refer to the link for more information.
#                 If the image provided is not related to Finance then just answer about the image and any caption if provided.
#                 """

#         prompt = message.caption if message.caption else ""  # Use the photo caption if available
#         # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + prompt
#         query = task + prompt 

#         image =  PIL.Image.open('data/uploaded_image.png') #(photo_file) 
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content(image)
#         await bot.send_message(chat_id,"I will describe the image that was uploaded")
#         await message.reply(response.text)


#         response = model.generate_content([query, image])
#         await message.reply(response.text) 
#     except Exception as e:
#         logging.error(f"Error generating response for the image: {e}")
#         await message.reply("There was an error generating response for the image. Please try again later.")
#     # await message.reply("Cant process the image")
#     # return



# from aiogram.filters import command
# from aiogram.types import bot_command


# # Function to process the analysis result
# def process_analysis_result(analysis_result):
#     return {
#         'initial_investment': analysis_result['initial_investment'],
#         'monthly_investment': analysis_result['monthly_investment'],
#         'return_min': analysis_result['return_min'],
#         'return_max': analysis_result['return_max'],
#         'growth_investment_min': analysis_result['growth_investment_min'],
#         'growth_investment_max': analysis_result['growth_investment_max'],
#         'fixed_income_min': analysis_result['fixed_income_min'],
#         'fixed_income_max': analysis_result['fixed_income_max'],
#         'cash_min': analysis_result['cash_min'],
#         'cash_max': analysis_result['cash_max'],
#         'time_horizon': analysis_result['time_horizon'],
#     }

# # Function to generate performance graphs
# def generate_performance_graph(data, years):
#     labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
#     sizes = [
#         (data['growth_investment_min'] + data['growth_investment_max']) / 2,
#         (data['fixed_income_min'] + data['fixed_income_max']) / 2,
#         (data['cash_min'] + data['cash_max']) / 2
#     ]
    
#     fig1, ax1 = plt.subplots()
#     ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
#     ax1.axis('equal')
#     pie_chart_buffer = io.BytesIO()
#     plt.savefig(pie_chart_buffer, format='png')
#     pie_chart_buffer.seek(0)

#     initial_investment = data['initial_investment']
#     monthly_investment = data['monthly_investment']
#     annual_return_rate = (data['return_min'] + data['return_max']) / 2 / 100

#     investment_values = [initial_investment]
#     for month in range(1, years * 12 + 1):
#         investment_values.append(investment_values[-1] * (1 + annual_return_rate / 12) + monthly_investment)
    
#     fig2, ax2 = plt.subplots()
#     ax2.plot(range(len(investment_values)), investment_values, marker='o')
#     ax2.set(title='Investment Growth Over Time', xlabel='Month', ylabel='Investment Value ($)')
#     ax2.grid(True)
#     investment_growth_buffer = io.BytesIO()
#     plt.savefig(investment_growth_buffer, format='png')
#     investment_growth_buffer.seek(0)

#     return pie_chart_buffer, investment_growth_buffer

        

# # if __name__ == "__main__":
# #     executor.start_polling(dispatcher, skip_updates=True)

# async def main() -> None:
#     # Initialize Bot instance with default bot properties which will be passed to all API calls
#     bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

#     # And the run events dispatching
#     await dp.start_polling(bot)


# if __name__ == "__main__":
#     logging.basicConfig(level=logging.INFO, stream=sys.stdout)
#     asyncio.run(main())






# # #best code so far now it can upload files and receive various forms of messages as well and provide us graphs that we want 
# # # and also reply to images

# # import os
# # import filetype
# # import docx
# # import PyPDF2
# # import re
# # from aiogram import Bot, Dispatcher, types
# # from dotenv import load_dotenv
# # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # from langchain_community.vectorstores import Chroma
# # from langchain_community.document_loaders import Docx2txtLoader
# # from langchain_core.prompts import ChatPromptTemplate
# # from langchain.chains import create_retrieval_chain
# # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # from langchain.chains.combine_documents import create_stuff_documents_chain
# # from langchain.memory import ConversationSummaryMemory
# # import asyncio
# # import numpy as np
# # import json
# # import re
# # import google.generativeai as genai
# # import pathlib
# # # Import things that are needed generically
# # from langchain.pydantic_v1 import BaseModel, Field
# # from langchain.tools import BaseTool, StructuredTool, tool

# # from aiogram.client.default import DefaultBotProperties
# # from aiogram.enums import ParseMode
# # from aiogram.filters import CommandStart
# # from aiogram.types import Message
# # from aiogram import F
# # from aiogram import Router
# # import logging
# # import sys
# # from aiogram.filters import Command
# # from aiogram.types import FSInputFile
# # # from aiogram.utils import executor
# # import io
# # import matplotlib.pyplot as plt
# # import seaborn as sns
# # import aiohttp
# # from aiogram.types import InputFile , BufferedInputFile
# # import PIL.Image

# # router = Router(name=__name__)

# # load_dotenv()

# # TOKEN = os.getenv("TOKEN")
# # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # os.environ["LANGCHAIN_TRACING_V2"]="true"
# # os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # # Configure generativeai with your API key
# # genai.configure(api_key=GOOGLE_API_KEY)

# # # Initialize bot
# # bot = Bot(token=TOKEN)
# # dp = Dispatcher()

# # rag_on = False
# # retriever = None  # Store retriever globally
# # summary = ""
# # investment_personality = ""
# # chat_history = ""

# # class Reference:
# #     def __init__(self):
# #         self.response = ""


# # reference = Reference()


# # def clear_past():
# #     reference.response = ""


# # @router.message(F.text == "clear")
# # async def clear(message: types.Message):
# #     """
# #     A handler to clear the previous conversation and context.
# #     """
# #     clear_past()
# #     await message.reply("I've cleared the past conversation and context.")


# # # Store user states
# # states = {}
# # # Dictionary to hold question-answer pairs
# # user_responses = {}

# # # Define Questions for assessment
# # questions = [
# #     """
# #         1. Singapore plans to build a new observation tower called 'The Rook'.
# #         How many steps do you think it will take to go to the top floor?

# #         a) Less than 500 
# #         b) More than 500

# #     """,
# #     "2. Now Guess the number of steps" ,
# #     """
# #     3. How confident are you that the real number is in the range you have selected? 
# #     Answer within a range of 100.  
# #     """,
# #     """ 
# #     4. You and your friend are betting on a series of coin tosses.

# #     He always bets ₹2,000 on Heads

# #     You always bet ₹2,000 on Tails

# #     Winner of last 8 turns

# #     You lost ₹8,000 in the last 4 turns!

# #     If you were to bet one last time, what would you bet on heads or tails ?
# #     """ ,
# #     """
# #     5. How confident are you that your bet will win this time?
# #     Answer how confident you are. 
# #     (Example: Not confident at all, somewhat confident, confident, or Very confident)
# #     """,
# #     """
# #     6. What do you think are the chances that you will achieve and maintain your financial goals within the next 
# #     10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
# #     Answer how likely you are to achieve your goal.
# #     (Example: Less likely than others, likely than others, or More likely than others)
# #     """,
# #     """
# #     7. Imagine you are a contestant in a game show, and you are presented the following choices.

# #     What would you prefer?
# #     a) 50 percent chance of winning 15 gold coins 
# #     b) 100 percent chance of winning 8 gold coins
# #     """,
# #     """
# #     8. Ok, one last choice...

# #     What would you prefer?
# #     a) 50 percent chance of winning 15 gold coins 
# #     b) 100 percent chance of winning 2 gold coins
# #     """,
# #     """
# #     9. In general, how would your best friend describe your risk-taking tendencies?
# #     a) A real gambler
# #     b) Willing to take risks after completing adequate research
# #     c) Cautious
# #     d) Avoids risk as much as possible
# #     """,
# #     """
# #     10. Suppose you could replace your current investment portfolio with this new one:
# #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# #     Answer between the range of -5 to -35.
# #     """,
# #     """
# #     11. Suppose that in the next 7 years,

# #     YOUR INCOME

# #     grows 8% each year

# #     VS
# #     INFLATION

# #     grows 10% a year

# #     At the end of 7 years, how much will you be able to buy with your income?
# #     Options:
# #     a) More than today
# #     b) Exactly the same
# #     c) Less than today
# #     d) Cannot say
# #     """,
# #     """
# #     12. If somebody buys a bond of Company B, which of the following statements seems correct:
# #     a) She owns part of Company B
# #     b) She has lent money to Company B
# #     c) She is liable for Company B's debt
# #     d) Cannot say
# #     """,
# #     """
# #     13. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
# #     a) More than ₹2 lakhs
# #     b) Less than ₹2 lakhs
# #     c) Exactly ₹2 lakhs
# #     d) Cannot say
# #     """,
# #     """
# #     14. When an investor spreads money across different asset classes, what happens to the risk of losing money:
# #     a) Increases
# #     b) Decreases
# #     c) Stays the same
# #     d) Cannot say
# #     """,
# #     """
# #     15. When a country's central bank reduces interest rates, it makes:

# #     a) Borrowing more attractive and saving less attractive
# #     b) Borrowing less attractive and saving more attractive
# #     c) Both borrowing and saving less attractive
# #     d) Cannot say
# #     """,
# #     """
# #     16. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
# #     a) Option A : Min Value is 9.2L and Max Value is 11.8L
# #     b) Option B : Min Value is 8.8L and Max Value is 12.3L
# #     c) Option C : Min Value is 8.5L and Max Value is 12.8L
# #     d) Option D : Min Value is 8.1L and Max Value is 13.3L
# #     e) Option E : Min Value is 7.8L and Max Value is 13.8L
# #     """,
# #     """
# #     17. From Sept 2008 to Nov 2008, Stock market went down by 31%.

# #     If you owned a stock investment that lost about 31 percent in 3 months, you would:
# #     a) Sell all of the remaining investment
# #     b) Sell a portion of the remaining investment
# #     c) Hold on to the investment and sell nothing
# #     d) Buy little
# #     e) Buy more of the investment
# #     """,
# #     """
# #     18. Over any 1-year period, what would be the maximum drop in the value of your investment 
# #     portfolio that you would be comfortable with?
# #     a) <5%
# #     b) 5 - 10%
# #     c) 10 - 15%
# #     d) 15 - 20%
# #     e) >20%
# #     """,
# #     """
# #     19. When investing, what do you consider the most?

# #     a) Risk 
# #     b) Return
# #     """,
# #     """
# #     20. What best describes your attitude?

# #     a) Prefer reasonable returns, can take reasonable risk
# #     b) Like higher returns, can take slightly higher risk
# #     c) Want to maximize returns, can take significant high risk
# #     """,
# #     """
# #     21. How much monthly investment you want to do?
# #     """,
# #     """
# #     22. What is the time horizon for your investment?
# #     You can answer in any range, example 1-5 years."""  
# # ]


# # # Handler for /start command
# # @dp.message(CommandStart())
# # async def handle_start(message: types.Message):
# #     """
# #     This handler receives messages with /start command
# #     """
# #     chat_id = message.chat.id
# #     # Start asking questions
# #     await start_assessment(chat_id)


# # # Function to start the assessment
# # async def start_assessment(chat_id):
# #     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
# #     await ask_next_question(chat_id, 0)

# # # Function to ask the next question
# # async def ask_next_question(chat_id, question_index):
# #     if question_index < len(questions):
# #         # Ask the next question
# #         await bot.send_message(chat_id, questions[question_index])
# #         # Update state to indicate the next expected answer
# #         states[chat_id] = question_index
# #     else:
# #         # No more questions, finish assessment
# #         await finish_assessment(chat_id)

# # # Handler for receiving assessment answers
# # assessment_in_progress = True

# # from aiogram.types import FSInputFile
# # async def finish_assessment(chat_id):
# #     if chat_id in states and states[chat_id] == len(questions):
# #         # All questions have been answered, now process the assessment
# #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# #         # Determine investment personality based on collected responses
# #         global investment_personality
# #         investment_personality = await determine_investment_personality(user_responses)

# #         # Inform the user about their investment personality
# #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# #         # Summarize collected information
# #         global summary
# #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# #         # Ensure to await the determination of investment personality
# #         await send_summary_chunks(chat_id, summary)
# #         global assessment_in_progress 
# #         assessment_in_progress = False
# #         # Prompt the user to begin financial advice process
# #         # await bot.send_message(chat_id, "To receive financial advice, please upload a document with your financial details using /begin.")

# #         await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
# #         file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

# #         await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
# #         # await bot.send_message(chat_id,file)

# # async def send_summary_chunks(chat_id, summary):
# #     # Split the summary into chunks that fit within Telegram's message limits
# #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# #     # Send each chunk as a separate message
# #     for chunk in chunks:
# #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # async def determine_investment_personality(assessment_data):
# #     try:
# #         # Prepare input text for the chatbot based on assessment data
# #         input_text = "User Profile:\n"
# #         for question, answer in assessment_data.items():
# #             input_text += f"{question}: {answer}\n"

# #         # Introduce the chatbot's task and prompt for classification
# #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# #                       "- Conservative Investor\n" \
# #                       "- Moderate Investor\n" \
# #                       "- Aggressive Investor"

# #         # Use your generative AI model to generate a response
# #         # print(input_text)
# #         model = genai.GenerativeModel('gemini-pro')
# #         response = model.generate_content(input_text)

# #         # Determine the investment personality from the chatbot's response
# #         response_text = response.text.lower()
# #         if "conservative" in response_text:
# #             personality = "Conservative Investor"
# #         elif "moderate" in response_text:
# #             personality = "Moderate Investor"
# #         elif "aggressive" in response_text:
# #             personality = "Aggressive Investor"
# #         else:
# #             personality = "Unknown"

# #         return personality
# #         # Send the determined investment personality back to the user
# #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# #     except Exception as e:
# #         print(f"Error generating response: {e}")
# #         #await bot.send_message(chat_id, "Error processing investment personality classification.")




# # @router.message(F.text == "help")
# # async def helper(message: types.Message):
# #     """
# #     A handler to display the help menu.
# #     """
# #     help_command = """
# #     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
# #     /start - to start the investment personality assessment.
# #     /clear - to clear the past conversation and context.
# #     /help - to get this help menu.
# #     /begin - to start the Financial Suggestion Conversation with the ChatBot.
# #     I hope this helps. :)
# #     """
# #     await message.reply(help_command)

# # # Handler for /begin command to initiate financial advice

# # @router.message(F.text == "begin")
# # async def handle_begin(message: types.Message):
# #     chat_id = message.chat.id
# #     file_instructions ="""Hello there!My name is Finbot and I am a Wealth Management Advisor Chatbot.I need more details related to your Financial Profile so that I can give you 
# #     personalised Financial Advice.Please follow this drive link : https://docs.google.com/document/d/1JaqWUXcq3MNrPTCvEy_2RngbGIY0rybX/edit?usp=sharing&ouid=101252223236130106095&rtpof=true&sd=true
# #     ,there you will find a Word Document.Please fill in your details with correct Information and then upload the Document in the chat. """

# #     await message.reply(file_instructions)
    

# # # Function to handle image messages
# # # @dp.message(F.photo)
# # # @router.message(F.photo)
# # import PIL.Image

# # async def handle_image(message: types.Message):
# #     global investment_personality, chat_history

# #     chat_id = message.chat.id
# #     # Handle image inputs
# #     try:
# #         # Obtain file information
# #         try:
# #             photo_id = message.document.file_id
# #             photo = await bot.get_file(photo_id)
# #             photo_path = photo.file_path
# #             # Download the file
# #             photo_file = await bot.download_file(photo_path, "data/uploaded_image.png")

# #         except Exception as e:
# #             print(f"Error downloading image: {e}")
# #             await bot.send_message(chat_id, "Error processing image. Please try again.")
# #             return
        
# #         model = genai.GenerativeModel('gemini-1.5-flash')
  
# #         task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# #                     Also give the user detailed information about the investment how to invest, where to invest and how much they
# #                     should invest in terms of percentage of their investment amount. Give the user detailed information about the returns on their 
# #                     investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compounded returns on their 
# #                     investment. Also explain the user why you are giving them that particular
# #                     investment suggestion. Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# #                     User should also invest as per their risk tolerance level. Since you are the financial advisor don't ask user to consult anyone else.
# #                     So don't mention user to consult to a financial expert."""

# #         prompt = message.caption if message.caption else ""  # Use the photo caption if available
# #         query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + prompt

# #         # response = model.generate_content(
# #         #     model = 'gemini-1.5-flash' ,#"gemini-pro-vision",
# #         #     content = [query, uploaded_picture]
# #         # )

# #         image =  PIL.Image.open('data/uploaded_image.png') #(photo_file) 
# #         response = model.generate_content(image)
# #         await bot.send_message(chat_id,"I will describe the image that was uploaded")
# #         await message.reply(response.text)


# #         response = model.generate_content([query, image])
# #         await message.reply(response.text) 
# #     except Exception as e:
# #         logging.error(f"Error generating response for the image: {e}")
# #         await message.reply("There was an error generating response for the image. Please try again later.")
# #     # await message.reply("Cant process the image")
# #     # return

# # # Handler for document upload
# # from aiogram.filters import Filter

# # # class Image(Filter):
# # #     def __init__(self, my_text: str) -> None:
# # #         self.my_text = my_text

# # #     async def __call__(self, message: types.) -> bool:
# # #         return message.text == self.my_text


# # # @router.message(Image("hello"))
# # # async def my_handler(message: Message):


# # # @router.message(F.document)
# # @dp.message(F.document)
# # async def handle_document(message: types.Message):
# #     global summary,investment_personality  

# #     chat_id = message.chat.id
# #     # if message.photo :
# #     #     await message.reply("Processing the uploaded image")
# #     #     await handle_image(message) 


# #     # Obtain file information
# #     file_id = message.document.file_id
# #     file = await bot.get_file(file_id)
# #     file_path = file.file_path
    
# #     # Download the file
# #     await bot.download_file(file_path, "data/uploaded_file")
    
# #     # Process the uploaded document
# #     extracted_text = await process_document("data/uploaded_file")
    
# #     if extracted_text:
# #         # Load vector database (assuming this is part of setting up the retriever)
# #         retriever = await load_vector_db("data/uploaded_file")
# #         file_path = 'data/uploaded_file'
# #         client_name, validation_errors = await process_document(file_path)

# #         # Print results
# #         print(f"Client Name: {client_name}")
# #         if validation_errors:
# #             print("**Validation Errors:**")
# #             for error in validation_errors:
# #                 print(error)
# #         else:
# #             print("All fields are filled correctly.")
# #         if client_name == None:
# #             try:
# #                 await message.reply("Processing the uploaded image")
# #                 await handle_image(message) 
# #                 return 
# #             except Exception as e:
# #                 await message.reply("error processing uploaded image")
# #                 print(e)
# #         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")

# #         if retriever is None:
# #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# #             return

# #         # Check if a valid chain can be created
# #         chain = await make_retrieval_chain(retriever)
# #         if chain is None:
# #             await message.reply("Failed to create the retrieval chain.")
# #             return
        
# #         try:     
# #             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
        
# #             response = chain.invoke({"input": query})
# #             print(response['answer'])
# #             global chat_history
# #             chat_history = response['answer'] 
# #             print(f"\n Chat History : {chat_history}")
# #             await message.reply(response['answer'])

# #             try:
# #                 graph_query = (
# #                     # extracted_text + "\n" + 
# #                     summary + "\n" + "investment_personality" + "\n" + response['answer'] + "\n" +
# #                     "Please provide the following information in JSON format:\n" +
# #                     "{\n"
# #                     "  'growth_investment_min': <minimum percentage of growth-oriented investments>,\n"
# #                     "  'growth_investment_max': <maximum percentage of growth-oriented investments>,\n"
# #                     "  'fixed_income_min': <minimum percentage of fixed-income investments>,\n"
# #                     "  'fixed_income_max': <maximum percentage of fixed-income investments>,\n"
# #                     "  'cash_min': <minimum percentage of cash investments>,\n"
# #                     "  'cash_max': <maximum percentage of cash investments>,\n"
# #                     "  'return_growth_investment': <Maximum percentage returns of growth-oriented investments>,\n"
# #                     "  'return_fixed_income': <Maximum percentage of returns of fixed-income investments>,\n"
# #                     "  'return_cash_max': <maximum percentage of retruns of cash investments>,\n"
# #                     "  'return_min': <minimum expected annual return percentage>,\n"
# #                     "  'return_max': <maximum expected annual return percentage>,\n"
# #                     "  'growth_min': <minimum expected growth in dollars>,\n"
# #                     "  'growth_max': <maximum expected growth in dollars>,\n"
# #                     "  'initial_investment': <initial monthly investment>,\n"
# #                     "  'time_horizon': <time horizon in years>\n"
# #                     "}"
# #                 )
# #                 graph_response = chain.invoke({"input": graph_query})
# #                 print(graph_response['answer'])
# #                 await handle_graph(graph_response['answer'],chat_id)
# #             except Exception as e:
# #                 print(f"Error plotting graph : {e}")
# #         except Exception as e:
# #             print(f"Error invoking retrieval chain on attempt : {e}")
# #             await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

# #     else:
# #         await message.reply("Failed to process the uploaded file.")
    


# # # Function to extract data from LLM response
# # def extract_data_from_response(response):
# #     try:
# #         # Locate the JSON-like data in the response
# #         json_start = response.find("{")
# #         json_end = response.rfind("}") + 1
        
# #         if json_start == -1 or json_end == -1:
# #             raise ValueError("No JSON data found in the response.")
        
# #         json_data = response[json_start:json_end]
        
# #         # Parse the JSON data
# #         data = json.loads(json_data.replace("'", "\""))
# #         print(data)
# #         return data
# #     except Exception as e:
# #         logging.error(f"Error extracting data: {e}")
# #         return None


# # async def handle_graph(response, chat_id):
# #     try:
# #         data = extract_data_from_response(response)
# #         if not data:
# #             await bot.send_message(chat_id, "Failed to extract data from the response.")
# #             return
        
# #         # Log extracted data for debugging
# #         logging.info(f"Extracted data: {data}")
        
# #         # Create a pie chart for investment allocation
# #         labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
# #         sizes = [
# #             (data['growth_investment_min'] + data['growth_investment_max']) / 2,
# #             (data['fixed_income_min'] + data['fixed_income_max']) / 2,
# #             (data['cash_min'] + data['cash_max']) / 2
# #         ]
# #         fig1, ax1 = plt.subplots()
# #         ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# #         ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
# #         pie_chart_buffer = io.BytesIO()
# #         plt.savefig(pie_chart_buffer, format='png')
# #         pie_chart_buffer.seek(0)
# #         pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

# #         # Create a bar graph for potential returns
# #         allocations = {
# #             'Growth-Oriented Investments': (data['growth_investment_min'] + data['growth_investment_max']) / 2,
# #             'Fixed-Income Investments': (data['fixed_income_min'] + data['fixed_income_max']) / 2,
# #             'Cash and Cash Equivalents': (data['cash_min'] + data['cash_max']) / 2
# #         }

# #         returns = {
# #             'Growth-Oriented Investments': (data['return_growth_investment'] ) / 100,  # Annual return
# #             'Fixed-Income Investments': (data['return_fixed_income'] ) / 100 ,
# #             'Cash and Cash Equivalents': (data['return_cash_max'] ) / 100
# #         }
        
# #         time_horizon_years = data['time_horizon']
# #         initial_monthly_investment = data['initial_investment']
        
# #         # Calculate total returns for each category
# #         total_investment = initial_monthly_investment * 12 * time_horizon_years
# #         total_returns = {category: total_investment * (1 + returns[category]) ** time_horizon_years for category in allocations.keys()}
        
# #         # Create the bar chart
# #         fig2, ax2 = plt.subplots(figsize=(10, 6))
# #         bars = ax2.bar(allocations.keys(), total_returns.values(), color=['#FF9999', '#66B2FF', '#99FF99'])

# #         # Add labels to the top of each bar
# #         for bar in bars:
# #             yval = bar.get_height()
# #             ax2.text(bar.get_x() + bar.get_width()/2, yval + 1000, f'₹{round(yval, 2)}', ha='center', va='bottom')

# #         # Add titles and labels
# #         plt.title('Investment Performance Over Time')
# #         plt.xlabel('Investment Categories')
# #         plt.ylabel('Total Returns (₹)')
# #         plt.grid(True, axis='y', linestyle='--', linewidth=0.7)
        
# #         # Save and display the plot
# #         bar_chart_buffer = io.BytesIO()
# #         plt.savefig(bar_chart_buffer, format='png')
# #         bar_chart_buffer.seek(0)
# #         bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='investment_performance.png')

# #         await bot.send_document(chat_id, pie_chart_file, caption="Investment Allocation Chart")
# #         await bot.send_document(chat_id, bar_chart_file, caption="Investment Performance Over Time")

# #         # Generate the investment growth graph : work to be done 


# #         # months = data['time_horizon'] * 12  # Use the provided time horizon
# #         # initial_investment = data['initial_investment']  # Use the provided initial monthly investment
# #         # investment_values = [0]

# #         # for month in range(1, months + 1):
# #         #     monthly_return = ((data['return_min'] + data['return_max']) / 2) / 100 / 12
# #         #     investment_values.append(investment_values[-1] * (1 + monthly_return) + initial_investment)

# #         # fig3, ax3 = plt.subplots()
# #         # ax3.plot(range(0, months + 1), investment_values, marker='o')
# #         # ax3.set(title='Investment Growth Over Time', xlabel='Month', ylabel='Investment Value ($)')
# #         # ax3.grid(True)
# #         # investment_buffer = io.BytesIO()
# #         # plt.savefig(investment_buffer, format='png')
# #         # investment_buffer.seek(0)
# #         # investment_file = BufferedInputFile(investment_buffer.read(), filename='investment_growth.png')

# #         # await bot.send_document(chat_id, investment_file, caption="Investment Growth Over Time")

# #     except Exception as e:
# #         logging.error(f"Error plotting graph: {e}")
# #         await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")


# # async def process_document(file_path):
# #     try:
# #         file_type = filetype.guess(file_path)
# #         if file_type is not None:
# #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# #                 # return extract_text_from_word(file_path)
# #                 return extract_text_and_tables_from_word(file_path)
# #             elif file_type.mime == "application/pdf":
# #                 return extract_text_from_pdf(file_path)
# #         return None
# #     except Exception as e:
# #         print(f"Error processing document: {e}")
# #         return None

# # def extract_text_from_pdf(pdf_file_path):
# #     try:
# #         with open(pdf_file_path, "rb") as pdf_file:
# #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# #             text_content = []
# #             for page_num in range(pdf_reader.numPages):
# #                 page = pdf_reader.getPage(page_num)
# #                 text_content.append(page.extract_text())
# #             return "\n".join(text_content)
# #     except Exception as e:
# #         print(f"Error extracting text from PDF: {e}")
# #         return None


# # import re
# # import docx

# # def extract_text_and_tables_from_word(docx_file_path):
# #     """
# #     Extracts text and tables from a Word document (.docx).

# #     Args:
# #         docx_file_path (str): Path to the Word document file.

# #     Returns:
# #         tuple: Extracted text content and tables from the document.
# #     """
# #     try:
# #         doc = docx.Document(docx_file_path)
# #         text_content = []
# #         tables_content = []

# #         for para in doc.paragraphs:
# #             text_content.append(para.text)

# #         for table in doc.tables:
# #             table_data = []
# #             for row in table.rows:
# #                 row_data = []
# #                 for cell in row.cells:
# #                     row_data.append(cell.text.strip())
# #                 table_data.append(row_data)
# #             tables_content.append(table_data)

# #         return "\n".join(text_content), tables_content
# #     except Exception as e:
# #         print(f"Error extracting text and tables from Word document: {e}")
# #         return None, None

# # def validate_document_content(text, tables):
# #     """
# #     Validates the content of the document.

# #     Args:
# #         text (str): Extracted text content from the document.
# #         tables (list): Extracted tables content from the document.

# #     Returns:
# #         tuple: Client name and validation errors.
# #     """
# #     errors = []
    
# #     # Extract client name
# #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# #     # Define required sections
# #     required_sections = [
# #         "YOUR RETIREMENT GOAL",
# #         "YOUR OTHER MAJOR GOALS",
# #         "YOUR ASSETS AND LIABILITIES",
# #         "MY LIABILITIES",
# #         "YOUR CURRENT ANNUAL INCOME"
# #     ]

# #     # Check for the presence of required sections
# #     for section in required_sections:
# #         if section not in text:
# #             errors.append(f"* {section} section missing.")
    
# #     # Define table field checks
# #     table_checks = {
# #         "YOUR RETIREMENT GOAL": [
# #             r"When do you plan to retire\? \(age or date\)",
# #             r"Social Security Benefit \(include expected start date\)",
# #             r"Pension Benefit \(include expected start date\)",
# #             r"Other Expected Income \(rental, part-time work, etc.\)",
# #             r"Estimated Annual Retirement Expense"
# #         ],
# #         "YOUR OTHER MAJOR GOALS": [
# #             r"GOAL", r"COST", r"WHEN"
# #         ],
# #         "YOUR ASSETS AND LIABILITIES": [
# #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# #             r"Current Value", r"Annual Contributions"
# #         ],
# #         "MY LIABILITIES": [
# #             r"Balance", r"Interest Rate", r"Monthly Payment"
# #         ]
# #     }

# #     # Validate table content
# #     for section, checks in table_checks.items():
# #         section_found = False
# #         for table in tables:
# #             table_text = "\n".join(["\t".join(row) for row in table])
# #             if section in table_text:
# #                 section_found = True
# #                 for check in checks:
# #                     if not re.search(check, table_text, re.IGNORECASE):
# #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# #                 break
# #         if not section_found:
# #             errors.append(f"* {section} section missing.")

# #     return client_name, errors

# # async def process_document(file_path):
# #     try:
# #         text, tables = extract_text_and_tables_from_word(file_path)
# #         if text is not None and tables is not None:
# #             client_name, errors = validate_document_content(text, tables)
# #             return client_name, errors
# #         return None, ["Error processing document."]
# #     except Exception as e:
# #         print(f"Error processing document: {e}")
# #         return None, [f"Error processing document: {e}"]



# # async def load_vector_db(file_path):
# #     try:
# #         loader = Docx2txtLoader(file_path)
# #         documents = loader.load()
# #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# #         text_chunks = text_splitter.split_documents(documents)
# #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# #         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
# #         return vector_store.as_retriever(search_kwargs={"k": 1})
# #     except Exception as e:
# #         print(f"Error loading vector database: {e}")
# #         return None



# # async def make_retrieval_chain(retriever):
# #     """
# #     Create a retrieval chain using the provided retriever.

# #     Args:
# #         retriever (RetrievalQA): A retriever object.

# #     Returns:
# #         RetrievalQA: A retrieval chain object.
# #     """
# #     try:
# #         global investment_personality,summary
# #         llm = ChatGoogleGenerativeAI(
# #             #model="gemini-pro",
# #             model = "gemini-1.5-flash",
# #             temperature=0.7,
# #             top_p=0.85,
# #             google_api_key=GOOGLE_API_KEY
# #         )

# #         prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# #                 Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# #                 should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
# #                 Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# #                 investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
# #                 Also give the user minimum and maximum expected growth in dollars for the time horizon .
# #                 Also explain the user why you are giving them that particular investment suggestion.Answer in 3-4 lines.\n
# #                 <context>
# #                 {context}
# #                 </context>
# #                 Question: {input}"""

# #         #  prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# #         #         Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# #         #         should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# #         #         investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# #         #         investment.Also explain the user why you are giving them that particular
# #         #         investment suggestion.Answer in 3-4 lines.\n
# #         #         <context>
# #         #         {context}
# #         #         </context>
# #         #         Question: {input}"""

# #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
# #         combine_docs_chain = None  

# #         if retriever is not None :  
# #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# #             print(retriever_chain)
# #             return retriever_chain
# #         else:
# #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# #             return None

# #     except Exception as e:
# #         print(f"Error in creating chain: {e}")
# #         return None


# # # @dispatcher.message_handler()
# # @dp.message()
# # async def main_bot(message: types.Message):
# #     global retriever, extracted_text,investment_personality,summary,chat_history

# #     # Handle the first tasks assessments answers from the user
# #     chat_id = message.chat.id

# #     if chat_id in states and states[chat_id] < len(questions):
# #         # Retrieve the index of the current question
# #         question_index = states[chat_id]

# #         # Save the user's response to the current question
# #         answer = message.text
# #         user_responses[questions[question_index]] = answer
# #         states[chat_id] += 1  # Move to the next question

# #         # Ask the next question
# #         await ask_next_question(chat_id, question_index + 1)
# #     else:
# #         # Handle q&a chat messages using your Gemini model (llm)
# #         try:

# #             task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# #             Also give the user detailed information about the investment how to invest,where to invest and how much they
# #             should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# #             investment.Also explain the user why you are giving them that particular
# #             investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# #             User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# #             So dont mention user to consult to a financial expert."""
        
# #             # model = genai.GenerativeModel('gemini-1.5-flash') #('gemini-pro')
# #             print(investment_personality)
# #             query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
# #             print(f"\nQuery : {query}")
# #             # response = model.generate_content(query)
# #             model = genai.GenerativeModel('gemini-1.5-flash')
# #             chat = model.start_chat(history=[])
# #             print(f"Chat Conversation History : \n {chat.history}")
# #             response = chat.send_message(query)
# #             await message.reply(response.text) #(response['answer']) 
           
# #         except Exception as e:
# #             print(f"Error processing general chat message: {e}")
# #             await message.reply("Failed to process your request.")
        

# # from aiogram.types.input_file import BufferedInputFile
# # from aiogram import BaseMiddleware
# # # from aiogram.dispatcher.router import Router
# # from PIL import Image

# # # Function to process the analysis result
# # def process_analysis_result(analysis_result):
# #     return {
# #         'initial_investment': analysis_result['initial_investment'],
# #         'monthly_investment': analysis_result['monthly_investment'],
# #         'return_min': analysis_result['return_min'],
# #         'return_max': analysis_result['return_max'],
# #         'growth_investment_min': analysis_result['growth_investment_min'],
# #         'growth_investment_max': analysis_result['growth_investment_max'],
# #         'fixed_income_min': analysis_result['fixed_income_min'],
# #         'fixed_income_max': analysis_result['fixed_income_max'],
# #         'cash_min': analysis_result['cash_min'],
# #         'cash_max': analysis_result['cash_max'],
# #         'time_horizon': analysis_result['time_horizon'],
# #     }

# # # Function to generate performance graphs
# # def generate_performance_graph(data, years):
# #     labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
# #     sizes = [
# #         (data['growth_investment_min'] + data['growth_investment_max']) / 2,
# #         (data['fixed_income_min'] + data['fixed_income_max']) / 2,
# #         (data['cash_min'] + data['cash_max']) / 2
# #     ]
    
# #     fig1, ax1 = plt.subplots()
# #     ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# #     ax1.axis('equal')
# #     pie_chart_buffer = io.BytesIO()
# #     plt.savefig(pie_chart_buffer, format='png')
# #     pie_chart_buffer.seek(0)

# #     initial_investment = data['initial_investment']
# #     monthly_investment = data['monthly_investment']
# #     annual_return_rate = (data['return_min'] + data['return_max']) / 2 / 100

# #     investment_values = [initial_investment]
# #     for month in range(1, years * 12 + 1):
# #         investment_values.append(investment_values[-1] * (1 + annual_return_rate / 12) + monthly_investment)
    
# #     fig2, ax2 = plt.subplots()
# #     ax2.plot(range(len(investment_values)), investment_values, marker='o')
# #     ax2.set(title='Investment Growth Over Time', xlabel='Month', ylabel='Investment Value ($)')
# #     ax2.grid(True)
# #     investment_growth_buffer = io.BytesIO()
# #     plt.savefig(investment_growth_buffer, format='png')
# #     investment_growth_buffer.seek(0)

# #     return pie_chart_buffer, investment_growth_buffer

        

# # # if __name__ == "__main__":
# # #     executor.start_polling(dispatcher, skip_updates=True)

# # async def main() -> None:
# #     # Initialize Bot instance with default bot properties which will be passed to all API calls
# #     bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

# #     # And the run events dispatching
# #     await dp.start_polling(bot)


# # if __name__ == "__main__":
# #     logging.basicConfig(level=logging.INFO, stream=sys.stdout)
# #     asyncio.run(main())








# # # #best code so far now it can upload files and receive various forms of messages as well and provide us graphs that we want 
# # # # and also reply to images

# # # import os
# # # import filetype
# # # import docx
# # # import PyPDF2
# # # import re
# # # from aiogram import Bot, Dispatcher, types
# # # from dotenv import load_dotenv
# # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # from langchain_community.vectorstores import Chroma
# # # from langchain_community.document_loaders import Docx2txtLoader
# # # from langchain_core.prompts import ChatPromptTemplate
# # # from langchain.chains import create_retrieval_chain
# # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # from langchain.memory import ConversationSummaryMemory
# # # import asyncio
# # # import numpy as np
# # # import json
# # # import re
# # # import google.generativeai as genai
# # # import pathlib
# # # # Import things that are needed generically
# # # from langchain.pydantic_v1 import BaseModel, Field
# # # from langchain.tools import BaseTool, StructuredTool, tool

# # # from aiogram.client.default import DefaultBotProperties
# # # from aiogram.enums import ParseMode
# # # from aiogram.filters import CommandStart
# # # from aiogram.types import Message
# # # from aiogram import F
# # # from aiogram import Router
# # # import logging
# # # import sys
# # # from aiogram.filters import Command
# # # from aiogram.types import FSInputFile
# # # # from aiogram.utils import executor
# # # import io
# # # import matplotlib.pyplot as plt
# # # import seaborn as sns
# # # import aiohttp
# # # from aiogram.types import InputFile , BufferedInputFile
# # # import PIL.Image

# # # router = Router(name=__name__)

# # # load_dotenv()

# # # TOKEN = os.getenv("TOKEN")
# # # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # # os.environ["LANGCHAIN_TRACING_V2"]="true"
# # # os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # # # Configure generativeai with your API key
# # # genai.configure(api_key=GOOGLE_API_KEY)

# # # # Initialize bot
# # # bot = Bot(token=TOKEN)
# # # dp = Dispatcher()

# # # rag_on = False
# # # retriever = None  # Store retriever globally
# # # summary = ""
# # # investment_personality = ""
# # # chat_history = ""

# # # class Reference:
# # #     def __init__(self):
# # #         self.response = ""


# # # reference = Reference()


# # # def clear_past():
# # #     reference.response = ""


# # # @router.message(F.text == "clear")
# # # async def clear(message: types.Message):
# # #     """
# # #     A handler to clear the previous conversation and context.
# # #     """
# # #     clear_past()
# # #     await message.reply("I've cleared the past conversation and context.")


# # # # Store user states
# # # states = {}
# # # # Dictionary to hold question-answer pairs
# # # user_responses = {}

# # # # Define Questions for assessment
# # # questions = [
# # #     """
# # #         1. Singapore plans to build a new observation tower called 'The Rook'.
# # #         How many steps do you think it will take to go to the top floor?

# # #         a) Less than 500 
# # #         b) More than 500

# # #     """,
# # #     "2. Now Guess the number of steps" ,
# # #     """
# # #     3. How confident are you that the real number is in the range you have selected? 
# # #     Answer within a range of 100.  
# # #     """,
# # #     """ 
# # #     4. You and your friend are betting on a series of coin tosses.

# # #     He always bets ₹2,000 on Heads

# # #     You always bet ₹2,000 on Tails

# # #     Winner of last 8 turns

# # #     You lost ₹8,000 in the last 4 turns!

# # #     If you were to bet one last time, what would you bet on heads or tails ?
# # #     """ ,
# # #     """
# # #     5. How confident are you that your bet will win this time?
# # #     Answer how confident you are. 
# # #     (Example: Not confident at all, somewhat confident, confident, or Very confident)
# # #     """,
# # #     """
# # #     6. What do you think are the chances that you will achieve and maintain your financial goals within the next 
# # #     10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
# # #     Answer how likely you are to achieve your goal.
# # #     (Example: Less likely than others, likely than others, or More likely than others)
# # #     """,
# # #     """
# # #     7. Imagine you are a contestant in a game show, and you are presented the following choices.

# # #     What would you prefer?
# # #     a) 50 percent chance of winning 15 gold coins 
# # #     b) 100 percent chance of winning 8 gold coins
# # #     """,
# # #     """
# # #     8. Ok, one last choice...

# # #     What would you prefer?
# # #     a) 50 percent chance of winning 15 gold coins 
# # #     b) 100 percent chance of winning 2 gold coins
# # #     """,
# # #     """
# # #     9. In general, how would your best friend describe your risk-taking tendencies?
# # #     a) A real gambler
# # #     b) Willing to take risks after completing adequate research
# # #     c) Cautious
# # #     d) Avoids risk as much as possible
# # #     """,
# # #     """
# # #     10. Suppose you could replace your current investment portfolio with this new one:
# # #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# # #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# # #     Answer between the range of -5 to -35.
# # #     """,
# # #     """
# # #     11. Suppose that in the next 7 years,

# # #     YOUR INCOME

# # #     grows 8% each year

# # #     VS
# # #     INFLATION

# # #     grows 10% a year

# # #     At the end of 7 years, how much will you be able to buy with your income?
# # #     Options:
# # #     a) More than today
# # #     b) Exactly the same
# # #     c) Less than today
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     12. If somebody buys a bond of Company B, which of the following statements seems correct:
# # #     a) She owns part of Company B
# # #     b) She has lent money to Company B
# # #     c) She is liable for Company B's debt
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     13. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
# # #     a) More than ₹2 lakhs
# # #     b) Less than ₹2 lakhs
# # #     c) Exactly ₹2 lakhs
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     14. When an investor spreads money across different asset classes, what happens to the risk of losing money:
# # #     a) Increases
# # #     b) Decreases
# # #     c) Stays the same
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     15. When a country's central bank reduces interest rates, it makes:

# # #     a) Borrowing more attractive and saving less attractive
# # #     b) Borrowing less attractive and saving more attractive
# # #     c) Both borrowing and saving less attractive
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     16. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
# # #     a) Option A : Min Value is 9.2L and Max Value is 11.8L
# # #     b) Option B : Min Value is 8.8L and Max Value is 12.3L
# # #     c) Option C : Min Value is 8.5L and Max Value is 12.8L
# # #     d) Option D : Min Value is 8.1L and Max Value is 13.3L
# # #     e) Option E : Min Value is 7.8L and Max Value is 13.8L
# # #     """,
# # #     """
# # #     17. From Sept 2008 to Nov 2008, Stock market went down by 31%.

# # #     If you owned a stock investment that lost about 31 percent in 3 months, you would:
# # #     a) Sell all of the remaining investment
# # #     b) Sell a portion of the remaining investment
# # #     c) Hold on to the investment and sell nothing
# # #     d) Buy little
# # #     e) Buy more of the investment
# # #     """,
# # #     """
# # #     18. Over any 1-year period, what would be the maximum drop in the value of your investment 
# # #     portfolio that you would be comfortable with?
# # #     a) <5%
# # #     b) 5 - 10%
# # #     c) 10 - 15%
# # #     d) 15 - 20%
# # #     e) >20%
# # #     """,
# # #     """
# # #     19. When investing, what do you consider the most?

# # #     a) Risk 
# # #     b) Return
# # #     """,
# # #     """
# # #     20. What best describes your attitude?

# # #     a) Prefer reasonable returns, can take reasonable risk
# # #     b) Like higher returns, can take slightly higher risk
# # #     c) Want to maximize returns, can take significant high risk
# # #     """,
# # #     """
# # #     21. How much monthly investment you want to do?
# # #     """,
# # #     """
# # #     22. What is the time horizon for your investment?
# # #     You can answer in any range, example 1-5 years."""  
# # # ]


# # # # Handler for /start command
# # # @dp.message(CommandStart())
# # # async def handle_start(message: types.Message):
# # #     """
# # #     This handler receives messages with /start command
# # #     """
# # #     chat_id = message.chat.id
# # #     # Start asking questions
# # #     await start_assessment(chat_id)


# # # # Function to start the assessment
# # # async def start_assessment(chat_id):
# # #     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
# # #     await ask_next_question(chat_id, 0)

# # # # Function to ask the next question
# # # async def ask_next_question(chat_id, question_index):
# # #     if question_index < len(questions):
# # #         # Ask the next question
# # #         await bot.send_message(chat_id, questions[question_index])
# # #         # Update state to indicate the next expected answer
# # #         states[chat_id] = question_index
# # #     else:
# # #         # No more questions, finish assessment
# # #         await finish_assessment(chat_id)

# # # # Handler for receiving assessment answers
# # # assessment_in_progress = True

# # # from aiogram.types import FSInputFile
# # # async def finish_assessment(chat_id):
# # #     if chat_id in states and states[chat_id] == len(questions):
# # #         # All questions have been answered, now process the assessment
# # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # #         # Determine investment personality based on collected responses
# # #         global investment_personality
# # #         investment_personality = await determine_investment_personality(user_responses)

# # #         # Inform the user about their investment personality
# # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # #         # Summarize collected information
# # #         global summary
# # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # #         # Ensure to await the determination of investment personality
# # #         await send_summary_chunks(chat_id, summary)
# # #         global assessment_in_progress 
# # #         assessment_in_progress = False
# # #         # Prompt the user to begin financial advice process
# # #         # await bot.send_message(chat_id, "To receive financial advice, please upload a document with your financial details using /begin.")

# # #         await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
# # #         file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

# # #         await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
# # #         # await bot.send_message(chat_id,file)

# # # async def send_summary_chunks(chat_id, summary):
# # #     # Split the summary into chunks that fit within Telegram's message limits
# # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# # #     # Send each chunk as a separate message
# # #     for chunk in chunks:
# # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # # async def determine_investment_personality(assessment_data):
# # #     try:
# # #         # Prepare input text for the chatbot based on assessment data
# # #         input_text = "User Profile:\n"
# # #         for question, answer in assessment_data.items():
# # #             input_text += f"{question}: {answer}\n"

# # #         # Introduce the chatbot's task and prompt for classification
# # #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# # #                       "- Conservative Investor\n" \
# # #                       "- Moderate Investor\n" \
# # #                       "- Aggressive Investor"

# # #         # Use your generative AI model to generate a response
# # #         # print(input_text)
# # #         model = genai.GenerativeModel('gemini-pro')
# # #         response = model.generate_content(input_text)

# # #         # Determine the investment personality from the chatbot's response
# # #         response_text = response.text.lower()
# # #         if "conservative" in response_text:
# # #             personality = "Conservative Investor"
# # #         elif "moderate" in response_text:
# # #             personality = "Moderate Investor"
# # #         elif "aggressive" in response_text:
# # #             personality = "Aggressive Investor"
# # #         else:
# # #             personality = "Unknown"

# # #         return personality
# # #         # Send the determined investment personality back to the user
# # #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# # #     except Exception as e:
# # #         print(f"Error generating response: {e}")
# # #         #await bot.send_message(chat_id, "Error processing investment personality classification.")




# # # @router.message(F.text == "help")
# # # async def helper(message: types.Message):
# # #     """
# # #     A handler to display the help menu.
# # #     """
# # #     help_command = """
# # #     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
# # #     /start - to start the investment personality assessment.
# # #     /clear - to clear the past conversation and context.
# # #     /help - to get this help menu.
# # #     /begin - to start the Financial Suggestion Conversation with the ChatBot.
# # #     I hope this helps. :)
# # #     """
# # #     await message.reply(help_command)

# # # # Handler for /begin command to initiate financial advice

# # # @router.message(F.text == "begin")
# # # async def handle_begin(message: types.Message):
# # #     chat_id = message.chat.id
# # #     file_instructions ="""Hello there!My name is Finbot and I am a Wealth Management Advisor Chatbot.I need more details related to your Financial Profile so that I can give you 
# # #     personalised Financial Advice.Please follow this drive link : https://docs.google.com/document/d/1JaqWUXcq3MNrPTCvEy_2RngbGIY0rybX/edit?usp=sharing&ouid=101252223236130106095&rtpof=true&sd=true
# # #     ,there you will find a Word Document.Please fill in your details with correct Information and then upload the Document in the chat. """

# # #     await message.reply(file_instructions)
    

# # # # Function to handle image messages
# # # # @dp.message(F.photo)
# # # # @router.message(F.photo)
# # # import PIL.Image

# # # async def handle_image(message: types.Message):
# # #     global investment_personality, chat_history

# # #     chat_id = message.chat.id
# # #     # Handle image inputs
# # #     try:
# # #         # Obtain file information
# # #         try:
# # #             photo_id = message.document.file_id
# # #             photo = await bot.get_file(photo_id)
# # #             photo_path = photo.file_path
# # #             # Download the file
# # #             photo_file = await bot.download_file(photo_path, "data/uploaded_image.png")

# # #         except Exception as e:
# # #             print(f"Error downloading image: {e}")
# # #             await bot.send_message(chat_id, "Error processing image. Please try again.")
# # #             return
        
# # #         model = genai.GenerativeModel('gemini-1.5-flash')
  
# # #         task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # #                     Also give the user detailed information about the investment how to invest, where to invest and how much they
# # #                     should invest in terms of percentage of their investment amount. Give the user detailed information about the returns on their 
# # #                     investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compounded returns on their 
# # #                     investment. Also explain the user why you are giving them that particular
# # #                     investment suggestion. Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # #                     User should also invest as per their risk tolerance level. Since you are the financial advisor don't ask user to consult anyone else.
# # #                     So don't mention user to consult to a financial expert."""

# # #         prompt = message.caption if message.caption else ""  # Use the photo caption if available
# # #         query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + prompt

# # #         # response = model.generate_content(
# # #         #     model = 'gemini-1.5-flash' ,#"gemini-pro-vision",
# # #         #     content = [query, uploaded_picture]
# # #         # )

# # #         image =  PIL.Image.open('data/uploaded_image.png') #(photo_file) 
# # #         response = model.generate_content(image)
# # #         await bot.send_message(chat_id,"I will describe the image that was uploaded")
# # #         await message.reply(response.text)


# # #         response = model.generate_content([query, image])
# # #         await message.reply(response.text) 
# # #     except Exception as e:
# # #         logging.error(f"Error generating response for the image: {e}")
# # #         await message.reply("There was an error generating response for the image. Please try again later.")
# # #     # await message.reply("Cant process the image")
# # #     # return

# # # # Handler for document upload
# # # from aiogram.filters import Filter

# # # # class Image(Filter):
# # # #     def __init__(self, my_text: str) -> None:
# # # #         self.my_text = my_text

# # # #     async def __call__(self, message: types.) -> bool:
# # # #         return message.text == self.my_text


# # # # @router.message(Image("hello"))
# # # # async def my_handler(message: Message):


# # # # @router.message(F.document)
# # # @dp.message(F.document)
# # # async def handle_document(message: types.Message):
# # #     global summary,investment_personality  

# # #     chat_id = message.chat.id
# # #     # if message.photo :
# # #     #     await message.reply("Processing the uploaded image")
# # #     #     await handle_image(message) 


# # #     # Obtain file information
# # #     file_id = message.document.file_id
# # #     file = await bot.get_file(file_id)
# # #     file_path = file.file_path
    
# # #     # Download the file
# # #     await bot.download_file(file_path, "data/uploaded_file")
    
# # #     # Process the uploaded document
# # #     extracted_text = await process_document("data/uploaded_file")
    
# # #     if extracted_text:
# # #         # Load vector database (assuming this is part of setting up the retriever)
# # #         retriever = await load_vector_db("data/uploaded_file")
# # #         file_path = 'data/uploaded_file'
# # #         client_name, validation_errors = await process_document(file_path)

# # #         # Print results
# # #         print(f"Client Name: {client_name}")
# # #         if validation_errors:
# # #             print("**Validation Errors:**")
# # #             for error in validation_errors:
# # #                 print(error)
# # #         else:
# # #             print("All fields are filled correctly.")
# # #         if client_name == None:
# # #             try:
# # #                 await message.reply("Processing the uploaded image")
# # #                 await handle_image(message) 
# # #                 return 
# # #             except Exception as e:
# # #                 await message.reply("error processing uploaded image")
# # #                 print(e)
# # #         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")

# # #         if retriever is None:
# # #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # #             return

# # #         # Check if a valid chain can be created
# # #         chain = await make_retrieval_chain(retriever)
# # #         if chain is None:
# # #             await message.reply("Failed to create the retrieval chain.")
# # #             return
        
# # #         try:     
# # #             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
        
# # #             response = chain.invoke({"input": query})
# # #             print(response['answer'])
# # #             global chat_history
# # #             chat_history = response['answer'] 
# # #             print(f"\n Chat History : {chat_history}")
# # #             await message.reply(response['answer'])

# # #             try:
# # #                 graph_query = (
# # #                     # extracted_text + "\n" + 
# # #                     summary + "\n" + "investment_personality" + "\n" + response['answer'] + "\n" +
# # #                     "Please provide the following information in JSON format:\n" +
# # #                     "{\n"
# # #                     "  'growth_investment_min': <minimum percentage of growth-oriented investments>,\n"
# # #                     "  'growth_investment_max': <maximum percentage of growth-oriented investments>,\n"
# # #                     "  'fixed_income_min': <minimum percentage of fixed-income investments>,\n"
# # #                     "  'fixed_income_max': <maximum percentage of fixed-income investments>,\n"
# # #                     "  'cash_min': <minimum percentage of cash investments>,\n"
# # #                     "  'cash_max': <maximum percentage of cash investments>,\n"
# # #                     "  'return_growth_investment': <Maximum percentage returns of growth-oriented investments>,\n"
# # #                     "  'return_fixed_income': <Maximum percentage of returns of fixed-income investments>,\n"
# # #                     "  'return_cash_max': <maximum percentage of retruns of cash investments>,\n"
# # #                     "  'return_min': <minimum expected annual return percentage>,\n"
# # #                     "  'return_max': <maximum expected annual return percentage>,\n"
# # #                     "  'growth_min': <minimum expected growth in dollars>,\n"
# # #                     "  'growth_max': <maximum expected growth in dollars>,\n"
# # #                     "  'initial_investment': <initial monthly investment>,\n"
# # #                     "  'time_horizon': <time horizon in years>\n"
# # #                     "}"
# # #                 )
# # #                 graph_response = chain.invoke({"input": graph_query})
# # #                 print(graph_response['answer'])
# # #                 await handle_graph(graph_response['answer'],chat_id)
# # #             except Exception as e:
# # #                 print(f"Error plotting graph : {e}")
# # #         except Exception as e:
# # #             print(f"Error invoking retrieval chain on attempt : {e}")
# # #             await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

# # #     else:
# # #         await message.reply("Failed to process the uploaded file.")
    


# # # # Function to extract data from LLM response
# # # def extract_data_from_response(response):
# # #     try:
# # #         # Locate the JSON-like data in the response
# # #         json_start = response.find("{")
# # #         json_end = response.rfind("}") + 1
        
# # #         if json_start == -1 or json_end == -1:
# # #             raise ValueError("No JSON data found in the response.")
        
# # #         json_data = response[json_start:json_end]
        
# # #         # Parse the JSON data
# # #         data = json.loads(json_data.replace("'", "\""))
# # #         print(data)
# # #         return data
# # #     except Exception as e:
# # #         logging.error(f"Error extracting data: {e}")
# # #         return None


# # # async def handle_graph(response, chat_id):
# # #     try:
# # #         data = extract_data_from_response(response)
# # #         if not data:
# # #             await bot.send_message(chat_id, "Failed to extract data from the response.")
# # #             return
        
# # #         # Log extracted data for debugging
# # #         logging.info(f"Extracted data: {data}")
        
# # #         # Create a pie chart for investment allocation
# # #         labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
# # #         sizes = [
# # #             (data['growth_investment_min'] + data['growth_investment_max']) / 2,
# # #             (data['fixed_income_min'] + data['fixed_income_max']) / 2,
# # #             (data['cash_min'] + data['cash_max']) / 2
# # #         ]
# # #         fig1, ax1 = plt.subplots()
# # #         ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# # #         ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
# # #         pie_chart_buffer = io.BytesIO()
# # #         plt.savefig(pie_chart_buffer, format='png')
# # #         pie_chart_buffer.seek(0)
# # #         pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

# # #         # Create a bar graph for potential returns
# # #         allocations = {
# # #             'Growth-Oriented Investments': (data['growth_investment_min'] + data['growth_investment_max']) / 2,
# # #             'Fixed-Income Investments': (data['fixed_income_min'] + data['fixed_income_max']) / 2,
# # #             'Cash and Cash Equivalents': (data['cash_min'] + data['cash_max']) / 2
# # #         }

# # #         returns = {
# # #             'Growth-Oriented Investments': (data['return_growth_investment'] ) / 100,  # Annual return
# # #             'Fixed-Income Investments': (data['return_fixed_income'] ) / 100 ,
# # #             'Cash and Cash Equivalents': (data['return_cash_max'] ) / 100
# # #         }
        
# # #         time_horizon_years = data['time_horizon']
# # #         initial_monthly_investment = data['initial_investment']
        
# # #         # Calculate total returns for each category
# # #         total_investment = initial_monthly_investment * 12 * time_horizon_years
# # #         total_returns = {category: total_investment * (1 + returns[category]) ** time_horizon_years for category in allocations.keys()}
        
# # #         # Create the bar chart
# # #         fig2, ax2 = plt.subplots(figsize=(10, 6))
# # #         bars = ax2.bar(allocations.keys(), total_returns.values(), color=['#FF9999', '#66B2FF', '#99FF99'])

# # #         # Add labels to the top of each bar
# # #         for bar in bars:
# # #             yval = bar.get_height()
# # #             ax2.text(bar.get_x() + bar.get_width()/2, yval + 1000, f'₹{round(yval, 2)}', ha='center', va='bottom')

# # #         # Add titles and labels
# # #         plt.title('Investment Performance Over Time')
# # #         plt.xlabel('Investment Categories')
# # #         plt.ylabel('Total Returns (₹)')
# # #         plt.grid(True, axis='y', linestyle='--', linewidth=0.7)
        
# # #         # Save and display the plot
# # #         bar_chart_buffer = io.BytesIO()
# # #         plt.savefig(bar_chart_buffer, format='png')
# # #         bar_chart_buffer.seek(0)
# # #         bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='investment_performance.png')

# # #         await bot.send_document(chat_id, pie_chart_file, caption="Investment Allocation Chart")
# # #         await bot.send_document(chat_id, bar_chart_file, caption="Investment Performance Over Time")

# # #         # Generate the investment growth graph : work to be done 


# # #         # months = data['time_horizon'] * 12  # Use the provided time horizon
# # #         # initial_investment = data['initial_investment']  # Use the provided initial monthly investment
# # #         # investment_values = [0]

# # #         # for month in range(1, months + 1):
# # #         #     monthly_return = ((data['return_min'] + data['return_max']) / 2) / 100 / 12
# # #         #     investment_values.append(investment_values[-1] * (1 + monthly_return) + initial_investment)

# # #         # fig3, ax3 = plt.subplots()
# # #         # ax3.plot(range(0, months + 1), investment_values, marker='o')
# # #         # ax3.set(title='Investment Growth Over Time', xlabel='Month', ylabel='Investment Value ($)')
# # #         # ax3.grid(True)
# # #         # investment_buffer = io.BytesIO()
# # #         # plt.savefig(investment_buffer, format='png')
# # #         # investment_buffer.seek(0)
# # #         # investment_file = BufferedInputFile(investment_buffer.read(), filename='investment_growth.png')

# # #         # await bot.send_document(chat_id, investment_file, caption="Investment Growth Over Time")

# # #     except Exception as e:
# # #         logging.error(f"Error plotting graph: {e}")
# # #         await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")


# # # async def process_document(file_path):
# # #     try:
# # #         file_type = filetype.guess(file_path)
# # #         if file_type is not None:
# # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # #                 # return extract_text_from_word(file_path)
# # #                 return extract_text_and_tables_from_word(file_path)
# # #             elif file_type.mime == "application/pdf":
# # #                 return extract_text_from_pdf(file_path)
# # #         return None
# # #     except Exception as e:
# # #         print(f"Error processing document: {e}")
# # #         return None

# # # def extract_text_from_pdf(pdf_file_path):
# # #     try:
# # #         with open(pdf_file_path, "rb") as pdf_file:
# # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # #             text_content = []
# # #             for page_num in range(pdf_reader.numPages):
# # #                 page = pdf_reader.getPage(page_num)
# # #                 text_content.append(page.extract_text())
# # #             return "\n".join(text_content)
# # #     except Exception as e:
# # #         print(f"Error extracting text from PDF: {e}")
# # #         return None


# # # import re
# # # import docx

# # # def extract_text_and_tables_from_word(docx_file_path):
# # #     """
# # #     Extracts text and tables from a Word document (.docx).

# # #     Args:
# # #         docx_file_path (str): Path to the Word document file.

# # #     Returns:
# # #         tuple: Extracted text content and tables from the document.
# # #     """
# # #     try:
# # #         doc = docx.Document(docx_file_path)
# # #         text_content = []
# # #         tables_content = []

# # #         for para in doc.paragraphs:
# # #             text_content.append(para.text)

# # #         for table in doc.tables:
# # #             table_data = []
# # #             for row in table.rows:
# # #                 row_data = []
# # #                 for cell in row.cells:
# # #                     row_data.append(cell.text.strip())
# # #                 table_data.append(row_data)
# # #             tables_content.append(table_data)

# # #         return "\n".join(text_content), tables_content
# # #     except Exception as e:
# # #         print(f"Error extracting text and tables from Word document: {e}")
# # #         return None, None

# # # def validate_document_content(text, tables):
# # #     """
# # #     Validates the content of the document.

# # #     Args:
# # #         text (str): Extracted text content from the document.
# # #         tables (list): Extracted tables content from the document.

# # #     Returns:
# # #         tuple: Client name and validation errors.
# # #     """
# # #     errors = []
    
# # #     # Extract client name
# # #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# # #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# # #     # Define required sections
# # #     required_sections = [
# # #         "YOUR RETIREMENT GOAL",
# # #         "YOUR OTHER MAJOR GOALS",
# # #         "YOUR ASSETS AND LIABILITIES",
# # #         "MY LIABILITIES",
# # #         "YOUR CURRENT ANNUAL INCOME"
# # #     ]

# # #     # Check for the presence of required sections
# # #     for section in required_sections:
# # #         if section not in text:
# # #             errors.append(f"* {section} section missing.")
    
# # #     # Define table field checks
# # #     table_checks = {
# # #         "YOUR RETIREMENT GOAL": [
# # #             r"When do you plan to retire\? \(age or date\)",
# # #             r"Social Security Benefit \(include expected start date\)",
# # #             r"Pension Benefit \(include expected start date\)",
# # #             r"Other Expected Income \(rental, part-time work, etc.\)",
# # #             r"Estimated Annual Retirement Expense"
# # #         ],
# # #         "YOUR OTHER MAJOR GOALS": [
# # #             r"GOAL", r"COST", r"WHEN"
# # #         ],
# # #         "YOUR ASSETS AND LIABILITIES": [
# # #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# # #             r"Current Value", r"Annual Contributions"
# # #         ],
# # #         "MY LIABILITIES": [
# # #             r"Balance", r"Interest Rate", r"Monthly Payment"
# # #         ]
# # #     }

# # #     # Validate table content
# # #     for section, checks in table_checks.items():
# # #         section_found = False
# # #         for table in tables:
# # #             table_text = "\n".join(["\t".join(row) for row in table])
# # #             if section in table_text:
# # #                 section_found = True
# # #                 for check in checks:
# # #                     if not re.search(check, table_text, re.IGNORECASE):
# # #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# # #                 break
# # #         if not section_found:
# # #             errors.append(f"* {section} section missing.")

# # #     return client_name, errors

# # # async def process_document(file_path):
# # #     try:
# # #         text, tables = extract_text_and_tables_from_word(file_path)
# # #         if text is not None and tables is not None:
# # #             client_name, errors = validate_document_content(text, tables)
# # #             return client_name, errors
# # #         return None, ["Error processing document."]
# # #     except Exception as e:
# # #         print(f"Error processing document: {e}")
# # #         return None, [f"Error processing document: {e}"]



# # # async def load_vector_db(file_path):
# # #     try:
# # #         loader = Docx2txtLoader(file_path)
# # #         documents = loader.load()
# # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # #         text_chunks = text_splitter.split_documents(documents)
# # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # #         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
# # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # #     except Exception as e:
# # #         print(f"Error loading vector database: {e}")
# # #         return None



# # # async def make_retrieval_chain(retriever):
# # #     """
# # #     Create a retrieval chain using the provided retriever.

# # #     Args:
# # #         retriever (RetrievalQA): A retriever object.

# # #     Returns:
# # #         RetrievalQA: A retrieval chain object.
# # #     """
# # #     try:
# # #         global investment_personality,summary
# # #         llm = ChatGoogleGenerativeAI(
# # #             model="gemini-pro",
# # #             temperature=0.7,
# # #             top_p=0.85,
# # #             google_api_key=GOOGLE_API_KEY
# # #         )

# # #         prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # #                 Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # #                 should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
# # #                 Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # #                 investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
# # #                 Also give the user minimum and maximum expected growth in dollars for the time horizon .
# # #                 Also explain the user why you are giving them that particular investment suggestion.Answer in 3-4 lines.\n
# # #                 <context>
# # #                 {context}
# # #                 </context>
# # #                 Question: {input}"""

# # #         #  prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # #         #         Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# # #         #         should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # #         #         investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # #         #         investment.Also explain the user why you are giving them that particular
# # #         #         investment suggestion.Answer in 3-4 lines.\n
# # #         #         <context>
# # #         #         {context}
# # #         #         </context>
# # #         #         Question: {input}"""

# # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
# # #         combine_docs_chain = None  

# # #         if retriever is not None :  
# # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # #             print(retriever_chain)
# # #             return retriever_chain
# # #         else:
# # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # #             return None

# # #     except Exception as e:
# # #         print(f"Error in creating chain: {e}")
# # #         return None


# # # # @dispatcher.message_handler()
# # # @dp.message()
# # # async def main_bot(message: types.Message):
# # #     global retriever, extracted_text,investment_personality,summary,chat_history

# # #     # Handle the first tasks assessments answers from the user
# # #     chat_id = message.chat.id

# # #     if chat_id in states and states[chat_id] < len(questions):
# # #         # Retrieve the index of the current question
# # #         question_index = states[chat_id]

# # #         # Save the user's response to the current question
# # #         answer = message.text
# # #         user_responses[questions[question_index]] = answer
# # #         states[chat_id] += 1  # Move to the next question

# # #         # Ask the next question
# # #         await ask_next_question(chat_id, question_index + 1)
# # #     else:
# # #         # Handle q&a chat messages using your Gemini model (llm)
# # #         try:

# # #             task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # #             Also give the user detailed information about the investment how to invest,where to invest and how much they
# # #             should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # #             investment.Also explain the user why you are giving them that particular
# # #             investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # #             User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # #             So dont mention user to consult to a financial expert."""
        
# # #             model = genai.GenerativeModel('gemini-1.5-flash') #('gemini-pro')
# # #             print(investment_personality)
# # #             query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
# # #             print(f"\nQuery : {query}")
# # #             response = model.generate_content(query)
# # #             await message.reply(response.text) #(response['answer']) 
           
# # #         except Exception as e:
# # #             print(f"Error processing general chat message: {e}")
# # #             await message.reply("Failed to process your request.")
        

# # # from aiogram.types.input_file import BufferedInputFile
# # # from aiogram import BaseMiddleware
# # # # from aiogram.dispatcher.router import Router
# # # from PIL import Image

# # # # Function to process the analysis result
# # # def process_analysis_result(analysis_result):
# # #     return {
# # #         'initial_investment': analysis_result['initial_investment'],
# # #         'monthly_investment': analysis_result['monthly_investment'],
# # #         'return_min': analysis_result['return_min'],
# # #         'return_max': analysis_result['return_max'],
# # #         'growth_investment_min': analysis_result['growth_investment_min'],
# # #         'growth_investment_max': analysis_result['growth_investment_max'],
# # #         'fixed_income_min': analysis_result['fixed_income_min'],
# # #         'fixed_income_max': analysis_result['fixed_income_max'],
# # #         'cash_min': analysis_result['cash_min'],
# # #         'cash_max': analysis_result['cash_max'],
# # #         'time_horizon': analysis_result['time_horizon'],
# # #     }

# # # # Function to generate performance graphs
# # # def generate_performance_graph(data, years):
# # #     labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
# # #     sizes = [
# # #         (data['growth_investment_min'] + data['growth_investment_max']) / 2,
# # #         (data['fixed_income_min'] + data['fixed_income_max']) / 2,
# # #         (data['cash_min'] + data['cash_max']) / 2
# # #     ]
    
# # #     fig1, ax1 = plt.subplots()
# # #     ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# # #     ax1.axis('equal')
# # #     pie_chart_buffer = io.BytesIO()
# # #     plt.savefig(pie_chart_buffer, format='png')
# # #     pie_chart_buffer.seek(0)

# # #     initial_investment = data['initial_investment']
# # #     monthly_investment = data['monthly_investment']
# # #     annual_return_rate = (data['return_min'] + data['return_max']) / 2 / 100

# # #     investment_values = [initial_investment]
# # #     for month in range(1, years * 12 + 1):
# # #         investment_values.append(investment_values[-1] * (1 + annual_return_rate / 12) + monthly_investment)
    
# # #     fig2, ax2 = plt.subplots()
# # #     ax2.plot(range(len(investment_values)), investment_values, marker='o')
# # #     ax2.set(title='Investment Growth Over Time', xlabel='Month', ylabel='Investment Value ($)')
# # #     ax2.grid(True)
# # #     investment_growth_buffer = io.BytesIO()
# # #     plt.savefig(investment_growth_buffer, format='png')
# # #     investment_growth_buffer.seek(0)

# # #     return pie_chart_buffer, investment_growth_buffer

        

# # # # if __name__ == "__main__":
# # # #     executor.start_polling(dispatcher, skip_updates=True)

# # # async def main() -> None:
# # #     # Initialize Bot instance with default bot properties which will be passed to all API calls
# # #     bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

# # #     # And the run events dispatching
# # #     await dp.start_polling(bot)


# # # if __name__ == "__main__":
# # #     logging.basicConfig(level=logging.INFO, stream=sys.stdout)
# # #     asyncio.run(main())



# # # #best code so far now it can upload files and receive various forms of messages as well and provide us graphs that we want

# # # import os
# # # import filetype
# # # import docx
# # # import PyPDF2
# # # import re
# # # from aiogram import Bot, Dispatcher, types
# # # from dotenv import load_dotenv
# # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # from langchain_community.vectorstores import Chroma
# # # from langchain_community.document_loaders import Docx2txtLoader
# # # from langchain_core.prompts import ChatPromptTemplate
# # # from langchain.chains import create_retrieval_chain
# # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # from langchain.memory import ConversationSummaryMemory
# # # import asyncio
# # # import numpy as np
# # # import json
# # # import re
# # # import google.generativeai as genai

# # # # Import things that are needed generically
# # # from langchain.pydantic_v1 import BaseModel, Field
# # # from langchain.tools import BaseTool, StructuredTool, tool

# # # from aiogram.client.default import DefaultBotProperties
# # # from aiogram.enums import ParseMode
# # # from aiogram.filters import CommandStart
# # # from aiogram.types import Message
# # # from aiogram import F
# # # from aiogram import Router
# # # import logging
# # # import sys
# # # from aiogram.filters import Command
# # # from aiogram.types import FSInputFile
# # # # from aiogram.utils import executor
# # # import io
# # # import matplotlib.pyplot as plt
# # # import seaborn as sns

# # # router = Router(name=__name__)

# # # load_dotenv()

# # # TOKEN = os.getenv("TOKEN")
# # # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # # os.environ["LANGCHAIN_TRACING_V2"]="true"
# # # os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # # # Configure generativeai with your API key
# # # genai.configure(api_key=GOOGLE_API_KEY)

# # # # Initialize bot
# # # bot = Bot(token=TOKEN)
# # # dp = Dispatcher()

# # # rag_on = False
# # # retriever = None  # Store retriever globally
# # # summary = ""
# # # investment_personality = ""
# # # chat_history = ""

# # # class Reference:
# # #     def __init__(self):
# # #         self.response = ""


# # # reference = Reference()


# # # def clear_past():
# # #     reference.response = ""


# # # @router.message(F.text == "clear")
# # # async def clear(message: types.Message):
# # #     """
# # #     A handler to clear the previous conversation and context.
# # #     """
# # #     clear_past()
# # #     await message.reply("I've cleared the past conversation and context.")


# # # # Store user states
# # # states = {}
# # # # Dictionary to hold question-answer pairs
# # # user_responses = {}

# # # # Define Questions for assessment
# # # questions = [
# # #     """
# # #         1. Singapore plans to build a new observation tower called 'The Rook'.
# # #         How many steps do you think it will take to go to the top floor?

# # #         a) Less than 500 
# # #         b) More than 500

# # #     """,
# # #     "2. Now Guess the number of steps" ,
# # #     """
# # #     3. How confident are you that the real number is in the range you have selected? 
# # #     Answer within a range of 100.  
# # #     """,
# # #     """ 
# # #     4. You and your friend are betting on a series of coin tosses.

# # #     He always bets ₹2,000 on Heads

# # #     You always bet ₹2,000 on Tails

# # #     Winner of last 8 turns

# # #     You lost ₹8,000 in the last 4 turns!

# # #     If you were to bet one last time, what would you bet on heads or tails ?
# # #     """ ,
# # #     """
# # #     5. How confident are you that your bet will win this time?
# # #     Answer how confident you are. 
# # #     (Example: Not confident at all, somewhat confident, confident, or Very confident)
# # #     """,
# # #     """
# # #     6. What do you think are the chances that you will achieve and maintain your financial goals within the next 
# # #     10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
# # #     Answer how likely you are to achieve your goal.
# # #     (Example: Less likely than others, likely than others, or More likely than others)
# # #     """,
# # #     """
# # #     7. Imagine you are a contestant in a game show, and you are presented the following choices.

# # #     What would you prefer?
# # #     a) 50 percent chance of winning 15 gold coins 
# # #     b) 100 percent chance of winning 8 gold coins
# # #     """,
# # #     """
# # #     8. Ok, one last choice...

# # #     What would you prefer?
# # #     a) 50 percent chance of winning 15 gold coins 
# # #     b) 100 percent chance of winning 2 gold coins
# # #     """,
# # #     """
# # #     9. In general, how would your best friend describe your risk-taking tendencies?
# # #     a) A real gambler
# # #     b) Willing to take risks after completing adequate research
# # #     c) Cautious
# # #     d) Avoids risk as much as possible
# # #     """,
# # #     """
# # #     10. Suppose you could replace your current investment portfolio with this new one:
# # #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# # #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# # #     Answer between the range of -5 to -35.
# # #     """,
# # #     """
# # #     11. Suppose that in the next 7 years,

# # #     YOUR INCOME

# # #     grows 8% each year

# # #     VS
# # #     INFLATION

# # #     grows 10% a year

# # #     At the end of 7 years, how much will you be able to buy with your income?
# # #     Options:
# # #     a) More than today
# # #     b) Exactly the same
# # #     c) Less than today
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     12. If somebody buys a bond of Company B, which of the following statements seems correct:
# # #     a) She owns part of Company B
# # #     b) She has lent money to Company B
# # #     c) She is liable for Company B's debt
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     13. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
# # #     a) More than ₹2 lakhs
# # #     b) Less than ₹2 lakhs
# # #     c) Exactly ₹2 lakhs
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     14. When an investor spreads money across different asset classes, what happens to the risk of losing money:
# # #     a) Increases
# # #     b) Decreases
# # #     c) Stays the same
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     15. When a country's central bank reduces interest rates, it makes:

# # #     a) Borrowing more attractive and saving less attractive
# # #     b) Borrowing less attractive and saving more attractive
# # #     c) Both borrowing and saving less attractive
# # #     d) Cannot say
# # #     """,
# # #     """
# # #     16. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
# # #     a) Option A : Min Value is 9.2L and Max Value is 11.8L
# # #     b) Option B : Min Value is 8.8L and Max Value is 12.3L
# # #     c) Option C : Min Value is 8.5L and Max Value is 12.8L
# # #     d) Option D : Min Value is 8.1L and Max Value is 13.3L
# # #     e) Option E : Min Value is 7.8L and Max Value is 13.8L
# # #     """,
# # #     """
# # #     17. From Sept 2008 to Nov 2008, Stock market went down by 31%.

# # #     If you owned a stock investment that lost about 31 percent in 3 months, you would:
# # #     a) Sell all of the remaining investment
# # #     b) Sell a portion of the remaining investment
# # #     c) Hold on to the investment and sell nothing
# # #     d) Buy little
# # #     e) Buy more of the investment
# # #     """,
# # #     """
# # #     18. Over any 1-year period, what would be the maximum drop in the value of your investment 
# # #     portfolio that you would be comfortable with?
# # #     a) <5%
# # #     b) 5 - 10%
# # #     c) 10 - 15%
# # #     d) 15 - 20%
# # #     e) >20%
# # #     """,
# # #     """
# # #     19. When investing, what do you consider the most?

# # #     a) Risk 
# # #     b) Return
# # #     """,
# # #     """
# # #     20. What best describes your attitude?

# # #     a) Prefer reasonable returns, can take reasonable risk
# # #     b) Like higher returns, can take slightly higher risk
# # #     c) Want to maximize returns, can take significant high risk
# # #     """,
# # #     """
# # #     21. How much monthly investment you want to do?
# # #     """,
# # #     """
# # #     22. What is the time horizon for your investment?
# # #     You can answer in any range, example 1-5 years."""  
# # # ]


# # # # Handler for /start command
# # # @dp.message(CommandStart())
# # # async def handle_start(message: types.Message):
# # #     """
# # #     This handler receives messages with /start command
# # #     """
# # #     chat_id = message.chat.id
# # #     # Start asking questions
# # #     await start_assessment(chat_id)


# # # # Function to start the assessment
# # # async def start_assessment(chat_id):
# # #     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
# # #     await ask_next_question(chat_id, 0)

# # # # Function to ask the next question
# # # async def ask_next_question(chat_id, question_index):
# # #     if question_index < len(questions):
# # #         # Ask the next question
# # #         await bot.send_message(chat_id, questions[question_index])
# # #         # Update state to indicate the next expected answer
# # #         states[chat_id] = question_index
# # #     else:
# # #         # No more questions, finish assessment
# # #         await finish_assessment(chat_id)

# # # # Handler for receiving assessment answers
# # # assessment_in_progress = True

# # # from aiogram.types import FSInputFile
# # # async def finish_assessment(chat_id):
# # #     if chat_id in states and states[chat_id] == len(questions):
# # #         # All questions have been answered, now process the assessment
# # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # #         # Determine investment personality based on collected responses
# # #         global investment_personality
# # #         investment_personality = await determine_investment_personality(user_responses)

# # #         # Inform the user about their investment personality
# # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # #         # Summarize collected information
# # #         global summary
# # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # #         # Ensure to await the determination of investment personality
# # #         await send_summary_chunks(chat_id, summary)
# # #         global assessment_in_progress 
# # #         assessment_in_progress = False
# # #         # Prompt the user to begin financial advice process
# # #         # await bot.send_message(chat_id, "To receive financial advice, please upload a document with your financial details using /begin.")

# # #         await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
# # #         file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

# # #         await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
# # #         # await bot.send_message(chat_id,file)

# # # async def send_summary_chunks(chat_id, summary):
# # #     # Split the summary into chunks that fit within Telegram's message limits
# # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# # #     # Send each chunk as a separate message
# # #     for chunk in chunks:
# # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # # async def determine_investment_personality(assessment_data):
# # #     try:
# # #         # Prepare input text for the chatbot based on assessment data
# # #         input_text = "User Profile:\n"
# # #         for question, answer in assessment_data.items():
# # #             input_text += f"{question}: {answer}\n"

# # #         # Introduce the chatbot's task and prompt for classification
# # #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# # #                       "- Conservative Investor\n" \
# # #                       "- Moderate Investor\n" \
# # #                       "- Aggressive Investor"

# # #         # Use your generative AI model to generate a response
# # #         # print(input_text)
# # #         model = genai.GenerativeModel('gemini-pro')
# # #         response = model.generate_content(input_text)

# # #         # Determine the investment personality from the chatbot's response
# # #         response_text = response.text.lower()
# # #         if "conservative" in response_text:
# # #             personality = "Conservative Investor"
# # #         elif "moderate" in response_text:
# # #             personality = "Moderate Investor"
# # #         elif "aggressive" in response_text:
# # #             personality = "Aggressive Investor"
# # #         else:
# # #             personality = "Unknown"

# # #         return personality
# # #         # Send the determined investment personality back to the user
# # #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# # #     except Exception as e:
# # #         print(f"Error generating response: {e}")
# # #         #await bot.send_message(chat_id, "Error processing investment personality classification.")




# # # @router.message(F.text == "help")
# # # async def helper(message: types.Message):
# # #     """
# # #     A handler to display the help menu.
# # #     """
# # #     help_command = """
# # #     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
# # #     /start - to start the investment personality assessment.
# # #     /clear - to clear the past conversation and context.
# # #     /help - to get this help menu.
# # #     /begin - to start the Financial Suggestion Conversation with the ChatBot.
# # #     I hope this helps. :)
# # #     """
# # #     await message.reply(help_command)

# # # # Handler for /begin command to initiate financial advice

# # # @router.message(F.text == "begin")
# # # async def handle_begin(message: types.Message):
# # #     chat_id = message.chat.id
# # #     file_instructions ="""Hello there!My name is Finbot and I am a Wealth Management Advisor Chatbot.I need more details related to your Financial Profile so that I can give you 
# # #     personalised Financial Advice.Please follow this drive link : https://docs.google.com/document/d/1JaqWUXcq3MNrPTCvEy_2RngbGIY0rybX/edit?usp=sharing&ouid=101252223236130106095&rtpof=true&sd=true
# # #     ,there you will find a Word Document.Please fill in your details with correct Information and then upload the Document in the chat. """

# # #     await message.reply(file_instructions)
    
    

# # # # Handler for document upload

# # # # @dispatcher.message_handler(content_types=['document'])

# # # # @router.message(F.document)
# # # @dp.message(F.document)
# # # async def handle_document(message: types.Message):
# # #     global summary,investment_personality  

# # #     chat_id = message.chat.id

# # #     # Obtain file information
# # #     file_id = message.document.file_id
# # #     file = await bot.get_file(file_id)
# # #     file_path = file.file_path
    
# # #     # Download the file
# # #     await bot.download_file(file_path, "data/uploaded_file")
    
# # #     # Process the uploaded document
# # #     extracted_text = await process_document("data/uploaded_file")
    
# # #     if extracted_text:
# # #         # Load vector database (assuming this is part of setting up the retriever)
# # #         retriever = await load_vector_db("data/uploaded_file")
# # #         file_path = 'data/uploaded_file'
# # #         client_name, validation_errors = await process_document(file_path)

# # #         # Print results
# # #         print(f"Client Name: {client_name}")
# # #         if validation_errors:
# # #             print("**Validation Errors:**")
# # #             for error in validation_errors:
# # #                 print(error)
# # #         else:
# # #             print("All fields are filled correctly.")

# # #         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")
 
# # #         if retriever is None:
# # #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # #             return

# # #         # Check if a valid chain can be created
# # #         chain = await make_retrieval_chain(retriever)
# # #         if chain is None:
# # #             await message.reply("Failed to create the retrieval chain.")
# # #             return
        
# # #         try:     
# # #             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
           
# # #             response = chain.invoke({"input": query})
# # #             print(response['answer'])
# # #             global chat_history
# # #             chat_history = response['answer'] 
# # #             print(f"\n Chat History : {chat_history}")
# # #             await message.reply(response['answer'])

# # #             try:
# # #                 graph_query = (
# # #                     # summary + "\n" + investment_personality + "\n" + response['answer'] + "\n" +  
# # #                     # "Please provide the following information in JSON format for the given response :\n" +
# # #                     # "{\n"
# # #                     # "  'growth_investment_min': <minimum percentage of growth-oriented investments>,\n"
# # #                     # "  'growth_investment_max': <maximum percentage of growth-oriented investments>,\n"
# # #                     # "  'return_min': <minimum expected annual return percentage>,\n"
# # #                     # "  'return_max': <maximum expected annual return percentage>,\n"
# # #                     # "  'growth_min': <minimum expected growth in dollars>,\n"
# # #                     # "  'growth_max': <maximum expected growth in dollars>\n"
# # #                     # "}"

# # #                     # extracted_text + "\n" + 
# # #                     summary + "\n" + "investment_personality" + "\n" + response['answer'] + "\n" +
# # #                     "Please provide the following information in JSON format:\n" +
# # #                     "{\n"
# # #                     "  'growth_investment_min': <minimum percentage of growth-oriented investments>,\n"
# # #                     "  'growth_investment_max': <maximum percentage of growth-oriented investments>,\n"
# # #                     "  'fixed_income_min': <minimum percentage of fixed-income investments>,\n"
# # #                     "  'fixed_income_max': <maximum percentage of fixed-income investments>,\n"
# # #                     "  'cash_min': <minimum percentage of cash investments>,\n"
# # #                     "  'cash_max': <maximum percentage of cash investments>,\n"
# # #                     "  'return_growth_investment': <Maximum percentage returns of growth-oriented investments>,\n"
# # #                     "  'return_fixed_income': <Maximum percentage of returns of fixed-income investments>,\n"
# # #                     "  'return_cash_max': <maximum percentage of retruns of cash investments>,\n"
# # #                     "  'return_min': <minimum expected annual return percentage>,\n"
# # #                     "  'return_max': <maximum expected annual return percentage>,\n"
# # #                     "  'growth_min': <minimum expected growth in dollars>,\n"
# # #                     "  'growth_max': <maximum expected growth in dollars>,\n"
# # #                     "  'initial_investment': <initial monthly investment>,\n"
# # #                     "  'time_horizon': <time horizon in years>\n"
# # #                     "}"
# # #                 )
# # #                 graph_response = chain.invoke({"input": graph_query})
# # #                 print(graph_response['answer'])
# # #                 await handle_graph(graph_response['answer'],chat_id)
# # #             except Exception as e:
# # #                 print(f"Error plotting graph : {e}")
# # #         except Exception as e:
# # #             print(f"Error invoking retrieval chain on attempt : {e}")
# # #             await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")
   
# # #     else:
# # #         await message.reply("Failed to process the uploaded file.")


# # # # Function to extract data from LLM response
# # # def extract_data_from_response(response):
# # #     try:
# # #         # Locate the JSON-like data in the response
# # #         json_start = response.find("{")
# # #         json_end = response.rfind("}") + 1
        
# # #         if json_start == -1 or json_end == -1:
# # #             raise ValueError("No JSON data found in the response.")
        
# # #         json_data = response[json_start:json_end]
        
# # #         # Parse the JSON data
# # #         data = json.loads(json_data.replace("'", "\""))
# # #         print(data)
# # #         return data
# # #     except Exception as e:
# # #         logging.error(f"Error extracting data: {e}")
# # #         return None


# # # async def handle_graph(response, chat_id):
# # #     try:
# # #         data = extract_data_from_response(response)
# # #         if not data:
# # #             await bot.send_message(chat_id, "Failed to extract data from the response.")
# # #             return
        
# # #         # Log extracted data for debugging
# # #         logging.info(f"Extracted data: {data}")
        
# # #         # Create a pie chart for investment allocation
# # #         labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
# # #         sizes = [
# # #             (data['growth_investment_min'] + data['growth_investment_max']) / 2,
# # #             (data['fixed_income_min'] + data['fixed_income_max']) / 2,
# # #             (data['cash_min'] + data['cash_max']) / 2
# # #         ]
# # #         fig1, ax1 = plt.subplots()
# # #         ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# # #         ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
# # #         pie_chart_buffer = io.BytesIO()
# # #         plt.savefig(pie_chart_buffer, format='png')
# # #         pie_chart_buffer.seek(0)
# # #         pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

# # #         # Create a bar graph for potential returns
# # #         allocations = {
# # #             'Growth-Oriented Investments': (data['growth_investment_min'] + data['growth_investment_max']) / 2,
# # #             'Fixed-Income Investments': (data['fixed_income_min'] + data['fixed_income_max']) / 2,
# # #             'Cash and Cash Equivalents': (data['cash_min'] + data['cash_max']) / 2
# # #         }
        
# # #         # returns = {
# # #         #     'Growth-Oriented Investments': ((data['return_min'] + data['return_max']) / 2) / 100,  # Annual return
# # #         #     'Fixed-Income Investments': 0.04,  # Assume a 4% annual return for fixed income
# # #         #     'Cash and Cash Equivalents': 0.01  # Assume a 1% annual return for cash equivalents
# # #         # }

# # #         returns = {
# # #             'Growth-Oriented Investments': (data['return_growth_investment'] ) / 100,  # Annual return
# # #             'Fixed-Income Investments': (data['return_fixed_income'] ) / 100 ,
# # #             'Cash and Cash Equivalents': (data['return_cash_max'] ) / 100
# # #         }
        
# # #         time_horizon_years = data['time_horizon']
# # #         initial_monthly_investment = data['initial_investment']
        
# # #         # Calculate total returns for each category
# # #         total_investment = initial_monthly_investment * 12 * time_horizon_years
# # #         total_returns = {category: total_investment * (1 + returns[category]) ** time_horizon_years for category in allocations.keys()}
        
# # #         # Create the bar chart
# # #         fig2, ax2 = plt.subplots(figsize=(10, 6))
# # #         bars = ax2.bar(allocations.keys(), total_returns.values(), color=['#FF9999', '#66B2FF', '#99FF99'])

# # #         # Add labels to the top of each bar
# # #         for bar in bars:
# # #             yval = bar.get_height()
# # #             ax2.text(bar.get_x() + bar.get_width()/2, yval + 1000, f'₹{round(yval, 2)}', ha='center', va='bottom')

# # #         # Add titles and labels
# # #         plt.title('Investment Performance Over Time')
# # #         plt.xlabel('Investment Categories')
# # #         plt.ylabel('Total Returns (₹)')
# # #         plt.grid(True, axis='y', linestyle='--', linewidth=0.7)
        
# # #         # Save and display the plot
# # #         bar_chart_buffer = io.BytesIO()
# # #         plt.savefig(bar_chart_buffer, format='png')
# # #         bar_chart_buffer.seek(0)
# # #         bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='investment_performance.png')

# # #         await bot.send_document(chat_id, pie_chart_file, caption="Investment Allocation Chart")
# # #         await bot.send_document(chat_id, bar_chart_file, caption="Investment Performance Over Time")

# # #         # Generate the investment growth graph : work to be done 


# # #         # months = data['time_horizon'] * 12  # Use the provided time horizon
# # #         # initial_investment = data['initial_investment']  # Use the provided initial monthly investment
# # #         # investment_values = [0]

# # #         # for month in range(1, months + 1):
# # #         #     monthly_return = ((data['return_min'] + data['return_max']) / 2) / 100 / 12
# # #         #     investment_values.append(investment_values[-1] * (1 + monthly_return) + initial_investment)

# # #         # fig3, ax3 = plt.subplots()
# # #         # ax3.plot(range(0, months + 1), investment_values, marker='o')
# # #         # ax3.set(title='Investment Growth Over Time', xlabel='Month', ylabel='Investment Value ($)')
# # #         # ax3.grid(True)
# # #         # investment_buffer = io.BytesIO()
# # #         # plt.savefig(investment_buffer, format='png')
# # #         # investment_buffer.seek(0)
# # #         # investment_file = BufferedInputFile(investment_buffer.read(), filename='investment_growth.png')

# # #         # await bot.send_document(chat_id, investment_file, caption="Investment Growth Over Time")

# # #     except Exception as e:
# # #         logging.error(f"Error plotting graph: {e}")
# # #         await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")







# # # # Generate sample plots and send them to the user (previous version)
# # # # async def handle_graph(response, chat_id):
# # # #     try:
# # # #         data = extract_data_from_response(response)
# # # #         if not data:
# # # #             await bot.send_message(chat_id, "Failed to extract data from the response.")
# # # #             return
        
# # # #         # Log extracted data for debugging
# # # #         logging.info(f"Extracted data: {data}")
        
# # # #         # Create a pie chart for investment allocation
# # # #         labels = ['Growth-Oriented Investments', 'Fixed-Income Investments', 'Cash and Cash Equivalents']
# # # #         sizes = [
# # # #             (data['growth_investment_min'] + data['growth_investment_max']) / 2,
# # # #             (data['fixed_income_min'] + data['fixed_income_max']) / 2,
# # # #             (data['cash_min'] + data['cash_max']) / 2
# # # #         ]
# # # #         fig1, ax1 = plt.subplots()
# # # #         ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# # # #         ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
# # # #         pie_chart_buffer = io.BytesIO()
# # # #         plt.savefig(pie_chart_buffer, format='png')
# # # #         pie_chart_buffer.seek(0)
# # # #         pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

# # # #         # Create a bar graph for potential returns
# # # #         years = list(range(1, data['time_horizon'] + 1))
# # # #         min_returns = [data['return_min']] * data['time_horizon']
# # # #         max_returns = [data['return_max']] * data['time_horizon']

# # # #         fig2, ax2 = plt.subplots()
# # # #         ax2.bar(years, min_returns, label='Min Return')
# # # #         ax2.bar(years, max_returns, bottom=min_returns, label='Max Return')
# # # #         ax2.set(xlabel='Year', ylabel='Return (%)', title='Expected Returns Over Time')
# # # #         ax2.legend()
# # # #         bar_chart_buffer = io.BytesIO()
# # # #         plt.savefig(bar_chart_buffer, format='png')
# # # #         bar_chart_buffer.seek(0)
# # # #         bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='bar_chart.png')

# # # #         # Send the generated graphs to the user with retry mechanism
# # # #         await bot.send_document(chat_id, pie_chart_file, caption="Investment Allocation Chart")
# # # #         await bot.send_document(chat_id, bar_chart_file, caption="Expected Returns Over Time")

# # # #         # Generate the investment growth graph
# # # #         months = data['time_horizon'] * 12  # Use the provided time horizon
# # # #         initial_investment = data['initial_investment']  # Use the provided initial monthly investment
# # # #         investment_values = [0]

# # # #         for month in range(1, months + 1):
# # # #             monthly_return = ((data['return_min'] + data['return_max']) / 2) / 100 / 12
# # # #             investment_values.append(investment_values[-1] * (1 + monthly_return) + initial_investment)

# # # #         fig3, ax3 = plt.subplots()
# # # #         ax3.plot(range(0, months + 1), investment_values, marker='o')
# # # #         ax3.set(title='Investment Growth Over Time', xlabel='Month', ylabel='Investment Value ($)')
# # # #         ax3.grid(True)
# # # #         investment_buffer = io.BytesIO()
# # # #         plt.savefig(investment_buffer, format='png')
# # # #         investment_buffer.seek(0)
# # # #         investment_file = BufferedInputFile(investment_buffer.read(), filename='investment_growth.png')

# # # #         await bot.send_document(chat_id, investment_file, caption="Investment Growth Over Time")

# # # #     except Exception as e:
# # # #         logging.error(f"Error plotting graph: {e}")
# # # #         await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")


# # # # import json
# # # # # Function to send document with retry mechanism 
# # # # import re
# # # # async def send_document_with_retry(chat_id, document, caption, retries=3):
# # # #     for attempt in range(retries):
# # # #         try:
# # # #             await bot.send_document(chat_id, document, caption=caption)
# # # #             return
# # # #         except aiohttp.ClientConnectionError as e:
# # # #             logging.error(f"Attempt {attempt + 1} failed: {e}")
# # # #             if attempt < retries - 1:
# # # #                 continue
# # # #             else:
# # # #                 raise e
            
# # # # Function to extract data from LLM response
# # # # # Function to extract data from LLM response
# # # # def extract_data_from_response(response):
# # # #     try:
# # # #         # Locate the JSON-like data in the response
# # # #         json_start = response.find("{")
# # # #         json_end = response.rfind("}") + 1
        
# # # #         if json_start == -1 or json_end == -1:
# # # #             raise ValueError("No JSON data found in the response.")
        
# # # #         json_data = response[json_start:json_end]
        
# # # #         # Parse the JSON data
# # # #         data = json.loads(json_data.replace("'", "\""))
# # # #         return data
# # # #     except Exception as e:
# # # #         logging.error(f"Error extracting data: {e}")
# # # #         return None



# # # # # Generate sample plots and send them to the user
# # # # async def handle_graph(response, chat_id):
# # # #     try:
# # # #         data = extract_data_from_response(response)
# # # #         if not data:
# # # #             await bot.send_message(chat_id, "Failed to extract data from the response.")
# # # #             return
            
        
# # # #         # Create a pie chart for investment allocation
# # # #         labels = 'Growth-Oriented Investments', 'Conservative Investments'
# # # #         sizes = [data['growth_investment_max'], 100 - data['growth_investment_max']]
# # # #         fig1, ax1 = plt.subplots()
# # # #         ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# # # #         ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
# # # #         pie_chart_buffer = io.BytesIO()
# # # #         plt.savefig(pie_chart_buffer, format='png')
# # # #         pie_chart_buffer.seek(0)
# # # #         pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

# # # #         # Create a bar graph for potential returns
# # # #         return_labels = ['Min Return', 'Max Return']
# # # #         return_values = [data['return_min'], data['return_max']]
# # # #         fig2, ax2 = plt.subplots()
# # # #         sns.barplot(x=return_labels, y=return_values, ax=ax2)
# # # #         ax2.set(xlabel='Return Type', ylabel='Percentage')
# # # #         bar_chart_buffer = io.BytesIO()
# # # #         plt.savefig(bar_chart_buffer, format='png')
# # # #         bar_chart_buffer.seek(0)
# # # #         bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='bar_chart.png')

# # # #         # Send the generated graphs to the user with retry mechanism
# # # #         await send_document_with_retry(chat_id, pie_chart_file, caption="Investment Allocation Chart")
# # # #         await send_document_with_retry(chat_id, bar_chart_file, caption="Potential Returns Chart")

# # # #         # Generate the investment growth graph
# # # #         fig3, ax3 = plt.subplots()
# # # #         growth_labels = ['Min Growth', 'Max Growth']
# # # #         growth_values = [data['growth_min'], data['growth_max']]
# # # #         ax3.bar(growth_labels, growth_values)
# # # #         ax3.set(title='Investment Growth Over Time', xlabel='Growth Type', ylabel='Value ($)')
# # # #         ax3.grid(True)
# # # #         investment_buffer = io.BytesIO()
# # # #         plt.savefig(investment_buffer, format='png')
# # # #         investment_buffer.seek(0)
# # # #         investment_file = BufferedInputFile(investment_buffer.read(), filename='investment_growth.png')

# # # #         await send_document_with_retry(chat_id, investment_file, caption="Investment Growth Over Time")

# # # #     except Exception as e:
# # # #         logging.error(f"Error plotting graph: {e}")
# # # #         await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")


# # # import aiohttp
# # # from aiogram.types import InputFile , BufferedInputFile



# # # async def process_document(file_path):
# # #     try:
# # #         file_type = filetype.guess(file_path)
# # #         if file_type is not None:
# # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # #                 # return extract_text_from_word(file_path)
# # #                 return extract_text_and_tables_from_word(file_path)
# # #             elif file_type.mime == "application/pdf":
# # #                 return extract_text_from_pdf(file_path)
# # #         return None
# # #     except Exception as e:
# # #         print(f"Error processing document: {e}")
# # #         return None

# # # def extract_text_from_pdf(pdf_file_path):
# # #     try:
# # #         with open(pdf_file_path, "rb") as pdf_file:
# # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # #             text_content = []
# # #             for page_num in range(pdf_reader.numPages):
# # #                 page = pdf_reader.getPage(page_num)
# # #                 text_content.append(page.extract_text())
# # #             return "\n".join(text_content)
# # #     except Exception as e:
# # #         print(f"Error extracting text from PDF: {e}")
# # #         return None


# # # import re
# # # import docx

# # # def extract_text_and_tables_from_word(docx_file_path):
# # #     """
# # #     Extracts text and tables from a Word document (.docx).

# # #     Args:
# # #         docx_file_path (str): Path to the Word document file.

# # #     Returns:
# # #         tuple: Extracted text content and tables from the document.
# # #     """
# # #     try:
# # #         doc = docx.Document(docx_file_path)
# # #         text_content = []
# # #         tables_content = []

# # #         for para in doc.paragraphs:
# # #             text_content.append(para.text)

# # #         for table in doc.tables:
# # #             table_data = []
# # #             for row in table.rows:
# # #                 row_data = []
# # #                 for cell in row.cells:
# # #                     row_data.append(cell.text.strip())
# # #                 table_data.append(row_data)
# # #             tables_content.append(table_data)

# # #         return "\n".join(text_content), tables_content
# # #     except Exception as e:
# # #         print(f"Error extracting text and tables from Word document: {e}")
# # #         return None, None

# # # def validate_document_content(text, tables):
# # #     """
# # #     Validates the content of the document.

# # #     Args:
# # #         text (str): Extracted text content from the document.
# # #         tables (list): Extracted tables content from the document.

# # #     Returns:
# # #         tuple: Client name and validation errors.
# # #     """
# # #     errors = []
    
# # #     # Extract client name
# # #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# # #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# # #     # Define required sections
# # #     required_sections = [
# # #         "YOUR RETIREMENT GOAL",
# # #         "YOUR OTHER MAJOR GOALS",
# # #         "YOUR ASSETS AND LIABILITIES",
# # #         "MY LIABILITIES",
# # #         "YOUR CURRENT ANNUAL INCOME"
# # #     ]

# # #     # Check for the presence of required sections
# # #     for section in required_sections:
# # #         if section not in text:
# # #             errors.append(f"* {section} section missing.")
    
# # #     # Define table field checks
# # #     table_checks = {
# # #         "YOUR RETIREMENT GOAL": [
# # #             r"When do you plan to retire\? \(age or date\)",
# # #             r"Social Security Benefit \(include expected start date\)",
# # #             r"Pension Benefit \(include expected start date\)",
# # #             r"Other Expected Income \(rental, part-time work, etc.\)",
# # #             r"Estimated Annual Retirement Expense"
# # #         ],
# # #         "YOUR OTHER MAJOR GOALS": [
# # #             r"GOAL", r"COST", r"WHEN"
# # #         ],
# # #         "YOUR ASSETS AND LIABILITIES": [
# # #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# # #             r"Current Value", r"Annual Contributions"
# # #         ],
# # #         "MY LIABILITIES": [
# # #             r"Balance", r"Interest Rate", r"Monthly Payment"
# # #         ]
# # #     }

# # #     # Validate table content
# # #     for section, checks in table_checks.items():
# # #         section_found = False
# # #         for table in tables:
# # #             table_text = "\n".join(["\t".join(row) for row in table])
# # #             if section in table_text:
# # #                 section_found = True
# # #                 for check in checks:
# # #                     if not re.search(check, table_text, re.IGNORECASE):
# # #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# # #                 break
# # #         if not section_found:
# # #             errors.append(f"* {section} section missing.")

# # #     return client_name, errors

# # # async def process_document(file_path):
# # #     try:
# # #         text, tables = extract_text_and_tables_from_word(file_path)
# # #         if text is not None and tables is not None:
# # #             client_name, errors = validate_document_content(text, tables)
# # #             return client_name, errors
# # #         return None, ["Error processing document."]
# # #     except Exception as e:
# # #         print(f"Error processing document: {e}")
# # #         return None, [f"Error processing document: {e}"]



# # # async def load_vector_db(file_path):
# # #     try:
# # #         loader = Docx2txtLoader(file_path)
# # #         documents = loader.load()
# # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # #         text_chunks = text_splitter.split_documents(documents)
# # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # #         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
# # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # #     except Exception as e:
# # #         print(f"Error loading vector database: {e}")
# # #         return None



# # # async def make_retrieval_chain(retriever):
# # #     """
# # #     Create a retrieval chain using the provided retriever.

# # #     Args:
# # #         retriever (RetrievalQA): A retriever object.

# # #     Returns:
# # #         RetrievalQA: A retrieval chain object.
# # #     """
# # #     try:
# # #         global investment_personality,summary
# # #         llm = ChatGoogleGenerativeAI(
# # #             model="gemini-pro",
# # #             temperature=0.7,
# # #             top_p=0.85,
# # #             google_api_key=GOOGLE_API_KEY
# # #         )

# # #         prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # #                 Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # #                 should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
# # #                 Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # #                 investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
# # #                 Also give the user minimum and maximum expected growth in dollars for the time horizon .
# # #                 Also explain the user why you are giving them that particular investment suggestion.Answer in 3-4 lines.\n
# # #                 <context>
# # #                 {context}
# # #                 </context>
# # #                 Question: {input}"""

# # #         #  prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # #         #         Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# # #         #         should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # #         #         investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # #         #         investment.Also explain the user why you are giving them that particular
# # #         #         investment suggestion.Answer in 3-4 lines.\n
# # #         #         <context>
# # #         #         {context}
# # #         #         </context>
# # #         #         Question: {input}"""

# # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
# # #         # Update combine_docs_chain with your actual document combining logic
# # #         combine_docs_chain = None  # Replace this with your combine_docs_chain

# # #         if retriever is not None :  #and combine_docs_chain is not None:
# # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # #             print(retriever_chain)
# # #             return retriever_chain
# # #         else:
# # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # #             return None

# # #     except Exception as e:
# # #         print(f"Error in creating chain: {e}")
# # #         return None


# # # # @dispatcher.message_handler()
# # # @dp.message()
# # # async def main_bot(message: types.Message):
# # #     global retriever, extracted_text,investment_personality,summary,chat_history

# # #     # Handle the first tasks assessments answers from the user
# # #     chat_id = message.chat.id

# # #     if chat_id in states and states[chat_id] < len(questions):
# # #         # Retrieve the index of the current question
# # #         question_index = states[chat_id]

# # #         # Save the user's response to the current question
# # #         answer = message.text
# # #         user_responses[questions[question_index]] = answer
# # #         states[chat_id] += 1  # Move to the next question

# # #         # Ask the next question
# # #         await ask_next_question(chat_id, question_index + 1)
# # #     else:
# # #         # Handle q&a chat messages using your Gemini model (llm)
# # #         try:

# # #             task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # #             Also give the user detailed information about the investment how to invest,where to invest and how much they
# # #             should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # #             investment.Also explain the user why you are giving them that particular
# # #             investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # #             User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # #             So dont mention user to consult to a financial expert."""
        
# # #             model = genai.GenerativeModel('gemini-pro')
# # #             print(investment_personality)
# # #             # query = task + "\n" + investment_personality + "\n" + summary + "\n" +  extracted_text + "\n"  +   message.text

# # #             # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +  extracted_text + "\n"  +   message.text
# # #             query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
# # #             print(f"\nQuery : {query}")
# # #             response = model.generate_content(query)
# # #             await message.reply(response.text) #(response['answer']) 
           
# # #         except Exception as e:
# # #             print(f"Error processing general chat message: {e}")
# # #             await message.reply("Failed to process your request.")
        


# # # # if __name__ == "__main__":
# # # #     executor.start_polling(dispatcher, skip_updates=True)

# # # async def main() -> None:
# # #     # Initialize Bot instance with default bot properties which will be passed to all API calls
# # #     bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

# # #     # And the run events dispatching
# # #     await dp.start_polling(bot)


# # # if __name__ == "__main__":
# # #     logging.basicConfig(level=logging.INFO, stream=sys.stdout)
# # #     asyncio.run(main())





# # # # #Best Code it can upload files generate graphs and can chat with the user as well :)"""
# # # # #best code so far now it can upload files and receive various forms of messages as well 

# # # # import os
# # # # import filetype
# # # # import docx
# # # # import PyPDF2
# # # # import re
# # # # from aiogram import Bot, Dispatcher, types
# # # # from dotenv import load_dotenv
# # # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # # from langchain_community.vectorstores import Chroma
# # # # from langchain_community.document_loaders import Docx2txtLoader
# # # # from langchain_core.prompts import ChatPromptTemplate
# # # # from langchain.chains import create_retrieval_chain
# # # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # # from langchain.memory import ConversationSummaryMemory
# # # # import asyncio
# # # # import numpy as np

# # # # import google.generativeai as genai

# # # # # Import things that are needed generically
# # # # from langchain.pydantic_v1 import BaseModel, Field
# # # # from langchain.tools import BaseTool, StructuredTool, tool

# # # # from aiogram.client.default import DefaultBotProperties
# # # # from aiogram.enums import ParseMode
# # # # from aiogram.filters import CommandStart
# # # # from aiogram.types import Message
# # # # from aiogram import F
# # # # from aiogram import Router
# # # # import logging
# # # # import sys
# # # # from aiogram.filters import Command
# # # # from aiogram.types import FSInputFile
# # # # # from aiogram.utils import executor
# # # # import io
# # # # import matplotlib.pyplot as plt
# # # # import seaborn as sns

# # # # router = Router(name=__name__)

# # # # load_dotenv()

# # # # TOKEN = os.getenv("TOKEN")
# # # # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # # # os.environ["LANGCHAIN_TRACING_V2"]="true"
# # # # os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # # # # Configure generativeai with your API key
# # # # genai.configure(api_key=GOOGLE_API_KEY)

# # # # # Initialize bot
# # # # bot = Bot(token=TOKEN)
# # # # dp = Dispatcher()

# # # # rag_on = False
# # # # retriever = None  # Store retriever globally
# # # # summary = ""
# # # # investment_personality = ""
# # # # chat_history = ""

# # # # class Reference:
# # # #     def __init__(self):
# # # #         self.response = ""


# # # # reference = Reference()


# # # # def clear_past():
# # # #     reference.response = ""


# # # # @router.message(F.text == "clear")
# # # # async def clear(message: types.Message):
# # # #     """
# # # #     A handler to clear the previous conversation and context.
# # # #     """
# # # #     clear_past()
# # # #     await message.reply("I've cleared the past conversation and context.")


# # # # # Store user states
# # # # states = {}
# # # # # Dictionary to hold question-answer pairs
# # # # user_responses = {}

# # # # # Define Questions for assessment
# # # # questions = [
# # # #     """
# # # #         1. Singapore plans to build a new observation tower called 'The Rook'.
# # # #         How many steps do you think it will take to go to the top floor?

# # # #         a) Less than 500 
# # # #         b) More than 500

# # # #     """,
# # # #     "2. Now Guess the number of steps" ,
# # # #     """
# # # #     3. How confident are you that the real number is in the range you have selected? 
# # # #     Answer within a range of 100.  
# # # #     """,
# # # #     """ 
# # # #     4. You and your friend are betting on a series of coin tosses.

# # # #     He always bets ₹2,000 on Heads

# # # #     You always bet ₹2,000 on Tails

# # # #     Winner of last 8 turns

# # # #     You lost ₹8,000 in the last 4 turns!

# # # #     If you were to bet one last time, what would you bet on heads or tails ?
# # # #     """ ,
# # # #     """
# # # #     5. How confident are you that your bet will win this time?
# # # #     Answer how confident you are. 
# # # #     (Example: Not confident at all, somewhat confident, confident, or Very confident)
# # # #     """,
# # # #     """
# # # #     6. What do you think are the chances that you will achieve and maintain your financial goals within the next 
# # # #     10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
# # # #     Answer how likely you are to achieve your goal.
# # # #     (Example: Less likely than others, likely than others, or More likely than others)
# # # #     """,
# # # #     """
# # # #     7. Imagine you are a contestant in a game show, and you are presented the following choices.

# # # #     What would you prefer?
# # # #     a) 50 percent chance of winning 15 gold coins 
# # # #     b) 100 percent chance of winning 8 gold coins
# # # #     """,
# # # #     """
# # # #     8. Ok, one last choice...

# # # #     What would you prefer?
# # # #     a) 50 percent chance of winning 15 gold coins 
# # # #     b) 100 percent chance of winning 2 gold coins
# # # #     """,
# # # #     """
# # # #     9. In general, how would your best friend describe your risk-taking tendencies?
# # # #     a) A real gambler
# # # #     b) Willing to take risks after completing adequate research
# # # #     c) Cautious
# # # #     d) Avoids risk as much as possible
# # # #     """,
# # # #     """
# # # #     10. Suppose you could replace your current investment portfolio with this new one:
# # # #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# # # #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# # # #     Answer between the range of -5 to -35.
# # # #     """,
# # # #     """
# # # #     11. Suppose that in the next 7 years,

# # # #     YOUR INCOME

# # # #     grows 8% each year

# # # #     VS
# # # #     INFLATION

# # # #     grows 10% a year

# # # #     At the end of 7 years, how much will you be able to buy with your income?
# # # #     Options:
# # # #     a) More than today
# # # #     b) Exactly the same
# # # #     c) Less than today
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     12. If somebody buys a bond of Company B, which of the following statements seems correct:
# # # #     a) She owns part of Company B
# # # #     b) She has lent money to Company B
# # # #     c) She is liable for Company B's debt
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     13. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
# # # #     a) More than ₹2 lakhs
# # # #     b) Less than ₹2 lakhs
# # # #     c) Exactly ₹2 lakhs
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     14. When an investor spreads money across different asset classes, what happens to the risk of losing money:
# # # #     a) Increases
# # # #     b) Decreases
# # # #     c) Stays the same
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     15. When a country's central bank reduces interest rates, it makes:

# # # #     a) Borrowing more attractive and saving less attractive
# # # #     b) Borrowing less attractive and saving more attractive
# # # #     c) Both borrowing and saving less attractive
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     16. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
# # # #     a) Option A : Min Value is 9.2L and Max Value is 11.8L
# # # #     b) Option B : Min Value is 8.8L and Max Value is 12.3L
# # # #     c) Option C : Min Value is 8.5L and Max Value is 12.8L
# # # #     d) Option D : Min Value is 8.1L and Max Value is 13.3L
# # # #     e) Option E : Min Value is 7.8L and Max Value is 13.8L
# # # #     """,
# # # #     """
# # # #     17. From Sept 2008 to Nov 2008, Stock market went down by 31%.

# # # #     If you owned a stock investment that lost about 31 percent in 3 months, you would:
# # # #     a) Sell all of the remaining investment
# # # #     b) Sell a portion of the remaining investment
# # # #     c) Hold on to the investment and sell nothing
# # # #     d) Buy little
# # # #     e) Buy more of the investment
# # # #     """,
# # # #     """
# # # #     18. Over any 1-year period, what would be the maximum drop in the value of your investment 
# # # #     portfolio that you would be comfortable with?
# # # #     a) <5%
# # # #     b) 5 - 10%
# # # #     c) 10 - 15%
# # # #     d) 15 - 20%
# # # #     e) >20%
# # # #     """,
# # # #     """
# # # #     19. When investing, what do you consider the most?

# # # #     a) Risk 
# # # #     b) Return
# # # #     """,
# # # #     """
# # # #     20. What best describes your attitude?

# # # #     a) Prefer reasonable returns, can take reasonable risk
# # # #     b) Like higher returns, can take slightly higher risk
# # # #     c) Want to maximize returns, can take significant high risk
# # # #     """,
# # # #     """
# # # #     21. How much monthly investment you want to do?
# # # #     """,
# # # #     """
# # # #     22. What is the time horizon for your investment?
# # # #     You can answer in any range, example 1-5 years."""  
# # # # ]


# # # # # Handler for /start command
# # # # @dp.message(CommandStart())
# # # # async def handle_start(message: types.Message):
# # # #     """
# # # #     This handler receives messages with /start command
# # # #     """
# # # #     chat_id = message.chat.id
# # # #     # Start asking questions
# # # #     await start_assessment(chat_id)


# # # # # Function to start the assessment
# # # # async def start_assessment(chat_id):
# # # #     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
# # # #     await ask_next_question(chat_id, 0)

# # # # # Function to ask the next question
# # # # async def ask_next_question(chat_id, question_index):
# # # #     if question_index < len(questions):
# # # #         # Ask the next question
# # # #         await bot.send_message(chat_id, questions[question_index])
# # # #         # Update state to indicate the next expected answer
# # # #         states[chat_id] = question_index
# # # #     else:
# # # #         # No more questions, finish assessment
# # # #         await finish_assessment(chat_id)

# # # # # Handler for receiving assessment answers
# # # # assessment_in_progress = True

# # # # from aiogram.types import FSInputFile
# # # # async def finish_assessment(chat_id):
# # # #     if chat_id in states and states[chat_id] == len(questions):
# # # #         # All questions have been answered, now process the assessment
# # # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # # #         # Determine investment personality based on collected responses
# # # #         global investment_personality
# # # #         investment_personality = await determine_investment_personality(user_responses)

# # # #         # Inform the user about their investment personality
# # # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # # #         # Summarize collected information
# # # #         global summary
# # # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # # #         # Ensure to await the determination of investment personality
# # # #         await send_summary_chunks(chat_id, summary)
# # # #         global assessment_in_progress 
# # # #         assessment_in_progress = False
# # # #         # Prompt the user to begin financial advice process
# # # #         # await bot.send_message(chat_id, "To receive financial advice, please upload a document with your financial details using /begin.")

# # # #         await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
# # # #         file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

# # # #         await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
# # # #         # await bot.send_message(chat_id,file)

# # # # async def send_summary_chunks(chat_id, summary):
# # # #     # Split the summary into chunks that fit within Telegram's message limits
# # # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# # # #     # Send each chunk as a separate message
# # # #     for chunk in chunks:
# # # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # # # async def determine_investment_personality(assessment_data):
# # # #     try:
# # # #         # Prepare input text for the chatbot based on assessment data
# # # #         input_text = "User Profile:\n"
# # # #         for question, answer in assessment_data.items():
# # # #             input_text += f"{question}: {answer}\n"

# # # #         # Introduce the chatbot's task and prompt for classification
# # # #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# # # #                       "- Conservative Investor\n" \
# # # #                       "- Moderate Investor\n" \
# # # #                       "- Aggressive Investor"

# # # #         # Use your generative AI model to generate a response
# # # #         # print(input_text)
# # # #         model = genai.GenerativeModel('gemini-pro')
# # # #         response = model.generate_content(input_text)

# # # #         # Determine the investment personality from the chatbot's response
# # # #         response_text = response.text.lower()
# # # #         if "conservative" in response_text:
# # # #             personality = "Conservative Investor"
# # # #         elif "moderate" in response_text:
# # # #             personality = "Moderate Investor"
# # # #         elif "aggressive" in response_text:
# # # #             personality = "Aggressive Investor"
# # # #         else:
# # # #             personality = "Unknown"

# # # #         return personality
# # # #         # Send the determined investment personality back to the user
# # # #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# # # #     except Exception as e:
# # # #         print(f"Error generating response: {e}")
# # # #         #await bot.send_message(chat_id, "Error processing investment personality classification.")




# # # # @router.message(F.text == "help")
# # # # async def helper(message: types.Message):
# # # #     """
# # # #     A handler to display the help menu.
# # # #     """
# # # #     help_command = """
# # # #     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
# # # #     /start - to start the investment personality assessment.
# # # #     /clear - to clear the past conversation and context.
# # # #     /help - to get this help menu.
# # # #     /begin - to start the Financial Suggestion Conversation with the ChatBot.
# # # #     I hope this helps. :)
# # # #     """
# # # #     await message.reply(help_command)

# # # # # Handler for /begin command to initiate financial advice

# # # # @router.message(F.text == "begin")
# # # # async def handle_begin(message: types.Message):
# # # #     chat_id = message.chat.id
# # # #     file_instructions ="""Hello there!My name is Finbot and I am a Wealth Management Advisor Chatbot.I need more details related to your Financial Profile so that I can give you 
# # # #     personalised Financial Advice.Please follow this drive link : https://docs.google.com/document/d/1JaqWUXcq3MNrPTCvEy_2RngbGIY0rybX/edit?usp=sharing&ouid=101252223236130106095&rtpof=true&sd=true
# # # #     ,there you will find a Word Document.Please fill in your details with correct Information and then upload the Document in the chat. """

# # # #     await message.reply(file_instructions)
    
    

# # # # # Handler for document upload

# # # # # @dispatcher.message_handler(content_types=['document'])

# # # # # @router.message(F.document)
# # # # @dp.message(F.document)
# # # # async def handle_document(message: types.Message):
# # # #     global summary,investment_personality  

# # # #     chat_id = message.chat.id

# # # #     # Obtain file information
# # # #     file_id = message.document.file_id
# # # #     file = await bot.get_file(file_id)
# # # #     file_path = file.file_path
    
# # # #     # Download the file
# # # #     await bot.download_file(file_path, "data/uploaded_file")
    
# # # #     # Process the uploaded document
# # # #     extracted_text = await process_document("data/uploaded_file")
    
# # # #     if extracted_text:
# # # #         # Load vector database (assuming this is part of setting up the retriever)
# # # #         retriever = await load_vector_db("data/uploaded_file")
# # # #         file_path = 'data/uploaded_file'
# # # #         client_name, validation_errors = await process_document(file_path)

# # # #         # Print results
# # # #         print(f"Client Name: {client_name}")
# # # #         if validation_errors:
# # # #             print("**Validation Errors:**")
# # # #             for error in validation_errors:
# # # #                 print(error)
# # # #         else:
# # # #             print("All fields are filled correctly.")

# # # #         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")
 
# # # #         if retriever is None:
# # # #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # #             return

# # # #         # Check if a valid chain can be created
# # # #         chain = await make_retrieval_chain(retriever)
# # # #         if chain is None:
# # # #             await message.reply("Failed to create the retrieval chain.")
# # # #             return
        
# # # #         try:     
# # # #             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
           
# # # #             response = chain.invoke({"input": query})
# # # #             print(response['answer'])
# # # #             global chat_history
# # # #             chat_history = response['answer'] 
# # # #             print(f"\n Chat History : {chat_history}")
# # # #             await message.reply(response['answer'])

# # # #             try:
# # # #                 graph_query = (
# # # #                     summary + "\n" + investment_personality + "\n" + response['answer'] + "\n" +  
# # # #                     "Please provide the following information in JSON format for the given response :\n" +
# # # #                     "{\n"
# # # #                     "  'growth_investment_min': <minimum percentage of growth-oriented investments>,\n"
# # # #                     "  'growth_investment_max': <maximum percentage of growth-oriented investments>,\n"
# # # #                     "  'return_min': <minimum expected annual return percentage>,\n"
# # # #                     "  'return_max': <maximum expected annual return percentage>,\n"
# # # #                     "  'growth_min': <minimum expected growth in dollars>,\n"
# # # #                     "  'growth_max': <maximum expected growth in dollars>\n"
# # # #                     "}"
# # # #                 )
# # # #                 graph_response = chain.invoke({"input": graph_query})
# # # #                 print(graph_response['answer'])
# # # #                 await handle_graph(graph_response['answer'],chat_id)
# # # #             except Exception as e:
# # # #                 print(f"Error plotting graph : {e}")
# # # #         except Exception as e:
# # # #             print(f"Error invoking retrieval chain on attempt : {e}")
   
# # # #     else:
# # # #         await message.reply("Failed to process the uploaded file.")





# # # # import json
# # # # # Function to send document with retry mechanism 
# # # # import re
# # # # async def send_document_with_retry(chat_id, document, caption, retries=3):
# # # #     for attempt in range(retries):
# # # #         try:
# # # #             await bot.send_document(chat_id, document, caption=caption)
# # # #             return
# # # #         except aiohttp.ClientConnectionError as e:
# # # #             logging.error(f"Attempt {attempt + 1} failed: {e}")
# # # #             if attempt < retries - 1:
# # # #                 continue
# # # #             else:
# # # #                 raise e
            
# # # # # Function to extract data from LLM response
# # # # def extract_data_from_response(response):
# # # #     try:
# # # #         # Locate the JSON-like data in the response
# # # #         json_data = response.split("{", 1)[1].rsplit("}", 1)[0]
# # # #         json_data = "{" + json_data + "}"
        
# # # #         # Parse the JSON data
# # # #         data = json.loads(json_data.replace("'", "\""))
# # # #         return data
# # # #     except Exception as e:
# # # #         logging.error(f"Error extracting data: {e}")
# # # #         return None

# # # # # Generate sample plots and send them to the user
# # # # async def handle_graph(response, chat_id):
# # # #     try:
# # # #         data = extract_data_from_response(response)
# # # #         if not data:
# # # #             await bot.send_message(chat_id, "Failed to extract data from the response.")
# # # #             return
        
# # # #         # Create a pie chart for investment allocation
# # # #         labels = 'Growth-Oriented Investments', 'Conservative Investments'
# # # #         sizes = [data['growth_investment_max'], 100 - data['growth_investment_max']]
# # # #         fig1, ax1 = plt.subplots()
# # # #         ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# # # #         ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
# # # #         pie_chart_buffer = io.BytesIO()
# # # #         plt.savefig(pie_chart_buffer, format='png')
# # # #         pie_chart_buffer.seek(0)
# # # #         pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

# # # #         # Create a bar graph for potential returns
# # # #         return_labels = ['Min Return', 'Max Return']
# # # #         return_values = [data['return_min'], data['return_max']]
# # # #         fig2, ax2 = plt.subplots()
# # # #         sns.barplot(x=return_labels, y=return_values, ax=ax2)
# # # #         ax2.set(xlabel='Return Type', ylabel='Percentage')
# # # #         bar_chart_buffer = io.BytesIO()
# # # #         plt.savefig(bar_chart_buffer, format='png')
# # # #         bar_chart_buffer.seek(0)
# # # #         bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='bar_chart.png')

# # # #         # Send the generated graphs to the user with retry mechanism
# # # #         await send_document_with_retry(chat_id, pie_chart_file, caption="Investment Allocation Chart")
# # # #         await send_document_with_retry(chat_id, bar_chart_file, caption="Potential Returns Chart")

# # # #         # Generate the investment growth graph
# # # #         fig3, ax3 = plt.subplots()
# # # #         growth_labels = ['Min Growth', 'Max Growth']
# # # #         growth_values = [data['growth_min'], data['growth_max']]
# # # #         ax3.bar(growth_labels, growth_values)
# # # #         ax3.set(title='Investment Growth Over Time', xlabel='Growth Type', ylabel='Value ($)')
# # # #         ax3.grid(True)
# # # #         investment_buffer = io.BytesIO()
# # # #         plt.savefig(investment_buffer, format='png')
# # # #         investment_buffer.seek(0)
# # # #         investment_file = BufferedInputFile(investment_buffer.read(), filename='investment_growth.png')

# # # #         await send_document_with_retry(chat_id, investment_file, caption="Investment Growth Over Time")

# # # #     except Exception as e:
# # # #         logging.error(f"Error plotting graph: {e}")
# # # #         await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")


# # # # # This approach failed 

# # # # # Function to send document with retry mechanism 
# # # # # import re
# # # # # async def send_document_with_retry(chat_id, document, caption, retries=3):
# # # # #     for attempt in range(retries):
# # # # #         try:
# # # # #             await bot.send_document(chat_id, document, caption=caption)
# # # # #             return
# # # # #         except aiohttp.ClientConnectionError as e:
# # # # #             logging.error(f"Attempt {attempt + 1} failed: {e}")
# # # # #             if attempt < retries - 1:
# # # # #                 continue
# # # # #             else:
# # # # #                 raise e

# # # # # # Function to extract data from LLM response
# # # # # def extract_data_from_response(response):
# # # # #     data = {}

# # # # #     try:
# # # # #         # Extract investment allocation
# # # # #         match = re.search(r'(\d+)-(\d+)% to growth-oriented investments', response)
# # # # #         if match:
# # # # #             data['growth_investment_min'] = int(match.group(1))
# # # # #             data['growth_investment_max'] = int(match.group(2))
        
# # # # #         # Extract potential returns
# # # # #         match = re.search(r'range of (\d+)-(\d+)% per year', response)
# # # # #         if match:
# # # # #             data['return_min'] = int(match.group(1))
# # # # #             data['return_max'] = int(match.group(2))
        
# # # # #         # Extract expected growth
# # # # #         match = re.search(r'grow to approximately \$(\d+)-\$(\d+)', response)
# # # # #         if match:
# # # # #             data['growth_min'] = int(match.group(1))
# # # # #             data['growth_max'] = int(match.group(2))
        
# # # # #         if not data:
# # # # #             raise ValueError("No data extracted from the response.")
# # # # #     except Exception as e:
# # # # #         logging.error(f"Error extracting data: {e}")
# # # # #         return None

# # # # #     return data

# # # # # # Generate sample plots and send them to the user
# # # # # async def handle_graph(response, chat_id):
# # # # #     try:
# # # # #         data = extract_data_from_response(response)
# # # # #         if not data:
# # # # #             await bot.send_message(chat_id, "Failed to extract data from the response.")
# # # # #             return
        
# # # # #         # Create a pie chart for investment allocation
# # # # #         labels = 'Growth-Oriented Investments', 'Conservative Investments'
# # # # #         sizes = [data['growth_investment_max'], 100 - data['growth_investment_max']]
# # # # #         fig1, ax1 = plt.subplots()
# # # # #         ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# # # # #         ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
# # # # #         pie_chart_buffer = io.BytesIO()
# # # # #         plt.savefig(pie_chart_buffer, format='png')
# # # # #         pie_chart_buffer.seek(0)
# # # # #         pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

# # # # #         # Create a bar graph for potential returns
# # # # #         return_labels = ['Min Return', 'Max Return']
# # # # #         return_values = [data['return_min'], data['return_max']]
# # # # #         fig2, ax2 = plt.subplots()
# # # # #         sns.barplot(x=return_labels, y=return_values, ax=ax2)
# # # # #         ax2.set(xlabel='Return Type', ylabel='Percentage')
# # # # #         bar_chart_buffer = io.BytesIO()
# # # # #         plt.savefig(bar_chart_buffer, format='png')
# # # # #         bar_chart_buffer.seek(0)
# # # # #         bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='bar_chart.png')

# # # # #         # Send the generated graphs to the user with retry mechanism
# # # # #         await send_document_with_retry(chat_id, pie_chart_file, caption="Investment Allocation Chart")
# # # # #         await send_document_with_retry(chat_id, bar_chart_file, caption="Potential Returns Chart")

# # # # #         # Generate the investment growth graph
# # # # #         fig3, ax3 = plt.subplots()
# # # # #         growth_labels = ['Min Growth', 'Max Growth']
# # # # #         growth_values = [data['growth_min'], data['growth_max']]
# # # # #         ax3.bar(growth_labels, growth_values)
# # # # #         ax3.set(title='Investment Growth Over Time', xlabel='Growth Type', ylabel='Value ($)')
# # # # #         ax3.grid(True)
# # # # #         investment_buffer = io.BytesIO()
# # # # #         plt.savefig(investment_buffer, format='png')
# # # # #         investment_buffer.seek(0)
# # # # #         investment_file = BufferedInputFile(investment_buffer.read(), filename='investment_growth.png')

# # # # #         await send_document_with_retry(chat_id, investment_file, caption="Investment Growth Over Time")

# # # # #     except Exception as e:
# # # # #         logging.error(f"Error plotting graph: {e}")
# # # # #         await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")





# # # # # Generate sample plots and send them to the user able to generate predetermined plots
# # # # import aiohttp
# # # # from aiogram.types import InputFile , BufferedInputFile
# # # # # # Function to send document with retry mechanism
# # # # # async def send_document_with_retry(chat_id, document, caption, retries=3):
# # # # #     for attempt in range(retries):
# # # # #         try:
# # # # #             await bot.send_document(chat_id, document, caption=caption)
# # # # #             return
# # # # #         except aiohttp.ClientConnectionError as e:
# # # # #             logging.error(f"Attempt {attempt + 1} failed: {e}")
# # # # #             if attempt < retries - 1:
# # # # #                 continue
# # # # #             else:
# # # # #                 raise e

# # # # # # Generate sample plots and send them to the user
# # # # # async def handle_graph(response, chat_id):
# # # # #     try:
# # # # #         # Example data (replace this with data from your LLM RAG chain)
# # # # #         labels = 'Category A', 'Category B', 'Category C', 'Category D'
# # # # #         sizes = [15, 30, 45, 10]
# # # # #         bar_labels = ['Q1', 'Q2', 'Q3', 'Q4']
# # # # #         bar_values = [20, 35, 30, 35]

# # # # #         # Create a pie chart
# # # # #         fig1, ax1 = plt.subplots()
# # # # #         ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
# # # # #         ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
# # # # #         pie_chart_buffer = io.BytesIO()
# # # # #         plt.savefig(pie_chart_buffer, format='png')
# # # # #         pie_chart_buffer.seek(0)
# # # # #         pie_chart_file = BufferedInputFile(pie_chart_buffer.read(), filename='pie_chart.png')

# # # # #         # Create a bar graph
# # # # #         fig2, ax2 = plt.subplots()
# # # # #         sns.barplot(x=bar_labels, y=bar_values, ax=ax2)
# # # # #         ax2.set(xlabel='Quarters', ylabel='Values')
# # # # #         bar_chart_buffer = io.BytesIO()
# # # # #         plt.savefig(bar_chart_buffer, format='png')
# # # # #         bar_chart_buffer.seek(0)
# # # # #         bar_chart_file = BufferedInputFile(bar_chart_buffer.read(), filename='bar_chart.png')

# # # # #         # Send the generated graphs to the user with retry mechanism
# # # # #         await send_document_with_retry(chat_id, pie_chart_file, caption="Here is the pie chart.")
# # # # #         await send_document_with_retry(chat_id, bar_chart_file, caption="Here is the bar graph.")

# # # # #         # Investment Growth Over Time Example
# # # # #         investment_amount = 3000  # Monthly investment
# # # # #         investment_horizon = 3  # Years
# # # # #         annualized_return = 0.06  # 6%
# # # # #         months = investment_horizon * 12
# # # # #         investment_values = [0]
# # # # #         for month in range(1, months + 1):
# # # # #             investment_values.append(investment_values[-1] * (1 + annualized_return / 12) + investment_amount)

# # # # #         # Generate the investment growth graph
# # # # #         fig3, ax3 = plt.subplots()
# # # # #         ax3.plot(range(0, months + 1), investment_values, marker='o')
# # # # #         ax3.set(title='Investment Growth Over Time', xlabel='Months', ylabel='Investment Value ($)')
# # # # #         ax3.grid(True)
# # # # #         investment_buffer = io.BytesIO()
# # # # #         plt.savefig(investment_buffer, format='png')
# # # # #         investment_buffer.seek(0)
# # # # #         investment_file = BufferedInputFile(investment_buffer.read(), filename='investment_growth.png')

# # # # #         # Expected Annualized Returns Example
# # # # #         expected_returns = [0.05, 0.06, 0.07, 0.08]  # Expected annualized returns in range 5-8%
# # # # #         return_labels = ['5%', '6%', '7%', '8%']
# # # # #         fig4, ax4 = plt.subplots()
# # # # #         ax4.bar(return_labels, expected_returns)
# # # # #         ax4.set(title='Expected Annualized Returns', xlabel='Return Percentage', ylabel='Annualized Return')
# # # # #         ax4.grid(True)
# # # # #         returns_buffer = io.BytesIO()
# # # # #         plt.savefig(returns_buffer, format='png')
# # # # #         returns_buffer.seek(0)
# # # # #         returns_file = BufferedInputFile(returns_buffer.read(), filename='expected_returns.png')

# # # # #         # Send the investment and returns graphs to the user with retry mechanism
# # # # #         await send_document_with_retry(chat_id, investment_file, caption="Investment Growth Over Time")
# # # # #         await send_document_with_retry(chat_id, returns_file, caption="Expected Annualized Returns")

# # # # #     except Exception as e:
# # # # #         logging.error(f"Error plotting graph: {e}")
# # # # #         await bot.send_message(chat_id, "There was an error generating the graphs. Please try again later.")


# # # # async def process_document(file_path):
# # # #     try:
# # # #         file_type = filetype.guess(file_path)
# # # #         if file_type is not None:
# # # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # #                 # return extract_text_from_word(file_path)
# # # #                 return extract_text_and_tables_from_word(file_path)
# # # #             elif file_type.mime == "application/pdf":
# # # #                 return extract_text_from_pdf(file_path)
# # # #         return None
# # # #     except Exception as e:
# # # #         print(f"Error processing document: {e}")
# # # #         return None

# # # # def extract_text_from_pdf(pdf_file_path):
# # # #     try:
# # # #         with open(pdf_file_path, "rb") as pdf_file:
# # # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # # #             text_content = []
# # # #             for page_num in range(pdf_reader.numPages):
# # # #                 page = pdf_reader.getPage(page_num)
# # # #                 text_content.append(page.extract_text())
# # # #             return "\n".join(text_content)
# # # #     except Exception as e:
# # # #         print(f"Error extracting text from PDF: {e}")
# # # #         return None


# # # # import re
# # # # import docx

# # # # def extract_text_and_tables_from_word(docx_file_path):
# # # #     """
# # # #     Extracts text and tables from a Word document (.docx).

# # # #     Args:
# # # #         docx_file_path (str): Path to the Word document file.

# # # #     Returns:
# # # #         tuple: Extracted text content and tables from the document.
# # # #     """
# # # #     try:
# # # #         doc = docx.Document(docx_file_path)
# # # #         text_content = []
# # # #         tables_content = []

# # # #         for para in doc.paragraphs:
# # # #             text_content.append(para.text)

# # # #         for table in doc.tables:
# # # #             table_data = []
# # # #             for row in table.rows:
# # # #                 row_data = []
# # # #                 for cell in row.cells:
# # # #                     row_data.append(cell.text.strip())
# # # #                 table_data.append(row_data)
# # # #             tables_content.append(table_data)

# # # #         return "\n".join(text_content), tables_content
# # # #     except Exception as e:
# # # #         print(f"Error extracting text and tables from Word document: {e}")
# # # #         return None, None

# # # # def validate_document_content(text, tables):
# # # #     """
# # # #     Validates the content of the document.

# # # #     Args:
# # # #         text (str): Extracted text content from the document.
# # # #         tables (list): Extracted tables content from the document.

# # # #     Returns:
# # # #         tuple: Client name and validation errors.
# # # #     """
# # # #     errors = []
    
# # # #     # Extract client name
# # # #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# # # #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# # # #     # Define required sections
# # # #     required_sections = [
# # # #         "YOUR RETIREMENT GOAL",
# # # #         "YOUR OTHER MAJOR GOALS",
# # # #         "YOUR ASSETS AND LIABILITIES",
# # # #         "MY LIABILITIES",
# # # #         "YOUR CURRENT ANNUAL INCOME"
# # # #     ]

# # # #     # Check for the presence of required sections
# # # #     for section in required_sections:
# # # #         if section not in text:
# # # #             errors.append(f"* {section} section missing.")
    
# # # #     # Define table field checks
# # # #     table_checks = {
# # # #         "YOUR RETIREMENT GOAL": [
# # # #             r"When do you plan to retire\? \(age or date\)",
# # # #             r"Social Security Benefit \(include expected start date\)",
# # # #             r"Pension Benefit \(include expected start date\)",
# # # #             r"Other Expected Income \(rental, part-time work, etc.\)",
# # # #             r"Estimated Annual Retirement Expense"
# # # #         ],
# # # #         "YOUR OTHER MAJOR GOALS": [
# # # #             r"GOAL", r"COST", r"WHEN"
# # # #         ],
# # # #         "YOUR ASSETS AND LIABILITIES": [
# # # #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# # # #             r"Current Value", r"Annual Contributions"
# # # #         ],
# # # #         "MY LIABILITIES": [
# # # #             r"Balance", r"Interest Rate", r"Monthly Payment"
# # # #         ]
# # # #     }

# # # #     # Validate table content
# # # #     for section, checks in table_checks.items():
# # # #         section_found = False
# # # #         for table in tables:
# # # #             table_text = "\n".join(["\t".join(row) for row in table])
# # # #             if section in table_text:
# # # #                 section_found = True
# # # #                 for check in checks:
# # # #                     if not re.search(check, table_text, re.IGNORECASE):
# # # #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# # # #                 break
# # # #         if not section_found:
# # # #             errors.append(f"* {section} section missing.")

# # # #     return client_name, errors

# # # # async def process_document(file_path):
# # # #     try:
# # # #         text, tables = extract_text_and_tables_from_word(file_path)
# # # #         if text is not None and tables is not None:
# # # #             client_name, errors = validate_document_content(text, tables)
# # # #             return client_name, errors
# # # #         return None, ["Error processing document."]
# # # #     except Exception as e:
# # # #         print(f"Error processing document: {e}")
# # # #         return None, [f"Error processing document: {e}"]



# # # # async def load_vector_db(file_path):
# # # #     try:
# # # #         loader = Docx2txtLoader(file_path)
# # # #         documents = loader.load()
# # # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # # #         text_chunks = text_splitter.split_documents(documents)
# # # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # # #         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
# # # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # # #     except Exception as e:
# # # #         print(f"Error loading vector database: {e}")
# # # #         return None


# # # # async def make_retrieval_chain(retriever):
# # # #     """
# # # #     Create a retrieval chain using the provided retriever.

# # # #     Args:
# # # #         retriever (RetrievalQA): A retriever object.

# # # #     Returns:
# # # #         RetrievalQA: A retrieval chain object.
# # # #     """
# # # #     try:
# # # #         global investment_personality,summary
# # # #         llm = ChatGoogleGenerativeAI(
# # # #             model="gemini-pro",
# # # #             temperature=0.7,
# # # #             top_p=0.85,
# # # #             google_api_key=GOOGLE_API_KEY
# # # #         )

# # # #         prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # #                 Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # #                 should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # #                 investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # #                 investment.Also explain the user why you are giving them that particular
# # # #                 investment suggestion.Answer in 3-4 lines.\n
# # # #                 <context>
# # # #                 {context}
# # # #                 </context>
# # # #                 Question: {input}"""

# # # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
# # # #         # Update combine_docs_chain with your actual document combining logic
# # # #         combine_docs_chain = None  # Replace this with your combine_docs_chain

# # # #         if retriever is not None :  #and combine_docs_chain is not None:
# # # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # # #             print(retriever_chain)
# # # #             return retriever_chain
# # # #         else:
# # # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # # #             return None

# # # #     except Exception as e:
# # # #         print(f"Error in creating chain: {e}")
# # # #         return None


# # # # # @dispatcher.message_handler()
# # # # @dp.message()
# # # # async def main_bot(message: types.Message):
# # # #     global retriever, extracted_text,investment_personality,summary,chat_history

# # # #     # Handle the first tasks assessments answers from the user
# # # #     chat_id = message.chat.id

# # # #     if chat_id in states and states[chat_id] < len(questions):
# # # #         # Retrieve the index of the current question
# # # #         question_index = states[chat_id]

# # # #         # Save the user's response to the current question
# # # #         answer = message.text
# # # #         user_responses[questions[question_index]] = answer
# # # #         states[chat_id] += 1  # Move to the next question

# # # #         # Ask the next question
# # # #         await ask_next_question(chat_id, question_index + 1)
# # # #     else:
# # # #         # Handle q&a chat messages using your Gemini model (llm)
# # # #         try:

# # # #             task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # #             Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # #             should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # #             investment.Also explain the user why you are giving them that particular
# # # #             investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # #             User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # # #             So dont mention user to consult to a financial expert."""
        
# # # #             model = genai.GenerativeModel('gemini-pro')
# # # #             print(investment_personality)
# # # #             # query = task + "\n" + investment_personality + "\n" + summary + "\n" +  extracted_text + "\n"  +   message.text

# # # #             # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +  extracted_text + "\n"  +   message.text
# # # #             query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
# # # #             print(f"\nQuery : {query}")
# # # #             response = model.generate_content(query)
# # # #             await message.reply(response.text) #(response['answer']) 
           
# # # #         except Exception as e:
# # # #             print(f"Error processing general chat message: {e}")
# # # #             await message.reply("Failed to process your request.")
        


# # # # # if __name__ == "__main__":
# # # # #     executor.start_polling(dispatcher, skip_updates=True)

# # # # async def main() -> None:
# # # #     # Initialize Bot instance with default bot properties which will be passed to all API calls
# # # #     bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

# # # #     # And the run events dispatching
# # # #     await dp.start_polling(bot)


# # # # if __name__ == "__main__":
# # # #     logging.basicConfig(level=logging.INFO, stream=sys.stdout)
# # # #     asyncio.run(main())




# # # # #best code so far now it can upload files and receive various forms of messages as well 

# # # # import os
# # # # import filetype
# # # # import docx
# # # # import PyPDF2
# # # # import re
# # # # from aiogram import Bot, Dispatcher, types
# # # # from dotenv import load_dotenv
# # # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # # from langchain_community.vectorstores import Chroma
# # # # from langchain_community.document_loaders import Docx2txtLoader
# # # # from langchain_core.prompts import ChatPromptTemplate
# # # # from langchain.chains import create_retrieval_chain
# # # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # # from langchain.memory import ConversationSummaryMemory
# # # # import asyncio

# # # # import google.generativeai as genai

# # # # # Import things that are needed generically
# # # # from langchain.pydantic_v1 import BaseModel, Field
# # # # from langchain.tools import BaseTool, StructuredTool, tool

# # # # from aiogram.client.default import DefaultBotProperties
# # # # from aiogram.enums import ParseMode
# # # # from aiogram.filters import CommandStart
# # # # from aiogram.types import Message
# # # # from aiogram import F
# # # # from aiogram import Router
# # # # import logging
# # # # import sys


# # # # router = Router(name=__name__)

# # # # load_dotenv()

# # # # TOKEN = os.getenv("TOKEN")
# # # # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # # # os.environ["LANGCHAIN_TRACING_V2"]="true"
# # # # os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # # # # Configure generativeai with your API key
# # # # genai.configure(api_key=GOOGLE_API_KEY)

# # # # # Initialize bot
# # # # bot = Bot(token=TOKEN)
# # # # dp = Dispatcher()

# # # # rag_on = False
# # # # retriever = None  # Store retriever globally
# # # # summary = ""
# # # # investment_personality = ""
# # # # chat_history = ""

# # # # class Reference:
# # # #     def __init__(self):
# # # #         self.response = ""


# # # # reference = Reference()


# # # # def clear_past():
# # # #     reference.response = ""


# # # # @router.message(F.text == "clear")
# # # # async def clear(message: types.Message):
# # # #     """
# # # #     A handler to clear the previous conversation and context.
# # # #     """
# # # #     clear_past()
# # # #     await message.reply("I've cleared the past conversation and context.")


# # # # # Store user states
# # # # states = {}
# # # # # Dictionary to hold question-answer pairs
# # # # user_responses = {}

# # # # # Define Questions for assessment
# # # # questions = [
# # # #     """
# # # #         1. Singapore plans to build a new observation tower called 'The Rook'.
# # # #         How many steps do you think it will take to go to the top floor?

# # # #         a) Less than 500 
# # # #         b) More than 500

# # # #     """,
# # # #     "2. Now Guess the number of steps" ,
# # # #     """
# # # #     3. How confident are you that the real number is in the range you have selected? 
# # # #     Answer within a range of 100.  
# # # #     """,
# # # #     """ 
# # # #     4. You and your friend are betting on a series of coin tosses.

# # # #     He always bets ₹2,000 on Heads

# # # #     You always bet ₹2,000 on Tails

# # # #     Winner of last 8 turns

# # # #     You lost ₹8,000 in the last 4 turns!

# # # #     If you were to bet one last time, what would you bet on heads or tails ?
# # # #     """ ,
# # # #     """
# # # #     5. How confident are you that your bet will win this time?
# # # #     Answer how confident you are. 
# # # #     (Example: Not confident at all, somewhat confident, confident, or Very confident)
# # # #     """,
# # # #     """
# # # #     6. What do you think are the chances that you will achieve and maintain your financial goals within the next 
# # # #     10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
# # # #     Answer how likely you are to achieve your goal.
# # # #     (Example: Less likely than others, likely than others, or More likely than others)
# # # #     """,
# # # #     """
# # # #     7. Imagine you are a contestant in a game show, and you are presented the following choices.

# # # #     What would you prefer?
# # # #     a) 50 percent chance of winning 15 gold coins 
# # # #     b) 100 percent chance of winning 8 gold coins
# # # #     """,
# # # #     """
# # # #     8. Ok, one last choice...

# # # #     What would you prefer?
# # # #     a) 50 percent chance of winning 15 gold coins 
# # # #     b) 100 percent chance of winning 2 gold coins
# # # #     """,
# # # #     """
# # # #     9. In general, how would your best friend describe your risk-taking tendencies?
# # # #     a) A real gambler
# # # #     b) Willing to take risks after completing adequate research
# # # #     c) Cautious
# # # #     d) Avoids risk as much as possible
# # # #     """,
# # # #     """
# # # #     10. Suppose you could replace your current investment portfolio with this new one:
# # # #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# # # #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# # # #     Answer between the range of -5 to -35.
# # # #     """,
# # # #     """
# # # #     11. Suppose that in the next 7 years,

# # # #     YOUR INCOME

# # # #     grows 8% each year

# # # #     VS
# # # #     INFLATION

# # # #     grows 10% a year

# # # #     At the end of 7 years, how much will you be able to buy with your income?
# # # #     Options:
# # # #     a) More than today
# # # #     b) Exactly the same
# # # #     c) Less than today
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     12. If somebody buys a bond of Company B, which of the following statements seems correct:
# # # #     a) She owns part of Company B
# # # #     b) She has lent money to Company B
# # # #     c) She is liable for Company B's debt
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     13. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
# # # #     a) More than ₹2 lakhs
# # # #     b) Less than ₹2 lakhs
# # # #     c) Exactly ₹2 lakhs
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     14. When an investor spreads money across different asset classes, what happens to the risk of losing money:
# # # #     a) Increases
# # # #     b) Decreases
# # # #     c) Stays the same
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     15. When a country's central bank reduces interest rates, it makes:

# # # #     a) Borrowing more attractive and saving less attractive
# # # #     b) Borrowing less attractive and saving more attractive
# # # #     c) Both borrowing and saving less attractive
# # # #     d) Cannot say
# # # #     """,
# # # #     """
# # # #     16. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
# # # #     a) Option A : Min Value is 9.2L and Max Value is 11.8L
# # # #     b) Option B : Min Value is 8.8L and Max Value is 12.3L
# # # #     c) Option C : Min Value is 8.5L and Max Value is 12.8L
# # # #     d) Option D : Min Value is 8.1L and Max Value is 13.3L
# # # #     e) Option E : Min Value is 7.8L and Max Value is 13.8L
# # # #     """,
# # # #     """
# # # #     17. From Sept 2008 to Nov 2008, Stock market went down by 31%.

# # # #     If you owned a stock investment that lost about 31 percent in 3 months, you would:
# # # #     a) Sell all of the remaining investment
# # # #     b) Sell a portion of the remaining investment
# # # #     c) Hold on to the investment and sell nothing
# # # #     d) Buy little
# # # #     e) Buy more of the investment
# # # #     """,
# # # #     """
# # # #     18. Over any 1-year period, what would be the maximum drop in the value of your investment 
# # # #     portfolio that you would be comfortable with?
# # # #     a) <5%
# # # #     b) 5 - 10%
# # # #     c) 10 - 15%
# # # #     d) 15 - 20%
# # # #     e) >20%
# # # #     """,
# # # #     """
# # # #     19. When investing, what do you consider the most?

# # # #     a) Risk 
# # # #     b) Return
# # # #     """,
# # # #     """
# # # #     20. What best describes your attitude?

# # # #     a) Prefer reasonable returns, can take reasonable risk
# # # #     b) Like higher returns, can take slightly higher risk
# # # #     c) Want to maximize returns, can take significant high risk
# # # #     """,
# # # #     """
# # # #     21. How much monthly investment you want to do?
# # # #     """,
# # # #     """
# # # #     22. What is the time horizon for your investment?
# # # #     You can answer in any range, example 1-5 years."""  
# # # # ]


# # # # # Handler for /start command
# # # # @dp.message(CommandStart())
# # # # async def handle_start(message: types.Message):
# # # #     """
# # # #     This handler receives messages with /start command
# # # #     """
# # # #     chat_id = message.chat.id
# # # #     # Start asking questions
# # # #     await start_assessment(chat_id)


# # # # # Function to start the assessment
# # # # async def start_assessment(chat_id):
# # # #     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
# # # #     await ask_next_question(chat_id, 0)

# # # # # Function to ask the next question
# # # # async def ask_next_question(chat_id, question_index):
# # # #     if question_index < len(questions):
# # # #         # Ask the next question
# # # #         await bot.send_message(chat_id, questions[question_index])
# # # #         # Update state to indicate the next expected answer
# # # #         states[chat_id] = question_index
# # # #     else:
# # # #         # No more questions, finish assessment
# # # #         await finish_assessment(chat_id)

# # # # # Handler for receiving assessment answers
# # # # assessment_in_progress = True

# # # # from aiogram.types import FSInputFile
# # # # async def finish_assessment(chat_id):
# # # #     if chat_id in states and states[chat_id] == len(questions):
# # # #         # All questions have been answered, now process the assessment
# # # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # # #         # Determine investment personality based on collected responses
# # # #         global investment_personality
# # # #         investment_personality = await determine_investment_personality(user_responses)

# # # #         # Inform the user about their investment personality
# # # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # # #         # Summarize collected information
# # # #         global summary
# # # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # # #         # Ensure to await the determination of investment personality
# # # #         await send_summary_chunks(chat_id, summary)
# # # #         global assessment_in_progress 
# # # #         assessment_in_progress = False
# # # #         # Prompt the user to begin financial advice process
# # # #         # await bot.send_message(chat_id, "To receive financial advice, please upload a document with your financial details using /begin.")

# # # #         await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
# # # #         file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

# # # #         await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
# # # #         # await bot.send_message(chat_id,file)

# # # # async def send_summary_chunks(chat_id, summary):
# # # #     # Split the summary into chunks that fit within Telegram's message limits
# # # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# # # #     # Send each chunk as a separate message
# # # #     for chunk in chunks:
# # # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # # # async def determine_investment_personality(assessment_data):
# # # #     try:
# # # #         # Prepare input text for the chatbot based on assessment data
# # # #         input_text = "User Profile:\n"
# # # #         for question, answer in assessment_data.items():
# # # #             input_text += f"{question}: {answer}\n"

# # # #         # Introduce the chatbot's task and prompt for classification
# # # #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# # # #                       "- Conservative Investor\n" \
# # # #                       "- Moderate Investor\n" \
# # # #                       "- Aggressive Investor"

# # # #         # Use your generative AI model to generate a response
# # # #         # print(input_text)
# # # #         model = genai.GenerativeModel('gemini-pro')
# # # #         response = model.generate_content(input_text)

# # # #         # Determine the investment personality from the chatbot's response
# # # #         response_text = response.text.lower()
# # # #         if "conservative" in response_text:
# # # #             personality = "Conservative Investor"
# # # #         elif "moderate" in response_text:
# # # #             personality = "Moderate Investor"
# # # #         elif "aggressive" in response_text:
# # # #             personality = "Aggressive Investor"
# # # #         else:
# # # #             personality = "Unknown"

# # # #         return personality
# # # #         # Send the determined investment personality back to the user
# # # #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# # # #     except Exception as e:
# # # #         print(f"Error generating response: {e}")
# # # #         #await bot.send_message(chat_id, "Error processing investment personality classification.")




# # # # @router.message(F.text == "help")
# # # # async def helper(message: types.Message):
# # # #     """
# # # #     A handler to display the help menu.
# # # #     """
# # # #     help_command = """
# # # #     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
# # # #     /start - to start the investment personality assessment.
# # # #     /clear - to clear the past conversation and context.
# # # #     /help - to get this help menu.
# # # #     /begin - to start the Financial Suggestion Conversation with the ChatBot.
# # # #     I hope this helps. :)
# # # #     """
# # # #     await message.reply(help_command)

# # # # # Handler for /begin command to initiate financial advice

# # # # @router.message(F.text == "begin")
# # # # async def handle_begin(message: types.Message):
# # # #     chat_id = message.chat.id
# # # #     file_instructions ="""Hello there!My name is Finbot and I am a Wealth Management Advisor Chatbot.I need more details related to your Financial Profile so that I can give you 
# # # #     personalised Financial Advice.Please follow this drive link : https://docs.google.com/document/d/1JaqWUXcq3MNrPTCvEy_2RngbGIY0rybX/edit?usp=sharing&ouid=101252223236130106095&rtpof=true&sd=true
# # # #     ,there you will find a Word Document.Please fill in your details with correct Information and then upload the Document in the chat. """

# # # #     await message.reply(file_instructions)
    
    

# # # # # Handler for document upload

# # # # # @dispatcher.message_handler(content_types=['document'])

# # # # # @router.message(F.document)
# # # # @dp.message(F.document)
# # # # async def handle_document(message: types.Message):
# # # #     global summary,investment_personality  
    
# # # #     # Obtain file information
# # # #     file_id = message.document.file_id
# # # #     file = await bot.get_file(file_id)
# # # #     file_path = file.file_path
    
# # # #     # Download the file
# # # #     await bot.download_file(file_path, "data/uploaded_file")
    
# # # #     # Process the uploaded document
# # # #     extracted_text = await process_document("data/uploaded_file")
    
# # # #     if extracted_text:
# # # #         # Load vector database (assuming this is part of setting up the retriever)
# # # #         retriever = await load_vector_db("data/uploaded_file")
# # # #         file_path = 'data/uploaded_file'
# # # #         client_name, validation_errors = await process_document(file_path)

# # # #         # Print results
# # # #         print(f"Client Name: {client_name}")
# # # #         if validation_errors:
# # # #             print("**Validation Errors:**")
# # # #             for error in validation_errors:
# # # #                 print(error)
# # # #         else:
# # # #             print("All fields are filled correctly.")

# # # #         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")
 
# # # #         if retriever is None:
# # # #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # #             return

# # # #         # Check if a valid chain can be created
# # # #         chain = await make_retrieval_chain(retriever)
# # # #         if chain is None:
# # # #             await message.reply("Failed to create the retrieval chain.")
# # # #             return
        
# # # #         try:     
# # # #             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
# # # #             response = chain.invoke({"input": query})
# # # #             print(response['answer'])
# # # #             global chat_history
# # # #             chat_history = response['answer'] 
# # # #             print(f"\n Chat History : {chat_history}")
# # # #             await message.reply(response['answer'])

# # # #         except Exception as e:
# # # #             print(f"Error invoking retrieval chain on attempt : {e}")
   
# # # #     else:
# # # #         await message.reply("Failed to process the uploaded file.")



# # # # async def process_document(file_path):
# # # #     try:
# # # #         file_type = filetype.guess(file_path)
# # # #         if file_type is not None:
# # # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # #                 # return extract_text_from_word(file_path)
# # # #                 return extract_text_and_tables_from_word(file_path)
# # # #             elif file_type.mime == "application/pdf":
# # # #                 return extract_text_from_pdf(file_path)
# # # #         return None
# # # #     except Exception as e:
# # # #         print(f"Error processing document: {e}")
# # # #         return None

# # # # def extract_text_from_pdf(pdf_file_path):
# # # #     try:
# # # #         with open(pdf_file_path, "rb") as pdf_file:
# # # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # # #             text_content = []
# # # #             for page_num in range(pdf_reader.numPages):
# # # #                 page = pdf_reader.getPage(page_num)
# # # #                 text_content.append(page.extract_text())
# # # #             return "\n".join(text_content)
# # # #     except Exception as e:
# # # #         print(f"Error extracting text from PDF: {e}")
# # # #         return None


# # # # import re
# # # # import docx

# # # # def extract_text_and_tables_from_word(docx_file_path):
# # # #     """
# # # #     Extracts text and tables from a Word document (.docx).

# # # #     Args:
# # # #         docx_file_path (str): Path to the Word document file.

# # # #     Returns:
# # # #         tuple: Extracted text content and tables from the document.
# # # #     """
# # # #     try:
# # # #         doc = docx.Document(docx_file_path)
# # # #         text_content = []
# # # #         tables_content = []

# # # #         for para in doc.paragraphs:
# # # #             text_content.append(para.text)

# # # #         for table in doc.tables:
# # # #             table_data = []
# # # #             for row in table.rows:
# # # #                 row_data = []
# # # #                 for cell in row.cells:
# # # #                     row_data.append(cell.text.strip())
# # # #                 table_data.append(row_data)
# # # #             tables_content.append(table_data)

# # # #         return "\n".join(text_content), tables_content
# # # #     except Exception as e:
# # # #         print(f"Error extracting text and tables from Word document: {e}")
# # # #         return None, None

# # # # def validate_document_content(text, tables):
# # # #     """
# # # #     Validates the content of the document.

# # # #     Args:
# # # #         text (str): Extracted text content from the document.
# # # #         tables (list): Extracted tables content from the document.

# # # #     Returns:
# # # #         tuple: Client name and validation errors.
# # # #     """
# # # #     errors = []
    
# # # #     # Extract client name
# # # #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# # # #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# # # #     # Define required sections
# # # #     required_sections = [
# # # #         "YOUR RETIREMENT GOAL",
# # # #         "YOUR OTHER MAJOR GOALS",
# # # #         "YOUR ASSETS AND LIABILITIES",
# # # #         "MY LIABILITIES",
# # # #         "YOUR CURRENT ANNUAL INCOME"
# # # #     ]

# # # #     # Check for the presence of required sections
# # # #     for section in required_sections:
# # # #         if section not in text:
# # # #             errors.append(f"* {section} section missing.")
    
# # # #     # Define table field checks
# # # #     table_checks = {
# # # #         "YOUR RETIREMENT GOAL": [
# # # #             r"When do you plan to retire\? \(age or date\)",
# # # #             r"Social Security Benefit \(include expected start date\)",
# # # #             r"Pension Benefit \(include expected start date\)",
# # # #             r"Other Expected Income \(rental, part-time work, etc.\)",
# # # #             r"Estimated Annual Retirement Expense"
# # # #         ],
# # # #         "YOUR OTHER MAJOR GOALS": [
# # # #             r"GOAL", r"COST", r"WHEN"
# # # #         ],
# # # #         "YOUR ASSETS AND LIABILITIES": [
# # # #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# # # #             r"Current Value", r"Annual Contributions"
# # # #         ],
# # # #         "MY LIABILITIES": [
# # # #             r"Balance", r"Interest Rate", r"Monthly Payment"
# # # #         ]
# # # #     }

# # # #     # Validate table content
# # # #     for section, checks in table_checks.items():
# # # #         section_found = False
# # # #         for table in tables:
# # # #             table_text = "\n".join(["\t".join(row) for row in table])
# # # #             if section in table_text:
# # # #                 section_found = True
# # # #                 for check in checks:
# # # #                     if not re.search(check, table_text, re.IGNORECASE):
# # # #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# # # #                 break
# # # #         if not section_found:
# # # #             errors.append(f"* {section} section missing.")

# # # #     return client_name, errors

# # # # async def process_document(file_path):
# # # #     try:
# # # #         text, tables = extract_text_and_tables_from_word(file_path)
# # # #         if text is not None and tables is not None:
# # # #             client_name, errors = validate_document_content(text, tables)
# # # #             return client_name, errors
# # # #         return None, ["Error processing document."]
# # # #     except Exception as e:
# # # #         print(f"Error processing document: {e}")
# # # #         return None, [f"Error processing document: {e}"]



# # # # async def load_vector_db(file_path):
# # # #     try:
# # # #         loader = Docx2txtLoader(file_path)
# # # #         documents = loader.load()
# # # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # # #         text_chunks = text_splitter.split_documents(documents)
# # # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # # #         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
# # # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # # #     except Exception as e:
# # # #         print(f"Error loading vector database: {e}")
# # # #         return None


# # # # async def make_retrieval_chain(retriever):
# # # #     """
# # # #     Create a retrieval chain using the provided retriever.

# # # #     Args:
# # # #         retriever (RetrievalQA): A retriever object.

# # # #     Returns:
# # # #         RetrievalQA: A retrieval chain object.
# # # #     """
# # # #     try:
# # # #         global investment_personality,summary
# # # #         llm = ChatGoogleGenerativeAI(
# # # #             model="gemini-pro",
# # # #             temperature=0.7,
# # # #             top_p=0.85,
# # # #             google_api_key=GOOGLE_API_KEY
# # # #         )

# # # #         prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # #                 Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # #                 should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # #                 investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # #                 investment.Also explain the user why you are giving them that particular
# # # #                 investment suggestion.Answer in 3-4 lines.\n
# # # #                 <context>
# # # #                 {context}
# # # #                 </context>
# # # #                 Question: {input}"""

# # # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
# # # #         # Update combine_docs_chain with your actual document combining logic
# # # #         combine_docs_chain = None  # Replace this with your combine_docs_chain

# # # #         if retriever is not None :  #and combine_docs_chain is not None:
# # # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # # #             print(retriever_chain)
# # # #             return retriever_chain
# # # #         else:
# # # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # # #             return None

# # # #     except Exception as e:
# # # #         print(f"Error in creating chain: {e}")
# # # #         return None


# # # # # @dispatcher.message_handler()
# # # # @dp.message()
# # # # async def main_bot(message: types.Message):
# # # #     global retriever, extracted_text,investment_personality,summary,chat_history

# # # #     # Handle the first tasks assessments answers from the user
# # # #     chat_id = message.chat.id

# # # #     if chat_id in states and states[chat_id] < len(questions):
# # # #         # Retrieve the index of the current question
# # # #         question_index = states[chat_id]

# # # #         # Save the user's response to the current question
# # # #         answer = message.text
# # # #         user_responses[questions[question_index]] = answer
# # # #         states[chat_id] += 1  # Move to the next question

# # # #         # Ask the next question
# # # #         await ask_next_question(chat_id, question_index + 1)
# # # #     else:
# # # #         # Handle q&a chat messages using your Gemini model (llm)
# # # #         try:

# # # #             task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # #             Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # #             should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # #             investment.Also explain the user why you are giving them that particular
# # # #             investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # #             User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # # #             So dont mention user to consult to a financial expert."""
        
# # # #             model = genai.GenerativeModel('gemini-pro')
# # # #             print(investment_personality)
# # # #             # query = task + "\n" + investment_personality + "\n" + summary + "\n" +  extracted_text + "\n"  +   message.text

# # # #             # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +  extracted_text + "\n"  +   message.text
# # # #             query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
# # # #             print(f"\nQuery : {query}")
# # # #             response = model.generate_content(query)
# # # #             await message.reply(response.text) #(response['answer']) 
           
# # # #         except Exception as e:
# # # #             print(f"Error processing general chat message: {e}")
# # # #             await message.reply("Failed to process your request.")
        


# # # # # if __name__ == "__main__":
# # # # #     executor.start_polling(dispatcher, skip_updates=True)

# # # # async def main() -> None:
# # # #     # Initialize Bot instance with default bot properties which will be passed to all API calls
# # # #     bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

# # # #     # And the run events dispatching
# # # #     await dp.start_polling(bot)


# # # # if __name__ == "__main__":
# # # #     logging.basicConfig(level=logging.INFO, stream=sys.stdout)
# # # #     asyncio.run(main())




# # # # # #best code so far correctly takes the assessment and gives financial advice along 
# # # # # #with returns with proper disclaimer and in brief

# # # # # import os
# # # # # import filetype
# # # # # import docx
# # # # # import PyPDF2
# # # # # import re
# # # # # from aiogram import Bot, Dispatcher, executor, types
# # # # # from dotenv import load_dotenv
# # # # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # # # from langchain_community.vectorstores import Chroma
# # # # # from langchain_community.document_loaders import Docx2txtLoader
# # # # # from langchain_core.prompts import ChatPromptTemplate
# # # # # from langchain.chains import create_retrieval_chain
# # # # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # # # from langchain.memory import ConversationSummaryMemory
# # # # # import asyncio
# # # # # from aiogram.utils.exceptions import NetworkError, RetryAfter, TelegramAPIError
# # # # # import google.generativeai as genai

# # # # # # Import things that are needed generically
# # # # # from langchain.pydantic_v1 import BaseModel, Field
# # # # # from langchain.tools import BaseTool, StructuredTool, tool

# # # # # load_dotenv()

# # # # # TOKEN = os.getenv("TOKEN")
# # # # # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # # # # os.environ["LANGCHAIN_TRACING_V2"]="true"
# # # # # os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # # # # # Configure generativeai with your API key
# # # # # genai.configure(api_key=GOOGLE_API_KEY)

# # # # # # Initialize bot
# # # # # bot = Bot(token=TOKEN)
# # # # # dispatcher = Dispatcher(bot)

# # # # # rag_on = False
# # # # # retriever = None  # Store retriever globally
# # # # # summary = ""
# # # # # investment_personality = ""
# # # # # chat_history = ""

# # # # # class Reference:
# # # # #     def __init__(self):
# # # # #         self.response = ""


# # # # # reference = Reference()


# # # # # def clear_past():
# # # # #     reference.response = ""


# # # # # @dispatcher.message_handler(commands=['clear'])
# # # # # async def clear(message: types.Message):
# # # # #     """
# # # # #     A handler to clear the previous conversation and context.
# # # # #     """
# # # # #     clear_past()
# # # # #     await message.reply("I've cleared the past conversation and context.")


# # # # # # Store user states
# # # # # states = {}
# # # # # # Dictionary to hold question-answer pairs
# # # # # user_responses = {}

# # # # # # Define Questions for assessment
# # # # # questions = [
# # # # #     """
# # # # #         1. Singapore plans to build a new observation tower called 'The Rook'.
# # # # #         How many steps do you think it will take to go to the top floor?

# # # # #         a) Less than 500 
# # # # #         b) More than 500

# # # # #     """,
# # # # #     "2. Now Guess the number of steps" ,
# # # # #     """
# # # # #     3. How confident are you that the real number is in the range you have selected? 
# # # # #     Answer within a range of 100.  
# # # # #     """,
# # # # #     """ 
# # # # #     4. You and your friend are betting on a series of coin tosses.

# # # # #     He always bets ₹2,000 on Heads

# # # # #     You always bet ₹2,000 on Tails

# # # # #     Winner of last 8 turns

# # # # #     You lost ₹8,000 in the last 4 turns!

# # # # #     If you were to bet one last time, what would you bet on heads or tails ?
# # # # #     """ ,
# # # # #     """
# # # # #     5. How confident are you that your bet will win this time?
# # # # #     Answer how confident you are. 
# # # # #     (Example: Not confident at all, somewhat confident, confident, or Very confident)
# # # # #     """,
# # # # #     """
# # # # #     6. What do you think are the chances that you will achieve and maintain your financial goals within the next 
# # # # #     10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
# # # # #     Answer how likely you are to achieve your goal.
# # # # #     (Example: Less likely than others, likely than others, or More likely than others)
# # # # #     """,
# # # # #     """
# # # # #     7. Imagine you are a contestant in a game show, and you are presented the following choices.

# # # # #     What would you prefer?
# # # # #     a) 50 percent chance of winning 15 gold coins 
# # # # #     b) 100 percent chance of winning 8 gold coins
# # # # #     """,
# # # # #     """
# # # # #     8. Ok, one last choice...

# # # # #     What would you prefer?
# # # # #     a) 50 percent chance of winning 15 gold coins 
# # # # #     b) 100 percent chance of winning 2 gold coins
# # # # #     """,
# # # # #     """
# # # # #     9. In general, how would your best friend describe your risk-taking tendencies?
# # # # #     a) A real gambler
# # # # #     b) Willing to take risks after completing adequate research
# # # # #     c) Cautious
# # # # #     d) Avoids risk as much as possible
# # # # #     """,
# # # # #     """
# # # # #     10. Suppose you could replace your current investment portfolio with this new one:
# # # # #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# # # # #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# # # # #     Answer between the range of -5 to -35.
# # # # #     """,
# # # # #     """
# # # # #     11. Suppose that in the next 7 years,

# # # # #     YOUR INCOME

# # # # #     grows 8% each year

# # # # #     VS
# # # # #     INFLATION

# # # # #     grows 10% a year

# # # # #     At the end of 7 years, how much will you be able to buy with your income?
# # # # #     Options:
# # # # #     a) More than today
# # # # #     b) Exactly the same
# # # # #     c) Less than today
# # # # #     d) Cannot say
# # # # #     """,
# # # # #     """
# # # # #     12. If somebody buys a bond of Company B, which of the following statements seems correct:
# # # # #     a) She owns part of Company B
# # # # #     b) She has lent money to Company B
# # # # #     c) She is liable for Company B's debt
# # # # #     d) Cannot say
# # # # #     """,
# # # # #     """
# # # # #     13. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
# # # # #     a) More than ₹2 lakhs
# # # # #     b) Less than ₹2 lakhs
# # # # #     c) Exactly ₹2 lakhs
# # # # #     d) Cannot say
# # # # #     """,
# # # # #     """
# # # # #     14. When an investor spreads money across different asset classes, what happens to the risk of losing money:
# # # # #     a) Increases
# # # # #     b) Decreases
# # # # #     c) Stays the same
# # # # #     d) Cannot say
# # # # #     """,
# # # # #     """
# # # # #     15. When a country's central bank reduces interest rates, it makes:

# # # # #     a) Borrowing more attractive and saving less attractive
# # # # #     b) Borrowing less attractive and saving more attractive
# # # # #     c) Both borrowing and saving less attractive
# # # # #     d) Cannot say
# # # # #     """,
# # # # #     """
# # # # #     16. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
# # # # #     a) Option A : Min Value is 9.2L and Max Value is 11.8L
# # # # #     b) Option B : Min Value is 8.8L and Max Value is 12.3L
# # # # #     c) Option C : Min Value is 8.5L and Max Value is 12.8L
# # # # #     d) Option D : Min Value is 8.1L and Max Value is 13.3L
# # # # #     e) Option E : Min Value is 7.8L and Max Value is 13.8L
# # # # #     """,
# # # # #     """
# # # # #     17. From Sept 2008 to Nov 2008, Stock market went down by 31%.

# # # # #     If you owned a stock investment that lost about 31 percent in 3 months, you would:
# # # # #     a) Sell all of the remaining investment
# # # # #     b) Sell a portion of the remaining investment
# # # # #     c) Hold on to the investment and sell nothing
# # # # #     d) Buy little
# # # # #     e) Buy more of the investment
# # # # #     """,
# # # # #     """
# # # # #     18. Over any 1-year period, what would be the maximum drop in the value of your investment 
# # # # #     portfolio that you would be comfortable with?
# # # # #     a) <5%
# # # # #     b) 5 - 10%
# # # # #     c) 10 - 15%
# # # # #     d) 15 - 20%
# # # # #     e) >20%
# # # # #     """,
# # # # #     """
# # # # #     19. When investing, what do you consider the most?

# # # # #     a) Risk 
# # # # #     b) Return
# # # # #     """,
# # # # #     """
# # # # #     20. What best describes your attitude?

# # # # #     a) Prefer reasonable returns, can take reasonable risk
# # # # #     b) Like higher returns, can take slightly higher risk
# # # # #     c) Want to maximize returns, can take significant high risk
# # # # #     """,
# # # # #     """
# # # # #     21. How much monthly investment you want to do?
# # # # #     """,
# # # # #     """
# # # # #     22. What is the time horizon for your investment?
# # # # #     You can answer in any range, example 1-5 years."""  
# # # # # ]


# # # # # # Handler for /start command
# # # # # @dispatcher.message_handler(commands=['start'])
# # # # # async def handle_start(message: types.Message):
# # # # #     """
# # # # #     This handler receives messages with /start command
# # # # #     """
# # # # #     chat_id = message.chat.id
# # # # #     # Start asking questions
# # # # #     await start_assessment(chat_id)


# # # # # # Function to start the assessment
# # # # # async def start_assessment(chat_id):
# # # # #     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Managemnet Advisor ChatBot! Let's start a quick personality assessment.")
# # # # #     await ask_next_question(chat_id, 0)

# # # # # # Function to ask the next question
# # # # # async def ask_next_question(chat_id, question_index):
# # # # #     if question_index < len(questions):
# # # # #         # Ask the next question
# # # # #         await bot.send_message(chat_id, questions[question_index])
# # # # #         # Update state to indicate the next expected answer
# # # # #         states[chat_id] = question_index
# # # # #     else:
# # # # #         # No more questions, finish assessment
# # # # #         await finish_assessment(chat_id)

# # # # # # Handler for receiving assessment answers
# # # # # assessment_in_progress = True


# # # # # async def finish_assessment(chat_id):
# # # # #     if chat_id in states and states[chat_id] == len(questions):
# # # # #         # All questions have been answered, now process the assessment
# # # # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # # # #         # Determine investment personality based on collected responses
# # # # #         global investment_personality
# # # # #         investment_personality = await determine_investment_personality(user_responses)

# # # # #         # Inform the user about their investment personality
# # # # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # # # #         # Summarize collected information
# # # # #         global summary
# # # # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # # # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # # # #         # Ensure to await the determination of investment personality
# # # # #         await send_summary_chunks(chat_id, summary)
# # # # #         global assessment_in_progress 
# # # # #         assessment_in_progress = False
# # # # #         # Prompt the user to begin financial advice process
# # # # #         await bot.send_message(chat_id, "To receive financial advice, please upload a document with your financial details using /begin.")

# # # # # async def send_summary_chunks(chat_id, summary):
# # # # #     # Split the summary into chunks that fit within Telegram's message limits
# # # # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# # # # #     # Send each chunk as a separate message
# # # # #     for chunk in chunks:
# # # # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # # # # async def determine_investment_personality(assessment_data):
# # # # #     try:
# # # # #         # Prepare input text for the chatbot based on assessment data
# # # # #         input_text = "User Profile:\n"
# # # # #         for question, answer in assessment_data.items():
# # # # #             input_text += f"{question}: {answer}\n"

# # # # #         # Introduce the chatbot's task and prompt for classification
# # # # #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# # # # #                       "- Conservative Investor\n" \
# # # # #                       "- Moderate Investor\n" \
# # # # #                       "- Aggressive Investor"

# # # # #         # Use your generative AI model to generate a response
# # # # #         # print(input_text)
# # # # #         model = genai.GenerativeModel('gemini-pro')
# # # # #         response = model.generate_content(input_text)

# # # # #         # Determine the investment personality from the chatbot's response
# # # # #         response_text = response.text.lower()
# # # # #         if "conservative" in response_text:
# # # # #             personality = "Conservative Investor"
# # # # #         elif "moderate" in response_text:
# # # # #             personality = "Moderate Investor"
# # # # #         elif "aggressive" in response_text:
# # # # #             personality = "Aggressive Investor"
# # # # #         else:
# # # # #             personality = "Unknown"

# # # # #         return personality
# # # # #         # Send the determined investment personality back to the user
# # # # #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# # # # #     except Exception as e:
# # # # #         print(f"Error generating response: {e}")
# # # # #         #await bot.send_message(chat_id, "Error processing investment personality classification.")




# # # # # @dispatcher.message_handler(commands=['help'])
# # # # # async def helper(message: types.Message):
# # # # #     """
# # # # #     A handler to display the help menu.
# # # # #     """
# # # # #     help_command = """
# # # # #     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
# # # # #     /start - to start the investment personality assessment.
# # # # #     /clear - to clear the past conversation and context.
# # # # #     /help - to get this help menu.
# # # # #     /begin - to start the Financial Suggestion Conversation with the ChatBot.
# # # # #     I hope this helps. :)
# # # # #     """
# # # # #     await message.reply(help_command)

# # # # # # Handler for /begin command to initiate financial advice
# # # # # @dispatcher.message_handler(commands=['begin'])
# # # # # async def handle_begin(message: types.Message):
# # # # #     chat_id = message.chat.id
# # # # #     file_instructions ="""Hello there!My name is Finbot and I am a Wealth Management Advisor Chatbot.I need more details related to your Financial Profile so that I can give you 
# # # # #     personalised Financial Advice.Please follow this drive link : https://docs.google.com/document/d/1JaqWUXcq3MNrPTCvEy_2RngbGIY0rybX/edit?usp=sharing&ouid=101252223236130106095&rtpof=true&sd=true
# # # # #     ,there you will find a Word Document.Please fill in your details with correct Information and then upload the Document in the chat. """

# # # # #     # """
# # # # #     # Hi there! I'm now in Financial Advisor mode. Please upload a document with your financial details.
# # # # #     # """
# # # # #     await message.reply(file_instructions)
# # # # #     #await bot.send_message(chat_id, "Please upload a document with your financial details.")

# # # # # # Handler for document upload
# # # # # @dispatcher.message_handler(content_types=['document'])
# # # # # async def handle_document(message: types.Message):
# # # # #     global summary,investment_personality  
    
# # # # #     # Obtain file information
# # # # #     file_id = message.document.file_id
# # # # #     file = await bot.get_file(file_id)
# # # # #     file_path = file.file_path
    
# # # # #     # Download the file
# # # # #     await bot.download_file(file_path, "data/uploaded_file")
    
# # # # #     # Process the uploaded document
# # # # #     extracted_text = await process_document("data/uploaded_file")
    
# # # # #     if extracted_text:
# # # # #         # Load vector database (assuming this is part of setting up the retriever)
# # # # #         retriever = await load_vector_db("data/uploaded_file")

# # # # #         # Extract client name and age from the uploaded document
                
# # # # #         # Example usage:
# # # # #         # file_path = "data/uploaded_file"
# # # # #         # client_name, validation_errors = validate_document_content(extracted_text)    #(file_path)
# # # # #         # if validation_errors:
# # # # #         #     validation_message = "**Validation Errors:**\n" + "\n".join([f"* {error}" for error in validation_errors])
# # # # #         #     await message.reply(validation_message)
# # # # #         #     await message.reply("Please fill in the details with correct information and then try again by clicking the /begin command and reupload the file.")
# # # # #         #     return

# # # # #             # Usage
# # # # #         file_path = 'data/uploaded_file'
# # # # #         client_name, validation_errors = await process_document(file_path)

# # # # #         # Print results
# # # # #         print(f"Client Name: {client_name}")
# # # # #         if validation_errors:
# # # # #             print("**Validation Errors:**")
# # # # #             for error in validation_errors:
# # # # #                 print(error)
# # # # #         else:
# # # # #             print("All fields are filled correctly.")

# # # # #         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")
        

# # # # #         # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # #         #         should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # #         #         investment by giving them an approximate return amount based on the time horizon of the investment the user had, based on which calculate the compunded returns on their 
# # # # #         #         investment and give the approximate return amount.Also explain the user why you are giving them that particular
# # # # #         #         investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # #         #         User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # # # #         #         So dont mention user to consult to a financial expert."""

        
# # # # #         if retriever is None:
# # # # #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # # #             return

# # # # #         # Check if a valid chain can be created
# # # # #         chain = await make_retrieval_chain(retriever)
# # # # #         if chain is None:
# # # # #             await message.reply("Failed to create the retrieval chain.")
# # # # #             return
        
# # # # #         try:
# # # # #             # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # #             #     Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # #             #     should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # #             #     investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # #             #     investment.Also explain the user why you are giving them that particular
# # # # #             #     investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # #             #     User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # # # #             #     So dont mention user to consult to a financial expert."""
                 
# # # # #             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
# # # # #             response = chain.invoke({"input": query})
# # # # #             print(response['answer'])
# # # # #             global chat_history
# # # # #             chat_history = response['answer'] 
# # # # #             print(f"\n Chat History : {chat_history}")
# # # # #             await message.reply(response['answer'])

# # # # #         except Exception as e:
# # # # #             print(f"Error invoking retrieval chain on attempt : {e}")
   
# # # # #     else:
# # # # #         await message.reply("Failed to process the uploaded file.")



# # # # # async def process_document(file_path):
# # # # #     try:
# # # # #         file_type = filetype.guess(file_path)
# # # # #         if file_type is not None:
# # # # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # # #                 # return extract_text_from_word(file_path)
# # # # #                 return extract_text_and_tables_from_word(file_path)
# # # # #             elif file_type.mime == "application/pdf":
# # # # #                 return extract_text_from_pdf(file_path)
# # # # #         return None
# # # # #     except Exception as e:
# # # # #         print(f"Error processing document: {e}")
# # # # #         return None

# # # # # # def extract_text_from_word(docx_file_path):
# # # # # #     try:
# # # # # #         doc = docx.Document(docx_file_path)
# # # # # #         text_content = []
# # # # # #         for para in doc.paragraphs:
# # # # # #             text_content.append(para.text)
# # # # # #         return "\n".join(text_content)
# # # # # #     except Exception as e:
# # # # # #         print(f"Error extracting text from Word document: {e}")
# # # # # #         return None

# # # # # def extract_text_from_pdf(pdf_file_path):
# # # # #     try:
# # # # #         with open(pdf_file_path, "rb") as pdf_file:
# # # # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # # # #             text_content = []
# # # # #             for page_num in range(pdf_reader.numPages):
# # # # #                 page = pdf_reader.getPage(page_num)
# # # # #                 text_content.append(page.extract_text())
# # # # #             return "\n".join(text_content)
# # # # #     except Exception as e:
# # # # #         print(f"Error extracting text from PDF: {e}")
# # # # #         return None


# # # # # # def validate_document_content(text):
# # # # # #     errors = []
# # # # # #     client_name = None

# # # # # #     # Validate client and co-client information
# # # # # #     client_info_pattern = re.compile(r"Client Name:\s*(.+?)\s*Age:\s*(\d+)")
# # # # # #     client_info_match = client_info_pattern.search(text)
# # # # # #     if client_info_match:
# # # # # #         client_name = client_info_match.group(1)
# # # # # #         client_age = client_info_match.group(2)
# # # # # #         if not client_age.isdigit():
# # # # # #             errors.append("Invalid or missing client age.")
# # # # # #     else:
# # # # # #         errors.append("Client Name or Age missing.")
    
# # # # # #     co_client_info_pattern = re.compile(r"Co-Client Name:\s*(.+?)\s*Age:\s*(\d+)")
# # # # # #     co_client_info_match = co_client_info_pattern.search(text)
# # # # # #     if co_client_info_match:
# # # # # #         co_client_name = co_client_info_match.group(1)
# # # # # #         co_client_age = co_client_info_match.group(2)
# # # # # #         if not co_client_age.isdigit():
# # # # # #             errors.append("Invalid or missing co-client age.")
# # # # # #     else:
# # # # # #         errors.append("Co-Client Name or Age missing.")
    
# # # # # #     # Validate Retirement Goal
# # # # # #     retirement_goal_section = re.search(r"YOUR RETIREMENT GOAL(.*?)YOUR OTHER MAJOR GOALS", text, re.DOTALL)
# # # # # #     if retirement_goal_section:
# # # # # #         retirement_goal_text = retirement_goal_section.group(1)
# # # # # #         required_fields = [
# # # # # #             r"When do you plan to retire\?\s*\(age or date\)\s*(\d+)",
# # # # # #             r"Social Security Benefit\s*\(include expected start date\)\s*(.+?)\s*Pension Benefit",
# # # # # #             r"Pension Benefit\s*\(include expected start date\)\s*(.+?)\s*Other Expected Income",
# # # # # #             r"Other Expected Income\s*\(rental, part-time work, etc.\)\s*(.+?)\s*Estimated Annual Retirement Expense",
# # # # # #             r"Estimated Annual Retirement Expense\s*\(\$\s*or\s*%\s*of\s*current\s*salary\)\s*(\$\d+|\d+%)"
# # # # # #         ]

# # # # # #         for field in required_fields:
# # # # # #             if not re.search(field, retirement_goal_text, re.DOTALL):
# # # # # #                 errors.append(f"Missing or incomplete data for: {field.split(' ')[0]}")
# # # # # #     else:
# # # # # #         errors.append("YOUR RETIREMENT GOAL section missing.")
    
# # # # # #     # Validate Other Major Goals
# # # # # #     other_goals_section = re.search(r"YOUR OTHER MAJOR GOALS(.*?)YOUR ASSETS AND LIABILITIES", text, re.DOTALL)
# # # # # #     if other_goals_section:
# # # # # #         goal_text = other_goals_section.group(1)
# # # # # #         goal_pattern = re.compile(r"GOAL:\s*(.+?)\s*COST:\s*(\$\d+|\d+)\s*WHEN:\s*(.+)")
# # # # # #         if not goal_pattern.search(goal_text):
# # # # # #             errors.append("At least one goal, along with its cost and timeframe, must be filled in YOUR OTHER MAJOR GOALS.")
# # # # # #     else:
# # # # # #         errors.append("YOUR OTHER MAJOR GOALS section missing.")
    
# # # # # #     # Validate Assets and Liabilities
# # # # # #     assets_section = re.search(r"YOUR ASSETS AND LIABILITIES(.*?)MY LIABILITIES", text, re.DOTALL)
# # # # # #     if assets_section:
# # # # # #         asset_text = assets_section.group(1)
# # # # # #         required_assets = {
# # # # # #             "Cash/bank accounts": r"Cash/bank accounts\s*Current Value:\s*\$\d+\s*Annual Contributions:\s*\$\d+",
# # # # # #             "Home": r"Home\s*Current Value:\s*\$\d+\s*Annual Contributions:\s*\$\d+",
# # # # # #             "Other Real Estate": r"Other Real Estate\s*Current Value:\s*\$\d+\s*Annual Contributions:\s*\$\d+",
# # # # # #             "Business": r"Business\s*Current Value:\s*\$\d+\s*Annual Contributions:\s*\$\d+"
# # # # # #         }

# # # # # #         for asset, pattern in required_assets.items():
# # # # # #             if not re.search(pattern, asset_text, re.DOTALL):
# # # # # #                 errors.append(f"{asset} details missing in YOUR ASSETS AND LIABILITIES.")
# # # # # #     else:
# # # # # #         errors.append("YOUR ASSETS AND LIABILITIES section missing.")
    
# # # # # #     # Validate Liabilities
# # # # # #     liabilities_section = re.search(r"MY LIABILITIES(.*?)(?:\n\s*\n|$)", text, re.DOTALL)
# # # # # #     if liabilities_section:
# # # # # #         liability_text = liabilities_section.group(1)
# # # # # #         required_liabilities = ["Mortgage(s)", "Credit Card(s)", "Other loans"]
# # # # # #         for liability in required_liabilities:
# # # # # #             if not re.search(fr"{liability}\s*Current Balance:\s*\$\d+\s*Monthly Payment:\s*\$\d+", liability_text, re.DOTALL):
# # # # # #                 errors.append(f"{liability} details missing in MY LIABILITIES.")
# # # # # #     else:
# # # # # #         errors.append("MY LIABILITIES section missing.")
    
# # # # # #     return client_name, "\n".join(errors) if errors else None

# # # # # import re
# # # # # import docx

# # # # # def extract_text_and_tables_from_word(docx_file_path):
# # # # #     """
# # # # #     Extracts text and tables from a Word document (.docx).

# # # # #     Args:
# # # # #         docx_file_path (str): Path to the Word document file.

# # # # #     Returns:
# # # # #         tuple: Extracted text content and tables from the document.
# # # # #     """
# # # # #     try:
# # # # #         doc = docx.Document(docx_file_path)
# # # # #         text_content = []
# # # # #         tables_content = []

# # # # #         for para in doc.paragraphs:
# # # # #             text_content.append(para.text)

# # # # #         for table in doc.tables:
# # # # #             table_data = []
# # # # #             for row in table.rows:
# # # # #                 row_data = []
# # # # #                 for cell in row.cells:
# # # # #                     row_data.append(cell.text.strip())
# # # # #                 table_data.append(row_data)
# # # # #             tables_content.append(table_data)

# # # # #         return "\n".join(text_content), tables_content
# # # # #     except Exception as e:
# # # # #         print(f"Error extracting text and tables from Word document: {e}")
# # # # #         return None, None

# # # # # def validate_document_content(text, tables):
# # # # #     """
# # # # #     Validates the content of the document.

# # # # #     Args:
# # # # #         text (str): Extracted text content from the document.
# # # # #         tables (list): Extracted tables content from the document.

# # # # #     Returns:
# # # # #         tuple: Client name and validation errors.
# # # # #     """
# # # # #     errors = []
    
# # # # #     # Extract client name
# # # # #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# # # # #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# # # # #     # Define required sections
# # # # #     required_sections = [
# # # # #         "YOUR RETIREMENT GOAL",
# # # # #         "YOUR OTHER MAJOR GOALS",
# # # # #         "YOUR ASSETS AND LIABILITIES",
# # # # #         "MY LIABILITIES",
# # # # #         "YOUR CURRENT ANNUAL INCOME"
# # # # #     ]

# # # # #     # Check for the presence of required sections
# # # # #     for section in required_sections:
# # # # #         if section not in text:
# # # # #             errors.append(f"* {section} section missing.")
    
# # # # #     # Define table field checks
# # # # #     table_checks = {
# # # # #         "YOUR RETIREMENT GOAL": [
# # # # #             r"When do you plan to retire\? \(age or date\)",
# # # # #             r"Social Security Benefit \(include expected start date\)",
# # # # #             r"Pension Benefit \(include expected start date\)",
# # # # #             r"Other Expected Income \(rental, part-time work, etc.\)",
# # # # #             r"Estimated Annual Retirement Expense"
# # # # #         ],
# # # # #         "YOUR OTHER MAJOR GOALS": [
# # # # #             r"GOAL", r"COST", r"WHEN"
# # # # #         ],
# # # # #         "YOUR ASSETS AND LIABILITIES": [
# # # # #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# # # # #             r"Current Value", r"Annual Contributions"
# # # # #         ],
# # # # #         "MY LIABILITIES": [
# # # # #             r"Balance", r"Interest Rate", r"Monthly Payment"
# # # # #         ]
# # # # #     }

# # # # #     # Validate table content
# # # # #     for section, checks in table_checks.items():
# # # # #         section_found = False
# # # # #         for table in tables:
# # # # #             table_text = "\n".join(["\t".join(row) for row in table])
# # # # #             if section in table_text:
# # # # #                 section_found = True
# # # # #                 for check in checks:
# # # # #                     if not re.search(check, table_text, re.IGNORECASE):
# # # # #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# # # # #                 break
# # # # #         if not section_found:
# # # # #             errors.append(f"* {section} section missing.")

# # # # #     return client_name, errors

# # # # # async def process_document(file_path):
# # # # #     try:
# # # # #         text, tables = extract_text_and_tables_from_word(file_path)
# # # # #         if text is not None and tables is not None:
# # # # #             client_name, errors = validate_document_content(text, tables)
# # # # #             return client_name, errors
# # # # #         return None, ["Error processing document."]
# # # # #     except Exception as e:
# # # # #         print(f"Error processing document: {e}")
# # # # #         return None, [f"Error processing document: {e}"]



# # # # # async def load_vector_db(file_path):
# # # # #     try:
# # # # #         loader = Docx2txtLoader(file_path)
# # # # #         documents = loader.load()
# # # # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # # # #         text_chunks = text_splitter.split_documents(documents)
# # # # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # # # #         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
# # # # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # # # #     except Exception as e:
# # # # #         print(f"Error loading vector database: {e}")
# # # # #         return None


# # # # # async def make_retrieval_chain(retriever):
# # # # #     """
# # # # #     Create a retrieval chain using the provided retriever.

# # # # #     Args:
# # # # #         retriever (RetrievalQA): A retriever object.

# # # # #     Returns:
# # # # #         RetrievalQA: A retrieval chain object.
# # # # #     """
# # # # #     try:
# # # # #         global investment_personality,summary
# # # # #         llm = ChatGoogleGenerativeAI(
# # # # #             model="gemini-pro",
# # # # #             temperature=0.7,
# # # # #             top_p=0.85,
# # # # #             google_api_key=GOOGLE_API_KEY
# # # # #         )

# # # # #         # prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # # #         #         Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # #         #         should invest in terms of percentage of their investment amount.Consider monthly investment unless user mentions annual.Give the user detailed information about the returns on their 
# # # # #         #         investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # #         #         investment and give approximate return amount.Also explain the user why you are giving them that particular
# # # # #         #         investment suggestion.ive user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # #         #         User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.So dont 
# # # # #         #         mention user to consult to a financial expert.Just give the response in 5-6 lines and not the whole paragraph containing previous information in response,
# # # # #         #         unless user asks for detailed answer.\n
# # # # #         #         <context>
# # # # #         #         {context}
# # # # #         #         </context>
# # # # #         #         Question: {input}"""

# # # # #         prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # # #                 Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # #                 should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # #                 investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # #                 investment.Also explain the user why you are giving them that particular
# # # # #                 investment suggestion.Answer in 3-4 lines.\n
# # # # #                 <context>
# # # # #                 {context}
# # # # #                 </context>
# # # # #                 Question: {input}"""


# # # # #         # finacial_suggestion = None 
# # # # #         # finacial_suggestion = store_financial_advice(finacial_suggestion)
# # # # #         # prompt_template = investment_personality + "\n" + summary + "\n" + finacial_suggestion + """\n" +""
# # # # #         #         <context>
# # # # #         #         {context}
# # # # #         #         </context>
# # # # #         #         Question: {input}"""

    
        

# # # # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # # # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
# # # # #         # Update combine_docs_chain with your actual document combining logic
# # # # #         combine_docs_chain = None  # Replace this with your combine_docs_chain

# # # # #         if retriever is not None :  #and combine_docs_chain is not None:
# # # # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # # # #             #combine_docs_chain)
# # # # #             # response = retriever_chain.invoke({"input":"Give me client detail "})
# # # # #             # print(response['answer'])
# # # # #             print(retriever_chain)
# # # # #             return retriever_chain
# # # # #         else:
# # # # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # # # #             return None

# # # # #     except Exception as e:
# # # # #         print(f"Error in creating chain: {e}")
# # # # #         return None


# # # # # # Compund Interest Calculator :

# # # # # # from typing import Optional, Type
# # # # # # from pydantic import BaseModel, Field
# # # # # # from langchain.tools import BaseTool
# # # # # # from langchain.callbacks.manager import AsyncCallbackManagerForToolRun, CallbackManagerForToolRun
# # # # # # from langchain_community.chat_models import ChatVertexAI
# # # # # # import vertexai

# # # # # # class CalculatorInput(BaseModel):
# # # # # #     a: float = Field(description="Principal amount")
# # # # # #     b: float = Field(description="Annual interest rate (percentage)")
# # # # # #     c: int = Field(description="Number of years")

# # # # # # class CustomCalculatorTool(BaseTool):
# # # # # #     name = "Calculator"
# # # # # #     description = "Use this tool to calculate compound interest"
# # # # # #     args_schema: Type[BaseModel] = CalculatorInput
# # # # # #     return_direct: bool = True

# # # # # #     def _run(self, a: float, b: float, c: int, run_manager: Optional[CallbackManagerForToolRun] = None) -> str:
# # # # # #         """Use the tool to calculate compound interest."""
# # # # # #         result = a * ((1 + b/100) ** c)
# # # # # #         return f"The compound interest result is: {result}"

# # # # # #     async def _arun(self, a: float, b: float, c: int, run_manager: Optional[AsyncCallbackManagerForToolRun] = None) -> str:
# # # # # #         """Use the tool asynchronously."""
# # # # # #         raise NotImplementedError("Calculator does not support async")

# # # # # # tools = [CustomCalculatorTool()]

# # # # # # # Initialize Google Cloud credentials
# # # # # # from google.auth import default
# # # # # # from google.auth.exceptions import DefaultCredentialsError
# # # # # # # from google.auth import google.auth.default()
# # # # # # try:
# # # # # #     credentials, project = google.auth.default()
# # # # # #     vertexai.init(project="finance-bot-424312", location="us-central1")
# # # # # # except DefaultCredentialsError as e:
# # # # # #     print("Google Cloud credentials not found. Please authenticate using `gcloud auth login` or set the environment variables.")



# # # # # @dispatcher.message_handler()
# # # # # async def main_bot(message: types.Message):
# # # # #     global retriever, extracted_text,investment_personality,summary,chat_history

# # # # #     # Handle the first tasks assessments answers from the user
# # # # #     chat_id = message.chat.id

# # # # #     if chat_id in states and states[chat_id] < len(questions):
# # # # #         # Retrieve the index of the current question
# # # # #         question_index = states[chat_id]

# # # # #         # Save the user's response to the current question
# # # # #         answer = message.text
# # # # #         user_responses[questions[question_index]] = answer
# # # # #         states[chat_id] += 1  # Move to the next question

# # # # #         # Ask the next question
# # # # #         await ask_next_question(chat_id, question_index + 1)
# # # # #     else:
# # # # #         # Handle q&a chat messages using your Gemini model (llm)
# # # # #         try:

# # # # #             task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # #             Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # #             should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # #             investment.Also explain the user why you are giving them that particular
# # # # #             investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # #             User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # # # #             So dont mention user to consult to a financial expert."""
        
# # # # #             model = genai.GenerativeModel('gemini-pro')
# # # # #             print(investment_personality)
# # # # #             # query = task + "\n" + investment_personality + "\n" + summary + "\n" +  extracted_text + "\n"  +   message.text

# # # # #             # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +  extracted_text + "\n"  +   message.text
# # # # #             query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
# # # # #             print(f"\nQuery : {query}")
# # # # #             response = model.generate_content(query)
# # # # #             await message.reply(response.text) #(response['answer']) 
# # # # #             # try:
# # # # #             #     ans = llm_with_tools.invoke(query).tool_calls
# # # # #             #     result = ans[0].output if ans else None

# # # # #             #     # Send the result back to the user
# # # # #             #     if result is not None:
# # # # #             #         await message.reply(result)
# # # # #             #     ans = llm_with_tools.invoke(query).tool_calls
# # # # #             #     await bot.send_message(ans)
# # # # #             # except Exception as e:
# # # # #             #     print(f"Error: {e}")
# # # # #             #     await message.reply("Failed to provide compound interest your request.")

# # # # #         except (NetworkError, RetryAfter) as e:
# # # # #             print(f"Network error: {e}. Retrying...")
# # # # #             await asyncio.sleep(5)  # Wait before retrying
# # # # #             await main_bot(message)  # Retry the message handling
# # # # #         except TelegramAPIError as e:
# # # # #             print(f"Telegram API error: {e}")
# # # # #             await message.reply("An error occurred while communicating with Telegram. Please try again later.")
# # # # #         except Exception as e:
# # # # #             print(f"Error processing general chat message: {e}")
# # # # #             await message.reply("Failed to process your request.")
        


# # # # # if __name__ == "__main__":
# # # # #     executor.start_polling(dispatcher, skip_updates=True)





# # # # # # #best code so far correctly takes the assessment and gives financial advice along 
# # # # # # #with returns with proper disclaimer and in brief

# # # # # # import os
# # # # # # import filetype
# # # # # # import docx
# # # # # # import PyPDF2
# # # # # # import re
# # # # # # from aiogram import Bot, Dispatcher, executor, types
# # # # # # from dotenv import load_dotenv
# # # # # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # # # # from langchain_community.vectorstores import Chroma
# # # # # # from langchain_community.document_loaders import Docx2txtLoader
# # # # # # from langchain_core.prompts import ChatPromptTemplate
# # # # # # from langchain.chains import create_retrieval_chain
# # # # # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # # # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # # # # from langchain.memory import ConversationSummaryMemory

# # # # # # import google.generativeai as genai

# # # # # # load_dotenv()

# # # # # # TOKEN = os.getenv("TOKEN")
# # # # # # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # # # # # os.environ["LANGCHAIN_TRACING_V2"]="true"
# # # # # # os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # # # # # # Configure generativeai with your API key
# # # # # # genai.configure(api_key=GOOGLE_API_KEY)

# # # # # # # Initialize bot
# # # # # # bot = Bot(token=TOKEN)
# # # # # # dispatcher = Dispatcher(bot)

# # # # # # rag_on = False
# # # # # # retriever = None  # Store retriever globally
# # # # # # summary = ""
# # # # # # investment_personality = ""
# # # # # # chat_history = ""

# # # # # # class Reference:
# # # # # #     def __init__(self):
# # # # # #         self.response = ""


# # # # # # reference = Reference()


# # # # # # def clear_past():
# # # # # #     reference.response = ""


# # # # # # @dispatcher.message_handler(commands=['clear'])
# # # # # # async def clear(message: types.Message):
# # # # # #     """
# # # # # #     A handler to clear the previous conversation and context.
# # # # # #     """
# # # # # #     clear_past()
# # # # # #     await message.reply("I've cleared the past conversation and context.")


# # # # # # # Store user states
# # # # # # states = {}
# # # # # # # Dictionary to hold question-answer pairs
# # # # # # user_responses = {}

# # # # # # # Define Questions for assessment
# # # # # # questions = [
# # # # # #     """
# # # # # #         1. Singapore plans to build a new observation tower called 'The Rook'.
# # # # # #         How many steps do you think it will take to go to the top floor?

# # # # # #         a) Less than 500 
# # # # # #         b) More than 500

# # # # # #     """,
# # # # # #     "2. Now Guess the number of steps" ,
# # # # # #     """
# # # # # #     3. How confident are you that the real number is in the range you have selected? 
# # # # # #     Answer within a range of 100.  
# # # # # #     """,
# # # # # #     """ 
# # # # # #     4. You and your friend are betting on a series of coin tosses.

# # # # # #     He always bets ₹2,000 on Heads

# # # # # #     You always bet ₹2,000 on Tails

# # # # # #     Winner of last 8 turns

# # # # # #     You lost ₹8,000 in the last 4 turns!

# # # # # #     If you were to bet one last time, what would you bet on heads or tails ?
# # # # # #     """ ,
# # # # # #     """
# # # # # #     5. How confident are you that your bet will win this time?
# # # # # #     Answer how confident you are. 
# # # # # #     (Example: Not confident at all, somewhat confident, confident, or Very confident)
# # # # # #     """,
# # # # # #     """
# # # # # #     6. What do you think are the chances that you will achieve and maintain your financial goals within the next 
# # # # # #     10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
# # # # # #     Answer how likely you are to achieve your goal.
# # # # # #     (Example: Less likely than others, likely than others, or More likely than others)
# # # # # #     """,
# # # # # #     """
# # # # # #     7. Imagine you are a contestant in a game show, and you are presented the following choices.

# # # # # #     What would you prefer?
# # # # # #     a) 50 percent chance of winning 15 gold coins 
# # # # # #     b) 100 percent chance of winning 8 gold coins
# # # # # #     """,
# # # # # #     """
# # # # # #     8. Ok, one last choice...

# # # # # #     What would you prefer?
# # # # # #     a) 50 percent chance of winning 15 gold coins 
# # # # # #     b) 100 percent chance of winning 2 gold coins
# # # # # #     """,
# # # # # #     """
# # # # # #     9. In general, how would your best friend describe your risk-taking tendencies?
# # # # # #     a) A real gambler
# # # # # #     b) Willing to take risks after completing adequate research
# # # # # #     c) Cautious
# # # # # #     d) Avoids risk as much as possible
# # # # # #     """,
# # # # # #     """
# # # # # #     10. Suppose you could replace your current investment portfolio with this new one:
# # # # # #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# # # # # #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# # # # # #     Answer between the range of -5 to -35.
# # # # # #     """,
# # # # # #     """
# # # # # #     11. Suppose that in the next 7 years,

# # # # # #     YOUR INCOME

# # # # # #     grows 8% each year

# # # # # #     VS
# # # # # #     INFLATION

# # # # # #     grows 10% a year

# # # # # #     At the end of 7 years, how much will you be able to buy with your income?
# # # # # #     Options:
# # # # # #     a) More than today
# # # # # #     b) Exactly the same
# # # # # #     c) Less than today
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     12. If somebody buys a bond of Company B, which of the following statements seems correct:
# # # # # #     a) She owns part of Company B
# # # # # #     b) She has lent money to Company B
# # # # # #     c) She is liable for Company B's debt
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     13. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
# # # # # #     a) More than ₹2 lakhs
# # # # # #     b) Less than ₹2 lakhs
# # # # # #     c) Exactly ₹2 lakhs
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     14. When an investor spreads money across different asset classes, what happens to the risk of losing money:
# # # # # #     a) Increases
# # # # # #     b) Decreases
# # # # # #     c) Stays the same
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     15. When a country's central bank reduces interest rates, it makes:

# # # # # #     a) Borrowing more attractive and saving less attractive
# # # # # #     b) Borrowing less attractive and saving more attractive
# # # # # #     c) Both borrowing and saving less attractive
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     16. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
# # # # # #     a) Option A : Min Value is 9.2L and Max Value is 11.8L
# # # # # #     b) Option B : Min Value is 8.8L and Max Value is 12.3L
# # # # # #     c) Option C : Min Value is 8.5L and Max Value is 12.8L
# # # # # #     d) Option D : Min Value is 8.1L and Max Value is 13.3L
# # # # # #     e) Option E : Min Value is 7.8L and Max Value is 13.8L
# # # # # #     """,
# # # # # #     """
# # # # # #     17. From Sept 2008 to Nov 2008, Stock market went down by 31%.

# # # # # #     If you owned a stock investment that lost about 31 percent in 3 months, you would:
# # # # # #     a) Sell all of the remaining investment
# # # # # #     b) Sell a portion of the remaining investment
# # # # # #     c) Hold on to the investment and sell nothing
# # # # # #     d) Buy little
# # # # # #     e) Buy more of the investment
# # # # # #     """,
# # # # # #     """
# # # # # #     18. Over any 1-year period, what would be the maximum drop in the value of your investment 
# # # # # #     portfolio that you would be comfortable with?
# # # # # #     a) <5%
# # # # # #     b) 5 - 10%
# # # # # #     c) 10 - 15%
# # # # # #     d) 15 - 20%
# # # # # #     e) >20%
# # # # # #     """,
# # # # # #     """
# # # # # #     19. When investing, what do you consider the most?

# # # # # #     a) Risk 
# # # # # #     b) Return
# # # # # #     """,
# # # # # #     """
# # # # # #     20. What best describes your attitude?

# # # # # #     a) Prefer reasonable returns, can take reasonable risk
# # # # # #     b) Like higher returns, can take slightly higher risk
# # # # # #     c) Want to maximize returns, can take significant high risk
# # # # # #     """,
# # # # # #     """
# # # # # #     21. How much monthly investment you want to do?
# # # # # #     """,
# # # # # #     """
# # # # # #     22. What is the time horizon for your investment?
# # # # # #     You can answer in any range, example 1-5 years."""  
# # # # # # ]


# # # # # # # Handler for /start command
# # # # # # @dispatcher.message_handler(commands=['start'])
# # # # # # async def handle_start(message: types.Message):
# # # # # #     """
# # # # # #     This handler receives messages with /start command
# # # # # #     """
# # # # # #     chat_id = message.chat.id
# # # # # #     # Start asking questions
# # # # # #     await start_assessment(chat_id)


# # # # # # # Function to start the assessment
# # # # # # async def start_assessment(chat_id):
# # # # # #     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Finance ChatBot! Let's start a quick personality assessment.")
# # # # # #     await ask_next_question(chat_id, 0)

# # # # # # # Function to ask the next question
# # # # # # async def ask_next_question(chat_id, question_index):
# # # # # #     if question_index < len(questions):
# # # # # #         # Ask the next question
# # # # # #         await bot.send_message(chat_id, questions[question_index])
# # # # # #         # Update state to indicate the next expected answer
# # # # # #         states[chat_id] = question_index
# # # # # #     else:
# # # # # #         # No more questions, finish assessment
# # # # # #         await finish_assessment(chat_id)

# # # # # # # Handler for receiving assessment answers
# # # # # # assessment_in_progress = True


# # # # # # async def finish_assessment(chat_id):
# # # # # #     if chat_id in states and states[chat_id] == len(questions):
# # # # # #         # All questions have been answered, now process the assessment
# # # # # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # # # # #         # Determine investment personality based on collected responses
# # # # # #         global investment_personality
# # # # # #         investment_personality = await determine_investment_personality(user_responses)

# # # # # #         # Inform the user about their investment personality
# # # # # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # # # # #         # Summarize collected information
# # # # # #         global summary
# # # # # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # # # # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # # # # #         # Ensure to await the determination of investment personality
# # # # # #         await send_summary_chunks(chat_id, summary)
# # # # # #         global assessment_in_progress 
# # # # # #         assessment_in_progress = False
# # # # # #         # Prompt the user to begin financial advice process
# # # # # #         await bot.send_message(chat_id, "To receive financial advice, please upload a document with your financial details using /begin.")

# # # # # # async def send_summary_chunks(chat_id, summary):
# # # # # #     # Split the summary into chunks that fit within Telegram's message limits
# # # # # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# # # # # #     # Send each chunk as a separate message
# # # # # #     for chunk in chunks:
# # # # # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # # # # # async def determine_investment_personality(assessment_data):
# # # # # #     try:
# # # # # #         # Prepare input text for the chatbot based on assessment data
# # # # # #         input_text = "User Profile:\n"
# # # # # #         for question, answer in assessment_data.items():
# # # # # #             input_text += f"{question}: {answer}\n"

# # # # # #         # Introduce the chatbot's task and prompt for classification
# # # # # #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# # # # # #                       "- Conservative Investor\n" \
# # # # # #                       "- Moderate Investor\n" \
# # # # # #                       "- Aggressive Investor"

# # # # # #         # Use your generative AI model to generate a response
# # # # # #         # print(input_text)
# # # # # #         model = genai.GenerativeModel('gemini-pro')
# # # # # #         response = model.generate_content(input_text)

# # # # # #         # Determine the investment personality from the chatbot's response
# # # # # #         response_text = response.text.lower()
# # # # # #         if "conservative" in response_text:
# # # # # #             personality = "Conservative Investor"
# # # # # #         elif "moderate" in response_text:
# # # # # #             personality = "Moderate Investor"
# # # # # #         elif "aggressive" in response_text:
# # # # # #             personality = "Aggressive Investor"
# # # # # #         else:
# # # # # #             personality = "Unknown"

# # # # # #         return personality
# # # # # #         # Send the determined investment personality back to the user
# # # # # #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# # # # # #     except Exception as e:
# # # # # #         print(f"Error generating response: {e}")
# # # # # #         #await bot.send_message(chat_id, "Error processing investment personality classification.")




# # # # # # @dispatcher.message_handler(commands=['help'])
# # # # # # async def helper(message: types.Message):
# # # # # #     """
# # # # # #     A handler to display the help menu.
# # # # # #     """
# # # # # #     help_command = """
# # # # # #     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
# # # # # #     /start - to start the investment personality assessment.
# # # # # #     /clear - to clear the past conversation and context.
# # # # # #     /help - to get this help menu.
# # # # # #     /begin - to start the Financial Suggestion Conversation with the ChatBot.
# # # # # #     I hope this helps. :)
# # # # # #     """
# # # # # #     await message.reply(help_command)

# # # # # # # Handler for /begin command to initiate financial advice
# # # # # # @dispatcher.message_handler(commands=['begin'])
# # # # # # async def handle_begin(message: types.Message):
# # # # # #     chat_id = message.chat.id
# # # # # #     file_instructions ="""Hello there!My name is Finbot and I am a Financial Advisor.I need more details related to your Financial Profile so that I can give you 
# # # # # #     personalised Financial Advice.Please follow this drive link : https://docs.google.com/document/d/1JaqWUXcq3MNrPTCvEy_2RngbGIY0rybX/edit?usp=sharing&ouid=101252223236130106095&rtpof=true&sd=true
# # # # # #     ,there you will find a Word Document.Please fill in your details with correct Information and then upload the Document in the chat. """

# # # # # #     # """
# # # # # #     # Hi there! I'm now in Financial Advisor mode. Please upload a document with your financial details.
# # # # # #     # """
# # # # # #     await message.reply(file_instructions)
# # # # # #     #await bot.send_message(chat_id, "Please upload a document with your financial details.")

# # # # # # # Handler for document upload
# # # # # # @dispatcher.message_handler(content_types=['document'])
# # # # # # async def handle_document(message: types.Message):
# # # # # #     global summary,investment_personality  
    
# # # # # #     # Obtain file information
# # # # # #     file_id = message.document.file_id
# # # # # #     file = await bot.get_file(file_id)
# # # # # #     file_path = file.file_path
    
# # # # # #     # Download the file
# # # # # #     await bot.download_file(file_path, "data/uploaded_file")
    
# # # # # #     # Process the uploaded document
# # # # # #     extracted_text = await process_document("data/uploaded_file")
    
# # # # # #     if extracted_text:
# # # # # #         # Load vector database (assuming this is part of setting up the retriever)
# # # # # #         retriever = await load_vector_db("data/uploaded_file")

# # # # # #         # Extract client name and age from the uploaded document
                
# # # # # #         # Example usage:
# # # # # #         file_path = "data/uploaded_file"
# # # # # #         client_name, validation_errors = validate_document_content(extracted_text)    #(file_path)
# # # # # #         # if validation_errors:
# # # # # #         #     validation_message = "**Validation Errors:**\n" + "\n".join([f"* {error}" for error in validation_errors])
# # # # # #         #     await message.reply(validation_message)
# # # # # #         #     await message.reply("Please fill in the details with correct information and then try again by clicking the /begin command and reupload the file.")
# # # # # #         #     return
        
# # # # # #         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some financial suggestions based on the details that you have provided.")
        

# # # # # #         # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # # #         #         should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # # #         #         investment by giving them an approximate return amount based on the time horizon of the investment the user had, based on which calculate the compunded returns on their 
# # # # # #         #         investment and give the approximate return amount.Also explain the user why you are giving them that particular
# # # # # #         #         investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # # #         #         User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # # # # #         #         So dont mention user to consult to a financial expert."""

        
# # # # # #         if retriever is None:
# # # # # #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # # # #             return

# # # # # #         # Check if a valid chain can be created
# # # # # #         chain = await make_retrieval_chain(retriever)
# # # # # #         if chain is None:
# # # # # #             await message.reply("Failed to create the retrieval chain.")
# # # # # #             return
        
# # # # # #         try:
# # # # # #             # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #             #     Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # # #             #     should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # # #             #     investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # # #             #     investment.Also explain the user why you are giving them that particular
# # # # # #             #     investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # # #             #     User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # # # # #             #     So dont mention user to consult to a financial expert."""
                 
# # # # # #             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
# # # # # #             response = chain.invoke({"input": query})
# # # # # #             print(response['answer'])
# # # # # #             global chat_history
# # # # # #             chat_history = response['answer'] 
# # # # # #             print(f"\n Chat History : {chat_history}")
# # # # # #             await message.reply(response['answer'])

# # # # # #         except Exception as e:
# # # # # #             print(f"Error invoking retrieval chain on attempt : {e}")

# # # # # #         # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # # #         #         should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # # #         #         investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # # #         #         investment.Also explain the user why you are giving them that particular
# # # # # #         #         investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # # #         #         User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # # # # #         #         So dont mention user to consult to a financial expert."""
        
# # # # # #         # model = genai.GenerativeModel('gemini-pro')
# # # # # #         # print(investment_personality)
# # # # # #         # query = extracted_text + "\n" + summary + "\n" + investment_personality + "\n" + task
# # # # # #         # response = model.generate_content(query)
# # # # # #         # await message.reply(response.text) #(response['answer']) 

# # # # # #         # store_financial_advice(response.text)
       
# # # # # #     else:
# # # # # #         await message.reply("Failed to process the uploaded file.")



# # # # # # async def process_document(file_path):
# # # # # #     try:
# # # # # #         file_type = filetype.guess(file_path)
# # # # # #         if file_type is not None:
# # # # # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # # # #                 return extract_text_from_word(file_path)
# # # # # #             elif file_type.mime == "application/pdf":
# # # # # #                 return extract_text_from_pdf(file_path)
# # # # # #         return None
# # # # # #     except Exception as e:
# # # # # #         print(f"Error processing document: {e}")
# # # # # #         return None

# # # # # # def extract_text_from_word(docx_file_path):
# # # # # #     try:
# # # # # #         doc = docx.Document(docx_file_path)
# # # # # #         text_content = []
# # # # # #         for para in doc.paragraphs:
# # # # # #             text_content.append(para.text)
# # # # # #         return "\n".join(text_content)
# # # # # #     except Exception as e:
# # # # # #         print(f"Error extracting text from Word document: {e}")
# # # # # #         return None

# # # # # # def extract_text_from_pdf(pdf_file_path):
# # # # # #     try:
# # # # # #         with open(pdf_file_path, "rb") as pdf_file:
# # # # # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # # # # #             text_content = []
# # # # # #             for page_num in range(pdf_reader.numPages):
# # # # # #                 page = pdf_reader.getPage(page_num)
# # # # # #                 text_content.append(page.extract_text())
# # # # # #             return "\n".join(text_content)
# # # # # #     except Exception as e:
# # # # # #         print(f"Error extracting text from PDF: {e}")
# # # # # #         return None


# # # # # # def validate_document_content(text):
# # # # # #     errors = []
# # # # # #     client_name = None

# # # # # #     # Validate client and co-client information
# # # # # #     client_info_pattern = re.compile(r"Client Name:\s*(.+?)\s*Age:\s*(\d+)")
# # # # # #     client_info_match = client_info_pattern.search(text)
# # # # # #     if client_info_match:
# # # # # #         client_name = client_info_match.group(1)
# # # # # #         client_age = client_info_match.group(2)
# # # # # #         if not client_age.isdigit():
# # # # # #             errors.append("Invalid or missing client age.")
# # # # # #     else:
# # # # # #         errors.append("Client Name or Age missing.")
    
# # # # # #     co_client_info_pattern = re.compile(r"Co-Client Name:\s*(.+?)\s*Age:\s*(\d+)")
# # # # # #     co_client_info_match = co_client_info_pattern.search(text)
# # # # # #     if co_client_info_match:
# # # # # #         co_client_name = co_client_info_match.group(1)
# # # # # #         co_client_age = co_client_info_match.group(2)
# # # # # #         if not co_client_age.isdigit():
# # # # # #             errors.append("Invalid or missing co-client age.")
# # # # # #     else:
# # # # # #         errors.append("Co-Client Name or Age missing.")
    
# # # # # #     # Validate Retirement Goal
# # # # # #     retirement_goal_section = re.search(r"YOUR RETIREMENT GOAL(.*?)YOUR OTHER MAJOR GOALS", text, re.DOTALL)
# # # # # #     if retirement_goal_section:
# # # # # #         retirement_goal_text = retirement_goal_section.group(1)
# # # # # #         required_fields = [
# # # # # #             r"When do you plan to retire\?\s*\(age or date\)\s*(\d+)",
# # # # # #             r"Social Security Benefit\s*\(include expected start date\)\s*(.+?)\s*Pension Benefit",
# # # # # #             r"Pension Benefit\s*\(include expected start date\)\s*(.+?)\s*Other Expected Income",
# # # # # #             r"Other Expected Income\s*\(rental, part-time work, etc.\)\s*(.+?)\s*Estimated Annual Retirement Expense",
# # # # # #             r"Estimated Annual Retirement Expense\s*\(\$\s*or\s*%\s*of\s*current\s*salary\)\s*(\$\d+|\d+%)"
# # # # # #         ]

# # # # # #         for field in required_fields:
# # # # # #             if not re.search(field, retirement_goal_text, re.DOTALL):
# # # # # #                 errors.append(f"Missing or incomplete data for: {field.split(' ')[0]}")
# # # # # #     else:
# # # # # #         errors.append("YOUR RETIREMENT GOAL section missing.")
    
# # # # # #     # Validate Other Major Goals
# # # # # #     other_goals_section = re.search(r"YOUR OTHER MAJOR GOALS(.*?)YOUR ASSETS AND LIABILITIES", text, re.DOTALL)
# # # # # #     if other_goals_section:
# # # # # #         goal_text = other_goals_section.group(1)
# # # # # #         goal_pattern = re.compile(r"GOAL:\s*(.+?)\s*COST:\s*(\$\d+|\d+)\s*WHEN:\s*(.+)")
# # # # # #         if not goal_pattern.search(goal_text):
# # # # # #             errors.append("At least one goal, along with its cost and timeframe, must be filled in YOUR OTHER MAJOR GOALS.")
# # # # # #     else:
# # # # # #         errors.append("YOUR OTHER MAJOR GOALS section missing.")
    
# # # # # #     # Validate Assets and Liabilities
# # # # # #     assets_section = re.search(r"YOUR ASSETS AND LIABILITIES(.*?)MY LIABILITIES", text, re.DOTALL)
# # # # # #     if assets_section:
# # # # # #         asset_text = assets_section.group(1)
# # # # # #         required_assets = {
# # # # # #             "Cash/bank accounts": r"Cash/bank accounts\s*Current Value:\s*\$\d+\s*Annual Contributions:\s*\$\d+",
# # # # # #             "Home": r"Home\s*Current Value:\s*\$\d+\s*Annual Contributions:\s*\$\d+",
# # # # # #             "Other Real Estate": r"Other Real Estate\s*Current Value:\s*\$\d+\s*Annual Contributions:\s*\$\d+",
# # # # # #             "Business": r"Business\s*Current Value:\s*\$\d+\s*Annual Contributions:\s*\$\d+"
# # # # # #         }

# # # # # #         for asset, pattern in required_assets.items():
# # # # # #             if not re.search(pattern, asset_text, re.DOTALL):
# # # # # #                 errors.append(f"{asset} details missing in YOUR ASSETS AND LIABILITIES.")
# # # # # #     else:
# # # # # #         errors.append("YOUR ASSETS AND LIABILITIES section missing.")
    
# # # # # #     # Validate Liabilities
# # # # # #     liabilities_section = re.search(r"MY LIABILITIES(.*?)(?:\n\s*\n|$)", text, re.DOTALL)
# # # # # #     if liabilities_section:
# # # # # #         liability_text = liabilities_section.group(1)
# # # # # #         required_liabilities = ["Mortgage(s)", "Credit Card(s)", "Other loans"]
# # # # # #         for liability in required_liabilities:
# # # # # #             if not re.search(fr"{liability}\s*Current Balance:\s*\$\d+\s*Monthly Payment:\s*\$\d+", liability_text, re.DOTALL):
# # # # # #                 errors.append(f"{liability} details missing in MY LIABILITIES.")
# # # # # #     else:
# # # # # #         errors.append("MY LIABILITIES section missing.")
    
# # # # # #     return client_name, "\n".join(errors) if errors else None


# # # # # # async def load_vector_db(file_path):
# # # # # #     try:
# # # # # #         loader = Docx2txtLoader(file_path)
# # # # # #         documents = loader.load()
# # # # # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # # # # #         text_chunks = text_splitter.split_documents(documents)
# # # # # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # # # # #         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
# # # # # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # # # # #     except Exception as e:
# # # # # #         print(f"Error loading vector database: {e}")
# # # # # #         return None


# # # # # # async def make_retrieval_chain(retriever):
# # # # # #     """
# # # # # #     Create a retrieval chain using the provided retriever.

# # # # # #     Args:
# # # # # #         retriever (RetrievalQA): A retriever object.

# # # # # #     Returns:
# # # # # #         RetrievalQA: A retrieval chain object.
# # # # # #     """
# # # # # #     try:
# # # # # #         global investment_personality,summary
# # # # # #         llm = ChatGoogleGenerativeAI(
# # # # # #             model="gemini-pro",
# # # # # #             temperature=0.7,
# # # # # #             top_p=0.85,
# # # # # #             google_api_key=GOOGLE_API_KEY
# # # # # #         )

# # # # # #         # prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # # # #         #         Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # # #         #         should invest in terms of percentage of their investment amount.Consider monthly investment unless user mentions annual.Give the user detailed information about the returns on their 
# # # # # #         #         investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # # #         #         investment and give approximate return amount.Also explain the user why you are giving them that particular
# # # # # #         #         investment suggestion.ive user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # # #         #         User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.So dont 
# # # # # #         #         mention user to consult to a financial expert.Just give the response in 5-6 lines and not the whole paragraph containing previous information in response,
# # # # # #         #         unless user asks for detailed answer.\n
# # # # # #         #         <context>
# # # # # #         #         {context}
# # # # # #         #         </context>
# # # # # #         #         Question: {input}"""

# # # # # #         prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # # # #                 Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # # #                 should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # # #                 investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # # #                 investment.Also explain the user why you are giving them that particular
# # # # # #                 investment suggestion.Answer in 3-4 lines.\n
# # # # # #                 <context>
# # # # # #                 {context}
# # # # # #                 </context>
# # # # # #                 Question: {input}"""


# # # # # #         # finacial_suggestion = None 
# # # # # #         # finacial_suggestion = store_financial_advice(finacial_suggestion)
# # # # # #         # prompt_template = investment_personality + "\n" + summary + "\n" + finacial_suggestion + """\n" +""
# # # # # #         #         <context>
# # # # # #         #         {context}
# # # # # #         #         </context>
# # # # # #         #         Question: {input}"""

    
        

# # # # # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # # # # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
# # # # # #         # Update combine_docs_chain with your actual document combining logic
# # # # # #         combine_docs_chain = None  # Replace this with your combine_docs_chain

# # # # # #         if retriever is not None :  #and combine_docs_chain is not None:
# # # # # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # # # # #             #combine_docs_chain)
# # # # # #             # response = retriever_chain.invoke({"input":"Give me client detail "})
# # # # # #             # print(response['answer'])
# # # # # #             print(retriever_chain)
# # # # # #             return retriever_chain
# # # # # #         else:
# # # # # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # # # # #             return None

# # # # # #     except Exception as e:
# # # # # #         print(f"Error in creating chain: {e}")
# # # # # #         return None

# # # # # # # get_file :

# # # # # # # get_gemini_response :
# # # # # # async def get_gemini_response(question):
# # # # # #     model = genai.GenerativeModel('gemini-pro')
# # # # # #     response = model.generate_content(question)
# # # # # #     return response.text


# # # # # # # model=genai.GenerativeModel("gemini-pro") 
# # # # # # # chat = model.start_chat(history=[])
# # # # # # # def get_gemini_chat_response(question):
    
# # # # # # #     response=chat.send_message(question,stream=True)
# # # # # # #     return response

# # # # # # import asyncio
# # # # # # from aiogram.utils.exceptions import NetworkError, RetryAfter, TelegramAPIError

# # # # # # def store_financial_advice(advice):
# # # # # #     if advice is not None:
# # # # # #         financial_suggestion = advice
# # # # # #         return
# # # # # #     advice = financial_suggestion
# # # # # #     print(f"Financial advice :{financial_suggestion}")
    

# # # # # # @dispatcher.message_handler()
# # # # # # async def main_bot(message: types.Message):
# # # # # #     global retriever, extracted_text,investment_personality,summary,chat_history

# # # # # #     # Handle the first tasks assessments answers from the user
# # # # # #     chat_id = message.chat.id

# # # # # #     if chat_id in states and states[chat_id] < len(questions):
# # # # # #         # Retrieve the index of the current question
# # # # # #         question_index = states[chat_id]

# # # # # #         # Save the user's response to the current question
# # # # # #         answer = message.text
# # # # # #         user_responses[questions[question_index]] = answer
# # # # # #         states[chat_id] += 1  # Move to the next question

# # # # # #         # Ask the next question
# # # # # #         await ask_next_question(chat_id, question_index + 1)
# # # # # #     else:
# # # # # #         # Handle q&a chat messages using your Gemini model (llm)
# # # # # #         try:

# # # # # #             task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #             Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # # #             should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # # #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # # #             investment.Also explain the user why you are giving them that particular
# # # # # #             investment suggestion.Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # # #             User should also invest as per their risk tolerance level.Since you are the financial advisor dont ask user to consult anyone else.
# # # # # #             So dont mention user to consult to a financial expert."""
        
# # # # # #             model = genai.GenerativeModel('gemini-pro')
# # # # # #             print(investment_personality)
# # # # # #             # query = task + "\n" + investment_personality + "\n" + summary + "\n" +  extracted_text + "\n"  +   message.text

# # # # # #             # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +  extracted_text + "\n"  +   message.text
# # # # # #             query = task + "\n" + investment_personality + "\n" + chat_history + "\n" +   message.text
# # # # # #             print(f"\nQuery : {query}")
# # # # # #             response = model.generate_content(query)
# # # # # #             await message.reply(response.text) #(response['answer']) 

# # # # # #             # if retriever is None:
# # # # # #             #     await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # # # #             #     return

# # # # # #             # # Check if a valid chain can be created
# # # # # #             # chain = await make_retrieval_chain(retriever)
# # # # # #             # if chain is None:
# # # # # #             #     await message.reply("Failed to create the retrieval chain.")
# # # # # #             #     return
            
# # # # # #             # try:
# # # # # #             #     response = chain.invoke({"input": message.text})
# # # # # #             #     print(response['answer'])
# # # # # #             #     await message.reply(response['answer'])
                    
# # # # # #             # except Exception as e:
# # # # # #             #     print(f"Error invoking retrieval chain on attempt {attempt + 1}: {e}")
# # # # # #             #     max_retries = 3
# # # # # #             #     for attempt in range(max_retries):
# # # # # #             #         try:
# # # # # #             #             response = chain.invoke({"input": message.text})
# # # # # #             #             print(response['answer'])
# # # # # #             #             await message.reply(response['answer'])
# # # # # #             #             break
# # # # # #             #         except Exception as e:
# # # # # #             #             print(f"Error invoking retrieval chain on attempt {attempt + 1}: {e}")
# # # # # #             #             if attempt == max_retries - 1:
# # # # # #             #                 # Investment suggestions logic
# # # # # #             #                 user_input = message.text.lower()
# # # # # #             #                 if "investment" in user_input or "suggestion" in user_input:
# # # # # #             #                     await message.reply("Based on your interest in investments, here are some suggestions:\n1. Diversify your portfolio.\n2. Consider low-cost index funds.\n3. Research individual stocks or bonds that match your risk tolerance and goals.\n4. Stay informed about market trends.\n5. Consult a financial advisor for personalized advice.")
# # # # # #             #                 await message.reply("Failed to process your request after multiple attempts.")

# # # # # #         except (NetworkError, RetryAfter) as e:
# # # # # #             print(f"Network error: {e}. Retrying...")
# # # # # #             await asyncio.sleep(5)  # Wait before retrying
# # # # # #             await main_bot(message)  # Retry the message handling
# # # # # #         except TelegramAPIError as e:
# # # # # #             print(f"Telegram API error: {e}")
# # # # # #             await message.reply("An error occurred while communicating with Telegram. Please try again later.")
# # # # # #         except Exception as e:
# # # # # #             print(f"Error processing general chat message: {e}")
# # # # # #             await message.reply("Failed to process your request.")
        


# # # # # # if __name__ == "__main__":
# # # # # #     executor.start_polling(dispatcher, skip_updates=True)








# # # # # # #best code so far correctly takes the assessment and gives financial advice along with returns

# # # # # # import os
# # # # # # import filetype
# # # # # # import docx
# # # # # # import PyPDF2
# # # # # # from aiogram import Bot, Dispatcher, executor, types
# # # # # # from dotenv import load_dotenv
# # # # # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # # # # from langchain_community.vectorstores import Chroma
# # # # # # from langchain_community.document_loaders import Docx2txtLoader
# # # # # # from langchain_core.prompts import ChatPromptTemplate
# # # # # # from langchain.chains import create_retrieval_chain
# # # # # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # # # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # # # # from langchain.memory import ConversationSummaryMemory

# # # # # # import google.generativeai as genai

# # # # # # load_dotenv()

# # # # # # TOKEN = os.getenv("TOKEN")
# # # # # # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # # # # # # Configure generativeai with your API key
# # # # # # genai.configure(api_key=GOOGLE_API_KEY)

# # # # # # # Initialize bot
# # # # # # bot = Bot(token=TOKEN)
# # # # # # dispatcher = Dispatcher(bot)

# # # # # # rag_on = False
# # # # # # retriever = None  # Store retriever globally
# # # # # # summary = ""
# # # # # # investment_personality = ""
# # # # # # chat_history = ""

# # # # # # class Reference:
# # # # # #     def __init__(self):
# # # # # #         self.response = ""


# # # # # # reference = Reference()


# # # # # # def clear_past():
# # # # # #     reference.response = ""


# # # # # # @dispatcher.message_handler(commands=['clear'])
# # # # # # async def clear(message: types.Message):
# # # # # #     """
# # # # # #     A handler to clear the previous conversation and context.
# # # # # #     """
# # # # # #     clear_past()
# # # # # #     await message.reply("I've cleared the past conversation and context.")


# # # # # # # Store user states
# # # # # # states = {}
# # # # # # # Dictionary to hold question-answer pairs
# # # # # # user_responses = {}

# # # # # # # Define Questions for assessment
# # # # # # questions = [
# # # # # #     "1. What is your name?",
# # # # # #     "2. What is your age?",
# # # # # #     "3. What is your gender?",
# # # # # #     "4. What is your marital status?",
# # # # # #     """
# # # # # #         5. Singapore plans to build a new observation tower called 'The Rook'.
# # # # # #         How many steps do you think it will take to go to the top floor?

# # # # # #         a) Less than 500 
# # # # # #         b) More than 500

# # # # # #     """,
# # # # # #     "6. Now Guess the number of steps" ,
# # # # # #     """
# # # # # #     7. How confident are you that the real number is in the range you have selected? 
# # # # # #     Answer within a range of 100.  
# # # # # #     """,
# # # # # #     """ 
# # # # # #     8. You and your friend are betting on a series of coin tosses.

# # # # # #     He always bets ₹2,000 on Heads

# # # # # #     You always bet ₹2,000 on Tails

# # # # # #     Winner of last 8 turns

# # # # # #     You lost ₹8,000 in the last 4 turns!

# # # # # #     If you were to bet one last time, what would you bet on heads or tails ?
# # # # # #     """ ,
# # # # # #     """
# # # # # #     9. How confident are you that your bet will win this time?
# # # # # #     Answer how confident you are. 
# # # # # #     (Example: Not confident at all, somewhat confident, confident, or Very confident)
# # # # # #     """,
# # # # # #     """
# # # # # #     10. What do you think are the chances that you will achieve and maintain your financial goals within the next 
# # # # # #     10 years, compared to others who are like you (in age, gender, wealth & stage of life)?
# # # # # #     Answer how likely you are to achieve your goal.
# # # # # #     (Example: Less likely than others, likely than others, or More likely than others)
# # # # # #     """,
# # # # # #     """
# # # # # #     11. Imagine you are a contestant in a game show, and you are presented the following choices.

# # # # # #     What would you prefer?
# # # # # #     a) 50 percent chance of winning 15 gold coins 
# # # # # #     b) 100 percent chance of winning 8 gold coins
# # # # # #     """,
# # # # # #     """
# # # # # #     12. Ok, one last choice...

# # # # # #     What would you prefer?
# # # # # #     a) 50 percent chance of winning 15 gold coins 
# # # # # #     b) 100 percent chance of winning 2 gold coins
# # # # # #     """,
# # # # # #     """
# # # # # #     13. In general, how would your best friend describe your risk-taking tendencies?
# # # # # #     a) A real gambler
# # # # # #     b) Willing to take risks after completing adequate research
# # # # # #     c) Cautious
# # # # # #     d) Avoids risk as much as possible
# # # # # #     """,
# # # # # #     """
# # # # # #     14. Suppose you could replace your current investment portfolio with this new one:
# # # # # #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# # # # # #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# # # # # #     Answer between the range of -5 to -35.
# # # # # #     """,
# # # # # #     """
# # # # # #     15. Suppose that in the next 7 years,

# # # # # #     YOUR INCOME

# # # # # #     grows 8% each year

# # # # # #     VS
# # # # # #     INFLATION

# # # # # #     grows 10% a year

# # # # # #     At the end of 7 years, how much will you be able to buy with your income?
# # # # # #     Options:
# # # # # #     a) More than today
# # # # # #     b) Exactly the same
# # # # # #     c) Less than today
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     16. If somebody buys a bond of Company B, which of the following statements seems correct:
# # # # # #     a) She owns part of Company B
# # # # # #     b) She has lent money to Company B
# # # # # #     c) She is liable for Company B's debt
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     17. An investment of ₹1 lakh compounded annually at an interest rate of 10 percent for 10 years will be worth:
# # # # # #     a) More than ₹2 lakhs
# # # # # #     b) Less than ₹2 lakhs
# # # # # #     c) Exactly ₹2 lakhs
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     18. When an investor spreads money across different asset classes, what happens to the risk of losing money:
# # # # # #     a) Increases
# # # # # #     b) Decreases
# # # # # #     c) Stays the same
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     19. When a country's central bank reduces interest rates, it makes:

# # # # # #     a) Borrowing more attractive and saving less attractive
# # # # # #     b) Borrowing less attractive and saving more attractive
# # # # # #     c) Both borrowing and saving less attractive
# # # # # #     d) Cannot say
# # # # # #     """,
# # # # # #     """
# # # # # #     20. If you had ₹10 lakhs to invest in a portfolio for a 1-year period which one would you choose?
# # # # # #     a) Option A : Min Value is 9.2L and Max Value is 11.8L
# # # # # #     b) Option B : Min Value is 8.8L and Max Value is 12.3L
# # # # # #     c) Option C : Min Value is 8.5L and Max Value is 12.8L
# # # # # #     d) Option D : Min Value is 8.1L and Max Value is 13.3L
# # # # # #     e) Option E : Min Value is 7.8L and Max Value is 13.8L
# # # # # #     """,
# # # # # #     """
# # # # # #     21. From Sept 2008 to Nov 2008, Stock market went down by 31%.

# # # # # #     If you owned a stock investment that lost about 31 percent in 3 months, you would:
# # # # # #     a) Sell all of the remaining investment
# # # # # #     b) Sell a portion of the remaining investment
# # # # # #     c) Hold on to the investment and sell nothing
# # # # # #     d) Buy little
# # # # # #     e) Buy more of the investment
# # # # # #     """,
# # # # # #     """
# # # # # #     22. Over any 1-year period, what would be the maximum drop in the value of your investment 
# # # # # #     portfolio that you would be comfortable with?
# # # # # #     a) <5%
# # # # # #     b) 5 - 10%
# # # # # #     c) 10 - 15%
# # # # # #     d) 15 - 20%
# # # # # #     e) >20%
# # # # # #     """,
# # # # # #     """
# # # # # #     23. When investing, what do you consider the most?

# # # # # #     a) Risk 
# # # # # #     b) Return
# # # # # #     """,
# # # # # #     """
# # # # # #     24. What best describes your attitude?

# # # # # #     a) Prefer reasonable returns, can take reasonable risk
# # # # # #     b) Like higher returns, can take slightly higher risk
# # # # # #     c) Want to maximize returns, can take significant high risk
# # # # # #     """,
# # # # # #     """
# # # # # #     25. What is your approximate monthly take-home salary/income?
# # # # # #     You may enter the salary that gets credited to your bank account or a rough estimate of your monthly income.
# # # # # #     """,
# # # # # #     """
# # # # # #     26. What is your monthly expenditure?

# # # # # #     How much do you end up roughly spending every month? Do not include EMIs.
# # # # # #     """,
# # # # # #     """
# # # # # #     27. How much are your assets worth?

# # # # # #     Some common assets include investments, deposits, real estate, etc.
# # # # # #     NOTE: For this assessment, please do not include the house you live in and the gold you use under "assets".
# # # # # #     """,
# # # # # #     """
# # # # # #     28. What is your total liability?

# # # # # #     A liability is something that you owe. Some common liabilities are home loans, car loans, debt, etc.
# # # # # #     """,
# # # # # #     """
# # # # # #     29. What is your total EMI amount?

# # # # # #     You usually pay EMIs for the loans you take for things such as education, home, car, etc.
# # # # # #     """,
# # # # # #     """
# # # # # #     30. How much monthly investment you want to do?
# # # # # #     """,
# # # # # #     """
# # # # # #     31. What is the time horizon for your investment?
# # # # # #     You can answer in any range, example 1-5 years."""  
# # # # # # ]


# # # # # # # Handler for /start command
# # # # # # @dispatcher.message_handler(commands=['start'])
# # # # # # async def handle_start(message: types.Message):
# # # # # #     """
# # # # # #     This handler receives messages with /start command
# # # # # #     """
# # # # # #     chat_id = message.chat.id
# # # # # #     # Start asking questions
# # # # # #     await start_assessment(chat_id)


# # # # # # # Function to start the assessment
# # # # # # async def start_assessment(chat_id):
# # # # # #     await bot.send_message(chat_id, "Hi, I am a Finance ChatBot! Let's start a quick personality assessment.")
# # # # # #     await ask_next_question(chat_id, 0)

# # # # # # # Function to ask the next question
# # # # # # async def ask_next_question(chat_id, question_index):
# # # # # #     if question_index < len(questions):
# # # # # #         # Ask the next question
# # # # # #         await bot.send_message(chat_id, questions[question_index])
# # # # # #         # Update state to indicate the next expected answer
# # # # # #         states[chat_id] = question_index
# # # # # #     else:
# # # # # #         # No more questions, finish assessment
# # # # # #         await finish_assessment(chat_id)

# # # # # # # Handler for receiving assessment answers
# # # # # # assessment_in_progress = True

# # # # # # # if assessment_in_progress:
# # # # # # #     @dispatcher.message_handler()
# # # # # # #     async def handle_answer(message: types.Message):
# # # # # # #         chat_id = message.chat.id

# # # # # # #         if chat_id in states and states[chat_id] < len(questions) :
# # # # # # #             # Retrieve the index of the current question
# # # # # # #             question_index = states[chat_id]

# # # # # # #             # Save the user's response to the current question
# # # # # # #             answer = message.text
# # # # # # #             user_responses[questions[question_index]] = answer
# # # # # # #             states[chat_id] += 1  # Move to the next question

# # # # # # #             # Ask the next question
# # # # # # #             await ask_next_question(chat_id, question_index + 1)



# # # # # # async def finish_assessment(chat_id):
# # # # # #     if chat_id in states and states[chat_id] == len(questions):
# # # # # #         # All questions have been answered, now process the assessment
# # # # # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # # # # #         # Determine investment personality based on collected responses
# # # # # #         global investment_personality
# # # # # #         investment_personality = await determine_investment_personality(user_responses)

# # # # # #         # Inform the user about their investment personality
# # # # # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # # # # #         # Summarize collected information
# # # # # #         global summary
# # # # # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # # # # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # # # # #         # global chat_history
# # # # # #         # chat_history=summary
# # # # # #         # Ensure to await the determination of investment personality
# # # # # #         await send_summary_chunks(chat_id, summary)
# # # # # #         global assessment_in_progress 
# # # # # #         assessment_in_progress = False
# # # # # #         # Prompt the user to begin financial advice process
# # # # # #         await bot.send_message(chat_id, "To receive financial advice, please upload a document with your financial details using /begin.")

# # # # # # async def send_summary_chunks(chat_id, summary):
# # # # # #     # Split the summary into chunks that fit within Telegram's message limits
# # # # # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# # # # # #     # Send each chunk as a separate message
# # # # # #     for chunk in chunks:
# # # # # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # # # # # async def determine_investment_personality(assessment_data):
# # # # # #     try:
# # # # # #         # Prepare input text for the chatbot based on assessment data
# # # # # #         input_text = "User Profile:\n"
# # # # # #         for question, answer in assessment_data.items():
# # # # # #             input_text += f"{question}: {answer}\n"

# # # # # #         # Introduce the chatbot's task and prompt for classification
# # # # # #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# # # # # #                       "- Conservative Investor\n" \
# # # # # #                       "- Moderate Investor\n" \
# # # # # #                       "- Aggressive Investor"

# # # # # #         # Use your generative AI model to generate a response
# # # # # #         print(input_text)
# # # # # #         model = genai.GenerativeModel('gemini-pro')
# # # # # #         response = model.generate_content(input_text)

# # # # # #         # Determine the investment personality from the chatbot's response
# # # # # #         response_text = response.text.lower()
# # # # # #         if "conservative" in response_text:
# # # # # #             personality = "Conservative Investor"
# # # # # #         elif "moderate" in response_text:
# # # # # #             personality = "Moderate Investor"
# # # # # #         elif "aggressive" in response_text:
# # # # # #             personality = "Aggressive Investor"
# # # # # #         else:
# # # # # #             personality = "Unknown"

# # # # # #         return personality
# # # # # #         # Send the determined investment personality back to the user
# # # # # #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# # # # # #     except Exception as e:
# # # # # #         print(f"Error generating response: {e}")
# # # # # #         #await bot.send_message(chat_id, "Error processing investment personality classification.")




# # # # # # @dispatcher.message_handler(commands=['help'])
# # # # # # async def helper(message: types.Message):
# # # # # #     """
# # # # # #     A handler to display the help menu.
# # # # # #     """
# # # # # #     help_command = """
# # # # # #     Hi there! I'm a bot created by Harshal Gidh. Please follow these commands:
# # # # # #     /start - to start the investment personality assessment.
# # # # # #     /clear - to clear the past conversation and context.
# # # # # #     /help - to get this help menu.
# # # # # #     /begin - to start the Financial Suggestion Conversation with the ChatBot.
# # # # # #     I hope this helps. :)
# # # # # #     """
# # # # # #     await message.reply(help_command)

# # # # # # # Handler for /begin command to initiate financial advice
# # # # # # @dispatcher.message_handler(commands=['begin'])
# # # # # # async def handle_begin(message: types.Message):
# # # # # #     chat_id = message.chat.id
# # # # # #     file_instructions = """
# # # # # #     Hi there! I'm now in Financial Advisor mode. Please upload a document with your financial details.
# # # # # #     """
# # # # # #     await message.reply(file_instructions)
# # # # # #     #await bot.send_message(chat_id, "Please upload a document with your financial details.")

# # # # # # # Handler for document upload
# # # # # # @dispatcher.message_handler(content_types=['document'])
# # # # # # async def handle_document(message: types.Message):
# # # # # #     global retriever,summary,investment_personality  
    
# # # # # #     # Obtain file information
# # # # # #     file_id = message.document.file_id
# # # # # #     file = await bot.get_file(file_id)
# # # # # #     file_path = file.file_path
    
# # # # # #     # Download the file
# # # # # #     await bot.download_file(file_path, "data/uploaded_file")
    
# # # # # #     # Process the uploaded document
# # # # # #     extracted_text = await process_document("data/uploaded_file")
    
# # # # # #     if extracted_text:
# # # # # #         # Load vector database (assuming this is part of setting up the retriever)
# # # # # #         retriever = await load_vector_db("data/uploaded_file")
        
# # # # # #         #if retriever:
# # # # # #         await message.reply("Processed file, I will now suggest you some Financial Suggestions based on the details you have provided.")

# # # # # #         task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # # #                 should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # # #                 investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # # #                 investment.Also explain the user why you are giving them that particular
# # # # # #                 investment suggestion."""
# # # # # #         model = genai.GenerativeModel('gemini-pro')
# # # # # #         print(investment_personality)
# # # # # #         query = extracted_text + "\n" + summary + "\n" + investment_personality + "\n" + task
# # # # # #         global chat_history
# # # # # #         chat_history = query
# # # # # #         print(query)
# # # # # #         response = model.generate_content(query)
# # # # # #         await message.reply(response.text) #(response['answer']) 
       
# # # # # #     else:
# # # # # #         await message.reply("Failed to process the uploaded file.")



# # # # # # async def process_document(file_path):
# # # # # #     try:
# # # # # #         file_type = filetype.guess(file_path)
# # # # # #         if file_type is not None:
# # # # # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # # # #                 return extract_text_from_word(file_path)
# # # # # #             elif file_type.mime == "application/pdf":
# # # # # #                 return extract_text_from_pdf(file_path)
# # # # # #         return None
# # # # # #     except Exception as e:
# # # # # #         print(f"Error processing document: {e}")
# # # # # #         return None

# # # # # # def extract_text_from_word(docx_file_path):
# # # # # #     """
# # # # # #     Extracts text content from a Word document (.docx).

# # # # # #     Args:
# # # # # #         docx_file_path (str): Path to the Word document file.

# # # # # #     Returns:
# # # # # #         str: Extracted text content from the document.
# # # # # #     """
# # # # # #     try:
# # # # # #         doc = docx.Document(docx_file_path)
# # # # # #         text_content = []
# # # # # #         for para in doc.paragraphs:
# # # # # #             text_content.append(para.text)
# # # # # #         return "\n".join(text_content)
# # # # # #     except Exception as e:
# # # # # #         print(f"Error extracting text from Word document: {e}")
# # # # # #         return None

# # # # # # def extract_text_from_pdf(pdf_file_path):
# # # # # #     """
# # # # # #     Extracts text content from a PDF file.

# # # # # #     Args:
# # # # # #         pdf_file_path (str): Path to the PDF file.

# # # # # #     Returns:
# # # # # #         str: Extracted text content from the PDF.
# # # # # #     """
# # # # # #     try:
# # # # # #         with open(pdf_file_path, "rb") as pdf_file:
# # # # # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # # # # #             text_content = []
# # # # # #             for page_num in range(pdf_reader.numPages):
# # # # # #                 page = pdf_reader.getPage(page_num)
# # # # # #                 text_content.append(page.extract_text())
# # # # # #             return "\n".join(text_content)
# # # # # #     except Exception as e:
# # # # # #         print(f"Error extracting text from PDF: {e}")
# # # # # #         return None

# # # # # # async def load_vector_db(file_path):
# # # # # #     try:
# # # # # #         loader = Docx2txtLoader(file_path)
# # # # # #         documents = loader.load()
# # # # # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # # # # #         text_chunks = text_splitter.split_documents(documents)
# # # # # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # # # # #         vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
# # # # # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # # # # #     except Exception as e:
# # # # # #         print(f"Error loading vector database: {e}")
# # # # # #         return None



# # # # # # async def make_retrieval_chain(retriever):
# # # # # #     """
# # # # # #     Create a retrieval chain using the provided retriever.

# # # # # #     Args:
# # # # # #         retriever (RetrievalQA): A retriever object.

# # # # # #     Returns:
# # # # # #         RetrievalQA: A retrieval chain object.
# # # # # #     """
# # # # # #     try:
# # # # # #         global investment_personality,summary
# # # # # #         llm = ChatGoogleGenerativeAI(
# # # # # #             model="gemini-pro",
# # # # # #             temperature=0.7,
# # # # # #             top_p=0.85,
# # # # # #             google_api_key=GOOGLE_API_KEY
# # # # # #         )

# # # # # #         prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # # # #                 Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # # #                 should invest in terms of percentage of their investment amount.Give the user detailed information about the returns on their 
# # # # # #                 investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # # #                 investment.Also explain the user why you are giving them that particular
# # # # # #                 investment suggestion.\n
# # # # # #                 <context>
# # # # # #                 {context}
# # # # # #                 </context>
# # # # # #                 Question: {input}"""
        
# # # # # #                 #Question: {question}\nContext: {context}\nAnswer:"""

# # # # # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # # # # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
# # # # # #         # Update combine_docs_chain with your actual document combining logic
# # # # # #         combine_docs_chain = None  # Replace this with your combine_docs_chain

# # # # # #         if retriever is not None :  #and combine_docs_chain is not None:
# # # # # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # # # # #             #combine_docs_chain)
# # # # # #             # response = retriever_chain.invoke({"input":"Give me client detail "})
# # # # # #             # print(response['answer'])
# # # # # #             print(retriever_chain)
# # # # # #             return retriever_chain
# # # # # #         else:
# # # # # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # # # # #             return None

# # # # # #     except Exception as e:
# # # # # #         print(f"Error in creating chain: {e}")
# # # # # #         return None

# # # # # # # get_file :

# # # # # # # get_gemini_response :
# # # # # # async def get_gemini_response(question):
# # # # # #     model = genai.GenerativeModel('gemini-pro')
# # # # # #     response = model.generate_content(question)
# # # # # #     return response.text


# # # # # # # model=genai.GenerativeModel("gemini-pro") 
# # # # # # # chat = model.start_chat(history=[])
# # # # # # # def get_gemini_chat_response(question):
    
# # # # # # #     response=chat.send_message(question,stream=True)
# # # # # # #     return response

# # # # # # import asyncio
# # # # # # from aiogram.utils.exceptions import NetworkError, RetryAfter, TelegramAPIError

# # # # # # @dispatcher.message_handler()
# # # # # # async def main_bot(message: types.Message):
# # # # # #     global chat_history, rag_on, retriever, extracted_text

# # # # # #     # Handle the first tasks assessments answers from the user
# # # # # #     chat_id = message.chat.id

# # # # # #     if chat_id in states and states[chat_id] < len(questions):
# # # # # #         # Retrieve the index of the current question
# # # # # #         question_index = states[chat_id]

# # # # # #         # Save the user's response to the current question
# # # # # #         answer = message.text
# # # # # #         user_responses[questions[question_index]] = answer
# # # # # #         states[chat_id] += 1  # Move to the next question

# # # # # #         # Ask the next question
# # # # # #         await ask_next_question(chat_id, question_index + 1)
# # # # # #     else:
# # # # # #         # Handle q&a chat messages using your Gemini model (llm)
# # # # # #         try:
# # # # # #             if retriever is None:
# # # # # #                 await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # # # #                 return

# # # # # #             # Check if a valid chain can be created
# # # # # #             chain = await make_retrieval_chain(retriever)
# # # # # #             if chain is None:
# # # # # #                 await message.reply("Failed to create the retrieval chain.")
# # # # # #                 return

# # # # # #             max_retries = 3
# # # # # #             for attempt in range(max_retries):
# # # # # #                 try:
# # # # # #                     response = chain.invoke({"input": message.text})
# # # # # #                     print(response['answer'])
# # # # # #                     await message.reply(response['answer'])
# # # # # #                     break
# # # # # #                 except Exception as e:
# # # # # #                     print(f"Error invoking retrieval chain on attempt {attempt + 1}: {e}")
# # # # # #                     if attempt == max_retries - 1:
# # # # # #                         await message.reply("Failed to process your request after multiple attempts.")

# # # # # #             # Investment suggestions logic
# # # # # #             user_input = message.text.lower()
# # # # # #             if "investment" in user_input or "suggestion" in user_input:
# # # # # #                 await message.reply("Based on your interest in investments, here are some suggestions:\n1. Diversify your portfolio.\n2. Consider low-cost index funds.\n3. Research individual stocks or bonds that match your risk tolerance and goals.\n4. Stay informed about market trends.\n5. Consult a financial advisor for personalized advice.")
# # # # # #         except (NetworkError, RetryAfter) as e:
# # # # # #             print(f"Network error: {e}. Retrying...")
# # # # # #             await asyncio.sleep(5)  # Wait before retrying
# # # # # #             await main_bot(message)  # Retry the message handling
# # # # # #         except TelegramAPIError as e:
# # # # # #             print(f"Telegram API error: {e}")
# # # # # #             await message.reply("An error occurred while communicating with Telegram. Please try again later.")
# # # # # #         except Exception as e:
# # # # # #             print(f"Error processing general chat message: {e}")
# # # # # #             await message.reply("Failed to process your request.")
        



# # # # # # # @dispatcher.message_handler()
# # # # # # # async def main_bot(message: types.Message):
# # # # # # #     global chat_history,rag_on, retriever,extracted_text

# # # # # # #     # if rag_on:
       
# # # # # # #     # else:

# # # # # # #     # to handle the first tasks assessments answers from the user
# # # # # # #     chat_id = message.chat.id

# # # # # # #     if chat_id in states and states[chat_id] < len(questions) :
# # # # # # #         # Retrieve the index of the current question
# # # # # # #         question_index = states[chat_id]

# # # # # # #         # Save the user's response to the current question
# # # # # # #         answer = message.text
# # # # # # #         user_responses[questions[question_index]] = answer
# # # # # # #         states[chat_id] += 1  # Move to the next question

# # # # # # #         # Ask the next question
# # # # # # #         await ask_next_question(chat_id, question_index + 1)
# # # # # # #     else :
# # # # # # #         # Handle q&a chat messages using your Gemini model (llm)
# # # # # # #         try:
# # # # # # #             if retriever is None:
# # # # # # #                 await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # # # # #                 return

# # # # # # #             # Check if a valid chain can be created
# # # # # # #             chain = await make_retrieval_chain(retriever)
# # # # # # #             if chain is None:
# # # # # # #                 await message.reply("Failed to create the retrieval chain.")
# # # # # # #                 return
        
# # # # # # #             try:
# # # # # # #                 response =  chain.invoke({"input":message.text}) 

# # # # # # #                 # response = await chain.invoke({"input":message.text}) 
# # # # # # #                 #(input_data) 
# # # # # # #                 #(({"context":chain,"question":message.text}))
# # # # # # #                 #({"context":retriever,"question":message.text}) 
# # # # # # #                 #({"input":message.text}) #(message.text)
# # # # # # #                 print(response['answer'])
# # # # # # #                 await message.reply(response['answer']) #(response.text)
# # # # # # #             except Exception as e:
# # # # # # #                 print(f"Error invoking retrieval chain: {e}")
# # # # # # #                 await message.reply("Failed to process your request.")


# # # # # # #             # # print(f"chat history : {chat_history}\n")
# # # # # # #             # print(f">>> USER: \n\t{message.text}")
# # # # # # #             # # chat_history = "\n".join([message.text])
# # # # # # #             # # response = get_gemini_chat_response(message.text)
# # # # # # #             # response = await get_gemini_response(chat_history) #(message.text)
# # # # # # #             # print(f">>> gemini: \n\t{response}")
# # # # # # #             # # chat_history = "\n".join([response])
# # # # # # #             # await bot.send_message(chat_id=message.chat.id, text=response)
# # # # # # #             # # await bot.send_message(chat_id=message.chat.id, text=response)

# # # # # # #         except Exception as e:
# # # # # # #             print(f"Error processing general chat message: {e}")
# # # # # # #             await message.reply("Failed to process your request.")


# # # # # # # @dispatcher.message_handler(content_types=['document'])
# # # # # # # async def get_file(message: types.Message):
# # # # # # #     global retriever
# # # # # # #     file_id = message.document.file_id
# # # # # # #     file = await bot.get_file(file_id)
# # # # # # #     file_path = file.file_path
# # # # # # #     await bot.download_file(file_path, "data/uploaded_file")
# # # # # # #     extracted_text = await process_document("data/uploaded_file")
# # # # # # #     if extracted_text:
# # # # # # #         retriever = await load_vector_db("data/uploaded_file")
# # # # # # #         if retriever:
# # # # # # #             await message.reply("Processed file, please ask your question.")
# # # # # # #         else:
# # # # # # #             await message.reply("Failed to load vector database.")
# # # # # # #     else:
# # # # # # #         await message.reply("Failed to process the uploaded file.")

# # # # # # # @dispatcher.message_handler()
# # # # # # # async def main_bot(message: types.Message):
# # # # # #     # global rag_on, retriever,extracted_text

# # # # # #     # if rag_on:
# # # # # #     #     if retriever is None:
# # # # # #     #         await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # # # #     #         return

# # # # # #     #     # Check if a valid chain can be created
# # # # # #     #     chain = await make_retrieval_chain(retriever)
# # # # # #     #     if chain is None:
# # # # # #     #         await message.reply("Failed to create the retrieval chain.")
# # # # # #     #         return
# # # # # #     #     # Prepare the input data for the retrieval chain
# # # # # #     #     # input_data = {
# # # # # #     #     #     "context": extracted_text,  # Use extracted text as the context
# # # # # #     #     #     "question": message.text
# # # # # #     #     # }
# # # # # #     #     # Invoke the retrieval chain with the user's message
# # # # # #     #     try:
# # # # # #     #         response =  chain.invoke({"input":message.text}) 

# # # # # #     #         # response = await chain.invoke({"input":message.text}) 
# # # # # #     #         #(input_data) 
# # # # # #     #         #(({"context":chain,"question":message.text}))
# # # # # #     #         #({"context":retriever,"question":message.text}) 
# # # # # #     #         #({"input":message.text}) #(message.text)
# # # # # #     #         print(response['answer'])
# # # # # #     #         await message.reply(response['answer']) #(response.text)
# # # # # #     #     except Exception as e:
# # # # # #     #         print(f"Error invoking retrieval chain: {e}")
# # # # # #     #         await message.reply("Failed to process your request.")
# # # # # #     # else:
# # # # # #     #     # to handle the first tasks assessments answers from the user
# # # # # #     #     chat_id = message.chat.id

# # # # # # #         if chat_id in states and states[chat_id] < len(questions) :
# # # # # # #             # Retrieve the index of the current question
# # # # # # #             question_index = states[chat_id]

# # # # # # #             # Save the user's response to the current question
# # # # # # #             answer = message.text
# # # # # # #             user_responses[questions[question_index]] = answer
# # # # # # #             states[chat_id] += 1  # Move to the next question

# # # # # # #             # Ask the next question
# # # # # # #             await ask_next_question(chat_id, question_index + 1)



# # # # # # if __name__ == "__main__":
# # # # # #     executor.start_polling(dispatcher, skip_updates=True)