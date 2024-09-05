#best code so far now it can upload files and receive various forms of messages as well and provide us graphs that we want 
# and also reply to images, give stock informations ,give stock analysis information and a prediction of the price

import os
import filetype
import docx
import PyPDF2
import re
from aiogram import Bot, Dispatcher, types
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma

import faiss
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import Docx2txtLoader

from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.memory import ConversationSummaryMemory
import asyncio
import numpy as np
import json
import re
import google.generativeai as genai
import pathlib
# Import things that are needed generically
from langchain.pydantic_v1 import BaseModel, Field
from langchain.tools import BaseTool, StructuredTool, tool

from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import CommandStart
from aiogram.types import Message
from aiogram import F
from aiogram import Router
import logging
import sys
from aiogram.filters import Command
from aiogram.types import FSInputFile
# from aiogram.utils import executor
import io
import matplotlib.pyplot as plt
import seaborn as sns
import aiohttp
from aiogram.types import InputFile , BufferedInputFile
import PIL.Image

router = Router(name=__name__)

load_dotenv()

TOKEN = os.getenv("TOKEN")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

os.environ["LANGCHAIN_TRACING_V2"]="true"
os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# Configure generativeai with your API key
genai.configure(api_key=GOOGLE_API_KEY)

# Initialize bot
bot = Bot(token=TOKEN)
dp = Dispatcher()

# Glbal variables
rag_on = False
retriever = None  # Store retriever globally
summary = ""
investment_personality = ""
# history = []
previous_suggestions = ""

CHAT_HISTORY_FILE = 'chat_history.json'

def read_chat_history(chat_id):
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, 'r') as file:
            chat_history = json.load(file)
            return chat_history.get(str(chat_id), [])
    return []

def write_chat_history(chat_id, message):
    chat_history = {}
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, 'r') as file:
            chat_history = json.load(file)
    if str(chat_id) not in chat_history:
        chat_history[str(chat_id)] = []
    chat_history[str(chat_id)].append(message)
    with open(CHAT_HISTORY_FILE, 'w') as file:
        json.dump(chat_history, file)

class Reference:
    def __init__(self):
        self.response = ""


reference = Reference()


def clear_past():
    reference.response = ""


@router.message(F.text == "clear")
async def clear(message: types.Message):
    """
    A handler to clear the previous conversation and context.
    """
    clear_past()
    await message.reply("I've cleared the past conversation and context.")

#Global Variables :

# Store user states
states = {}

# Dictionary to hold question-answer pairs
user_responses = {}
#
user_images = {}
# Define Questions for assessment
questions = [
    """ 
    1. You and your friend are betting on a series of coin tosses.

    He always bets ₹2,000 on Heads

    You always bet ₹2,000 on Tails

    Winner of last 8 turns

    You lost ₹8,000 in the last 4 turns!

    If you were to bet one last time, what would you bet on:
    a) heads or b) tails ?
    """ ,
    """
    2. Imagine you are a contestant in a game show, and you are presented the following choices.

    What would you prefer?
    a) 50 percent chance of winning 15 gold coins 
    b) 100 percent chance of winning 8 gold coins
    """,
    """
    3. In general, how would your best friend describe your risk-taking tendencies?
    a) A real gambler
    b) Willing to take risks after completing adequate research
    c) Cautious
    d) Avoids risk as much as possible
    """,
    """
    4. Suppose you could replace your current investment portfolio with this new one:
    50 percent chance of Gaining 35 percent or 50 percent chance of Loss
    In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
    a)-5 to -10
    b)-10 to -15
    c)-15 to -20
    d)-20 to -25
    e)-25 to -30
    f)-30 to -35
    """,
    """
    5. Over any 1-year period, what would be the maximum drop in the value of your investment 
    portfolio that you would be comfortable with?
    a) <5%
    b) 5 - 10%
    c) 10 - 15%
    d) 15 - 20%
    e) >20%
    """,
    """
    6. When investing, what do you consider the most?

    a) Risk 
    b) Return
    """,
    """
    7. What best describes your attitude?

    a) Prefer reasonable returns, can take reasonable risk
    b) Like higher returns, can take slightly higher risk
    c) Want to maximize returns, can take significant high risk
    """,
    """
    8. How much monthly investment you want to do?
    """,
    """
    9. What is the time horizon for your investment?
    You can answer in any range, example 1-5 years."""  
]



import logging
from aiogram import Bot, Dispatcher, types
# Register the router with the dispatcher
dp.include_router(router)

# from aiogram.utils import executor
from aiogram.filters import CommandStart
from aiogram.types import Poll, PollAnswer

# Command handler to start the poll
@dp.message(CommandStart())
async def handle_start(message: types.Message):
    chat_id = message.chat.id
    await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! ")
    question="How Can I Help you today ?"
    options = """\na. Know my Investment Personality \nb. Tax Related Quries \nc. Savings and Wealth Management \nd. Debt Repayment Strategies
              """
    await bot.send_message(chat_id, question + options)

    

# Function to start the assessment
async def start_assessment(chat_id):
    await bot.send_message(chat_id, """To analyse your investment personality I need to ask you some questions.\nLet's start a quick personality assessment.""")
    await ask_next_question(chat_id, 0)

# Function to ask the next question
async def ask_next_question(chat_id, question_index):
    if question_index < len(questions):
        # Ask the next question
        await bot.send_message(chat_id, questions[question_index])
        # Update state to indicate the next expected answer
        states[chat_id] = question_index
    else:
        # No more questions, finish assessment
        await finish_assessment(chat_id)

# Handler for receiving assessment answers
assessment_in_progress = True

from aiogram.types import FSInputFile
async def finish_assessment(chat_id):
    if chat_id in states and states[chat_id] == len(questions):
        # All questions have been answered, now process the assessment
        await bot.send_message(chat_id, "Assessment completed. Thank you!")

        # Determine investment personality based on collected responses
        global investment_personality
        investment_personality = await determine_investment_personality(user_responses)

        # Inform the user about their investment personality
        await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

        # Store the response in chat history
        write_chat_history(chat_id, {'role': 'bot', 'message': investment_personality})

        # Summarize collected information
        global summary
        summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
        summary = summary + "\n" + "Your investment personality:" + investment_personality
        # Ensure to await the determination of investment personality
        await send_summary_chunks(chat_id, summary)
        global assessment_in_progress 
        assessment_in_progress = False
       
        # await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
        # file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

        # await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
        # await bot.send_message(chat_id,file)

async def send_summary_chunks(chat_id, summary):
    # Split the summary into chunks that fit within Telegram's message limits
    chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

    # Send each chunk as a separate message
    for chunk in chunks:
        await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


async def determine_investment_personality(assessment_data):
    try:
        # Prepare input text for the chatbot based on assessment data
        input_text = "User Profile:\n"
        for question, answer in assessment_data.items():
            input_text += f"{question}: {answer}\n"

        # Introduce the chatbot's task and prompt for classification
        input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
                      "- Conservative Investor\n" \
                      "- Moderate Investor\n" \
                      "- Aggressive Investor"

        # Use your generative AI model to generate a response
        # print(input_text)
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(input_text)

        # Determine the investment personality from the chatbot's response
        response_text = response.text.lower()
        if "conservative" in response_text:
            personality = "Conservative Investor"
        elif "moderate" in response_text:
            personality = "Moderate Investor"
        elif "aggressive" in response_text:
            personality = "Aggressive Investor"
        else:
            personality = "Unknown"

        return personality
        # Send the determined investment personality back to the user
        #await bot.send_message(chat_id, f"Investment Personality: {personality}")

    except Exception as e:
        print(f"Error generating response: {e}")
        #await bot.send_message(chat_id, "Error processing investment personality classification.")


# Tax Related Queries :

# Define a global state to track the questions
tax_states = {}
tax_responses = {}

# Define the questions
tax_questions = [
    "What is your annual income?",
    "In which state do you live?",
    "Are you married or single?\n(a) Married\n(b) Single",
    "For which year do you wish to calculate tax?",
    "Do you have any mortgages or any tax reductions?"
]

# @dp.message()
async def tax_management(message: types.Message):
    chat_id = message.chat.id
    # # If the user is answering a tax question, process the response
    # if chat_id in tax_states and tax_states[chat_id] < len(tax_questions):
    #     question_index = tax_states[chat_id]
    #     answer = message.text
    #     tax_responses[tax_questions[question_index]] = answer
    #     tax_states[chat_id] += 1

    #     # Ask the next question
    #     if tax_states[chat_id] < len(tax_questions):
    #         await bot.send_message(chat_id, tax_questions[tax_states[chat_id]])
    #     else:
    #         # All questions answered, now process the data
    #         await calculate_taxes(chat_id)
    #         await bot.send_message(chat_id, "Thank you for your responses! Your tax-related queries have been processed.")
    #         # Reset the state
    #         del tax_states[chat_id]
    #         del tax_responses[chat_id]

    # # If the user starts the tax management flow, ask the first question
    # else:
    #     tax_states[chat_id] = 0
    #     tax_responses[chat_id] = {}
    #     await bot.send_message(chat_id, "Let's get started with your tax-related queries.")
    #     await bot.send_message(chat_id, tax_questions[0])

# async def calculate_taxes(chat_id):
    try:
        # Get the user's responses
        annual_income = "$100k" #tax_responses[tax_questions[0]]
        state = "texas" #tax_responses[tax_questions[1]]
        marital_status = "Married" #tax_responses[tax_questions[2]]
        tax_year = "2024" #tax_responses[tax_questions[3]]
        mortgages_or_deductions = "$200k" #tax_responses[tax_questions[4]]

        # Prepare the context for the LLM
        context = f"""
        Annual Income: {annual_income}
        State: {state}
        Marital Status: {marital_status}
        Tax Year: {tax_year}
        Mortgages or Deductions: {mortgages_or_deductions}
        """

        # Use the context to get tax calculation and advice
        task = """You are a Tax Calculations Expert in the entire world.
            Ask user tax related queries to help users with tax related queries.
            Consider user's investment personality  if provided.
            Address the user by their name(client_name: Emily in our case but if any other name is given refer to that) if provided.
            Help users to save tax on their income and earnings.
            If user asks queries related to saving taxes or calculating taxes refer to the 
            US Tax Laws given by the IRS and based on that information calculate the taxes for the user 
            consider the information shared by the user such as their annual income and their monthly investment if provided,
            also give advice to the user on how they can save their taxes.
            Include a disclaimer to monitor and rebalance investments regularly based on risk tolerance."""
        
        query = task + "\n" + context
        model = genai.GenerativeModel('gemini-1.5-flash')
        chat = model.start_chat(history=[])
        response = chat.send_message(query)

        # Enhanced logging for debugging
        logging.info(f"Model response: {response}")
        format_response = markdown_to_text(response.text)

        # Store the response in chat history
        write_chat_history(chat_id, {'role': 'bot', 'message': format_response})
        await bot.send_message(chat_id,"Here is your calculated tax as per the responses provided ")
        await bot.send_message(chat_id, format_response)

    except Exception as e:
        print(f"Error calculating taxes: {e}")
        await bot.send_message(chat_id, "Error calculating taxes. Please try again later.")

# async def tax_management(chat_id) :
#     try:
        # task = """You are a Tax Calculations Expert in the entire world.
        #     Ask user tax related queries to help users with tax related queries.
        #     Consider user's investment personality  if provided.
        #     Address the user by their name(client_name: Emily in our case but if any other name is given refer to that) if provided.
        #     Help users to save tax on their income and earnings.
        #     If user asks queries related to saving taxes or calculating taxes refer to the 
        #     US Tax Laws given by the IRS and based on that information calculate the taxes for the user 
        #     consider the information shared by the user such as their annual income and their monthly investment if provided,
        #     also give advice to the user on how they can save their taxes.
        #     Include a disclaimer to monitor and rebalance investments regularly based on risk tolerance."""
        
        
#         # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + message.text
#         # query = chat_history_text + "\n" + query

#         # Include chat history
#         chat_history = read_chat_history(chat_id)
#         chat_history_text = '\n'.join([f"{entry['role']}: {entry['message']}" for entry in chat_history])
#         # history.append(chat_history_text)

#         # query = task + "\n" + investment_personality + "\n" + chat_history_text # + "\n" + message.text
#         query = task + "\n" + chat_history_text # + "\n" + message.text


#         model = genai.GenerativeModel('gemini-1.5-flash')
#         chat = model.start_chat(history=[])
#         response = chat.send_message(query)

#         # Enhanced logging for debugging
#         logging.info(f"Model response: {response}")
#         format_response = markdown_to_text(response.text) #(response_text) #response.result

#         # Store the response in chat history
#         write_chat_history(chat_id, {'role': 'bot', 'message': format_response})
#         await bot.send_message(chat_id,format_response)
#         # await message.reply(format_response)

#     except Exception as e:
#         print(f"Error invoking retrieval chain on attempt : {e}")
#         await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")
        

# Savings and Wealth Management :
async def savings_management(chat_id) :
    # Savings and Wealth Management related questions
    await bot.send_message(chat_id,"Hello there here is a Simple Personal Budget Excel File.Please fill in your details with correct Information and then upload it in the chat")
    file = FSInputFile("data\Emily_Budget.xlsx", filename="Your Simple Personal Budget.xlsx")


# Handler for document upload
async def load_vector_db(file_path):
    try:
        print("Loading vector database...")
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
        vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

        # vector_store = FAISS(
        #     embedding_function=embeddings,
        #     index=index,
        #     docstore=InMemoryDocstore(),
        #     index_to_docstore_id={},
        # )
        
        print("Vector database loaded successfully.") 
        return vector_store.as_retriever(search_kwargs={"k": 1})
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None


# change prompt template :
async def make_retrieval_chain(retriever):
    """
    Create a retrieval chain using the provided retriever.

    Args:
        retriever (RetrievalQA): A retriever object.

    Returns:
        RetrievalQA: A retrieval chain object.
    """
    try:
        global investment_personality,summary
        llm = ChatGoogleGenerativeAI(
            #model="gemini-pro",
            model = "gemini-1.5-flash",
            temperature=0.7,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )

        # prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
        #         Respond to the client by the client name.
        #         Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
        #         Also give the user detailed information about the investment how to invest,where to invest and how much they
        #         should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
        #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
        #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
        #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
        #         Also explain the user why you are giving them that particular investment suggestion.
        #         Give the client suggestions of Investment based on their investment personality that can also help them manage their mortages and other liablilities if they have any,ignore if they dont have any liabilities.
        #         Answer in 3-4 lines.\n
        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""


        prompt_template = investment_personality + "\n" + summary + "\n" + """ Role : You are a Top class highly professional and world's best Savings Advisor for 
                savings related question-answering tasks related to the document.
                Respond to the client by the client name.
                Give Savings Suggestions to the user so that they could do proper responsible savings and save their expenses based on their investment personality and budget if provided.
                Also give the user detailed information about their savings such that they could save more money and save their expenses.
                Give the user minimum and maximum percentage of savings the user can do by reducing their expenses. If the users have given a budget then analyse it and give suggestions based on that.
                Try to imitate human language and talk to the user/client like a human and give personal savings suggestions.
                If the user is having many unnecessary expenses then give the user some advice in a gentle manner without offending the user or hurt their feelings and suggest and advice them to stop or reduce their unnecessary expenses 
                in order to increase their savings.
                Also explain the user why you are giving them that particular savings suggestion.
                Give the client suggestions of savings based on their investment personality that can also help them manage their mortages and other liablilities if they have any,ignore if they dont have any liabilities.
                Answer in 3-4 lines.\n
                <context>
                {context}
                </context>
                Question: {input}"""

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in creating chain: {e}")
        return None

from aiogram.filters import Filter

# @router.message(F.document)

import os
import pandas as pd
from openpyxl import load_workbook

@dp.message(F.document)
async def handle_document(message: types.Message):
    global summary, investment_personality  

    chat_id = message.chat.id
    await message.reply("File Received") 
    
    # Obtain file information
    file_id = message.document.file_id
    file = await bot.get_file(file_id)
    file_path = file.file_path

    # Get the file extension
    file_extension = os.path.splitext(message.document.file_name)[-1].lower()

    # Download the file
    local_file_path = "data/uploaded_file" + file_extension
    await bot.download_file(file_path, local_file_path)

    # Process the uploaded document based on the file type
    if file_extension in ['.xlsx', '.xls']:
        extracted_text = await process_excel_file(local_file_path)
    else:
        extracted_text = await process_document(local_file_path)

    if extracted_text:
        # Proceed with further processing
        print("Retriever being loaded ")
        retriever = await load_vector_db(local_file_path)
        client_name, validation_errors = await validate_process_document(local_file_path)

        print(f"Client Name: {client_name}")
        if validation_errors:
            print("**Validation Errors:**")
            for error in validation_errors:
                print(error)
        else:
            print("All fields are filled correctly.")
        # if client_name is None:
        #     try:
        #         await message.reply("Processing the uploaded image")
        #         await handle_image(message) 
        #         return 
        #     except Exception as e:
        #         await message.reply("Error processing uploaded image")
        #         print(e)
        if client_name == None : client_name = "Emilly"
        await message.reply(f"Thanks for providing me the details, {client_name}. I have processed the file and now I will provide you some Savings suggestions based on the details that you have provided.")

        if retriever is None:
            await message.reply("The retrieval chain is not set up. Please upload a document first.")
            return

        chain = await make_retrieval_chain(retriever)
        if chain is None:
            await message.reply("Failed to create the retrieval chain.")
            return

        try:     
            query = summary + "\n" + investment_personality
     
            response = chain.invoke({"input": query})
            print(response['answer'])
            global chat_history
            chat_history = response['answer'] 
            print(f"\n Chat History : {chat_history}")
            format_response = markdown_to_text(response['answer'])

            write_chat_history(chat_id, {'role': 'bot', 'message': extracted_text})
            write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

            await message.reply(format_response)

        except Exception as e:
            print(f"Error invoking retrieval chain on attempt : {e}")
            await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

    else:
        await message.reply("Failed to process the uploaded file.")

async def process_excel_file(file_path):
    try:
        # Reading the Excel file using pandas
        df = pd.read_excel(file_path)
        # Extracting relevant information (This is just an example, customize it as needed)
        extracted_text = df.to_string(index=False)
        return extracted_text
    except Exception as e:
        print(f"Error processing Excel file: {e}")
        return None


# @dp.message(F.document)
# async def handle_document(message: types.Message):
#     global summary,investment_personality  

#     chat_id = message.chat.id
#     await message.reply("File Received") 
#     # Obtain file information
#     file_id = message.document.file_id
#     file = await bot.get_file(file_id)
#     file_path = file.file_path
    
#     # Download the file
#     await bot.download_file(file_path, "data/uploaded_file")
    
#     # Process the uploaded document
#     extracted_text = await process_document("data/uploaded_file")
#     # print(extracted_text)

#     if extracted_text:
#         # Load vector database (assuming this is part of setting up the retriever)
#         print("Retriever being loaded ")
#         retriever = await load_vector_db("data/uploaded_file")
#         file_path = 'data/uploaded_file'
#         client_name, validation_errors = await validate_process_document(file_path)

#         # Print results
#         print(f"Client Name: {client_name}")
#         if validation_errors:
#             print("**Validation Errors:**")
#             for error in validation_errors:
#                 print(error)
#         else:
#             print("All fields are filled correctly.")
#         if client_name == None:
#             try:
#                 await message.reply("Processing the uploaded image")
#                 await handle_image(message) 
#                 return 
#             except Exception as e:
#                 await message.reply("error processing uploaded image")
#                 print(e)
#         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some Savings suggestions based on the details that you have provided.")

#         if retriever is None:
#             await message.reply("The retrieval chain is not set up. Please upload a document first.")
#             return

#         # Check if a valid chain can be created
#         chain = await make_retrieval_chain(retriever)
#         if chain is None:
#             await message.reply("Failed to create the retrieval chain.")
#             return
        
#         try:     
#             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
        
#             response = chain.invoke({"input": query})
#             print(response['answer'])
#             global chat_history
#             chat_history = response['answer'] 
#             print(f"\n Chat History : {chat_history}")
#             format_response = markdown_to_text(response['answer'])

#             # Store the extracted_text in chat history
#             write_chat_history(chat_id, {'role': 'bot', 'message': extracted_text})
        
#             # Store the response in chat history
#             write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

#             await message.reply(format_response)
#             # await message.reply(response['answer'])

#         except Exception as e:
#             print(f"Error invoking retrieval chain on attempt : {e}")
#             await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

#     else:
#         await message.reply("Failed to process the uploaded file.")
    

# Function to extract data from LLM response
def extract_data_from_response(response):
    try:
        # Locate the JSON-like data in the response
        json_start = response.find("{")
        json_end = response.rfind("}") + 1
        
        if json_start == -1 or json_end == -1:
            raise ValueError("No JSON data found in the response.")
        
        json_data = response[json_start:json_end]
        
        # Parse the JSON data
        data = json.loads(json_data.replace("'", "\""))
        print(data)
        return data
    except Exception as e:
        logging.error(f"Error extracting data: {e}")
        return None

 


def extract_allocations_from_json(json_data,chat_id):
    allocations = {}
    for entry in json_data.get(str(chat_id), []):
        if entry['role'] == 'bot':
            message = entry['message']
            lines = message.split('\n')
            current_category = None

            for line in lines:
                match = re.match(r'^(.*?):\s*(\d+)%$', line)
                if match:
                    category, percent = match.groups()
                    allocations[category] = []
                    current_category = category
                elif current_category and re.match(r'.*\d+%', line):
                    subcategory_match = re.match(r'^(.*?)(\d+)%$', line)
                    if subcategory_match:
                        subcategory, percent = subcategory_match.groups()
                        allocations[current_category].append((subcategory.strip(), float(percent)))

    return allocations


def create_pie_chart(allocations, chat_id):
    labels = []
    sizes = []
    for category, subcategories in allocations.items():
        for subcategory, percent in subcategories:
            labels.append(f"{category} - {subcategory}")
            sizes.append(percent)
    
    if sizes:
        fig, ax = plt.subplots()
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
        ax.axis('equal')
        
        plt.title("Investment Allocation")
        chart_path = f"data/investment_allocation_{chat_id}.png"
        plt.savefig(chart_path)
        plt.close()
        
        return chart_path
    else:
        return None
  

async def process_document(file_path):
    try:
        print("Processing the document")
        file_type = filetype.guess(file_path)
        if file_type is not None:
            if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # return extract_text_from_word(file_path)
                return extract_text_and_tables_from_word(file_path)
            elif file_type.mime == "application/pdf":
                return extract_text_from_pdf(file_path)
        return None
    except Exception as e:
        print(f"Error processing document: {e}")
        return None

def extract_text_from_pdf(pdf_file_path):
    try:
        print("Processing pdf file")
        with open(pdf_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            text_content = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text_content.append(page.extract_text())
            return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None


import re
import docx

def extract_text_and_tables_from_word(docx_file_path):
    """
    Extracts text and tables from a Word document (.docx).

    Args:
        docx_file_path (str): Path to the Word document file.

    Returns:
        tuple: Extracted text content and tables from the document.
    """
    try:
        print("Extracting text and tables from word file")
        doc = docx.Document(docx_file_path)
        text_content = []
        tables_content = []

        for para in doc.paragraphs:
            text_content.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)
        print("Extracted text from word file")
        return "\n".join(text_content), tables_content
    except Exception as e:
        print(f"Error extracting text and tables from Word document: {e}")
        return None, None

def validate_document_content(text, tables):
    """
    Validates the content of the document.

    Args:
        text (str): Extracted text content from the document.
        tables (list): Extracted tables content from the document.

    Returns:
        tuple: Client name and validation errors.
    """
    errors = []
    
    # Extract client name
    client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
    client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

    # Define required sections
    required_sections = [
        "YOUR RETIREMENT GOAL",
        "YOUR OTHER MAJOR GOALS",
        "YOUR ASSETS AND LIABILITIES",
        "MY LIABILITIES",
        "YOUR CURRENT ANNUAL INCOME"
    ]

    # Check for the presence of required sections
    for section in required_sections:
        if section not in text:
            errors.append(f"* {section} section missing.")
    
    # Define table field checks
    table_checks = {
        "YOUR RETIREMENT GOAL": [
            r"When do you plan to retire\? \(age or date\)",
            r"Social Security Benefit \(include expected start date\)",
            r"Pension Benefit \(include expected start date\)",
            r"Other Expected Income \(rental, part-time work, etc.\)",
            r"Estimated Annual Retirement Expense"
        ],
        "YOUR OTHER MAJOR GOALS": [
            r"GOAL", r"COST", r"WHEN"
        ],
        "YOUR ASSETS AND LIABILITIES": [
            r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
            r"Current Value", r"Annual Contributions"
        ],
        "MY LIABILITIES": [
            r"Balance", r"Interest Rate", r"Monthly Payment"
        ]
    }

    # Validate table content
    for section, checks in table_checks.items():
        section_found = False
        for table in tables:
            table_text = "\n".join(["\t".join(row) for row in table])
            if section in table_text:
                section_found = True
                for check in checks:
                    if not re.search(check, table_text, re.IGNORECASE):
                        errors.append(f"* Missing or empty field in {section} section: {check}")
                break
        if not section_found:
            errors.append(f"* {section} section missing.")

    return client_name, errors

async def validate_process_document(file_path):
    try:
        print("Validating process document : ")
        text, tables = extract_text_and_tables_from_word(file_path)
        if text is not None and tables is not None:
            client_name, errors = validate_document_content(text, tables)
            return client_name, errors
        return None, ["Error processing document."]
    except Exception as e:
        print(f"Error processing document: {e}")
        return None, [f"Error processing document: {e}"]



# somwheat working :
@dp.message()
async def main_bot(message: types.Message):
    global retriever, extracted_text, investment_personality, summary, chat_history

    chat_id = message.chat.id
    question="How Can I Help you today ?"
    options = """\na. Know my Investment Personality \nb. Tax Related Queries \nc. Savings and Wealth Management \nd. Debt Repayment Strategies
              """
    
    # if chat_id in states and states[chat_id] < len(questions):
    #     question_index = states[chat_id]
    #     answer = message.text
    #     user_responses[questions[question_index]] = answer
    #     states[chat_id] += 1
    #     await ask_next_question(chat_id, question_index + 1)

    if chat_id in states and states[chat_id] < len(questions):
        question_index = states[chat_id]
        answer = message.text
        user_responses[questions[question_index]] = answer
        states[chat_id] += 1
        if states[chat_id] < len(questions):
            await ask_next_question(chat_id, question_index + 1)
        else:
            await ask_next_question(chat_id, question_index + 1)
            
            # await bot.send_message(chat_id, "Assessment Completed.")
            await bot.send_message(chat_id, "What do you want to do next?\n" + question + options)


    elif message.text:
        lower_text = message.text.lower()

        # Investment Personality :
        if any(variant in lower_text for variant in ["a", "a.", "a)", "(a)", "1", "1.", "1)", "(1)"]):
            await start_assessment(chat_id)

        # Tax Related Queries :
        elif any(variant in lower_text for variant in ["b", "b.", "b)", "(b)", "2", "2.", "2)", "(2)"]):
            # await bot.send_message(chat_id,"Hello, I will ask you some Tax Related Questions.\nPlease answer them correctly so that I can calculate your tax")
            await tax_management(message) #(chat_id) #pass
            await bot.send_message(chat_id, "What do you want to do next?\n" + question + options)
        
        # Savings and Wealth Management :
        elif any(variant in lower_text for variant in ["c", "c.", "c)", "(c)", "3", "3.", "3)", "(3)"]):
            await savings_management(chat_id) #pass  #  
        
        # Debt Repayment Strategies :
        elif any(variant in lower_text for variant in ["d", "d.", "d)", "(d)", "4", "4.", "4)", "(4)"]):
            pass  # 

        elif lower_text in ["yes", "y"]:
            await start_assessment(chat_id)

        else:
            await bot.send_message(chat_id, "Assessment Completed. Do you wish to retake the assessment? Type 'yes' or 'no'.")
            await bot.send_message(chat_id, "Thank you for your response.")
            await bot.send_message(chat_id, "What do you want to do next?\n" + question + options)

            try:
                task = """You are a Financial Expert and Wealth Advisor.
                    You also a Stock Market Expert. You know everything about stock market trends and patterns.
                    Provide financial advice or Stock Related advice and suggestions based on the user's query.
                    Consider user's investment personality and Financial Details if provided.
                    Address the user by their name(client_name: Emily in our case but if any other name is give refer to that) if provided.
                    Include detailed information about the investment, where to invest, how much to invest, 
                    expected returns, and why you are giving this advice.
                    As you are a Wealth Advisor if user asks queries related to saving taxes or calculating taxes refer to the 
                    US Tax Laws given by the IRS and based on that information calculate the taxes for the user 
                    consider the information shared by the user such as their annual income and their monthly investment if provided,
                    also give advice to the user on how they can save their taxes.
                    Include a disclaimer to monitor and rebalance investments regularly based on risk tolerance."""
                
                # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + message.text
                # query = chat_history_text + "\n" + query

                # Include chat history
                chat_history = read_chat_history(chat_id)
                chat_history_text = '\n'.join([f"{entry['role']}: {entry['message']}" for entry in chat_history])
                # history.append(chat_history_text)
                query = task + "\n" + investment_personality + "\n" + chat_history_text + "\n" + message.text

                model = genai.GenerativeModel('gemini-1.5-flash')
                chat = model.start_chat(history=[])
                response = chat.send_message(query)

                # Enhanced logging for debugging
                logging.info(f"Model response: {response}")
                format_response = markdown_to_text(response.text) #(response_text) #response.result

                # Store the response in chat history
                write_chat_history(chat_id, {'role': 'bot', 'message': format_response})
                await message.reply(format_response)

            except Exception as e:
                logging.error(f"Error processing general chat message: {e}")
                await message.reply("Failed to process your request.")



# markdown to text :
def markdown_to_text(md):
    # Simple conversion for markdown to plain text
    md = md.replace('**', '')
    md = md.replace('*', '')
    md = md.replace('_', '')
    md = md.replace('#', '')
    md = md.replace('`', '')
    return md.strip()


from aiogram.types.input_file import BufferedInputFile
from aiogram import BaseMiddleware
# from aiogram.dispatcher.router import Router
from PIL import Image

# Function to handle image messages
# @dp.message(F.photo)
# @router.message(F.photo)
import PIL.Image

async def handle_image(message: types.Message):
    global investment_personality, chat_history

    chat_id = message.chat.id
    # Handle image inputs
    try:
        # Obtain file information
        try:
            photo_id = message.document.file_id
            photo = await bot.get_file(photo_id)
            photo_path = photo.file_path
            # Download the file
            photo_file = await bot.download_file(photo_path, "data/uploaded_image.png")

        except Exception as e:
            print(f"Error downloading image: {e}")
            await bot.send_message(chat_id, "Error processing image. Please try again.")
            return
        
        # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
        #             Also give the user detailed information about the investment how to invest, where to invest and how much they
        #             should invest in terms of percentage of their investment amount. Give the user detailed information about the returns on their 
        #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compounded returns on their 
        #             investment. Also explain the user why you are giving them that particular
        #             investment suggestion. Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
        #             User should also invest as per their risk tolerance level. Since you are the financial advisor don't ask user to consult anyone else.
        #             So don't mention user to consult to a financial expert."""

        task = """You are a Financial Expert.You will be provided with a Financial Form from Boston Harbor.
                If you recieve any other image tell the user to Upload the Images of the form or upload the word document of the form.
                You are supposed to Respond to the user's Image query and If they ask for any information provide them the information in Detail.
                Be helpful and informative.Give proper information of any Financial terms the user may ask you.Address the user by their Client Name if provided.
                Also provide the user helpful links so that they can refer to the link for more information.
                If the image provided is not related to Finance then just answer about the image and any caption if provided.
                """

        prompt = message.caption if message.caption else ""  # Use the photo caption if available
        # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + prompt
        query = task + prompt 

        image =  PIL.Image.open('data/uploaded_image.png') #(photo_file) 
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(image)
        await bot.send_message(chat_id,"I will describe the image that was uploaded")
        format_response = markdown_to_text(response.text)
        await message.reply(format_response)
        # await message.reply(response.text)

        # chat = model.start_chat(history=[])
        # response = chat.send_message(query)
        # format_response = markdown_to_text(response.result)
        # await message.reply(format_response)

        response = model.generate_content([query, image])
        format_response = markdown_to_text(response.text)

        # Store the response in chat history
        write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

        await message.reply(format_response)
        # await message.reply(response.text) 
    except Exception as e:
        logging.error(f"Error generating response for the image: {e}")
        await message.reply("There was an error generating response for the image. Please try again later.")
    # await message.reply("Cant process the image")
    # return



from aiogram.filters import command
from aiogram.types import bot_command
import markdown
from bs4 import BeautifulSoup

def markdown_to_text(markdown_text):
    # Convert markdown to HTML
    html = markdown.markdown(markdown_text)
    # Parse the HTML
    soup = BeautifulSoup(html, 'html.parser')
    # Extract plain text
    text = soup.get_text()
    return text


# if __name__ == "__main__":
#     executor.start_polling(dispatcher, skip_updates=True)

async def main() -> None:
    # Initialize Bot instance with default bot properties which will be passed to all API calls
    bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

    # And the run events dispatching
    await dp.start_polling(bot)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, stream=sys.stdout)
    asyncio.run(main())








