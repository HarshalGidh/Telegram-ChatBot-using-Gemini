# Latest version of the code : 

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt

import os
import filetype
import docx
import PyPDF2
import re
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.memory import ConversationSummaryMemory
import asyncio
import numpy as np
import json

import google.generativeai as genai
import pathlib
import logging
import sys
import io
import matplotlib.pyplot as plt
import seaborn as sns
# Import things that are needed generically
from langchain.pydantic_v1 import BaseModel, Field
from langchain.tools import BaseTool, StructuredTool, tool
# Define functions to generate investment suggestions :

load_dotenv()
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

from flask import Flask, request, jsonify

app = Flask(__name__)

# Configure generativeai with your API key
genai.configure(api_key=GOOGLE_API_KEY)

import markdown
# def convert_to_markdown(raw_text):
#     # Replace specific text patterns with markdown syntax
#     formatted_text = raw_text.replace('\n', '\n\n')  # Ensure newlines create paragraphs
    
#     # Convert text into markdown format
#     html = markdown.markdown(formatted_text)

#     return html


import markdown2
from bs4 import BeautifulSoup

def markdown_to_readable_text(md_text):
    # Convert markdown to HTML
    html = markdown2.markdown(md_text)

    # Parse the HTML
    soup = BeautifulSoup(html, "html.parser")

    # Function to format plain text from tags
    def format_text_from_html(soup):
        formatted_text = ''
        for element in soup:
            if element.name == "h1":
                formatted_text += f"\n\n# {element.text.upper()} #\n\n"
            elif element.name == "h2":
                formatted_text += f"\n\n## {element.text} ##\n\n"
            elif element.name == "h3":
                formatted_text += f"\n\n### {element.text} ###\n\n"
            elif element.name == "strong":
                formatted_text += f"**{element.text}**"
            elif element.name == "em":
                formatted_text += f"_{element.text}_"
            elif element.name == "ul":
                for li in element.find_all("li"):
                    formatted_text += f"\n - {li.text}"
            elif element.name == "ol":
                for idx, li in enumerate(element.find_all("li"), 1):
                    formatted_text += f"\n {idx}. {li.text}"
            elif element.name == "table":
                rows = element.find_all("tr")
                for row in rows:
                    cols = row.find_all(["th", "td"])
                    row_text = ' | '.join(col.text.strip() for col in cols)
                    formatted_text += f"{row_text}\n"
                formatted_text += "\n"
            else:
                formatted_text += element.text

        return formatted_text.strip()

    return format_text_from_html(soup)

def markdown_to_text(md): # og solution code 
    # Simple conversion for markdown to plain text
    md = md.replace('**', '')
    md = md.replace('*', '')
    md = md.replace('_', '')
    md = md.replace('#', '')
    md = md.replace('`', '')
    return md.strip()


# import docx

# def extract_responses_from_docx(personality_file):
#     """
#     Extracts responses from a Word document (.docx) where answers are typed in.

#     Args:
#         personality_file (UploadedFile): The file object uploaded via Streamlit.

#     Returns:
#         dict: A dictionary containing the questions and the typed answers.
#     """
#     try:
#         doc = docx.Document(personality_file)
#         responses = {}
#         current_question = None

#         # Check paragraphs
#         for para in doc.paragraphs:
#             text = para.text.strip()
#             if text:
#                 # Check if the paragraph contains a question
#                 if "?" in text or text.endswith(":"):
#                     current_question = text
#                     st.write(f"Identified question: {current_question}")  # Debugging log
#                 else:
#                     # This is a typed answer
#                     typed_answer = text.strip()
#                     st.write(f"Identified typed answer: {typed_answer}")  # Debugging log
#                     if current_question:
#                         # If the question already has an answer, append to it (handles multiple responses)
#                         if current_question in responses:
#                             responses[current_question] += "; " + typed_answer
#                         else:
#                             responses[current_question] = typed_answer

#             # Debugging log to understand document structure
#             st.write(f"Processing paragraph: {text}")  # Console log for local testing

#         # Check tables for additional responses
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     text = cell.text.strip()
#                     if text:
#                         if "?" in text or text.endswith(":"):
#                             current_question = text
#                             st.write(f"Identified question in table: {current_question}")  # Debugging log
#                         else:
#                             typed_answer = text.strip()
#                             st.write(f"Identified typed answer in table: {typed_answer}")  # Debugging log
#                             if current_question:
#                                 if current_question in responses:
#                                     responses[current_question] += "; " + typed_answer
#                                 else:
#                                     responses[current_question] = typed_answer

#         if responses:
#             st.write("Extracted Responses:")
#             for question, answer in responses.items():
#                 st.write(f"**{question}**: {answer}")
#         else:
#             st.write("No responses captured. Please check the document formatting or symbols used.")

#         return responses

#     except Exception as e:
#         st.write(f"Error extracting responses: {e}")  # Console log for local testing
#         return None

# def determine_investment_personality(responses):
#     """
#     Determines the investment personality based on extracted responses.

#     Args:
#         responses (dict): A dictionary containing the questions and the selected answers.

#     Returns:
#         str: The determined investment personality.
#     """
#     try:
#         # Prepare input text for the chatbot based on extracted responses
#         input_text = "User Profile:\n"
#         for question, response in responses.items():
#             input_text += f"{question}: {response}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYour task is to determine the investment personality based on the above profile."

#         # Here you would send the input_text to your chatbot or classification model
#         # For demonstration, we'll just return the input_text
#         return input_text

#     except Exception as e:
#         st.write(f"Error determining investment personality: {e}")  # Console log for local testing
#         return None

# def extract_responses_from_docx(personality_file):
#     try:
#         doc = docx.Document(personality_file)
#         responses = {}
#         current_question = None

#         # Check paragraphs
#         for para in doc.paragraphs:
#             text = para.text.strip()
#             if text:
#                 # Check if the paragraph contains a question
#                 if "?" in text or text.endswith(":"):
#                     current_question = text
#                 else:
#                     # This is a typed answer
#                     typed_answer = text.strip()
#                     if current_question:
#                         # If the question already has an answer, append to it (handles multiple responses)
#                         if current_question in responses:
#                             responses[current_question] += "; " + typed_answer
#                         else:
#                             responses[current_question] = typed_answer

#         # Check tables for additional responses
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     text = cell.text.strip()
#                     if text:
#                         if "?" in text or text.endswith(":"):
#                             current_question = text
#                         else:
#                             typed_answer = text.strip()
#                             if current_question:
#                                 if current_question in responses:
#                                     responses[current_question] += "; " + typed_answer
#                                 else:
#                                     responses[current_question] = typed_answer

#         return responses

#     except Exception as e:
#         print(f"Error extracting responses: {e}")
#         return None

import docx

# # GET Method for me POST method for Frontend
def extract_responses_from_docx(personality_file): # Using text responses parsing
    """
    Extracts responses from a Word document (.docx) where the selected answers are listed as text after the options.

    Args:
        personality_file (str): Path to the Word document file.

    Returns:
        dict: A dictionary containing the questions and the selected answers.
    """
    try:
        doc = docx.Document(personality_file)
        responses = {}
        current_question = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect the beginning of a question
                if "?" in text:
                    current_question = text
                # Detect a chosen response (assuming it follows the question and options)
                elif current_question and not text.startswith(("a.", "b.", "c.", "d.")):
                    selected_answer = text
                    responses[current_question] = selected_answer
                    current_question = None  # Reset for the next question

        if responses:
            print(responses)
            # st.write(responses)
        else:
            print("\nNo responses captured")
            st.write("No responses captured")
        return responses
    except Exception as e:
        print(f"Error extracting responses: {e}")
        return None

# def extract_responses_from_assessment(personality_file): # using boxes
#     # Load the document
#     # doc = Document(docx_filename)
#     doc = docx.Document(personality_file)
    
#     # Initialize a list to store responses
#     responses = []
    
#     # Iterate through each paragraph in the document
#     for para in doc.paragraphs:
#         text = para.text.strip()
#         # Check if the paragraph contains a checkbox
#         if '☒' in text or '☐' in text:
#             # Extract the response marked with ☒
#             if '☒' in text:
#                 response = text.split('☒')[1].strip()
#                 responses.append(response)
    
#     return responses

# import asyncio
# # from some_generative_ai_library import GenerativeModel  # Replace with actual import

# async def determine_investment_personality(assessment_data):
#     try:
#         # Prepare input text for the chatbot based on assessment data
#         input_text = "User Profile:\n"
#         for question, answer in assessment_data.items():
#             input_text += f"{question}: {answer}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
#                       "- Conservative Investor\n" \
#                       "- Moderate Investor\n" \
#                       "- Aggressive Investor\n\n" \
#                       "Please provide the classification below:\n"

#         # Use your generative AI model to generate a response
#         model = GenerativeModel('gemini-1.5-flash')
#         response = await model.generate_content(input_text)

#         # Determine the investment personality from the chatbot's response
#         response_text = response.text.lower()

#         if "conservative investor" in response_text:
#             personality = "Conservative Investor"
#         elif "moderate investor" in response_text:
#             personality = "Moderate Investor"
#         elif "aggressive investor" in response_text:
#             personality = "Aggressive Investor"
#         else:
#             personality = "Unknown"

#         return personality
#     except Exception as e:
#         print(f"Error generating response: {e}")
#         return "Unknown"


# GET Method
async def determine_investment_personality(assessment_data): # proper code 
    try:
        # Prepare input text for the chatbot based on assessment data
        input_text = "User Profile:\n"
        for question, answer in assessment_data.items():
            input_text += f"{question}: {answer}\n"

        # Introduce the chatbot's task and prompt for classification
        input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
                      "- Conservative Investor\n" \
                      "- Moderate Investor\n" \
                      "- Aggressive Investor\n\n" \
                      "Please provide the classification below:\n"

        # Use your generative AI model to generate a response
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(input_text)

        # Determine the investment personality from the chatbot's response
        response_text = response.text.lower()

        if "conservative investor" in response_text:
            personality = "Conservative Investor"
        elif "moderate investor" in response_text:
            personality = "Moderate Investor"
        elif "aggressive investor" in response_text:
            personality = "Aggressive Investor"
        else:
            personality = "Unknown"

        return personality
    except Exception as e:
        print(f"Error generating response: {e}")
        return "Unknown"




#Load the Vector DataBase : # current version :
async def load_vector_db(file_path): # # GET Method 
    try:
        print("Loading vector database...")
        # file_path = os.path.basename(file_path)
        
        # Verify the file path
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        print(f"File path: {file_path}")
        
        # Check file permissions
        if not os.access(file_path, os.R_OK):
            raise PermissionError(f"File is not readable: {file_path}")
        
        # print(file_path)
        
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
        vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

        # vector_store = FAISS(
        #     embedding_function=embeddings,
        #     index=index,
        #     docstore=InMemoryDocstore(),
        #     index_to_docstore_id={},
        # )
        
        print("Vector database loaded successfully.") 
        return vector_store.as_retriever(search_kwargs={"k": 1})
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None

# import os

# async def load_vector_db(file_storage): 
#     try:
#         # Define the destination folder and ensure it exists
#         destination_folder = 'path/to/your/destination/folder'
#         if not os.path.exists(destination_folder):
#             os.makedirs(destination_folder)
        
#         # Construct the destination file path
#         file_path = os.path.join(destination_folder, file_storage.filename)
        
#         # Save the file to the destination folder
#         file_storage.save(file_path)
        
#         print("Loading vector database...")
#         print(f"File path: {file_path}")
        
#         loader = Docx2txtLoader(file_path)
#         documents = loader.load()
        
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         text_chunks = text_splitter.split_documents(documents)
        
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
#         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        
#         print("Vector database loaded successfully.") 
#         return vector_store.as_retriever(search_kwargs={"k": 1})
#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None



# investment_personality = "Moderate Investor"
async def make_retrieval_chain(retriever,investmentPersonality,clientName,monthly_investment=10000,investment_period=3): # GET Method
    """
    Create a retrieval chain using the provided retriever.

    Args:
        retriever (RetrievalQA): A retriever object.

    Returns:
        RetrievalQA: A retrieval chain object.
    """
    try:
        # global investment_personality #,summary
        
        print(f"{retriever}\n {investmentPersonality}\n {clientName}\n {monthly_investment}")
        # try:
        #     print(type(investmentPersonality))
        # except Exception as e:
        #     print(f"Error in personality: {e}")
        #     return None
        
        # print(clientName)
        
        llm = ChatGoogleGenerativeAI(
            #model="gemini-pro",
            model = "gemini-1.5-flash",
            temperature=0.7,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )
        # New Template 
        investment_period = str(investment_period)
        print(investmentPersonality)
        monthly_investment = str(monthly_investment)
        print(monthly_investment)
        print(investment_period)
        
        # New Prompt Template :
        
        prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
                Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
                Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
                Also give the user detailed information about the investment how to invest,where to invest and how much they
                should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
                Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
                investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
                Also give the user minimum and maximum expected growth in dollars for the time horizon .
                Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
                
                You are a Financial Advisor for question-answering tasks related to the document. Based on the client's investment personality and financial details provided, generate responsible investment suggestions to achieve their financial goals while managing debts.

                Step-by-Step Guidance:
                1. Assets: Calculate total assets by analyzing the provided financial document in the My Assets section. Ensure you include cash, real estate, retirement accounts, brokerage accounts, and any other relevant asset types from the document.
                2. Liabilities: Calculate total liabilities by analyzing the provided financial document in the My Liabilities section. Consider mortgages, credit card debts, student loans, car loans, and other liabilities. 
                3. Monthly Investment Feasibility: Use the client's assets and liabilities to assess whether their planned monthly investment is feasible. If not feasible, suggest a more realistic monthly investment amount.
                4. Analyze Liabilities: Determine if the client's monthly investment plan is feasible after covering liabilities and expected expenses and also considering some amount for savings. If the client's monthly investment plan is not feasible after covering expenses and savings, generate investment suggestions on a smaller monthly investment plan amount if it can help the client else mention amount is too small for the client's requirementys to be made.
                5. Investment Strategy: Suggest a strategy where monthly investments can both generate returns and pay off debts effectively and helps client to achieve their financial goals.
                6. Allocation: Provide detailed allocations between growth-oriented investments and conservative investments, ensuring the client can meet their monthly debt obligations and save for their future financial goals.
                7. Returns: Include minimum and maximum compounded returns over 5-10 years, along with inflation-adjusted returns for clarity.
                8. Suggestions: Offer advice on how to use remaining funds to build wealth after clearing liabilities and achive their financial goal.
                
                
                Here's an example for the required Output Format(if there are comments indicated by # in the example output format then thats a side note for your reference dont write it in the response that will be generated ) :
                
                Investment Suggestions : 
                
                
                Client Name: """ + clientName + """

                Financial Overview: (#the data presented is just an example for your reference do not consider it as factual refere to the document provided to you and generate data based on the provided data and only when nothing is provided assume some data for analysis)
                
                - Total Assets: (# Sum of all client assets and Annual Income . Mention all assets and their respected values.if non consider the example assets)
                
                - Total Liabilities: (# Sum of all liabilities. Mention all liabilities and their respected values if non consider the example liabilities)
                
                
                - Monthly Liabilities: (# Monthly payments derived from liabilities)
                
                - Total Annual Income : (# Sum of all client's anual income)
                
                - Monthly Investment Amount : """ + monthly_investment + """ (# if no specific amount is specified to you then only assume  10,000 else consider the amount mention to you and just display the amount)
                
                - Investment Period : """ + investment_period + """  (# if no specific period is specified to you then only assume 3 years else consider the period mention to you and just display the period)

                Financial Analysis :(#Analyse the assets and liabilities and based on that give a suggestion for analysis generate suggestions for one of the following conditions:)
                (#1st condition : Everything is Positive)Based on the given Financial Conditions the client is having a good and stable income with great assets and manageable debt and liabilities.
                Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the clients monthly income.
                (# if this condition is true then ignore the other conditions and start with the Investment Suggestions)
                
                (#2nd condition : Everything is temporarily Negative) Based on the given Financial Conditions the client is facing a bad income for now but have great assets and manageable debt and liabilities.
                Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the client's monthly income but the client might not be able to sustain the monthly investment amount that they are planning.)
                Instead I would like to recommend this amount to the client for their monthly investment : (#Mention a feasible amount to the client for monthly investment and start suggesting investments based on this amount and not the previous amount being taken into consideration)
                
                (#3rd condition : Everything is Negative) Based on the given Financial Conditions the client is facing a bad income and doesnt have good assets to manage the debts and liabilities of the client and in such a condition this monthly investment amount is not feasible.
                Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is not manageable for the client's monthly income and so the client might not be able to sustain the monthly investment amount that they are planning to do.)
                I would like to recommend this amount to the client for monthly investment : (# Mention a minimum amount to the client for monthly investment if possible else just say the client should first prioritize on savings and generating more income to manage their debts and liabilities first and so dont give any investment suggestions to the client.)
                
                (#If the financial is 1 or 2 only then give investment suggestions to the client)
                
                
                Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

                Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

                Investment Allocation: (#remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

                Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
                How to Invest: Diversify across various asset classes like:  (#Give allocations % as well)
                
                Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                
                ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                
                Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                
                Consider investing in blue-chip companies or growth sectors like technology. 
                
                
                Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


                Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
                Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

                Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
                How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
                Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
                Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
                Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
                Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
                Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
                Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.


                Time Horizon and Expected Returns:

                Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
                Minimum Expected Annual Return: 4% - 6% 
                
                
                Maximum Expected Annual Return: 8% - 10% 
                
                
                Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, (# consider the monthly investment amount and give returns based on that only) $10,000 could grow to approximately 17,908 in 10 years.
                Minimum Expected Growth in Dollars: 
                
                4,000−6,000 (over 10 years) 
                
                
                Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
                Inflation Adjusted Returns:(#do not write this part inside the bracket just give answer,assume US inflation rate assume 3% if you dont know, and give the investment returns value that was suggested by you for the considered monthly investment amount after 3,5,10years of growth mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
                Rationale for Investment Suggestions:

                This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
                Important Considerations:

                Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

                Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

                Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
                <context>
                {context}
                </context>
                Question: {input}"""
    
        print("Retriever Created ")
        print(f"Investment Personality :{investmentPersonality}")
        
        # Good but inaccurate in terms of assets and liability data calculation
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
        #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
        #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
        #         Also give the user detailed information about the investment how to invest,where to invest and how much they
        #         should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
        #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
        #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
        #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
        #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
                
        #         You are a Financial Advisor for question-answering tasks related to the document. Based on the client's investment personality and financial details provided, generate responsible investment suggestions to achieve their financial goals while managing debts.

        #         Step-by-Step Guidance:
        #         1. Analyze Liabilities: Determine if the client's monthly investment plan is feasible after covering liabilities and expected expenses and also considering some amount for savings. If the client's monthly investment plan is not feasible after covering expenses and savings, generate investment suggestions on a smaller monthly investment plan amount if it can help the client else mention amount is too small for the client's requirementys to be made.
        #         2. Investment Strategy: Suggest a strategy where monthly investments can both generate returns and pay off debts effectively and helps client to achieve their financial goals.
        #         3. Allocation: Provide detailed allocations between growth-oriented investments and conservative investments, ensuring the client can meet their monthly debt obligations and save for their future financial goals.
        #         4. Returns: Include minimum and maximum compounded returns over 5-10 years, along with inflation-adjusted returns for clarity.
        #         5. Suggestions: Offer advice on how to use remaining funds to build wealth after clearing liabilities.
                
                
        #         Here's an example for the required Output Format(if there are comments indicated by # in the example output format then thats a side note for your reference dont write it in the response that will be generated ) :
                
        #         Investment Suggestions : 
                
                
        #         Client Name: """ + clientName + """

        #         Financial Overview: (#the data presented is just an example for your reference do not consider it as factual refere to the document provided to you and generate data based on the provided data only)
                
        #         - Total Assets: (# Consider the data available to you and use it  for display. For ex : $100,000 (cash), $150,000 (home), $12,000 (other assets) )
                
        #         - Liabilities: (# Consider the data available to you and use it  for display. For ex : $200,000 mortgage at 12% interest, $400 credit card debt at 3.5%, $15,000 other loans at 10%.)
                
        #         - Total Monthly Liabilities: (# Consider the data available to you and use it  for display. For ex : $1,650 (Mortgage, Credit Card, Other Loans).)
                
        #         - Monthly Investment Amount : """ + monthly_investment + """ (# if no specific amount is specified to you then only assume  10,000 else consider the amount mention to you and just display the amount)
                
        #         - Investment Period : """ + investment_period + """  (# if no specific period is specified to you then only assume 3 years else consider the period mention to you and just display the period)

        #         Financial Analysis :(#Analyse the assets and liabilities and based on that give a suggestion for analysis generate suggestions for one of the following conditions:)
        #         (#1st condition : Everything is Positive)Based on the given Financial Conditions the client is having a good and stable income with great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the clients monthly income.
        #         (# if this condition is true then ignore the other conditions and start with the Investment Suggestions)
                
        #         (#2nd condition : Everything is temporarily Negative) Based on the given Financial Conditions the client is facing a bad income for now but have great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the client's monthly income but the client might not be able to sustain the monthly investment amount that they are planning.)
        #         Instead I would like to recommend this amount to the client for their monthly investment : (#Mention a feasible amount to the client for monthly investment and start suggesting investments based on this amount and not the previous amount being taken into consideration)
                
        #         (#3rd condition : Everything is Negative) Based on the given Financial Conditions the client is facing a bad income and doesnt have good assets to manage the debts and liabilities of the client and in such a condition this monthly investment amount is not feasible.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is not manageable for the client's monthly income and so the client might not be able to sustain the monthly investment amount that they are planning to do.)
        #         I would like to recommend this amount to the client for monthly investment : (# Mention a minimum amount to the client for monthly investment if possible else just say the client should first prioritize on savings and generating more income to manage their debts and liabilities first and so dont give any investment suggestions to the client.)
                
        #         (#If the financial is 1 or 2 only then give investment suggestions to the client)
                
                
        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

        #         Investment Allocation: (#remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (#Give allocations % as well)
                
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                
        #         Consider investing in blue-chip companies or growth sectors like technology. 
                
                
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.


        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, (# consider the monthly investment amount and give returns based on that only) $10,000 could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(#do not write this part inside the bracket just give answer,assume US inflation rate assume 3% if you dont know, and give the investment returns value that was suggested by you for the considered monthly investment amount after 3,5,10years of growth mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""
        
        
                
        # # latest version gives some info about the client conditions but not in detail 
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
        #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
        #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
        #         Also give the user detailed information about the investment how to invest,where to invest and how much they
        #         should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
        #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
        #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
        #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
        #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
        #         Here's an example for the required Output Format :

        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

        #         Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
        #         Consider investing in blue-chip companies or growth sectors like technology. 
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, a 10,000 investment could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(do not write this part inside the bracket just give answer,assume US inflation rate, and give the investment returns value that was suggested by you  for $10k investment after 3,5,10years of growth  mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""
        
        
        
        
        # # Working code but gives clientname in brackets 
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
        #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality provided to you.
        #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
        #         Also give the user detailed information about the investment how to invest,where to invest and how much they
        #         should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
        #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
        #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
        #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
        #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
        #         Here's an example for the required Output Format :

        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

        #         Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
        #         Consider investing in blue-chip companies or growth sectors like technology. 
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, a 10,000 investment could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(do not write this part inside the bracket just give answer,assume US inflation rate, and give the investment returns value that was suggested by you  for $10k investment after 3,5,10years of growth  mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""
                

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            # print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in creating chain: {e}")
        return None


import json
import io


async def process_document(file_path): # GET Method
    try:
        print("Processing the document")
        file_type = filetype.guess(file_path)
        if file_type is not None:
            if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Await the coroutine to extract text and tables
                return await extract_text_and_tables_from_word(file_path)
            elif file_type.mime == "application/pdf":
                return await extract_text_from_pdf(file_path)
        return None
    except Exception as e:
        print(f"Error processing document: {e}")
        return None

# Async function to extract text from a PDF file
async def extract_text_from_pdf(pdf_file_path): # GET Method
    try:
        print("Processing pdf file")
        with open(pdf_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            text_content = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text_content.append(page.extract_text())
            return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

# Async function to extract text and tables from a Word document
async def extract_text_and_tables_from_word(docx_file_path): # GET Method
    try:
        print("Extracting text and tables from word file")
        doc = docx.Document(docx_file_path)
        text_content = []
        tables_content = []

        for para in doc.paragraphs:
            text_content.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)
        print("Extracted text from word file")
        return "\n".join(text_content), tables_content
    except Exception as e:
        print(f"Error extracting text and tables from Word document: {e}")
        return None, None



async def validate_document_content(text, tables):
    """
    Validates the content of the document.

    Args:
        text (str): Extracted text content from the document.
        tables (list): Extracted tables content from the document.

    Returns:
        tuple: Client name and validation errors.
    """
    errors = []
    
    # Extract client name
    client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
    client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

    # Define required sections
    required_sections = [
        "YOUR RETIREMENT GOAL",
        "YOUR OTHER MAJOR GOALS",
        "YOUR ASSETS AND LIABILITIES",
        "MY LIABILITIES",
        "YOUR CURRENT ANNUAL INCOME"
    ]

    # Check for the presence of required sections
    for section in required_sections:
        if section not in text:
            errors.append(f"* {section} section missing.")
    
    # Define table field checks
    table_checks = {
        "YOUR RETIREMENT GOAL": [
            r"When do you plan to retire\? \(age or date\)",
            r"Social Security Benefit \(include expected start date\)",
            r"Pension Benefit \(include expected start date\)",
            r"Other Expected Income \(rental, part-time work, etc.\)",
            r"Estimated Annual Retirement Expense"
        ],
        "YOUR OTHER MAJOR GOALS": [
            r"GOAL", r"COST", r"WHEN"
        ],
        "YOUR ASSETS AND LIABILITIES": [
            r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
            r"Current Value", r"Annual Contributions"
        ],
        "MY LIABILITIES": [
            r"Balance", r"Interest Rate", r"Monthly Payment"
        ]
    }

    # Validate table content
    for section, checks in table_checks.items():
        section_found = False
        for table in tables:
            table_text = "\n".join(["\t".join(row) for row in table])
            if section in table_text:
                section_found = True
                for check in checks:
                    if not re.search(check, table_text, re.IGNORECASE):
                        errors.append(f"* Missing or empty field in {section} section: {check}")
                break
        if not section_found:
            errors.append(f"* {section} section missing.")

    return client_name, errors


# RUN Button :
# async def generate_investment_suggestions(investment_personality, context): # # GET Method for py , for front end its Post API
    
#     # retriever = asyncio.run(load_vector_db("uploaded_file"))

#     retriever = await load_vector_db("uploaded_file")
#     # retriever = await load_vector_db("data\Financial_Investment_1.docx") 

#     chain = await make_retrieval_chain(retriever)

#     # chain = asyncio.run(make_retrieval_chain(retriever))
    
#     if chain is not None:
#         # summary = context
#         # query = summary + "\n" + investment_personality
#         query = str(investment_personality)
#         response = chain.invoke({"input": query})
#         format_response = markdown_to_text(response['answer'])
#         return format_response
#         # st.write(format_response)

#         # handle_graph(response['answer'])

#     else:
#         st.error("Failed to create the retrieval chain. Please upload a valid document.")



# Generate Infographics : Best Code so far:

import re
from collections import defaultdict
import numpy as np



def extract_numerical_data(response): # curr version but cant capture annual return 
    # Define patterns to match different sections and their respective allocations
    patterns = {
        'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
        'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
        'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
    }

    data = defaultdict(dict)

    for section, pattern in patterns.items():
        match = pattern.search(response)
        if match:
            investments_text = match.group(1)
            # Extract individual investment types and their allocations
            investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
            for investment_match in investment_pattern.findall(investments_text):
                investment_type, min_allocation, max_allocation = investment_match
                data[section][investment_type.strip()] = {
                    'min': min_allocation,
                    'max': max_allocation
                }

    # Extract time horizon and expected returns
    time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
    min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
    max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
    min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
    max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

    time_horizon_match = time_horizon_pattern.search(response)
    min_return_match = min_return_pattern.search(response)
    max_return_match = max_return_pattern.search(response)
    min_growth_match = min_growth_pattern.search(response)
    max_growth_match = max_growth_pattern.search(response)

    if time_horizon_match:
        data['Time Horizon'] = {
            'min_years': time_horizon_match.group(1),
            'max_years': time_horizon_match.group(2)
        }

    if min_return_match:
        data['Expected Annual Return'] = {
            'min': min_return_match.group(1),
            'max': min_return_match.group(2)
        }

    if max_return_match:
        data['Expected Annual Return'] = {
            'min': max_return_match.group(1),
            'max': max_return_match.group(2)
        }

    if min_growth_match:
        data['Expected Growth in Dollars'] = {
            'min': min_growth_match.group(1),
            'max': min_growth_match.group(2)
        }

    if max_growth_match:
        data['Expected Growth in Dollars'] = {
            'min': max_growth_match.group(1),
            'max': max_growth_match.group(2)
        }

    return data

def normalize_allocations(allocations):
    total = sum(allocations)
    if total == 100:
        return allocations
    return [round((allocation / total) * 100, 2) for allocation in allocations]


import datetime  # Import the datetime module to get the current year
# uodated to have current year
def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
    try:
        # Get the current year
        curr_year = datetime.datetime.now().year

        # Print data_extracted to debug the structure
        print("Data extracted:", data_extracted)

        # Check if 'Expected Annual Return' and 'Time Horizon' exist and have the expected keys
        if 'Expected Annual Return' not in data_extracted:
            print("'Expected Annual Return' missing in data_extracted")
            data_extracted['Expected Annual Return']['min'] = 6
            data_extracted['Expected Annual Return']['max'] = 8
            min_return = 6
            max_return = 8
        else:
            min_return = float(data_extracted['Expected Annual Return'].get('min', '0').strip('%'))
            max_return = float(data_extracted['Expected Annual Return'].get('max', '0').strip('%'))

        min_years = int(data_extracted['Time Horizon'].get('min_years', 1))  # Default to 1 year if missing
        max_years = int(data_extracted['Time Horizon'].get('max_years', 10))  # Default to 10 years if missing

        def calculate_compounded_return(principal, rate, years):
            return principal * (1 + rate / 100) ** years

        def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
            return nominal_return / (1 + inflation_rate / 100) ** years

        # Create labels for the next 10 years starting from the current year
        labels = list(range(curr_year, curr_year + max_years))

        min_compounded = []
        max_compounded = []
        min_inflation_adjusted = []
        max_inflation_adjusted = []

        for year in range(1, max_years + 1):
            # Calculate nominal compounded returns
            min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
            max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

            # Calculate inflation-adjusted compounded returns
            min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
            max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

            # Append results
            min_compounded.append(min_compounded_value)
            max_compounded.append(max_compounded_value)
            min_inflation_adjusted.append(min_inflation_value)
            max_inflation_adjusted.append(max_inflation_value)

        # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
        combined_chart_data = {
            'labels': labels,  # Current year and the next 10 years
            'datasets': [
                {
                    'label': 'Minimum Compounded Return',
                    'data': min_compounded,
                    'borderColor': 'rgb(255, 99, 132)',  # Red color
                    'fill': False
                },
                {
                    'label': 'Maximum Compounded Return',
                    'data': max_compounded,
                    'borderColor': 'rgb(54, 162, 235)',  # Blue color
                    'fill': False
                },
                {
                    'label': 'Min Inflation Adjusted Return',
                    'data': min_inflation_adjusted,
                    'borderColor': 'rgb(75, 192, 192)',  # Light blue
                    'borderDash': [5, 5],  # Dashed line for distinction
                    'fill': False
                },
                {
                    'label': 'Max Inflation Adjusted Return',
                    'data': max_inflation_adjusted,
                    'borderColor': 'rgb(153, 102, 255)',  # Light purple
                    'borderDash': [5, 5],  # Dashed line for distinction
                    'fill': False
                }
            ]
        }
    except KeyError as e:
        print(f"KeyError occurred: {e}")
        return jsonify({'message': f'Key Error: {e}'}), 400
    except Exception as e:
        print(f"Error occurred while preparing data for combined line chart: {e}")
        return jsonify({'message': 'Internal Server Error in creating line chart'}), 500

    return combined_chart_data



# def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
#     try:
#         min_return = float(data_extracted['Expected Annual Return']['min'].strip('%'))
#         max_return = float(data_extracted['Expected Annual Return']['max'].strip('%'))
#         min_years = int(data_extracted['Time Horizon']['min_years'])
#         max_years = int(data_extracted['Time Horizon']['max_years'])

#         def calculate_compounded_return(principal, rate, years):
#             return principal * (1 + rate / 100) ** years

#         def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
#             return nominal_return / (1 + inflation_rate / 100) ** years

#         labels = list(range(1, max_years + 1))  # Years for the x-axis
#         min_compounded = []
#         max_compounded = []
#         min_inflation_adjusted = []
#         max_inflation_adjusted = []

#         for year in labels:
#             # Calculate nominal compounded returns
#             min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
#             max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

#             # Calculate inflation-adjusted compounded returns
#             min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
#             max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

#             # Append results
#             min_compounded.append(min_compounded_value)
#             max_compounded.append(max_compounded_value)
#             min_inflation_adjusted.append(min_inflation_value)
#             max_inflation_adjusted.append(max_inflation_value)

#         # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
#         combined_chart_data = {
#             'labels': labels,
#             'datasets': [
#                 {
#                     'label': 'Minimum Compounded Return',
#                     'data': min_compounded,
#                     'borderColor': 'rgb(255, 99, 132)',  # Red color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Maximum Compounded Return',
#                     'data': max_compounded,
#                     'borderColor': 'rgb(54, 162, 235)',  # Blue color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Min Inflation Adjusted Return',
#                     'data': min_inflation_adjusted,
#                     'borderColor': 'rgb(75, 192, 192)',  # Light blue
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 },
#                 {
#                     'label': 'Max Inflation Adjusted Return',
#                     'data': max_inflation_adjusted,
#                     'borderColor': 'rgb(153, 102, 255)',  # Light purple
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 }
#             ]
#         }
#     except Exception as e:
#         print(f"Error occurred while preparing data for combined line chart: {e}")
#         return jsonify({'message': 'Internal Server Error in creating line chart'}), 500
#     return combined_chart_data



def plot_investment_allocations(data):
    # Create subplots with a large figure size
    fig, axes = plt.subplots(2, 1, figsize= (16,10)) #(28, 15))  # Adjust size as needed

    # Plot Growth-Oriented Investments
    growth_data = data['Growth-Oriented Investments']
    growth_labels = list(growth_data.keys())
    growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
    growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

    axes[0].bar(growth_labels, growth_min, color='skyblue', label='Min Allocation')
    axes[0].bar(growth_labels, growth_max, bottom=growth_min, color='lightgreen', label='Max Allocation')
    axes[0].set_title('Growth-Oriented Investments', fontsize=16)
    axes[0].set_ylabel('Percentage Allocation', fontsize=14)
    axes[0].set_xlabel('Investment Types', fontsize=14)
    axes[0].tick_params(axis='x', rotation=45, labelsize=12)
    axes[0].tick_params(axis='y', labelsize=12)
    axes[0].legend()

    # Plot Conservative Investments
    conservative_data = data['Conservative Investments']
    conservative_labels = list(conservative_data.keys())
    conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
    conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

    axes[1].bar(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
    axes[1].bar(conservative_labels, conservative_max, bottom=conservative_min, color='lightgreen', label='Max Allocation')
    axes[1].set_title('Conservative Investments', fontsize=16)
    axes[1].set_ylabel('Percentage Allocation', fontsize=14)
    axes[1].set_xlabel('Investment Types', fontsize=14)
    axes[1].tick_params(axis='x', rotation=45, labelsize=12)
    axes[1].tick_params(axis='y', labelsize=12)
    axes[1].legend()

    # Tight layout for better spacing
    plt.tight_layout()
    plt.show()
    return fig


def plot_pie_chart(data):
    fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

    # Combine all investment data for pie chart
    all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
    labels = list(all_data.keys())
    sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
    colors = plt.cm.Paired(range(len(labels)))

    wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    ax.set_title('Investment Allocation')

    # Add legend
    ax.legend(wedges, labels, title="Investment Types", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

    return fig



def bar_chart(data):
    fig, ax = plt.subplots(figsize=(12, 8))  # Increased size

    # Data for plotting
    categories = list(data.keys())
    values_min = [int(data[cat]['min'].strip('%')) for cat in categories]
    values_max = [int(data[cat]['max'].strip('%')) for cat in categories]

    x = range(len(categories))

    ax.bar(x, values_min, width=0.4, label='Min Allocation', color='skyblue', align='center')
    ax.bar(x, values_max, width=0.4, label='Max Allocation', color='lightgreen', align='edge')

    ax.set_xticks(x)
    ax.set_xticklabels(categories, rotation=45, ha='right')
    ax.set_xlabel('Investment Categories')
    ax.set_ylabel('Percentage Allocation')
    ax.set_title('Investment Allocation')
    ax.legend()

    plt.tight_layout()
    return fig


import random

def generate_colors(n):
    """
    Generate 'n' random RGB colors.

    Args:
        n (int): Number of colors to generate.
    
    Returns:
        list: A list of RGB colors in 'rgb(r, g, b)' format.
    """
    colors = []
    for _ in range(n):
        r = random.randint(0, 255)
        g = random.randint(0, 255)
        b = random.randint(0, 255)
        colors.append(f'rgb({r}, {g}, {b})')
    
    return colors


import plotly.graph_objects as go
import numpy as np

 
# def client_form():
#     st.title("Client Details Form")

#     with st.form("client_form"):
#         st.header("Personal Information")
#         client_name = st.text_input("Client Name")
#         co_client_name = st.text_input("Co-Client Name")
#         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
#         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
#         today_date = st.date_input("Today's Date")
        
#         st.header("Financial Information")
#         current_assets = st.text_area("Current Assets (e.g., type and value)")
#         liabilities = st.text_area("Liabilities (e.g., type and amount)")
#         annual_income = st.text_area("Current Annual Income (source and amount)")
#         annual_contributions = st.text_area("Annual Contributions (e.g., retirement savings)")

#         st.header("Insurance Information")
#         life_insurance = st.text_input("Life Insurance (e.g., coverage amount)")
#         disability_insurance = st.text_input("Disability Insurance (e.g., coverage amount)")
#         long_term_care = st.text_input("Long-Term Care Insurance (e.g., coverage amount)")

#         st.header("Estate Planning")
#         will_status = st.radio("Do you have a will?", ["Yes", "No"])
#         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
#         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
#         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

#         # Submit button
#         submitted = st.form_submit_button("Submit")

#         if submitted:
#             # Save form data
#             form_data = {
#                 "Client Name": client_name,
#                 "Co-Client Name": co_client_name,
#                 "Client Age": client_age,
#                 "Co-Client Age": co_client_age,
#                 "Today's Date": str(today_date),
#                 "Current Assets": current_assets,
#                 "Liabilities": liabilities,
#                 "Annual Income": annual_income,
#                 "Annual Contributions": annual_contributions,
#                 "Life Insurance": life_insurance,
#                 "Disability Insurance": disability_insurance,
#                 "Long-Term Care Insurance": long_term_care,
#                 "Will Status": will_status,
#                 "Trust Status": trust_status,
#                 "Power of Attorney": power_of_attorney,
#                 "Healthcare Proxy": healthcare_proxy,
#             }
            
#             # Save to a file or database
#             with open("client_data.txt", "a") as f:
#                 f.write(str(form_data) + "\n")
            
#             st.success("Form submitted successfully!")
#             st.session_state.page = "main"  # Redirect back to main page after form submission


from datetime import date  # Make sure to import the date class


# Function to parse financial data from the text
import re

def parse_financial_data(text_content):
    assets = []
    liabilities = []

    # Define regex patterns to capture text following headings
    asset_pattern = re.compile(r"MY ASSETS:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)
    liability_pattern = re.compile(r"LIABILITIES:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)

    # Extract assets
    asset_matches = asset_pattern.findall(text_content)
    if asset_matches:
        asset_text = asset_matches[0]
        # Further processing to extract individual asset values if they are detailed
        asset_lines = asset_text.split('\n')
        for line in asset_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                asset_value = float(match.group().replace(",", ""))
                assets.append(asset_value)

    # Extract liabilities
    liability_matches = liability_pattern.findall(text_content)
    if liability_matches:
        liability_text = liability_matches[0]
        # Further processing to extract individual liability values if they are detailed
        liability_lines = liability_text.split('\n')
        for line in liability_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                liability_value = float(match.group().replace(",", ""))
                liabilities.append(liability_value)

    print("Assets Found:", assets)
    print("Liabilities Found:", liabilities)

    return assets, liabilities



# Function to extract numerical values from a text input
def extract_numeric(value):
    try:
        return float(re.sub(r'[^\d.]', '', value))  # Remove non-numeric characters and convert to float
    except ValueError:
        return 0


# plots graph from the details of the form :


def is_float(value):
    try:
        float(value)
        return True
    except ValueError:
        return False



def plot_assets_liabilities_pie_chart(assets, liabilities, threshold=50): # best plot 
    """
    Plots separate pie charts for assets and liabilities. If there are any categories
    below a specified threshold, they are plotted in an additional small pie chart.
    
    Parameters:
    - assets: dict, keys are asset names, values are their amounts.
    - liabilities: dict, keys are liability names, values are their amounts.
    - threshold: int, percentage threshold below which segments are considered small.
    """
    # Update matplotlib settings to increase the font size globally
    # plt.rcParams.update({'font.size': 32})

    plt.rcParams.update({'font.size': 16})

    def plot_pie(data, title):
        # Filter out zero values and create a summary for small segments
        total = sum(data.values())
        filtered_data = {k: v for k, v in data.items() if (v / total) >= threshold / 100}
        small_segments = {k: v for k, v in data.items() if (v / total) < threshold / 100}
        small_total = sum(small_segments.values())

        # Plotting logic
        if small_segments:
            fig, (ax_main, ax_small) = plt.subplots(1, 2, figsize=(30, 15))  # Side-by-side layout
        else:
            fig, ax_main = plt.subplots(figsize=(30, 20))  # Only main chart with larger size

            # fig, ax_main = plt.subplots(figsize=(10, 10))  # Only main chart with larger size

        # Plot main pie chart
        labels_main = list(filtered_data.keys()) + ([f"Other small {title}"] if small_segments else [])
        values_main = list(filtered_data.values()) + ([small_total] if small_segments else [])
        wedges_main, texts_main, autotexts_main = ax_main.pie(
            values_main, labels=labels_main, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
            startangle=140, textprops={'fontsize': 28} #18}  # Larger font size for labels
        )

        ax_main.set_title(title, fontsize=20)
        # Position legend to the right of the plot to avoid overlapping
        ax_main.legend(wedges_main, labels_main, title="Categories", loc="upper right", bbox_to_anchor=(0.001, 0.9), fontsize= 28)#14)

        if small_segments:
            # Plot additional small pie chart for small segments
            labels_small = list(small_segments.keys())
            values_small = list(small_segments.values())
            wedges_small, texts_small, autotexts_small = ax_small.pie(
                values_small, labels=labels_small, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
                startangle=140, textprops={'fontsize': 24} #14}  # Consistent label size for small chart
            )
            ax_small.set_title(f"Small Segments of {title}", fontsize=20)
            # Position legend to the right of the small pie chart but slightly lower to avoid overlap with the main chart's legend
            ax_small.legend(wedges_small, labels_small, title="Small Categories", loc="center left", bbox_to_anchor=(1.2, 0.3), fontsize= 22)#12)

        st.pyplot(fig)

    # Convert valid entries to float, ensuring only numeric values are considered
    assets = {k: float(v) for k, v in assets.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}
    liabilities = {k: float(v) for k, v in liabilities.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}

    # Plot pie charts
    plot_pie(assets, 'Distribution of Assets')
    plot_pie(liabilities, 'Distribution of Liabilities')


# def plot_assets_liabilities_pie_chart(assets, liabilities):# properly plots a big and 1 small pie chart for both assets and liability
#     # Filter and convert values to float, handle non-numeric or empty inputs
#     filtered_assets = {k: float(v) for k, v in assets.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}
#     filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}

#     # Combine assets and liabilities for total calculation
#     all_values = {**filtered_assets, **filtered_liabilities}
#     total_value = sum(all_values.values())

#     # Separate main and small segments
#     main_segments = {k: v for k, v in all_values.items() if (v / total_value) >= 0.05}
#     small_segments = {k: v for k, v in all_values.items() if (v / total_value) < 0.05}
#     small_total = sum(small_segments.values())

#     # Prepare data for main pie chart
#     main_labels = list(main_segments.keys()) + (["Others"] if small_segments else [])
#     main_values = list(main_segments.values()) + ([small_total] if small_segments else [])

#     # Prepare data for small pie chart (only if there are small segments)
#     small_labels = list(small_segments.keys())
#     small_values = list(small_segments.values())

#     fig, ax = plt.subplots(figsize=(8, 6))

#     # Plot main pie chart
#     wedges, texts, autotexts = ax.pie(
#         main_values,
#         labels=main_labels,
#         autopct='%1.1f%%',
#         startangle=140,
#         colors=plt.cm.Paired.colors,
#     )

#     # Explode the "Others" slice
#     if small_segments:
#         others_index = main_labels.index("Others")
#         wedges[others_index].set_edgecolor('white')
#         # wedges[others_index].set_linestyle('--')
#         wedges[others_index].set_linewidth(2)
#         wedges[others_index].set_hatch('/')

#     ax.set_title('Assets and Liabilities Distribution')

#     # Draw a second pie chart for "Others"
#     if small_segments:
#         fig2, ax2 = plt.subplots(figsize=(8, 6))
#         wedges_small, texts_small, autotexts_small = ax2.pie(
#             small_values,
#             labels=small_labels,
#             autopct='%1.1f%%',
#             startangle=140,
#             colors=plt.cm.Pastel1.colors
#         )

#         ax2.set_title('Detailed View of "Others" Categories')

#     plt.tight_layout()
#     st.pyplot(fig)
#     if small_segments:
#         st.pyplot(fig2)



def save_data_to_file(form_data):
    file_path = 'client_data.txt'
    with open(file_path, 'a') as file:
        file.write(str(form_data) + "\n")
    # st.success(f"Form data saved to {file_path}")
    print(f"Form data saved to {file_path}")
    

def client_form():
    st.title("Client Details Form")

    with st.form("client_form"):
        st.header("Personal Information")
        client_name = st.text_input("Client Name")
        co_client_name = st.text_input("Co-Client Name")
        client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
        co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
        today_date = st.date_input("Today's Date")

        st.header("Your Assets (in $)")

        assets = {
            # 'Annual Income': st.text_input("Annual Income (e.g. , Your Annual Salary Income or other source of income) "),
            'Cash/Bank Account': st.text_input("Cash/Bank Account"),
            '401(k), 403(b), 457 Plans': st.text_input("Your 401(k), 403(b), 457 Plans "),
            'Traditional, SEP and SIMPLE IRAs': st.text_input("Traditional, SEP and SIMPLE IRAs "),
            'Roth IRA,Roth 401(k)': st.text_input("Roth IRA, Roth 401(k)"),
            'Brokerage/non-qualified accounts': st.text_input("Brokerage/non-qualified accounts"),
            'Annuities': st.text_input("Annuities"),
            '529 Plans': st.text_input("529 Plans"),
            'Home': st.text_input("Home"),
            'Other Real Estate': st.text_input("Other Real Estate"),
            'Business': st.text_input("Business"),
            'Other': st.text_input("Other")
        }
        st.header("Your Liabilities (in $)")

        liabilities = {
            'Mortgage': st.text_input("Mortgage"),
            # 'Annual Mortgage Interest Rate': st.number_input("Annual Mortgage Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=12.0, step=0.5),
            # 'Mortagage Time Period': st.number_input("Mortagage Time Period (Mention the time period of the Mortgage in years)", min_value=0, max_value=100,value=10,step=1),

            'Home Loans': st.text_input("Home Loans"),
            # 'Home Loans Interest Rate': st.number_input("Home Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=10.0, step=0.5),
            # 'Home Loans Time Period': st.number_input("Home Loans Time Period (Mention the time period of the Home Loan in years)", min_value=0, max_value=100,value=15,step=1),

            'Vehicle Loans': st.text_input("Vehicle Loans"),
            # 'Vehicle Loans Interest Rate': st.number_input("Vehicle Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
            # 'Vehicle Loans Time Period': st.number_input("Vehicle Loans Time Period (Mention the time period of the Car/Vehicle Loan in years)", min_value=0, max_value=100,value=15,step=1),

            'Education Loans': st.text_input("Education Loans"),
            # 'Education Loans Interest Rate' : st.number_input("Education Loans Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
            # 'Education Loans Time Period': st.number_input("Education Loans Time Period (Mention the time period of the Education Loan in years)", min_value=0, max_value=100,value=15,step=1),

            # 'Credit Card': st.text_input("Monthly Credit Card Debt (Mention Amount)"),
            # 'Credit Card Debt Interest Rate': st.number_input("Credit Card Debt Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),

            'Miscellaneous': st.text_input("Miscellaneous"),
        }

        st.header("Your Retirement Goal")
        retirement_age = st.number_input("At what age do you plan to retire?", min_value=0, max_value=120, value=65, step=1)
        retirement_income = st.text_input("Desired annual retirement income")

        st.header("Your Other Goals")
        goal_name = st.text_input("Name of the Goal (e.g . , Dream House, Travel, Educational, etc.)")
        goal_amount = st.text_input("Amount needed for the goal (in $)")
        goal_timeframe = st.number_input("Timeframe to achieve the goal (in years)", min_value=0, max_value=100, value=5, step=1)

        st.header("Insurance Information")
        life_insurance_Benefit = st.text_input("Life Insurance-Benefit")
        life_insurance_Premium = st.text_input("Life Insurance-Premium")
        disability_insurance_Benefit = st.text_input("Disability Insurance-Benefit")
        disability_insurance_Premium = st.text_input("Disability Insurance-Premium")
        long_term_care_benefit = st.text_input("Long-Term Care Insurance-Benefit")
        long_term_care_premium = st.text_input("Long-Term Care Insurance-Premium")


        st.header("Estate Planning")
        will_status = st.radio("Do you have a will?", ["Yes", "No"])
        trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
        power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
        healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

        submitted = st.form_submit_button("Submit")

        if submitted:
            form_data = {
                "Client Name": client_name,
                "Co-Client Name": co_client_name,
                "Client Age": client_age,
                "Co-Client Age": co_client_age,
                "Today's Date": str(today_date),
                "Assets": assets,
                "Liabilities": liabilities,
                "Retirement Age": retirement_age,
                "Desired Retirement Income": retirement_income,
                "Goal Name": goal_name,
                "Goal Amount": goal_amount,
                "Goal Timeframe": goal_timeframe,
                "Life Insurance Benefit": life_insurance_Benefit,
                "Life Insurance Premium": life_insurance_Premium,
                "Disability Insurance Benefit": disability_insurance_Benefit,
                "Disability Insurance Premium": disability_insurance_Premium,
                "Long-Term Care Insurance Benefit": long_term_care_benefit,
                "Long-Term Care Insurance Premium": long_term_care_premium,
                "Will Status": will_status,
                "Trust Status": trust_status,
                "Power of Attorney": power_of_attorney,
                "Healthcare Proxy": healthcare_proxy,
            }

            save_data_to_file(form_data)
            
            # # Plot the pie chart
            # st.subheader("Assets and Liabilities Breakdown")
            # plot_assets_liabilities_pie_chart(assets, liabilities)

            # Store data in session state and redirect to main
            st.session_state.assets = assets
            st.session_state.liabilities = liabilities
            st.session_state.total_assets, st.session_state.total_liabilities = calculate_totals(assets, liabilities)
            st.session_state.page = "main"
            st.success("Data submitted!\nThank You for filling the form !\nReturning to main portal...")

import math
def calculate_compounded_amount(principal, rate, time):
    """
    Calculates the compounded amount using the formula:
    A = P * (1 + r/n)^(nt)
    Assuming n (compounding frequency) is 1 for simplicity (annually).
    """
    if principal == 0 or rate == 0 or time == 0:
        return principal
    else:
        # Using annual compounding
        return principal * (1 + rate / 100) ** time
    
def calculate_totals(assets, liabilities):
    total_assets = sum(extract_numeric(v) for v in assets.values())
    print(f"Total Assets : {total_assets}")
    total_liabilities = 0
    total_liabilities = sum(extract_numeric(v) for v in liabilities.values() )

    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Mortgage']),
    #     liabilities['Annual Mortgage Interest Rate'],
    #     liabilities['Mortagage Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Home Loans']),
    #     liabilities['Home Loans Interest Rate'],
    #     liabilities['Home Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Vehicle Loans']),
    #     liabilities['Vehicle Loans Interest Rate'],
    #     liabilities['Vehicle Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Education Loans']),
    #     liabilities['Education Loans Interest Rate'],
    #     liabilities['Education Loans Time Period']
    # )
    
    # For credit card debt, only calculate compounded amount if interest rate > 0

    # credit_card_balance = extract_numeric(liabilities['Credit Card'])
    # credit_card_interest = liabilities['Credit Card Debt Interest Rate']
    # if credit_card_interest > 0:
    #     # Assuming the time period for credit card debt is 1 year for compounding
    #     total_liabilities += calculate_compounded_amount(credit_card_balance, credit_card_interest, 1)
    # else:
    #     total_liabilities += credit_card_balance
    
    # Miscellaneous debts are taken directly as is
    total_liabilities += extract_numeric(liabilities['Miscellaneous'])
    rounded_liabilities = round(total_liabilities,2)

    print(f"Total liabilities :{total_liabilities}")
    print(f"Rounded of Total liabilities :{rounded_liabilities}")

    return total_assets, rounded_liabilities #total_liabilities

def create_financial_summary_table(assets, liabilities):
    # Filter out items with zero value
    filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
    filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

    # Create DataFrames for assets and liabilities with indices starting from 1
    assets_df = pd.DataFrame(
        list(filtered_assets.items()), 
        columns=['Assets', 'Amount ($)'], 
        index=range(1, len(filtered_assets) + 1)
    )
    liabilities_df = pd.DataFrame(
        list(filtered_liabilities.items()), 
        columns=['Liabilities', 'Amount ($)'], 
        index=range(1, len(filtered_liabilities) + 1)
    )

    # Calculate total
    total_assets, total_liabilities = calculate_totals(assets, liabilities)

    # Add total row with index incremented by 1
    total_assets_row = pd.DataFrame(
        [['TOTAL', total_assets]], 
        columns=['Assets', 'Amount ($)'], 
        index=[len(assets_df) + 1]
    )
    total_liabilities_row = pd.DataFrame(
        [['TOTAL', total_liabilities]], 
        columns=['Liabilities', 'Amount ($)'], 
        index=[len(liabilities_df) + 1]
    )

    # Append total rows to DataFrames
    assets_df = pd.concat([assets_df, total_assets_row])
    liabilities_df = pd.concat([liabilities_df, total_liabilities_row])

    # Display tables with formatted values
    st.subheader("Assets")
    st.table(assets_df.style.format({'Amount ($)': '{:,.2f}'}))

    st.subheader("Liabilities")
    st.table(liabilities_df.style.format({'Amount ($)': '{:,.2f}'}))


def plot_bar_graphs(assets, liabilities):
    # Filter out items with zero values
    filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
    filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

    # Calculate compounded liabilities
    # compounded_liabilities = {} 

    # for k, v in filtered_liabilities.items():
        # if 'Interest Rate' in k or 'Time Period' in k:
        #     continue  # Skip non-monetary entries

        # if k == 'Credit Card Payment' and liabilities['Credit Card Debt Interest Rate'] == 0.0:
        #     continue  # Skip if credit card interest rate is zero

        # if k == 'Mortgage':
        #     interest_rate = liabilities['Annual Mortgage Interest Rate']
        #     time_period = liabilities['Mortagage Time Period']

        # elif k == 'Home Loans':
        #     interest_rate = liabilities['Home Loans Interest Rate']
        #     time_period = liabilities['Home Loans Time Period']

        # elif k == 'Car/Vehicle Loans':
        #     interest_rate = liabilities['Car/Vehicle Loans Interest Rate']
        #     time_period = liabilities['Car/Vehicle Loans Time Period']

        # elif k == 'Education Loans':
        #     interest_rate = liabilities['Education Loans Interest Rate']
        #     time_period = liabilities['Education Loans Time Period']

        # elif k == 'Credit Card Payment':
        #     interest_rate = liabilities['Credit Card Debt Interest Rate']
        #     time_period = 1  # Assuming interest is calculated yearly

        # if interest_rate > 0:
        #     compounded_amount = float(v) * (1 + float(interest_rate) / 100) ** float(time_period)
        #     compounded_liabilities[k] = compounded_amount
        # else:
        #     compounded_liabilities[k] = float(v)

    # Plot bar graph for assets
    st.write("### All Assets ")
    fig1, ax1 = plt.subplots()
    ax1.bar(filtered_assets.keys(), filtered_assets.values(), color='green')
    ax1.set_ylabel('Amount ($)')
    ax1.set_xlabel('Asset Type')
    ax1.set_title(' All Assets ')
    plt.xticks(rotation=45)
    st.pyplot(fig1)

    # Plot bar graph for liabilities
    st.write("### All Liabilities ")
    # st.write("### All Liabilities with Compounded Interest")
    fig2, ax2 = plt.subplots()
    # ax2.bar(compounded_liabilities.keys(), compounded_liabilities.values(), color='red')
    ax2.bar(filtered_liabilities.keys(), filtered_liabilities.values(), color='red')    
    ax2.set_ylabel('Amount ($)')
    ax2.set_xlabel('Liability Type')
    ax2.set_title(' All Liabilities ')

    # ax2.set_title(' All Liabilities with Compounded Interest')
    plt.xticks(rotation=45)
    st.pyplot(fig2)


from docx import Document
# Define a helper function to read and extract text from a DOCX file
def read_docx(file_path):
    document = Document(file_path)
    extracted_text = "\n".join([para.text for para in document.paragraphs])
    return extracted_text



class TrieNode:
    def __init__(self):
        self.children = {}
        self.client_ids = []
        self.end_of_name = False  # Marks the end of a client's name

class Trie:
    def __init__(self):
        self.root = TrieNode()

    def insert(self, name, client_id):
        node = self.root
        for char in name:
            if char not in node.children:
                node.children[char] = TrieNode()
            node = node.children[char]
        node.client_ids.append(client_id)
        node.end_of_name = True

    def search(self, prefix):
        node = self.root
        for char in prefix:
            if char in node.children:
                node = node.children[char]
            else:
                return []  # Prefix not found
        return self._get_all_names_from_node(prefix, node)

    def _get_all_names_from_node(self, prefix, node):
        suggestions = []
        if node.end_of_name:
            suggestions.append((prefix, node.client_ids))
        for char, child_node in node.children.items():
            suggestions.extend(self._get_all_names_from_node(prefix + char, child_node))
        return suggestions



def preload_trie():
    trie = Trie()
    clients = {
        "John Doe": "C001",
        "Jane Smith": "C002",
        "James Brown": "C003",
        "Jill Johnson": "C004",
        "Jake White": "C005"
    }
    for name, client_id in clients.items():
        trie.insert(name.lower(), client_id)  # Insert in lowercase for case-insensitive search
    return trie

async def generate_investment_suggestions_for_investor(investment_personality,clientName,financial_data,financial_file,monthly_investment=10000,investment_period=3): # # GET Method for py , for front end its Post API
    
    # retriever = asyncio.run(load_vector_db("uploaded_file"))

    # retriever =  await load_vector_db("uploaded_file")
    try:
        retriever =  await load_vector_db(financial_file)
    except Exception as e :
        print(f"Error : {e}")
        return jsonify("Error : Failed to load vector database and to generate suggestions : {e}"),400
    
    if not retriever:
        await load_vector_db("data\Financial_Investment_new.docx")
        # await load_vector_db("data\Financial_Investment_1_new.docx") # doesnt works
        # await load_vector_db("data\Financial_Investment_1.docx")
        if not retriever:
            raise Exception("Failed to load vector database.")
    
    print("VectorDB is created successfully")
    # retriever = await load_vector_db("data\Financial_Investment_1.docx") 
    
    try:
        chain = await make_retrieval_chain(retriever,investment_personality,clientName,monthly_investment,investment_period)
    except Exception as e :
        print(f"Error : {e}")
        return jsonify("Error : Failed to create retrieval chain and generate suggestions : {e}"),400
    
    if not chain:
        raise Exception("Failed to create retrieval chain.")
    print("Chain is created to generate suggestions ")
    
    # chain = asyncio.run(make_retrieval_chain(retriever))
    
    print(f"Financial Data : {financial_data}")
    try :
        print(type(financial_data))
        query = f"The Investment Personality of {clientName} is : {investment_personality}" + f"Consider the Monthly Investment as {monthly_investment} and Investment period as {investment_period}" + f"Financial Data of client is : {financial_data[0]}"
        print(query)
    except Exception as e :
        print(f"Error : {e}")
        return "Error : Failed to load financial data"
    
    if chain is not None:
        # summary = context
        # query = summary + "\n" + investment_personality
        
        # query = str(investment_personality)
        response = chain.invoke({"input": query})
        
        # format_response = markdown_to_text(response['answer'])
        # return format_response
        
        # html_output = markdown.markdown(response['answer'])
        # return html_output
        
        # readable_text = markdown_to_readable_text(response['answer'])
        # print(readable_text)
        # return readable_text

        # format_text = convert_to_markdown(response['answer'])
        # return format_text
        
        return response['answer']
    
        

        # handle_graph(response['answer'])

    else:
        logging.INFO("response is not generated by llm model")
        return jsonify("response is not generated by llm model"),500
        # st.error("Failed to create the retrieval chain. Please upload a valid document.")

from flask import Flask, request, jsonify, send_file
import asyncio
from flask_cors import CORS
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})
# CORS(app,resources={r"/api/*":{"origins":"*"}})
# CORS(app)

# Initialize the Trie with preloaded clients
trie = preload_trie()

@app.route('/')
def home():
    return "Wealth Advisor Chatbot API"

@app.route('/investment-suggestions', methods=['POST'])
def investment_suggestions():
    # Get the input data (new or existing client)
    data = request.get_json()

    # Determine if it's a new client or existing client
    client_type = data.get("client_type")

    if client_type == "New Client":
        # Get form details and perform investment suggestions

        # Check if assets and liabilities are provided
        assets = data.get('assets', None)
        liabilities = data.get('liabilities', None)

        if assets and liabilities:
            financial_summary = create_financial_summary_table(assets, liabilities)
            bar_graphs = plot_bar_graphs(assets, liabilities)
            pie_chart = plot_assets_liabilities_pie_chart(assets, liabilities)

            return jsonify({
                "financial_summary": financial_summary,
                "bar_graphs": "Bar graphs generated.",
                "pie_chart": "Pie chart generated."
            })

        return jsonify({"message": "Please fill in the client details to view the assets and liabilities breakdown."})

    elif client_type == "Existing Client":
        # Search for an existing client in the Trie
        search_query = data.get("search_query", "").lower()
        matching_names = trie.search(search_query)

        if matching_names:
            suggestions = [{"name": name, "client_ids": client_ids} for name, client_ids in matching_names]
            return jsonify({"suggestions": suggestions})
        else:
            return jsonify({"message": "No matching clients found."})
    
    return jsonify({"message": "Invalid client type."})



import random

# Generate unique client ID

# def generate_unique_id(name):
#     name_parts = name.split(" ")
#     first_initial = name_parts[0][0] if len(name_parts) > 0 else ""
#     last_initial = name_parts[1][0] if len(name_parts) > 1 else ""
#     random_number = random.randint(1000, 9999)
#     unique_id = f"{first_initial}{last_initial}{random_number}"
#     return unique_id

# # Save details in a Word file
import docx
import os
# #Curr version :
def save_to_word_file(data, file_name):
    doc = docx.Document()
    doc.add_heading('Client Details', 0)

    # Adding client details
    client_details = data.get('clientDetail', {})
    doc.add_paragraph(f"Client Name: {client_details.get('clientName', '')}")
    doc.add_paragraph(f"Client Mobile: {client_details.get('clientMoNo', '')}")
    doc.add_paragraph(f"Client Age: {client_details.get('clientAge', '')}")
    doc.add_paragraph(f"Co-Client Name: {client_details.get('coClientName', '')}")
    doc.add_paragraph(f"Co-Client Mobile: {client_details.get('coMobileNo', '')}")
    doc.add_paragraph(f"Co-Client Age: {client_details.get('coClientAge', '')}")

    # Retirement Plan
    retirement_goal = data.get('retirementGoal', {})
    retirement_plan = retirement_goal.get('retirementPlan', {})
    doc.add_paragraph(f"Retirement Plan Client Age: {retirement_plan.get('retirementAgeClient', '')}")
    doc.add_paragraph(f"Retirement Plan Co-Client Age: {retirement_plan.get('retirementAgeCoClient', '')}")
    
    social_benefit = retirement_goal.get('socialBenefit', {})
    doc.add_paragraph(f"Social Benefit Client: {social_benefit.get('socialBenefitClient', '')}")
    doc.add_paragraph(f"Social Benefit Co-Client: {social_benefit.get('socialBenefitCoClient', '')}")
    
    pension_benefit = retirement_goal.get('pensionBenefit', {})
    doc.add_paragraph(f"Pension Benefit Client: {pension_benefit.get('pensionBenefitClient', '')}")
    doc.add_paragraph(f"Pension Benefit Co-Client: {pension_benefit.get('pensionBenefitCoClient', '')}")
    
    otherIncome = retirement_goal.get('otherIncome', {})
    doc.add_paragraph(f"Other IncomeClient Client: {otherIncome.get('otherIncomeClient', '')}")
    doc.add_paragraph(f"Other IncomeClient Co-Client: {otherIncome.get('otherIncomeCoClient', '')}")
   
    # Estimated Annual Retirement Expense ($ or % of current salary)
    annualRetirement = retirement_goal.get('annualRetirement', {})
    doc.add_paragraph(f"Estimated Annual Retirement Expense ($ or % of current salary) Client: {annualRetirement.get('annualRetireClient', '')}")
    doc.add_paragraph(f"Estimated Annual Retirement Expense ($ or % of current salary) Co-Client: {annualRetirement.get('annualRetireCoClient', '')}")
    

    # Assets and Liabilities
    assets_liabilities = data.get('assetsLiabilities', {})
    
    # Assets
    
    for asset_key, asset_info in assets_liabilities.items():
        current_value_key = [key for key in asset_info.keys() if key.startswith("current")][0]
        annual_value_key = [key for key in asset_info.keys() if key.startswith("annual")][0]
        assets_name_key = "assetsName"
        doc.add_paragraph(f"Assets - {asset_info[assets_name_key]} : Current Value - {asset_info[current_value_key]} , Annual Contributions - {asset_info[annual_value_key]}")
        
    # Liabilities
    myLiabilities = data.get('myLiabilities', {})
    for liability_key, liability_info in myLiabilities.items():
        balance_key = [key for key in liability_info.keys() if key.endswith("Balance")][0]
        interest_key = [key for key in liability_info.keys() if key.endswith("Interest")][0]
        monthly_key = [key for key in liability_info.keys() if key.endswith("Monthly")][0]
        liability_name_key = "liabilityName"
        doc.add_paragraph(f"Liabilities - {liability_info[liability_name_key]} : Balance - {liability_info[balance_key]} , Interest - {liability_info[interest_key]} , Monthly - {liability_info[monthly_key]}")
        
    # my_liabilities = data.get('myLiabilities', {})
    # for liability_type, liability_info in my_liabilities.items():
    #     doc.add_paragraph(f"Liabilities - {liability_info.get('liabilityName', '')}: Balance - {liability_info.get('mortgageBalance', '')} Interest - {liability_info.get('mortgageInterest', '')} Monthly - {liability_info.get('mortgageMonthly', '')}")

    # Protection Plan
    protection_plan = data.get('protectionPlan', {})
    doc.add_paragraph(f"Check Will: {protection_plan.get('checkWill', False)}")
    doc.add_paragraph(f"Check Healthcare: {protection_plan.get('checkHealthCare', False)}")
    doc.add_paragraph(f"Check Attorney: {protection_plan.get('checkAttorney', False)}")
    doc.add_paragraph(f"Check Trust: {protection_plan.get('checkTrust', False)}")

    # Insurance Coverage
    insurance_coverage = data.get('insuranceCoverage', {})
    life_insurance_client = insurance_coverage.get('lifeInsuranceClient', {})
    doc.add_paragraph(f"Life Insurance Client: Benefit - {life_insurance_client.get('benefitLIClient', '')} Monthly Pay - {life_insurance_client.get('monthlyPayLIClient', '')}")
    
    life_insurance_co_client = insurance_coverage.get('lifeInsuranceCoClient', {})
    doc.add_paragraph(f"Life Insurance Co-Client: Benefit - {life_insurance_co_client.get('benefitLICoClient', '')} Monthly Pay - {life_insurance_co_client.get('monthlyPayLICoClient', '')}")
 
    disableIncome = insurance_coverage.get('disableIncomeClient', {})
    disableIncomeClient = insurance_coverage.get('disableIncomeClient',{})
    doc.add_paragraph(f"Disable Income Client - {disableIncomeClient.get('benefitDisableClient', '')}")
    
    disableIncomeCoClient = insurance_coverage.get('disableIncomeCoClient', {})
    doc.add_paragraph(f"Disable Income Co-Client - {disableIncomeCoClient.get('benefitDisableCoClient', '')}")
    
    longTermCoClient = insurance_coverage.get('longTermCoClient')
    doc.add_paragraph(f"Long Term Client: Benefit - {longTermCoClient.get('benefitLongTermClient', '')} Monthly Pay - {longTermCoClient.get('monthlyPayLongTermClient', '')}")
    
    # Goal Fields
    goal_fields = data.get('goalFields', [])
    for goal in goal_fields:
        doc.add_paragraph(f"Goal: {goal.get('goal', '')} Cost: {goal.get('cost', '')} When: {goal.get('when', '')}")

    # Income Fields
    income_fields = data.get('incomeFields', [])
    for income in income_fields:
        doc.add_paragraph(f"Income Source: {income.get('sourceIncome', '')} Amount: {income.get('amountIncome', '')}")

    # Save file
    file_name = os.path.join("data", file_name)
    doc.save(f"{file_name}.docx")


@app.route('/submit-client-data', methods=['POST'])
def submit_client_data():
    try:
        data = request.get_json()
        print(data)
        # Generate the unique ID
        client_name = data['clientDetail']['clientName']
        
        # unique_id = generate_unique_id(client_name)
        # unique_id = data['clientDetail']['uniqueId']
        
        unique_id = data['uniqueId']
        print(unique_id)
        
        data['uniqueId'] = unique_id
        
        # Save the data to a Word file
        
        file_name = unique_id
        # save_to_word_file(data, file_name)
        save_to_word_file(data, file_name)
        
        return jsonify({
            'message': 'Client data received and saved successfully.'
        }), 200
    except Exception as e:
        return jsonify({'message': f"An error occurred: {e}"}), 500


# Determine Investment personality through the investor assesmnet tab : 
@app.route('/investor-personality-assessment', methods=['POST'])
def investor_personality_assessment():
    try:
        # Collecting client name and assessment data
        data = request.json  # Expecting JSON input
        # client_name = data.get('client_name')
        client_id = data.get('client_id')
        assessment_data = data.get('assessment_data')  
        
        # if not client_id or not assessment_data:
        #     return jsonify({'message': 'Client name and assessment data are required.'}), 400
        
        logging.info(f"Received assessment data for client with client id : {client_id}")

        # Pass the assessment data to determine the investment personality
        personality = asyncio.run(determine_investment_personality(assessment_data))
        logging.info(f"Determined personality for {client_id}: {personality}")

        # Return the personality and client id in response
        return jsonify({
            'client_id': client_id,
            'investment_personality': personality
        }), 200
    
    except Exception as e:
        logging.error(f"Error processing investor assessment: {e}")
        return jsonify({'message': 'Internal Server Error'}), 500
    

import logging
# global investmentPersonality  # Global Variable
# investmentPersonality = ""

async def make_suggestions(investmentPersonality,clientName,financial_file="data\Financial_Investment_1.docx",monthly_investment=10000,investment_period=3):
    try:
        try:
            # financial_file = financial_file
            print(f"Processing the File for the client: {clientName} and the file : {financial_file}")
            # financial_data = asyncio.run(process_document(financial_file))
            
            financial_data = await process_document(f"data\{financial_file}")
            
            print(f"Data passed : {financial_data}")
            
            financial_file = f"data\{financial_file}"
            
            print(f"Finished processing the File for the client : {financial_file}")
            
            # suggestions = await generate_investment_suggestions_for_investor(investmentPersonality,clientName, financial_file,monthly_investment,investment_period)
            
            suggestions = await generate_investment_suggestions_for_investor(investmentPersonality,clientName,financial_data,financial_file,monthly_investment,investment_period)
            
            
            # print(f"Finished processing the suggestions : {suggestions}")
            
            htmlSuggestions = markdown.markdown(suggestions)
            
            formatSuggestions = markdown_to_text(suggestions)
            
            # print(f"The suggestions generated for the client are :\n {formatSuggestions}")
            
            # need to change the data extraction process : 
            data_extracted = extract_numerical_data(suggestions)
            
            min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                            [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
            max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                            [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

            # Normalize allocations
            min_allocations = normalize_allocations(min_allocations)
            max_allocations = normalize_allocations(max_allocations)

            # Update Bar Chart Data
            bar_chart_data = {
                'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
                'datasets': [{
                    'label': 'Min Allocation',
                    'data': min_allocations,
                    'backgroundColor': 'skyblue'
                },
                {
                    'label': 'Max Allocation',
                    'data': max_allocations,
                    'backgroundColor': 'lightgreen'
                }]
            }

            # Similar changes can be made for the Pie Chart Data:
            all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
            num_labels = len(all_labels)
            max_allocations_for_pie = normalize_allocations(
                [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
                [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
            )
            
            # Generate colors based on the number of labels
            dynamic_colors = generate_colors(num_labels)

            # Update Pie Chart Data
            pie_chart_data = {
                'labels': all_labels,
                'datasets': [{
                    'label': 'Investment Allocation',
                    'data': max_allocations_for_pie,
                    'backgroundColor': dynamic_colors,
                    'hoverOffset': 4
                }]
            }
            
            
            # Prepare the data for the line chart with inflation adjustment
            initial_investment = 10000
            # compounded_chart_data, inflation_adjusted_chart_data = prepare_line_chart_data_with_inflation(data_extracted, initial_investment)
            combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
            print(f"\nThe combined chart data is : {combined_chart_data}")
            
            return htmlSuggestions, pie_chart_data, bar_chart_data, combined_chart_data
            
        except Exception as e:
            logging.info(f"Error occurred while generating investment suggestions: {e}")
            return jsonify({'message': f'Error occurred while considering preuploaded file : {e}'}), 500
        

        
        # return jsonify({
        #     "status": 200,
        #     "message": "Success",
        #     "investmentSuggestions": htmlSuggestions,
        #     "pieChartData": pie_chart_data,
        #     "barChartData": bar_chart_data,
        #     "compoundedChartData":combined_chart_data
        # }), 200
        
    except Exception as e:
        logging.error(f"Error processing personality assessment: {e}")
        return jsonify({'message': 'Error in generating suggestions with personality'}), 500
        

@app.route('/personality-assessment', methods=['POST'])
def personality_selected():
    try:
        data = request.json
        clientName = data.get('clientName')
        try :
            # clientId = data.get('client_id')
            clientId = data.get('clientId')
            investmentPersonality = data.get('investmentPersonality') # investment_personality
            financial_file = f"{clientId}.docx"
            # financial_file = f"data\{clientId}.docx" # data\EW2400.docx
            print(f"The clients ClientName is : {clientName} and their ClientId is : {clientId}")
            print(f"InvestmentPersonality received is : {investmentPersonality}")
            logging.info('Recieved Values')
            
        except Exception as e:
            logging.info(f"Error occurred while retrieving client id: {e}")
            return jsonify({'message': f'Error occurred while retrieving client id: {e}'}), 400

        try:
            # monthly_investment= data.get('monthly_investment') #10000
            # investment_period= data.get('investment_period')  #3
            monthly_investment= 10000
            investment_period= 3
            htmlSuggestions,pie_chart_data,bar_chart_data,combined_chart_data = asyncio.run(make_suggestions(investmentPersonality,clientName,financial_file,monthly_investment,investment_period))
            
        except Exception as e:
            logging.info(f"Error occurred while processing investment data: {e}")
            return jsonify({'message': f'Error occurred while processing investment data: {e}'}), 400
        
        # htmlSuggestions,pie_chart_data,bar_chart_data,combined_chart_data = asyncio.run(make_suggestions(investmentPersonality,clientName))
        
        return jsonify({
            "status": 200,
            "message": "Success",
            "investmentSuggestions": htmlSuggestions,
            "pieChartData": pie_chart_data,
            "barChartData": bar_chart_data,
            "compoundedChartData":combined_chart_data
        }), 200
                    
                    
        # return jsonify({'message':'Sab thik'}),200
        
        # if investmentPersonality == 'aggressiveInvestor':
        #     pass
    
    except Exception as e:
        logging.info(f"Error in personality assessment: {e}")
        print(f"Error occured in Investor Personality while collecting data :\n{e}")
        return jsonify({'message': 'Internal Server Error in Investor Personality'}), 500



# Route to handle generating investment suggestions
import shutil
import os
import os

def save_file_to_folder(file_storage, destination_folder):
    try:
        # Ensure the destination folder exists
        if not os.path.exists(destination_folder):
            os.makedirs(destination_folder)
        
        # Construct the destination file path
        destination_file_path = os.path.join(destination_folder, file_storage.filename)
        
        # Check if the file already exists
        if not os.path.exists(destination_file_path):
            # Save the file
            file_storage.save(destination_file_path)
            print(f"File saved to {destination_file_path}")
            return destination_file_path
        else:
            print(f"File already exists at {destination_file_path}")
            return destination_file_path
        
    except Exception as e:
        print(f"Error saving file: {e}")


@app.route('/generate-investment-suggestions', methods=['POST'])
def generate_investment_suggestions():
    try:
        assessment_file = request.files['assessmentFile']
        financial_file = request.files['financialFile']
        logging.info("Requested files")
        
        responses = extract_responses_from_docx(assessment_file)
        if not responses:
            raise Exception("Failed to extract responses from assessment file.")
        
        destination_folder = 'data'
        file_path = save_file_to_folder(financial_file, destination_folder)
        if not file_path:
            raise Exception("Failed to save financial file.")
        
        financial_data = asyncio.run(process_document(file_path))
        if not financial_data:
            raise Exception("Failed to process financial file.")
        
        logging.info(f"Received Responses from the file {responses}")
        
        personality = asyncio.run(determine_investment_personality(responses))
        if not personality:
            raise Exception("Failed to determine personality.")
        
        logging.info(f"Personality of the user is: {personality}")
        
        clientName = "Emilly Watts"
        suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality, clientName, financial_data, file_path))
        if "Error" in suggestions:
            raise Exception(suggestions)
        
        htmlSuggestions = markdown.markdown(suggestions)
        logging.info(f"Suggestions for investor: \n{suggestions}")
        
        formatSuggestions = markdown_to_text(suggestions)
        data_extracted = extract_numerical_data(suggestions)
        
        min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                          [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
        max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                          [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

        # Normalize allocations
        min_allocations = normalize_allocations(min_allocations)
        max_allocations = normalize_allocations(max_allocations)

        # Update Bar Chart Data
        bar_chart_data = {
            'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
            'datasets': [{
                'label': 'Min Allocation',
                'data': min_allocations,
                'backgroundColor': 'skyblue'
            },
            {
                'label': 'Max Allocation',
                'data': max_allocations,
                'backgroundColor': 'lightgreen'
            }]
        }

        # Similar changes can be made for the Pie Chart Data:
        all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
        num_labels = len(all_labels)
        max_allocations_for_pie = normalize_allocations(
            [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
            [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
        )
        
        # Generate colors based on the number of labels
        dynamic_colors = generate_colors(num_labels)

        # Update Pie Chart Data
        pie_chart_data = {
            'labels': all_labels,
            'datasets': [{
                'label': 'Investment Allocation',
                'data': max_allocations_for_pie,
                'backgroundColor': dynamic_colors,
                'hoverOffset': 4
            }]
        }
        
        # Prepare the data for the line chart with inflation adjustment
        initial_investment = 10000
        combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
        print(f"\nThe combined chart data is: {combined_chart_data}")
        
        return jsonify({
            "status": 200,
            "message": "Success",
            "investmentSuggestions": htmlSuggestions,
            "pieChartData": pie_chart_data,
            "barChartData": bar_chart_data,
            "compoundedChartData": combined_chart_data
        }), 200

    except Exception as e:
        logging.info(f"Error in generating investment suggestions: {e}")
        return jsonify({'message': f'Internal Server Error in Generating responses : {e}'}), 500


# #Current error code
# @app.route('/generate-investment-suggestions', methods=['POST'])
# def generate_investment_suggestions():
#     try:
#         try :
#             assessment_file = request.files['assessmentFile']
#             financial_file = request.files['financialFile']
#             logging.info(" Requested files")
#         except Exception as e:
#             logging.info(" Requested files not passed")
#             return jsonify({'message': f'Error occurred while retrieving files12: {e}'}), 400
                   
#         try :
#             try:
#                 responses = extract_responses_from_docx(assessment_file)
#             except Exception as e:
#                 logging.info(f"Failed to extract responses from assessment file: {e}")
#                 return jsonify({'message': 'Failed to extract responses from assessment file.'}), 400
            
#             try:
#                 financial_data = asyncio.run(process_document(financial_file))
#                 print(financial_data)
#                 print(f"The Financial Data is of type : {type(financial_data)}")
#                 print(financial_file)
#                 destination_folder = 'data'
#                 file_path = save_file_to_folder(financial_file, destination_folder)
#                 print(file_path)
#             except Exception as e:
#                 logging.info(f"Failed to process financial file: {e}")
#                 return jsonify({'message': 'Failed to process financial file.'}), 400

#             logging.info(f"Received Responses from the file {responses}")
#         except Exception as e:
#             logging.info("Failed to process files")
#             return jsonify({'message': f'Error occurred while processing files: {e}'}), 400

#         try:
#             # Determine investment personality
#             # monthly_investment= 10000
#             # investment_period= 3
#             # personality,monthly_investment,investment_period = asyncio.run(determine_investment_personality(responses))
            
#             personality= asyncio.run(determine_investment_personality(responses))
            
#             logging.info(f"\nPersonality of the user is : {personality}")
#             print(f"\nPersonality of the user is : {personality}")
#             print(f"Type of Investment Personality is : {type(personality)}")
            
#         except Exception as e:
#             logging.info("Failed to determine personality")
#             return jsonify({'message': f'Error occurred while determining personality: {e}'}), 400
        
#         try:
#             # Generate investment suggestions based on personality and financial data
#             clientName = "Emilly Watts"
#             print(f"{type(personality)} \n {type(financial_data)} \n {type(financial_file)} \n {type(file_path)}")
#             try:
#                 suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality,clientName, financial_data,file_path)) #monthly_investment,investment_period))
#             except Exception as e:
#                 logging.info(f"Failed to generate investment suggestions for investor: {e}")
#                 print(f"Failed to generate investment suggestions for investor: {e}")
#                 return jsonify({'message': f'Failed to generate investment suggestions for investor. : {e}'}), 400
            
#             # suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality,clientName, financial_data,monthly_investment,investment_period))
#             htmlSuggestions = markdown.markdown(suggestions)
#             logging.info(f"\Suggestions for investor : \n{suggestions}")
            
#         except Exception as e:
#             logging.info("Failed to generate suggestions")
#             return jsonify({'message': f'Error occurred while generating suggestions: {e}'}), 400


#         logging.info("Successfully generated")
#         formatSuggestions = markdown_to_text(suggestions)
#         data_extracted = extract_numerical_data(suggestions)
        
#         min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                         [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
#         max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                         [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

#         # Normalize allocations
#         min_allocations = normalize_allocations(min_allocations)
#         max_allocations = normalize_allocations(max_allocations)

#         # Update Bar Chart Data
#         bar_chart_data = {
#             'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
#             'datasets': [{
#                 'label': 'Min Allocation',
#                 'data': min_allocations,
#                 'backgroundColor': 'skyblue'
#             },
#             {
#                 'label': 'Max Allocation',
#                 'data': max_allocations,
#                 'backgroundColor': 'lightgreen'
#             }]
#         }

#         # Similar changes can be made for the Pie Chart Data:
#         all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
#         num_labels = len(all_labels)
#         max_allocations_for_pie = normalize_allocations(
#             [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
#             [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
#         )
        
#         # Generate colors based on the number of labels
#         dynamic_colors = generate_colors(num_labels)

#         # Update Pie Chart Data
#         pie_chart_data = {
#             'labels': all_labels,
#             'datasets': [{
#                 'label': 'Investment Allocation',
#                 'data': max_allocations_for_pie,
#                 'backgroundColor': dynamic_colors,
#                 'hoverOffset': 4
#             }]
#         }
        
        
#         # Prepare the data for the line chart with inflation adjustment
#         initial_investment = 10000
#         # compounded_chart_data, inflation_adjusted_chart_data = prepare_line_chart_data_with_inflation(data_extracted, initial_investment)
#         combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
#         print(f"\nThe combined chart data is : {combined_chart_data}")
        
#         return jsonify({
#             "status": 200,
#             "message": "Success",
#             "investmentSuggestions": htmlSuggestions,
#             "pieChartData": pie_chart_data,
#             "barChartData": bar_chart_data,
#             "compoundedChartData":combined_chart_data
#         }), 200

#     except Exception as e:
#         logging.info(f"Error in generating investment suggestions: {e}")
#         return jsonify({'message': 'Internal Server Error in Generating responses'}), 500




# Run the Flask application
if __name__ == '__main__':
    app.run(host='0.0.0.0',debug=True)


# # Investment suggestion template needs to be updated  :
# import streamlit as st
# import pandas as pd
# import matplotlib.pyplot as plt

# import os
# import filetype
# import docx
# import PyPDF2
# import re
# from dotenv import load_dotenv
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
# from langchain_community.docstore.in_memory import InMemoryDocstore
# from langchain_community.vectorstores import FAISS
# from langchain_community.document_loaders import Docx2txtLoader
# from langchain_core.prompts import ChatPromptTemplate
# from langchain.chains import create_retrieval_chain
# from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# from langchain.chains.combine_documents import create_stuff_documents_chain
# from langchain.memory import ConversationSummaryMemory
# import asyncio
# import numpy as np
# import json

# import google.generativeai as genai
# import pathlib
# import logging
# import sys
# import io
# import matplotlib.pyplot as plt
# import seaborn as sns
# # Import things that are needed generically
# from langchain.pydantic_v1 import BaseModel, Field
# from langchain.tools import BaseTool, StructuredTool, tool
# # Define functions to generate investment suggestions :

# load_dotenv()
# GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

# from flask import Flask, request, jsonify

# app = Flask(__name__)

# # Configure generativeai with your API key
# genai.configure(api_key=GOOGLE_API_KEY)

# import markdown
# # def convert_to_markdown(raw_text):
# #     # Replace specific text patterns with markdown syntax
# #     formatted_text = raw_text.replace('\n', '\n\n')  # Ensure newlines create paragraphs
    
# #     # Convert text into markdown format
# #     html = markdown.markdown(formatted_text)

# #     return html


# import markdown2
# from bs4 import BeautifulSoup

# def markdown_to_readable_text(md_text):
#     # Convert markdown to HTML
#     html = markdown2.markdown(md_text)

#     # Parse the HTML
#     soup = BeautifulSoup(html, "html.parser")

#     # Function to format plain text from tags
#     def format_text_from_html(soup):
#         formatted_text = ''
#         for element in soup:
#             if element.name == "h1":
#                 formatted_text += f"\n\n# {element.text.upper()} #\n\n"
#             elif element.name == "h2":
#                 formatted_text += f"\n\n## {element.text} ##\n\n"
#             elif element.name == "h3":
#                 formatted_text += f"\n\n### {element.text} ###\n\n"
#             elif element.name == "strong":
#                 formatted_text += f"**{element.text}**"
#             elif element.name == "em":
#                 formatted_text += f"_{element.text}_"
#             elif element.name == "ul":
#                 for li in element.find_all("li"):
#                     formatted_text += f"\n - {li.text}"
#             elif element.name == "ol":
#                 for idx, li in enumerate(element.find_all("li"), 1):
#                     formatted_text += f"\n {idx}. {li.text}"
#             elif element.name == "table":
#                 rows = element.find_all("tr")
#                 for row in rows:
#                     cols = row.find_all(["th", "td"])
#                     row_text = ' | '.join(col.text.strip() for col in cols)
#                     formatted_text += f"{row_text}\n"
#                 formatted_text += "\n"
#             else:
#                 formatted_text += element.text

#         return formatted_text.strip()

#     return format_text_from_html(soup)

# def markdown_to_text(md): # og solution code 
#     # Simple conversion for markdown to plain text
#     md = md.replace('**', '')
#     md = md.replace('*', '')
#     md = md.replace('_', '')
#     md = md.replace('#', '')
#     md = md.replace('`', '')
#     return md.strip()


# # import docx

# # def extract_responses_from_docx(personality_file):
# #     """
# #     Extracts responses from a Word document (.docx) where answers are typed in.

# #     Args:
# #         personality_file (UploadedFile): The file object uploaded via Streamlit.

# #     Returns:
# #         dict: A dictionary containing the questions and the typed answers.
# #     """
# #     try:
# #         doc = docx.Document(personality_file)
# #         responses = {}
# #         current_question = None

# #         # Check paragraphs
# #         for para in doc.paragraphs:
# #             text = para.text.strip()
# #             if text:
# #                 # Check if the paragraph contains a question
# #                 if "?" in text or text.endswith(":"):
# #                     current_question = text
# #                     st.write(f"Identified question: {current_question}")  # Debugging log
# #                 else:
# #                     # This is a typed answer
# #                     typed_answer = text.strip()
# #                     st.write(f"Identified typed answer: {typed_answer}")  # Debugging log
# #                     if current_question:
# #                         # If the question already has an answer, append to it (handles multiple responses)
# #                         if current_question in responses:
# #                             responses[current_question] += "; " + typed_answer
# #                         else:
# #                             responses[current_question] = typed_answer

# #             # Debugging log to understand document structure
# #             st.write(f"Processing paragraph: {text}")  # Console log for local testing

# #         # Check tables for additional responses
# #         for table in doc.tables:
# #             for row in table.rows:
# #                 for cell in row.cells:
# #                     text = cell.text.strip()
# #                     if text:
# #                         if "?" in text or text.endswith(":"):
# #                             current_question = text
# #                             st.write(f"Identified question in table: {current_question}")  # Debugging log
# #                         else:
# #                             typed_answer = text.strip()
# #                             st.write(f"Identified typed answer in table: {typed_answer}")  # Debugging log
# #                             if current_question:
# #                                 if current_question in responses:
# #                                     responses[current_question] += "; " + typed_answer
# #                                 else:
# #                                     responses[current_question] = typed_answer

# #         if responses:
# #             st.write("Extracted Responses:")
# #             for question, answer in responses.items():
# #                 st.write(f"**{question}**: {answer}")
# #         else:
# #             st.write("No responses captured. Please check the document formatting or symbols used.")

# #         return responses

# #     except Exception as e:
# #         st.write(f"Error extracting responses: {e}")  # Console log for local testing
# #         return None

# # def determine_investment_personality(responses):
# #     """
# #     Determines the investment personality based on extracted responses.

# #     Args:
# #         responses (dict): A dictionary containing the questions and the selected answers.

# #     Returns:
# #         str: The determined investment personality.
# #     """
# #     try:
# #         # Prepare input text for the chatbot based on extracted responses
# #         input_text = "User Profile:\n"
# #         for question, response in responses.items():
# #             input_text += f"{question}: {response}\n"

# #         # Introduce the chatbot's task and prompt for classification
# #         input_text += "\nYour task is to determine the investment personality based on the above profile."

# #         # Here you would send the input_text to your chatbot or classification model
# #         # For demonstration, we'll just return the input_text
# #         return input_text

# #     except Exception as e:
# #         st.write(f"Error determining investment personality: {e}")  # Console log for local testing
# #         return None

# # def extract_responses_from_docx(personality_file):
# #     try:
# #         doc = docx.Document(personality_file)
# #         responses = {}
# #         current_question = None

# #         # Check paragraphs
# #         for para in doc.paragraphs:
# #             text = para.text.strip()
# #             if text:
# #                 # Check if the paragraph contains a question
# #                 if "?" in text or text.endswith(":"):
# #                     current_question = text
# #                 else:
# #                     # This is a typed answer
# #                     typed_answer = text.strip()
# #                     if current_question:
# #                         # If the question already has an answer, append to it (handles multiple responses)
# #                         if current_question in responses:
# #                             responses[current_question] += "; " + typed_answer
# #                         else:
# #                             responses[current_question] = typed_answer

# #         # Check tables for additional responses
# #         for table in doc.tables:
# #             for row in table.rows:
# #                 for cell in row.cells:
# #                     text = cell.text.strip()
# #                     if text:
# #                         if "?" in text or text.endswith(":"):
# #                             current_question = text
# #                         else:
# #                             typed_answer = text.strip()
# #                             if current_question:
# #                                 if current_question in responses:
# #                                     responses[current_question] += "; " + typed_answer
# #                                 else:
# #                                     responses[current_question] = typed_answer

# #         return responses

# #     except Exception as e:
# #         print(f"Error extracting responses: {e}")
# #         return None

# import docx

# # # GET Method for me POST method for Frontend
# def extract_responses_from_docx(personality_file): # Using text responses parsing
#     """
#     Extracts responses from a Word document (.docx) where the selected answers are listed as text after the options.

#     Args:
#         personality_file (str): Path to the Word document file.

#     Returns:
#         dict: A dictionary containing the questions and the selected answers.
#     """
#     try:
#         doc = docx.Document(personality_file)
#         responses = {}
#         current_question = None

#         for para in doc.paragraphs:
#             text = para.text.strip()
#             if text:
#                 # Detect the beginning of a question
#                 if "?" in text:
#                     current_question = text
#                 # Detect a chosen response (assuming it follows the question and options)
#                 elif current_question and not text.startswith(("a.", "b.", "c.", "d.")):
#                     selected_answer = text
#                     responses[current_question] = selected_answer
#                     current_question = None  # Reset for the next question

#         if responses:
#             print(responses)
#             # st.write(responses)
#         else:
#             print("\nNo responses captured")
#             st.write("No responses captured")
#         return responses
#     except Exception as e:
#         print(f"Error extracting responses: {e}")
#         return None

# # def extract_responses_from_assessment(personality_file): # using boxes
# #     # Load the document
# #     # doc = Document(docx_filename)
# #     doc = docx.Document(personality_file)
    
# #     # Initialize a list to store responses
# #     responses = []
    
# #     # Iterate through each paragraph in the document
# #     for para in doc.paragraphs:
# #         text = para.text.strip()
# #         # Check if the paragraph contains a checkbox
# #         if '☒' in text or '☐' in text:
# #             # Extract the response marked with ☒
# #             if '☒' in text:
# #                 response = text.split('☒')[1].strip()
# #                 responses.append(response)
    
# #     return responses

# # import asyncio
# # # from some_generative_ai_library import GenerativeModel  # Replace with actual import

# # async def determine_investment_personality(assessment_data):
# #     try:
# #         # Prepare input text for the chatbot based on assessment data
# #         input_text = "User Profile:\n"
# #         for question, answer in assessment_data.items():
# #             input_text += f"{question}: {answer}\n"

# #         # Introduce the chatbot's task and prompt for classification
# #         input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
# #                       "- Conservative Investor\n" \
# #                       "- Moderate Investor\n" \
# #                       "- Aggressive Investor\n\n" \
# #                       "Please provide the classification below:\n"

# #         # Use your generative AI model to generate a response
# #         model = GenerativeModel('gemini-1.5-flash')
# #         response = await model.generate_content(input_text)

# #         # Determine the investment personality from the chatbot's response
# #         response_text = response.text.lower()

# #         if "conservative investor" in response_text:
# #             personality = "Conservative Investor"
# #         elif "moderate investor" in response_text:
# #             personality = "Moderate Investor"
# #         elif "aggressive investor" in response_text:
# #             personality = "Aggressive Investor"
# #         else:
# #             personality = "Unknown"

# #         return personality
# #     except Exception as e:
# #         print(f"Error generating response: {e}")
# #         return "Unknown"


# # GET Method
# async def determine_investment_personality(assessment_data): # proper code 
#     try:
#         # Prepare input text for the chatbot based on assessment data
#         input_text = "User Profile:\n"
#         for question, answer in assessment_data.items():
#             input_text += f"{question}: {answer}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
#                       "- Conservative Investor\n" \
#                       "- Moderate Investor\n" \
#                       "- Aggressive Investor\n\n" \
#                       "Please provide the classification below:\n"

#         # Use your generative AI model to generate a response
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content(input_text)

#         # Determine the investment personality from the chatbot's response
#         response_text = response.text.lower()

#         if "conservative investor" in response_text:
#             personality = "Conservative Investor"
#         elif "moderate investor" in response_text:
#             personality = "Moderate Investor"
#         elif "aggressive investor" in response_text:
#             personality = "Aggressive Investor"
#         else:
#             personality = "Unknown"

#         return personality
#     except Exception as e:
#         print(f"Error generating response: {e}")
#         return "Unknown"



# # #Load the Vector DataBase :
# async def load_vector_db(file_path): # # GET Method 
#     try:
#         print("Loading vector database...")
#         loader = Docx2txtLoader(file_path)
#         documents = loader.load()
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         text_chunks = text_splitter.split_documents(documents)
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
#         # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
#         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
#         # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

#         # vector_store = FAISS(
#         #     embedding_function=embeddings,
#         #     index=index,
#         #     docstore=InMemoryDocstore(),
#         #     index_to_docstore_id={},
#         # )
        
#         print("Vector database loaded successfully.") 
#         return vector_store.as_retriever(search_kwargs={"k": 1})
#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None


# # async def load_vector_db(file_path="client_data.txt"):
# #     try:
# #         print("Loading vector database...")
# #         with open(file_path, "r") as file:
# #             text = file.read()
        
# #         loader = Docx2txtLoader(file_path)
# #         documents = loader.load()
# #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# #         text_chunks = text_splitter.split_documents(documents) #([Document(text=text)])
# #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        
# #         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        
# #         print("Vector database loaded successfully.") 
# #         return vector_store.as_retriever(search_kwargs={"k": 1})
# #     except Exception as e:
# #         print(f"Error loading vector database: {e}")
# #         return None


# # investment_personality = "Moderate Investor"
# async def make_retrieval_chain(retriever,investmentPersonality,clientName,monthly_investment=10000,investment_period=3): # GET Method
#     """
#     Create a retrieval chain using the provided retriever.

#     Args:
#         retriever (RetrievalQA): A retriever object.

#     Returns:
#         RetrievalQA: A retrieval chain object.
#     """
#     try:
#         # global investment_personality #,summary
#         print(investmentPersonality)
#         print(clientName)
#         llm = ChatGoogleGenerativeAI(
#             #model="gemini-pro",
#             model = "gemini-1.5-flash",
#             temperature=0.7,
#             top_p=0.85,
#             google_api_key=GOOGLE_API_KEY
#         )
#         # New Template 
#         monthly_investment = str(monthly_investment)
#         investment_period = str(investment_period)
#         print(monthly_investment)
#         print(investment_period)
        
#         prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
#                 Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
#                 Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
#                 Also give the user detailed information about the investment how to invest,where to invest and how much they
#                 should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
#                 Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
#                 investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
#                 Also give the user minimum and maximum expected growth in dollars for the time horizon .
#                 Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
                
#                 Client Name: """ + clientName + """

#                 Financial Overview:
#                 - Total Assets: $100,000 (cash), $150,000 (home), $12,000 (other assets).
                
#                 - Liabilities: $200,000 mortgage at 12% interest, $400 credit card debt at 3.5%, $15,000 other loans at 10%.
                
#                 - Total Monthly Liabilities: $1,650 (Mortgage, Credit Card, Other Loans).
                

#                 You are a Financial Advisor for question-answering tasks related to the document. Based on the client's investment personality and financial details provided, generate responsible investment suggestions to achieve their financial goals while managing debts.

#                 Step-by-Step Guidance:
#                 1. Analyze Liabilities: Determine if the client's monthly investment plan is feasible after covering liabilities and expected expenses and also considering some amount for savings. If the client's monthly investment plan is not feasible after covering expenses and savings, generate investment suggestions on a smaller monthly investment plan amount if it can help the client else mention amount is too small for the client's requirementys to be made.
#                 2. Investment Strategy: Suggest a strategy where monthly investments can both generate returns and pay off debts effectively and helps client to achieve their financial goals.
#                 3. Allocation: Provide detailed allocations between growth-oriented investments and conservative investments, ensuring the client can meet their monthly debt obligations and save for their future financial goals.
#                 4. Returns: Include minimum and maximum compounded returns over 5-10 years, along with inflation-adjusted returns for clarity.
#                 5. Suggestions: Offer advice on how to use remaining funds to build wealth after clearing liabilities.
                
                
#                 Here's an example for the required Output Format(if there are comments indicated by # in the example output format then thats a side note for your reference dont write it in the response that will be generated ) :
                
#                 Investment Suggestions : 
                
                
#                 Client Name: """ + clientName + """

#                 Financial Overview: (#the data presented is just an example for your reference do not consider it as factual refere to the document provided to you and generate data based on the provided data only)
                
#                 - Total Assets: (# For ex : $100,000 (cash), $150,000 (home), $12,000 (other assets) display data available to you.)
                
#                 - Liabilities: (# For ex : $200,000 mortgage at 12% interest, $400 credit card debt at 3.5%, $15,000 other loans at 10%. display data available to you)
                
#                 - Total Monthly Liabilities: (# For ex : $1,650 (Mortgage, Credit Card, Other Loans). display data available to you)
                
#                 - Monthly Investment Amount : """ + monthly_investment + """ (# if no specific amount is specified to you then only assume  10,000 else consider the amount mention to you and just display the amount)
                
#                 - Investment Period : """ + investment_period + """  (# if no specific period is specified to you then only assume 3 years else consider the period mention to you and just display the period)

#                 Financial Analysis :(#Analyse the assets and liabilities and based on that give a suggestion for analysis generate suggestions for one of the following conditions:)
#                 (#1st condition : Everything is Positive)Based on the given Financial Conditions the client is having a good and stable income with great assets and manageable debt and liabilities.
#                 Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the clients monthly income.
#                 (# if this condition is true then ignore the other conditions and start with the Investment Suggestions)
                
#                 (#2nd condition : Everything is temporarily Negative) Based on the given Financial Conditions the client is facing a bad income for now but have great assets and manageable debt and liabilities.
#                 Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the client's monthly income but the client might not be able to sustain the monthly investment amount that they are planning.)
#                 Instead I would like to recommend this amount to the client for their monthly investment : (#Mention a feasible amount to the client for monthly investment and start suggesting investments based on this amount and not the previous amount being taken into consideration)
                
#                 (#3rd condition : Everything is Negative) Based on the given Financial Conditions the client is facing a bad income and doesnt have good assets to manage the debts and liabilities of the client and in such a condition this monthly investment amount is not feasible.
#                 Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is not manageable for the client's monthly income and so the client might not be able to sustain the monthly investment amount that they are planning to do.)
#                 I would like to recommend this amount to the client for monthly investment : (# Mention a minimum amount to the client for monthly investment if possible else just say the client should first prioritize on savings and generating more income to manage their debts and liabilities first and so dont give any investment suggestions to the client.)
                
#                 (#If the financial is 1 or 2 only then give investment suggestions to the client)
#                 Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

#                 Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

#                 Investment Allocation: (#remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

#                 Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
#                 How to Invest: Diversify across various asset classes like:  (#Give allocations % as well)
                
#                 Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                
#                 ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                
#                 Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                
#                 Consider investing in blue-chip companies or growth sectors like technology. 
                
                
#                 Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


#                 Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
#                 Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

#                 Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
#                 How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
#                 Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
#                 Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
#                 Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
#                 Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
#                 Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
#                 Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.


#                 Time Horizon and Expected Returns:

#                 Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
#                 Minimum Expected Annual Return: 4% - 6% 
                
                
#                 Maximum Expected Annual Return: 8% - 10% 
                
                
#                 Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, (# consider the monthly investment amount and give returns based on that only) $10,000 could grow to approximately 17,908 in 10 years.
#                 Minimum Expected Growth in Dollars: 
                
#                 4,000−6,000 (over 10 years) 
                
                
#                 Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
#                 Inflation Adjusted Returns:(#do not write this part inside the bracket just give answer,assume US inflation rate assume 3% if you dont know, and give the investment returns value that was suggested by you for the considered monthly investment amount after 3,5,10years of growth mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
#                 Rationale for Investment Suggestions:

#                 This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
#                 Important Considerations:

#                 Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

#                 Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

#                 Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
#                 <context>
#                 {context}
#                 </context>
#                 Question: {input}"""
                
#         # # latest version gives some info about the client conditions but not in detail 
#         # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
#         #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
#         #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
#         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
#         #         should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
#         #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
#         #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
#         #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
#         #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
#         #         Here's an example for the required Output Format :

#         #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

#         #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

#         #         Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

#         #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
#         #         How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
#         #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
#         #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
#         #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
#         #         Consider investing in blue-chip companies or growth sectors like technology. 
#         #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


#         #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
#         #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

#         #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
#         #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
#         #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
#         #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
#         #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
#         #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
#         #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
#         #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

#         #         Time Horizon and Expected Returns:

#         #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
#         #         Minimum Expected Annual Return: 4% - 6% 
                
                
#         #         Maximum Expected Annual Return: 8% - 10% 
                
#         #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, a 10,000 investment could grow to approximately 17,908 in 10 years.
#         #         Minimum Expected Growth in Dollars: 
                
#         #         4,000−6,000 (over 10 years) 
                
                
#         #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
#         #         Inflation Adjusted Returns:(do not write this part inside the bracket just give answer,assume US inflation rate, and give the investment returns value that was suggested by you  for $10k investment after 3,5,10years of growth  mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
#         #         Rationale for Investment Suggestions:

#         #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
#         #         Important Considerations:

#         #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

#         #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

#         #         Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
#         #         <context>
#         #         {context}
#         #         </context>
#         #         Question: {input}"""
        
        
        
        
#         # # Working code but gives clientname in brackets 
#         # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
#         #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality provided to you.
#         #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
#         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
#         #         should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
#         #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
#         #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
#         #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
#         #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
#         #         Here's an example for the required Output Format :

#         #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

#         #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

#         #         Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

#         #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
#         #         How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
#         #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
#         #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
#         #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
#         #         Consider investing in blue-chip companies or growth sectors like technology. 
#         #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


#         #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
#         #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

#         #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
#         #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
#         #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
#         #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
#         #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
#         #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
#         #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
#         #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

#         #         Time Horizon and Expected Returns:

#         #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
#         #         Minimum Expected Annual Return: 4% - 6% 
                
                
#         #         Maximum Expected Annual Return: 8% - 10% 
                
#         #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, a 10,000 investment could grow to approximately 17,908 in 10 years.
#         #         Minimum Expected Growth in Dollars: 
                
#         #         4,000−6,000 (over 10 years) 
                
                
#         #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
#         #         Inflation Adjusted Returns:(do not write this part inside the bracket just give answer,assume US inflation rate, and give the investment returns value that was suggested by you  for $10k investment after 3,5,10years of growth  mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
#         #         Rationale for Investment Suggestions:

#         #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
#         #         Important Considerations:

#         #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

#         #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

#         #         <context>
#         #         {context}
#         #         </context>
#         #         Question: {input}"""
                

#         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

#         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
#         combine_docs_chain = None  

#         if retriever is not None :  
#             retriever_chain = create_retrieval_chain(retriever,document_chain) 
#             print(retriever_chain)
#             return retriever_chain
#         else:
#             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
#             return None

#     except Exception as e:
#         print(f"Error in creating chain: {e}")
#         return None

# import streamlit as st
# import json
# import matplotlib.pyplot as plt
# import io


# async def process_document(file_path): # GET Method
#     try:
#         print("Processing the document")
#         file_type = filetype.guess(file_path)
#         if file_type is not None:
#             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 # Await the coroutine to extract text and tables
#                 return await extract_text_and_tables_from_word(file_path)
#             elif file_type.mime == "application/pdf":
#                 return await extract_text_from_pdf(file_path)
#         return None
#     except Exception as e:
#         print(f"Error processing document: {e}")
#         return None

# # Async function to extract text from a PDF file
# async def extract_text_from_pdf(pdf_file_path): # GET Method
#     try:
#         print("Processing pdf file")
#         with open(pdf_file_path, "rb") as pdf_file:
#             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#             text_content = []
#             for page_num in range(pdf_reader.numPages):
#                 page = pdf_reader.getPage(page_num)
#                 text_content.append(page.extract_text())
#             return "\n".join(text_content)
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None

# # Async function to extract text and tables from a Word document
# async def extract_text_and_tables_from_word(docx_file_path): # GET Method
#     try:
#         print("Extracting text and tables from word file")
#         doc = docx.Document(docx_file_path)
#         text_content = []
#         tables_content = []

#         for para in doc.paragraphs:
#             text_content.append(para.text)

#         for table in doc.tables:
#             table_data = []
#             for row in table.rows:
#                 row_data = []
#                 for cell in row.cells:
#                     row_data.append(cell.text.strip())
#                 table_data.append(row_data)
#             tables_content.append(table_data)
#         print("Extracted text from word file")
#         return "\n".join(text_content), tables_content
#     except Exception as e:
#         print(f"Error extracting text and tables from Word document: {e}")
#         return None, None



# async def validate_document_content(text, tables):
#     """
#     Validates the content of the document.

#     Args:
#         text (str): Extracted text content from the document.
#         tables (list): Extracted tables content from the document.

#     Returns:
#         tuple: Client name and validation errors.
#     """
#     errors = []
    
#     # Extract client name
#     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
#     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

#     # Define required sections
#     required_sections = [
#         "YOUR RETIREMENT GOAL",
#         "YOUR OTHER MAJOR GOALS",
#         "YOUR ASSETS AND LIABILITIES",
#         "MY LIABILITIES",
#         "YOUR CURRENT ANNUAL INCOME"
#     ]

#     # Check for the presence of required sections
#     for section in required_sections:
#         if section not in text:
#             errors.append(f"* {section} section missing.")
    
#     # Define table field checks
#     table_checks = {
#         "YOUR RETIREMENT GOAL": [
#             r"When do you plan to retire\? \(age or date\)",
#             r"Social Security Benefit \(include expected start date\)",
#             r"Pension Benefit \(include expected start date\)",
#             r"Other Expected Income \(rental, part-time work, etc.\)",
#             r"Estimated Annual Retirement Expense"
#         ],
#         "YOUR OTHER MAJOR GOALS": [
#             r"GOAL", r"COST", r"WHEN"
#         ],
#         "YOUR ASSETS AND LIABILITIES": [
#             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
#             r"Current Value", r"Annual Contributions"
#         ],
#         "MY LIABILITIES": [
#             r"Balance", r"Interest Rate", r"Monthly Payment"
#         ]
#     }

#     # Validate table content
#     for section, checks in table_checks.items():
#         section_found = False
#         for table in tables:
#             table_text = "\n".join(["\t".join(row) for row in table])
#             if section in table_text:
#                 section_found = True
#                 for check in checks:
#                     if not re.search(check, table_text, re.IGNORECASE):
#                         errors.append(f"* Missing or empty field in {section} section: {check}")
#                 break
#         if not section_found:
#             errors.append(f"* {section} section missing.")

#     return client_name, errors


# # RUN Button :
# async def generate_investment_suggestions(investment_personality, context): # # GET Method for py , for front end its Post API
    
#     # retriever = asyncio.run(load_vector_db("uploaded_file"))

#     retriever = await load_vector_db("uploaded_file")
#     # retriever = await load_vector_db("data\Financial_Investment_1.docx") 

#     chain = await make_retrieval_chain(retriever)

#     # chain = asyncio.run(make_retrieval_chain(retriever))
    
#     if chain is not None:
#         # summary = context
#         # query = summary + "\n" + investment_personality
#         query = str(investment_personality)
#         response = chain.invoke({"input": query})
#         format_response = markdown_to_text(response['answer'])
#         return format_response
#         # st.write(format_response)

#         # handle_graph(response['answer'])

#     else:
#         st.error("Failed to create the retrieval chain. Please upload a valid document.")



# # Generate Infographics : Best Code so far:


# import seaborn as sns
# import re
# from collections import defaultdict
# import matplotlib.pyplot as plt
# import streamlit as st
# import numpy as np

# # def extract_numerical_data(response):
# #     # Define patterns to match different sections and their respective allocations
# #     patterns = {
# #         'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
# #         'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
# #         'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
# #     }

# #     data = defaultdict(dict)

# #     for section, pattern in patterns.items():
# #         match = pattern.search(response)
# #         if match:
# #             investments_text = match.group(1)
# #             # Extract individual investment types and their allocations
# #             investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
# #             for investment_match in investment_pattern.findall(investments_text):
# #                 investment_type, min_allocation, max_allocation = investment_match
# #                 data[section][investment_type.strip()] = {
# #                     'min': min_allocation,
# #                     'max': max_allocation
# #                 }

# #     # Update: Extracting time horizon and expected returns using a more flexible pattern
# #     time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
# #     expected_return_pattern = re.compile(r'(Minimum|Maximum) Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
# #     min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
# #     max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

# #     time_horizon_match = time_horizon_pattern.search(response)
# #     expected_return_matches = expected_return_pattern.findall(response)
# #     min_growth_match = min_growth_pattern.search(response)
# #     max_growth_match = max_growth_pattern.search(response)

# #     if time_horizon_match:
# #         data['Time Horizon'] = {
# #             'min_years': time_horizon_match.group(1),
# #             'max_years': time_horizon_match.group(2)
# #         }

# #     # Correct extraction for both Minimum and Maximum Expected Annual Return
# #     for match in expected_return_matches:
# #         if match[0].lower() == "minimum":
# #             data['Expected Annual Return'] = {
# #                 'min': match[1],
# #                 'max': match[2]  # Extract min and max for minimum expected return
# #             }
# #         elif match[0].lower() == "maximum":
# #             data['Expected Annual Return'].update({
# #                 'min': match[1],
# #                 'max': match[2]  # Update for maximum expected return
# #             })

# #     if min_growth_match:
# #         data['Expected Growth in Dollars'] = {
# #             'min': min_growth_match.group(1),
# #             'max': min_growth_match.group(2)
# #         }

# #     if max_growth_match:
# #         data['Expected Growth in Dollars'] = {
# #             'min': max_growth_match.group(1),
# #             'max': max_growth_match.group(2)
# #         }

# #     return data


# def extract_numerical_data(response): # curr version but cant capture annual return 
#     # Define patterns to match different sections and their respective allocations
#     patterns = {
#         'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
#         'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
#         'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
#     }

#     data = defaultdict(dict)

#     for section, pattern in patterns.items():
#         match = pattern.search(response)
#         if match:
#             investments_text = match.group(1)
#             # Extract individual investment types and their allocations
#             investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
#             for investment_match in investment_pattern.findall(investments_text):
#                 investment_type, min_allocation, max_allocation = investment_match
#                 data[section][investment_type.strip()] = {
#                     'min': min_allocation,
#                     'max': max_allocation
#                 }

#     # Extract time horizon and expected returns
#     time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
#     min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
#     max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

#     time_horizon_match = time_horizon_pattern.search(response)
#     min_return_match = min_return_pattern.search(response)
#     max_return_match = max_return_pattern.search(response)
#     min_growth_match = min_growth_pattern.search(response)
#     max_growth_match = max_growth_pattern.search(response)

#     if time_horizon_match:
#         data['Time Horizon'] = {
#             'min_years': time_horizon_match.group(1),
#             'max_years': time_horizon_match.group(2)
#         }

#     if min_return_match:
#         data['Expected Annual Return'] = {
#             'min': min_return_match.group(1),
#             'max': min_return_match.group(2)
#         }

#     if max_return_match:
#         data['Expected Annual Return'] = {
#             'min': max_return_match.group(1),
#             'max': max_return_match.group(2)
#         }

#     if min_growth_match:
#         data['Expected Growth in Dollars'] = {
#             'min': min_growth_match.group(1),
#             'max': min_growth_match.group(2)
#         }

#     if max_growth_match:
#         data['Expected Growth in Dollars'] = {
#             'min': max_growth_match.group(1),
#             'max': max_growth_match.group(2)
#         }

#     return data

# def normalize_allocations(allocations):
#     total = sum(allocations)
#     if total == 100:
#         return allocations
#     return [round((allocation / total) * 100, 2) for allocation in allocations]



# # def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4): # previous version
# #     try:
# #         # Print data_extracted to debug the structure
# #         print("Data extracted:", data_extracted)

# #         # Check if 'Expected Annual Return' and 'Time Horizon' exist and have the expected keys
# #         if 'Expected Annual Return' not in data_extracted :
# #             print("'Expected Annual Return' missing in data_extracted")
# #             data_extracted['Expected Annual Return']['min'] = 6
# #             data_extracted['Expected Annual Return']['max'] = 8
# #             min_return = 6
# #             max_return = 8
# #         else :
# #             min_return = float(data_extracted['Expected Annual Return'].get('min', '0').strip('%'))
# #             max_return = float(data_extracted['Expected Annual Return'].get('max', '0').strip('%'))
                
# #         min_years = int(data_extracted['Time Horizon'].get('min_years', 1))  # Default to 1 year if missing
# #         max_years = int(data_extracted['Time Horizon'].get('max_years', 10))  # Default to 10 years if missing

# #         def calculate_compounded_return(principal, rate, years):
# #             return principal * (1 + rate / 100) ** years

# #         def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
# #             return nominal_return / (1 + inflation_rate / 100) ** years

# #         labels = list(range(1, max_years + 1))  # Years for the x-axis
# #         min_compounded = []
# #         max_compounded = []
# #         min_inflation_adjusted = []
# #         max_inflation_adjusted = []

# #         for year in labels:
# #             # Calculate nominal compounded returns
# #             min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
# #             max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

# #             # Calculate inflation-adjusted compounded returns
# #             min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
# #             max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

# #             # Append results
# #             min_compounded.append(min_compounded_value)
# #             max_compounded.append(max_compounded_value)
# #             min_inflation_adjusted.append(min_inflation_value)
# #             max_inflation_adjusted.append(max_inflation_value)

# #         # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
# #         combined_chart_data = {
# #             'labels': labels,
# #             'datasets': [
# #                 {
# #                     'label': 'Minimum Compounded Return',
# #                     'data': min_compounded,
# #                     'borderColor': 'rgb(255, 99, 132)',  # Red color
# #                     'fill': False
# #                 },
# #                 {
# #                     'label': 'Maximum Compounded Return',
# #                     'data': max_compounded,
# #                     'borderColor': 'rgb(54, 162, 235)',  # Blue color
# #                     'fill': False
# #                 },
# #                 {
# #                     'label': 'Min Inflation Adjusted Return',
# #                     'data': min_inflation_adjusted,
# #                     'borderColor': 'rgb(75, 192, 192)',  # Light blue
# #                     'borderDash': [5, 5],  # Dashed line for distinction
# #                     'fill': False
# #                 },
# #                 {
# #                     'label': 'Max Inflation Adjusted Return',
# #                     'data': max_inflation_adjusted,
# #                     'borderColor': 'rgb(153, 102, 255)',  # Light purple
# #                     'borderDash': [5, 5],  # Dashed line for distinction
# #                     'fill': False
# #                 }
# #             ]
# #         }
# #     except KeyError as e:
# #         print(f"KeyError occurred: {e}")
# #         return jsonify({'message': f'Key Error: {e}'}), 400
# #     except Exception as e:
# #         print(f"Error occurred while preparing data for combined line chart: {e}")
# #         return jsonify({'message': 'Internal Server Error in creating line chart'}), 500

# #     return combined_chart_data


# import datetime  # Import the datetime module to get the current year
# # uodated to have current year
# def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
#     try:
#         # Get the current year
#         curr_year = datetime.datetime.now().year

#         # Print data_extracted to debug the structure
#         print("Data extracted:", data_extracted)

#         # Check if 'Expected Annual Return' and 'Time Horizon' exist and have the expected keys
#         if 'Expected Annual Return' not in data_extracted:
#             print("'Expected Annual Return' missing in data_extracted")
#             data_extracted['Expected Annual Return']['min'] = 6
#             data_extracted['Expected Annual Return']['max'] = 8
#             min_return = 6
#             max_return = 8
#         else:
#             min_return = float(data_extracted['Expected Annual Return'].get('min', '0').strip('%'))
#             max_return = float(data_extracted['Expected Annual Return'].get('max', '0').strip('%'))

#         min_years = int(data_extracted['Time Horizon'].get('min_years', 1))  # Default to 1 year if missing
#         max_years = int(data_extracted['Time Horizon'].get('max_years', 10))  # Default to 10 years if missing

#         def calculate_compounded_return(principal, rate, years):
#             return principal * (1 + rate / 100) ** years

#         def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
#             return nominal_return / (1 + inflation_rate / 100) ** years

#         # Create labels for the next 10 years starting from the current year
#         labels = list(range(curr_year, curr_year + max_years))

#         min_compounded = []
#         max_compounded = []
#         min_inflation_adjusted = []
#         max_inflation_adjusted = []

#         for year in range(1, max_years + 1):
#             # Calculate nominal compounded returns
#             min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
#             max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

#             # Calculate inflation-adjusted compounded returns
#             min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
#             max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

#             # Append results
#             min_compounded.append(min_compounded_value)
#             max_compounded.append(max_compounded_value)
#             min_inflation_adjusted.append(min_inflation_value)
#             max_inflation_adjusted.append(max_inflation_value)

#         # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
#         combined_chart_data = {
#             'labels': labels,  # Current year and the next 10 years
#             'datasets': [
#                 {
#                     'label': 'Minimum Compounded Return',
#                     'data': min_compounded,
#                     'borderColor': 'rgb(255, 99, 132)',  # Red color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Maximum Compounded Return',
#                     'data': max_compounded,
#                     'borderColor': 'rgb(54, 162, 235)',  # Blue color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Min Inflation Adjusted Return',
#                     'data': min_inflation_adjusted,
#                     'borderColor': 'rgb(75, 192, 192)',  # Light blue
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 },
#                 {
#                     'label': 'Max Inflation Adjusted Return',
#                     'data': max_inflation_adjusted,
#                     'borderColor': 'rgb(153, 102, 255)',  # Light purple
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 }
#             ]
#         }
#     except KeyError as e:
#         print(f"KeyError occurred: {e}")
#         return jsonify({'message': f'Key Error: {e}'}), 400
#     except Exception as e:
#         print(f"Error occurred while preparing data for combined line chart: {e}")
#         return jsonify({'message': 'Internal Server Error in creating line chart'}), 500

#     return combined_chart_data



# # def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
# #     try:
# #         min_return = float(data_extracted['Expected Annual Return']['min'].strip('%'))
# #         max_return = float(data_extracted['Expected Annual Return']['max'].strip('%'))
# #         min_years = int(data_extracted['Time Horizon']['min_years'])
# #         max_years = int(data_extracted['Time Horizon']['max_years'])

# #         def calculate_compounded_return(principal, rate, years):
# #             return principal * (1 + rate / 100) ** years

# #         def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
# #             return nominal_return / (1 + inflation_rate / 100) ** years

# #         labels = list(range(1, max_years + 1))  # Years for the x-axis
# #         min_compounded = []
# #         max_compounded = []
# #         min_inflation_adjusted = []
# #         max_inflation_adjusted = []

# #         for year in labels:
# #             # Calculate nominal compounded returns
# #             min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
# #             max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

# #             # Calculate inflation-adjusted compounded returns
# #             min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
# #             max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

# #             # Append results
# #             min_compounded.append(min_compounded_value)
# #             max_compounded.append(max_compounded_value)
# #             min_inflation_adjusted.append(min_inflation_value)
# #             max_inflation_adjusted.append(max_inflation_value)

# #         # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
# #         combined_chart_data = {
# #             'labels': labels,
# #             'datasets': [
# #                 {
# #                     'label': 'Minimum Compounded Return',
# #                     'data': min_compounded,
# #                     'borderColor': 'rgb(255, 99, 132)',  # Red color
# #                     'fill': False
# #                 },
# #                 {
# #                     'label': 'Maximum Compounded Return',
# #                     'data': max_compounded,
# #                     'borderColor': 'rgb(54, 162, 235)',  # Blue color
# #                     'fill': False
# #                 },
# #                 {
# #                     'label': 'Min Inflation Adjusted Return',
# #                     'data': min_inflation_adjusted,
# #                     'borderColor': 'rgb(75, 192, 192)',  # Light blue
# #                     'borderDash': [5, 5],  # Dashed line for distinction
# #                     'fill': False
# #                 },
# #                 {
# #                     'label': 'Max Inflation Adjusted Return',
# #                     'data': max_inflation_adjusted,
# #                     'borderColor': 'rgb(153, 102, 255)',  # Light purple
# #                     'borderDash': [5, 5],  # Dashed line for distinction
# #                     'fill': False
# #                 }
# #             ]
# #         }
# #     except Exception as e:
# #         print(f"Error occurred while preparing data for combined line chart: {e}")
# #         return jsonify({'message': 'Internal Server Error in creating line chart'}), 500
# #     return combined_chart_data



# def plot_investment_allocations(data):
#     # Create subplots with a large figure size
#     fig, axes = plt.subplots(2, 1, figsize= (16,10)) #(28, 15))  # Adjust size as needed

#     # Plot Growth-Oriented Investments
#     growth_data = data['Growth-Oriented Investments']
#     growth_labels = list(growth_data.keys())
#     growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
#     growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

#     axes[0].bar(growth_labels, growth_min, color='skyblue', label='Min Allocation')
#     axes[0].bar(growth_labels, growth_max, bottom=growth_min, color='lightgreen', label='Max Allocation')
#     axes[0].set_title('Growth-Oriented Investments', fontsize=16)
#     axes[0].set_ylabel('Percentage Allocation', fontsize=14)
#     axes[0].set_xlabel('Investment Types', fontsize=14)
#     axes[0].tick_params(axis='x', rotation=45, labelsize=12)
#     axes[0].tick_params(axis='y', labelsize=12)
#     axes[0].legend()

#     # Plot Conservative Investments
#     conservative_data = data['Conservative Investments']
#     conservative_labels = list(conservative_data.keys())
#     conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
#     conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

#     axes[1].bar(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
#     axes[1].bar(conservative_labels, conservative_max, bottom=conservative_min, color='lightgreen', label='Max Allocation')
#     axes[1].set_title('Conservative Investments', fontsize=16)
#     axes[1].set_ylabel('Percentage Allocation', fontsize=14)
#     axes[1].set_xlabel('Investment Types', fontsize=14)
#     axes[1].tick_params(axis='x', rotation=45, labelsize=12)
#     axes[1].tick_params(axis='y', labelsize=12)
#     axes[1].legend()

#     # Tight layout for better spacing
#     plt.tight_layout()
#     plt.show()
#     return fig


# def plot_pie_chart(data):
#     fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

#     # Combine all investment data for pie chart
#     all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
#     labels = list(all_data.keys())
#     sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
#     colors = plt.cm.Paired(range(len(labels)))

#     wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
#     ax.set_title('Investment Allocation')

#     # Add legend
#     ax.legend(wedges, labels, title="Investment Types", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

#     return fig



# def bar_chart(data):
#     fig, ax = plt.subplots(figsize=(12, 8))  # Increased size

#     # Data for plotting
#     categories = list(data.keys())
#     values_min = [int(data[cat]['min'].strip('%')) for cat in categories]
#     values_max = [int(data[cat]['max'].strip('%')) for cat in categories]

#     x = range(len(categories))

#     ax.bar(x, values_min, width=0.4, label='Min Allocation', color='skyblue', align='center')
#     ax.bar(x, values_max, width=0.4, label='Max Allocation', color='lightgreen', align='edge')

#     ax.set_xticks(x)
#     ax.set_xticklabels(categories, rotation=45, ha='right')
#     ax.set_xlabel('Investment Categories')
#     ax.set_ylabel('Percentage Allocation')
#     ax.set_title('Investment Allocation')
#     ax.legend()

#     plt.tight_layout()
#     return fig


# import random

# def generate_colors(n):
#     """
#     Generate 'n' random RGB colors.

#     Args:
#         n (int): Number of colors to generate.
    
#     Returns:
#         list: A list of RGB colors in 'rgb(r, g, b)' format.
#     """
#     colors = []
#     for _ in range(n):
#         r = random.randint(0, 255)
#         g = random.randint(0, 255)
#         b = random.randint(0, 255)
#         colors.append(f'rgb({r}, {g}, {b})')
    
#     return colors


# import plotly.graph_objects as go
# import numpy as np

 
# # def client_form():
# #     st.title("Client Details Form")

# #     with st.form("client_form"):
# #         st.header("Personal Information")
# #         client_name = st.text_input("Client Name")
# #         co_client_name = st.text_input("Co-Client Name")
# #         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
# #         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
# #         today_date = st.date_input("Today's Date")
        
# #         st.header("Financial Information")
# #         current_assets = st.text_area("Current Assets (e.g., type and value)")
# #         liabilities = st.text_area("Liabilities (e.g., type and amount)")
# #         annual_income = st.text_area("Current Annual Income (source and amount)")
# #         annual_contributions = st.text_area("Annual Contributions (e.g., retirement savings)")

# #         st.header("Insurance Information")
# #         life_insurance = st.text_input("Life Insurance (e.g., coverage amount)")
# #         disability_insurance = st.text_input("Disability Insurance (e.g., coverage amount)")
# #         long_term_care = st.text_input("Long-Term Care Insurance (e.g., coverage amount)")

# #         st.header("Estate Planning")
# #         will_status = st.radio("Do you have a will?", ["Yes", "No"])
# #         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
# #         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
# #         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

# #         # Submit button
# #         submitted = st.form_submit_button("Submit")

# #         if submitted:
# #             # Save form data
# #             form_data = {
# #                 "Client Name": client_name,
# #                 "Co-Client Name": co_client_name,
# #                 "Client Age": client_age,
# #                 "Co-Client Age": co_client_age,
# #                 "Today's Date": str(today_date),
# #                 "Current Assets": current_assets,
# #                 "Liabilities": liabilities,
# #                 "Annual Income": annual_income,
# #                 "Annual Contributions": annual_contributions,
# #                 "Life Insurance": life_insurance,
# #                 "Disability Insurance": disability_insurance,
# #                 "Long-Term Care Insurance": long_term_care,
# #                 "Will Status": will_status,
# #                 "Trust Status": trust_status,
# #                 "Power of Attorney": power_of_attorney,
# #                 "Healthcare Proxy": healthcare_proxy,
# #             }
            
# #             # Save to a file or database
# #             with open("client_data.txt", "a") as f:
# #                 f.write(str(form_data) + "\n")
            
# #             st.success("Form submitted successfully!")
# #             st.session_state.page = "main"  # Redirect back to main page after form submission


# from datetime import date  # Make sure to import the date class


# # Function to parse financial data from the text
# import re

# def parse_financial_data(text_content):
#     assets = []
#     liabilities = []

#     # Define regex patterns to capture text following headings
#     asset_pattern = re.compile(r"MY ASSETS:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)
#     liability_pattern = re.compile(r"LIABILITIES:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)

#     # Extract assets
#     asset_matches = asset_pattern.findall(text_content)
#     if asset_matches:
#         asset_text = asset_matches[0]
#         # Further processing to extract individual asset values if they are detailed
#         asset_lines = asset_text.split('\n')
#         for line in asset_lines:
#             match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
#             if match:
#                 asset_value = float(match.group().replace(",", ""))
#                 assets.append(asset_value)

#     # Extract liabilities
#     liability_matches = liability_pattern.findall(text_content)
#     if liability_matches:
#         liability_text = liability_matches[0]
#         # Further processing to extract individual liability values if they are detailed
#         liability_lines = liability_text.split('\n')
#         for line in liability_lines:
#             match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
#             if match:
#                 liability_value = float(match.group().replace(",", ""))
#                 liabilities.append(liability_value)

#     print("Assets Found:", assets)
#     print("Liabilities Found:", liabilities)

#     return assets, liabilities



# # Function to extract numerical values from a text input
# def extract_numeric(value):
#     try:
#         return float(re.sub(r'[^\d.]', '', value))  # Remove non-numeric characters and convert to float
#     except ValueError:
#         return 0


# # plots graph from the details of the form :


# def is_float(value):
#     try:
#         float(value)
#         return True
#     except ValueError:
#         return False



# def plot_assets_liabilities_pie_chart(assets, liabilities, threshold=50): # best plot 
#     """
#     Plots separate pie charts for assets and liabilities. If there are any categories
#     below a specified threshold, they are plotted in an additional small pie chart.
    
#     Parameters:
#     - assets: dict, keys are asset names, values are their amounts.
#     - liabilities: dict, keys are liability names, values are their amounts.
#     - threshold: int, percentage threshold below which segments are considered small.
#     """
#     # Update matplotlib settings to increase the font size globally
#     # plt.rcParams.update({'font.size': 32})

#     plt.rcParams.update({'font.size': 16})

#     def plot_pie(data, title):
#         # Filter out zero values and create a summary for small segments
#         total = sum(data.values())
#         filtered_data = {k: v for k, v in data.items() if (v / total) >= threshold / 100}
#         small_segments = {k: v for k, v in data.items() if (v / total) < threshold / 100}
#         small_total = sum(small_segments.values())

#         # Plotting logic
#         if small_segments:
#             fig, (ax_main, ax_small) = plt.subplots(1, 2, figsize=(30, 15))  # Side-by-side layout
#         else:
#             fig, ax_main = plt.subplots(figsize=(30, 20))  # Only main chart with larger size

#             # fig, ax_main = plt.subplots(figsize=(10, 10))  # Only main chart with larger size

#         # Plot main pie chart
#         labels_main = list(filtered_data.keys()) + ([f"Other small {title}"] if small_segments else [])
#         values_main = list(filtered_data.values()) + ([small_total] if small_segments else [])
#         wedges_main, texts_main, autotexts_main = ax_main.pie(
#             values_main, labels=labels_main, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
#             startangle=140, textprops={'fontsize': 28} #18}  # Larger font size for labels
#         )

#         ax_main.set_title(title, fontsize=20)
#         # Position legend to the right of the plot to avoid overlapping
#         ax_main.legend(wedges_main, labels_main, title="Categories", loc="upper right", bbox_to_anchor=(0.001, 0.9), fontsize= 28)#14)

#         if small_segments:
#             # Plot additional small pie chart for small segments
#             labels_small = list(small_segments.keys())
#             values_small = list(small_segments.values())
#             wedges_small, texts_small, autotexts_small = ax_small.pie(
#                 values_small, labels=labels_small, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
#                 startangle=140, textprops={'fontsize': 24} #14}  # Consistent label size for small chart
#             )
#             ax_small.set_title(f"Small Segments of {title}", fontsize=20)
#             # Position legend to the right of the small pie chart but slightly lower to avoid overlap with the main chart's legend
#             ax_small.legend(wedges_small, labels_small, title="Small Categories", loc="center left", bbox_to_anchor=(1.2, 0.3), fontsize= 22)#12)

#         st.pyplot(fig)

#     # Convert valid entries to float, ensuring only numeric values are considered
#     assets = {k: float(v) for k, v in assets.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}
#     liabilities = {k: float(v) for k, v in liabilities.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}

#     # Plot pie charts
#     plot_pie(assets, 'Distribution of Assets')
#     plot_pie(liabilities, 'Distribution of Liabilities')

# # def plot_assets_liabilities_pie_chart(assets, liabilities):# properly plots a big and 1 small pie chart for both assets and liability
# #     # Filter and convert values to float, handle non-numeric or empty inputs
# #     filtered_assets = {k: float(v) for k, v in assets.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}
# #     filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}

# #     # Combine assets and liabilities for total calculation
# #     all_values = {**filtered_assets, **filtered_liabilities}
# #     total_value = sum(all_values.values())

# #     # Separate main and small segments
# #     main_segments = {k: v for k, v in all_values.items() if (v / total_value) >= 0.05}
# #     small_segments = {k: v for k, v in all_values.items() if (v / total_value) < 0.05}
# #     small_total = sum(small_segments.values())

# #     # Prepare data for main pie chart
# #     main_labels = list(main_segments.keys()) + (["Others"] if small_segments else [])
# #     main_values = list(main_segments.values()) + ([small_total] if small_segments else [])

# #     # Prepare data for small pie chart (only if there are small segments)
# #     small_labels = list(small_segments.keys())
# #     small_values = list(small_segments.values())

# #     fig, ax = plt.subplots(figsize=(8, 6))

# #     # Plot main pie chart
# #     wedges, texts, autotexts = ax.pie(
# #         main_values,
# #         labels=main_labels,
# #         autopct='%1.1f%%',
# #         startangle=140,
# #         colors=plt.cm.Paired.colors,
# #     )

# #     # Explode the "Others" slice
# #     if small_segments:
# #         others_index = main_labels.index("Others")
# #         wedges[others_index].set_edgecolor('white')
# #         # wedges[others_index].set_linestyle('--')
# #         wedges[others_index].set_linewidth(2)
# #         wedges[others_index].set_hatch('/')

# #     ax.set_title('Assets and Liabilities Distribution')

# #     # Draw a second pie chart for "Others"
# #     if small_segments:
# #         fig2, ax2 = plt.subplots(figsize=(8, 6))
# #         wedges_small, texts_small, autotexts_small = ax2.pie(
# #             small_values,
# #             labels=small_labels,
# #             autopct='%1.1f%%',
# #             startangle=140,
# #             colors=plt.cm.Pastel1.colors
# #         )

# #         ax2.set_title('Detailed View of "Others" Categories')

# #     plt.tight_layout()
# #     st.pyplot(fig)
# #     if small_segments:
# #         st.pyplot(fig2)



# def save_data_to_file(form_data):
#     file_path = 'client_data.txt'
#     with open(file_path, 'a') as file:
#         file.write(str(form_data) + "\n")
#     # st.success(f"Form data saved to {file_path}")
#     print(f"Form data saved to {file_path}")
    

# def client_form():
#     st.title("Client Details Form")

#     with st.form("client_form"):
#         st.header("Personal Information")
#         client_name = st.text_input("Client Name")
#         co_client_name = st.text_input("Co-Client Name")
#         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
#         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
#         today_date = st.date_input("Today's Date")

#         st.header("Your Assets (in $)")

#         assets = {
#             # 'Annual Income': st.text_input("Annual Income (e.g. , Your Annual Salary Income or other source of income) "),
#             'Cash/Bank Account': st.text_input("Cash/Bank Account"),
#             '401(k), 403(b), 457 Plans': st.text_input("Your 401(k), 403(b), 457 Plans "),
#             'Traditional, SEP and SIMPLE IRAs': st.text_input("Traditional, SEP and SIMPLE IRAs "),
#             'Roth IRA,Roth 401(k)': st.text_input("Roth IRA, Roth 401(k)"),
#             'Brokerage/non-qualified accounts': st.text_input("Brokerage/non-qualified accounts"),
#             'Annuities': st.text_input("Annuities"),
#             '529 Plans': st.text_input("529 Plans"),
#             'Home': st.text_input("Home"),
#             'Other Real Estate': st.text_input("Other Real Estate"),
#             'Business': st.text_input("Business"),
#             'Other': st.text_input("Other")
#         }
#         st.header("Your Liabilities (in $)")

#         liabilities = {
#             'Mortgage': st.text_input("Mortgage"),
#             # 'Annual Mortgage Interest Rate': st.number_input("Annual Mortgage Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=12.0, step=0.5),
#             # 'Mortagage Time Period': st.number_input("Mortagage Time Period (Mention the time period of the Mortgage in years)", min_value=0, max_value=100,value=10,step=1),

#             'Home Loans': st.text_input("Home Loans"),
#             # 'Home Loans Interest Rate': st.number_input("Home Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=10.0, step=0.5),
#             # 'Home Loans Time Period': st.number_input("Home Loans Time Period (Mention the time period of the Home Loan in years)", min_value=0, max_value=100,value=15,step=1),

#             'Vehicle Loans': st.text_input("Vehicle Loans"),
#             # 'Vehicle Loans Interest Rate': st.number_input("Vehicle Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
#             # 'Vehicle Loans Time Period': st.number_input("Vehicle Loans Time Period (Mention the time period of the Car/Vehicle Loan in years)", min_value=0, max_value=100,value=15,step=1),

#             'Education Loans': st.text_input("Education Loans"),
#             # 'Education Loans Interest Rate' : st.number_input("Education Loans Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
#             # 'Education Loans Time Period': st.number_input("Education Loans Time Period (Mention the time period of the Education Loan in years)", min_value=0, max_value=100,value=15,step=1),

#             # 'Credit Card': st.text_input("Monthly Credit Card Debt (Mention Amount)"),
#             # 'Credit Card Debt Interest Rate': st.number_input("Credit Card Debt Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),

#             'Miscellaneous': st.text_input("Miscellaneous"),
#         }

#         st.header("Your Retirement Goal")
#         retirement_age = st.number_input("At what age do you plan to retire?", min_value=0, max_value=120, value=65, step=1)
#         retirement_income = st.text_input("Desired annual retirement income")

#         st.header("Your Other Goals")
#         goal_name = st.text_input("Name of the Goal (e.g . , Dream House, Travel, Educational, etc.)")
#         goal_amount = st.text_input("Amount needed for the goal (in $)")
#         goal_timeframe = st.number_input("Timeframe to achieve the goal (in years)", min_value=0, max_value=100, value=5, step=1)

#         st.header("Insurance Information")
#         life_insurance_Benefit = st.text_input("Life Insurance-Benefit")
#         life_insurance_Premium = st.text_input("Life Insurance-Premium")
#         disability_insurance_Benefit = st.text_input("Disability Insurance-Benefit")
#         disability_insurance_Premium = st.text_input("Disability Insurance-Premium")
#         long_term_care_benefit = st.text_input("Long-Term Care Insurance-Benefit")
#         long_term_care_premium = st.text_input("Long-Term Care Insurance-Premium")


#         st.header("Estate Planning")
#         will_status = st.radio("Do you have a will?", ["Yes", "No"])
#         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
#         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
#         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

#         submitted = st.form_submit_button("Submit")

#         if submitted:
#             form_data = {
#                 "Client Name": client_name,
#                 "Co-Client Name": co_client_name,
#                 "Client Age": client_age,
#                 "Co-Client Age": co_client_age,
#                 "Today's Date": str(today_date),
#                 "Assets": assets,
#                 "Liabilities": liabilities,
#                 "Retirement Age": retirement_age,
#                 "Desired Retirement Income": retirement_income,
#                 "Goal Name": goal_name,
#                 "Goal Amount": goal_amount,
#                 "Goal Timeframe": goal_timeframe,
#                 "Life Insurance Benefit": life_insurance_Benefit,
#                 "Life Insurance Premium": life_insurance_Premium,
#                 "Disability Insurance Benefit": disability_insurance_Benefit,
#                 "Disability Insurance Premium": disability_insurance_Premium,
#                 "Long-Term Care Insurance Benefit": long_term_care_benefit,
#                 "Long-Term Care Insurance Premium": long_term_care_premium,
#                 "Will Status": will_status,
#                 "Trust Status": trust_status,
#                 "Power of Attorney": power_of_attorney,
#                 "Healthcare Proxy": healthcare_proxy,
#             }

#             save_data_to_file(form_data)
            
#             # # Plot the pie chart
#             # st.subheader("Assets and Liabilities Breakdown")
#             # plot_assets_liabilities_pie_chart(assets, liabilities)

#             # Store data in session state and redirect to main
#             st.session_state.assets = assets
#             st.session_state.liabilities = liabilities
#             st.session_state.total_assets, st.session_state.total_liabilities = calculate_totals(assets, liabilities)
#             st.session_state.page = "main"
#             st.success("Data submitted!\nThank You for filling the form !\nReturning to main portal...")

# import math
# def calculate_compounded_amount(principal, rate, time):
#     """
#     Calculates the compounded amount using the formula:
#     A = P * (1 + r/n)^(nt)
#     Assuming n (compounding frequency) is 1 for simplicity (annually).
#     """
#     if principal == 0 or rate == 0 or time == 0:
#         return principal
#     else:
#         # Using annual compounding
#         return principal * (1 + rate / 100) ** time
    
# def calculate_totals(assets, liabilities):
#     total_assets = sum(extract_numeric(v) for v in assets.values())
#     print(f"Total Assets : {total_assets}")
#     total_liabilities = 0
#     total_liabilities = sum(extract_numeric(v) for v in liabilities.values() )

#     # total_liabilities += calculate_compounded_amount(
#     #     extract_numeric(liabilities['Mortgage']),
#     #     liabilities['Annual Mortgage Interest Rate'],
#     #     liabilities['Mortagage Time Period']
#     # )
#     # total_liabilities += calculate_compounded_amount(
#     #     extract_numeric(liabilities['Home Loans']),
#     #     liabilities['Home Loans Interest Rate'],
#     #     liabilities['Home Loans Time Period']
#     # )
#     # total_liabilities += calculate_compounded_amount(
#     #     extract_numeric(liabilities['Vehicle Loans']),
#     #     liabilities['Vehicle Loans Interest Rate'],
#     #     liabilities['Vehicle Loans Time Period']
#     # )
#     # total_liabilities += calculate_compounded_amount(
#     #     extract_numeric(liabilities['Education Loans']),
#     #     liabilities['Education Loans Interest Rate'],
#     #     liabilities['Education Loans Time Period']
#     # )
    
#     # For credit card debt, only calculate compounded amount if interest rate > 0

#     # credit_card_balance = extract_numeric(liabilities['Credit Card'])
#     # credit_card_interest = liabilities['Credit Card Debt Interest Rate']
#     # if credit_card_interest > 0:
#     #     # Assuming the time period for credit card debt is 1 year for compounding
#     #     total_liabilities += calculate_compounded_amount(credit_card_balance, credit_card_interest, 1)
#     # else:
#     #     total_liabilities += credit_card_balance
    
#     # Miscellaneous debts are taken directly as is
#     total_liabilities += extract_numeric(liabilities['Miscellaneous'])
#     rounded_liabilities = round(total_liabilities,2)

#     print(f"Total liabilities :{total_liabilities}")
#     print(f"Rounded of Total liabilities :{rounded_liabilities}")

#     return total_assets, rounded_liabilities #total_liabilities

# def create_financial_summary_table(assets, liabilities):
#     # Filter out items with zero value
#     filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
#     filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

#     # Create DataFrames for assets and liabilities with indices starting from 1
#     assets_df = pd.DataFrame(
#         list(filtered_assets.items()), 
#         columns=['Assets', 'Amount ($)'], 
#         index=range(1, len(filtered_assets) + 1)
#     )
#     liabilities_df = pd.DataFrame(
#         list(filtered_liabilities.items()), 
#         columns=['Liabilities', 'Amount ($)'], 
#         index=range(1, len(filtered_liabilities) + 1)
#     )

#     # Calculate total
#     total_assets, total_liabilities = calculate_totals(assets, liabilities)

#     # Add total row with index incremented by 1
#     total_assets_row = pd.DataFrame(
#         [['TOTAL', total_assets]], 
#         columns=['Assets', 'Amount ($)'], 
#         index=[len(assets_df) + 1]
#     )
#     total_liabilities_row = pd.DataFrame(
#         [['TOTAL', total_liabilities]], 
#         columns=['Liabilities', 'Amount ($)'], 
#         index=[len(liabilities_df) + 1]
#     )

#     # Append total rows to DataFrames
#     assets_df = pd.concat([assets_df, total_assets_row])
#     liabilities_df = pd.concat([liabilities_df, total_liabilities_row])

#     # Display tables with formatted values
#     st.subheader("Assets")
#     st.table(assets_df.style.format({'Amount ($)': '{:,.2f}'}))

#     st.subheader("Liabilities")
#     st.table(liabilities_df.style.format({'Amount ($)': '{:,.2f}'}))


# def plot_bar_graphs(assets, liabilities):
#     # Filter out items with zero values
#     filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
#     filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

#     # Calculate compounded liabilities
#     # compounded_liabilities = {} 

#     # for k, v in filtered_liabilities.items():
#         # if 'Interest Rate' in k or 'Time Period' in k:
#         #     continue  # Skip non-monetary entries

#         # if k == 'Credit Card Payment' and liabilities['Credit Card Debt Interest Rate'] == 0.0:
#         #     continue  # Skip if credit card interest rate is zero

#         # if k == 'Mortgage':
#         #     interest_rate = liabilities['Annual Mortgage Interest Rate']
#         #     time_period = liabilities['Mortagage Time Period']

#         # elif k == 'Home Loans':
#         #     interest_rate = liabilities['Home Loans Interest Rate']
#         #     time_period = liabilities['Home Loans Time Period']

#         # elif k == 'Car/Vehicle Loans':
#         #     interest_rate = liabilities['Car/Vehicle Loans Interest Rate']
#         #     time_period = liabilities['Car/Vehicle Loans Time Period']

#         # elif k == 'Education Loans':
#         #     interest_rate = liabilities['Education Loans Interest Rate']
#         #     time_period = liabilities['Education Loans Time Period']

#         # elif k == 'Credit Card Payment':
#         #     interest_rate = liabilities['Credit Card Debt Interest Rate']
#         #     time_period = 1  # Assuming interest is calculated yearly

#         # if interest_rate > 0:
#         #     compounded_amount = float(v) * (1 + float(interest_rate) / 100) ** float(time_period)
#         #     compounded_liabilities[k] = compounded_amount
#         # else:
#         #     compounded_liabilities[k] = float(v)

#     # Plot bar graph for assets
#     st.write("### All Assets ")
#     fig1, ax1 = plt.subplots()
#     ax1.bar(filtered_assets.keys(), filtered_assets.values(), color='green')
#     ax1.set_ylabel('Amount ($)')
#     ax1.set_xlabel('Asset Type')
#     ax1.set_title(' All Assets ')
#     plt.xticks(rotation=45)
#     st.pyplot(fig1)

#     # Plot bar graph for liabilities
#     st.write("### All Liabilities ")
#     # st.write("### All Liabilities with Compounded Interest")
#     fig2, ax2 = plt.subplots()
#     # ax2.bar(compounded_liabilities.keys(), compounded_liabilities.values(), color='red')
#     ax2.bar(filtered_liabilities.keys(), filtered_liabilities.values(), color='red')    
#     ax2.set_ylabel('Amount ($)')
#     ax2.set_xlabel('Liability Type')
#     ax2.set_title(' All Liabilities ')

#     # ax2.set_title(' All Liabilities with Compounded Interest')
#     plt.xticks(rotation=45)
#     st.pyplot(fig2)


# from docx import Document
# # Define a helper function to read and extract text from a DOCX file
# def read_docx(file_path):
#     document = Document(file_path)
#     extracted_text = "\n".join([para.text for para in document.paragraphs])
#     return extracted_text



# class TrieNode:
#     def __init__(self):
#         self.children = {}
#         self.client_ids = []
#         self.end_of_name = False  # Marks the end of a client's name

# class Trie:
#     def __init__(self):
#         self.root = TrieNode()

#     def insert(self, name, client_id):
#         node = self.root
#         for char in name:
#             if char not in node.children:
#                 node.children[char] = TrieNode()
#             node = node.children[char]
#         node.client_ids.append(client_id)
#         node.end_of_name = True

#     def search(self, prefix):
#         node = self.root
#         for char in prefix:
#             if char in node.children:
#                 node = node.children[char]
#             else:
#                 return []  # Prefix not found
#         return self._get_all_names_from_node(prefix, node)

#     def _get_all_names_from_node(self, prefix, node):
#         suggestions = []
#         if node.end_of_name:
#             suggestions.append((prefix, node.client_ids))
#         for char, child_node in node.children.items():
#             suggestions.extend(self._get_all_names_from_node(prefix + char, child_node))
#         return suggestions



# def preload_trie():
#     trie = Trie()
#     clients = {
#         "John Doe": "C001",
#         "Jane Smith": "C002",
#         "James Brown": "C003",
#         "Jill Johnson": "C004",
#         "Jake White": "C005"
#     }
#     for name, client_id in clients.items():
#         trie.insert(name.lower(), client_id)  # Insert in lowercase for case-insensitive search
#     return trie

# async def generate_investment_suggestions_for_investor(investment_personality,clientName ,context,monthly_investment=10000,investment_period=3): # # GET Method for py , for front end its Post API
    
#     # retriever = asyncio.run(load_vector_db("uploaded_file"))

#     retriever =  await load_vector_db("uploaded_file")
#     # retriever = await load_vector_db("data\Financial_Investment_1.docx") 

#     chain = await make_retrieval_chain(retriever,investmentPersonality,clientName,monthly_investment,investment_period)

#     # chain = asyncio.run(make_retrieval_chain(retriever))
    
#     if chain is not None:
#         # summary = context
#         # query = summary + "\n" + investment_personality
#         query = str(investment_personality)
#         response = chain.invoke({"input": query})
        
#         # format_response = markdown_to_text(response['answer'])
#         # return format_response
        
#         # html_output = markdown.markdown(response['answer'])
#         # return html_output
        
#         # readable_text = markdown_to_readable_text(response['answer'])
#         # print(readable_text)
#         # return readable_text

#         # format_text = convert_to_markdown(response['answer'])
#         # return format_text
        
#         return response['answer']
    
        

#         # handle_graph(response['answer'])

#     else:
#         logging.INFO("response is not generated by llm model")
#         return jsonify("response is not generated by llm model"),500
#         # st.error("Failed to create the retrieval chain. Please upload a valid document.")

# from flask import Flask, request, jsonify, send_file
# import asyncio
# from flask_cors import CORS
# app = Flask(__name__)
# CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})
# # CORS(app,resources={r"/api/*":{"origins":"*"}})
# # CORS(app)

# # Initialize the Trie with preloaded clients
# trie = preload_trie()

# @app.route('/')
# def home():
#     return "Wealth Advisor Chatbot API"

# @app.route('/investment-suggestions', methods=['POST'])
# def investment_suggestions():
#     # Get the input data (new or existing client)
#     data = request.get_json()

#     # Determine if it's a new client or existing client
#     client_type = data.get("client_type")

#     if client_type == "New Client":
#         # Get form details and perform investment suggestions

#         # Check if assets and liabilities are provided
#         assets = data.get('assets', None)
#         liabilities = data.get('liabilities', None)

#         if assets and liabilities:
#             financial_summary = create_financial_summary_table(assets, liabilities)
#             bar_graphs = plot_bar_graphs(assets, liabilities)
#             pie_chart = plot_assets_liabilities_pie_chart(assets, liabilities)

#             return jsonify({
#                 "financial_summary": financial_summary,
#                 "bar_graphs": "Bar graphs generated.",
#                 "pie_chart": "Pie chart generated."
#             })

#         return jsonify({"message": "Please fill in the client details to view the assets and liabilities breakdown."})

#     elif client_type == "Existing Client":
#         # Search for an existing client in the Trie
#         search_query = data.get("search_query", "").lower()
#         matching_names = trie.search(search_query)

#         if matching_names:
#             suggestions = [{"name": name, "client_ids": client_ids} for name, client_ids in matching_names]
#             return jsonify({"suggestions": suggestions})
#         else:
#             return jsonify({"message": "No matching clients found."})
    
#     return jsonify({"message": "Invalid client type."})



# # @app.route('/upload-personal-details', methods=['POST']) # not necessary
# # def upload_personal_details():
# #     file = request.files['file']

# #     if file:
# #         # Process the uploaded file (personal details document)
# #         document_data = asyncio.run(process_document(file))
        
# #         if isinstance(document_data, tuple) and len(document_data) == 2:
# #             extracted_text, tables_content = document_data
# #             return jsonify({"message": "File processed", "extracted_text": extracted_text})
        
# #         return jsonify({"message": "Unexpected data format returned from document processing."})
    
# #     return jsonify({"message": "No file uploaded."})


# # Determine Investment personality through the investor assesmnet tab : 
# @app.route('/investor-personality-assessment', methods=['POST'])
# def investor_personality_assessment():
#     try:
#         # Collecting client name and assessment data
#         data = request.json  # Expecting JSON input
#         # client_name = data.get('client_name')
#         client_id = data.get('client_id')
#         assessment_data = data.get('assessment_data')  
        
#         # if not client_id or not assessment_data:
#         #     return jsonify({'message': 'Client name and assessment data are required.'}), 400
        
#         logging.info(f"Received assessment data for client with client id : {client_id}")

#         # Pass the assessment data to determine the investment personality
#         personality = asyncio.run(determine_investment_personality(assessment_data))
#         logging.info(f"Determined personality for {client_id}: {personality}")

#         # Return the personality and client id in response
#         return jsonify({
#             'client_id': client_id,
#             'investment_personality': personality
#         }), 200
    
#     except Exception as e:
#         logging.error(f"Error processing investor assessment: {e}")
#         return jsonify({'message': 'Internal Server Error'}), 500
    

# import logging
# global investmentPersonality  # Global Variable
# investmentPersonality = ""

# async def make_suggestions(investmentPersonality,clientName,finacial_file="data\Financial_Investment_1.docx",monthly_investment=10000,investment_period=3):
#     try:
#         try:
#             financial_file = "data\Financial_Investment_1.docx"
#             # financial_data = asyncio.run(process_document(financial_file))
#             financial_data = await process_document(financial_file)
#             # suggestions = asyncio.run(generate_investment_suggestions_for_investor(investmentPersonality,clientName, financial_data))
#             suggestions = await generate_investment_suggestions_for_investor(investmentPersonality,clientName, financial_data,monthly_investment,investment_period)
            
#             htmlSuggestions = markdown.markdown(suggestions)
            
#             formatSuggestions = markdown_to_text(suggestions)
#             data_extracted = extract_numerical_data(suggestions)
            
#             min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                             [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
#             max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                             [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

#             # Normalize allocations
#             min_allocations = normalize_allocations(min_allocations)
#             max_allocations = normalize_allocations(max_allocations)

#             # Update Bar Chart Data
#             bar_chart_data = {
#                 'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
#                 'datasets': [{
#                     'label': 'Min Allocation',
#                     'data': min_allocations,
#                     'backgroundColor': 'skyblue'
#                 },
#                 {
#                     'label': 'Max Allocation',
#                     'data': max_allocations,
#                     'backgroundColor': 'lightgreen'
#                 }]
#             }

#             # Similar changes can be made for the Pie Chart Data:
#             all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
#             num_labels = len(all_labels)
#             max_allocations_for_pie = normalize_allocations(
#                 [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
#                 [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
#             )
            
#             # Generate colors based on the number of labels
#             dynamic_colors = generate_colors(num_labels)

#             # Update Pie Chart Data
#             pie_chart_data = {
#                 'labels': all_labels,
#                 'datasets': [{
#                     'label': 'Investment Allocation',
#                     'data': max_allocations_for_pie,
#                     'backgroundColor': dynamic_colors,
#                     'hoverOffset': 4
#                 }]
#             }
            
            
#             # Prepare the data for the line chart with inflation adjustment
#             initial_investment = 10000
#             # compounded_chart_data, inflation_adjusted_chart_data = prepare_line_chart_data_with_inflation(data_extracted, initial_investment)
#             combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
#             print(f"\nThe combined chart data is : {combined_chart_data}")
            
#             return htmlSuggestions, pie_chart_data, bar_chart_data, combined_chart_data
            
#         except Exception as e:
#             logging.info(f"Error occurred while generating investment suggestions: {e}")
#             return jsonify({'message': f'Error occurred while considering preuploaded file : {e}'}), 500
        

        
#         # return jsonify({
#         #     "status": 200,
#         #     "message": "Success",
#         #     "investmentSuggestions": htmlSuggestions,
#         #     "pieChartData": pie_chart_data,
#         #     "barChartData": bar_chart_data,
#         #     "compoundedChartData":combined_chart_data
#         # }), 200
        
#     except Exception as e:
#         logging.error(f"Error processing personality assessment: {e}")
#         return jsonify({'message': 'Error in generating suggestions with personality'}), 500
        

# @app.route('/personality-assessment', methods=['POST'])
# def personality_selected():
#     try:
#         data = request.json
#         clientName = data.get('clientName')
#         try :
#             clientId = data.get('client_id')
#             investmentPersonality = data.get('investmentPersonality') # investment_personality
#             print(f"InvestmentPersonality is : {investmentPersonality}")
#             logging.info('Recieved Values')
#         except Exception as e:
#             logging.info(f"Error occurred while retrieving client id: {e}")
#             return jsonify({'message': f'Error occurred while retrieving client id: {e}'}), 400

#         try:
#             # monthly_investment= data.get('monthly_investment') #10000
#             # investment_period= data.get('investment_period')  #3
#             monthly_investment= 10000
#             investment_period= 3
#             htmlSuggestions,pie_chart_data,bar_chart_data,combined_chart_data = asyncio.run(make_suggestions(investmentPersonality,clientName,monthly_investment,investment_period))
            
#         except Exception as e:
#             logging.info(f"Error occurred while processing investment data: {e}")
#             return jsonify({'message': f'Error occurred while processing investment data: {e}'}), 400
        
#         # htmlSuggestions,pie_chart_data,bar_chart_data,combined_chart_data = asyncio.run(make_suggestions(investmentPersonality,clientName))
        
#         return jsonify({
#             "status": 200,
#             "message": "Success",
#             "investmentSuggestions": htmlSuggestions,
#             "pieChartData": pie_chart_data,
#             "barChartData": bar_chart_data,
#             "compoundedChartData":combined_chart_data
#         }), 200
                    
                    
#         # return jsonify({'message':'Sab thik'}),200
        
#         # if investmentPersonality == 'aggressiveInvestor':
#         #     pass
    
#     except Exception as e:
#         logging.info(f"Error in personality assessment: {e}")
#         print(f"Error occured in Investor Personality while collecting data :\n{e}")
#         return jsonify({'message': 'Internal Server Error in Investor Personality'}), 500



# # Route to handle generating investment suggestions
# @app.route('/generate-investment-suggestions', methods=['POST'])
# def generate_investment_suggestions():
#     try:
#         try :
#             if investmentPersonality:
#                 try:
#                     financial_file = "data\Financial_Investment_1.docx"
#                     financial_data = asyncio.run(process_document(financial_file))
#                     suggestions = asyncio.run(generate_investment_suggestions_for_investor(investmentPersonality, financial_data))
#                 except Exception as e:
#                     logging.info(f"Error occurred while generating investment suggestions: {e}")
#                     return jsonify({'message': f'Error occurred while considering preuploaded file : {e}'}), 500
#                 htmlSuggestions = markdown.markdown(suggestions)
                
#                 formatSuggestions = markdown_to_text(suggestions)
#                 data_extracted = extract_numerical_data(suggestions)
                
#                 min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                                 [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
#                 max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                                 [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

#                 # Normalize allocations
#                 min_allocations = normalize_allocations(min_allocations)
#                 max_allocations = normalize_allocations(max_allocations)

#                 # Update Bar Chart Data
#                 bar_chart_data = {
#                     'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
#                     'datasets': [{
#                         'label': 'Min Allocation',
#                         'data': min_allocations,
#                         'backgroundColor': 'skyblue'
#                     },
#                     {
#                         'label': 'Max Allocation',
#                         'data': max_allocations,
#                         'backgroundColor': 'lightgreen'
#                     }]
#                 }

#                 # Similar changes can be made for the Pie Chart Data:
#                 all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
#                 num_labels = len(all_labels)
#                 max_allocations_for_pie = normalize_allocations(
#                     [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
#                     [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
#                 )
                
#                 # Generate colors based on the number of labels
#                 dynamic_colors = generate_colors(num_labels)

#                 # Update Pie Chart Data
#                 pie_chart_data = {
#                     'labels': all_labels,
#                     'datasets': [{
#                         'label': 'Investment Allocation',
#                         'data': max_allocations_for_pie,
#                         'backgroundColor': dynamic_colors,
#                         'hoverOffset': 4
#                     }]
#                 }
                
                
#                 # Prepare the data for the line chart with inflation adjustment
#                 initial_investment = 10000
#                 # compounded_chart_data, inflation_adjusted_chart_data = prepare_line_chart_data_with_inflation(data_extracted, initial_investment)
#                 combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
#                 print(f"\nThe combined chart data is : {combined_chart_data}")
                
#                 return jsonify({
#                     "status": 200,
#                     "message": "Success",
#                     "investmentSuggestions": htmlSuggestions,
#                     "pieChartData": pie_chart_data,
#                     "barChartData": bar_chart_data,
#                     "compoundedChartData":combined_chart_data
#                 }), 200
            
#             assessment_file = request.files['assessmentFile']
#             financial_file = request.files['financialFile']
#             logging.info(" Requested files")
#         except Exception as e:
#             logging.info(" Requested files not passed")
#             return jsonify({'message': f'Error occurred while retrieving files12: {e}'}), 400
            
        
#         try :
#             try:
#                 responses = extract_responses_from_docx(assessment_file)
#             except Exception as e:
#                 logging.info(f"Failed to extract responses from assessment file: {e}")
#                 return jsonify({'message': 'Failed to extract responses from assessment file.'}), 400
#             try:
#                 financial_data = asyncio.run(process_document(financial_file))
#             except Exception as e:
#                 logging.info(f"Failed to process financial file: {e}")
#                 return jsonify({'message': 'Failed to process financial file.'}), 400

#             logging.info(f"Received Responses from the file {responses}")
#         except Exception as e:
#             logging.info("Failed to process files")
#             return jsonify({'message': f'Error occurred while processing files: {e}'}), 400

#         try:
#             # Determine investment personality
#             # monthly_investment= 10000
#             # investment_period= 3
#             # personality,monthly_investment,investment_period = asyncio.run(determine_investment_personality(responses))
            
#             personality= asyncio.run(determine_investment_personality(responses))
            
#             logging.info(f"\nPersonality of the user is : {personality}")
#         except Exception as e:
#             logging.info("Failed to determine personality")
#             return jsonify({'message': f'Error occurred while determining personality: {e}'}), 400
        
#         try:
#             # Generate investment suggestions based on personality and financial data
#             clientName = "Harshal Gidh"
#             suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality,clientName, financial_data))
#             # suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality,clientName, financial_data,monthly_investment,investment_period))
#             htmlSuggestions = markdown.markdown(suggestions)
#             logging.info(f"\Suggestions for investor : \n{suggestions}")
#         except Exception as e:
#             logging.info("Failed to generate suggestions")
#             return jsonify({'message': f'Error occurred while generating suggestions: {e}'}), 400

#         logging.info("Successfully generated")
#         formatSuggestions = markdown_to_text(suggestions)
#         data_extracted = extract_numerical_data(suggestions)
        
#         min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                         [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
#         max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                         [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

#         # Normalize allocations
#         min_allocations = normalize_allocations(min_allocations)
#         max_allocations = normalize_allocations(max_allocations)

#         # Update Bar Chart Data
#         bar_chart_data = {
#             'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
#             'datasets': [{
#                 'label': 'Min Allocation',
#                 'data': min_allocations,
#                 'backgroundColor': 'skyblue'
#             },
#             {
#                 'label': 'Max Allocation',
#                 'data': max_allocations,
#                 'backgroundColor': 'lightgreen'
#             }]
#         }

#         # Similar changes can be made for the Pie Chart Data:
#         all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
#         num_labels = len(all_labels)
#         max_allocations_for_pie = normalize_allocations(
#             [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
#             [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
#         )
        
#         # Generate colors based on the number of labels
#         dynamic_colors = generate_colors(num_labels)

#         # Update Pie Chart Data
#         pie_chart_data = {
#             'labels': all_labels,
#             'datasets': [{
#                 'label': 'Investment Allocation',
#                 'data': max_allocations_for_pie,
#                 'backgroundColor': dynamic_colors,
#                 'hoverOffset': 4
#             }]
#         }
        
        
#         # Prepare the data for the line chart with inflation adjustment
#         initial_investment = 10000
#         # compounded_chart_data, inflation_adjusted_chart_data = prepare_line_chart_data_with_inflation(data_extracted, initial_investment)
#         combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
#         print(f"\nThe combined chart data is : {combined_chart_data}")
        
#         return jsonify({
#             "status": 200,
#             "message": "Success",
#             "investmentSuggestions": htmlSuggestions,
#             "pieChartData": pie_chart_data,
#             "barChartData": bar_chart_data,
#             "compoundedChartData":combined_chart_data
#         }), 200

#     except Exception as e:
#         logging.info(f"Error in generating investment suggestions: {e}")
#         return jsonify({'message': 'Internal Server Error in Generating responses'}), 500




# # Run the Flask application
# if __name__ == '__main__':
#     app.run(host='0.0.0.0',debug=True)





# # # Flask Code of APP.py

# # import streamlit as st
# # import pandas as pd
# # import matplotlib.pyplot as plt

# # import os
# # import filetype
# # import docx
# # import PyPDF2
# # import re
# # from dotenv import load_dotenv
# # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # from langchain_community.vectorstores import Chroma
# # from langchain_community.docstore.in_memory import InMemoryDocstore
# # from langchain_community.vectorstores import FAISS
# # from langchain_community.document_loaders import Docx2txtLoader
# # from langchain_core.prompts import ChatPromptTemplate
# # from langchain.chains import create_retrieval_chain
# # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # from langchain.chains.combine_documents import create_stuff_documents_chain
# # from langchain.memory import ConversationSummaryMemory
# # import asyncio
# # import numpy as np
# # import json

# # import google.generativeai as genai
# # import pathlib
# # import logging
# # import sys
# # import io
# # import matplotlib.pyplot as plt
# # import seaborn as sns
# # # Import things that are needed generically
# # from langchain.pydantic_v1 import BaseModel, Field
# # from langchain.tools import BaseTool, StructuredTool, tool
# # # Define functions to generate investment suggestions :

# # load_dotenv()
# # GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

# # from flask import Flask, request, jsonify

# # app = Flask(__name__)

# # # Configure generativeai with your API key
# # genai.configure(api_key=GOOGLE_API_KEY)

# # import markdown
# # # def convert_to_markdown(raw_text):
# # #     # Replace specific text patterns with markdown syntax
# # #     formatted_text = raw_text.replace('\n', '\n\n')  # Ensure newlines create paragraphs
    
# # #     # Convert text into markdown format
# # #     html = markdown.markdown(formatted_text)

# # #     return html


# # import markdown2
# # from bs4 import BeautifulSoup

# # def markdown_to_readable_text(md_text):
# #     # Convert markdown to HTML
# #     html = markdown2.markdown(md_text)

# #     # Parse the HTML
# #     soup = BeautifulSoup(html, "html.parser")

# #     # Function to format plain text from tags
# #     def format_text_from_html(soup):
# #         formatted_text = ''
# #         for element in soup:
# #             if element.name == "h1":
# #                 formatted_text += f"\n\n# {element.text.upper()} #\n\n"
# #             elif element.name == "h2":
# #                 formatted_text += f"\n\n## {element.text} ##\n\n"
# #             elif element.name == "h3":
# #                 formatted_text += f"\n\n### {element.text} ###\n\n"
# #             elif element.name == "strong":
# #                 formatted_text += f"**{element.text}**"
# #             elif element.name == "em":
# #                 formatted_text += f"_{element.text}_"
# #             elif element.name == "ul":
# #                 for li in element.find_all("li"):
# #                     formatted_text += f"\n - {li.text}"
# #             elif element.name == "ol":
# #                 for idx, li in enumerate(element.find_all("li"), 1):
# #                     formatted_text += f"\n {idx}. {li.text}"
# #             elif element.name == "table":
# #                 rows = element.find_all("tr")
# #                 for row in rows:
# #                     cols = row.find_all(["th", "td"])
# #                     row_text = ' | '.join(col.text.strip() for col in cols)
# #                     formatted_text += f"{row_text}\n"
# #                 formatted_text += "\n"
# #             else:
# #                 formatted_text += element.text

# #         return formatted_text.strip()

# #     return format_text_from_html(soup)

# # def markdown_to_text(md): # og solution code 
# #     # Simple conversion for markdown to plain text
# #     md = md.replace('**', '')
# #     md = md.replace('*', '')
# #     md = md.replace('_', '')
# #     md = md.replace('#', '')
# #     md = md.replace('`', '')
# #     return md.strip()


# # # import docx

# # # def extract_responses_from_docx(personality_file):
# # #     """
# # #     Extracts responses from a Word document (.docx) where answers are typed in.

# # #     Args:
# # #         personality_file (UploadedFile): The file object uploaded via Streamlit.

# # #     Returns:
# # #         dict: A dictionary containing the questions and the typed answers.
# # #     """
# # #     try:
# # #         doc = docx.Document(personality_file)
# # #         responses = {}
# # #         current_question = None

# # #         # Check paragraphs
# # #         for para in doc.paragraphs:
# # #             text = para.text.strip()
# # #             if text:
# # #                 # Check if the paragraph contains a question
# # #                 if "?" in text or text.endswith(":"):
# # #                     current_question = text
# # #                     st.write(f"Identified question: {current_question}")  # Debugging log
# # #                 else:
# # #                     # This is a typed answer
# # #                     typed_answer = text.strip()
# # #                     st.write(f"Identified typed answer: {typed_answer}")  # Debugging log
# # #                     if current_question:
# # #                         # If the question already has an answer, append to it (handles multiple responses)
# # #                         if current_question in responses:
# # #                             responses[current_question] += "; " + typed_answer
# # #                         else:
# # #                             responses[current_question] = typed_answer

# # #             # Debugging log to understand document structure
# # #             st.write(f"Processing paragraph: {text}")  # Console log for local testing

# # #         # Check tables for additional responses
# # #         for table in doc.tables:
# # #             for row in table.rows:
# # #                 for cell in row.cells:
# # #                     text = cell.text.strip()
# # #                     if text:
# # #                         if "?" in text or text.endswith(":"):
# # #                             current_question = text
# # #                             st.write(f"Identified question in table: {current_question}")  # Debugging log
# # #                         else:
# # #                             typed_answer = text.strip()
# # #                             st.write(f"Identified typed answer in table: {typed_answer}")  # Debugging log
# # #                             if current_question:
# # #                                 if current_question in responses:
# # #                                     responses[current_question] += "; " + typed_answer
# # #                                 else:
# # #                                     responses[current_question] = typed_answer

# # #         if responses:
# # #             st.write("Extracted Responses:")
# # #             for question, answer in responses.items():
# # #                 st.write(f"**{question}**: {answer}")
# # #         else:
# # #             st.write("No responses captured. Please check the document formatting or symbols used.")

# # #         return responses

# # #     except Exception as e:
# # #         st.write(f"Error extracting responses: {e}")  # Console log for local testing
# # #         return None

# # # def determine_investment_personality(responses):
# # #     """
# # #     Determines the investment personality based on extracted responses.

# # #     Args:
# # #         responses (dict): A dictionary containing the questions and the selected answers.

# # #     Returns:
# # #         str: The determined investment personality.
# # #     """
# # #     try:
# # #         # Prepare input text for the chatbot based on extracted responses
# # #         input_text = "User Profile:\n"
# # #         for question, response in responses.items():
# # #             input_text += f"{question}: {response}\n"

# # #         # Introduce the chatbot's task and prompt for classification
# # #         input_text += "\nYour task is to determine the investment personality based on the above profile."

# # #         # Here you would send the input_text to your chatbot or classification model
# # #         # For demonstration, we'll just return the input_text
# # #         return input_text

# # #     except Exception as e:
# # #         st.write(f"Error determining investment personality: {e}")  # Console log for local testing
# # #         return None

# # # def extract_responses_from_docx(personality_file):
# # #     try:
# # #         doc = docx.Document(personality_file)
# # #         responses = {}
# # #         current_question = None

# # #         # Check paragraphs
# # #         for para in doc.paragraphs:
# # #             text = para.text.strip()
# # #             if text:
# # #                 # Check if the paragraph contains a question
# # #                 if "?" in text or text.endswith(":"):
# # #                     current_question = text
# # #                 else:
# # #                     # This is a typed answer
# # #                     typed_answer = text.strip()
# # #                     if current_question:
# # #                         # If the question already has an answer, append to it (handles multiple responses)
# # #                         if current_question in responses:
# # #                             responses[current_question] += "; " + typed_answer
# # #                         else:
# # #                             responses[current_question] = typed_answer

# # #         # Check tables for additional responses
# # #         for table in doc.tables:
# # #             for row in table.rows:
# # #                 for cell in row.cells:
# # #                     text = cell.text.strip()
# # #                     if text:
# # #                         if "?" in text or text.endswith(":"):
# # #                             current_question = text
# # #                         else:
# # #                             typed_answer = text.strip()
# # #                             if current_question:
# # #                                 if current_question in responses:
# # #                                     responses[current_question] += "; " + typed_answer
# # #                                 else:
# # #                                     responses[current_question] = typed_answer

# # #         return responses

# # #     except Exception as e:
# # #         print(f"Error extracting responses: {e}")
# # #         return None

# # import docx

# # # # GET Method for me POST method for Frontend
# # def extract_responses_from_docx(personality_file): # Using text responses parsing
# #     """
# #     Extracts responses from a Word document (.docx) where the selected answers are listed as text after the options.

# #     Args:
# #         personality_file (str): Path to the Word document file.

# #     Returns:
# #         dict: A dictionary containing the questions and the selected answers.
# #     """
# #     try:
# #         doc = docx.Document(personality_file)
# #         responses = {}
# #         current_question = None

# #         for para in doc.paragraphs:
# #             text = para.text.strip()
# #             if text:
# #                 # Detect the beginning of a question
# #                 if "?" in text:
# #                     current_question = text
# #                 # Detect a chosen response (assuming it follows the question and options)
# #                 elif current_question and not text.startswith(("a.", "b.", "c.", "d.")):
# #                     selected_answer = text
# #                     responses[current_question] = selected_answer
# #                     current_question = None  # Reset for the next question

# #         if responses:
# #             print(responses)
# #             # st.write(responses)
# #         else:
# #             print("\nNo responses captured")
# #             st.write("No responses captured")
# #         return responses
# #     except Exception as e:
# #         print(f"Error extracting responses: {e}")
# #         return None

# # # def extract_responses_from_assessment(personality_file): # using boxes
# # #     # Load the document
# # #     # doc = Document(docx_filename)
# # #     doc = docx.Document(personality_file)
    
# # #     # Initialize a list to store responses
# # #     responses = []
    
# # #     # Iterate through each paragraph in the document
# # #     for para in doc.paragraphs:
# # #         text = para.text.strip()
# # #         # Check if the paragraph contains a checkbox
# # #         if '☒' in text or '☐' in text:
# # #             # Extract the response marked with ☒
# # #             if '☒' in text:
# # #                 response = text.split('☒')[1].strip()
# # #                 responses.append(response)
    
# # #     return responses

# # # import asyncio
# # # # from some_generative_ai_library import GenerativeModel  # Replace with actual import

# # # async def determine_investment_personality(assessment_data):
# # #     try:
# # #         # Prepare input text for the chatbot based on assessment data
# # #         input_text = "User Profile:\n"
# # #         for question, answer in assessment_data.items():
# # #             input_text += f"{question}: {answer}\n"

# # #         # Introduce the chatbot's task and prompt for classification
# # #         input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
# # #                       "- Conservative Investor\n" \
# # #                       "- Moderate Investor\n" \
# # #                       "- Aggressive Investor\n\n" \
# # #                       "Please provide the classification below:\n"

# # #         # Use your generative AI model to generate a response
# # #         model = GenerativeModel('gemini-1.5-flash')
# # #         response = await model.generate_content(input_text)

# # #         # Determine the investment personality from the chatbot's response
# # #         response_text = response.text.lower()

# # #         if "conservative investor" in response_text:
# # #             personality = "Conservative Investor"
# # #         elif "moderate investor" in response_text:
# # #             personality = "Moderate Investor"
# # #         elif "aggressive investor" in response_text:
# # #             personality = "Aggressive Investor"
# # #         else:
# # #             personality = "Unknown"

# # #         return personality
# # #     except Exception as e:
# # #         print(f"Error generating response: {e}")
# # #         return "Unknown"


# # # GET Method
# # async def determine_investment_personality(assessment_data): # proper code 
# #     try:
# #         # Prepare input text for the chatbot based on assessment data
# #         input_text = "User Profile:\n"
# #         for question, answer in assessment_data.items():
# #             input_text += f"{question}: {answer}\n"

# #         # Introduce the chatbot's task and prompt for classification
# #         input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
# #                       "- Conservative Investor\n" \
# #                       "- Moderate Investor\n" \
# #                       "- Aggressive Investor\n\n" \
# #                       "Please provide the classification below:\n"

# #         # Use your generative AI model to generate a response
# #         model = genai.GenerativeModel('gemini-1.5-flash')
# #         response = model.generate_content(input_text)

# #         # Determine the investment personality from the chatbot's response
# #         response_text = response.text.lower()

# #         if "conservative investor" in response_text:
# #             personality = "Conservative Investor"
# #         elif "moderate investor" in response_text:
# #             personality = "Moderate Investor"
# #         elif "aggressive investor" in response_text:
# #             personality = "Aggressive Investor"
# #         else:
# #             personality = "Unknown"

# #         return personality
# #     except Exception as e:
# #         print(f"Error generating response: {e}")
# #         return "Unknown"



# # # #Load the Vector DataBase :
# # async def load_vector_db(file_path): # # GET Method 
# #     try:
# #         print("Loading vector database...")
# #         loader = Docx2txtLoader(file_path)
# #         documents = loader.load()
# #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# #         text_chunks = text_splitter.split_documents(documents)
# #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# #         # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
# #         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
# #         # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

# #         # vector_store = FAISS(
# #         #     embedding_function=embeddings,
# #         #     index=index,
# #         #     docstore=InMemoryDocstore(),
# #         #     index_to_docstore_id={},
# #         # )
        
# #         print("Vector database loaded successfully.") 
# #         return vector_store.as_retriever(search_kwargs={"k": 1})
# #     except Exception as e:
# #         print(f"Error loading vector database: {e}")
# #         return None


# # # async def load_vector_db(file_path="client_data.txt"):
# # #     try:
# # #         print("Loading vector database...")
# # #         with open(file_path, "r") as file:
# # #             text = file.read()
        
# # #         loader = Docx2txtLoader(file_path)
# # #         documents = loader.load()
# # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # #         text_chunks = text_splitter.split_documents(documents) #([Document(text=text)])
# # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        
# # #         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        
# # #         print("Vector database loaded successfully.") 
# # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # #     except Exception as e:
# # #         print(f"Error loading vector database: {e}")
# # #         return None


# # investment_personality = "Moderate Investor"
# # async def make_retrieval_chain(retriever): # GET Method
# #     """
# #     Create a retrieval chain using the provided retriever.

# #     Args:
# #         retriever (RetrievalQA): A retriever object.

# #     Returns:
# #         RetrievalQA: A retrieval chain object.
# #     """
# #     try:
# #         global investment_personality #,summary
# #         llm = ChatGoogleGenerativeAI(
# #             #model="gemini-pro",
# #             model = "gemini-1.5-flash",
# #             temperature=0.7,
# #             top_p=0.85,
# #             google_api_key=GOOGLE_API_KEY
# #         )
# #                                                     # \n + summary 
# #         prompt_template = investment_personality +  "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# #                 Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality.
# #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# #                 should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
# #                 Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# #                 investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
# #                 Also give the user minimum and maximum expected growth in dollars for the time horizon .
# #                 Also explain the user why you are giving them that particular investment suggestion.
# #                 Here's an example for the required Output Format :

# #                 Investment Suggestions for a Moderate Investor(This is for a Moderate Investor but you need to generate for any investor)

# #                 Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

# #                 Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

# #                 Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
# #                 How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
# #                 Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
# #                 ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
# #                 Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
# #                 Consider investing in blue-chip companies or growth sectors like technology. 
# #                 Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


# #                 Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
# #                 Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

# #                 Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
# #                 How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
# #                 Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
# #                 Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
# #                 Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
# #                 Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
# #                 Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
# #                 Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

# #                 Time Horizon and Expected Returns:

# #                 Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
# #                 Minimum Expected Annual Return: 4% - 6% 
                
                
# #                 \nMaximum Expected Annual Return: 8% - 10% 
                
# #                 Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, a 10,000 investment could grow to approximately 17,908 in 10 years.
# #                 Minimum Expected Growth in Dollars: 
                
# #                 4,000−6,000 (over 10 years) 
                
                
# #                 Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
# #                 Rationale for Investment Suggestions:

# #                 This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
# #                 Important Considerations:

# #                 Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

# #                 Inflation Adjusted Returns:(do not write this part inside the bracket just give answer,assume US inflation rate, and give the investment returns value that was suggested by you  for $10k investment after 3,5,10years of growth  mention the values before adjusting and after adjusting with inflation )

# #                 Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

# #                 <context>
# #                 {context}
# #                 </context>
# #                 Question: {input}"""

# #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
# #         combine_docs_chain = None  

# #         if retriever is not None :  
# #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# #             print(retriever_chain)
# #             return retriever_chain
# #         else:
# #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# #             return None

# #     except Exception as e:
# #         print(f"Error in creating chain: {e}")
# #         return None

# # import streamlit as st
# # import json
# # import matplotlib.pyplot as plt
# # import io


# # async def process_document(file_path): # GET Method
# #     try:
# #         print("Processing the document")
# #         file_type = filetype.guess(file_path)
# #         if file_type is not None:
# #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# #                 # Await the coroutine to extract text and tables
# #                 return await extract_text_and_tables_from_word(file_path)
# #             elif file_type.mime == "application/pdf":
# #                 return await extract_text_from_pdf(file_path)
# #         return None
# #     except Exception as e:
# #         print(f"Error processing document: {e}")
# #         return None

# # # Async function to extract text from a PDF file
# # async def extract_text_from_pdf(pdf_file_path): # GET Method
# #     try:
# #         print("Processing pdf file")
# #         with open(pdf_file_path, "rb") as pdf_file:
# #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# #             text_content = []
# #             for page_num in range(pdf_reader.numPages):
# #                 page = pdf_reader.getPage(page_num)
# #                 text_content.append(page.extract_text())
# #             return "\n".join(text_content)
# #     except Exception as e:
# #         print(f"Error extracting text from PDF: {e}")
# #         return None

# # # Async function to extract text and tables from a Word document
# # async def extract_text_and_tables_from_word(docx_file_path): # GET Method
# #     try:
# #         print("Extracting text and tables from word file")
# #         doc = docx.Document(docx_file_path)
# #         text_content = []
# #         tables_content = []

# #         for para in doc.paragraphs:
# #             text_content.append(para.text)

# #         for table in doc.tables:
# #             table_data = []
# #             for row in table.rows:
# #                 row_data = []
# #                 for cell in row.cells:
# #                     row_data.append(cell.text.strip())
# #                 table_data.append(row_data)
# #             tables_content.append(table_data)
# #         print("Extracted text from word file")
# #         return "\n".join(text_content), tables_content
# #     except Exception as e:
# #         print(f"Error extracting text and tables from Word document: {e}")
# #         return None, None



# # async def validate_document_content(text, tables):
# #     """
# #     Validates the content of the document.

# #     Args:
# #         text (str): Extracted text content from the document.
# #         tables (list): Extracted tables content from the document.

# #     Returns:
# #         tuple: Client name and validation errors.
# #     """
# #     errors = []
    
# #     # Extract client name
# #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# #     # Define required sections
# #     required_sections = [
# #         "YOUR RETIREMENT GOAL",
# #         "YOUR OTHER MAJOR GOALS",
# #         "YOUR ASSETS AND LIABILITIES",
# #         "MY LIABILITIES",
# #         "YOUR CURRENT ANNUAL INCOME"
# #     ]

# #     # Check for the presence of required sections
# #     for section in required_sections:
# #         if section not in text:
# #             errors.append(f"* {section} section missing.")
    
# #     # Define table field checks
# #     table_checks = {
# #         "YOUR RETIREMENT GOAL": [
# #             r"When do you plan to retire\? \(age or date\)",
# #             r"Social Security Benefit \(include expected start date\)",
# #             r"Pension Benefit \(include expected start date\)",
# #             r"Other Expected Income \(rental, part-time work, etc.\)",
# #             r"Estimated Annual Retirement Expense"
# #         ],
# #         "YOUR OTHER MAJOR GOALS": [
# #             r"GOAL", r"COST", r"WHEN"
# #         ],
# #         "YOUR ASSETS AND LIABILITIES": [
# #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# #             r"Current Value", r"Annual Contributions"
# #         ],
# #         "MY LIABILITIES": [
# #             r"Balance", r"Interest Rate", r"Monthly Payment"
# #         ]
# #     }

# #     # Validate table content
# #     for section, checks in table_checks.items():
# #         section_found = False
# #         for table in tables:
# #             table_text = "\n".join(["\t".join(row) for row in table])
# #             if section in table_text:
# #                 section_found = True
# #                 for check in checks:
# #                     if not re.search(check, table_text, re.IGNORECASE):
# #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# #                 break
# #         if not section_found:
# #             errors.append(f"* {section} section missing.")

# #     return client_name, errors


# # # RUN Button :
# # async def generate_investment_suggestions(investment_personality, context): # # GET Method for py , for front end its Post API
    
# #     # retriever = asyncio.run(load_vector_db("uploaded_file"))

# #     retriever = await load_vector_db("uploaded_file")
# #     # retriever = await load_vector_db("data\Financial_Investment_1.docx") 

# #     chain = await make_retrieval_chain(retriever)

# #     # chain = asyncio.run(make_retrieval_chain(retriever))
    
# #     if chain is not None:
# #         # summary = context
# #         # query = summary + "\n" + investment_personality
# #         query = str(investment_personality)
# #         response = chain.invoke({"input": query})
# #         format_response = markdown_to_text(response['answer'])
# #         return format_response
# #         # st.write(format_response)

# #         # handle_graph(response['answer'])

# #     else:
# #         st.error("Failed to create the retrieval chain. Please upload a valid document.")



# # # Generate Infographics : Best Code so far:


# # import seaborn as sns
# # import re
# # from collections import defaultdict
# # import matplotlib.pyplot as plt
# # import streamlit as st
# # import numpy as np

# # def extract_numerical_data(response):
# #     # Define patterns to match different sections and their respective allocations
# #     patterns = {
# #         'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
# #         'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
# #         'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
# #     }

# #     data = defaultdict(dict)

# #     for section, pattern in patterns.items():
# #         match = pattern.search(response)
# #         if match:
# #             investments_text = match.group(1)
# #             # Extract individual investment types and their allocations
# #             investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
# #             for investment_match in investment_pattern.findall(investments_text):
# #                 investment_type, min_allocation, max_allocation = investment_match
# #                 data[section][investment_type.strip()] = {
# #                     'min': min_allocation,
# #                     'max': max_allocation
# #                 }

# #     # Extract time horizon and expected returns
# #     time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
# #     min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
# #     max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
# #     min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
# #     max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

# #     time_horizon_match = time_horizon_pattern.search(response)
# #     min_return_match = min_return_pattern.search(response)
# #     max_return_match = max_return_pattern.search(response)
# #     min_growth_match = min_growth_pattern.search(response)
# #     max_growth_match = max_growth_pattern.search(response)

# #     if time_horizon_match:
# #         data['Time Horizon'] = {
# #             'min_years': time_horizon_match.group(1),
# #             'max_years': time_horizon_match.group(2)
# #         }

# #     if min_return_match:
# #         data['Expected Annual Return'] = {
# #             'min': min_return_match.group(1),
# #             'max': min_return_match.group(2)
# #         }

# #     if max_return_match:
# #         data['Expected Annual Return'] = {
# #             'min': max_return_match.group(1),
# #             'max': max_return_match.group(2)
# #         }

# #     if min_growth_match:
# #         data['Expected Growth in Dollars'] = {
# #             'min': min_growth_match.group(1),
# #             'max': min_growth_match.group(2)
# #         }

# #     if max_growth_match:
# #         data['Expected Growth in Dollars'] = {
# #             'min': max_growth_match.group(1),
# #             'max': max_growth_match.group(2)
# #         }

# #     return data

# # def normalize_allocations(allocations):
# #     total = sum(allocations)
# #     if total == 100:
# #         return allocations
# #     return [round((allocation / total) * 100, 2) for allocation in allocations]


# # # def prepare_line_chart_data_with_inflation(data_extracted, initial_investment=10000, inflation_rate=4):
# # #     min_return = data_extracted['Expected Annual Return']['min']
# # #     max_return = data_extracted['Expected Annual Return']['max']
# # #     min_years = data_extracted['Time Horizon']['min_years']
# # #     max_years = data_extracted['Time Horizon']['max_years']

# # #     def calculate_compounded_return(principal, rate, years):
# # #         return principal * (1 + rate / 100) ** years

# # #     def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
# # #         return nominal_return / (1 + inflation_rate / 100) ** years

# # #     labels = list(range(1, max_years + 1))  # Years for the x-axis
# # #     min_compounded = []
# # #     max_compounded = []
# # #     min_inflation_adjusted = []
# # #     max_inflation_adjusted = []

# # #     for year in labels:
# # #         # Calculate compounded returns (nominal)
# # #         min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
# # #         max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

# # #         # Calculate inflation-adjusted returns
# # #         min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
# # #         max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

# # #         # Append results
# # #         min_compounded.append(min_compounded_value)
# # #         max_compounded.append(max_compounded_value)
# # #         min_inflation_adjusted.append(min_inflation_value)
# # #         max_inflation_adjusted.append(max_inflation_value)

# # #     # Line Chart Data for Compounded Returns
# # #     compounded_return_chart_data = {
# # #         'labels': labels,
# # #         'datasets': [
# # #             {
# # #                 'label': 'Minimum Compounded Return',
# # #                 'data': min_compounded,
# # #                 'borderColor': 'rgb(255, 99, 132)',  # Red color
# # #                 'fill': False
# # #             },
# # #             {
# # #                 'label': 'Maximum Compounded Return',
# # #                 'data': max_compounded,
# # #                 'borderColor': 'rgb(54, 162, 235)',  # Blue color
# # #                 'fill': False
# # #             }
# # #         ]
# # #     }

# # #     # Line Chart Data for Inflation-Adjusted Returns
# # #     inflation_adjusted_chart_data = {
# # #         'labels': labels,
# # #         'datasets': [
# # #             {
# # #                 'label': 'Min Inflation Adjusted Return',
# # #                 'data': min_inflation_adjusted,
# # #                 'borderColor': 'rgb(75, 192, 192)',  # Light blue
# # #                 'fill': False
# # #             },
# # #             {
# # #                 'label': 'Max Inflation Adjusted Return',
# # #                 'data': max_inflation_adjusted,
# # #                 'borderColor': 'rgb(153, 102, 255)',  # Light purple
# # #                 'fill': False
# # #             }
# # #         ]
# # #     }

# # #     return compounded_return_chart_data, inflation_adjusted_chart_data


# # def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
# #     try:
# #         # Print data_extracted to debug the structure
# #         print("Data extracted:", data_extracted)

# #         # Check if 'Expected Annual Return' and 'Time Horizon' exist and have the expected keys
# #         if 'Expected Annual Return' not in data_extracted :
# #             print("'Expected Annual Return' missing in data_extracted")
# #             data_extracted['Expected Annual Return']['min'] = 6
# #             data_extracted['Expected Annual Return']['max'] = 8
# #             min_return = 6
# #             max_return = 8
# #         else :
# #             min_return = float(data_extracted['Expected Annual Return'].get('min', '0').strip('%'))
# #             max_return = float(data_extracted['Expected Annual Return'].get('max', '0').strip('%'))
                
# #         min_years = int(data_extracted['Time Horizon'].get('min_years', 1))  # Default to 1 year if missing
# #         max_years = int(data_extracted['Time Horizon'].get('max_years', 10))  # Default to 10 years if missing

# #         def calculate_compounded_return(principal, rate, years):
# #             return principal * (1 + rate / 100) ** years

# #         def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
# #             return nominal_return / (1 + inflation_rate / 100) ** years

# #         labels = list(range(1, max_years + 1))  # Years for the x-axis
# #         min_compounded = []
# #         max_compounded = []
# #         min_inflation_adjusted = []
# #         max_inflation_adjusted = []

# #         for year in labels:
# #             # Calculate nominal compounded returns
# #             min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
# #             max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

# #             # Calculate inflation-adjusted compounded returns
# #             min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
# #             max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

# #             # Append results
# #             min_compounded.append(min_compounded_value)
# #             max_compounded.append(max_compounded_value)
# #             min_inflation_adjusted.append(min_inflation_value)
# #             max_inflation_adjusted.append(max_inflation_value)

# #         # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
# #         combined_chart_data = {
# #             'labels': labels,
# #             'datasets': [
# #                 {
# #                     'label': 'Minimum Compounded Return',
# #                     'data': min_compounded,
# #                     'borderColor': 'rgb(255, 99, 132)',  # Red color
# #                     'fill': False
# #                 },
# #                 {
# #                     'label': 'Maximum Compounded Return',
# #                     'data': max_compounded,
# #                     'borderColor': 'rgb(54, 162, 235)',  # Blue color
# #                     'fill': False
# #                 },
# #                 {
# #                     'label': 'Min Inflation Adjusted Return',
# #                     'data': min_inflation_adjusted,
# #                     'borderColor': 'rgb(75, 192, 192)',  # Light blue
# #                     'borderDash': [5, 5],  # Dashed line for distinction
# #                     'fill': False
# #                 },
# #                 {
# #                     'label': 'Max Inflation Adjusted Return',
# #                     'data': max_inflation_adjusted,
# #                     'borderColor': 'rgb(153, 102, 255)',  # Light purple
# #                     'borderDash': [5, 5],  # Dashed line for distinction
# #                     'fill': False
# #                 }
# #             ]
# #         }
# #     except KeyError as e:
# #         print(f"KeyError occurred: {e}")
# #         return jsonify({'message': f'Key Error: {e}'}), 400
# #     except Exception as e:
# #         print(f"Error occurred while preparing data for combined line chart: {e}")
# #         return jsonify({'message': 'Internal Server Error in creating line chart'}), 500

# #     return combined_chart_data


# # # def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
# # #     try:
# # #         min_return = float(data_extracted['Expected Annual Return']['min'].strip('%'))
# # #         max_return = float(data_extracted['Expected Annual Return']['max'].strip('%'))
# # #         min_years = int(data_extracted['Time Horizon']['min_years'])
# # #         max_years = int(data_extracted['Time Horizon']['max_years'])

# # #         def calculate_compounded_return(principal, rate, years):
# # #             return principal * (1 + rate / 100) ** years

# # #         def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
# # #             return nominal_return / (1 + inflation_rate / 100) ** years

# # #         labels = list(range(1, max_years + 1))  # Years for the x-axis
# # #         min_compounded = []
# # #         max_compounded = []
# # #         min_inflation_adjusted = []
# # #         max_inflation_adjusted = []

# # #         for year in labels:
# # #             # Calculate nominal compounded returns
# # #             min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
# # #             max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

# # #             # Calculate inflation-adjusted compounded returns
# # #             min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
# # #             max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

# # #             # Append results
# # #             min_compounded.append(min_compounded_value)
# # #             max_compounded.append(max_compounded_value)
# # #             min_inflation_adjusted.append(min_inflation_value)
# # #             max_inflation_adjusted.append(max_inflation_value)

# # #         # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
# # #         combined_chart_data = {
# # #             'labels': labels,
# # #             'datasets': [
# # #                 {
# # #                     'label': 'Minimum Compounded Return',
# # #                     'data': min_compounded,
# # #                     'borderColor': 'rgb(255, 99, 132)',  # Red color
# # #                     'fill': False
# # #                 },
# # #                 {
# # #                     'label': 'Maximum Compounded Return',
# # #                     'data': max_compounded,
# # #                     'borderColor': 'rgb(54, 162, 235)',  # Blue color
# # #                     'fill': False
# # #                 },
# # #                 {
# # #                     'label': 'Min Inflation Adjusted Return',
# # #                     'data': min_inflation_adjusted,
# # #                     'borderColor': 'rgb(75, 192, 192)',  # Light blue
# # #                     'borderDash': [5, 5],  # Dashed line for distinction
# # #                     'fill': False
# # #                 },
# # #                 {
# # #                     'label': 'Max Inflation Adjusted Return',
# # #                     'data': max_inflation_adjusted,
# # #                     'borderColor': 'rgb(153, 102, 255)',  # Light purple
# # #                     'borderDash': [5, 5],  # Dashed line for distinction
# # #                     'fill': False
# # #                 }
# # #             ]
# # #         }
# # #     except Exception as e:
# # #         print(f"Error occurred while preparing data for combined line chart: {e}")
# # #         return jsonify({'message': 'Internal Server Error in creating line chart'}), 500
# # #     return combined_chart_data



# # def plot_investment_allocations(data):
# #     # Create subplots with a large figure size
# #     fig, axes = plt.subplots(2, 1, figsize= (16,10)) #(28, 15))  # Adjust size as needed

# #     # Plot Growth-Oriented Investments
# #     growth_data = data['Growth-Oriented Investments']
# #     growth_labels = list(growth_data.keys())
# #     growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
# #     growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

# #     axes[0].bar(growth_labels, growth_min, color='skyblue', label='Min Allocation')
# #     axes[0].bar(growth_labels, growth_max, bottom=growth_min, color='lightgreen', label='Max Allocation')
# #     axes[0].set_title('Growth-Oriented Investments', fontsize=16)
# #     axes[0].set_ylabel('Percentage Allocation', fontsize=14)
# #     axes[0].set_xlabel('Investment Types', fontsize=14)
# #     axes[0].tick_params(axis='x', rotation=45, labelsize=12)
# #     axes[0].tick_params(axis='y', labelsize=12)
# #     axes[0].legend()

# #     # Plot Conservative Investments
# #     conservative_data = data['Conservative Investments']
# #     conservative_labels = list(conservative_data.keys())
# #     conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
# #     conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

# #     axes[1].bar(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
# #     axes[1].bar(conservative_labels, conservative_max, bottom=conservative_min, color='lightgreen', label='Max Allocation')
# #     axes[1].set_title('Conservative Investments', fontsize=16)
# #     axes[1].set_ylabel('Percentage Allocation', fontsize=14)
# #     axes[1].set_xlabel('Investment Types', fontsize=14)
# #     axes[1].tick_params(axis='x', rotation=45, labelsize=12)
# #     axes[1].tick_params(axis='y', labelsize=12)
# #     axes[1].legend()

# #     # Tight layout for better spacing
# #     plt.tight_layout()
# #     plt.show()
# #     return fig


# # def plot_pie_chart(data):
# #     fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

# #     # Combine all investment data for pie chart
# #     all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
# #     labels = list(all_data.keys())
# #     sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
# #     colors = plt.cm.Paired(range(len(labels)))

# #     wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
# #     ax.set_title('Investment Allocation')

# #     # Add legend
# #     ax.legend(wedges, labels, title="Investment Types", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

# #     return fig



# # def bar_chart(data):
# #     fig, ax = plt.subplots(figsize=(12, 8))  # Increased size

# #     # Data for plotting
# #     categories = list(data.keys())
# #     values_min = [int(data[cat]['min'].strip('%')) for cat in categories]
# #     values_max = [int(data[cat]['max'].strip('%')) for cat in categories]

# #     x = range(len(categories))

# #     ax.bar(x, values_min, width=0.4, label='Min Allocation', color='skyblue', align='center')
# #     ax.bar(x, values_max, width=0.4, label='Max Allocation', color='lightgreen', align='edge')

# #     ax.set_xticks(x)
# #     ax.set_xticklabels(categories, rotation=45, ha='right')
# #     ax.set_xlabel('Investment Categories')
# #     ax.set_ylabel('Percentage Allocation')
# #     ax.set_title('Investment Allocation')
# #     ax.legend()

# #     plt.tight_layout()
# #     return fig


# # import random

# # def generate_colors(n):
# #     """
# #     Generate 'n' random RGB colors.

# #     Args:
# #         n (int): Number of colors to generate.
    
# #     Returns:
# #         list: A list of RGB colors in 'rgb(r, g, b)' format.
# #     """
# #     colors = []
# #     for _ in range(n):
# #         r = random.randint(0, 255)
# #         g = random.randint(0, 255)
# #         b = random.randint(0, 255)
# #         colors.append(f'rgb({r}, {g}, {b})')
    
# #     return colors


# # import plotly.graph_objects as go
# # import streamlit as st
# # import numpy as np

# # def plot_3d_bar_graph(data):
# #     # Initialize a Plotly figure
# #     fig = go.Figure()

# #     if not data:
# #         st.write("No data available to plot.")
# #         return

# #     # Extracting data for plotting
# #     x = []
# #     y1 = []  # Min values for plotting
# #     y2 = []  # Max values for plotting
# #     categories = []

# #     for i, (key, value) in enumerate(data.items()):
# #         categories.append(key)  # Categories
# #         x.append(i)  # X-axis value
# #         min_val = int(value.get('min', '0').replace('%', ''))  # Default to 0 if 'min' is missing
# #         max_val = int(value.get('max', '0').replace('%', ''))  # Default to 0 if 'max' is missing
# #         y1.append(min_val)
# #         y2.append(max_val)

# #     # Adding bars for Min Compounded Returns
# #     for i in range(len(x)):
# #         fig.add_trace(go.Scatter3d(
# #             x=[x[i], x[i]],
# #             y=[0, 0.5],
# #             z=[0, y1[i]],
# #             mode='lines',
# #             line=dict(color='red', width=10),
# #             name='Min Compounded Returns'
# #         ))

# #     # Adding bars for Max Compounded Returns
# #     for i in range(len(x)):
# #         fig.add_trace(go.Scatter3d(
# #             x=[x[i] + 0.5, x[i] + 0.5],
# #             y=[0, 0.5],
# #             z=[0, y2[i]],
# #             mode='lines',
# #             line=dict(color='blue', width=10),
# #             name='Max Compounded Returns'
# #         ))

# #     # Update layout to match the desired appearance
# #     fig.update_layout(
# #         scene=dict(
# #             xaxis=dict(
# #                 tickvals=np.arange(len(categories)) + 0.25,
# #                 ticktext=categories,
# #                 title='Investment Types'
# #             ),
# #             yaxis=dict(title=''),
# #             zaxis=dict(title='Allocation (%)')
# #         ),
# #         title='Investment Allocation 3D Bar Graph',
# #         legend=dict(x=0.1, y=0.9)
# #     )

# #     st.plotly_chart(fig)




# # # def plot_3d_bar_graph(data):
# # #     # Initialize a Plotly figure
# # #     fig = go.Figure()

# # #     if not data:
# # #         print("No data available to plot.")
# # #         return

# # #     # Extracting data for plotting
# # #     x = []
# # #     y1 = []  # Min values for plotting
# # #     y2 = []  # Max values for plotting
# # #     categories = []

# # #     for i, (key, value) in enumerate(data.items()):
# # #         categories.append(key)  # Categories
# # #         x.append(i)  # X-axis value
# # #         min_val = int(value.get('min', '0').replace('%', ''))  # Default to 0 if 'min' is missing
# # #         max_val = int(value.get('max', '0').replace('%', ''))  # Default to 0 if 'max' is missing
# # #         y1.append(min_val)
# # #         y2.append(max_val)

# # #     # Adding bars for Min Compounded Returns
# # #     for i in range(len(x)):
# # #         fig.add_trace(go.Scatter3d(
# # #             x=[x[i], x[i]],
# # #             y=[0, 0.5],
# # #             z=[0, y1[i]],
# # #             mode='lines',
# # #             line=dict(color='red', width=10),
# # #             name='Min Compounded Returns'
# # #         ))

# # #     # Adding bars for Max Compounded Returns
# # #     for i in range(len(x)):
# # #         fig.add_trace(go.Scatter3d(
# # #             x=[x[i] + 0.5, x[i] + 0.5],
# # #             y=[0, 0.5],
# # #             z=[0, y2[i]],
# # #             mode='lines',
# # #             line=dict(color='blue', width=10),
# # #             name='Max Compounded Returns'
# # #         ))

# # #     # Update layout to match the desired appearance
# # #     fig.update_layout(
# # #         scene=dict(
# # #             xaxis=dict(
# # #                 tickvals=np.arange(len(categories)) + 0.25,
# # #                 ticktext=categories,
# # #                 title='Investment Types'
# # #             ),
# # #             yaxis=dict(title=''),
# # #             zaxis=dict(title='Allocation (%)')
# # #         ),
# # #         title='Investment Allocation 3D Bar Graph',
# # #         legend=dict(x=0.1, y=0.9)
# # #     )

# # #     fig.show()
# # #     return fig


 
# # # def client_form():
# # #     st.title("Client Details Form")

# # #     with st.form("client_form"):
# # #         st.header("Personal Information")
# # #         client_name = st.text_input("Client Name")
# # #         co_client_name = st.text_input("Co-Client Name")
# # #         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
# # #         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
# # #         today_date = st.date_input("Today's Date")
        
# # #         st.header("Financial Information")
# # #         current_assets = st.text_area("Current Assets (e.g., type and value)")
# # #         liabilities = st.text_area("Liabilities (e.g., type and amount)")
# # #         annual_income = st.text_area("Current Annual Income (source and amount)")
# # #         annual_contributions = st.text_area("Annual Contributions (e.g., retirement savings)")

# # #         st.header("Insurance Information")
# # #         life_insurance = st.text_input("Life Insurance (e.g., coverage amount)")
# # #         disability_insurance = st.text_input("Disability Insurance (e.g., coverage amount)")
# # #         long_term_care = st.text_input("Long-Term Care Insurance (e.g., coverage amount)")

# # #         st.header("Estate Planning")
# # #         will_status = st.radio("Do you have a will?", ["Yes", "No"])
# # #         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
# # #         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
# # #         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

# # #         # Submit button
# # #         submitted = st.form_submit_button("Submit")

# # #         if submitted:
# # #             # Save form data
# # #             form_data = {
# # #                 "Client Name": client_name,
# # #                 "Co-Client Name": co_client_name,
# # #                 "Client Age": client_age,
# # #                 "Co-Client Age": co_client_age,
# # #                 "Today's Date": str(today_date),
# # #                 "Current Assets": current_assets,
# # #                 "Liabilities": liabilities,
# # #                 "Annual Income": annual_income,
# # #                 "Annual Contributions": annual_contributions,
# # #                 "Life Insurance": life_insurance,
# # #                 "Disability Insurance": disability_insurance,
# # #                 "Long-Term Care Insurance": long_term_care,
# # #                 "Will Status": will_status,
# # #                 "Trust Status": trust_status,
# # #                 "Power of Attorney": power_of_attorney,
# # #                 "Healthcare Proxy": healthcare_proxy,
# # #             }
            
# # #             # Save to a file or database
# # #             with open("client_data.txt", "a") as f:
# # #                 f.write(str(form_data) + "\n")
            
# # #             st.success("Form submitted successfully!")
# # #             st.session_state.page = "main"  # Redirect back to main page after form submission


# # from datetime import date  # Make sure to import the date class


# # # Function to parse financial data from the text
# # import re

# # def parse_financial_data(text_content):
# #     assets = []
# #     liabilities = []

# #     # Define regex patterns to capture text following headings
# #     asset_pattern = re.compile(r"MY ASSETS:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)
# #     liability_pattern = re.compile(r"LIABILITIES:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)

# #     # Extract assets
# #     asset_matches = asset_pattern.findall(text_content)
# #     if asset_matches:
# #         asset_text = asset_matches[0]
# #         # Further processing to extract individual asset values if they are detailed
# #         asset_lines = asset_text.split('\n')
# #         for line in asset_lines:
# #             match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
# #             if match:
# #                 asset_value = float(match.group().replace(",", ""))
# #                 assets.append(asset_value)

# #     # Extract liabilities
# #     liability_matches = liability_pattern.findall(text_content)
# #     if liability_matches:
# #         liability_text = liability_matches[0]
# #         # Further processing to extract individual liability values if they are detailed
# #         liability_lines = liability_text.split('\n')
# #         for line in liability_lines:
# #             match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
# #             if match:
# #                 liability_value = float(match.group().replace(",", ""))
# #                 liabilities.append(liability_value)

# #     print("Assets Found:", assets)
# #     print("Liabilities Found:", liabilities)

# #     return assets, liabilities



# # # Function to extract numerical values from a text input
# # def extract_numeric(value):
# #     try:
# #         return float(re.sub(r'[^\d.]', '', value))  # Remove non-numeric characters and convert to float
# #     except ValueError:
# #         return 0


# # # plots graph from the details of the form :


# # def is_float(value):
# #     try:
# #         float(value)
# #         return True
# #     except ValueError:
# #         return False



# # def plot_assets_liabilities_pie_chart(assets, liabilities, threshold=50): # best plot 
# #     """
# #     Plots separate pie charts for assets and liabilities. If there are any categories
# #     below a specified threshold, they are plotted in an additional small pie chart.
    
# #     Parameters:
# #     - assets: dict, keys are asset names, values are their amounts.
# #     - liabilities: dict, keys are liability names, values are their amounts.
# #     - threshold: int, percentage threshold below which segments are considered small.
# #     """
# #     # Update matplotlib settings to increase the font size globally
# #     # plt.rcParams.update({'font.size': 32})

# #     plt.rcParams.update({'font.size': 16})

# #     def plot_pie(data, title):
# #         # Filter out zero values and create a summary for small segments
# #         total = sum(data.values())
# #         filtered_data = {k: v for k, v in data.items() if (v / total) >= threshold / 100}
# #         small_segments = {k: v for k, v in data.items() if (v / total) < threshold / 100}
# #         small_total = sum(small_segments.values())

# #         # Plotting logic
# #         if small_segments:
# #             fig, (ax_main, ax_small) = plt.subplots(1, 2, figsize=(30, 15))  # Side-by-side layout
# #         else:
# #             fig, ax_main = plt.subplots(figsize=(30, 20))  # Only main chart with larger size

# #             # fig, ax_main = plt.subplots(figsize=(10, 10))  # Only main chart with larger size

# #         # Plot main pie chart
# #         labels_main = list(filtered_data.keys()) + ([f"Other small {title}"] if small_segments else [])
# #         values_main = list(filtered_data.values()) + ([small_total] if small_segments else [])
# #         wedges_main, texts_main, autotexts_main = ax_main.pie(
# #             values_main, labels=labels_main, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
# #             startangle=140, textprops={'fontsize': 28} #18}  # Larger font size for labels
# #         )

# #         ax_main.set_title(title, fontsize=20)
# #         # Position legend to the right of the plot to avoid overlapping
# #         ax_main.legend(wedges_main, labels_main, title="Categories", loc="upper right", bbox_to_anchor=(0.001, 0.9), fontsize= 28)#14)

# #         if small_segments:
# #             # Plot additional small pie chart for small segments
# #             labels_small = list(small_segments.keys())
# #             values_small = list(small_segments.values())
# #             wedges_small, texts_small, autotexts_small = ax_small.pie(
# #                 values_small, labels=labels_small, autopct='%1.1f%%', colors=plt.cm.Paired.colors, 
# #                 startangle=140, textprops={'fontsize': 24} #14}  # Consistent label size for small chart
# #             )
# #             ax_small.set_title(f"Small Segments of {title}", fontsize=20)
# #             # Position legend to the right of the small pie chart but slightly lower to avoid overlap with the main chart's legend
# #             ax_small.legend(wedges_small, labels_small, title="Small Categories", loc="center left", bbox_to_anchor=(1.2, 0.3), fontsize= 22)#12)

# #         st.pyplot(fig)

# #     # Convert valid entries to float, ensuring only numeric values are considered
# #     assets = {k: float(v) for k, v in assets.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}
# #     liabilities = {k: float(v) for k, v in liabilities.items() if isinstance(v, (str, float)) and is_float(v) and float(v) > 0.0}

# #     # Plot pie charts
# #     plot_pie(assets, 'Distribution of Assets')
# #     plot_pie(liabilities, 'Distribution of Liabilities')

# # # def plot_assets_liabilities_pie_chart(assets, liabilities):# properly plots a big and 1 small pie chart for both assets and liability
# # #     # Filter and convert values to float, handle non-numeric or empty inputs
# # #     filtered_assets = {k: float(v) for k, v in assets.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}
# # #     filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and is_float(v) and float(v) > 0 and 'interest' not in k.lower() and 'time' not in k.lower()}

# # #     # Combine assets and liabilities for total calculation
# # #     all_values = {**filtered_assets, **filtered_liabilities}
# # #     total_value = sum(all_values.values())

# # #     # Separate main and small segments
# # #     main_segments = {k: v for k, v in all_values.items() if (v / total_value) >= 0.05}
# # #     small_segments = {k: v for k, v in all_values.items() if (v / total_value) < 0.05}
# # #     small_total = sum(small_segments.values())

# # #     # Prepare data for main pie chart
# # #     main_labels = list(main_segments.keys()) + (["Others"] if small_segments else [])
# # #     main_values = list(main_segments.values()) + ([small_total] if small_segments else [])

# # #     # Prepare data for small pie chart (only if there are small segments)
# # #     small_labels = list(small_segments.keys())
# # #     small_values = list(small_segments.values())

# # #     fig, ax = plt.subplots(figsize=(8, 6))

# # #     # Plot main pie chart
# # #     wedges, texts, autotexts = ax.pie(
# # #         main_values,
# # #         labels=main_labels,
# # #         autopct='%1.1f%%',
# # #         startangle=140,
# # #         colors=plt.cm.Paired.colors,
# # #     )

# # #     # Explode the "Others" slice
# # #     if small_segments:
# # #         others_index = main_labels.index("Others")
# # #         wedges[others_index].set_edgecolor('white')
# # #         # wedges[others_index].set_linestyle('--')
# # #         wedges[others_index].set_linewidth(2)
# # #         wedges[others_index].set_hatch('/')

# # #     ax.set_title('Assets and Liabilities Distribution')

# # #     # Draw a second pie chart for "Others"
# # #     if small_segments:
# # #         fig2, ax2 = plt.subplots(figsize=(8, 6))
# # #         wedges_small, texts_small, autotexts_small = ax2.pie(
# # #             small_values,
# # #             labels=small_labels,
# # #             autopct='%1.1f%%',
# # #             startangle=140,
# # #             colors=plt.cm.Pastel1.colors
# # #         )

# # #         ax2.set_title('Detailed View of "Others" Categories')

# # #     plt.tight_layout()
# # #     st.pyplot(fig)
# # #     if small_segments:
# # #         st.pyplot(fig2)



# # def save_data_to_file(form_data):
# #     file_path = 'client_data.txt'
# #     with open(file_path, 'a') as file:
# #         file.write(str(form_data) + "\n")
# #     # st.success(f"Form data saved to {file_path}")
# #     print(f"Form data saved to {file_path}")
    

# # def client_form():
# #     st.title("Client Details Form")

# #     with st.form("client_form"):
# #         st.header("Personal Information")
# #         client_name = st.text_input("Client Name")
# #         co_client_name = st.text_input("Co-Client Name")
# #         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
# #         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
# #         today_date = st.date_input("Today's Date")

# #         st.header("Your Assets (in $)")

# #         assets = {
# #             # 'Annual Income': st.text_input("Annual Income (e.g. , Your Annual Salary Income or other source of income) "),
# #             'Cash/Bank Account': st.text_input("Cash/Bank Account"),
# #             '401(k), 403(b), 457 Plans': st.text_input("Your 401(k), 403(b), 457 Plans "),
# #             'Traditional, SEP and SIMPLE IRAs': st.text_input("Traditional, SEP and SIMPLE IRAs "),
# #             'Roth IRA,Roth 401(k)': st.text_input("Roth IRA, Roth 401(k)"),
# #             'Brokerage/non-qualified accounts': st.text_input("Brokerage/non-qualified accounts"),
# #             'Annuities': st.text_input("Annuities"),
# #             '529 Plans': st.text_input("529 Plans"),
# #             'Home': st.text_input("Home"),
# #             'Other Real Estate': st.text_input("Other Real Estate"),
# #             'Business': st.text_input("Business"),
# #             'Other': st.text_input("Other")
# #         }
# #         st.header("Your Liabilities (in $)")

# #         liabilities = {
# #             'Mortgage': st.text_input("Mortgage"),
# #             # 'Annual Mortgage Interest Rate': st.number_input("Annual Mortgage Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=12.0, step=0.5),
# #             # 'Mortagage Time Period': st.number_input("Mortagage Time Period (Mention the time period of the Mortgage in years)", min_value=0, max_value=100,value=10,step=1),

# #             'Home Loans': st.text_input("Home Loans"),
# #             # 'Home Loans Interest Rate': st.number_input("Home Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0, value=10.0, step=0.5),
# #             # 'Home Loans Time Period': st.number_input("Home Loans Time Period (Mention the time period of the Home Loan in years)", min_value=0, max_value=100,value=15,step=1),

# #             'Vehicle Loans': st.text_input("Vehicle Loans"),
# #             # 'Vehicle Loans Interest Rate': st.number_input("Vehicle Loan Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
# #             # 'Vehicle Loans Time Period': st.number_input("Vehicle Loans Time Period (Mention the time period of the Car/Vehicle Loan in years)", min_value=0, max_value=100,value=15,step=1),

# #             'Education Loans': st.text_input("Education Loans"),
# #             # 'Education Loans Interest Rate' : st.number_input("Education Loans Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),
# #             # 'Education Loans Time Period': st.number_input("Education Loans Time Period (Mention the time period of the Education Loan in years)", min_value=0, max_value=100,value=15,step=1),

# #             # 'Credit Card': st.text_input("Monthly Credit Card Debt (Mention Amount)"),
# #             # 'Credit Card Debt Interest Rate': st.number_input("Credit Card Debt Interest Rate (in Percentage%)", min_value=0.0, max_value=100.0,value=10.0, step=0.5),

# #             'Miscellaneous': st.text_input("Miscellaneous"),
# #         }

# #         st.header("Your Retirement Goal")
# #         retirement_age = st.number_input("At what age do you plan to retire?", min_value=0, max_value=120, value=65, step=1)
# #         retirement_income = st.text_input("Desired annual retirement income")

# #         st.header("Your Other Goals")
# #         goal_name = st.text_input("Name of the Goal (e.g . , Dream House, Travel, Educational, etc.)")
# #         goal_amount = st.text_input("Amount needed for the goal (in $)")
# #         goal_timeframe = st.number_input("Timeframe to achieve the goal (in years)", min_value=0, max_value=100, value=5, step=1)

# #         st.header("Insurance Information")
# #         life_insurance_Benefit = st.text_input("Life Insurance-Benefit")
# #         life_insurance_Premium = st.text_input("Life Insurance-Premium")
# #         disability_insurance_Benefit = st.text_input("Disability Insurance-Benefit")
# #         disability_insurance_Premium = st.text_input("Disability Insurance-Premium")
# #         long_term_care_benefit = st.text_input("Long-Term Care Insurance-Benefit")
# #         long_term_care_premium = st.text_input("Long-Term Care Insurance-Premium")


# #         st.header("Estate Planning")
# #         will_status = st.radio("Do you have a will?", ["Yes", "No"])
# #         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
# #         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
# #         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

# #         submitted = st.form_submit_button("Submit")

# #         if submitted:
# #             form_data = {
# #                 "Client Name": client_name,
# #                 "Co-Client Name": co_client_name,
# #                 "Client Age": client_age,
# #                 "Co-Client Age": co_client_age,
# #                 "Today's Date": str(today_date),
# #                 "Assets": assets,
# #                 "Liabilities": liabilities,
# #                 "Retirement Age": retirement_age,
# #                 "Desired Retirement Income": retirement_income,
# #                 "Goal Name": goal_name,
# #                 "Goal Amount": goal_amount,
# #                 "Goal Timeframe": goal_timeframe,
# #                 "Life Insurance Benefit": life_insurance_Benefit,
# #                 "Life Insurance Premium": life_insurance_Premium,
# #                 "Disability Insurance Benefit": disability_insurance_Benefit,
# #                 "Disability Insurance Premium": disability_insurance_Premium,
# #                 "Long-Term Care Insurance Benefit": long_term_care_benefit,
# #                 "Long-Term Care Insurance Premium": long_term_care_premium,
# #                 "Will Status": will_status,
# #                 "Trust Status": trust_status,
# #                 "Power of Attorney": power_of_attorney,
# #                 "Healthcare Proxy": healthcare_proxy,
# #             }

# #             save_data_to_file(form_data)
            
# #             # # Plot the pie chart
# #             # st.subheader("Assets and Liabilities Breakdown")
# #             # plot_assets_liabilities_pie_chart(assets, liabilities)

# #             # Store data in session state and redirect to main
# #             st.session_state.assets = assets
# #             st.session_state.liabilities = liabilities
# #             st.session_state.total_assets, st.session_state.total_liabilities = calculate_totals(assets, liabilities)
# #             st.session_state.page = "main"
# #             st.success("Data submitted!\nThank You for filling the form !\nReturning to main portal...")

# # import math
# # def calculate_compounded_amount(principal, rate, time):
# #     """
# #     Calculates the compounded amount using the formula:
# #     A = P * (1 + r/n)^(nt)
# #     Assuming n (compounding frequency) is 1 for simplicity (annually).
# #     """
# #     if principal == 0 or rate == 0 or time == 0:
# #         return principal
# #     else:
# #         # Using annual compounding
# #         return principal * (1 + rate / 100) ** time
    
# # def calculate_totals(assets, liabilities):
# #     total_assets = sum(extract_numeric(v) for v in assets.values())
# #     print(f"Total Assets : {total_assets}")
# #     total_liabilities = 0
# #     total_liabilities = sum(extract_numeric(v) for v in liabilities.values() )

# #     # total_liabilities += calculate_compounded_amount(
# #     #     extract_numeric(liabilities['Mortgage']),
# #     #     liabilities['Annual Mortgage Interest Rate'],
# #     #     liabilities['Mortagage Time Period']
# #     # )
# #     # total_liabilities += calculate_compounded_amount(
# #     #     extract_numeric(liabilities['Home Loans']),
# #     #     liabilities['Home Loans Interest Rate'],
# #     #     liabilities['Home Loans Time Period']
# #     # )
# #     # total_liabilities += calculate_compounded_amount(
# #     #     extract_numeric(liabilities['Vehicle Loans']),
# #     #     liabilities['Vehicle Loans Interest Rate'],
# #     #     liabilities['Vehicle Loans Time Period']
# #     # )
# #     # total_liabilities += calculate_compounded_amount(
# #     #     extract_numeric(liabilities['Education Loans']),
# #     #     liabilities['Education Loans Interest Rate'],
# #     #     liabilities['Education Loans Time Period']
# #     # )
    
# #     # For credit card debt, only calculate compounded amount if interest rate > 0

# #     # credit_card_balance = extract_numeric(liabilities['Credit Card'])
# #     # credit_card_interest = liabilities['Credit Card Debt Interest Rate']
# #     # if credit_card_interest > 0:
# #     #     # Assuming the time period for credit card debt is 1 year for compounding
# #     #     total_liabilities += calculate_compounded_amount(credit_card_balance, credit_card_interest, 1)
# #     # else:
# #     #     total_liabilities += credit_card_balance
    
# #     # Miscellaneous debts are taken directly as is
# #     total_liabilities += extract_numeric(liabilities['Miscellaneous'])
# #     rounded_liabilities = round(total_liabilities,2)

# #     print(f"Total liabilities :{total_liabilities}")
# #     print(f"Rounded of Total liabilities :{rounded_liabilities}")

# #     return total_assets, rounded_liabilities #total_liabilities

# # def create_financial_summary_table(assets, liabilities):
# #     # Filter out items with zero value
# #     filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
# #     filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

# #     # Create DataFrames for assets and liabilities with indices starting from 1
# #     assets_df = pd.DataFrame(
# #         list(filtered_assets.items()), 
# #         columns=['Assets', 'Amount ($)'], 
# #         index=range(1, len(filtered_assets) + 1)
# #     )
# #     liabilities_df = pd.DataFrame(
# #         list(filtered_liabilities.items()), 
# #         columns=['Liabilities', 'Amount ($)'], 
# #         index=range(1, len(filtered_liabilities) + 1)
# #     )

# #     # Calculate total
# #     total_assets, total_liabilities = calculate_totals(assets, liabilities)

# #     # Add total row with index incremented by 1
# #     total_assets_row = pd.DataFrame(
# #         [['TOTAL', total_assets]], 
# #         columns=['Assets', 'Amount ($)'], 
# #         index=[len(assets_df) + 1]
# #     )
# #     total_liabilities_row = pd.DataFrame(
# #         [['TOTAL', total_liabilities]], 
# #         columns=['Liabilities', 'Amount ($)'], 
# #         index=[len(liabilities_df) + 1]
# #     )

# #     # Append total rows to DataFrames
# #     assets_df = pd.concat([assets_df, total_assets_row])
# #     liabilities_df = pd.concat([liabilities_df, total_liabilities_row])

# #     # Display tables with formatted values
# #     st.subheader("Assets")
# #     st.table(assets_df.style.format({'Amount ($)': '{:,.2f}'}))

# #     st.subheader("Liabilities")
# #     st.table(liabilities_df.style.format({'Amount ($)': '{:,.2f}'}))


# # def plot_bar_graphs(assets, liabilities):
# #     # Filter out items with zero values
# #     filtered_assets = {k: float(v) for k, v in assets.items() if v and float(v) > 0.0}
# #     filtered_liabilities = {k: float(v) for k, v in liabilities.items() if v and float(v) > 0.0}

# #     # Calculate compounded liabilities
# #     # compounded_liabilities = {} 

# #     # for k, v in filtered_liabilities.items():
# #         # if 'Interest Rate' in k or 'Time Period' in k:
# #         #     continue  # Skip non-monetary entries

# #         # if k == 'Credit Card Payment' and liabilities['Credit Card Debt Interest Rate'] == 0.0:
# #         #     continue  # Skip if credit card interest rate is zero

# #         # if k == 'Mortgage':
# #         #     interest_rate = liabilities['Annual Mortgage Interest Rate']
# #         #     time_period = liabilities['Mortagage Time Period']

# #         # elif k == 'Home Loans':
# #         #     interest_rate = liabilities['Home Loans Interest Rate']
# #         #     time_period = liabilities['Home Loans Time Period']

# #         # elif k == 'Car/Vehicle Loans':
# #         #     interest_rate = liabilities['Car/Vehicle Loans Interest Rate']
# #         #     time_period = liabilities['Car/Vehicle Loans Time Period']

# #         # elif k == 'Education Loans':
# #         #     interest_rate = liabilities['Education Loans Interest Rate']
# #         #     time_period = liabilities['Education Loans Time Period']

# #         # elif k == 'Credit Card Payment':
# #         #     interest_rate = liabilities['Credit Card Debt Interest Rate']
# #         #     time_period = 1  # Assuming interest is calculated yearly

# #         # if interest_rate > 0:
# #         #     compounded_amount = float(v) * (1 + float(interest_rate) / 100) ** float(time_period)
# #         #     compounded_liabilities[k] = compounded_amount
# #         # else:
# #         #     compounded_liabilities[k] = float(v)

# #     # Plot bar graph for assets
# #     st.write("### All Assets ")
# #     fig1, ax1 = plt.subplots()
# #     ax1.bar(filtered_assets.keys(), filtered_assets.values(), color='green')
# #     ax1.set_ylabel('Amount ($)')
# #     ax1.set_xlabel('Asset Type')
# #     ax1.set_title(' All Assets ')
# #     plt.xticks(rotation=45)
# #     st.pyplot(fig1)

# #     # Plot bar graph for liabilities
# #     st.write("### All Liabilities ")
# #     # st.write("### All Liabilities with Compounded Interest")
# #     fig2, ax2 = plt.subplots()
# #     # ax2.bar(compounded_liabilities.keys(), compounded_liabilities.values(), color='red')
# #     ax2.bar(filtered_liabilities.keys(), filtered_liabilities.values(), color='red')    
# #     ax2.set_ylabel('Amount ($)')
# #     ax2.set_xlabel('Liability Type')
# #     ax2.set_title(' All Liabilities ')

# #     # ax2.set_title(' All Liabilities with Compounded Interest')
# #     plt.xticks(rotation=45)
# #     st.pyplot(fig2)


# # from docx import Document
# # # Define a helper function to read and extract text from a DOCX file
# # def read_docx(file_path):
# #     document = Document(file_path)
# #     extracted_text = "\n".join([para.text for para in document.paragraphs])
# #     return extracted_text



# # # # Ask for investment personality
# # # investment_personality = st.selectbox(
# # #     "Select the investment personality of the client:",
# # #     ("Conservative Investor", "Moderate Investor", "Aggressive Investor")
# # # )

# # # # Step 4: Generate investment suggestions
# # # if st.button("Generate Investment Suggestions"):
# # #     st.write("Generating investment suggestions...")

# # #     # Replace this with the path to your DOCX file
# # #     file_path = "data/Financial_Investment_1.docx"

# # #     # New logic to read the file directly
# # #     try:
# # #         st.write("Loading client document from predefined path...")
# # #         extracted_text = read_docx(file_path)
# # #         st.write("Client document loaded successfully.")
# # #     except FileNotFoundError:
# # #         st.write("Client document file not found.")
# # #         st.stop()  # Stop the Streamlit app if the file is not found


# # #     # Assuming generate_investment_suggestions is an async function
    
# # #     suggestions = asyncio.run(generate_investment_suggestions(investment_personality, extracted_text))
# # #     st.write(suggestions)

# # #     data_extracted = extract_numerical_data(suggestions)

# # #     # Streamlit app for visualizing investment allocation
# # #     st.title('Investment Allocation Infographics')

# # #     st.write('## Investment Allocation Charts')
# # #     fig = plot_investment_allocations(data_extracted)
# # #     st.pyplot(fig)

# # #     st.write('## Pie Chart of Investment Allocation')
# # #     fig = plot_pie_chart(data_extracted)
# # #     st.pyplot(fig)

# # #     st.write('## Bar Chart of Compounded Returns')
# # #     fig = bar_chart(data_extracted['Growth-Oriented Investments'])
# # #     st.pyplot(fig)

# # #     plot_3d_bar_graph(data_extracted)


# # class TrieNode:
# #     def __init__(self):
# #         self.children = {}
# #         self.client_ids = []
# #         self.end_of_name = False  # Marks the end of a client's name

# # class Trie:
# #     def __init__(self):
# #         self.root = TrieNode()

# #     def insert(self, name, client_id):
# #         node = self.root
# #         for char in name:
# #             if char not in node.children:
# #                 node.children[char] = TrieNode()
# #             node = node.children[char]
# #         node.client_ids.append(client_id)
# #         node.end_of_name = True

# #     def search(self, prefix):
# #         node = self.root
# #         for char in prefix:
# #             if char in node.children:
# #                 node = node.children[char]
# #             else:
# #                 return []  # Prefix not found
# #         return self._get_all_names_from_node(prefix, node)

# #     def _get_all_names_from_node(self, prefix, node):
# #         suggestions = []
# #         if node.end_of_name:
# #             suggestions.append((prefix, node.client_ids))
# #         for char, child_node in node.children.items():
# #             suggestions.extend(self._get_all_names_from_node(prefix + char, child_node))
# #         return suggestions



# # def preload_trie():
# #     trie = Trie()
# #     clients = {
# #         "John Doe": "C001",
# #         "Jane Smith": "C002",
# #         "James Brown": "C003",
# #         "Jill Johnson": "C004",
# #         "Jake White": "C005"
# #     }
# #     for name, client_id in clients.items():
# #         trie.insert(name.lower(), client_id)  # Insert in lowercase for case-insensitive search
# #     return trie

# # async def generate_investment_suggestions_for_investor(investment_personality, context): # # GET Method for py , for front end its Post API
    
# #     # retriever = asyncio.run(load_vector_db("uploaded_file"))

# #     retriever =  await load_vector_db("uploaded_file")
# #     # retriever = await load_vector_db("data\Financial_Investment_1.docx") 

# #     chain = await make_retrieval_chain(retriever)

# #     # chain = asyncio.run(make_retrieval_chain(retriever))
    
# #     if chain is not None:
# #         # summary = context
# #         # query = summary + "\n" + investment_personality
# #         query = str(investment_personality)
# #         response = chain.invoke({"input": query})
        
# #         # format_response = markdown_to_text(response['answer'])
# #         # return format_response
        
# #         # html_output = markdown.markdown(response['answer'])
# #         # return html_output
        
# #         # readable_text = markdown_to_readable_text(response['answer'])
# #         # print(readable_text)
# #         # return readable_text

# #         # format_text = convert_to_markdown(response['answer'])
# #         # return format_text
        
# #         return response['answer']
    
        

# #         # handle_graph(response['answer'])

# #     else:
# #         logging.INFO("response is not generated by llm model")
# #         return jsonify("response is not generated by llm model"),500
# #         # st.error("Failed to create the retrieval chain. Please upload a valid document.")

# # from flask import Flask, request, jsonify, send_file
# # import asyncio
# # from flask_cors import CORS
# # app = Flask(__name__)
# # CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})
# # # CORS(app,resources={r"/api/*":{"origins":"*"}})
# # # CORS(app)

# # # Initialize the Trie with preloaded clients
# # trie = preload_trie()

# # @app.route('/')
# # def home():
# #     return "Wealth Advisor Chatbot API"

# # @app.route('/investment-suggestions', methods=['POST'])
# # def investment_suggestions():
# #     # Get the input data (new or existing client)
# #     data = request.get_json()

# #     # Determine if it's a new client or existing client
# #     client_type = data.get("client_type")

# #     if client_type == "New Client":
# #         # Get form details and perform investment suggestions

# #         # Check if assets and liabilities are provided
# #         assets = data.get('assets', None)
# #         liabilities = data.get('liabilities', None)

# #         if assets and liabilities:
# #             financial_summary = create_financial_summary_table(assets, liabilities)
# #             bar_graphs = plot_bar_graphs(assets, liabilities)
# #             pie_chart = plot_assets_liabilities_pie_chart(assets, liabilities)

# #             return jsonify({
# #                 "financial_summary": financial_summary,
# #                 "bar_graphs": "Bar graphs generated.",
# #                 "pie_chart": "Pie chart generated."
# #             })

# #         return jsonify({"message": "Please fill in the client details to view the assets and liabilities breakdown."})

# #     elif client_type == "Existing Client":
# #         # Search for an existing client in the Trie
# #         search_query = data.get("search_query", "").lower()
# #         matching_names = trie.search(search_query)

# #         if matching_names:
# #             suggestions = [{"name": name, "client_ids": client_ids} for name, client_ids in matching_names]
# #             return jsonify({"suggestions": suggestions})
# #         else:
# #             return jsonify({"message": "No matching clients found."})
    
# #     return jsonify({"message": "Invalid client type."})




# # # @app.route('/investment-personality-assessment', methods=['POST']) # 1st code
# # # def investment_personality_assessment():
# # #     file = request.files['file']
    
# # #     if file:
# # #         # Process the uploaded file and extract responses
# # #         responses = extract_responses_from_docx(file)

# # #         if responses:
# # #             # Determine investment personality
# # #             personality = asyncio.run(determine_investment_personality(responses))
# # #             return jsonify({
# # #                 "responses": responses,
# # #                 "investment_personality": personality
# # #             })
        
# # #         return jsonify({"message": "No responses found in the document."})
    
# # #     return jsonify({"message": "No file uploaded."})

# # # @app.route('/investment-personality-assessment', methods=['POST']) # 1st test code for investment-personality-assessment
# # # def investment_personality_assessment():
# # #     try:
# # #         # if 'personalityFile' not in request.files:
# # #         #     return jsonify({'message': 'Personality file is required!'}), 400
        
# # #         personality_file = request.files['personalityFile']
        
# # #         # Process the personality file (DOCX/PDF) here
# # #         # Call the function to extract responses
# # #         responses = extract_responses_from_docx(personality_file)
        
# # #         if not responses:
# # #             return jsonify({'message': 'No responses found in the document.'}), 400
        
# # #         # Determine the investment personality using the extracted responses
# # #         personality = asyncio.run(determine_investment_personality(responses))
        
# # #         return jsonify({'investmentPersonality': personality}), 200
    
# # #     except Exception as e:
# # #         print(f"Error in personality assessment: {e}")
# # #         return jsonify({'message': 'Internal Server Error'}), 500

# # # @app.route('/generate-investment-suggestions', methods=['POST']) # test code for POST requests
# # # def generate_investment_suggestions_api():
# # #     if request.is_json:   

# # #         data = request.get_json()
# # #         # Your processing logic
# # #         return jsonify({"message": "Success"})

# # #     return jsonify({"message": "Invalid content type, expecting application/json"}), 415


# # # @app.route('/generate-investment-suggestions', methods=['POST']) # 1st test code
# # # def generate_investment_suggestions():
# # #     try:
# # #         if 'personalityFile' not in request.files or 'financialFile' not in request.files:
# # #             return jsonify({'message': 'Both personality and financial files are required!'}), 400

# # #         personality_file = request.files['personalityFile']
# # #         financial_file = request.files['financialFile']

# # #         # Extract personality data and financial data
# # #         responses = extract_responses_from_docx(personality_file)
# # #         financial_data = asyncio.run(process_document(financial_file))

# # #         if not responses or not financial_data:
# # #             return jsonify({'message': 'Failed to process one or both of the documents.'}), 400

# # #         # Determine investment personality
# # #         personality = asyncio.run(determine_investment_personality(responses))

# # #         # Generate investment suggestions based on personality and financial data
# # #         suggestions = asyncio.run(generate_investment_suggestions(personality, financial_data))

# # #         return jsonify({'investmentSuggestions': suggestions}), 200

# # #     except Exception as e:
# # #         print(f"Error in generating investment suggestions: {e}")
# # #         return jsonify({'message': 'Internal Server Error'}), 500





# # # @app.route('/upload-personal-details', methods=['POST']) # not necessary
# # # def upload_personal_details():
# # #     file = request.files['file']

# # #     if file:
# # #         # Process the uploaded file (personal details document)
# # #         document_data = asyncio.run(process_document(file))
        
# # #         if isinstance(document_data, tuple) and len(document_data) == 2:
# # #             extracted_text, tables_content = document_data
# # #             return jsonify({"message": "File processed", "extracted_text": extracted_text})
        
# # #         return jsonify({"message": "Unexpected data format returned from document processing."})
    
# # #     return jsonify({"message": "No file uploaded."})

# # import logging

# # @app.route('/investment-personality-assessment', methods=['POST'])
# # def investment_personality_assessment():
# #     try:
# #         # Check if files are present in the request
# #         # if 'assessmentFile' not in request.files:
# #         #     return jsonify({'message': 'Assessment file is required!'}), 400

# #         assessment_file = request.files['assessmentFile']
        
# #         # Process the assessment file here (DOCX/PDF)
# #         responses = extract_responses_from_docx(assessment_file)
        
# #         if not responses:
# #             logging.info("Assessment file not found")
# #             return jsonify({'message': 'No responses found in the assessment file.'}), 400
        
# #         # Determine investment personality based on the extracted responses
# #         personality = asyncio.run(determine_investment_personality(responses))
        
# #         return jsonify({'investmentPersonality': personality}), 200
    
# #     except Exception as e:
# #         logging.info(f"Error in personality assessment: {e}")
# #         return jsonify({'message': 'Internal Server Error in Investor Personality'}), 500

# # # Route to handle generating investment suggestions
# # @app.route('/generate-investment-suggestions', methods=['POST'])
# # def generate_investment_suggestions():
# #     try:
# #         # Check if files are present in the request
# #         # if 'assessmentFile' not in request.files or 'financialFile' not in request.files:
# #         #     return jsonify({'message': 'Both assessment and financial files are required!'}), 400
# #         try :
# #             assessment_file = request.files['assessmentFile']
# #             financial_file = request.files['financialFile']
# #             logging.info(" Requested files")
# #         except Exception as e:
# #             logging.info(" Requested files not passed")
# #             return jsonify({'message': f'Error occurred while retrieving files12: {e}'}), 400
            
        
# #         try :
# #             # Extract personality data from assessment file and financial data
# #             try:
# #                 responses = extract_responses_from_docx(assessment_file)
# #             except Exception as e:
# #                 logging.info(f"Failed to extract responses from assessment file: {e}")
# #                 return jsonify({'message': 'Failed to extract responses from assessment file.'}), 400
# #             try:
# #                 financial_data = asyncio.run(process_document(financial_file))
# #             except Exception as e:
# #                 logging.info(f"Failed to process financial file: {e}")
# #                 return jsonify({'message': 'Failed to process financial file.'}), 400

# #             logging.info(f"Received Responses from the file {responses}")
# #         except Exception as e:
# #             logging.info("Failed to process files")
# #             return jsonify({'message': f'Error occurred while processing files: {e}'}), 400

# #         # if not responses or not financial_data:
# #         #     return jsonify({'message': 'Failed to process one or both of the documents.'}), 400
# #         try:
# #             # Determine investment personality
# #             personality = asyncio.run(determine_investment_personality(responses))
# #             logging.info(f"\nPersonality of the user is : {personality}")
# #         except Exception as e:
# #             logging.info("Failed to determine personality")
# #             return jsonify({'message': f'Error occurred while determining personality: {e}'}), 400
        
# #         try:
# #             # Generate investment suggestions based on personality and financial data
# #             suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality, financial_data))
# #             htmlSuggestions = markdown.markdown(suggestions)
# #             logging.info(f"\Suggestions for investor : \n{suggestions}")
# #         except Exception as e:
# #             logging.info("Failed to generate suggestions")
# #             return jsonify({'message': f'Error occurred while generating suggestions: {e}'}), 400

# #         logging.info("Successfully generated")
# #         formatSuggestions = markdown_to_text(suggestions)
# #         data_extracted = extract_numerical_data(suggestions)
        
# #         min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
# #                         [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
# #         max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
# #                         [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

# #         # Normalize allocations
# #         min_allocations = normalize_allocations(min_allocations)
# #         max_allocations = normalize_allocations(max_allocations)

# #         # Update Bar Chart Data
# #         bar_chart_data = {
# #             'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
# #             'datasets': [{
# #                 'label': 'Min Allocation',
# #                 'data': min_allocations,
# #                 'backgroundColor': 'skyblue'
# #             },
# #             {
# #                 'label': 'Max Allocation',
# #                 'data': max_allocations,
# #                 'backgroundColor': 'lightgreen'
# #             }]
# #         }

# #         # Similar changes can be made for the Pie Chart Data:
# #         all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
# #         num_labels = len(all_labels)
# #         max_allocations_for_pie = normalize_allocations(
# #             [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
# #             [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
# #         )
        
# #         # Generate colors based on the number of labels
# #         dynamic_colors = generate_colors(num_labels)

# #         # Update Pie Chart Data
# #         pie_chart_data = {
# #             'labels': all_labels,
# #             'datasets': [{
# #                 'label': 'Investment Allocation',
# #                 'data': max_allocations_for_pie,
# #                 'backgroundColor': dynamic_colors,
# #                 'hoverOffset': 4
# #             }]
# #         }
        
# #         # all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
# #         # num_labels = len(all_labels)
        
        
        
# #         # # Pie Chart Data
# #         # pie_chart_data = {
# #         #     'labels': all_labels,
# #         #     'datasets': [{
# #         #         'label': 'Investment Allocation',
# #         #         'data': [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) 
# #         #                 for label in data_extracted['Growth-Oriented Investments']] + 
# #         #                 [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) 
# #         #                 for label in data_extracted['Conservative Investments']],
# #         #         'backgroundColor': dynamic_colors,  # Use the dynamically generated colors
# #         #         'hoverOffset': 4
# #         #     }]
# #         # }
        
        
# #         #  # Bar Chart Data
# #         # bar_chart_data = {
# #         #     'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
# #         #     'datasets': [{
# #         #         'label': 'Min Allocation',
# #         #         'data': [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
# #         #                 [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']],
# #         #         'backgroundColor': 'skyblue'
# #         #     },
# #         #     {
# #         #         'label': 'Max Allocation',
# #         #         'data': [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
# #         #                 [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']],
# #         #         'backgroundColor': 'lightgreen'
# #         #     }]
# #         # }
        
        
# #         # pie_chart_data = { # 1st version
# #         #     'labels': list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys()),
# #         #     'datasets': [{
# #         #         'label': 'Investment Allocation',
# #         #         'data': [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
# #         #                 [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']],
# #         #         'backgroundColor': ['rgb(255, 99, 132)', 'rgb(54, 162, 235)', 'rgb(255, 205, 86)', 'rgb(75, 192, 192)', 'rgb(153, 102, 255)'],
# #         #         'hoverOffset': 4
# #         #     }]
# #         # }

        
# #         # Prepare the data for the line chart with inflation adjustment
# #         initial_investment = 10000
# #         # compounded_chart_data, inflation_adjusted_chart_data = prepare_line_chart_data_with_inflation(data_extracted, initial_investment)
# #         combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
# #         print(f"\nThe combined chart data is : {combined_chart_data}")
# #         # return jsonify({
# #         #     "status": 200,
# #         #     "message": "Success",
# #         #     "investmentSuggestions": htmlSuggestions,
# #         #     "pieChartData": pie_chart_data,
# #         #     "barChartData": bar_chart_data
# #         # }), 200
        
# #         return jsonify({
# #             "status": 200,
# #             "message": "Success",
# #             "investmentSuggestions": htmlSuggestions,
# #             "pieChartData": pie_chart_data,
# #             "barChartData": bar_chart_data,
# #             "compoundedChartData":combined_chart_data
# #         }), 200

# #     except Exception as e:
# #         logging.info(f"Error in generating investment suggestions: {e}")
# #         return jsonify({'message': 'Internal Server Error in Generating responses'}), 500




# # # Run the Flask application
# # if __name__ == '__main__':
# #     app.run(host='0.0.0.0',debug=True)







# # # last updated streamlit code :
# # # import streamlit as st 
# # # import pandas as pd
# # # import matplotlib.pyplot as plt

# # # import os
# # # import filetype
# # # import docx
# # # import PyPDF2
# # # import re
# # # from dotenv import load_dotenv
# # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # from langchain_community.vectorstores import Chroma
# # # from langchain_community.docstore.in_memory import InMemoryDocstore
# # # from langchain_community.vectorstores import FAISS
# # # from langchain_community.document_loaders import Docx2txtLoader
# # # from langchain_core.prompts import ChatPromptTemplate
# # # from langchain.chains import create_retrieval_chain
# # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # from langchain.memory import ConversationSummaryMemory
# # # import asyncio
# # # import numpy as np
# # # import json

# # # import google.generativeai as genai
# # # import pathlib
# # # import logging
# # # import sys
# # # import io
# # # import matplotlib.pyplot as plt
# # # import seaborn as sns
# # # # Import things that are needed generically
# # # from langchain.pydantic_v1 import BaseModel, Field
# # # from langchain.tools import BaseTool, StructuredTool, tool
# # # # Define functions to generate investment suggestions :

# # # load_dotenv()
# # # GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

# # # def markdown_to_text(md):
# # #     # Simple conversion for markdown to plain text
# # #     md = md.replace('**', '')
# # #     md = md.replace('*', '')
# # #     md = md.replace('_', '')
# # #     md = md.replace('#', '')
# # #     md = md.replace('`', '')
# # #     return md.strip()

# # # # Load the Vector DataBase :
# # # async def load_vector_db(file_path):
# # #     try:
# # #         print("Loading vector database...")
# # #         loader = Docx2txtLoader(file_path)
# # #         documents = loader.load()
# # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # #         text_chunks = text_splitter.split_documents(documents)
# # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # #         # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
# # #         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
# # #         # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

# # #         # vector_store = FAISS(
# # #         #     embedding_function=embeddings,
# # #         #     index=index,
# # #         #     docstore=InMemoryDocstore(),
# # #         #     index_to_docstore_id={},
# # #         # )
        
# # #         print("Vector database loaded successfully.") 
# # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # #     except Exception as e:
# # #         print(f"Error loading vector database: {e}")
# # #         return None


# # # investment_personality = "Moderate Investor"
# # # async def make_retrieval_chain(retriever):
# # #     """
# # #     Create a retrieval chain using the provided retriever.

# # #     Args:
# # #         retriever (RetrievalQA): A retriever object.

# # #     Returns:
# # #         RetrievalQA: A retrieval chain object.
# # #     """
# # #     try:
# # #         global investment_personality #,summary
# # #         llm = ChatGoogleGenerativeAI(
# # #             #model="gemini-pro",
# # #             model = "gemini-1.5-flash",
# # #             temperature=0.7,
# # #             top_p=0.85,
# # #             google_api_key=GOOGLE_API_KEY
# # #         )
# # #                                                     # \n + summary 
# # #         prompt_template = investment_personality +  "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # #                 Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality.
# # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # #                 should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
# # #                 Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # #                 investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
# # #                 Also give the user minimum and maximum expected growth in dollars for the time horizon .
# # #                 Also explain the user why you are giving them that particular investment suggestion.
# # #                 Here's an example for the required Output Format :

# # #                 Investment Suggestions for a Moderate Investor(This is for a Moderate Investor but you need to generate for any investor)

# # #                 Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

# # #                 Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

# # #                 Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
# # #                 How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
# # #                 Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
# # #                 ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
# # #                 Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
# # #                 Consider investing in blue-chip companies or growth sectors like technology. 
# # #                 Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


# # #                 Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
# # #                 Percentage Allocation: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

# # #                 Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
# # #                 How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
# # #                 Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
# # #                 Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
# # #                 Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
# # #                 Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
# # #                 Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
# # #                 Percentage Allocation: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

# # #                 Time Horizon and Expected Returns:

# # #                 Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
# # #                 Minimum Expected Annual Return: 4% - 6% 
# # #                 Maximum Expected Annual Return: 8% - 10% 
# # #                 Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, a 10,000 investment could grow to approximately 17,908 in 10 years.
# # #                 Minimum Expected Growth in Dollars: 
                
# # #                 4,000−6,000 (over 10 years) Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
# # #                 Rationale for Investment Suggestions:

# # #                 This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
# # #                 Important Considerations:

# # #                 Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

# # #                 Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

# # #                 <context>
# # #                 {context}
# # #                 </context>
# # #                 Question: {input}"""

# # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
# # #         combine_docs_chain = None  

# # #         if retriever is not None :  
# # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # #             print(retriever_chain)
# # #             return retriever_chain
# # #         else:
# # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # #             return None

# # #     except Exception as e:
# # #         print(f"Error in creating chain: {e}")
# # #         return None

# # # import streamlit as st
# # # import json
# # # import matplotlib.pyplot as plt
# # # import io



# # # # async def process_document(file_path):
# # # #     try:
# # # #         print("Processing the document")
# # # #         file_type = filetype.guess(file_path)
# # # #         if file_type is not None:
# # # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # #                 # return extract_text_from_word(file_path)
# # # #                 return extract_text_and_tables_from_word(file_path)
# # # #             elif file_type.mime == "application/pdf":
# # # #                 return extract_text_from_pdf(file_path)
# # # #         return None
# # # #     except Exception as e:
# # # #         print(f"Error processing document: {e}")
# # # #         return None


# # # # async def extract_text_from_pdf(pdf_file_path):
# # # #     try:
# # # #         print("Processing pdf file")
# # # #         with open(pdf_file_path, "rb") as pdf_file:
# # # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # # #             text_content = []
# # # #             for page_num in range(pdf_reader.numPages):
# # # #                 page = pdf_reader.getPage(page_num)
# # # #                 text_content.append(page.extract_text())
# # # #             return "\n".join(text_content)
# # # #     except Exception as e:
# # # #         print(f"Error extracting text from PDF: {e}")
# # # #         return None



# # # # async def extract_text_and_tables_from_word(docx_file_path):
# # # #     """
# # # #     Extracts text and tables from a Word document (.docx).

# # # #     Args:
# # # #         docx_file_path (str): Path to the Word document file.

# # # #     Returns:
# # # #         tuple: Extracted text content and tables from the document.
# # # #     """
# # # #     try:
# # # #         print("Extracting text and tables from word file")
# # # #         doc = docx.Document(docx_file_path)
# # # #         text_content = []
# # # #         tables_content = []

# # # #         for para in doc.paragraphs:
# # # #             text_content.append(para.text)

# # # #         for table in doc.tables:
# # # #             table_data = []
# # # #             for row in table.rows:
# # # #                 row_data = []
# # # #                 for cell in row.cells:
# # # #                     row_data.append(cell.text.strip())
# # # #                 table_data.append(row_data)
# # # #             tables_content.append(table_data)
# # # #         print("Extracted text from word file")
# # # #         return "\n".join(text_content), tables_content
# # # #     except Exception as e:
# # # #         print(f"Error extracting text and tables from Word document: {e}")
# # # #         return None, None


# # # # def parse_financial_data(text_content, tables_content):
# # # #     assets = {}
# # # #     liabilities = {}

# # # #     # Extract assets and liabilities from tables
# # # #     if tables_content:
# # # #         for table in tables_content:
# # # #             for row in table:
# # # #                 # Basic keyword matching (e.g., 'Asset' or 'Liability' detection)
# # # #                 if any("asset" in cell.lower() for cell in row):
# # # #                     key = row[0]
# # # #                     value = float(row[1].replace('$', '').replace(',', ''))
# # # #                     assets[key] = value
# # # #                 elif any("liability" in cell.lower() for cell in row):
# # # #                     key = row[0]
# # # #                     value = float(row[1].replace('$', '').replace(',', ''))
# # # #                     liabilities[key] = value

# # # #     # Additional parsing logic can be added here based on the document structure

# # # #     return assets, liabilities

# # # # # Step 3: Plot pie chart
# # # # def plot_pie_chart(assets, liabilities):
# # # #     total_assets = sum(assets.values())
# # # #     total_liabilities = sum(liabilities.values())

# # # #     labels = ['Assets', 'Liabilities']
# # # #     sizes = [total_assets, total_liabilities]

# # # #     fig, ax = plt.subplots(figsize=(8, 8))
# # # #     ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=['green', 'red'])
# # # #     ax.set_title('Assets vs Liabilities Distribution')

# # # #     st.pyplot(fig)

# # # async def process_document(file_path):
# # #     try:
# # #         print("Processing the document")
# # #         file_type = filetype.guess(file_path)
# # #         if file_type is not None:
# # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # #                 # Await the coroutine to extract text and tables
# # #                 return await extract_text_and_tables_from_word(file_path)
# # #             elif file_type.mime == "application/pdf":
# # #                 return await extract_text_from_pdf(file_path)
# # #         return None
# # #     except Exception as e:
# # #         print(f"Error processing document: {e}")
# # #         return None

# # # # Async function to extract text from a PDF file
# # # async def extract_text_from_pdf(pdf_file_path):
# # #     try:
# # #         print("Processing pdf file")
# # #         with open(pdf_file_path, "rb") as pdf_file:
# # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # #             text_content = []
# # #             for page_num in range(pdf_reader.numPages):
# # #                 page = pdf_reader.getPage(page_num)
# # #                 text_content.append(page.extract_text())
# # #             return "\n".join(text_content)
# # #     except Exception as e:
# # #         print(f"Error extracting text from PDF: {e}")
# # #         return None

# # # # Async function to extract text and tables from a Word document
# # # async def extract_text_and_tables_from_word(docx_file_path):
# # #     try:
# # #         print("Extracting text and tables from word file")
# # #         doc = docx.Document(docx_file_path)
# # #         text_content = []
# # #         tables_content = []

# # #         for para in doc.paragraphs:
# # #             text_content.append(para.text)

# # #         for table in doc.tables:
# # #             table_data = []
# # #             for row in table.rows:
# # #                 row_data = []
# # #                 for cell in row.cells:
# # #                     row_data.append(cell.text.strip())
# # #                 table_data.append(row_data)
# # #             tables_content.append(table_data)
# # #         print("Extracted text from word file")
# # #         return "\n".join(text_content), tables_content
# # #     except Exception as e:
# # #         print(f"Error extracting text and tables from Word document: {e}")
# # #         return None, None



# # # async def validate_document_content(text, tables):
# # #     """
# # #     Validates the content of the document.

# # #     Args:
# # #         text (str): Extracted text content from the document.
# # #         tables (list): Extracted tables content from the document.

# # #     Returns:
# # #         tuple: Client name and validation errors.
# # #     """
# # #     errors = []
    
# # #     # Extract client name
# # #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# # #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# # #     # Define required sections
# # #     required_sections = [
# # #         "YOUR RETIREMENT GOAL",
# # #         "YOUR OTHER MAJOR GOALS",
# # #         "YOUR ASSETS AND LIABILITIES",
# # #         "MY LIABILITIES",
# # #         "YOUR CURRENT ANNUAL INCOME"
# # #     ]

# # #     # Check for the presence of required sections
# # #     for section in required_sections:
# # #         if section not in text:
# # #             errors.append(f"* {section} section missing.")
    
# # #     # Define table field checks
# # #     table_checks = {
# # #         "YOUR RETIREMENT GOAL": [
# # #             r"When do you plan to retire\? \(age or date\)",
# # #             r"Social Security Benefit \(include expected start date\)",
# # #             r"Pension Benefit \(include expected start date\)",
# # #             r"Other Expected Income \(rental, part-time work, etc.\)",
# # #             r"Estimated Annual Retirement Expense"
# # #         ],
# # #         "YOUR OTHER MAJOR GOALS": [
# # #             r"GOAL", r"COST", r"WHEN"
# # #         ],
# # #         "YOUR ASSETS AND LIABILITIES": [
# # #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# # #             r"Current Value", r"Annual Contributions"
# # #         ],
# # #         "MY LIABILITIES": [
# # #             r"Balance", r"Interest Rate", r"Monthly Payment"
# # #         ]
# # #     }

# # #     # Validate table content
# # #     for section, checks in table_checks.items():
# # #         section_found = False
# # #         for table in tables:
# # #             table_text = "\n".join(["\t".join(row) for row in table])
# # #             if section in table_text:
# # #                 section_found = True
# # #                 for check in checks:
# # #                     if not re.search(check, table_text, re.IGNORECASE):
# # #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# # #                 break
# # #         if not section_found:
# # #             errors.append(f"* {section} section missing.")

# # #     return client_name, errors



# # # async def generate_investment_suggestions(investment_personality, context):
    
# # #     # retriever = asyncio.run(load_vector_db("uploaded_file"))

# # #     retriever = await load_vector_db("uploaded_file")
# # #     chain = await make_retrieval_chain(retriever)

# # #     # chain = asyncio.run(make_retrieval_chain(retriever))
    
# # #     if chain is not None:
# # #         # summary = context
# # #         # query = summary + "\n" + investment_personality
# # #         query = str(investment_personality)
# # #         response = chain.invoke({"input": query})
# # #         format_response = markdown_to_text(response['answer'])
# # #         return format_response
# # #         # st.write(format_response)

# # #         # handle_graph(response['answer'])

# # #     else:
# # #         st.error("Failed to create the retrieval chain. Please upload a valid document.")



# # # # Generate Infographics : Best Code so far:


# # # import seaborn as sns
# # # import re
# # # from collections import defaultdict
# # # import matplotlib.pyplot as plt
# # # import streamlit as st
# # # import numpy as np

# # # def extract_numerical_data(response):
# # #     # Define patterns to match different sections and their respective allocations
# # #     patterns = {
# # #         'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
# # #         'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
# # #         'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
# # #     }

# # #     data = defaultdict(dict)

# # #     for section, pattern in patterns.items():
# # #         match = pattern.search(response)
# # #         if match:
# # #             investments_text = match.group(1)
# # #             # Extract individual investment types and their allocations
# # #             investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
# # #             for investment_match in investment_pattern.findall(investments_text):
# # #                 investment_type, min_allocation, max_allocation = investment_match
# # #                 data[section][investment_type.strip()] = {
# # #                     'min': min_allocation,
# # #                     'max': max_allocation
# # #                 }

# # #     # Extract time horizon and expected returns
# # #     time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
# # #     min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
# # #     max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
# # #     min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
# # #     max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

# # #     time_horizon_match = time_horizon_pattern.search(response)
# # #     min_return_match = min_return_pattern.search(response)
# # #     max_return_match = max_return_pattern.search(response)
# # #     min_growth_match = min_growth_pattern.search(response)
# # #     max_growth_match = max_growth_pattern.search(response)

# # #     if time_horizon_match:
# # #         data['Time Horizon'] = {
# # #             'min_years': time_horizon_match.group(1),
# # #             'max_years': time_horizon_match.group(2)
# # #         }

# # #     if min_return_match:
# # #         data['Expected Annual Return'] = {
# # #             'min': min_return_match.group(1),
# # #             'max': min_return_match.group(2)
# # #         }

# # #     if max_return_match:
# # #         data['Expected Annual Return'] = {
# # #             'min': max_return_match.group(1),
# # #             'max': max_return_match.group(2)
# # #         }

# # #     if min_growth_match:
# # #         data['Expected Growth in Dollars'] = {
# # #             'min': min_growth_match.group(1),
# # #             'max': min_growth_match.group(2)
# # #         }

# # #     if max_growth_match:
# # #         data['Expected Growth in Dollars'] = {
# # #             'min': max_growth_match.group(1),
# # #             'max': max_growth_match.group(2)
# # #         }

# # #     return data


# # # # def plot_investment_allocations(data):
# # # #     # fig, axes = plt.subplots(1, 2, figsize=(14, 7))
# # # #     # fig, axes = plt.subplots(1, 2, figsize=(18, 9))

# # # #     # fig, axes = plt.subplots(2, 1, figsize=(18, 9))
# # # #     fig, axes = plt.subplots(2, 1, figsize=(28, 15))

# # # #     # Plot Growth-Oriented Investments
# # # #     growth_data = data['Growth-Oriented Investments']
# # # #     growth_labels = list(growth_data.keys())
# # # #     growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
# # # #     growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

# # # #     axes[0].barh(growth_labels, growth_min, color='skyblue', label='Min Allocation')
# # # #     axes[0].barh(growth_labels, growth_max, left=growth_min, color='lightgreen', label='Max Allocation')
# # # #     axes[0].set_title('Growth-Oriented Investments')
# # # #     axes[0].set_xlabel('Percentage Allocation')
# # # #     axes[0].legend()

# # # #     # Plot Conservative Investments
# # # #     conservative_data = data['Conservative Investments']
# # # #     conservative_labels = list(conservative_data.keys())
# # # #     conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
# # # #     conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

# # # #     axes[1].barh(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
# # # #     axes[1].barh(conservative_labels, conservative_max, left=conservative_min, color='lightgreen', label='Max Allocation')
# # # #     axes[1].set_title('Conservative Investments')
# # # #     axes[1].set_xlabel('Percentage Allocation')
# # # #     axes[1].legend()

# # # #     plt.tight_layout()
# # # #     return fig

# # # def plot_investment_allocations(data):
# # #     # Create subplots with a large figure size
# # #     fig, axes = plt.subplots(2, 1, figsize= (16,10)) #(28, 15))  # Adjust size as needed

# # #     # Plot Growth-Oriented Investments
# # #     growth_data = data['Growth-Oriented Investments']
# # #     growth_labels = list(growth_data.keys())
# # #     growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
# # #     growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

# # #     axes[0].bar(growth_labels, growth_min, color='skyblue', label='Min Allocation')
# # #     axes[0].bar(growth_labels, growth_max, bottom=growth_min, color='lightgreen', label='Max Allocation')
# # #     axes[0].set_title('Growth-Oriented Investments', fontsize=16)
# # #     axes[0].set_ylabel('Percentage Allocation', fontsize=14)
# # #     axes[0].set_xlabel('Investment Types', fontsize=14)
# # #     axes[0].tick_params(axis='x', rotation=45, labelsize=12)
# # #     axes[0].tick_params(axis='y', labelsize=12)
# # #     axes[0].legend()

# # #     # Plot Conservative Investments
# # #     conservative_data = data['Conservative Investments']
# # #     conservative_labels = list(conservative_data.keys())
# # #     conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
# # #     conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

# # #     axes[1].bar(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
# # #     axes[1].bar(conservative_labels, conservative_max, bottom=conservative_min, color='lightgreen', label='Max Allocation')
# # #     axes[1].set_title('Conservative Investments', fontsize=16)
# # #     axes[1].set_ylabel('Percentage Allocation', fontsize=14)
# # #     axes[1].set_xlabel('Investment Types', fontsize=14)
# # #     axes[1].tick_params(axis='x', rotation=45, labelsize=12)
# # #     axes[1].tick_params(axis='y', labelsize=12)
# # #     axes[1].legend()

# # #     # Tight layout for better spacing
# # #     plt.tight_layout()
# # #     plt.show()
# # #     return fig


# # # def plot_pie_chart(data):
# # #     fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

# # #     # Combine all investment data for pie chart
# # #     all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
# # #     labels = list(all_data.keys())
# # #     sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
# # #     colors = plt.cm.Paired(range(len(labels)))

# # #     wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
# # #     ax.set_title('Investment Allocation')

# # #     # Add legend
# # #     ax.legend(wedges, labels, title="Investment Types", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

# # #     return fig




# # # # def plot_pie_chart(data):
# # # #     fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

# # # #     # Combine all investment data for pie chart
# # # #     all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
# # # #     labels = list(all_data.keys())
# # # #     sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
# # # #     colors = plt.cm.Paired(range(len(labels)))

# # # #     ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
# # # #     ax.legend(loc='lower left',frameon= True)
# # # #     ax.set_title('Investment Allocation')

# # # # #         myexplode = [0.2, 0, 0, 0]

# # # # # plt.pie(y, labels = mylabels, explode = myexplode, shadow = True)

# # # #     return fig


# # # def bar_chart(data):
# # #     fig, ax = plt.subplots(figsize=(12, 8))  # Increased size

# # #     # Data for plotting
# # #     categories = list(data.keys())
# # #     values_min = [int(data[cat]['min'].strip('%')) for cat in categories]
# # #     values_max = [int(data[cat]['max'].strip('%')) for cat in categories]

# # #     x = range(len(categories))

# # #     ax.bar(x, values_min, width=0.4, label='Min Allocation', color='skyblue', align='center')
# # #     ax.bar(x, values_max, width=0.4, label='Max Allocation', color='lightgreen', align='edge')

# # #     ax.set_xticks(x)
# # #     ax.set_xticklabels(categories, rotation=45, ha='right')
# # #     ax.set_xlabel('Investment Categories')
# # #     ax.set_ylabel('Percentage Allocation')
# # #     ax.set_title('Investment Allocation')
# # #     ax.legend()

# # #     plt.tight_layout()
# # #     return fig



# # # import plotly.graph_objects as go
# # # import numpy as np

# # # # import plotly.graph_objects as go
# # # # import streamlit as st
# # # # import numpy as np

# # # # def plot_3d_bar_graph(data):
# # # #     # Initialize a Plotly figure
# # # #     fig = go.Figure()

# # # #     if not data:
# # # #         st.write("No data available to plot.")
# # # #         return

# # # #     # Extracting data for plotting
# # # #     x = []
# # # #     y1 = []  # Min values for plotting
# # # #     y2 = []  # Max values for plotting
# # # #     categories = []

# # # #     for i, (key, value) in enumerate(data.items()):
# # # #         categories.append(key)  # Categories
# # # #         x.append(i)  # X-axis value
# # # #         min_val = int(value.get('min', '0').replace('%', ''))  # Default to 0 if 'min' is missing
# # # #         max_val = int(value.get('max', '0').replace('%', ''))  # Default to 0 if 'max' is missing
# # # #         y1.append(min_val)
# # # #         y2.append(max_val)

# # # #     # Adding bars for Min Compounded Returns
# # # #     fig.add_trace(go.Bar(
# # # #         x=x, y=y1, name='Min Compounded Returns',
# # # #         marker=dict(color='red')
# # # #     ))

# # # #     # Adding bars for Max Compounded Returns
# # # #     fig.add_trace(go.Bar(
# # # #         x=x, y=y2, name='Max Compounded Returns',
# # # #         marker=dict(color='blue')
# # # #     ))

# # # #     # Update layout to match the desired appearance
# # # #     fig.update_layout(
# # # #         barmode='group',
# # # #         xaxis=dict(
# # # #             tickvals=np.arange(len(categories)),
# # # #             ticktext=categories,
# # # #             title='Investment Types'
# # # #         ),
# # # #         yaxis=dict(title='Compounded Returns (%)'),
# # # #         title='Investment Allocation Bar Graph',
# # # #         legend=dict(x=0.1, y=0.9)
# # # #     )

# # # #     st.plotly_chart(fig)


# # # import plotly.graph_objects as go
# # # import streamlit as st
# # # import numpy as np

# # # def plot_3d_bar_graph(data):
# # #     # Initialize a Plotly figure
# # #     fig = go.Figure()

# # #     if not data:
# # #         st.write("No data available to plot.")
# # #         return

# # #     # Extracting data for plotting
# # #     x = []
# # #     y1 = []  # Min values for plotting
# # #     y2 = []  # Max values for plotting
# # #     categories = []

# # #     for i, (key, value) in enumerate(data.items()):
# # #         categories.append(key)  # Categories
# # #         x.append(i)  # X-axis value
# # #         min_val = int(value.get('min', '0').replace('%', ''))  # Default to 0 if 'min' is missing
# # #         max_val = int(value.get('max', '0').replace('%', ''))  # Default to 0 if 'max' is missing
# # #         y1.append(min_val)
# # #         y2.append(max_val)

# # #     # Adding bars for Min Compounded Returns
# # #     for i in range(len(x)):
# # #         fig.add_trace(go.Scatter3d(
# # #             x=[x[i], x[i]],
# # #             y=[0, 0.5],
# # #             z=[0, y1[i]],
# # #             mode='lines',
# # #             line=dict(color='red', width=10),
# # #             name='Min Compounded Returns'
# # #         ))

# # #     # Adding bars for Max Compounded Returns
# # #     for i in range(len(x)):
# # #         fig.add_trace(go.Scatter3d(
# # #             x=[x[i] + 0.5, x[i] + 0.5],
# # #             y=[0, 0.5],
# # #             z=[0, y2[i]],
# # #             mode='lines',
# # #             line=dict(color='blue', width=10),
# # #             name='Max Compounded Returns'
# # #         ))

# # #     # Update layout to match the desired appearance
# # #     fig.update_layout(
# # #         scene=dict(
# # #             xaxis=dict(
# # #                 tickvals=np.arange(len(categories)) + 0.25,
# # #                 ticktext=categories,
# # #                 title='Investment Types'
# # #             ),
# # #             yaxis=dict(title=''),
# # #             zaxis=dict(title='Allocation (%)')
# # #         ),
# # #         title='Investment Allocation 3D Bar Graph',
# # #         legend=dict(x=0.1, y=0.9)
# # #     )

# # #     st.plotly_chart(fig)




# # # # def plot_3d_bar_graph(data):
# # # #     # Initialize a Plotly figure
# # # #     fig = go.Figure()

# # # #     if not data:
# # # #         print("No data available to plot.")
# # # #         return

# # # #     # Extracting data for plotting
# # # #     x = []
# # # #     y1 = []  # Min values for plotting
# # # #     y2 = []  # Max values for plotting
# # # #     categories = []

# # # #     for i, (key, value) in enumerate(data.items()):
# # # #         categories.append(key)  # Categories
# # # #         x.append(i)  # X-axis value
# # # #         min_val = int(value.get('min', '0').replace('%', ''))  # Default to 0 if 'min' is missing
# # # #         max_val = int(value.get('max', '0').replace('%', ''))  # Default to 0 if 'max' is missing
# # # #         y1.append(min_val)
# # # #         y2.append(max_val)

# # # #     # Adding bars for Min Compounded Returns
# # # #     for i in range(len(x)):
# # # #         fig.add_trace(go.Scatter3d(
# # # #             x=[x[i], x[i]],
# # # #             y=[0, 0.5],
# # # #             z=[0, y1[i]],
# # # #             mode='lines',
# # # #             line=dict(color='red', width=10),
# # # #             name='Min Compounded Returns'
# # # #         ))

# # # #     # Adding bars for Max Compounded Returns
# # # #     for i in range(len(x)):
# # # #         fig.add_trace(go.Scatter3d(
# # # #             x=[x[i] + 0.5, x[i] + 0.5],
# # # #             y=[0, 0.5],
# # # #             z=[0, y2[i]],
# # # #             mode='lines',
# # # #             line=dict(color='blue', width=10),
# # # #             name='Max Compounded Returns'
# # # #         ))

# # # #     # Update layout to match the desired appearance
# # # #     fig.update_layout(
# # # #         scene=dict(
# # # #             xaxis=dict(
# # # #                 tickvals=np.arange(len(categories)) + 0.25,
# # # #                 ticktext=categories,
# # # #                 title='Investment Types'
# # # #             ),
# # # #             yaxis=dict(title=''),
# # # #             zaxis=dict(title='Allocation (%)')
# # # #         ),
# # # #         title='Investment Allocation 3D Bar Graph',
# # # #         legend=dict(x=0.1, y=0.9)
# # # #     )

# # # #     fig.show()
# # # #     return fig


 
# # # # def client_form():
# # # #     st.title("Client Details Form")

# # # #     with st.form("client_form"):
# # # #         st.header("Personal Information")
# # # #         client_name = st.text_input("Client Name")
# # # #         co_client_name = st.text_input("Co-Client Name")
# # # #         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
# # # #         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
# # # #         today_date = st.date_input("Today's Date")
        
# # # #         st.header("Financial Information")
# # # #         current_assets = st.text_area("Current Assets (e.g., type and value)")
# # # #         liabilities = st.text_area("Liabilities (e.g., type and amount)")
# # # #         annual_income = st.text_area("Current Annual Income (source and amount)")
# # # #         annual_contributions = st.text_area("Annual Contributions (e.g., retirement savings)")

# # # #         st.header("Insurance Information")
# # # #         life_insurance = st.text_input("Life Insurance (e.g., coverage amount)")
# # # #         disability_insurance = st.text_input("Disability Insurance (e.g., coverage amount)")
# # # #         long_term_care = st.text_input("Long-Term Care Insurance (e.g., coverage amount)")

# # # #         st.header("Estate Planning")
# # # #         will_status = st.radio("Do you have a will?", ["Yes", "No"])
# # # #         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
# # # #         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
# # # #         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

# # # #         # Submit button
# # # #         submitted = st.form_submit_button("Submit")

# # # #         if submitted:
# # # #             # Save form data
# # # #             form_data = {
# # # #                 "Client Name": client_name,
# # # #                 "Co-Client Name": co_client_name,
# # # #                 "Client Age": client_age,
# # # #                 "Co-Client Age": co_client_age,
# # # #                 "Today's Date": str(today_date),
# # # #                 "Current Assets": current_assets,
# # # #                 "Liabilities": liabilities,
# # # #                 "Annual Income": annual_income,
# # # #                 "Annual Contributions": annual_contributions,
# # # #                 "Life Insurance": life_insurance,
# # # #                 "Disability Insurance": disability_insurance,
# # # #                 "Long-Term Care Insurance": long_term_care,
# # # #                 "Will Status": will_status,
# # # #                 "Trust Status": trust_status,
# # # #                 "Power of Attorney": power_of_attorney,
# # # #                 "Healthcare Proxy": healthcare_proxy,
# # # #             }
            
# # # #             # Save to a file or database
# # # #             with open("client_data.txt", "a") as f:
# # # #                 f.write(str(form_data) + "\n")
            
# # # #             st.success("Form submitted successfully!")
# # # #             st.session_state.page = "main"  # Redirect back to main page after form submission


# # # from datetime import date  # Make sure to import the date class

# # # # def client_form():
# # # #     st.title("Client Details Form")

# # # #     with st.form("client_form"):
# # # #         st.header("Personal Information")
# # # #         client_name = st.text_input("Client Name")
# # # #         co_client_name = st.text_input("Co-Client Name")
# # # #         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
# # # #         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
# # # #         today_date = st.date_input("Today's Date")
# # # #         # today_date = st.date_input("Today's Date", value=date.today())

# # # #         st.header("Your Assets and Liabilities")
        
# # # #         # Assets
# # # #         assets = {}
# # # #         # assets['Checking Accounts'] = st.number_input("Checking Accounts", min_value=0.0, step=0.01)
# # # #         assets['Annual Income'] = st.text_input("Annual Income (e.g. , Your Annual Salary Income or other source of income)" )
# # # #         assets['Savings Accounts'] = st.text_input("Savings Accounts (e.g. , Your Annual Savings)" )
# # # #         assets['Retirement Accounts'] = st.text_input("Retirement Accounts (e.g. , Your Retirement Savings)" )
# # # #         assets['Investment Accounts'] = st.text_input("Investment Accounts (e.g. , Your Investment Savings)" )
        
# # # #         # Liabilities
# # # #         liabilities = {}
# # # #         liabilities['Mortgage'] = st.text_input("Mortgage (Mention the balance amount, Interest rate and Monthly Payment)" )
# # # #         liabilities['Home Loans'] = st.text_input("Home Loans (Mention the balance amount, Interest rate and Monthly Payment)")
# # # #         liabilities['Car/Vehicle Loans'] = st.text_input("Car/Vehicle Loans (Mention the balance amount, Interest rate and Monthly Payment)")
# # # #         liabilities['Education Loans'] = st.text_input("Education Loans (Mention the balance amount, Interest rate and Monthly Payment)")
# # # #         liabilities['Credit Card Debt'] = st.text_input("Credit Card Debt")
# # # #         liabilities['Miscellaneous'] = st.text_input("Miscellaneous (Mention the balance amount, Interest rate and Monthly Payment)" )

# # # #         st.header("Your Retirement Goal")
# # # #         retirement_age = st.number_input("At what age do you plan to retire?", min_value=0, max_value=120, value=65, step=1)
# # # #         retirement_income = st.text_input("Desired annual retirement income" )

# # # #         st.header("Your Other Goals")
# # # #         goal_name = st.text_input("Name of the Goal (e.g., Buy a House, Education,Travel,etc)" )
# # # #         goal_amount = st.text_input("Amount needed for the goal")
# # # #         goal_timeframe = st.number_input("Timeframe to achieve the goal (in years)", min_value=0, max_value=100, value=5, step=1)

# # # #         st.header("Insurance Information")
# # # #         life_insurance = st.text_input("Life Insurance (e.g., coverage amount)")
# # # #         disability_insurance = st.text_input("Disability Insurance (e.g., coverage amount)")
# # # #         long_term_care = st.text_input("Long-Term Care Insurance (e.g., coverage amount)")

# # # #         st.header("Estate Planning")
# # # #         will_status = st.radio("Do you have a will?", ["Yes", "No"])
# # # #         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
# # # #         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
# # # #         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

# # # #         # Submit button
# # # #         submitted = st.form_submit_button("Submit")

# # # #         if submitted:
# # # #             # Save form data
# # # #             form_data = {
# # # #                 "Client Name": client_name,
# # # #                 "Co-Client Name": co_client_name,
# # # #                 "Client Age": client_age,
# # # #                 "Co-Client Age": co_client_age,
# # # #                 "Today's Date": str(today_date),
# # # #                 "Assets": assets,
# # # #                 "Liabilities": liabilities,
# # # #                 "Retirement Age": retirement_age,
# # # #                 "Desired Retirement Income": retirement_income,
# # # #                 "Goal Name": goal_name,
# # # #                 "Goal Amount": goal_amount,
# # # #                 "Goal Timeframe": goal_timeframe,
# # # #                 "Life Insurance": life_insurance,
# # # #                 "Disability Insurance": disability_insurance,
# # # #                 "Long-Term Care Insurance": long_term_care,
# # # #                 "Will Status": will_status,
# # # #                 "Trust Status": trust_status,
# # # #                 "Power of Attorney": power_of_attorney,
# # # #                 "Healthcare Proxy": healthcare_proxy,
# # # #             }
            
# # # #             # Save to a file or database
# # # #             with open("client_data.txt", "a") as f:
# # # #                 f.write(str(form_data) + "\n")
            
# # # #             st.success("Form submitted successfully!")
# # # #             st.session_state.page = "main"  # Redirect back to main page after form submission

# # # # Function to parse financial data from the text
# # # import re

# # # def parse_financial_data(text_content):
# # #     assets = []
# # #     liabilities = []

# # #     # Define regex patterns to capture text following headings
# # #     asset_pattern = re.compile(r"MY ASSETS:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)
# # #     liability_pattern = re.compile(r"LIABILITIES:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)

# # #     # Extract assets
# # #     asset_matches = asset_pattern.findall(text_content)
# # #     if asset_matches:
# # #         asset_text = asset_matches[0]
# # #         # Further processing to extract individual asset values if they are detailed
# # #         asset_lines = asset_text.split('\n')
# # #         for line in asset_lines:
# # #             match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
# # #             if match:
# # #                 asset_value = float(match.group().replace(",", ""))
# # #                 assets.append(asset_value)

# # #     # Extract liabilities
# # #     liability_matches = liability_pattern.findall(text_content)
# # #     if liability_matches:
# # #         liability_text = liability_matches[0]
# # #         # Further processing to extract individual liability values if they are detailed
# # #         liability_lines = liability_text.split('\n')
# # #         for line in liability_lines:
# # #             match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
# # #             if match:
# # #                 liability_value = float(match.group().replace(",", ""))
# # #                 liabilities.append(liability_value)

# # #     print("Assets Found:", assets)
# # #     print("Liabilities Found:", liabilities)

# # #     return assets, liabilities



# # # # Function to extract numerical values from a text input
# # # def extract_numeric(value):
# # #     try:
# # #         return float(re.sub(r'[^\d.]', '', value))  # Remove non-numeric characters and convert to float
# # #     except ValueError:
# # #         return 0

# # # # Function to plot pie chart for assets and liabilities
# # # # def plot_assets_liabilities_pie_chart(assets, liabilities):
# # # #     data = {}

# # # #     for key, value in assets.items():
# # # #         numeric_value = extract_numeric(value)
# # # #         if numeric_value > 0:
# # # #             data[key] = numeric_value

# # # #     for key, value in liabilities.items():
# # # #         numeric_value = extract_numeric(value)
# # # #         if numeric_value > 0:
# # # #             data[key] = numeric_value

# # # #     if not data:
# # # #         st.warning("No valid financial data available to plot.")
# # # #         return

# # # #     labels = list(data.keys())
# # # #     values = list(data.values())

# # # #     fig, ax = plt.subplots()
# # # #     ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=140, colors=plt.cm.Paired.colors)
# # # #     ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

# # # #     plt.title("Distribution of Assets and Liabilities")
# # # #     st.pyplot(fig)

# # # # def plot_assets_and_liabilities_pie_chart(assets, liabilities): default plot 
# # # #     # total_assets = sum(assets.values())
# # # #     # total_liabilities = sum(liabilities.values())

# # # #     # labels = ['Assets', 'Liabilities']
# # # #     # sizes = [total_assets, total_liabilities]

# # # #     # fig, ax = plt.subplots(figsize=(8, 8))
# # # #     # ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=['green', 'red'])
# # # #     # ax.set_title('Assets vs Liabilities Distribution')

# # # #     total_assets = sum(assets)
# # # #     total_liabilities = sum(liabilities)

# # # #     if total_assets == 0 and total_liabilities == 0:
# # # #         print("No data to plot.")
# # # #         return

# # # #     labels = ['Assets', 'Liabilities']
# # # #     sizes = [total_assets, total_liabilities]

# # # #     fig, ax = plt.subplots(figsize=(8, 8))
# # # #     ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=['green', 'red'])
# # # #     ax.set_title('Assets vs Liabilities Distribution')

# # # #     # plt.show()

# # # #     st.pyplot(fig)


# # # def plot_assets_liabilities_pie_chart(assets, liabilities):
# # #     labels = list(assets.keys()) + list(liabilities.keys())
# # #     values = [float(value) for value in assets.values()] + [float(value) for value in liabilities.values()]

# # #     fig, ax = plt.subplots()
# # #     ax.pie(values, labels=labels, autopct='%1.1f%%', colors=plt.cm.Paired.colors)
# # #     ax.set_title('Distribution of Assets and Liabilities')
    
# # #     # Create a legend to avoid overlapping labels
# # #     ax.legend(labels, loc="best", bbox_to_anchor=(1, 0, 0.5, 1))
# # #     plt.axis('equal')
    
# # #     st.pyplot(fig)

# # # def save_data_to_file(form_data):
# # #     file_path = 'client_data.txt'
# # #     with open(file_path, 'a') as file:
# # #         file.write(str(form_data) + "\n")
# # #     st.success(f"Form data saved to {file_path}")

# # # def client_form():
# # #     st.title("Client Details Form")

# # #     with st.form("client_form"):
# # #         st.header("Personal Information")
# # #         client_name = st.text_input("Client Name")
# # #         co_client_name = st.text_input("Co-Client Name")
# # #         client_age = st.number_input("Client Age", min_value=0, max_value=120, value=30, step=1)
# # #         co_client_age = st.number_input("Co-Client Age", min_value=0, max_value=120, value=30, step=1)
# # #         today_date = st.date_input("Today's Date")

# # #         st.header("Your Assets and Liabilities")
        
# # #         assets = {
# # #             'Annual Income': st.text_input("Annual Income"),
# # #             'Savings Accounts': st.text_input("Savings Accounts"),
# # #             'Retirement Accounts': st.text_input("Retirement Accounts"),
# # #             'Investment Accounts': st.text_input("Investment Accounts")
# # #         }
        
# # #         liabilities = {
# # #             'Mortgage': st.text_input("Mortgage"),
# # #             'Home Loans': st.text_input("Home Loans"),
# # #             'Car/Vehicle Loans': st.text_input("Car/Vehicle Loans"),
# # #             'Education Loans': st.text_input("Education Loans"),
# # #             'Credit Card Debt': st.text_input("Credit Card Debt"),
# # #             'Miscellaneous': st.text_input("Miscellaneous")
# # #         }

# # #         st.header("Your Retirement Goal")
# # #         retirement_age = st.number_input("At what age do you plan to retire?", min_value=0, max_value=120, value=65, step=1)
# # #         retirement_income = st.text_input("Desired annual retirement income")

# # #         st.header("Your Other Goals")
# # #         goal_name = st.text_input("Name of the Goal")
# # #         goal_amount = st.text_input("Amount needed for the goal")
# # #         goal_timeframe = st.number_input("Timeframe to achieve the goal (in years)", min_value=0, max_value=100, value=5, step=1)

# # #         st.header("Insurance Information")
# # #         life_insurance = st.text_input("Life Insurance")
# # #         disability_insurance = st.text_input("Disability Insurance")
# # #         long_term_care = st.text_input("Long-Term Care Insurance")

# # #         st.header("Estate Planning")
# # #         will_status = st.radio("Do you have a will?", ["Yes", "No"])
# # #         trust_status = st.radio("Do you have any trusts?", ["Yes", "No"])
# # #         power_of_attorney = st.radio("Do you have a Power of Attorney?", ["Yes", "No"])
# # #         healthcare_proxy = st.radio("Do you have a Healthcare Proxy?", ["Yes", "No"])

# # #         submitted = st.form_submit_button("Submit")

# # #         if submitted:
# # #             form_data = {
# # #                 "Client Name": client_name,
# # #                 "Co-Client Name": co_client_name,
# # #                 "Client Age": client_age,
# # #                 "Co-Client Age": co_client_age,
# # #                 "Today's Date": str(today_date),
# # #                 "Assets": assets,
# # #                 "Liabilities": liabilities,
# # #                 "Retirement Age": retirement_age,
# # #                 "Desired Retirement Income": retirement_income,
# # #                 "Goal Name": goal_name,
# # #                 "Goal Amount": goal_amount,
# # #                 "Goal Timeframe": goal_timeframe,
# # #                 "Life Insurance": life_insurance,
# # #                 "Disability Insurance": disability_insurance,
# # #                 "Long-Term Care Insurance": long_term_care,
# # #                 "Will Status": will_status,
# # #                 "Trust Status": trust_status,
# # #                 "Power of Attorney": power_of_attorney,
# # #                 "Healthcare Proxy": healthcare_proxy,
# # #             }

# # #             save_data_to_file(form_data)
            
# # #             # Plot the pie chart
# # #             st.subheader("Assets and Liabilities Breakdown")
# # #             plot_assets_liabilities_pie_chart(assets, liabilities)


# # # def main():
# # #     st.title("Wealth Advisor Chatbot")

# # #     # Step 1: Choose between Investment Suggestions or Stock Analysis
# # #     task_choice = st.radio("Select the task you want to perform:", ["Investment Suggestions", "Stock Analysis"])

# # #     # Step 2: Choose between New Client or Existing Client (for Investment Suggestions)
# # #     if task_choice == "Investment Suggestions":
# # #         client_type = st.radio("Is this for a new client or an existing client?", ["New Client", "Existing Client"])

# # #         if client_type == "New Client":
# # #             # Button to redirect to the form page
# # #             if st.button("Fill in the Client Details"):
# # #                 st.session_state.page = "form"
# # #                 # st.query_params(page="form")  #set_query_params(page="form")
# # #                 # Display the pie chart if assets and liabilities data is available

# # #             if 'assets' in st.session_state and 'liabilities' in st.session_state:
# # #                 st.subheader("Assets and Liabilities Breakdown")
# # #                 plot_assets_liabilities_pie_chart(st.session_state.assets, st.session_state.liabilities)
# # #             else:
# # #                 st.info("Please fill in the client details to view the assets and liabilities breakdown.")

# # #             uploaded_file = st.file_uploader("Upload the client's document", type=["docx", "pdf"])

# # #             if uploaded_file is not None:
# # #                 st.write("Extracting text from the document...")
# # #                 document_data = asyncio.run(process_document(uploaded_file))

# # #                 if document_data:
# # #                     if isinstance(document_data, tuple) and len(document_data) == 2:
# # #                         extracted_text, tables_content = document_data
# # #                         st.write("Text extracted from the document")

# # #                         # Print extracted text for debugging
# # #                         print("Extracted Text Content:", extracted_text)

# # #                         # Parse extracted data to get assets and liabilities
# # #                         # assets, liabilities = parse_financial_data(extracted_text)

# # #                         # if assets or liabilities:
# # #                         #     plot_pie_chart(assets, liabilities)
# # #                         # else:
# # #                         #     # Extracted data from the word file
# # #                         #     st.write("Financial Data Plotted :")
# # #                         #     assets = [100000, 150000, 12000]  # Cash/bank accounts, Home, Other (e.g. car, boat, art, etc.)
# # #                         #     liabilities = [200000, 400, 15000]  # Mortgage(s), Credit Card(s), Other loans (car, education, etc.)
# # #                         #     # plot_assets_and_liabilities_pie_chart(assets, liabilities)
# # #                         #     plot_assets_liabilities_pie_chart()
# # #                         #     # st.write("No financial data found to plot.")
# # #                     else:
# # #                         st.write("Unexpected data format returned from document processing.")



# # #                 # extracted_text = asyncio.run(process_document(uploaded_file))
# # #                 # extracted_text, tables_content = asyncio.run(process_document(uploaded_file))
# # #                 # # st.write("Text extracted from the document")
# # #                 # if extracted_text:
# # #                 #     st.write("Text extracted from the document")
                    
# # #                 #     # Parse extracted data to get assets and liabilities
# # #                 #     assets, liabilities = parse_financial_data(extracted_text)
                    
# # #                 #     # Check if there is any data to plot
# # #                 #     if assets or liabilities:
# # #                 #         plot_pie_chart(assets, liabilities)
# # #                 #     else:
# # #                 #         st.write("No financial data found to plot.")
# # #                 # else:
# # #                 #     st.write("Failed to extract data from the document.")

# # #                 # Ask for investment personality
# # #                 investment_personality = st.selectbox(
# # #                     "Select the investment personality of the client:",
# # #                     ("Conservative Investor", "Moderate Investor", "Aggressive Investor")
# # #                 )

# # #                 # Step 4: Generate investment suggestions
# # #                 if st.button("Generate Investment Suggestions"):
# # #                     st.write("Generating investment suggestions...")
# # #                     # Assuming generate_investment_suggestions is an async function
                    
# # #                     suggestions = asyncio.run(generate_investment_suggestions(investment_personality, extracted_text))
# # #                     st.write(suggestions)

# # #                     # # Generate infographics
# # #                     # generate_infographics(suggestions, investment_personality)

# # #                     data_extracted = extract_numerical_data(suggestions)

# # #                     # Streamlit app
# # #                     st.title('Investment Allocation Infographics')
# # #                     # st.write('### Extracted Investment Data')
# # #                     # st.json(data_extracted)

# # #                     st.write('## Investment Allocation Charts')
# # #                     fig = plot_investment_allocations(data_extracted)
# # #                     st.pyplot(fig)

# # #                     st.write('## Pie Chart of Investment Allocation')

# # #                     # plot_pie_chart(data_extracted)
# # #                     fig = plot_pie_chart(data_extracted)
# # #                     st.pyplot(fig)

# # #                     st.write('## Bar Chart of Compounded Returns')
# # #                     fig = bar_chart(data_extracted['Growth-Oriented Investments'])
# # #                     st.pyplot(fig)

# # #                     plot_3d_bar_graph(data_extracted)
# # #                     # fig = plot_3d_bar_graph(data_extracted)
# # #                     # st.pyplot(fig)

# # #         elif client_type == "Existing Client":
# # #             st.write("Fetching existing client data...")

# # #     elif task_choice == "Stock Analysis":
# # #         st.write("Performing stock analysis...")


# # # if __name__ == "__main__":
# # #     if 'page' not in st.session_state:
# # #         st.session_state.page = "main"
    
# # #     if st.session_state.page == "form":
# # #         client_form()
# # #     else:
# # #         main()

# # # # if __name__ == "__main__":
# # # #     main()




# # # # # infographics oplotting maanger side code |


# # # # import streamlit as st
# # # # import pandas as pd
# # # # import matplotlib.pyplot as plt

# # # # import os
# # # # import filetype
# # # # import docx
# # # # import PyPDF2
# # # # import re
# # # # from dotenv import load_dotenv
# # # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # # from langchain_community.vectorstores import Chroma
# # # # from langchain_community.docstore.in_memory import InMemoryDocstore
# # # # from langchain_community.vectorstores import FAISS
# # # # from langchain_community.document_loaders import Docx2txtLoader
# # # # from langchain_core.prompts import ChatPromptTemplate
# # # # from langchain.chains import create_retrieval_chain
# # # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # # from langchain.memory import ConversationSummaryMemory
# # # # import asyncio
# # # # import numpy as np
# # # # import json

# # # # import google.generativeai as genai
# # # # import pathlib
# # # # import logging
# # # # import sys
# # # # import io
# # # # import matplotlib.pyplot as plt
# # # # import seaborn as sns
# # # # # Import things that are needed generically
# # # # from langchain.pydantic_v1 import BaseModel, Field
# # # # from langchain.tools import BaseTool, StructuredTool, tool
# # # # # Define functions to generate investment suggestions :

# # # # load_dotenv()
# # # # GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

# # # # def markdown_to_text(md):
# # # #     # Simple conversion for markdown to plain text
# # # #     md = md.replace('**', '')
# # # #     md = md.replace('*', '')
# # # #     md = md.replace('_', '')
# # # #     md = md.replace('#', '')
# # # #     md = md.replace('`', '')
# # # #     return md.strip()

# # # # # Load the Vector DataBase :
# # # # async def load_vector_db(file_path):
# # # #     try:
# # # #         print("Loading vector database...")
# # # #         loader = Docx2txtLoader(file_path)
# # # #         documents = loader.load()
# # # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # # #         text_chunks = text_splitter.split_documents(documents)
# # # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # # #         # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
# # # #         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
# # # #         # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

# # # #         # vector_store = FAISS(
# # # #         #     embedding_function=embeddings,
# # # #         #     index=index,
# # # #         #     docstore=InMemoryDocstore(),
# # # #         #     index_to_docstore_id={},
# # # #         # )
        
# # # #         print("Vector database loaded successfully.") 
# # # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # # #     except Exception as e:
# # # #         print(f"Error loading vector database: {e}")
# # # #         return None


# # # # investment_personality = "Moderate Investor"
# # # # async def make_retrieval_chain(retriever):
# # # #     """
# # # #     Create a retrieval chain using the provided retriever.

# # # #     Args:
# # # #         retriever (RetrievalQA): A retriever object.

# # # #     Returns:
# # # #         RetrievalQA: A retrieval chain object.
# # # #     """
# # # #     try:
# # # #         global investment_personality #,summary
# # # #         llm = ChatGoogleGenerativeAI(
# # # #             #model="gemini-pro",
# # # #             model = "gemini-1.5-flash",
# # # #             temperature=0.7,
# # # #             top_p=0.85,
# # # #             google_api_key=GOOGLE_API_KEY
# # # #         )
# # # #                                                     # \n + summary 
# # # #         prompt_template = investment_personality +  "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # #                 Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality.
# # # #                 Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # #                 should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
# # # #                 Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # #                 investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
# # # #                 Also give the user minimum and maximum expected growth in dollars for the time horizon .
# # # #                 Also explain the user why you are giving them that particular investment suggestion.
# # # #                 Here's an example for the required Output Format :

# # # #                 Investment Suggestions for a Moderate Investor(This is for a Moderate Investor but you need to generate for any investor)

# # # #                 Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

# # # #                 Investment Allocation: (remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

# # # #                 Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
# # # #                 How to Invest: Diversify across various asset classes like:  (Give allocations % as well)
# # # #                 Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
# # # #                 ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
# # # #                 Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
# # # #                 Consider investing in blue-chip companies or growth sectors like technology. 
# # # #                 Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


# # # #                 Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
# # # #                 Percentage Allocation: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

# # # #                 Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
# # # #                 How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
# # # #                 Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
# # # #                 Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
# # # #                 Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
# # # #                 Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
# # # #                 Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
# # # #                 Percentage Allocation: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.

# # # #                 Time Horizon and Expected Returns:

# # # #                 Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
# # # #                 Minimum Expected Annual Return: 4% - 6% Maximum Expected Annual Return: 8% - 10% Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, 
# # # #                 a 10,000 investment could growto approximately 17,908 in 10 years. Minimum Expected Growth in Dollars: 
                
# # # #                 4,000−6,000 (over 10 years) Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
# # # #                 Rationale for Investment Suggestions:

# # # #                 This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
# # # #                 Important Considerations:

# # # #                 Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

# # # #                 Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

# # # #                 <context>
# # # #                 {context}
# # # #                 </context>
# # # #                 Question: {input}"""

# # # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
# # # #         combine_docs_chain = None  

# # # #         if retriever is not None :  
# # # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # # #             print(retriever_chain)
# # # #             return retriever_chain
# # # #         else:
# # # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # # #             return None

# # # #     except Exception as e:
# # # #         print(f"Error in creating chain: {e}")
# # # #         return None

# # # # import streamlit as st
# # # # import json
# # # # import matplotlib.pyplot as plt
# # # # import io



# # # # # Create InfoGraphics :


# # # # # import re
# # # # # import matplotlib.pyplot as plt
# # # # # import seaborn as sns
# # # # # from collections import defaultdict

# # # # # def extract_numerical_data(response):
# # # # #     # Extract percentage ranges, dollar values, and other numbers from the text
# # # # #     pattern = re.compile(r'(\b[A-Za-z\s]+\b):\s*([\d.,]+%-[\d.,]+%|[\d.,]+%-[\d.,]+|\d+-\d+|\d+%|\d+|[$]\d+,?\d*)')
# # # # #     data = defaultdict(list)
    
# # # # #     for match in pattern.findall(response):
# # # # #         category, values = match
# # # # #         values = values.replace('$', '').replace(',', '')
        
# # # # #         if '-' in values:
# # # # #             if '%' in values:
# # # # #                 min_val, max_val = map(lambda x: float(x.replace('%', '')), values.split('-'))
# # # # #             else:
# # # # #                 min_val, max_val = map(float, values.split('-'))
# # # # #         else:
# # # # #             min_val = max_val = float(values.replace('%', ''))

# # # # #         data[category.strip()].append((min_val, max_val))
    
# # # # #     return data

# # # # # def create_bar_chart(data, title):
# # # # #     categories = list(data.keys())
# # # # #     min_values = [v[0][0] for v in data.values()]
# # # # #     max_values = [v[0][1] for v in data.values()]

# # # # #     plt.figure(figsize=(10, 6))
# # # # #     plt.barh(categories, max_values, color='lightblue', edgecolor='blue', label='Max Value')
# # # # #     plt.barh(categories, min_values, color='lightgreen', edgecolor='green', label='Min Value')

# # # # #     plt.xlabel('Values')
# # # # #     plt.ylabel('Categories')
# # # # #     plt.title(title)
# # # # #     plt.legend()
# # # # #     plt.show()

# # # # # def create_pie_chart(data, title):
# # # # #     labels = list(data.keys())
# # # # #     sizes = [sum(v[0]) / 2 for v in data.values()]

# # # # #     plt.figure(figsize=(8, 8))
# # # # #     plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140, colors=sns.color_palette('Set3', len(labels)))
# # # # #     plt.axis('equal')
# # # # #     plt.title(title)
# # # # #     plt.show()

# # # # # def create_expected_return_chart(data, title):
# # # # #     categories = list(data.keys())
# # # # #     min_values = [v[0][0] for v in data.values()]
# # # # #     max_values = [v[0][1] for v in data.values()]

# # # # #     plt.figure(figsize=(10, 6))
# # # # #     plt.bar(categories, max_values, color='coral', edgecolor='red', label='Max Value')
# # # # #     plt.bar(categories, min_values, color='lightcoral', edgecolor='darkred', label='Min Value')
    
# # # # #     plt.ylabel('Values')
# # # # #     plt.title(title)
# # # # #     plt.legend()
# # # # #     plt.show()

# # # # # def generate_infographics(response, title_prefix="Investment Strategy"):
# # # # #     # Extract the data
# # # # #     data = extract_numerical_data(response)
    
# # # # #     # Generate the charts
# # # # #     create_bar_chart(data, f"{title_prefix} - Asset Allocation")
# # # # #     create_pie_chart(data, f"{title_prefix} - Asset Distribution")
# # # # #     create_expected_return_chart(data, f"{title_prefix} - Expected Returns")






# # # # async def process_document(file_path):
# # # #     try:
# # # #         print("Processing the document")
# # # #         file_type = filetype.guess(file_path)
# # # #         if file_type is not None:
# # # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # #                 # return extract_text_from_word(file_path)
# # # #                 return extract_text_and_tables_from_word(file_path)
# # # #             elif file_type.mime == "application/pdf":
# # # #                 return extract_text_from_pdf(file_path)
# # # #         return None
# # # #     except Exception as e:
# # # #         print(f"Error processing document: {e}")
# # # #         return None


# # # # async def extract_text_from_pdf(pdf_file_path):
# # # #     try:
# # # #         print("Processing pdf file")
# # # #         with open(pdf_file_path, "rb") as pdf_file:
# # # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # # #             text_content = []
# # # #             for page_num in range(pdf_reader.numPages):
# # # #                 page = pdf_reader.getPage(page_num)
# # # #                 text_content.append(page.extract_text())
# # # #             return "\n".join(text_content)
# # # #     except Exception as e:
# # # #         print(f"Error extracting text from PDF: {e}")
# # # #         return None



# # # # async def extract_text_and_tables_from_word(docx_file_path):
# # # #     """
# # # #     Extracts text and tables from a Word document (.docx).

# # # #     Args:
# # # #         docx_file_path (str): Path to the Word document file.

# # # #     Returns:
# # # #         tuple: Extracted text content and tables from the document.
# # # #     """
# # # #     try:
# # # #         print("Extracting text and tables from word file")
# # # #         doc = docx.Document(docx_file_path)
# # # #         text_content = []
# # # #         tables_content = []

# # # #         for para in doc.paragraphs:
# # # #             text_content.append(para.text)

# # # #         for table in doc.tables:
# # # #             table_data = []
# # # #             for row in table.rows:
# # # #                 row_data = []
# # # #                 for cell in row.cells:
# # # #                     row_data.append(cell.text.strip())
# # # #                 table_data.append(row_data)
# # # #             tables_content.append(table_data)
# # # #         print("Extracted text from word file")
# # # #         return "\n".join(text_content), tables_content
# # # #     except Exception as e:
# # # #         print(f"Error extracting text and tables from Word document: {e}")
# # # #         return None, None

# # # # async def validate_document_content(text, tables):
# # # #     """
# # # #     Validates the content of the document.

# # # #     Args:
# # # #         text (str): Extracted text content from the document.
# # # #         tables (list): Extracted tables content from the document.

# # # #     Returns:
# # # #         tuple: Client name and validation errors.
# # # #     """
# # # #     errors = []
    
# # # #     # Extract client name
# # # #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# # # #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# # # #     # Define required sections
# # # #     required_sections = [
# # # #         "YOUR RETIREMENT GOAL",
# # # #         "YOUR OTHER MAJOR GOALS",
# # # #         "YOUR ASSETS AND LIABILITIES",
# # # #         "MY LIABILITIES",
# # # #         "YOUR CURRENT ANNUAL INCOME"
# # # #     ]

# # # #     # Check for the presence of required sections
# # # #     for section in required_sections:
# # # #         if section not in text:
# # # #             errors.append(f"* {section} section missing.")
    
# # # #     # Define table field checks
# # # #     table_checks = {
# # # #         "YOUR RETIREMENT GOAL": [
# # # #             r"When do you plan to retire\? \(age or date\)",
# # # #             r"Social Security Benefit \(include expected start date\)",
# # # #             r"Pension Benefit \(include expected start date\)",
# # # #             r"Other Expected Income \(rental, part-time work, etc.\)",
# # # #             r"Estimated Annual Retirement Expense"
# # # #         ],
# # # #         "YOUR OTHER MAJOR GOALS": [
# # # #             r"GOAL", r"COST", r"WHEN"
# # # #         ],
# # # #         "YOUR ASSETS AND LIABILITIES": [
# # # #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# # # #             r"Current Value", r"Annual Contributions"
# # # #         ],
# # # #         "MY LIABILITIES": [
# # # #             r"Balance", r"Interest Rate", r"Monthly Payment"
# # # #         ]
# # # #     }

# # # #     # Validate table content
# # # #     for section, checks in table_checks.items():
# # # #         section_found = False
# # # #         for table in tables:
# # # #             table_text = "\n".join(["\t".join(row) for row in table])
# # # #             if section in table_text:
# # # #                 section_found = True
# # # #                 for check in checks:
# # # #                     if not re.search(check, table_text, re.IGNORECASE):
# # # #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# # # #                 break
# # # #         if not section_found:
# # # #             errors.append(f"* {section} section missing.")

# # # #     return client_name, errors



# # # # async def generate_investment_suggestions(investment_personality, context):
    
# # # #     # retriever = asyncio.run(load_vector_db("uploaded_file"))

# # # #     retriever = await load_vector_db("uploaded_file")
# # # #     chain = await make_retrieval_chain(retriever)

# # # #     # chain = asyncio.run(make_retrieval_chain(retriever))
    
# # # #     if chain is not None:
# # # #         # summary = context
# # # #         # query = summary + "\n" + investment_personality
# # # #         query = str(investment_personality)
# # # #         response = chain.invoke({"input": query})
# # # #         format_response = markdown_to_text(response['answer'])
# # # #         return format_response
# # # #         # st.write(format_response)

# # # #         # handle_graph(response['answer'])

# # # #     else:
# # # #         st.error("Failed to create the retrieval chain. Please upload a valid document.")


# # # # # Generating Infographics : however got st.pyplot warnings:

# # # # # import re
# # # # # import matplotlib.pyplot as plt
# # # # # import seaborn as sns
# # # # # from collections import defaultdict
# # # # # import streamlit as st

# # # # # def extract_numerical_data(response):
# # # # #     pattern = re.compile(r'(\b[A-Za-z\s]+\b):\s*([\d.,]+%-[\d.,]+%|[\d.,]+%-[\d.,]+|\d+-\d+|\d+%|\d+|[$]\d+,?\d*)')
# # # # #     data = defaultdict(list)
    
# # # # #     for match in pattern.findall(response):
# # # # #         category, values = match
# # # # #         values = values.replace('$', '').replace(',', '')
        
# # # # #         if '-' in values:
# # # # #             if '%' in values:
# # # # #                 min_val, max_val = map(lambda x: float(x.replace('%', '')), values.split('-'))
# # # # #             else:
# # # # #                 min_val, max_val = map(float, values.split('-'))
# # # # #         else:
# # # # #             min_val = max_val = float(values.replace('%', ''))

# # # # #         data[category.strip()].append((min_val, max_val))
    
# # # # #     return data

# # # # # def create_bar_chart(data, title):
# # # # #     categories = list(data.keys())
# # # # #     min_values = [v[0][0] for v in data.values()]
# # # # #     max_values = [v[0][1] for v in data.values()]

# # # # #     plt.figure(figsize=(10, 6))
# # # # #     plt.barh(categories, max_values, color='lightblue', edgecolor='blue', label='Max Value')
# # # # #     plt.barh(categories, min_values, color='lightgreen', edgecolor='green', label='Min Value')

# # # # #     plt.xlabel('Values')
# # # # #     plt.ylabel('Categories')
# # # # #     plt.title(title)
# # # # #     plt.legend()
# # # # #     st.pyplot()  # Display the plot in Streamlit

# # # # # def create_pie_chart(data, title):
# # # # #     labels = list(data.keys())
# # # # #     sizes = [sum(v[0]) / 2 for v in data.values()]

# # # # #     plt.figure(figsize=(8, 8))
# # # # #     plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140, colors=sns.color_palette('Set3', len(labels)))
# # # # #     plt.axis('equal')
# # # # #     plt.title(title)
# # # # #     st.pyplot()  # Display the plot in Streamlit

# # # # # def create_expected_return_chart(data, title):
# # # # #     categories = list(data.keys())
# # # # #     min_values = [v[0][0] for v in data.values()]
# # # # #     max_values = [v[0][1] for v in data.values()]

# # # # #     plt.figure(figsize=(10, 6))
# # # # #     plt.bar(categories, max_values, color='coral', edgecolor='red', label='Max Value')
# # # # #     plt.bar(categories, min_values, color='lightcoral', edgecolor='darkred', label='Min Value')
    
# # # # #     plt.ylabel('Values')
# # # # #     plt.title(title)
# # # # #     plt.legend()
# # # # #     st.pyplot()  # Display the plot in Streamlit

# # # # # def generate_infographics(response, title_prefix="Investment Strategy"):
# # # # #     data = extract_numerical_data(response)
    
# # # # #     create_bar_chart(data, f"{title_prefix} - Asset Allocation")
# # # # #     create_pie_chart(data, f"{title_prefix} - Asset Distribution")
# # # # #     create_expected_return_chart(data, f"{title_prefix} - Expected Returns")




# # # # # # Genrating Inforgaphics :
# # # # # # generated unreadable graphs:

# # # # # import re
# # # # # import matplotlib.pyplot as plt
# # # # # import seaborn as sns
# # # # # from collections import defaultdict
# # # # # import streamlit as st

# # # # # def extract_numerical_data(response):
# # # # #     pattern = re.compile(r'(\b[A-Za-z\s]+\b):\s*([\d.,]+%-[\d.,]+%|[\d.,]+%-[\d.,]+|\d+-\d+|\d+%|\d+|[$]\d+,?\d*)')
# # # # #     data = defaultdict(list)
    
# # # # #     for match in pattern.findall(response):
# # # # #         category, values = match
# # # # #         values = values.replace('$', '').replace(',', '')
        
# # # # #         if '-' in values:
# # # # #             if '%' in values:
# # # # #                 min_val, max_val = map(lambda x: float(x.replace('%', '')), values.split('-'))
# # # # #             else:
# # # # #                 min_val, max_val = map(float, values.split('-'))
# # # # #         else:
# # # # #             min_val = max_val = float(values.replace('%', ''))

# # # # #         data[category.strip()].append((min_val, max_val))
    
# # # # #     return data

# # # # # def create_bar_chart(data, title):
# # # # #     categories = list(data.keys())
# # # # #     min_values = [v[0][0] for v in data.values()]
# # # # #     max_values = [v[0][1] for v in data.values()]

# # # # #     fig, ax = plt.subplots(figsize=(10, 6))
# # # # #     ax.barh(categories, max_values, color='lightblue', edgecolor='blue', label='Max Value')
# # # # #     ax.barh(categories, min_values, color='lightgreen', edgecolor='green', label='Min Value')

# # # # #     ax.set_xlabel('Values')
# # # # #     ax.set_ylabel('Categories')
# # # # #     ax.set_title(title)
# # # # #     ax.legend()
# # # # #     st.pyplot(fig)  # Pass the figure to st.pyplot()

# # # # # def create_pie_chart(data, title):
# # # # #     labels = list(data.keys())
# # # # #     sizes = [sum(v[0]) / 2 for v in data.values()]

# # # # #     fig, ax = plt.subplots(figsize=(8, 8))
# # # # #     ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140, colors=sns.color_palette('Set3', len(labels)))
# # # # #     ax.axis('equal')
# # # # #     ax.set_title(title)
# # # # #     st.pyplot(fig)  # Pass the figure to st.pyplot()

# # # # # def create_expected_return_chart(data, title):
# # # # #     categories = list(data.keys())
# # # # #     min_values = [v[0][0] for v in data.values()]
# # # # #     max_values = [v[0][1] for v in data.values()]

# # # # #     fig, ax = plt.subplots(figsize=(10, 6))
# # # # #     width = 0.35  # Width of the bars

# # # # #     ax.bar(categories, max_values, color='coral', edgecolor='red', label='Max Value')
# # # # #     ax.bar(categories, min_values, color='lightcoral', edgecolor='darkred', label='Min Value')
    
# # # # #     ax.set_ylabel('Values')
# # # # #     ax.set_title(title)
# # # # #     ax.legend()
# # # # #     st.pyplot(fig)  # Pass the figure to st.pyplot()

# # # # # def generate_infographics(response, title_prefix="Investment Strategy"):
# # # # #     data = extract_numerical_data(response)
    
# # # # #     create_bar_chart(data, f"{title_prefix} - Asset Allocation")
# # # # #     create_pie_chart(data, f"{title_prefix} - Asset Distribution")
# # # # #     create_expected_return_chart(data, f"{title_prefix} - Expected Returns")


# # # # import re
# # # # import matplotlib.pyplot as plt
# # # # import seaborn as sns
# # # # import streamlit as st

# # # # # Function to extract numerical data from the response text
# # # # # def extract_numerical_data(response, investment_personality):
# # # # #     investment_types = [investment_personality]  # Wrap the personality in a list
# # # # #     allocation = []
# # # # #     returns = []
# # # # #     time_horizon = []

# # # # #     # Regex patterns to match relevant data
# # # # #     allocation_pattern = re.compile(r'Percentage Allocation:.*?(\d+)%', re.IGNORECASE)
# # # # #     return_pattern = re.compile(r'Expected Annual Return:.*?(\d+)%', re.IGNORECASE)
# # # # #     time_pattern = re.compile(r'Time Horizon:.*?(\d+) years', re.IGNORECASE)

# # # # #     # Extracting allocation percentages
# # # # #     allocation_matches = allocation_pattern.findall(response)
# # # # #     allocation = [int(a) for a in allocation_matches]

# # # # #     # Extracting expected returns
# # # # #     return_matches = return_pattern.findall(response)
# # # # #     returns = [int(r) for r in return_matches]

# # # # #     # Extracting time horizon
# # # # #     time_matches = time_pattern.findall(response)
# # # # #     time_horizon = [int(t) for t in time_matches]

# # # # #     return investment_types, allocation, returns, time_horizon

# # # # # Function to create a pie chart for asset allocation
# # # # # def create_pie_chart(investment_types, allocation):
# # # # #     if not investment_types or not allocation:
# # # # #         st.warning("Not enough data for Asset Allocation Pie Chart.")
# # # # #         return

# # # # #     plt.figure(figsize=(8, 6))
# # # # #     plt.pie(allocation, labels=investment_types, autopct='%1.1f%%', startangle=140, colors=sns.color_palette('Set3', len(investment_types)))
# # # # #     plt.title('Asset Allocation')
# # # # #     st.pyplot(plt)
# # # # #     plt.clf()

# # # # # def create_pie_chart(investment_types, allocation):
# # # # #     # Check if the lengths of the lists match
# # # # #     if not investment_types or not allocation:
# # # # #         st.warning("Not enough data for Asset Allocation Pie Chart.")
# # # # #         return

# # # # #     if len(investment_types) != len(allocation):
# # # # #         st.warning("Mismatch between investment types and allocation data.")
# # # # #         st.write(f"Investment types: {investment_types}")
# # # # #         st.write(f"Allocation: {allocation}")
# # # # #         return

# # # # #     plt.figure(figsize=(8, 6))
# # # # #     plt.pie(allocation, labels=investment_types, autopct='%1.1f%%', startangle=140, colors=sns.color_palette('Set3', len(investment_types)))
# # # # #     plt.title('Asset Allocation')
# # # # #     st.pyplot(plt)
# # # # #     plt.clf()

# # # # # # Function to create a bar chart for expected returns
# # # # # def create_compounded_return_chart(returns, time_horizon):
# # # # #     if not returns or not time_horizon:
# # # # #         st.warning("Not enough data for Compounded Returns Bar Chart.")
# # # # #         return

# # # # #     # Ensure we have at least two time horizons for the plot (e.g., 5 years and 10 years)
# # # # #     if len(time_horizon) < 2:
# # # # #         time_horizon.append(time_horizon[0] * 2)  # Add a second time horizon by doubling the first

# # # # #     years = ['5 Years', '10 Years']
# # # # #     plt.figure(figsize=(10, 6))

# # # # #     for i, rate in enumerate(returns):
# # # # #         compounded_5_years = ((1 + rate / 100) ** 5 - 1) * 100
# # # # #         compounded_10_years = ((1 + rate / 100) ** 10 - 1) * 100
# # # # #         plt.bar(years, [compounded_5_years, compounded_10_years], color=['blue', 'green'][i % 2], alpha=0.7, label=f'Return Rate {rate}%')

# # # # #     plt.xlabel('Time Horizon')
# # # # #     plt.ylabel('Returns (%)')
# # # # #     plt.title('Compounded Returns over Different Time Horizons')
# # # # #     plt.legend()
# # # # #     st.pyplot(plt)
# # # # #     plt.clf()

# # # # # # Main function to generate infographics
# # # # # def generate_infographics(response_text, investment_personality):
# # # # #     investment_types, allocation, returns, time_horizon = extract_numerical_data(response_text, investment_personality)

# # # # #     st.write(f"Investment Type: {investment_types}")
# # # # #     st.write(f"Allocation: {allocation}")
# # # # #     st.write(f"Returns: {returns}")
# # # # #     st.write(f"Time Horizon: {time_horizon}")

# # # # #     create_pie_chart(investment_types, allocation)
# # # # #     create_compounded_return_chart(returns, time_horizon)

# # # # import re
# # # # import matplotlib.pyplot as plt
# # # # import seaborn as sns
# # # # import streamlit as st

# # # # # Function to extract numerical data from the response text
# # # # # def extract_numerical_data(response):
# # # # #     allocation_pattern = re.compile(r'(\b[A-Za-z\s]+\b):\s*(\d+)%', re.IGNORECASE)
# # # # #     return_pattern = re.compile(r'Expected Annual Return:.*?(\d+)%', re.IGNORECASE)
# # # # #     time_pattern = re.compile(r'Time Horizon:.*?(\d+) years', re.IGNORECASE)

# # # # #     allocations = {}
# # # # #     returns = []
# # # # #     time_horizon = []

# # # # #     # Extracting allocations
# # # # #     allocation_matches = allocation_pattern.findall(response)
# # # # #     for match in allocation_matches:
# # # # #         investment_type, percentage = match
# # # # #         allocations[investment_type.strip()] = int(percentage)

# # # # #     # Extracting returns
# # # # #     return_matches = return_pattern.findall(response)
# # # # #     returns = [int(r) for r in return_matches]

# # # # #     # Extracting time horizon
# # # # #     time_matches = time_pattern.findall(response)
# # # # #     time_horizon = [int(t) for t in time_matches]

# # # # #     return allocations, returns, time_horizon

# # # # # # Function to create a pie chart for asset allocation
# # # # # def create_pie_chart(allocations):
# # # # #     if not allocations:
# # # # #         st.warning("Not enough data for Asset Allocation Pie Chart.")
# # # # #         return

# # # # #     labels = list(allocations.keys())
# # # # #     sizes = list(allocations.values())

# # # # #     plt.figure(figsize=(8, 6))
# # # # #     plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140, colors=sns.color_palette('Set3', len(labels)))
# # # # #     plt.title('Asset Allocation')
# # # # #     st.pyplot(plt)
# # # # #     plt.clf()

# # # # # # Function to create a bar chart for expected returns
# # # # # def create_compounded_return_chart(returns, time_horizon):
# # # # #     if not returns or not time_horizon:
# # # # #         st.warning("Not enough data for Compounded Returns Bar Chart.")
# # # # #         return

# # # # #     years = ['5 Years', '10 Years']
# # # # #     plt.figure(figsize=(10, 6))

# # # # #     for i, rate in enumerate(returns):
# # # # #         plt.bar(years, [rate * 5, rate * 10], color=['blue', 'green'][i % 2], alpha=0.7, label=f'Return Rate {rate}%')

# # # # #     plt.xlabel('Time Horizon')
# # # # #     plt.ylabel('Returns')
# # # # #     plt.title('Compounded Returns over Different Time Horizons')
# # # # #     plt.legend()
# # # # #     st.pyplot(plt)
# # # # #     plt.clf()

# # # # # # Main function to generate infographics
# # # # # def generate_infographics(response_text):
# # # # #     allocations, returns, time_horizon = extract_numerical_data(response_text)

# # # # #     st.write(f"Allocations: {allocations}")
# # # # #     st.write(f"Returns: {returns}")
# # # # #     st.write(f"Time Horizon: {time_horizon}")

# # # # #     create_pie_chart(allocations)
# # # # #     create_compounded_return_chart(returns, time_horizon)


# # # # # Genrate Infographics : static code :

# # # # # import re
# # # # # import matplotlib.pyplot as plt
# # # # # import seaborn as sns
# # # # # import streamlit as st

# # # # # # Function to extract numerical data from the response text
# # # # # def extract_numerical_data(response, investment_personality):
# # # # #     # Define default allocation percentages for each investment type based on personality
# # # # #     default_allocations = {
# # # # #         'Conservative Investor': {
# # # # #             'Growth-Oriented Investments': {'min': 20, 'max': 30},
# # # # #             'Conservative Investments': {'min': 70, 'max': 80},
# # # # #             'Growth Types': ['stocks', 'short term bonds', 'low-volatility ETFs'],
# # # # #             'Conservative Types': ['bonds', 'ETFs', 'mutual funds', 'real estate (REITs)']
# # # # #         },
# # # # #         'Moderate Investor': {
# # # # #             'Growth-Oriented Investments': {'min': 50, 'max': 60},
# # # # #             'Conservative Investments': {'min': 40, 'max': 50},
# # # # #             'Growth Types': ['index funds', 'mutual funds', 'ETFs', 'stocks'],
# # # # #             'Conservative Types': ['cash', 'REITs']
# # # # #         },
# # # # #         'Aggressive Investor': {
# # # # #             'Growth-Oriented Investments': {'min': 70, 'max': 80},
# # # # #             'Conservative Investments': {'min': 20, 'max': 30},
# # # # #             'Growth Types': ['stocks', 'ETFs', 'mutual funds', 'cryptocurrency'],
# # # # #             'Conservative Types': ['bonds', 'real estate']
# # # # #         }
# # # # #     }

# # # # #     allocations = {'Growth-Oriented Investments': {}, 'Conservative Investments': {}}

# # # # #     # Attempt to dynamically extract allocation percentages from the response
# # # # #     try:
# # # # #         # Search for allocation percentages in the response text
# # # # #         # growth_pattern = r"Growth-Oriented Investments: (\d+)%"
# # # # #         growth_pattern = re.compile("Growth-Oriented Investments: (\d+)%",re.IGNORECASE)
# # # # #         # conservative_pattern = r"Conservative Investments: (\d+)%"
# # # # #         conservative_pattern = re.compile("Conservative Investments: (\d+)%",re.IGNORECASE)

# # # # #         # allocation_pattern = re.compile(r'Percentage Allocation:.*?(\d+)%', re.IGNORECASE)
# # # # #         # return_pattern = re.compile(r'Expected Annual Return:.*?(\d+)%', re.IGNORECASE)
# # # # #         # time_pattern = re.compile(r'Time Horizon:.*?(\d+) years', re.IGNORECASE)
# # # # #         growth_match = re.search(growth_pattern, response)
# # # # #         conservative_match = re.search(conservative_pattern, response)

# # # # #         if growth_match and conservative_match:
# # # # #             print("Growth and conservative Patttern is found, original logic is working properly")
# # # # #             growth_allocation = int(growth_match.group(1))
# # # # #             conservative_allocation = int(conservative_match.group(1))
# # # # #         elif growth_pattern and conservative_pattern:
# # # # #             print("Growth and conservative Patttern is found")
# # # # #             # growth_allocation = int(growth_match.group(1))
# # # # #             # conservative_allocation = int(conservative_match.group(1))
# # # # #         else:
# # # # #             # Use default values if specific percentages not found
# # # # #             allocation_info = default_allocations[investment_personality]
# # # # #             growth_range = allocation_info['Growth-Oriented Investments']
# # # # #             conservative_range = allocation_info['Conservative Investments']

# # # # #             growth_allocation = (growth_range['min'] + growth_range['max']) / 2
# # # # #             conservative_allocation = (conservative_range['min'] + conservative_range['max']) / 2

# # # # #         # Distribute equally among types
# # # # #         growth_types = default_allocations[investment_personality]['Growth Types']
# # # # #         conservative_types = default_allocations[investment_personality]['Conservative Types']

# # # # #         for g_type in growth_types:
# # # # #             allocations['Growth-Oriented Investments'][g_type] = growth_allocation / len(growth_types)

# # # # #         for c_type in conservative_types:
# # # # #             allocations['Conservative Investments'][c_type] = conservative_allocation / len(conservative_types)

# # # # #     except Exception as e:
# # # # #         st.write(f"Error extracting data: {e}")
# # # # #         # Fallback to default allocation if error occurs
# # # # #         allocation_info = default_allocations[investment_personality]
# # # # #         growth_range = allocation_info['Growth-Oriented Investments']
# # # # #         conservative_range = allocation_info['Conservative Investments']

# # # # #         growth_allocation = (growth_range['min'] + growth_range['max']) / 2
# # # # #         conservative_allocation = (conservative_range['min'] + conservative_range['max']) / 2

# # # # #         # Distribute equally among types
# # # # #         growth_types = allocation_info['Growth Types']
# # # # #         conservative_types = allocation_info['Conservative Types']

# # # # #         for g_type in growth_types:
# # # # #             allocations['Growth-Oriented Investments'][g_type] = growth_allocation / len(growth_types)

# # # # #         for c_type in conservative_types:
# # # # #             allocations['Conservative Investments'][c_type] = conservative_allocation / len(conservative_types)
    
# # # # #     return allocations

# # # # # # Function to create pie charts for allocations
# # # # # def create_pie_charts(allocations):
# # # # #     for key in allocations:
# # # # #         plt.figure(figsize=(8, 6))
# # # # #         labels = list(allocations[key].keys())
# # # # #         sizes = list(allocations[key].values())
        
# # # # #         plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140, colors=sns.color_palette('Set3', len(labels)))
# # # # #         plt.title(f'{key} Allocation')
# # # # #         st.pyplot(plt)
# # # # #         plt.clf()

# # # # # # Function to plot overall split between Growth and Conservative investments
# # # # # def plot_overall_split(allocations):
# # # # #     plt.figure(figsize=(8, 6))
# # # # #     total_growth = sum(allocations['Growth-Oriented Investments'].values())
# # # # #     total_conservative = sum(allocations['Conservative Investments'].values())
    
# # # # #     plt.bar(['Growth-Oriented Investments', 'Conservative Investments'], [total_growth, total_conservative], color=['blue', 'green'])
# # # # #     plt.ylabel('Percentage')
# # # # #     plt.title('Overall Investment Split')
# # # # #     st.pyplot(plt)
# # # # #     plt.clf()

# # # # # # Function to create compounded return chart
# # # # # def create_compounded_return_chart(returns, time_horizon):
# # # # #     plt.figure(figsize=(8, 6))
# # # # #     years = list(range(1, time_horizon + 1))
# # # # #     compounded_returns = [returns[0] * (1 + returns[1] / 100) ** year for year in years]
    
# # # # #     plt.plot(years, compounded_returns, marker='o', linestyle='-', color='blue')
# # # # #     plt.xlabel('Years')
# # # # #     plt.ylabel('Returns')
# # # # #     plt.title('Compounded Returns Over Time')
# # # # #     st.pyplot(plt)
# # # # #     plt.clf()

# # # # # # Main function to generate infographics
# # # # # def generate_infographics(response_text, investment_personality):
# # # # #     allocations = extract_numerical_data(response_text, investment_personality)

# # # # #     st.write(f"Allocations: {allocations}")

# # # # #     create_pie_charts(allocations)
# # # # #     plot_overall_split(allocations)

# # # # #     # Example values for returns and time horizon
# # # # #     example_returns = [10000, 8]  # Principal amount and annual return rate
# # # # #     example_time_horizon = 10     # Number of years
# # # # #     create_compounded_return_chart(example_returns, example_time_horizon)


# # # # # Generate Infographics : Best Code so far:

# # # # import re
# # # # from collections import defaultdict
# # # # import matplotlib.pyplot as plt
# # # # import streamlit as st

# # # # def extract_numerical_data(response):
# # # #     # Define patterns to match different sections and their respective allocations
# # # #     patterns = {
# # # #         'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
# # # #         'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
# # # #         'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
# # # #     }

# # # #     data = defaultdict(dict)

# # # #     for section, pattern in patterns.items():
# # # #         match = pattern.search(response)
# # # #         if match:
# # # #             investments_text = match.group(1)
# # # #             # Extract individual investment types and their allocations
# # # #             investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
# # # #             for investment_match in investment_pattern.findall(investments_text):
# # # #                 investment_type, min_allocation, max_allocation = investment_match
# # # #                 data[section][investment_type.strip()] = {
# # # #                     'min': min_allocation,
# # # #                     'max': max_allocation
# # # #                 }

# # # #     # Extract time horizon and expected returns
# # # #     time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
# # # #     min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
# # # #     max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
# # # #     min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
# # # #     max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

# # # #     time_horizon_match = time_horizon_pattern.search(response)
# # # #     min_return_match = min_return_pattern.search(response)
# # # #     max_return_match = max_return_pattern.search(response)
# # # #     min_growth_match = min_growth_pattern.search(response)
# # # #     max_growth_match = max_growth_pattern.search(response)

# # # #     if time_horizon_match:
# # # #         data['Time Horizon'] = {
# # # #             'min_years': time_horizon_match.group(1),
# # # #             'max_years': time_horizon_match.group(2)
# # # #         }

# # # #     if min_return_match:
# # # #         data['Expected Annual Return'] = {
# # # #             'min': min_return_match.group(1),
# # # #             'max': min_return_match.group(2)
# # # #         }

# # # #     if max_return_match:
# # # #         data['Expected Annual Return'] = {
# # # #             'min': max_return_match.group(1),
# # # #             'max': max_return_match.group(2)
# # # #         }

# # # #     if min_growth_match:
# # # #         data['Expected Growth in Dollars'] = {
# # # #             'min': min_growth_match.group(1),
# # # #             'max': min_growth_match.group(2)
# # # #         }

# # # #     if max_growth_match:
# # # #         data['Expected Growth in Dollars'] = {
# # # #             'min': max_growth_match.group(1),
# # # #             'max': max_growth_match.group(2)
# # # #         }

# # # #     return data

# # # # def plot_investment_allocations(data):
# # # #     # fig, axes = plt.subplots(1, 2, figsize=(14, 7))
# # # #     # fig, axes = plt.subplots(1, 2, figsize=(18, 9))

# # # #     # fig, axes = plt.subplots(2, 1, figsize=(18, 9))
# # # #     fig, axes = plt.subplots(2, 1, figsize=(28, 15))

# # # #     # Plot Growth-Oriented Investments
# # # #     growth_data = data['Growth-Oriented Investments']
# # # #     growth_labels = list(growth_data.keys())
# # # #     growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
# # # #     growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

# # # #     axes[0].barh(growth_labels, growth_min, color='skyblue', label='Min Allocation')
# # # #     axes[0].barh(growth_labels, growth_max, left=growth_min, color='lightgreen', label='Max Allocation')
# # # #     axes[0].set_title('Growth-Oriented Investments')
# # # #     axes[0].set_xlabel('Percentage Allocation')
# # # #     axes[0].legend()

# # # #     # Plot Conservative Investments
# # # #     conservative_data = data['Conservative Investments']
# # # #     conservative_labels = list(conservative_data.keys())
# # # #     conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
# # # #     conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

# # # #     axes[1].barh(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
# # # #     axes[1].barh(conservative_labels, conservative_max, left=conservative_min, color='lightgreen', label='Max Allocation')
# # # #     axes[1].set_title('Conservative Investments')
# # # #     axes[1].set_xlabel('Percentage Allocation')
# # # #     axes[1].legend()

# # # #     plt.tight_layout()
# # # #     return fig

# # # # def plot_pie_chart(data):
# # # #     fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

# # # #     # Combine all investment data for pie chart
# # # #     all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
# # # #     labels = list(all_data.keys())
# # # #     sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
# # # #     colors = plt.cm.Paired(range(len(labels)))

# # # #     ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
# # # #     ax.set_title('Investment Allocation')

# # # #     return fig


# # # # def bar_chart(data):
# # # #     fig, ax = plt.subplots(figsize=(12, 8))  # Increased size

# # # #     # Data for plotting
# # # #     categories = list(data.keys())
# # # #     values_min = [int(data[cat]['min'].strip('%')) for cat in categories]
# # # #     values_max = [int(data[cat]['max'].strip('%')) for cat in categories]

# # # #     x = range(len(categories))

# # # #     ax.bar(x, values_min, width=0.4, label='Min Allocation', color='skyblue', align='center')
# # # #     ax.bar(x, values_max, width=0.4, label='Max Allocation', color='lightgreen', align='edge')

# # # #     ax.set_xticks(x)
# # # #     ax.set_xticklabels(categories, rotation=45, ha='right')
# # # #     ax.set_xlabel('Investment Categories')
# # # #     ax.set_ylabel('Percentage Allocation')
# # # #     ax.set_title('Investment Allocation')
# # # #     ax.legend()

# # # #     plt.tight_layout()
# # # #     return fig



# # # # import plotly.graph_objects as go
# # # # import numpy as np

# # # # import matplotlib.pyplot as plt
# # # # from mpl_toolkits.mplot3d import Axes3D

# # # # def plot_3d_bar_graph(data):
# # # #     # Ensure 'values' is a numpy array
# # # #     values_array = np.array(data['values'])

# # # #     # Check if 'values' can be raveled
# # # #     if values_array.ndim > 1:
# # # #         dz = values_array.ravel()
# # # #     else:
# # # #         dz = values_array

# # # #     # Creating a new figure
# # # #     fig = plt.figure()
# # # #     ax = fig.add_subplot(111, projection='3d')

# # # #     # Unpacking the data
# # # #     x_data = np.arange(len(data['x_labels']))
# # # #     y_data = np.arange(len(data['y_labels']))
# # # #     x, y = np.meshgrid(x_data, y_data)
# # # #     x, y = x.ravel(), y.ravel()
# # # #     z = np.zeros_like(x)
# # # #     dx = dy = 0.8

# # # #     # Creating the bar graph
# # # #     ax.bar3d(x, y, z, dx, dy, dz, color='b', zsort='average')

# # # #     # Setting labels
# # # #     ax.set_xlabel(data['x_label'])
# # # #     ax.set_ylabel(data['y_label'])
# # # #     ax.set_zlabel(data['z_label'])
# # # #     ax.set_xticks(x_data)
# # # #     ax.set_yticks(y_data)
# # # #     ax.set_xticklabels(data['x_labels'])
# # # #     ax.set_yticklabels(data['y_labels'])

# # # #     # Displaying the figure using Streamlit
# # # #     st.pyplot(fig)


# # # # # def plot_3d_bar_chart():
# # # # #     # Data for plotting
# # # # #     categories = ['Timeline', 'Min Compounded Returns', 'Max Compounded Returns']
# # # # #     x = np.arange(len(categories))
# # # # #     y1 = [5, 10, 15]  # Example data for 5 Yr Compounded Returns
# # # # #     y2 = [10, 20, 30]  # Example data for 10 Yr Compounded Returns

# # # # #     fig = go.Figure()

# # # # #     # Adding bars for 5 Yr Compounded Returns
# # # # #     for i in range(len(x)):
# # # # #         fig.add_trace(go.Scatter3d(
# # # # #             x=[x[i], x[i]],
# # # # #             y=[0, 0.5],
# # # # #             z=[0, y1[i]],
# # # # #             mode='lines',
# # # # #             line=dict(color='red', width=10),
# # # # #             name='5 Yr Compounded Returns'
# # # # #         ))

# # # # #     # Adding bars for 10 Yr Compounded Returns
# # # # #     for i in range(len(x)):
# # # # #         fig.add_trace(go.Scatter3d(
# # # # #             x=[x[i] + 0.5, x[i] + 0.5],
# # # # #             y=[0, 0.5],
# # # # #             z=[0, y2[i]],
# # # # #             mode='lines',
# # # # #             line=dict(color='blue', width=10),
# # # # #             name='10 Yr Compounded Returns'
# # # # #         ))

# # # # #     fig.update_layout(
# # # # #         scene=dict(
# # # # #             xaxis=dict(
# # # # #                 tickvals=x + 0.25,
# # # # #                 ticktext=categories,
# # # # #                 title='Categories'
# # # # #             ),
# # # # #             yaxis=dict(title='Returns'),
# # # # #             zaxis=dict(title='Percentage')
# # # # #         ),
# # # # #         title='COMPOUNDED RETURNS',
# # # # #         legend=dict(x=0.1, y=0.9)
# # # # #     )

# # # # #     fig.show() # return fig 
    





# # # # def main():
# # # #     st.title("Wealth Advisor Chatbot")

# # # #     # Step 1: Choose between Investment Suggestions or Stock Analysis
# # # #     task_choice = st.radio("Select the task you want to perform:", ["Investment Suggestions", "Stock Analysis"])

# # # #     # Step 2: Choose between New Client or Existing Client (for Investment Suggestions)
# # # #     if task_choice == "Investment Suggestions":
# # # #         client_type = st.radio("Is this for a new client or an existing client?", ["New Client", "Existing Client"])

# # # #         if client_type == "New Client":
# # # #             uploaded_file = st.file_uploader("Upload the client's document", type=["docx", "pdf"])

# # # #             if uploaded_file is not None:
# # # #                 st.write("Extracting text from the document...")
# # # #                 # Assuming process_document is an async function
# # # #                 extracted_text = asyncio.run(process_document(uploaded_file))
# # # #                 st.write("Text extracted from the document")

# # # #                 # Ask for investment personality
# # # #                 investment_personality = st.selectbox(
# # # #                     "Select the investment personality of the client:",
# # # #                     ("Conservative Investor", "Moderate Investor", "Aggressive Investor")
# # # #                 )

# # # #                 # Step 4: Generate investment suggestions
# # # #                 if st.button("Generate Investment Suggestions"):
# # # #                     st.write("Generating investment suggestions...")
# # # #                     # Assuming generate_investment_suggestions is an async function
                    
# # # #                     suggestions = asyncio.run(generate_investment_suggestions(investment_personality, extracted_text))
# # # #                     st.write(suggestions)

# # # #                     # # Generate infographics
# # # #                     # generate_infographics(suggestions, investment_personality)

# # # #                     data_extracted = extract_numerical_data(suggestions)

# # # #                     # Streamlit app
# # # #                     st.title('Investment Allocation Infographics')
# # # #                     # st.write('### Extracted Investment Data')
# # # #                     # st.json(data_extracted)

# # # #                     st.write('## Investment Allocation Charts')
# # # #                     fig = plot_investment_allocations(data_extracted)
# # # #                     st.pyplot(fig)

# # # #                     st.write('## Pie Chart of Investment Allocation')
# # # #                     fig = plot_pie_chart(data_extracted)
# # # #                     st.pyplot(fig)

# # # #                     st.write('## Bar Chart of Compounded Returns')
# # # #                     fig = bar_chart(data_extracted['Growth-Oriented Investments'])
# # # #                     st.pyplot(fig)

# # # #                     # fig = plot_compounded_returns(data_extracted)
# # # #                     # st.pyplot(fig)

# # # #                     # fig = plot_3d_bar_chart(data_extracted)
# # # #                     plot_3d_bar_graph(data_extracted)
# # # #                     # st.pyplot(fig)

# # # #         elif client_type == "Existing Client":
# # # #             st.write("Fetching existing client data...")

# # # #     elif task_choice == "Stock Analysis":
# # # #         st.write("Performing stock analysis...")

# # # # if __name__ == "__main__":
# # # #     main()


# # # # # def main():
# # # # #     st.title("Wealth Advisor Chatbot")

# # # # #     task_choice = st.radio("Select the task you want to perform:", ["Investment Suggestions", "Stock Analysis"])

# # # # #     if task_choice == "Investment Suggestions":
# # # # #         client_type = st.radio("Is this for a new client or an existing client?", ["New Client", "Existing Client"])

# # # # #         if client_type == "New Client":
# # # # #             uploaded_file = st.file_uploader("Upload the client's document", type=["docx", "pdf"])

# # # # #             if uploaded_file is not None:
# # # # #                 st.write("Extracting text from the document...")
# # # # #                 extracted_text = asyncio.run(process_document(uploaded_file))
# # # # #                 st.write("Text extracted from the document")

# # # # #                 investment_personality = st.selectbox(
# # # # #                     "Select the investment personality of the client:",
# # # # #                     ("Conservative Investor", "Moderate Investor", "Aggressive Investor")
# # # # #                 )

# # # # #                 if st.button("Generate Investment Suggestions"):
# # # # #                     st.write("Generating investment suggestions...")
# # # # #                     suggestions = asyncio.run(generate_investment_suggestions(investment_personality, extracted_text))
# # # # #                     st.write(suggestions)

# # # # #                     st.title("Investment Suggestions and Infographics")
# # # # #                     generate_infographics(suggestions, investment_personality)

# # # # #         elif client_type == "Existing Client":
# # # # #             st.write("Fetching existing client data...")

# # # # #     elif task_choice == "Stock Analysis":
# # # # #         st.write("Performing stock analysis...")

# # # # # if __name__ == "__main__":
# # # # #     main()





# # # # # Reponse :

# # # # # Financial Suggestions for a Moderate Investor

# # # # # Based on your provided information, you appear to be a moderate investor with a healthy financial foundation. You have a significant amount of assets, including a home and other real estate, but also carry a substantial mortgage and credit card debt. This suggests you're comfortable with some risk but also prioritize stability and security.

# # # # # Here's a suggested investment strategy tailored for your profile:

# # # # # Investment Allocation:

# # # # # Growth-Oriented Investments (50-70%): This portion of your portfolio will focus on investments with the potential for higher returns over the long term. Stocks: Invest in a diversified portfolio of stocks through index funds or ETFs tracking the S&P 500 or other broad market indices. This provides exposure to the overall stock market growth. (30-50%) Bonds: Allocate a portion to bonds, which offer lower risk and potential returns compared to stocks. Consider investing in investment-grade corporate bonds or government bonds (Treasury bonds). (10-20%) Real Estate: Maintain your existing real estate investments and consider adding to your portfolio if you have the financial capacity and desire. Real estate can offer steady rental income and potential appreciation. (10-20%)

# # # # # Conservative Investments (30-50%): This portion will focus on preserving capital and providing stability. Cash: Keep a substantial portion of your assets in cash (10-20%) for emergency funds, short-term goals, and potential market downturns. Fixed Annuities: Consider fixed annuities for a guaranteed return and principal protection. They offer lower returns but provide peace of mind. (10-20%) High-Yield Savings Accounts: Explore high-yield savings accounts for a slightly higher return than traditional savings accounts while maintaining FDIC insurance. (10-20%)

# # # # # Investment Time Horizon:

# # # # # As a moderate investor, you likely have a long-term investment horizon (5+ years). This allows for potential market fluctuations and provides time for your investments to compound and grow.

# # # # # Expected Returns:

# # # # # Minimum Expected Annual Return: 4-6% Maximum Expected Annual Return: 8-10%

# # # # # Expected Growth in Dollars:

# # # # # Minimum Expected Growth: 
# # # # # 20
# # # # # ,
# # # # # 000
# # # # # −
# # # # # 20,000−30,000 over 5 years Maximum Expected Growth: 
# # # # # 40
# # # # # ,
# # # # # 000
# # # # # −
# # # # # 40,000−50,000 over 5 years

# # # # # Why this Strategy?

# # # # # This strategy balances potential growth with risk mitigation. The allocation to growth-oriented investments allows for the potential for higher returns, while the conservative investments provide stability and a safety net.

# # # # # Key Considerations:

# # # # # Debt Management: Prioritize paying down your high-interest credit card debt. This will significantly improve your overall financial health and free up more funds for investing. Retirement Savings: Contribute regularly to your Roth IRA and Roth 401(k) to maximize tax benefits and build a strong retirement nest egg. Regular Review: Regularly review your portfolio and adjust your asset allocation as needed based on your risk tolerance, time horizon, and market conditions.

# # # # # Remember: This is a general investment strategy. It's crucial to consult with a qualified financial advisor to tailor a personalized plan that aligns with your specific goals, risk tolerance, and financial situation.


# # # # # def main():
# # # # #     st.title("Wealth Advisor Chatbot")

# # # # #     # Step 1: Upload Document
# # # # #     uploaded_file = st.file_uploader("Upload a document", type=["docx", "pdf"])

# # # # #     if uploaded_file is not None:
# # # # #         # Save the uploaded file temporarily
# # # # #         with open("uploaded_file", "wb") as f:
# # # # #             f.write(uploaded_file.getbuffer())

# # # # #         # Process the document
# # # # #         extracted_text = process_document("uploaded_file")
# # # # #         # st.write(extracted_text)

# # # # #         # Validate the document
# # # # #         # client_name, validation_errors = await validate_process_document("uploaded_file")
# # # # #         client_name, validation_errors = asyncio.run(validate_process_document("uploaded_file"))
# # # # #         if client_name:
# # # # #             st.success(f"Document processed successfully for client: {client_name}")
# # # # #         else:
# # # # #             st.error("Error processing document.")
# # # # #             if validation_errors:
# # # # #                 for error in validation_errors:
# # # # #                     st.error(error)

# # # # #         # Assume retriever and chain setup as in aiogram code
# # # # #         retriever = load_vector_db("uploaded_file")
# # # # #         chain = make_retrieval_chain(retriever)
        
# # # # #         if chain is not None:
# # # # #             summary = "Summary of the document"  # Placeholder for actual summary logic
# # # # #             investment_personality = "Investment Personality"  # Placeholder for actual logic
# # # # #             query = summary + "\n" + investment_personality
# # # # #             response = chain.invoke({"input": query})
# # # # #             format_response = markdown_to_text(response['answer'])
# # # # #             st.write(format_response)

# # # # #             handle_graph(response['answer'])

# # # # #         else:
# # # # #             st.error("Failed to create the retrieval chain. Please upload a valid document.")

# # # # # if __name__ == "__main__":
# # # # #     main()


# # # # # def generate_investment_suggestions(stock_prices) : pass

# # # # # def generate_stock_prices(stock_prices) : pass

# # # # # def predict_stock_prices(stock_prices) : pass

# # # # # Define the Streamlit app layout
# # # # # st.title("Wealth Advisor Chatbot")

# # # # # # Option selection for users
# # # # # st.sidebar.title("Menu")
# # # # # option = st.sidebar.selectbox(
# # # # #     "Choose an option",
# # # # #     ["Investment Suggestions", "Stock Analysis"]
# # # # # )

# # # # # # Investment Suggestions Section
# # # # # if option == "Investment Suggestions":
# # # # #     st.header("Investment Suggestions")
# # # # #     client_type = st.radio("Are you a New or Existing Client?", ("New Client", "Existing Client"))

# # # # #     if client_type == "New Client":
# # # # #         st.subheader("New Client Investment Suggestions")
# # # # #     elif client_type == "Existing Client":
# # # # #         st.subheader("Existing Client Investment Suggestions")
    
# # # # #     # Assuming the function takes in client_type and returns suggestions
# # # # #     investment_suggestions = generate_investment_suggestions(client_type)

# # # # #     st.write(investment_suggestions)

# # # # #     # Generate report with infographics
# # # # #     st.subheader("Investment Report")
# # # # #     st.write("Inflation-adjusted returns for suggested investments:")
# # # # #     # Mock data for illustration, replace with actual data
# # # # #     data = {
# # # # #         "Investment Type": ["Stocks", "Mutual Funds", "ETFs"],
# # # # #         "Inflation-adjusted Returns (%)": [10.5, 8.3, 9.2]
# # # # #     }
# # # # #     df = pd.DataFrame(data)

# # # # #     fig, ax = plt.subplots()
# # # # #     df.plot(kind="bar", x="Investment Type", y="Inflation-adjusted Returns (%)", ax=ax)
# # # # #     st.pyplot(fig)

# # # # # # Stock Analysis Section
# # # # # if option == "Stock Analysis":
# # # # #     st.header("Stock Analysis")
# # # # #     stock_name = st.text_input("Enter Stock Name/Ticker:")
# # # # #     time_period = st.number_input("Enter time period (in days):", min_value=1, max_value=365)

# # # # #     if st.button("Predict Future Prices"):
# # # # #         if stock_name:
# # # # #             predicted_prices = predict_stock_prices(stock_name, time_period)
# # # # #             st.write(f"Predicted Prices for {stock_name} over the next {time_period} days:")
# # # # #             st.line_chart(predicted_prices)
# # # # #         else:
# # # # #             st.error("Please enter a valid stock name/ticker.")

# # # # # Additional sections for Retirement Planning, Mortgage Planning, etc., to be added here




# # # # # # latest version pf client side :

# # # # # #best code so far now it can upload files and receive various forms of messages as well and provide us graphs that we want 
# # # # # # and also reply to images, give stock informations ,give stock analysis information and a prediction of the price

# # # # # import os
# # # # # import filetype
# # # # # import docx
# # # # # import PyPDF2
# # # # # import re
# # # # # from aiogram import Bot, Dispatcher, types
# # # # # from dotenv import load_dotenv
# # # # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # # # from langchain_community.vectorstores import Chroma

# # # # # import faiss
# # # # # from langchain_community.docstore.in_memory import InMemoryDocstore
# # # # # from langchain_community.vectorstores import FAISS
# # # # # from langchain_community.document_loaders import Docx2txtLoader

# # # # # from langchain_core.prompts import ChatPromptTemplate
# # # # # from langchain.chains import create_retrieval_chain
# # # # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # # # from langchain.memory import ConversationSummaryMemory
# # # # # import asyncio
# # # # # import numpy as np
# # # # # import json
# # # # # import re
# # # # # import google.generativeai as genai
# # # # # import pathlib
# # # # # # Import things that are needed generically
# # # # # from langchain.pydantic_v1 import BaseModel, Field
# # # # # from langchain.tools import BaseTool, StructuredTool, tool

# # # # # from aiogram.client.default import DefaultBotProperties
# # # # # from aiogram.enums import ParseMode
# # # # # from aiogram.filters import CommandStart
# # # # # from aiogram.types import Message
# # # # # from aiogram import F
# # # # # from aiogram import Router
# # # # # import logging
# # # # # import sys
# # # # # from aiogram.filters import Command
# # # # # from aiogram.types import FSInputFile
# # # # # # from aiogram.utils import executor
# # # # # import io
# # # # # import matplotlib.pyplot as plt
# # # # # import seaborn as sns
# # # # # import aiohttp
# # # # # from aiogram.types import InputFile , BufferedInputFile
# # # # # import PIL.Image

# # # # # router = Router(name=__name__)

# # # # # load_dotenv()

# # # # # TOKEN = os.getenv("TOKEN")
# # # # # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # # # # os.environ["LANGCHAIN_TRACING_V2"]="true"
# # # # # os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # # # # # Configure generativeai with your API key
# # # # # genai.configure(api_key=GOOGLE_API_KEY)

# # # # # # Initialize bot
# # # # # bot = Bot(token=TOKEN)
# # # # # dp = Dispatcher()

# # # # # # Glbal variables
# # # # # rag_on = False
# # # # # retriever = None  # Store retriever globally
# # # # # summary = ""
# # # # # investment_personality = ""
# # # # # # history = []
# # # # # previous_suggestions = ""

# # # # # CHAT_HISTORY_FILE = 'chat_history.json'

# # # # # def read_chat_history(chat_id):
# # # # #     if os.path.exists(CHAT_HISTORY_FILE):
# # # # #         with open(CHAT_HISTORY_FILE, 'r') as file:
# # # # #             chat_history = json.load(file)
# # # # #             return chat_history.get(str(chat_id), [])
# # # # #     return []

# # # # # def write_chat_history(chat_id, message):
# # # # #     chat_history = {}
# # # # #     if os.path.exists(CHAT_HISTORY_FILE):
# # # # #         with open(CHAT_HISTORY_FILE, 'r') as file:
# # # # #             chat_history = json.load(file)
# # # # #     if str(chat_id) not in chat_history:
# # # # #         chat_history[str(chat_id)] = []
# # # # #     chat_history[str(chat_id)].append(message)
# # # # #     with open(CHAT_HISTORY_FILE, 'w') as file:
# # # # #         json.dump(chat_history, file)

# # # # # class Reference:
# # # # #     def __init__(self):
# # # # #         self.response = ""


# # # # # reference = Reference()


# # # # # def clear_past():
# # # # #     reference.response = ""


# # # # # @router.message(F.text == "clear")
# # # # # async def clear(message: types.Message):
# # # # #     """
# # # # #     A handler to clear the previous conversation and context.
# # # # #     """
# # # # #     clear_past()
# # # # #     await message.reply("I've cleared the past conversation and context.")

# # # # # #Global Variables :

# # # # # # Store user states
# # # # # states = {}

# # # # # # Dictionary to hold question-answer pairs
# # # # # user_responses = {}
# # # # # #
# # # # # user_images = {}
# # # # # # Define Questions for assessment
# # # # # questions = [
# # # # #     """ 
# # # # #     1. You and your friend are betting on a series of coin tosses.

# # # # #     He always bets ₹2,000 on Heads

# # # # #     You always bet ₹2,000 on Tails

# # # # #     Winner of last 8 turns

# # # # #     You lost ₹8,000 in the last 4 turns!

# # # # #     If you were to bet one last time, what would you bet on:
# # # # #     a) heads or b) tails ?
# # # # #     """ ,
# # # # #     """
# # # # #     2. Imagine you are a contestant in a game show, and you are presented the following choices.

# # # # #     What would you prefer?
# # # # #     a) 50 percent chance of winning 15 gold coins 
# # # # #     b) 100 percent chance of winning 8 gold coins
# # # # #     """,
# # # # #     """
# # # # #     3. In general, how would your best friend describe your risk-taking tendencies?
# # # # #     a) A real gambler
# # # # #     b) Willing to take risks after completing adequate research
# # # # #     c) Cautious
# # # # #     d) Avoids risk as much as possible
# # # # #     """,
# # # # #     """
# # # # #     4. Suppose you could replace your current investment portfolio with this new one:
# # # # #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# # # # #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# # # # #     a)-5 to -10
# # # # #     b)-10 to -15
# # # # #     c)-15 to -20
# # # # #     d)-20 to -25
# # # # #     e)-25 to -30
# # # # #     f)-30 to -35
# # # # #     """,
# # # # #     """
# # # # #     5. Over any 1-year period, what would be the maximum drop in the value of your investment 
# # # # #     portfolio that you would be comfortable with?
# # # # #     a) <5%
# # # # #     b) 5 - 10%
# # # # #     c) 10 - 15%
# # # # #     d) 15 - 20%
# # # # #     e) >20%
# # # # #     """,
# # # # #     """
# # # # #     6. When investing, what do you consider the most?

# # # # #     a) Risk 
# # # # #     b) Return
# # # # #     """,
# # # # #     """
# # # # #     7. What best describes your attitude?

# # # # #     a) Prefer reasonable returns, can take reasonable risk
# # # # #     b) Like higher returns, can take slightly higher risk
# # # # #     c) Want to maximize returns, can take significant high risk
# # # # #     """,
# # # # #     """
# # # # #     8. How much monthly investment you want to do?
# # # # #     """,
# # # # #     """
# # # # #     9. What is the time horizon for your investment?
# # # # #     You can answer in any range, example 1-5 years."""  
# # # # # ]



# # # # # import logging
# # # # # from aiogram import Bot, Dispatcher, types
# # # # # # Register the router with the dispatcher
# # # # # dp.include_router(router)

# # # # # # from aiogram.utils import executor
# # # # # from aiogram.filters import CommandStart
# # # # # from aiogram.types import Poll, PollAnswer

# # # # # # Command handler to start the poll
# # # # # @dp.message(CommandStart())
# # # # # async def handle_start(message: types.Message):
# # # # #     chat_id = message.chat.id
# # # # #     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! ")
# # # # #     question="How Can I Help you today ?"
# # # # #     options = """\na. Know my Investment Personality \nb. Tax Related Queires \nc. Savings and Wealth Management \nd. Debt Repayment Strategies
# # # # #               """
# # # # #     await bot.send_message(chat_id, question + options)

    

# # # # # # Function to start the assessment
# # # # # async def start_assessment(chat_id):
# # # # #     await bot.send_message(chat_id, """To analyse your investment personality I need to ask you some questions.\nLet's start a quick personality assessment.""")
# # # # #     await ask_next_question(chat_id, 0)

# # # # # # Function to ask the next question
# # # # # async def ask_next_question(chat_id, question_index):
# # # # #     if question_index < len(questions):
# # # # #         # Ask the next question
# # # # #         await bot.send_message(chat_id, questions[question_index])
# # # # #         # Update state to indicate the next expected answer
# # # # #         states[chat_id] = question_index
# # # # #     else:
# # # # #         # No more questions, finish assessment
# # # # #         await finish_assessment(chat_id)

# # # # # # Handler for receiving assessment answers
# # # # # assessment_in_progress = True

# # # # # from aiogram.types import FSInputFile
# # # # # async def finish_assessment(chat_id):
# # # # #     if chat_id in states and states[chat_id] == len(questions):
# # # # #         # All questions have been answered, now process the assessment
# # # # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # # # #         # Determine investment personality based on collected responses
# # # # #         global investment_personality
# # # # #         investment_personality = await determine_investment_personality(user_responses)

# # # # #         # Inform the user about their investment personality
# # # # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # # # #         # Store the response in chat history
# # # # #         write_chat_history(chat_id, {'role': 'bot', 'message': investment_personality})

# # # # #         # Summarize collected information
# # # # #         global summary
# # # # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # # # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # # # #         # Ensure to await the determination of investment personality
# # # # #         await send_summary_chunks(chat_id, summary)
# # # # #         global assessment_in_progress 
# # # # #         assessment_in_progress = False
       
# # # # #         # await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
# # # # #         # file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

# # # # #         # await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
# # # # #         # await bot.send_message(chat_id,file)

# # # # # async def send_summary_chunks(chat_id, summary):
# # # # #     # Split the summary into chunks that fit within Telegram's message limits
# # # # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# # # # #     # Send each chunk as a separate message
# # # # #     for chunk in chunks:
# # # # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # # # # async def determine_investment_personality(assessment_data):
# # # # #     try:
# # # # #         # Prepare input text for the chatbot based on assessment data
# # # # #         input_text = "User Profile:\n"
# # # # #         for question, answer in assessment_data.items():
# # # # #             input_text += f"{question}: {answer}\n"

# # # # #         # Introduce the chatbot's task and prompt for classification
# # # # #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# # # # #                       "- Conservative Investor\n" \
# # # # #                       "- Moderate Investor\n" \
# # # # #                       "- Aggressive Investor"

# # # # #         # Use your generative AI model to generate a response
# # # # #         # print(input_text)
# # # # #         model = genai.GenerativeModel('gemini-pro')
# # # # #         response = model.generate_content(input_text)

# # # # #         # Determine the investment personality from the chatbot's response
# # # # #         response_text = response.text.lower()
# # # # #         if "conservative" in response_text:
# # # # #             personality = "Conservative Investor"
# # # # #         elif "moderate" in response_text:
# # # # #             personality = "Moderate Investor"
# # # # #         elif "aggressive" in response_text:
# # # # #             personality = "Aggressive Investor"
# # # # #         else:
# # # # #             personality = "Unknown"

# # # # #         return personality
# # # # #         # Send the determined investment personality back to the user
# # # # #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# # # # #     except Exception as e:
# # # # #         print(f"Error generating response: {e}")
# # # # #         #await bot.send_message(chat_id, "Error processing investment personality classification.")


# # # # # # Tax Related Queries :

# # # # # # Define a global state to track the questions
# # # # # tax_states = {}
# # # # # tax_responses = {}

# # # # # # Define the questions
# # # # # tax_questions = [
# # # # #     "What is your annual income?",
# # # # #     "In which state do you live?",
# # # # #     "Are you married or single?\n(a) Married\n(b) Single",
# # # # #     "For which year do you wish to calculate tax?",
# # # # #     "Do you have any mortgages or any tax reductions?"
# # # # # ]

# # # # # # @dp.message()
# # # # # async def tax_management(message: types.Message):
# # # # #     chat_id = message.chat.id
# # # # #     # If the user is answering a tax question, process the response
# # # # #     if chat_id in tax_states and tax_states[chat_id] < len(tax_questions):
# # # # #         question_index = tax_states[chat_id]
# # # # #         answer = message.text
# # # # #         tax_responses[tax_questions[question_index]] = answer
# # # # #         tax_states[chat_id] += 1

# # # # #         # Ask the next question
# # # # #         if tax_states[chat_id] < len(tax_questions):
# # # # #             await bot.send_message(chat_id, tax_questions[tax_states[chat_id]])
# # # # #         else:
# # # # #             # All questions answered, now process the data
# # # # #             await calculate_taxes(chat_id)
# # # # #             await bot.send_message(chat_id, "Thank you for your responses! Your tax-related queries have been processed.")
# # # # #             # Reset the state
# # # # #             del tax_states[chat_id]
# # # # #             del tax_responses[chat_id]

# # # # #     # If the user starts the tax management flow, ask the first question
# # # # #     else:
# # # # #         tax_states[chat_id] = 0
# # # # #         tax_responses[chat_id] = {}
# # # # #         await bot.send_message(chat_id, "Let's get started with your tax-related queries.")
# # # # #         await bot.send_message(chat_id, tax_questions[0])

# # # # # async def calculate_taxes(chat_id):
# # # # #     try:
# # # # #         # Get the user's responses
# # # # #         annual_income = tax_responses[tax_questions[0]]
# # # # #         state = tax_responses[tax_questions[1]]
# # # # #         marital_status = tax_responses[tax_questions[2]]
# # # # #         tax_year = tax_responses[tax_questions[3]]
# # # # #         mortgages_or_deductions = tax_responses[tax_questions[4]]

# # # # #         # Prepare the context for the LLM
# # # # #         context = f"""
# # # # #         Annual Income: {annual_income}
# # # # #         State: {state}
# # # # #         Marital Status: {marital_status}
# # # # #         Tax Year: {tax_year}
# # # # #         Mortgages or Deductions: {mortgages_or_deductions}
# # # # #         """

# # # # #         # Use the context to get tax calculation and advice
# # # # #         task = """You are a Tax Calculations Expert in the entire world.
# # # # #             Ask user tax related queries to help users with tax related queries.
# # # # #             Consider user's investment personality  if provided.
# # # # #             Address the user by their name(client_name: Emily in our case but if any other name is given refer to that) if provided.
# # # # #             Help users to save tax on their income and earnings.
# # # # #             If user asks queries related to saving taxes or calculating taxes refer to the 
# # # # #             US Tax Laws given by the IRS and based on that information calculate the taxes for the user 
# # # # #             consider the information shared by the user such as their annual income and their monthly investment if provided,
# # # # #             also give advice to the user on how they can save their taxes.
# # # # #             Include a disclaimer to monitor and rebalance investments regularly based on risk tolerance."""
        
# # # # #         query = task + "\n" + context
# # # # #         model = genai.GenerativeModel('gemini-1.5-flash')
# # # # #         chat = model.start_chat(history=[])
# # # # #         response = chat.send_message(query)

# # # # #         # Enhanced logging for debugging
# # # # #         logging.info(f"Model response: {response}")
# # # # #         format_response = markdown_to_text(response.text)

# # # # #         # Store the response in chat history
# # # # #         write_chat_history(chat_id, {'role': 'bot', 'message': format_response})
# # # # #         await bot.send_message(chat_id,"Here is your calculated tax as per the responses provided ")
# # # # #         await bot.send_message(chat_id, format_response)

# # # # #     except Exception as e:
# # # # #         print(f"Error calculating taxes: {e}")
# # # # #         await bot.send_message(chat_id, "Error calculating taxes. Please try again later.")

# # # # # # async def tax_management(chat_id) :
# # # # # #     try:
# # # # #         # task = """You are a Tax Calculations Expert in the entire world.
# # # # #         #     Ask user tax related queries to help users with tax related queries.
# # # # #         #     Consider user's investment personality  if provided.
# # # # #         #     Address the user by their name(client_name: Emily in our case but if any other name is given refer to that) if provided.
# # # # #         #     Help users to save tax on their income and earnings.
# # # # #         #     If user asks queries related to saving taxes or calculating taxes refer to the 
# # # # #         #     US Tax Laws given by the IRS and based on that information calculate the taxes for the user 
# # # # #         #     consider the information shared by the user such as their annual income and their monthly investment if provided,
# # # # #         #     also give advice to the user on how they can save their taxes.
# # # # #         #     Include a disclaimer to monitor and rebalance investments regularly based on risk tolerance."""
        
        
# # # # # #         # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + message.text
# # # # # #         # query = chat_history_text + "\n" + query

# # # # # #         # Include chat history
# # # # # #         chat_history = read_chat_history(chat_id)
# # # # # #         chat_history_text = '\n'.join([f"{entry['role']}: {entry['message']}" for entry in chat_history])
# # # # # #         # history.append(chat_history_text)

# # # # # #         # query = task + "\n" + investment_personality + "\n" + chat_history_text # + "\n" + message.text
# # # # # #         query = task + "\n" + chat_history_text # + "\n" + message.text


# # # # # #         model = genai.GenerativeModel('gemini-1.5-flash')
# # # # # #         chat = model.start_chat(history=[])
# # # # # #         response = chat.send_message(query)

# # # # # #         # Enhanced logging for debugging
# # # # # #         logging.info(f"Model response: {response}")
# # # # # #         format_response = markdown_to_text(response.text) #(response_text) #response.result

# # # # # #         # Store the response in chat history
# # # # # #         write_chat_history(chat_id, {'role': 'bot', 'message': format_response})
# # # # # #         await bot.send_message(chat_id,format_response)
# # # # # #         # await message.reply(format_response)

# # # # # #     except Exception as e:
# # # # # #         print(f"Error invoking retrieval chain on attempt : {e}")
# # # # # #         await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")
        

# # # # # # Savings and Wealth Management :
# # # # # async def savings_management(chat_id) :
# # # # #     # Savings and Wealth Management related questions
# # # # #     await bot.send_message(chat_id,"Hello there here is a Simple Personal Budget Excel File.Please fill in your details with correct Information and then upload it in the chat")
# # # # #     file = FSInputFile("data\Emily_Budget.xlsx", filename="Your Simple Personal Budget.xlsx")


# # # # # # Handler for document upload
# # # # # async def load_vector_db(file_path):
# # # # #     try:
# # # # #         print("Loading vector database...")
# # # # #         loader = Docx2txtLoader(file_path)
# # # # #         documents = loader.load()
# # # # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # # # #         text_chunks = text_splitter.split_documents(documents)
# # # # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # # # #         # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
# # # # #         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
# # # # #         # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

# # # # #         # vector_store = FAISS(
# # # # #         #     embedding_function=embeddings,
# # # # #         #     index=index,
# # # # #         #     docstore=InMemoryDocstore(),
# # # # #         #     index_to_docstore_id={},
# # # # #         # )
        
# # # # #         print("Vector database loaded successfully.") 
# # # # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # # # #     except Exception as e:
# # # # #         print(f"Error loading vector database: {e}")
# # # # #         return None


# # # # # # change prompt template :
# # # # # async def make_retrieval_chain(retriever):
# # # # #     """
# # # # #     Create a retrieval chain using the provided retriever.

# # # # #     Args:
# # # # #         retriever (RetrievalQA): A retriever object.

# # # # #     Returns:
# # # # #         RetrievalQA: A retrieval chain object.
# # # # #     """
# # # # #     try:
# # # # #         global investment_personality,summary
# # # # #         llm = ChatGoogleGenerativeAI(
# # # # #             #model="gemini-pro",
# # # # #             model = "gemini-1.5-flash",
# # # # #             temperature=0.7,
# # # # #             top_p=0.85,
# # # # #             google_api_key=GOOGLE_API_KEY
# # # # #         )

# # # # #         # prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # # #         #         Respond to the client by the client name.
# # # # #         #         Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # #         #         should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
# # # # #         #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # #         #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
# # # # #         #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
# # # # #         #         Also explain the user why you are giving them that particular investment suggestion.
# # # # #         #         Give the client suggestions of Investment based on their investment personality that can also help them manage their mortages and other liablilities if they have any,ignore if they dont have any liabilities.
# # # # #         #         Answer in 3-4 lines.\n
# # # # #         #         <context>
# # # # #         #         {context}
# # # # #         #         </context>
# # # # #         #         Question: {input}"""


# # # # #         prompt_template = investment_personality + "\n" + summary + "\n" + """ Role : You are a Top class highly professional and world's best Savings Advisor for 
# # # # #                 savings related question-answering tasks related to the document.
# # # # #                 Respond to the client by the client name.
# # # # #                 Give Savings Suggestions to the user so that they could do proper responsible savings and save their expenses based on their investment personality and budget if provided.
# # # # #                 Also give the user detailed information about their savings such that they could save more money and save their expenses.
# # # # #                 Give the user minimum and maximum percentage of savings the user can do by reducing their expenses. If the users have given a budget then analyse it and give suggestions based on that.
# # # # #                 Try to imitate human language and talk to the user/client like a human and give personal savings suggestions.
# # # # #                 If the user is having many unnecessary expenses then give the user some advice in a gentle manner without offending the user or hurt their feelings and suggest and advice them to stop or reduce their unnecessary expenses 
# # # # #                 in order to increase their savings.
# # # # #                 Also explain the user why you are giving them that particular savings suggestion.
# # # # #                 Give the client suggestions of savings based on their investment personality that can also help them manage their mortages and other liablilities if they have any,ignore if they dont have any liabilities.
# # # # #                 Answer in 3-4 lines.\n
# # # # #                 <context>
# # # # #                 {context}
# # # # #                 </context>
# # # # #                 Question: {input}"""

# # # # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # # # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
# # # # #         combine_docs_chain = None  

# # # # #         if retriever is not None :  
# # # # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # # # #             print(retriever_chain)
# # # # #             return retriever_chain
# # # # #         else:
# # # # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # # # #             return None

# # # # #     except Exception as e:
# # # # #         print(f"Error in creating chain: {e}")
# # # # #         return None

# # # # # from aiogram.filters import Filter

# # # # # # @router.message(F.document)

# # # # # import os
# # # # # import pandas as pd
# # # # # from openpyxl import load_workbook

# # # # # @dp.message(F.document)
# # # # # async def handle_document(message: types.Message):
# # # # #     global summary, investment_personality  

# # # # #     chat_id = message.chat.id
# # # # #     await message.reply("File Received") 
    
# # # # #     # Obtain file information
# # # # #     file_id = message.document.file_id
# # # # #     file = await bot.get_file(file_id)
# # # # #     file_path = file.file_path

# # # # #     # Get the file extension
# # # # #     file_extension = os.path.splitext(message.document.file_name)[-1].lower()

# # # # #     # Download the file
# # # # #     local_file_path = "data/uploaded_file" + file_extension
# # # # #     await bot.download_file(file_path, local_file_path)

# # # # #     # Process the uploaded document based on the file type
# # # # #     if file_extension in ['.xlsx', '.xls']:
# # # # #         extracted_text = await process_excel_file(local_file_path)
# # # # #     else:
# # # # #         extracted_text = await process_document(local_file_path)

# # # # #     if extracted_text:
# # # # #         # Proceed with further processing
# # # # #         print("Retriever being loaded ")
# # # # #         retriever = await load_vector_db(local_file_path)
# # # # #         client_name, validation_errors = await validate_process_document(local_file_path)

# # # # #         print(f"Client Name: {client_name}")
# # # # #         if validation_errors:
# # # # #             print("**Validation Errors:**")
# # # # #             for error in validation_errors:
# # # # #                 print(error)
# # # # #         else:
# # # # #             print("All fields are filled correctly.")
# # # # #         # if client_name is None:
# # # # #         #     try:
# # # # #         #         await message.reply("Processing the uploaded image")
# # # # #         #         await handle_image(message) 
# # # # #         #         return 
# # # # #         #     except Exception as e:
# # # # #         #         await message.reply("Error processing uploaded image")
# # # # #         #         print(e)
# # # # #         if client_name == None : client_name = "Emilly"
# # # # #         await message.reply(f"Thanks for providing me the details, {client_name}. I have processed the file and now I will provide you some Savings suggestions based on the details that you have provided.")

# # # # #         if retriever is None:
# # # # #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # # #             return

# # # # #         chain = await make_retrieval_chain(retriever)
# # # # #         if chain is None:
# # # # #             await message.reply("Failed to create the retrieval chain.")
# # # # #             return

# # # # #         try:     
# # # # #             query = summary + "\n" + investment_personality
     
# # # # #             response = chain.invoke({"input": query})
# # # # #             print(response['answer'])
# # # # #             global chat_history
# # # # #             chat_history = response['answer'] 
# # # # #             print(f"\n Chat History : {chat_history}")
# # # # #             format_response = markdown_to_text(response['answer'])

# # # # #             write_chat_history(chat_id, {'role': 'bot', 'message': extracted_text})
# # # # #             write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

# # # # #             await message.reply(format_response)

# # # # #         except Exception as e:
# # # # #             print(f"Error invoking retrieval chain on attempt : {e}")
# # # # #             await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

# # # # #     else:
# # # # #         await message.reply("Failed to process the uploaded file.")

# # # # # async def process_excel_file(file_path):
# # # # #     try:
# # # # #         # Reading the Excel file using pandas
# # # # #         df = pd.read_excel(file_path)
# # # # #         # Extracting relevant information (This is just an example, customize it as needed)
# # # # #         extracted_text = df.to_string(index=False)
# # # # #         return extracted_text
# # # # #     except Exception as e:
# # # # #         print(f"Error processing Excel file: {e}")
# # # # #         return None


# # # # # # @dp.message(F.document)
# # # # # # async def handle_document(message: types.Message):
# # # # # #     global summary,investment_personality  

# # # # # #     chat_id = message.chat.id
# # # # # #     await message.reply("File Received") 
# # # # # #     # Obtain file information
# # # # # #     file_id = message.document.file_id
# # # # # #     file = await bot.get_file(file_id)
# # # # # #     file_path = file.file_path
    
# # # # # #     # Download the file
# # # # # #     await bot.download_file(file_path, "data/uploaded_file")
    
# # # # # #     # Process the uploaded document
# # # # # #     extracted_text = await process_document("data/uploaded_file")
# # # # # #     # print(extracted_text)

# # # # # #     if extracted_text:
# # # # # #         # Load vector database (assuming this is part of setting up the retriever)
# # # # # #         print("Retriever being loaded ")
# # # # # #         retriever = await load_vector_db("data/uploaded_file")
# # # # # #         file_path = 'data/uploaded_file'
# # # # # #         client_name, validation_errors = await validate_process_document(file_path)

# # # # # #         # Print results
# # # # # #         print(f"Client Name: {client_name}")
# # # # # #         if validation_errors:
# # # # # #             print("**Validation Errors:**")
# # # # # #             for error in validation_errors:
# # # # # #                 print(error)
# # # # # #         else:
# # # # # #             print("All fields are filled correctly.")
# # # # # #         if client_name == None:
# # # # # #             try:
# # # # # #                 await message.reply("Processing the uploaded image")
# # # # # #                 await handle_image(message) 
# # # # # #                 return 
# # # # # #             except Exception as e:
# # # # # #                 await message.reply("error processing uploaded image")
# # # # # #                 print(e)
# # # # # #         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some Savings suggestions based on the details that you have provided.")

# # # # # #         if retriever is None:
# # # # # #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # # # #             return

# # # # # #         # Check if a valid chain can be created
# # # # # #         chain = await make_retrieval_chain(retriever)
# # # # # #         if chain is None:
# # # # # #             await message.reply("Failed to create the retrieval chain.")
# # # # # #             return
        
# # # # # #         try:     
# # # # # #             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
        
# # # # # #             response = chain.invoke({"input": query})
# # # # # #             print(response['answer'])
# # # # # #             global chat_history
# # # # # #             chat_history = response['answer'] 
# # # # # #             print(f"\n Chat History : {chat_history}")
# # # # # #             format_response = markdown_to_text(response['answer'])

# # # # # #             # Store the extracted_text in chat history
# # # # # #             write_chat_history(chat_id, {'role': 'bot', 'message': extracted_text})
        
# # # # # #             # Store the response in chat history
# # # # # #             write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

# # # # # #             await message.reply(format_response)
# # # # # #             # await message.reply(response['answer'])

# # # # # #         except Exception as e:
# # # # # #             print(f"Error invoking retrieval chain on attempt : {e}")
# # # # # #             await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

# # # # # #     else:
# # # # # #         await message.reply("Failed to process the uploaded file.")
    

# # # # # # Function to extract data from LLM response
# # # # # def extract_data_from_response(response):
# # # # #     try:
# # # # #         # Locate the JSON-like data in the response
# # # # #         json_start = response.find("{")
# # # # #         json_end = response.rfind("}") + 1
        
# # # # #         if json_start == -1 or json_end == -1:
# # # # #             raise ValueError("No JSON data found in the response.")
        
# # # # #         json_data = response[json_start:json_end]
        
# # # # #         # Parse the JSON data
# # # # #         data = json.loads(json_data.replace("'", "\""))
# # # # #         print(data)
# # # # #         return data
# # # # #     except Exception as e:
# # # # #         logging.error(f"Error extracting data: {e}")
# # # # #         return None

 


# # # # # def extract_allocations_from_json(json_data,chat_id):
# # # # #     allocations = {}
# # # # #     for entry in json_data.get(str(chat_id), []):
# # # # #         if entry['role'] == 'bot':
# # # # #             message = entry['message']
# # # # #             lines = message.split('\n')
# # # # #             current_category = None

# # # # #             for line in lines:
# # # # #                 match = re.match(r'^(.*?):\s*(\d+)%$', line)
# # # # #                 if match:
# # # # #                     category, percent = match.groups()
# # # # #                     allocations[category] = []
# # # # #                     current_category = category
# # # # #                 elif current_category and re.match(r'.*\d+%', line):
# # # # #                     subcategory_match = re.match(r'^(.*?)(\d+)%$', line)
# # # # #                     if subcategory_match:
# # # # #                         subcategory, percent = subcategory_match.groups()
# # # # #                         allocations[current_category].append((subcategory.strip(), float(percent)))

# # # # #     return allocations


# # # # # def create_pie_chart(allocations, chat_id):
# # # # #     labels = []
# # # # #     sizes = []
# # # # #     for category, subcategories in allocations.items():
# # # # #         for subcategory, percent in subcategories:
# # # # #             labels.append(f"{category} - {subcategory}")
# # # # #             sizes.append(percent)
    
# # # # #     if sizes:
# # # # #         fig, ax = plt.subplots()
# # # # #         ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
# # # # #         ax.axis('equal')
        
# # # # #         plt.title("Investment Allocation")
# # # # #         chart_path = f"data/investment_allocation_{chat_id}.png"
# # # # #         plt.savefig(chart_path)
# # # # #         plt.close()
        
# # # # #         return chart_path
# # # # #     else:
# # # # #         return None
  

# # # # # async def process_document(file_path):
# # # # #     try:
# # # # #         print("Processing the document")
# # # # #         file_type = filetype.guess(file_path)
# # # # #         if file_type is not None:
# # # # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # # #                 # return extract_text_from_word(file_path)
# # # # #                 return extract_text_and_tables_from_word(file_path)
# # # # #             elif file_type.mime == "application/pdf":
# # # # #                 return extract_text_from_pdf(file_path)
# # # # #         return None
# # # # #     except Exception as e:
# # # # #         print(f"Error processing document: {e}")
# # # # #         return None

# # # # # def extract_text_from_pdf(pdf_file_path):
# # # # #     try:
# # # # #         print("Processing pdf file")
# # # # #         with open(pdf_file_path, "rb") as pdf_file:
# # # # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # # # #             text_content = []
# # # # #             for page_num in range(pdf_reader.numPages):
# # # # #                 page = pdf_reader.getPage(page_num)
# # # # #                 text_content.append(page.extract_text())
# # # # #             return "\n".join(text_content)
# # # # #     except Exception as e:
# # # # #         print(f"Error extracting text from PDF: {e}")
# # # # #         return None


# # # # # import re
# # # # # import docx

# # # # # def extract_text_and_tables_from_word(docx_file_path):
# # # # #     """
# # # # #     Extracts text and tables from a Word document (.docx).

# # # # #     Args:
# # # # #         docx_file_path (str): Path to the Word document file.

# # # # #     Returns:
# # # # #         tuple: Extracted text content and tables from the document.
# # # # #     """
# # # # #     try:
# # # # #         print("Extracting text and tables from word file")
# # # # #         doc = docx.Document(docx_file_path)
# # # # #         text_content = []
# # # # #         tables_content = []

# # # # #         for para in doc.paragraphs:
# # # # #             text_content.append(para.text)

# # # # #         for table in doc.tables:
# # # # #             table_data = []
# # # # #             for row in table.rows:
# # # # #                 row_data = []
# # # # #                 for cell in row.cells:
# # # # #                     row_data.append(cell.text.strip())
# # # # #                 table_data.append(row_data)
# # # # #             tables_content.append(table_data)
# # # # #         print("Extracted text from word file")
# # # # #         return "\n".join(text_content), tables_content
# # # # #     except Exception as e:
# # # # #         print(f"Error extracting text and tables from Word document: {e}")
# # # # #         return None, None

# # # # # def validate_document_content(text, tables):
# # # # #     """
# # # # #     Validates the content of the document.

# # # # #     Args:
# # # # #         text (str): Extracted text content from the document.
# # # # #         tables (list): Extracted tables content from the document.

# # # # #     Returns:
# # # # #         tuple: Client name and validation errors.
# # # # #     """
# # # # #     errors = []
    
# # # # #     # Extract client name
# # # # #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# # # # #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# # # # #     # Define required sections
# # # # #     required_sections = [
# # # # #         "YOUR RETIREMENT GOAL",
# # # # #         "YOUR OTHER MAJOR GOALS",
# # # # #         "YOUR ASSETS AND LIABILITIES",
# # # # #         "MY LIABILITIES",
# # # # #         "YOUR CURRENT ANNUAL INCOME"
# # # # #     ]

# # # # #     # Check for the presence of required sections
# # # # #     for section in required_sections:
# # # # #         if section not in text:
# # # # #             errors.append(f"* {section} section missing.")
    
# # # # #     # Define table field checks
# # # # #     table_checks = {
# # # # #         "YOUR RETIREMENT GOAL": [
# # # # #             r"When do you plan to retire\? \(age or date\)",
# # # # #             r"Social Security Benefit \(include expected start date\)",
# # # # #             r"Pension Benefit \(include expected start date\)",
# # # # #             r"Other Expected Income \(rental, part-time work, etc.\)",
# # # # #             r"Estimated Annual Retirement Expense"
# # # # #         ],
# # # # #         "YOUR OTHER MAJOR GOALS": [
# # # # #             r"GOAL", r"COST", r"WHEN"
# # # # #         ],
# # # # #         "YOUR ASSETS AND LIABILITIES": [
# # # # #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# # # # #             r"Current Value", r"Annual Contributions"
# # # # #         ],
# # # # #         "MY LIABILITIES": [
# # # # #             r"Balance", r"Interest Rate", r"Monthly Payment"
# # # # #         ]
# # # # #     }

# # # # #     # Validate table content
# # # # #     for section, checks in table_checks.items():
# # # # #         section_found = False
# # # # #         for table in tables:
# # # # #             table_text = "\n".join(["\t".join(row) for row in table])
# # # # #             if section in table_text:
# # # # #                 section_found = True
# # # # #                 for check in checks:
# # # # #                     if not re.search(check, table_text, re.IGNORECASE):
# # # # #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# # # # #                 break
# # # # #         if not section_found:
# # # # #             errors.append(f"* {section} section missing.")

# # # # #     return client_name, errors

# # # # # async def validate_process_document(file_path):
# # # # #     try:
# # # # #         print("Validating process document : ")
# # # # #         text, tables = extract_text_and_tables_from_word(file_path)
# # # # #         if text is not None and tables is not None:
# # # # #             client_name, errors = validate_document_content(text, tables)
# # # # #             return client_name, errors
# # # # #         return None, ["Error processing document."]
# # # # #     except Exception as e:
# # # # #         print(f"Error processing document: {e}")
# # # # #         return None, [f"Error processing document: {e}"]

# # # # # @dp.message()
# # # # # async def main_bot(message: types.Message):
# # # # #     global retriever, extracted_text, investment_personality, summary, chat_history

# # # # #     chat_id = message.chat.id
# # # # #     question="How Can I Help you today ?"
# # # # #     options = """\na. Know my Investment Personality \nb. Tax Related Queires \nc. Savings and Wealth Management \nd. Debt Repayment Strategies
# # # # #               """
    
# # # # #     # if chat_id in states and states[chat_id] < len(questions):
# # # # #     #     question_index = states[chat_id]
# # # # #     #     answer = message.text
# # # # #     #     user_responses[questions[question_index]] = answer
# # # # #     #     states[chat_id] += 1
# # # # #     #     await ask_next_question(chat_id, question_index + 1)

# # # # #     if chat_id in states and states[chat_id] < len(questions):
# # # # #         question_index = states[chat_id]
# # # # #         answer = message.text
# # # # #         user_responses[questions[question_index]] = answer
# # # # #         states[chat_id] += 1
# # # # #         if states[chat_id] < len(questions):
# # # # #             await ask_next_question(chat_id, question_index + 1)
# # # # #         else:
# # # # #             await ask_next_question(chat_id, question_index + 1)
            
# # # # #             # await bot.send_message(chat_id, "Assessment Completed.")
# # # # #             await bot.send_message(chat_id, "What do you want to do next?\n" + question + options)


# # # # #     elif message.text:
# # # # #         lower_text = message.text.lower()

# # # # #         # Investment Personality :
# # # # #         if any(variant in lower_text for variant in ["a", "a.", "a)", "(a)", "1", "1.", "1)", "(1)"]):
# # # # #             await start_assessment(chat_id)

# # # # #         # Tax Related Queries :
# # # # #         elif any(variant in lower_text for variant in ["b", "b.", "b)", "(b)", "2", "2.", "2)", "(2)"]):
# # # # #             await bot.send_message(chat_id,"Hello, I will ask you some Tax Related Questions.\nPlease answer them correctly so that I can calculate your tax")
# # # # #             await tax_management(message) #(chat_id) #pass
        
# # # # #         # Savings and Wealth Management :
# # # # #         elif any(variant in lower_text for variant in ["c", "c.", "c)", "(c)", "3", "3.", "3)", "(3)"]):
# # # # #             await savings_management(chat_id) #pass  #  
        
# # # # #         # Debt Repayment Strategies :
# # # # #         elif any(variant in lower_text for variant in ["d", "d.", "d)", "(d)", "4", "4.", "4)", "(4)"]):
# # # # #             pass  # 

# # # # #         elif lower_text in ["yes", "y"]:
# # # # #             await start_assessment(chat_id)

# # # # #         else:
# # # # #             await bot.send_message(chat_id, "Assessment Completed. Do you wish to retake the assessment? Type 'yes' or 'no'.")
# # # # #             await bot.send_message(chat_id, "Thank you for your response.")
# # # # #             await bot.send_message(chat_id, "What do you want to do next?\n" + question + options)

# # # # #             try:
# # # # #                 task = """You are a Financial Expert and Wealth Advisor.
# # # # #                     You also a Stock Market Expert. You know everything about stock market trends and patterns.
# # # # #                     Provide financial advice or Stock Related advice and suggestions based on the user's query.
# # # # #                     Consider user's investment personality and Financial Details if provided.
# # # # #                     Address the user by their name(client_name: Emily in our case but if any other name is give refer to that) if provided.
# # # # #                     Include detailed information about the investment, where to invest, how much to invest, 
# # # # #                     expected returns, and why you are giving this advice.
# # # # #                     As you are a Wealth Advisor if user asks queries related to saving taxes or calculating taxes refer to the 
# # # # #                     US Tax Laws given by the IRS and based on that information calculate the taxes for the user 
# # # # #                     consider the information shared by the user such as their annual income and their monthly investment if provided,
# # # # #                     also give advice to the user on how they can save their taxes.
# # # # #                     Include a disclaimer to monitor and rebalance investments regularly based on risk tolerance."""
                
# # # # #                 # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + message.text
# # # # #                 # query = chat_history_text + "\n" + query

# # # # #                 # Include chat history
# # # # #                 chat_history = read_chat_history(chat_id)
# # # # #                 chat_history_text = '\n'.join([f"{entry['role']}: {entry['message']}" for entry in chat_history])
# # # # #                 # history.append(chat_history_text)
# # # # #                 query = task + "\n" + investment_personality + "\n" + chat_history_text + "\n" + message.text

# # # # #                 model = genai.GenerativeModel('gemini-1.5-flash')
# # # # #                 chat = model.start_chat(history=[])
# # # # #                 response = chat.send_message(query)

# # # # #                 # Enhanced logging for debugging
# # # # #                 logging.info(f"Model response: {response}")
# # # # #                 format_response = markdown_to_text(response.text) #(response_text) #response.result

# # # # #                 # Store the response in chat history
# # # # #                 write_chat_history(chat_id, {'role': 'bot', 'message': format_response})
# # # # #                 await message.reply(format_response)

# # # # #             except Exception as e:
# # # # #                 logging.error(f"Error processing general chat message: {e}")
# # # # #                 await message.reply("Failed to process your request.")



# # # # # # markdown to text :
# # # # # def markdown_to_text(md):
# # # # #     # Simple conversion for markdown to plain text
# # # # #     md = md.replace('**', '')
# # # # #     md = md.replace('*', '')
# # # # #     md = md.replace('_', '')
# # # # #     md = md.replace('#', '')
# # # # #     md = md.replace('`', '')
# # # # #     return md.strip()


# # # # # from aiogram.types.input_file import BufferedInputFile
# # # # # from aiogram import BaseMiddleware
# # # # # # from aiogram.dispatcher.router import Router
# # # # # from PIL import Image

# # # # # # Function to handle image messages
# # # # # # @dp.message(F.photo)
# # # # # # @router.message(F.photo)
# # # # # import PIL.Image

# # # # # async def handle_image(message: types.Message):
# # # # #     global investment_personality, chat_history

# # # # #     chat_id = message.chat.id
# # # # #     # Handle image inputs
# # # # #     try:
# # # # #         # Obtain file information
# # # # #         try:
# # # # #             photo_id = message.document.file_id
# # # # #             photo = await bot.get_file(photo_id)
# # # # #             photo_path = photo.file_path
# # # # #             # Download the file
# # # # #             photo_file = await bot.download_file(photo_path, "data/uploaded_image.png")

# # # # #         except Exception as e:
# # # # #             print(f"Error downloading image: {e}")
# # # # #             await bot.send_message(chat_id, "Error processing image. Please try again.")
# # # # #             return
        
# # # # #         # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # #         #             Also give the user detailed information about the investment how to invest, where to invest and how much they
# # # # #         #             should invest in terms of percentage of their investment amount. Give the user detailed information about the returns on their 
# # # # #         #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compounded returns on their 
# # # # #         #             investment. Also explain the user why you are giving them that particular
# # # # #         #             investment suggestion. Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # #         #             User should also invest as per their risk tolerance level. Since you are the financial advisor don't ask user to consult anyone else.
# # # # #         #             So don't mention user to consult to a financial expert."""

# # # # #         task = """You are a Financial Expert.You will be provided with a Financial Form from Boston Harbor.
# # # # #                 If you recieve any other image tell the user to Upload the Images of the form or upload the word document of the form.
# # # # #                 You are supposed to Respond to the user's Image query and If they ask for any information provide them the information in Detail.
# # # # #                 Be helpful and informative.Give proper information of any Financial terms the user may ask you.Address the user by their Client Name if provided.
# # # # #                 Also provide the user helpful links so that they can refer to the link for more information.
# # # # #                 If the image provided is not related to Finance then just answer about the image and any caption if provided.
# # # # #                 """

# # # # #         prompt = message.caption if message.caption else ""  # Use the photo caption if available
# # # # #         # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + prompt
# # # # #         query = task + prompt 

# # # # #         image =  PIL.Image.open('data/uploaded_image.png') #(photo_file) 
# # # # #         model = genai.GenerativeModel('gemini-1.5-flash')
# # # # #         response = model.generate_content(image)
# # # # #         await bot.send_message(chat_id,"I will describe the image that was uploaded")
# # # # #         format_response = markdown_to_text(response.text)
# # # # #         await message.reply(format_response)
# # # # #         # await message.reply(response.text)

# # # # #         # chat = model.start_chat(history=[])
# # # # #         # response = chat.send_message(query)
# # # # #         # format_response = markdown_to_text(response.result)
# # # # #         # await message.reply(format_response)

# # # # #         response = model.generate_content([query, image])
# # # # #         format_response = markdown_to_text(response.text)

# # # # #         # Store the response in chat history
# # # # #         write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

# # # # #         await message.reply(format_response)
# # # # #         # await message.reply(response.text) 
# # # # #     except Exception as e:
# # # # #         logging.error(f"Error generating response for the image: {e}")
# # # # #         await message.reply("There was an error generating response for the image. Please try again later.")
# # # # #     # await message.reply("Cant process the image")
# # # # #     # return



# # # # # from aiogram.filters import command
# # # # # from aiogram.types import bot_command
# # # # # import markdown
# # # # # from bs4 import BeautifulSoup

# # # # # def markdown_to_text(markdown_text):
# # # # #     # Convert markdown to HTML
# # # # #     html = markdown.markdown(markdown_text)
# # # # #     # Parse the HTML
# # # # #     soup = BeautifulSoup(html, 'html.parser')
# # # # #     # Extract plain text
# # # # #     text = soup.get_text()
# # # # #     return text


# # # # # # if __name__ == "__main__":
# # # # # #     executor.start_polling(dispatcher, skip_updates=True)

# # # # # async def main() -> None:
# # # # #     # Initialize Bot instance with default bot properties which will be passed to all API calls
# # # # #     bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

# # # # #     # And the run events dispatching
# # # # #     await dp.start_polling(bot)


# # # # # if __name__ == "__main__":
# # # # #     logging.basicConfig(level=logging.INFO, stream=sys.stdout)
# # # # #     asyncio.run(main())





# # # # # # #best code so far now it can work efficiently on Client side :

# # # # # # import os
# # # # # # import filetype
# # # # # # import docx
# # # # # # import PyPDF2
# # # # # # import re
# # # # # # from aiogram import Bot, Dispatcher, types
# # # # # # from dotenv import load_dotenv
# # # # # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # # # # from langchain_community.vectorstores import Chroma

# # # # # # import faiss
# # # # # # from langchain_community.docstore.in_memory import InMemoryDocstore
# # # # # # from langchain_community.vectorstores import FAISS
# # # # # # from langchain_community.document_loaders import Docx2txtLoader

# # # # # # from langchain_core.prompts import ChatPromptTemplate
# # # # # # from langchain.chains import create_retrieval_chain
# # # # # # from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
# # # # # # from langchain.chains.combine_documents import create_stuff_documents_chain
# # # # # # from langchain.memory import ConversationSummaryMemory
# # # # # # import asyncio
# # # # # # import numpy as np
# # # # # # import json
# # # # # # import re
# # # # # # import google.generativeai as genai
# # # # # # import pathlib
# # # # # # # Import things that are needed generically
# # # # # # from langchain.pydantic_v1 import BaseModel, Field
# # # # # # from langchain.tools import BaseTool, StructuredTool, tool

# # # # # # from aiogram.client.default import DefaultBotProperties
# # # # # # from aiogram.enums import ParseMode
# # # # # # from aiogram.filters import CommandStart
# # # # # # from aiogram.types import Message
# # # # # # from aiogram import F
# # # # # # from aiogram import Router
# # # # # # import logging
# # # # # # import sys
# # # # # # from aiogram.filters import Command
# # # # # # from aiogram.types import FSInputFile
# # # # # # # from aiogram.utils import executor
# # # # # # import io
# # # # # # import matplotlib.pyplot as plt
# # # # # # import seaborn as sns
# # # # # # import aiohttp
# # # # # # from aiogram.types import InputFile , BufferedInputFile
# # # # # # import PIL.Image

# # # # # # router = Router(name=__name__)

# # # # # # load_dotenv()

# # # # # # TOKEN = os.getenv("TOKEN")
# # # # # # GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# # # # # # os.environ["LANGCHAIN_TRACING_V2"]="true"
# # # # # # os.environ["LANGCHAIN_API_KEY"]=os.getenv("LANGCHAIN_API_KEY")

# # # # # # # Configure generativeai with your API key
# # # # # # genai.configure(api_key=GOOGLE_API_KEY)

# # # # # # # Initialize bot
# # # # # # bot = Bot(token=TOKEN)
# # # # # # dp = Dispatcher()

# # # # # # # Glbal variables
# # # # # # rag_on = False
# # # # # # retriever = None  # Store retriever globally
# # # # # # summary = ""
# # # # # # investment_personality = ""
# # # # # # # history = []
# # # # # # previous_suggestions = ""

# # # # # # CHAT_HISTORY_FILE = 'chat_history.json'

# # # # # # def read_chat_history(chat_id):
# # # # # #     if os.path.exists(CHAT_HISTORY_FILE):
# # # # # #         with open(CHAT_HISTORY_FILE, 'r') as file:
# # # # # #             chat_history = json.load(file)
# # # # # #             return chat_history.get(str(chat_id), [])
# # # # # #     return []

# # # # # # def write_chat_history(chat_id, message):
# # # # # #     chat_history = {}
# # # # # #     if os.path.exists(CHAT_HISTORY_FILE):
# # # # # #         with open(CHAT_HISTORY_FILE, 'r') as file:
# # # # # #             chat_history = json.load(file)
# # # # # #     if str(chat_id) not in chat_history:
# # # # # #         chat_history[str(chat_id)] = []
# # # # # #     chat_history[str(chat_id)].append(message)
# # # # # #     with open(CHAT_HISTORY_FILE, 'w') as file:
# # # # # #         json.dump(chat_history, file)

# # # # # # class Reference:
# # # # # #     def __init__(self):
# # # # # #         self.response = ""


# # # # # # reference = Reference()


# # # # # # def clear_past():
# # # # # #     reference.response = ""


# # # # # # @router.message(F.text == "clear")
# # # # # # async def clear(message: types.Message):
# # # # # #     """
# # # # # #     A handler to clear the previous conversation and context.
# # # # # #     """
# # # # # #     clear_past()
# # # # # #     await message.reply("I've cleared the past conversation and context.")

# # # # # # #Global Variables :

# # # # # # # Store user states
# # # # # # states = {}

# # # # # # # Dictionary to hold question-answer pairs
# # # # # # user_responses = {}
# # # # # # #
# # # # # # user_images = {}
# # # # # # # Define Questions for assessment
# # # # # # questions = [
# # # # # #     """ 
# # # # # #     1. You and your friend are betting on a series of coin tosses.

# # # # # #     He always bets ₹2,000 on Heads

# # # # # #     You always bet ₹2,000 on Tails

# # # # # #     Winner of last 8 turns

# # # # # #     You lost ₹8,000 in the last 4 turns!

# # # # # #     If you were to bet one last time, what would you bet on:
# # # # # #     a) heads or b) tails ?
# # # # # #     """ ,
# # # # # #     """
# # # # # #     2. Imagine you are a contestant in a game show, and you are presented the following choices.

# # # # # #     What would you prefer?
# # # # # #     a) 50 percent chance of winning 15 gold coins 
# # # # # #     b) 100 percent chance of winning 8 gold coins
# # # # # #     """,
# # # # # #     """
# # # # # #     3. In general, how would your best friend describe your risk-taking tendencies?
# # # # # #     a) A real gambler
# # # # # #     b) Willing to take risks after completing adequate research
# # # # # #     c) Cautious
# # # # # #     d) Avoids risk as much as possible
# # # # # #     """,
# # # # # #     """
# # # # # #     4. Suppose you could replace your current investment portfolio with this new one:
# # # # # #     50 percent chance of Gaining 35 percent or 50 percent chance of Loss
# # # # # #     In order to have a 50 percent chance of gaining +35 percent, how much loss are you willing to take?
# # # # # #     a)-5 to -10
# # # # # #     b)-10 to -15
# # # # # #     c)-15 to -20
# # # # # #     d)-20 to -25
# # # # # #     e)-25 to -30
# # # # # #     f)-30 to -35
# # # # # #     """,
# # # # # #     """
# # # # # #     5. Over any 1-year period, what would be the maximum drop in the value of your investment 
# # # # # #     portfolio that you would be comfortable with?
# # # # # #     a) <5%
# # # # # #     b) 5 - 10%
# # # # # #     c) 10 - 15%
# # # # # #     d) 15 - 20%
# # # # # #     e) >20%
# # # # # #     """,
# # # # # #     """
# # # # # #     6. When investing, what do you consider the most?

# # # # # #     a) Risk 
# # # # # #     b) Return
# # # # # #     """,
# # # # # #     """
# # # # # #     7. What best describes your attitude?

# # # # # #     a) Prefer reasonable returns, can take reasonable risk
# # # # # #     b) Like higher returns, can take slightly higher risk
# # # # # #     c) Want to maximize returns, can take significant high risk
# # # # # #     """,
# # # # # #     """
# # # # # #     8. How much monthly investment you want to do?
# # # # # #     """,
# # # # # #     """
# # # # # #     9. What is the time horizon for your investment?
# # # # # #     You can answer in any range, example 1-5 years."""  
# # # # # # ]


# # # # # # # Handler for /start command
# # # # # # @dp.message(CommandStart())
# # # # # # async def handle_start(message: types.Message):
# # # # # #     """
# # # # # #     This handler receives messages with /start command
# # # # # #     """
# # # # # #     chat_id = message.chat.id
# # # # # #     # Start asking questions
# # # # # #     await start_assessment(chat_id)


# # # # # # # Function to start the assessment
# # # # # # async def start_assessment(chat_id):
# # # # # #     await bot.send_message(chat_id, "Hi,My name is Finbot and I am a Wealth Management Advisor ChatBot! Let's start a quick personality assessment.")
# # # # # #     await ask_next_question(chat_id, 0)

# # # # # # # Function to ask the next question
# # # # # # async def ask_next_question(chat_id, question_index):
# # # # # #     if question_index < len(questions):
# # # # # #         # Ask the next question
# # # # # #         await bot.send_message(chat_id, questions[question_index])
# # # # # #         # Update state to indicate the next expected answer
# # # # # #         states[chat_id] = question_index
# # # # # #     else:
# # # # # #         # No more questions, finish assessment
# # # # # #         await finish_assessment(chat_id)

# # # # # # # Handler for receiving assessment answers
# # # # # # assessment_in_progress = True

# # # # # # from aiogram.types import FSInputFile
# # # # # # async def finish_assessment(chat_id):
# # # # # #     if chat_id in states and states[chat_id] == len(questions):
# # # # # #         # All questions have been answered, now process the assessment
# # # # # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # # # # #         # Determine investment personality based on collected responses
# # # # # #         global investment_personality
# # # # # #         investment_personality = await determine_investment_personality(user_responses)

# # # # # #         # Inform the user about their investment personality
# # # # # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # # # # #         # Store the response in chat history
# # # # # #         write_chat_history(chat_id, {'role': 'bot', 'message': investment_personality})

# # # # # #         # Summarize collected information
# # # # # #         global summary
# # # # # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # # # # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # # # # #         # Ensure to await the determination of investment personality
# # # # # #         await send_summary_chunks(chat_id, summary)
# # # # # #         global assessment_in_progress 
# # # # # #         assessment_in_progress = False
       
# # # # # #         await bot.send_message(chat_id,"Hello there here is the Word Document.Please fill in your details with correct Information and then upload it in the chat")
# # # # # #         file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")

# # # # # #         await bot.send_document(chat_id, document=file, caption="To receive financial advice,Please fill in the Financial Details in this document with correct information and upload it here.")
# # # # # #         # await bot.send_message(chat_id,file)

# # # # # # async def send_summary_chunks(chat_id, summary):
# # # # # #     # Split the summary into chunks that fit within Telegram's message limits
# # # # # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]

# # # # # #     # Send each chunk as a separate message
# # # # # #     for chunk in chunks:
# # # # # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")


# # # # # # async def determine_investment_personality(assessment_data):
# # # # # #     try:
# # # # # #         # Prepare input text for the chatbot based on assessment data
# # # # # #         input_text = "User Profile:\n"
# # # # # #         for question, answer in assessment_data.items():
# # # # # #             input_text += f"{question}: {answer}\n"

# # # # # #         # Introduce the chatbot's task and prompt for classification
# # # # # #         input_text += "\nYou are an investment personality identifier. Classify the user as:\n" \
# # # # # #                       "- Conservative Investor\n" \
# # # # # #                       "- Moderate Investor\n" \
# # # # # #                       "- Aggressive Investor"

# # # # # #         # Use your generative AI model to generate a response
# # # # # #         # print(input_text)
# # # # # #         model = genai.GenerativeModel('gemini-pro')
# # # # # #         response = model.generate_content(input_text)

# # # # # #         # Determine the investment personality from the chatbot's response
# # # # # #         response_text = response.text.lower()
# # # # # #         if "conservative" in response_text:
# # # # # #             personality = "Conservative Investor"
# # # # # #         elif "moderate" in response_text:
# # # # # #             personality = "Moderate Investor"
# # # # # #         elif "aggressive" in response_text:
# # # # # #             personality = "Aggressive Investor"
# # # # # #         else:
# # # # # #             personality = "Unknown"

# # # # # #         return personality
# # # # # #         # Send the determined investment personality back to the user
# # # # # #         #await bot.send_message(chat_id, f"Investment Personality: {personality}")

# # # # # #     except Exception as e:
# # # # # #         print(f"Error generating response: {e}")
# # # # # #         #await bot.send_message(chat_id, "Error processing investment personality classification.")


# # # # # # # Handler for document upload
# # # # # # async def load_vector_db(file_path):
# # # # # #     try:
# # # # # #         print("Loading vector database...")
# # # # # #         loader = Docx2txtLoader(file_path)
# # # # # #         documents = loader.load()
# # # # # #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # # # # #         text_chunks = text_splitter.split_documents(documents)
# # # # # #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
# # # # # #         # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
# # # # # #         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
# # # # # #         # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

# # # # # #         # vector_store = FAISS(
# # # # # #         #     embedding_function=embeddings,
# # # # # #         #     index=index,
# # # # # #         #     docstore=InMemoryDocstore(),
# # # # # #         #     index_to_docstore_id={},
# # # # # #         # )
        
# # # # # #         print("Vector database loaded successfully.") 
# # # # # #         return vector_store.as_retriever(search_kwargs={"k": 1})
# # # # # #     except Exception as e:
# # # # # #         print(f"Error loading vector database: {e}")
# # # # # #         return None


# # # # # # # change prompt template :
# # # # # # async def make_retrieval_chain(retriever):
# # # # # #     """
# # # # # #     Create a retrieval chain using the provided retriever.

# # # # # #     Args:
# # # # # #         retriever (RetrievalQA): A retriever object.

# # # # # #     Returns:
# # # # # #         RetrievalQA: A retrieval chain object.
# # # # # #     """
# # # # # #     try:
# # # # # #         global investment_personality,summary
# # # # # #         llm = ChatGoogleGenerativeAI(
# # # # # #             #model="gemini-pro",
# # # # # #             model = "gemini-1.5-flash",
# # # # # #             temperature=0.7,
# # # # # #             top_p=0.85,
# # # # # #             google_api_key=GOOGLE_API_KEY
# # # # # #         )

# # # # # #         # prompt_template = investment_personality + "\n" + summary + "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
# # # # # #         #         Respond to the client by the client name.
# # # # # #         #         Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #         #         Also give the user detailed information about the investment how to invest,where to invest and how much they
# # # # # #         #         should invest in terms of percentage of their investment amount.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
# # # # # #         #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
# # # # # #         #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon.
# # # # # #         #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
# # # # # #         #         Also explain the user why you are giving them that particular investment suggestion.
# # # # # #         #         Give the client suggestions of Investment based on their investment personality that can also help them manage their mortages and other liablilities if they have any,ignore if they dont have any liabilities.
# # # # # #         #         Answer in 3-4 lines.\n
# # # # # #         #         <context>
# # # # # #         #         {context}
# # # # # #         #         </context>
# # # # # #         #         Question: {input}"""


# # # # # #         prompt_template = investment_personality + "\n" + summary + "\n" + """ Role : You are a Top class highly professional and world's best Savings Advisor for 
# # # # # #                 savings related question-answering tasks related to the document.
# # # # # #                 Respond to the client by the client name.
# # # # # #                 Give Savings Suggestions to the user so that they could do proper responsible savings and save their expenses based on their investment personality and budget if provided.
# # # # # #                 Also give the user detailed information about their savings such that they could save more money and save their expenses.
# # # # # #                 Give the user minimum and maximum percentage of savings the user can do by reducing their expenses. If the users have given a budget then analyse it and give suggestions based on that.
# # # # # #                 Try to imitate human language and talk to the user/client like a human and give personal savings suggestions.
# # # # # #                 If the user is having many unnecessary expenses then give the user some advice in a gentle manner without offending the user or hurt their feelings and suggest and advice them to stop or reduce their unnecessary expenses 
# # # # # #                 in order to increase their savings.
# # # # # #                 Also explain the user why you are giving them that particular savings suggestion.
# # # # # #                 Give the client suggestions of savings based on their investment personality that can also help them manage their mortages and other liablilities if they have any,ignore if they dont have any liabilities.
# # # # # #                 Answer in 3-4 lines.\n
# # # # # #                 <context>
# # # # # #                 {context}
# # # # # #                 </context>
# # # # # #                 Question: {input}"""

# # # # # #         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

# # # # # #         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
# # # # # #         combine_docs_chain = None  

# # # # # #         if retriever is not None :  
# # # # # #             retriever_chain = create_retrieval_chain(retriever,document_chain) 
# # # # # #             print(retriever_chain)
# # # # # #             return retriever_chain
# # # # # #         else:
# # # # # #             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
# # # # # #             return None

# # # # # #     except Exception as e:
# # # # # #         print(f"Error in creating chain: {e}")
# # # # # #         return None

# # # # # # from aiogram.filters import Filter

# # # # # # # @router.message(F.document)
# # # # # # @dp.message(F.document)
# # # # # # async def handle_document(message: types.Message):
# # # # # #     global summary,investment_personality  

# # # # # #     chat_id = message.chat.id
# # # # # #     await message.reply("File Received") 
# # # # # #     # Obtain file information
# # # # # #     file_id = message.document.file_id
# # # # # #     file = await bot.get_file(file_id)
# # # # # #     file_path = file.file_path
    
# # # # # #     # Download the file
# # # # # #     await bot.download_file(file_path, "data/uploaded_file")
    
# # # # # #     # Process the uploaded document
# # # # # #     extracted_text = await process_document("data/uploaded_file")
# # # # # #     # print(extracted_text)

# # # # # #     if extracted_text:
# # # # # #         # Load vector database (assuming this is part of setting up the retriever)
# # # # # #         print("Retriever being loaded ")
# # # # # #         retriever = await load_vector_db("data/uploaded_file")
# # # # # #         file_path = 'data/uploaded_file'
# # # # # #         client_name, validation_errors = await validate_process_document(file_path)

# # # # # #         # Print results
# # # # # #         print(f"Client Name: {client_name}")
# # # # # #         if validation_errors:
# # # # # #             print("**Validation Errors:**")
# # # # # #             for error in validation_errors:
# # # # # #                 print(error)
# # # # # #         else:
# # # # # #             print("All fields are filled correctly.")
# # # # # #         if client_name == None:
# # # # # #             try:
# # # # # #                 await message.reply("Processing the uploaded image")
# # # # # #                 await handle_image(message) 
# # # # # #                 return 
# # # # # #             except Exception as e:
# # # # # #                 await message.reply("error processing uploaded image")
# # # # # #                 print(e)
# # # # # #         await message.reply(f"Thanks for providing me the details, {client_name}.I have processed the file and now I will provide you some Savings suggestions based on the details that you have provided.")

# # # # # #         if retriever is None:
# # # # # #             await message.reply("The retrieval chain is not set up. Please upload a document first.")
# # # # # #             return

# # # # # #         # Check if a valid chain can be created
# # # # # #         chain = await make_retrieval_chain(retriever)
# # # # # #         if chain is None:
# # # # # #             await message.reply("Failed to create the retrieval chain.")
# # # # # #             return
        
# # # # # #         try:     
# # # # # #             query = summary + "\n" + investment_personality # + "\n"  #extracted_text + "\n" + task
        
# # # # # #             response = chain.invoke({"input": query})
# # # # # #             print(response['answer'])
# # # # # #             global chat_history
# # # # # #             chat_history = response['answer'] 
# # # # # #             print(f"\n Chat History : {chat_history}")
# # # # # #             format_response = markdown_to_text(response['answer'])

# # # # # #             # Store the extracted_text in chat history
# # # # # #             write_chat_history(chat_id, {'role': 'bot', 'message': extracted_text})
        
# # # # # #             # Store the response in chat history
# # # # # #             write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

# # # # # #             await message.reply(format_response)
# # # # # #             # await message.reply(response['answer'])

# # # # # #         except Exception as e:
# # # # # #             print(f"Error invoking retrieval chain on attempt : {e}")
# # # # # #             await bot.send_message(chat_id, "Error invoking retrieval. Please try again later.")

# # # # # #     else:
# # # # # #         await message.reply("Failed to process the uploaded file.")
    

# # # # # # # Function to extract data from LLM response
# # # # # # def extract_data_from_response(response):
# # # # # #     try:
# # # # # #         # Locate the JSON-like data in the response
# # # # # #         json_start = response.find("{")
# # # # # #         json_end = response.rfind("}") + 1
        
# # # # # #         if json_start == -1 or json_end == -1:
# # # # # #             raise ValueError("No JSON data found in the response.")
        
# # # # # #         json_data = response[json_start:json_end]
        
# # # # # #         # Parse the JSON data
# # # # # #         data = json.loads(json_data.replace("'", "\""))
# # # # # #         print(data)
# # # # # #         return data
# # # # # #     except Exception as e:
# # # # # #         logging.error(f"Error extracting data: {e}")
# # # # # #         return None

 


# # # # # # def extract_allocations_from_json(json_data,chat_id):
# # # # # #     allocations = {}
# # # # # #     for entry in json_data.get(str(chat_id), []):
# # # # # #         if entry['role'] == 'bot':
# # # # # #             message = entry['message']
# # # # # #             lines = message.split('\n')
# # # # # #             current_category = None

# # # # # #             for line in lines:
# # # # # #                 match = re.match(r'^(.*?):\s*(\d+)%$', line)
# # # # # #                 if match:
# # # # # #                     category, percent = match.groups()
# # # # # #                     allocations[category] = []
# # # # # #                     current_category = category
# # # # # #                 elif current_category and re.match(r'.*\d+%', line):
# # # # # #                     subcategory_match = re.match(r'^(.*?)(\d+)%$', line)
# # # # # #                     if subcategory_match:
# # # # # #                         subcategory, percent = subcategory_match.groups()
# # # # # #                         allocations[current_category].append((subcategory.strip(), float(percent)))

# # # # # #     return allocations


# # # # # # def create_pie_chart(allocations, chat_id):
# # # # # #     labels = []
# # # # # #     sizes = []
# # # # # #     for category, subcategories in allocations.items():
# # # # # #         for subcategory, percent in subcategories:
# # # # # #             labels.append(f"{category} - {subcategory}")
# # # # # #             sizes.append(percent)
    
# # # # # #     if sizes:
# # # # # #         fig, ax = plt.subplots()
# # # # # #         ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
# # # # # #         ax.axis('equal')
        
# # # # # #         plt.title("Investment Allocation")
# # # # # #         chart_path = f"data/investment_allocation_{chat_id}.png"
# # # # # #         plt.savefig(chart_path)
# # # # # #         plt.close()
        
# # # # # #         return chart_path
# # # # # #     else:
# # # # # #         return None
  

# # # # # # async def process_document(file_path):
# # # # # #     try:
# # # # # #         print("Processing the document")
# # # # # #         file_type = filetype.guess(file_path)
# # # # # #         if file_type is not None:
# # # # # #             if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # # # #                 # return extract_text_from_word(file_path)
# # # # # #                 return extract_text_and_tables_from_word(file_path)
# # # # # #             elif file_type.mime == "application/pdf":
# # # # # #                 return extract_text_from_pdf(file_path)
# # # # # #         return None
# # # # # #     except Exception as e:
# # # # # #         print(f"Error processing document: {e}")
# # # # # #         return None

# # # # # # def extract_text_from_pdf(pdf_file_path):
# # # # # #     try:
# # # # # #         print("Processing pdf file")
# # # # # #         with open(pdf_file_path, "rb") as pdf_file:
# # # # # #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# # # # # #             text_content = []
# # # # # #             for page_num in range(pdf_reader.numPages):
# # # # # #                 page = pdf_reader.getPage(page_num)
# # # # # #                 text_content.append(page.extract_text())
# # # # # #             return "\n".join(text_content)
# # # # # #     except Exception as e:
# # # # # #         print(f"Error extracting text from PDF: {e}")
# # # # # #         return None


# # # # # # import re
# # # # # # import docx

# # # # # # def extract_text_and_tables_from_word(docx_file_path):
# # # # # #     """
# # # # # #     Extracts text and tables from a Word document (.docx).

# # # # # #     Args:
# # # # # #         docx_file_path (str): Path to the Word document file.

# # # # # #     Returns:
# # # # # #         tuple: Extracted text content and tables from the document.
# # # # # #     """
# # # # # #     try:
# # # # # #         print("Extracting text and tables from word file")
# # # # # #         doc = docx.Document(docx_file_path)
# # # # # #         text_content = []
# # # # # #         tables_content = []

# # # # # #         for para in doc.paragraphs:
# # # # # #             text_content.append(para.text)

# # # # # #         for table in doc.tables:
# # # # # #             table_data = []
# # # # # #             for row in table.rows:
# # # # # #                 row_data = []
# # # # # #                 for cell in row.cells:
# # # # # #                     row_data.append(cell.text.strip())
# # # # # #                 table_data.append(row_data)
# # # # # #             tables_content.append(table_data)
# # # # # #         print("Extracted text from word file")
# # # # # #         return "\n".join(text_content), tables_content
# # # # # #     except Exception as e:
# # # # # #         print(f"Error extracting text and tables from Word document: {e}")
# # # # # #         return None, None

# # # # # # def validate_document_content(text, tables):
# # # # # #     """
# # # # # #     Validates the content of the document.

# # # # # #     Args:
# # # # # #         text (str): Extracted text content from the document.
# # # # # #         tables (list): Extracted tables content from the document.

# # # # # #     Returns:
# # # # # #         tuple: Client name and validation errors.
# # # # # #     """
# # # # # #     errors = []
    
# # # # # #     # Extract client name
# # # # # #     client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
# # # # # #     client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

# # # # # #     # Define required sections
# # # # # #     required_sections = [
# # # # # #         "YOUR RETIREMENT GOAL",
# # # # # #         "YOUR OTHER MAJOR GOALS",
# # # # # #         "YOUR ASSETS AND LIABILITIES",
# # # # # #         "MY LIABILITIES",
# # # # # #         "YOUR CURRENT ANNUAL INCOME"
# # # # # #     ]

# # # # # #     # Check for the presence of required sections
# # # # # #     for section in required_sections:
# # # # # #         if section not in text:
# # # # # #             errors.append(f"* {section} section missing.")
    
# # # # # #     # Define table field checks
# # # # # #     table_checks = {
# # # # # #         "YOUR RETIREMENT GOAL": [
# # # # # #             r"When do you plan to retire\? \(age or date\)",
# # # # # #             r"Social Security Benefit \(include expected start date\)",
# # # # # #             r"Pension Benefit \(include expected start date\)",
# # # # # #             r"Other Expected Income \(rental, part-time work, etc.\)",
# # # # # #             r"Estimated Annual Retirement Expense"
# # # # # #         ],
# # # # # #         "YOUR OTHER MAJOR GOALS": [
# # # # # #             r"GOAL", r"COST", r"WHEN"
# # # # # #         ],
# # # # # #         "YOUR ASSETS AND LIABILITIES": [
# # # # # #             r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
# # # # # #             r"Current Value", r"Annual Contributions"
# # # # # #         ],
# # # # # #         "MY LIABILITIES": [
# # # # # #             r"Balance", r"Interest Rate", r"Monthly Payment"
# # # # # #         ]
# # # # # #     }

# # # # # #     # Validate table content
# # # # # #     for section, checks in table_checks.items():
# # # # # #         section_found = False
# # # # # #         for table in tables:
# # # # # #             table_text = "\n".join(["\t".join(row) for row in table])
# # # # # #             if section in table_text:
# # # # # #                 section_found = True
# # # # # #                 for check in checks:
# # # # # #                     if not re.search(check, table_text, re.IGNORECASE):
# # # # # #                         errors.append(f"* Missing or empty field in {section} section: {check}")
# # # # # #                 break
# # # # # #         if not section_found:
# # # # # #             errors.append(f"* {section} section missing.")

# # # # # #     return client_name, errors

# # # # # # async def validate_process_document(file_path):
# # # # # #     try:
# # # # # #         print("Validating process document : ")
# # # # # #         text, tables = extract_text_and_tables_from_word(file_path)
# # # # # #         if text is not None and tables is not None:
# # # # # #             client_name, errors = validate_document_content(text, tables)
# # # # # #             return client_name, errors
# # # # # #         return None, ["Error processing document."]
# # # # # #     except Exception as e:
# # # # # #         print(f"Error processing document: {e}")
# # # # # #         return None, [f"Error processing document: {e}"]

# # # # # # @dp.message()
# # # # # # async def main_bot(message: types.Message):
# # # # # #     global retriever, extracted_text, investment_personality, summary, chat_history

# # # # # #     chat_id = message.chat.id

# # # # # #     if chat_id in states and states[chat_id] < len(questions):
# # # # # #         question_index = states[chat_id]
# # # # # #         answer = message.text
# # # # # #         user_responses[questions[question_index]] = answer
# # # # # #         states[chat_id] += 1
# # # # # #         await ask_next_question(chat_id, question_index + 1)
# # # # # #     elif message.text:
# # # # # #         try:
# # # # # #             task = """You are a Financial Expert and Wealth Advisor.
# # # # # #                 You also a Stock Market Expert. You know everything about stock market trends and patterns.
# # # # # #                 Provide financial advice or Stock Related advice and suggestions based on the user's query.
# # # # # #                 Consider user's investment personality and Financial Details if provided.
# # # # # #                 Address the user by their name(client_name: Emily in our case but if any other name is give refer to that) if provided.
# # # # # #                 Include detailed information about the investment, where to invest, how much to invest, 
# # # # # #                 expected returns, and why you are giving this advice.
# # # # # #                 As you are a Wealth Advisor if user asks queries related to saving taxes or calculating taxes refer to the 
# # # # # #                 US Tax Laws given by the IRS and based on that information calculate the taxes for the user 
# # # # # #                 consider the information shared by the user such as their annual income and their monthly investment if provided,
# # # # # #                 also give advice to the user on how they can save their taxes.
# # # # # #                 Include a disclaimer to monitor and rebalance investments regularly based on risk tolerance."""
            
# # # # # #             # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + message.text
# # # # # #             # query = chat_history_text + "\n" + query

# # # # # #             # Include chat history
# # # # # #             chat_history = read_chat_history(chat_id)
# # # # # #             chat_history_text = '\n'.join([f"{entry['role']}: {entry['message']}" for entry in chat_history])
# # # # # #             # history.append(chat_history_text)
# # # # # #             query = task + "\n" + investment_personality + "\n" + chat_history_text + "\n" + message.text

# # # # # #             model = genai.GenerativeModel('gemini-1.5-flash')
# # # # # #             chat = model.start_chat(history=[])
# # # # # #             response = chat.send_message(query)

# # # # # #             # Enhanced logging for debugging
# # # # # #             logging.info(f"Model response: {response}")
# # # # # #             format_response = markdown_to_text(response.text) #(response_text) #response.result

# # # # # #             # Store the response in chat history
# # # # # #             write_chat_history(chat_id, {'role': 'bot', 'message': format_response})
# # # # # #             await message.reply(format_response)

# # # # # #         except Exception as e:
# # # # # #             logging.error(f"Error processing general chat message: {e}")
# # # # # #             await message.reply("Failed to process your request.")



# # # # # # # markdown to text :
# # # # # # def markdown_to_text(md):
# # # # # #     # Simple conversion for markdown to plain text
# # # # # #     md = md.replace('**', '')
# # # # # #     md = md.replace('*', '')
# # # # # #     md = md.replace('_', '')
# # # # # #     md = md.replace('#', '')
# # # # # #     md = md.replace('`', '')
# # # # # #     return md.strip()


# # # # # # from aiogram.types.input_file import BufferedInputFile
# # # # # # from aiogram import BaseMiddleware
# # # # # # # from aiogram.dispatcher.router import Router
# # # # # # from PIL import Image

# # # # # # # Function to handle image messages
# # # # # # # @dp.message(F.photo)
# # # # # # # @router.message(F.photo)
# # # # # # import PIL.Image

# # # # # # async def handle_image(message: types.Message):
# # # # # #     global investment_personality, chat_history

# # # # # #     chat_id = message.chat.id
# # # # # #     # Handle image inputs
# # # # # #     try:
# # # # # #         # Obtain file information
# # # # # #         try:
# # # # # #             photo_id = message.document.file_id
# # # # # #             photo = await bot.get_file(photo_id)
# # # # # #             photo_path = photo.file_path
# # # # # #             # Download the file
# # # # # #             photo_file = await bot.download_file(photo_path, "data/uploaded_image.png")

# # # # # #         except Exception as e:
# # # # # #             print(f"Error downloading image: {e}")
# # # # # #             await bot.send_message(chat_id, "Error processing image. Please try again.")
# # # # # #             return
        
# # # # # #         # task = """Give Financial Suggestions to the user so that they could do proper responsible investment based on their investment personality.
# # # # # #         #             Also give the user detailed information about the investment how to invest, where to invest and how much they
# # # # # #         #             should invest in terms of percentage of their investment amount. Give the user detailed information about the returns on their 
# # # # # #         #             investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compounded returns on their 
# # # # # #         #             investment. Also explain the user why you are giving them that particular
# # # # # #         #             investment suggestion. Give user Disclaimer to keep monitoring their investments regularly and rebalance it if necessary.
# # # # # #         #             User should also invest as per their risk tolerance level. Since you are the financial advisor don't ask user to consult anyone else.
# # # # # #         #             So don't mention user to consult to a financial expert."""

# # # # # #         task = """You are a Financial Expert.You will be provided with a Financial Form from Boston Harbor.
# # # # # #                 If you recieve any other image tell the user to Upload the Images of the form or upload the word document of the form.
# # # # # #                 You are supposed to Respond to the user's Image query and If they ask for any information provide them the information in Detail.
# # # # # #                 Be helpful and informative.Give proper information of any Financial terms the user may ask you.Address the user by their Client Name if provided.
# # # # # #                 Also provide the user helpful links so that they can refer to the link for more information.
# # # # # #                 If the image provided is not related to Finance then just answer about the image and any caption if provided.
# # # # # #                 """

# # # # # #         prompt = message.caption if message.caption else ""  # Use the photo caption if available
# # # # # #         # query = task + "\n" + investment_personality + "\n" + chat_history + "\n" + prompt
# # # # # #         query = task + prompt 

# # # # # #         image =  PIL.Image.open('data/uploaded_image.png') #(photo_file) 
# # # # # #         model = genai.GenerativeModel('gemini-1.5-flash')
# # # # # #         response = model.generate_content(image)
# # # # # #         await bot.send_message(chat_id,"I will describe the image that was uploaded")
# # # # # #         format_response = markdown_to_text(response.text)
# # # # # #         await message.reply(format_response)
# # # # # #         # await message.reply(response.text)

# # # # # #         # chat = model.start_chat(history=[])
# # # # # #         # response = chat.send_message(query)
# # # # # #         # format_response = markdown_to_text(response.result)
# # # # # #         # await message.reply(format_response)

# # # # # #         response = model.generate_content([query, image])
# # # # # #         format_response = markdown_to_text(response.text)

# # # # # #         # Store the response in chat history
# # # # # #         write_chat_history(chat_id, {'role': 'bot', 'message': format_response})

# # # # # #         await message.reply(format_response)
# # # # # #         # await message.reply(response.text) 
# # # # # #     except Exception as e:
# # # # # #         logging.error(f"Error generating response for the image: {e}")
# # # # # #         await message.reply("There was an error generating response for the image. Please try again later.")
# # # # # #     # await message.reply("Cant process the image")
# # # # # #     # return



# # # # # # from aiogram.filters import command
# # # # # # from aiogram.types import bot_command
# # # # # # import markdown
# # # # # # from bs4 import BeautifulSoup

# # # # # # def markdown_to_text(markdown_text):
# # # # # #     # Convert markdown to HTML
# # # # # #     html = markdown.markdown(markdown_text)
# # # # # #     # Parse the HTML
# # # # # #     soup = BeautifulSoup(html, 'html.parser')
# # # # # #     # Extract plain text
# # # # # #     text = soup.get_text()
# # # # # #     return text


# # # # # # # if __name__ == "__main__":
# # # # # # #     executor.start_polling(dispatcher, skip_updates=True)

# # # # # # async def main() -> None:
# # # # # #     # Initialize Bot instance with default bot properties which will be passed to all API calls
# # # # # #     bot = Bot(token=TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))

# # # # # #     # And the run events dispatching
# # # # # #     await dp.start_polling(bot)


# # # # # # if __name__ == "__main__":
# # # # # #     logging.basicConfig(level=logging.INFO, stream=sys.stdout)
# # # # # #     asyncio.run(main())





# # # # # # previous polling mechanism :



# # # # # # # Function to ask the next question
# # # # # # async def ask_next_question(chat_id, question_index):
# # # # # #     if question_index < len(questions):
# # # # # #         question_data = questions[question_index]

# # # # # #         if question_data["options"]:  # If there are predefined options, send a poll
# # # # # #             poll = await bot.send_poll(
# # # # # #                 chat_id=chat_id,
# # # # # #                 question=question_data["question"],
# # # # # #                 options=question_data["options"],
# # # # # #                 type="regular",
# # # # # #                 is_anonymous=False
# # # # # #             )
# # # # # #             logging.info(f"Poll sent with ID {poll.poll.id} for question {question_index + 1}")
# # # # # #             user_responses[chat_id] = {"poll_id": poll.poll.id, "question_index": question_index}
# # # # # #             # question_index += 1
# # # # # #         else:  # If there are no predefined options, ask as a simple text question
# # # # # #             await bot.send_message(chat_id, question_data["question"])
# # # # # #             states[chat_id] = question_index

# # # # # #     else:
# # # # # #         # No more questions, finish assessment
# # # # # #         await finish_assessment(chat_id)

# # # # # # # Poll Answer handler
# # # # # # @router.poll_answer()
# # # # # # async def handle_assessment_poll_answer(poll_answer: PollAnswer):
# # # # # #     logging.info(f"Poll answer received from user {poll_answer.user.id}: selected option {poll_answer.option_ids[0]}")
# # # # # #     chat_id = poll_answer.user.id
# # # # # #     question_index = user_responses[chat_id]["question_index"]
    
# # # # # #     selected_option = questions[question_index]["options"][poll_answer.option_ids[0]]
# # # # # #     user_responses[chat_id][f"answer_{question_index + 1}"] = selected_option

# # # # # #     # Proceed to the next question
# # # # # #     await ask_next_question(chat_id, question_index + 1)

# # # # # # async def finish_assessment(chat_id):
# # # # # #     if chat_id in states and states[chat_id] == len(questions):
# # # # # #         await bot.send_message(chat_id, "Assessment completed. Thank you!")

# # # # # #         # Process and determine investment personality
# # # # # #         investment_personality = await determine_investment_personality(user_responses)

# # # # # #         # Inform the user about their investment personality
# # # # # #         await bot.send_message(chat_id, f"Your investment personality: {investment_personality}")

# # # # # #         # Summarize collected information
# # # # # #         summary = "\n".join([f"{q}: {a}" for q, a in user_responses.items()])
# # # # # #         summary = summary + "\n" + "Your investment personality:" + investment_personality
# # # # # #         await send_summary_chunks(chat_id, summary)
        
# # # # # #         # Send document for financial details
# # # # # #         await bot.send_message(chat_id, "Please fill in your details with correct information and then upload the document.")
# # # # # #         file = FSInputFile("data\Your Financial Profile.docx", filename="Your Financial Profile.docx")
# # # # # #         await bot.send_document(chat_id, document=file, caption="Fill in the Financial Details in this document and upload it here.")

# # # # # # async def send_summary_chunks(chat_id, summary):
# # # # # #     chunks = [summary[i:i+4000] for i in range(0, len(summary), 4000)]
# # # # # #     for chunk in chunks:
# # # # # #         await bot.send_message(chat_id, f"Assessment Summary:\n{chunk}")

# # # # # # async def determine_investment_personality(assessment_data):
# # # # # #     try:
# # # # # #         input_text = "User Profile:\n"
# # # # # #         for question, answer in assessment_data.items():
# # # # # #             input_text += f"{question}: {answer}\n"

# # # # # #         input_text += "\nClassify the user as:\n" \
# # # # # #                       "- Conservative Investor\n" \
# # # # # #                       "- Moderate Investor\n" \
# # # # # #                       "- Aggressive Investor"

# # # # # #         model = genai.GenerativeModel('gemini-pro')
# # # # # #         response = model.generate_content(input_text)
# # # # # #         response_text = response.text.lower()

# # # # # #         if "conservative" in response_text:
# # # # # #             personality = "Conservative Investor"
# # # # # #         elif "moderate" in response_text:
# # # # # #             personality = "Moderate Investor"
# # # # # #         elif "aggressive" in response_text:
# # # # # #             personality = "Aggressive Investor"
# # # # # #         else:
# # # # # #             personality = "Unknown"

# # # # # #         return personality

# # # # # #     except Exception as e:
# # # # # #         logging.error(f"Error generating response: {e}")
# # # # # #         return "Unknown"



# # # # # # questions = [
# # # # # #     {
# # # # # #         "question": "1. You and your friend are betting on a series of coin tosses.\n\nHe always bets ₹2,000 on Heads\nYou always bet ₹2,000 on Tails\n\nWinner of last 8 turns\nYou lost ₹8,000 in the last 4 turns!\n\nIf you were to bet one last time, what would you bet on?",
# # # # # #         "options": ["Heads", "Tails"]
# # # # # #     },
# # # # # #     {
# # # # # #         "question": "2. Imagine you are a contestant in a game show, and you are presented with the following choices. What would you prefer?",
# # # # # #         "options": ["50% chance of winning 15 gold coins", "100% chance of winning 8 gold coins"]
# # # # # #     },
# # # # # #     {
# # # # # #         "question": "3. In general, how would your best friend describe your risk-taking tendencies?",
# # # # # #         "options": ["A real gambler", "Willing to take risks after completing adequate research", "Cautious", "Avoids risk as much as possible"]
# # # # # #     },
# # # # # #     {
# # # # # #         "question": "4. Suppose you could replace your current investment portfolio with this new one:\n\n50% chance of gaining 35% or 50% chance of Loss\n\nIn order to have a 50% chance of gaining +35%, how much loss are you willing to take?",
# # # # # #         "options": ["-5% to -10%", "-10% to -15%", "-15% to -20%", "-20% to -25%", "-25% to -30%", "-30% to -35%"]
# # # # # #     },
# # # # # #     {
# # # # # #         "question": "5. Over any 1-year period, what would be the maximum drop in the value of your investment portfolio that you would be comfortable with?",
# # # # # #         "options": ["<5%", "5% - 10%", "10% - 15%", "15% - 20%", ">20%"]
# # # # # #     },
# # # # # #     {
# # # # # #         "question": "6. When investing, what do you consider the most?",
# # # # # #         "options": ["Risk", "Return"]
# # # # # #     },
# # # # # #     {
# # # # # #         "question": "7. What best describes your attitude?",
# # # # # #         "options": ["Prefer reasonable returns, can take reasonable risk", "Like higher returns, can take slightly higher risk", "Want to maximize returns, can take significant high risk"]
# # # # # #     },
# # # # # #     {
# # # # # #         "question": "8. How much monthly investment do you want to do?",
# # # # # #         "options": ["$1000 to $2000","$3000 to $4000","$4000 to $5000","more than $5000"]
# # # # # #     },
# # # # # #     {
# # # # # #         "question": "9. What is the time horizon for your investment?\nChoose any Range :",
# # # # # #         "options": ["0 to 1 year","1 to 2 years","3 to 5 years","more than 5 years" ]
# # # # # #     }
# # # # # # ]



# # # # # # # Poll handler to send a poll
# # # # # # async def send_poll(message: types.Message):
# # # # # #     poll = await bot.send_poll(
# # # # # #         chat_id=message.chat.id,
# # # # # #         question="How Can I Help you today ?",
# # # # # #         options=["Investment Personality", "Tax Details", "Savings and Wealth Management", "Debt Repayment Strategies"],
# # # # # #         type="regular",
# # # # # #         is_anonymous=False
# # # # # #     ) # question="What would you like to explore?",

# # # # # #     logging.info(f"Poll sent with ID {poll.poll.id}")
# # # # # #     print(f"Poll sent with ID {poll.poll.id}")
# # # # # #     # Store poll id to track user's choice later
# # # # # #     user_responses[message.from_user.id] = {'poll_id': poll.poll.id}

# # # # # # # Poll Answer handler
# # # # # # @router.poll_answer()
# # # # # # async def handle_poll_answer(poll_answer: PollAnswer):
# # # # # #     logging.info(f"Poll answer received from user {poll_answer.user.id}: selected option {poll_answer.option_ids[0]}")
# # # # # #     print(f"Poll answer received from user {poll_answer.user.id} : selected option {poll_answer.option_ids[0]}")
# # # # # #     chat_id = poll_answer.user.id
# # # # # #     selected_option_id = poll_answer.option_ids[0]
    
# # # # # #     logging.info(f"Poll answer received from user {chat_id}: selected option {selected_option_id}")
# # # # # #     print(f"Poll answer received from user {chat_id} : selected option {selected_option_id}")

# # # # # #     # Mapping of options to functions with chat_id passed
# # # # # #     options_map = {
# # # # # #         0: lambda: start_assessment(chat_id),
# # # # # #         1: lambda: start_assessment(chat_id),
# # # # # #         2: lambda: start_assessment(chat_id),
# # # # # #         3: lambda: start_assessment(chat_id),
# # # # # #     }

# # # # # #     if selected_option_id in options_map:
# # # # # #         # Call the corresponding function
# # # # # #         logging.info(f"Starting process for option {selected_option_id}")
# # # # # #         print(f"Starting process for option {selected_option_id}")
# # # # # #         await options_map[selected_option_id]()
# # # # # #     else:
# # # # # #         logging.warning(f"Selected option {selected_option_id} is not in the options_map")
# # # # # #         print(f"Selected option {selected_option_id} is not in the options_map")